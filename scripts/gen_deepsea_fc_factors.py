#!/usr/bin/env python3
"""Generate docx: 深海多堆燃料电池系统工程因素全览"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 样式基础 ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35

# 标题样式
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── 辅助函数 ──────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    """简易表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    # data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return table


def add_bullet(doc, text, level=0, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ── 封面 / 标题 ──────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
run = title.add_run('深海多堆燃料电池系统\n工程因素全览')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('多目标优化框架下的工程因素映射 · 2026-07')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第一章：四个优化目标的工程意义
# ═══════════════════════════════════════════════════════════
doc.add_heading('一、四个优化目标的工程意义映射', level=1)

doc.add_paragraph(
    '本研究的四个优化目标（等效氢耗、功率波动频率、电堆启停次数、四堆效率偏差）'
    '分别映射到 7 个顶层工程维度。下表是开题总结对照表：'
)

add_table(
    doc,
    ['优化目标', '工程维度', '直接工程指标', '物理单位', '对应因素类型', '参考权重\n(深海)'],
    [
        ['1. 等效氢耗最小', '经济性',  '续航里程',               'km',           '目标', '★★★★☆'],
        ['',                '经济性',  '燃料成本',               '¥/任务',       '目标', ''],
        ['',                '储运',    '储氢/储氧系统体积',       'm³ 或 kg',     '约束', ''],
        ['',                '热管理',  '散热系统功耗占比',        '% of BOP',     '约束', ''],
        ['',                '安全性',  '储氢量→泄漏风险池',       '无量纲风险分', '约束', ''],
        ['2. 功率波动最小', '耐久性',  '电堆机械/热寿命',         'h',            '目标', '★★★★★'],
        ['',                '耐久性',  'BOP 附件寿命',            'h',            '目标', ''],
        ['',                '声学',    '声辐射/振动特征 ⚓',       'dB re 1μPa',  '目标', ''],
        ['',                '电气',    'DC/DC 可靠性 + EMC',      'MTBF /  dBμV', '约束', ''],
        ['',                '动态',    '负载跟随响应时间',        's',            '约束', ''],
        ['3. 启停次数最少', '耐久性',  '碳腐蚀寿命（最大杀手）',   '次 → h',       '目标', '★★★★★'],
        ['',                '安全性',  '任务可靠性（启动失败率）',  '[0,1]',        '约束', ''],
        ['',                '能耗',    '冷启动+待机功耗',          '% SOC/次',     '约束', ''],
        ['',                '消耗品',  '吹扫 N₂ 消耗量',          'L/次',         '约束', ''],
        ['4. 效率偏差最小', '维护性',  '堆间均衡老化→维护周期',    'h / 更换',     '目标', '★★★★☆'],
        ['',                '安全性',  '故障预警灵敏度',           'h 提前量',     '目标', ''],
        ['',                '物流',    '备件库存策略复杂度',       '种类/年',      '约束', ''],
        ['',                '可扩展',  '控制策略能否推广到6/8堆',   '是/否',        '设计验证', ''],
    ],
    col_widths=[2.5, 1.8, 3.5, 2.2, 2, 1.8]
)

doc.add_paragraph()
doc.add_heading('1.1 等效氢耗最小 → 续航里程', level=2)
doc.add_paragraph(
    '等效氢耗将电池净电耗折算为等效氢气消耗，统一度量后最小化。'
    '这是最直观的经济性指标——同样储氢量下跑得更远。'
    '在深海场景中，补氢几乎不可能，续航即任务半径。'
    '然而，单纯追求低氢耗往往需要频繁调整功率适配负载，与目标 2 和 3 冲突。'
)

doc.add_heading('1.2 功率波动频率最小 → 电堆寿命 + 系统可靠性', level=2)
doc.add_paragraph('燃料电池"怕变不怕稳"。频繁的功率波动会带来连锁损伤链：')
add_bullet(doc, '机械应力：MEA 反复膨胀收缩 → 膜电极机械疲劳', bold_prefix='')
add_bullet(doc, '气体饥饿：电流突变时供气滞后 → 局部缺气（starvation）→ 反极腐蚀催化剂', bold_prefix='')
add_bullet(doc, '热循环：功率变化 → 产热变化 → 温度梯度冲击 → 双极板密封失效加速', bold_prefix='')
add_bullet(doc, '电力电子：前端 DC/DC 变换器 IGBT 开关损耗和热应力随波动频次上升', bold_prefix='')
doc.add_paragraph(
    '深海场景中热管理惯性更大（海水导热但调节慢），功率波动的危害被放大。'
    '典型工程约束：电堆功率变化率 ≤ 3~5 kW/s（地面车载为 5~10 kW/s）。'
)

doc.add_heading('1.3 电堆启停次数最少 → 最关键的寿命杠杆', level=2)
doc.add_paragraph(
    '启停是电堆衰减的最大来源，通常贡献总衰减的 50~60%。其损伤机制如下：'
)
add_bullet(doc, '启动阶段：低电压 + 高电位 → 碳载体腐蚀（Carbon Corrosion），Pt 催化剂脱落/团聚', bold_prefix='')
add_bullet(doc, '停机阶段：阳极残留 H₂ 与穿透的 O₂ 形成氢-空界面（H₂/air front）→ 阳极高电位腐蚀', bold_prefix='')
add_bullet(doc, '数据参考：一次启停 ≈ 连续运行数小时的衰减量', bold_prefix='')
doc.add_paragraph(
    '在深海场景中，启停的危害更大——深度放电冷启动需要额外 15~30 分钟电加热，'
    '升温过程中的热应力会叠加损伤。工程上采用"不关机保温 idling"策略尽可能避免启停。'
)

doc.add_heading('1.4 四堆效率偏差最小 → 均衡老化 + 维护经济性', level=2)
doc.add_paragraph(
    '多堆系统特有的目标。核心工程逻辑是"短板效应"：'
)
add_bullet(doc, '四个电堆中只要有一个效率显著下降（如冷却不均导致膜干枯），整个系统被它拖累', bold_prefix='')
add_bullet(doc, '均衡老化 ≈ 各堆同时到寿、同时更换，而不是轮流进厂——维护成本大幅降低', bold_prefix='')
add_bullet(doc, '偏差往往是故障先兆：偏差扩大可能预告某堆正在静默失效（密封渗漏/催化剂中毒）', bold_prefix='')
doc.add_paragraph(
    '控制手段：通过功率分配策略（power-split）让各堆运行在相近的效率点，'
    '同时利用偏差监控作为故障早期预警的输入。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第二章：深海环境的特殊性
# ═══════════════════════════════════════════════════════════
doc.add_heading('二、深海环境的特殊性——附加工程考量', level=1)

doc.add_paragraph(
    '地面车载燃料电池的工程因素框架只是起点。深海环境引入了车载场景不存在的'
    '一系列硬约束，需要在优化模型中显式建模。'
)

doc.add_heading('2.1 高压环境（2000m+ → 20MPa+）', level=2)

add_table(
    doc,
    ['工程问题', '物理机制', '后果', '应对措施'],
    [
        ['气体压缩性非线性', '高压下 H₂/O₂ 密度剧变，质量流量计误差放大', '供气控制精度下降', '温度/压力补偿流量模型'],
        ['MEA 压差失衡', '外静水压 vs 内气体压差超标', '膜破裂或密封失效', '压力补偿系统（响应 ≤ 1s）'],
        ['高压气液分离', '高压下气泡更小更难分离', '水淹 → 局部缺气 → 反极', '高压气液分离器 / 氢气循环泵'],
        ['密封蠕变加速', '持续高压下 O-ring 蠕变', '海水渗入 → 短路烧毁', 'FFKM / 双道密封 + 泄漏监控'],
    ],
    col_widths=[3.5, 4, 3.5, 4]
)

doc.add_heading('2.2 氧气供应策略', level=2)
doc.add_paragraph('深海没有空气（21% O₂），氧气来源是两个截然不同的思路：')

add_table(
    doc,
    ['维度', '方案 A：携带纯氧', '方案 B：海水溶解氧提取'],
    [
        ['实现方式', '高压氧罐（液态/气态）', '膜接触器从海水萃取溶解 O₂'],
        ['能量密度', '高（体积占比 15~20%）', '低（萃取能耗高）'],
        ['续航限制', '储氧量封顶', '理论上无限（有海就有氧）'],
        ['安全性', '纯氧泄漏 → 助燃风险', '常压氧 → 安全'],
        ['技术成熟度', '成熟（已工程化）', '前沿（实验室阶段）'],
        ['适用场景', '浅海 < 500m / 短期任务', '深海长航时 AUV（远期）'],
    ],
    col_widths=[2.5, 5.5, 6]
)

doc.add_heading('2.3 低温海水热管理', level=2)
doc.add_paragraph(
    '深海海水 2~4°C vs 电堆工作温度 60~80°C，大温差既是优势也是陷阱：'
)
add_bullet(doc, '散热效率极高 → 散热器体积可显著缩小', bold_prefix='✅ 优势：')
add_bullet(doc, '冷启动困难 → 需要电加热或氢气催化燃烧预热（15~30 min）', bold_prefix='❌ 陷阱：')
add_bullet(doc, '局部过冷 → 冷凝水堆积 → 流道水淹（flooding）', bold_prefix='❌ 陷阱：')
add_bullet(doc, '控制措施：冷却回路三通比例阀，根据负载自动调节冷却水流量', bold_prefix='')

doc.add_heading('2.4 压力瞬变——浮潜过程', level=2)
doc.add_paragraph('上浮/下潜时环境压力快速变化，对电堆密封和压力补偿系统构成考验：')

add_table(
    doc,
    ['潜航速度 (m/s)', '压力变化率 (kPa/s)', '电堆影响评估', '工程要求'],
    [
        ['< 0.5', '< 5', '安全区', '常规设计即可'],
        ['0.5 ~ 1.0', '5 ~ 10', '密封可承受区', '双道密封'],
        ['1.0 ~ 2.0', '10 ~ 20', '需设计验证', '压力补偿器响应 ≤ 2s'],
        ['> 2.0', '> 20', '⚠ 危险区', '需要特殊快速补偿设计'],
    ],
    col_widths=[2.5, 3, 3.5, 5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第三章：多堆架构的工程深化
# ═══════════════════════════════════════════════════════════
doc.add_heading('三、多堆架构——深海场景的工程深化', level=1)

doc.add_heading('3.1 多堆的必然性', level=2)
doc.add_paragraph(
    '深海大功率场景（载人深潜器 / 深海工作站）单堆功率受限（当前单堆最大 150~200 kW），'
    '必须采用多堆串联/并联架构。多堆带来了额外的工程要求：'
)
add_bullet(doc, '各堆独立压力补偿，防止一个堆泄漏带崩整个系统')
add_bullet(doc, '故障隔离与热插拔：任一电堆故障自动切出，其余堆提升功率确保安全返航')
add_bullet(doc, '冗余管理：N+M 冗余设计（如 4+1），确保单点故障不影响任务')

doc.add_heading('3.2 密封系统——深海第一故障模式', level=2)
doc.add_paragraph('深海燃料电池的第一大故障不是催化剂衰减，而是密封失效。')

add_table(
    doc,
    ['密封位置', '失效后果', '工程方案'],
    [
        ['电堆端板密封', '海水渗入 → 短路烧毁', '双道 O-ring + 泄漏检测腔（中间真空监控）'],
        ['气体管路接头', 'H₂/O₂ 泄漏到舱内 → 爆炸风险', '全焊接管路 + VCR 金属密封接头'],
        ['压力补偿器膜片', '海水进入补偿油 → 污染电堆', '多层复合膜 + 位移传感器监控膜片位置'],
        ['水道密封', '冷却液泄漏到堆内 → 绝缘下降', '绝缘电阻在线监测 + 接地故障保护'],
    ],
    col_widths=[3.5, 4.5, 6.5]
)

doc.add_heading('3.3 深海排水与水管理', level=2)
doc.add_paragraph(
    '深海高压下的水管理比常压复杂得多。生成水在高压下溶解度、表面张力均变化，'
    'Nafion 膜的水含量-电导率关系需要重新标定。排水不能直接排海（压差过大），'
    '需要专门的排水舱逐级减压：'
)
add_bullet(doc, '收集到废水舱 → 浮到浅海再排出')
add_bullet(doc, '或利用海水负压喷射排水（需额外能耗）')
add_bullet(doc, '阳极侧氢气循环泵可帮助带出液态水，但循环泵本身在高压下的可靠性需验证')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第四章：系统集成工程约束
# ═══════════════════════════════════════════════════════════
doc.add_heading('四、系统集成工程约束清单', level=1)

doc.add_paragraph(
    '以下为深海多堆燃料电池系统的硬性约束，在优化模型中应作为约束条件而非优化目标处理：'
)

doc.add_heading('4.1 性能约束', level=2)
add_table(
    doc,
    ['约束项', '典型值', '超出后果', '与优化目标的关系'],
    [
        ['电堆输出功率范围', '额定功率 20%~100%', '<20% 水淹/效率崩；>100% 热失控', '约束 1（氢耗优化边界）'],
        ['功率变化率', '≤ 3~5 kW/s', '供气滞后 → 反极', '约束 2（波动优化边界）'],
        ['电堆温度范围', '60~80°C', '<60°C 效率低；>80°C 膜脱水', '约束 1 & 2（间接）'],
        ['堆间温度偏差', '≤ 3°C', '偏差扩大 → 单堆提前老化', '约束 4（偏差优化边界）'],
        ['阳极-阴极压差', '≤ 50 kPa', 'MEA 机械破裂', '安全硬约束'],
        ['单电池最低电压', '≥ 0.55 V', '反极 → 不可逆损伤', '安全硬约束'],
        ['系统总储氢量', '任务前确定', '不可逾越的物理上限', '约束 1'],
        ['电堆总重量/体积', '由舱体容量决定', '装不进去', '约束 4（影响堆数）'],
    ],
    col_widths=[3.5, 3, 4, 4.5]
)

doc.add_heading('4.2 安全性约束（深海无逃生通道）', level=2)

add_table(
    doc,
    ['风险源', '后果', '防护措施'],
    [
        ['氢气泄漏到舱内', '爆炸（空气中 H₂ 4~75%）', '全舱 H₂ 传感器 + 强制通风 + 点火源隔离'],
        ['氧气泄漏', '富氧环境 → 材料易燃', 'O₂ 浓度监测 + 材料 UL94 V-0 阻燃等级'],
        ['高压气体管路破裂', '瞬间减压冲击波', '爆破片 + 泄压管路引到舱外'],
        ['电堆热失控', '膜烧穿 → H₂/O₂ 混合 → 局部燃烧', '每片单电池电压监控 + 快速切断阀'],
        ['冷却液泄漏', '绝缘下降 → 短路 → 停机', '绝缘电阻在线监测'],
    ],
    col_widths=[3.5, 4, 7]
)

doc.add_heading('4.3 体积/重量约束（潜航器舱体寸土寸金）', level=2)
add_table(
    doc,
    ['子系统', '体积占比（典型）', '减重/紧凑化方向'],
    [
        ['氢气储罐', '35~45%', '金属氢化物储氢（体积密度↑）'],
        ['氧气储罐/供应系统', '15~20%', '液态氧 LOX 或海水提氧'],
        ['电堆本体', '12~18%', '超薄金属双极板'],
        ['热管理系统', '10~15%', '深海海水直冷（省略散热器）'],
        ['BOP（泵/阀/传感器/DC/DC）', '8~12%', '集成化阀组'],
        ['压力补偿系统', '5~8%', '共用补偿器'],
    ],
    col_widths=[3.5, 2.5, 8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第五章：优化目标权重在深海 vs 地面
# ═══════════════════════════════════════════════════════════
doc.add_heading('五、优化目标权重的场景迁移', level=1)

doc.add_paragraph(
    '地面车载与深海潜航器对四个优化目标的优先级显著不同。核心判断标准：'
    '深海场景中安全 > 寿命 > 效率。'
)

add_table(
    doc,
    ['目标', '地面车辆', '深海潜航器', '权重漂移原因'],
    [
        ['等效氢耗最小', '★★★★★', '★★★★☆', '深海补氢几乎不可能，但安全约束更硬'],
        ['功率波动最小', '★★★☆☆', '★★★★★', '深海热管理惯性大，波动易引发热失控;\n声学隐身要求放大波动权重'],
        ['启停最少', '★★★★☆', '★★★★★', '深海冷启动能耗极高 + 启动失败→不可返回'],
        ['偏差最小', '★★☆☆☆', '★★★★☆', '偏差是静默失效的先行指标→安全预警'],
    ],
    col_widths=[2.5, 2.5, 2.5, 6.5]
)

doc.add_paragraph()
doc.add_heading('5.1 额外工程维度的加权考虑', level=2)
doc.add_paragraph('除上述四个目标外，以下工程因素在深海场景中的重要性应被纳入优化决策：')

add_table(
    doc,
    ['工程维度', '地面车辆', '深海潜航器', '在优化框架中的位置'],
    [
        ['安全性（泄漏/绝缘/压差）',    '★★★☆☆', '★★★★★', '硬约束，不参与 Pareto 优化'],
        ['声学特征/振动噪声',          '★☆☆☆☆', '★★★★★ (军用)',   '军用：第 6 优化目标；民用：可忽略'],
        ['堆内温度均匀性',            '★★★☆☆', '★★★★☆', '建议作为第 5 优化目标'],
        ['负载跟随响应速度',           '★★★★☆', '★★★★☆', '硬约束，防止优化出"慢响应"解'],
        ['系统复杂度（传感器/执行器数）', '★★★☆☆', '★★★☆☆', '设计约束：可靠性 vs 性能权衡'],
        ['可维护性（模块化更换）',      '★★☆☆☆', '★★★★☆', '设计约束，影响全生命周期成本'],
    ],
    col_widths=[4, 2.5, 2.5, 5.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('核心结论：')
run.bold = True
run.font.size = Pt(12)
p.add_run(
    '地面场景优化排序为 氢耗 > 启停 > 波动 > 偏差；'
    '深海场景优化排序为 启停 > 波动 > 偏差 > 氢耗。'
    '安全性和任务可靠性压倒纯经济性。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第六章：全生命周期成本分解
# ═══════════════════════════════════════════════════════════
doc.add_heading('六、四个优化目标 → 全生命周期成本', level=1)

doc.add_paragraph(
    '四个优化目标最终服务的是同一个宏观经济指标——全生命周期成本。'
    '优化框架回答的是"一次能跑多远 + 整个寿命花多少钱"两个问题。'
)

doc.add_heading('6.1 成本分解模型', level=2)

p = doc.add_paragraph()
run = p.add_run(
    '总成本 = 燃料成本 + 电堆折旧成本 + 维护成本 + 任务失败风险成本'
)
run.bold = True

doc.add_paragraph()

add_table(
    doc,
    ['成本分量', '占比（经验估计）', '由哪个目标直接决定', '间接影响'],
    [
        ['燃料成本', '20~30%', '目标 1（氢耗最小）', '无'],
        ['电堆折旧成本', '30~45%', '目标 3（启停最少）贡献 ~60%', '目标 2（波动）贡献 ~20%'],
        ['维护成本', '15~25%', '目标 4（偏差最小）', '目标 3（启停）影响维护间隔'],
        ['任务失败风险', '5~15%', '目标 2 + 目标 4', '可靠性建模'],
    ],
    col_widths=[3, 2.5, 4, 4]
)

doc.add_paragraph()
doc.add_heading('6.2 因素关系总图', level=2)

doc.add_paragraph(
    '以下用文字描述多目标优化背后的因果关系链，供建立优化模型时参考：'
)

lines = [
    '目标 1（氢耗最小）    → 燃料费 ↓  ↛ 代价：波动↑ 启停↑',
    '目标 2（波动最小）    → 机械/热寿命↑ → 折旧费 ↓  +  可靠性 ↑',
    '目标 3（启停最少）    → 碳腐蚀寿命↑ → 折旧费 ↓↓↓（最显著）',
    '目标 4（偏差最小）    → 维护周期统一 → 维护费 ↓↓  +  故障预警 ↑',
]
for line in lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'  {line}')
    run.font.size = Pt(10.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════════════════════
doc.add_heading('附录：工程因素 → 优化变量映射', level=1)

doc.add_paragraph(
    '将工程因素转化为数学优化模型时，建议的映射关系：'
)

add_table(
    doc,
    ['工程因素', '优化变量类型', '数学模型表示', '备注'],
    [
        ['续航里程',              '目标函数', 'J₁ = ∫ Ḣ₂_eq dt', '等效氢耗累计'],
        ['电堆寿命（机械+碳腐蚀）', '目标函数', 'J₂ = 1/(αN_start + β∫|dP/dt|dt)', '启停 + 波动联合建模'],
        ['系统可靠性',             '目标函数', 'J₃ = σ_eff⁻¹ · (1 - P_fail)', '效率偏差→故障预警'],
        ['堆内温度均匀性',         '目标函数\n(建议新增)', 'J₄ = max(T_ij) - min(T_ij)', '第 5 优化目标'],
        ['声学特征（军用水下）',    '目标函数\n(军用场景)', 'J₅ = ∫ S(f) · W(f) df', '声压谱密度 × 听觉加权'],
        ['功率范围',               '约束',  'P_min ≤ P ≤ P_max', '硬件限制'],
        ['功率变化率',             '约束',  '|dP/dt| ≤ ΔP_max', '硬件限制'],
        ['负载跟随时间',           '约束',  't_response ≤ 2 s', '动态响应硬约束'],
        ['电堆温度窗口',           '约束',  'T_min ≤ T ≤ T_max', '硬件限制'],
        ['阳极-阴极压差',           '约束',  '|P_an - P_ca| ≤ 50 kPa', '安全红线'],
        ['单电池最低电压',          '约束',  'V_cell ≥ 0.55 V', '安全红线'],
        ['舱内 H₂ 浓度',           '约束',  'C_H2 ≤ 0.4% LFL', '安全红线'],
        ['绝缘电阻',               '约束',  'R_ins ≥ 1 MΩ', '安全红线'],
        ['储氢量上限',             '约束',  '∫ Ḣ_total dt ≤ V_H₂', '物理上限'],
        ['系统总重/体积',           '约束',  'W ≤ W_max, V ≤ V_max', '舱体限制'],
        ['过载能力',               '约束',  'P_max ≥ 120%·P_rated (30s)', '紧急避障需求'],
        ['母线电压波动',            '约束',  'ΔV_bus ≤ ±5%', '推进电机控制品质'],
        ['冷启动能耗',             '约束',  'E_start ≤ 5% SOC', '电池容量限制'],
    ],
    col_widths=[3, 2.5, 5, 3]
)

doc.add_paragraph()
doc.add_paragraph(
    '— 文档结束 —'
).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 保存 ──────────────────────────────────────────────────
out_path = r'F:\CLAUDE\research\ems-platform\docs\深海多堆燃料电池系统工程因素全览.docx'
doc.save(out_path)
print(f'[OK] 文档已生成：{out_path}')
print(f'    文件大小：{os.path.getsize(out_path) / 1024:.1f} KB')
