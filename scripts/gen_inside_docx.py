#!/usr/bin/env python3
"""生成体制内投递清单 DOCX"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# 封面
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EMS/BMS 算法方向 · 体制内投递清单')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026-06-18 | 211本+双非硕 | 控制算法方向')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

def add_table(doc, data):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    return table

# 一、央企
doc.add_heading('一、央企·总部/子公司（14家）', level=2)
add_table(doc, [
    ['#', '单位', '城市', '对口岗位', '年薪参考', '投递渠道'],
    ['1', '中电科蓝天/18所', '天津/成都', '电源系统研发(BMS/嵌入式)', '硕士优厚+六险二金', 'dklt.zhiye.com'],
    ['2', '中广核集团', '深圳/广东', '新能源/储能/控制', '硕士15-25万', 'cgnpc.iguopin.com'],
    ['3', '南方电网储能', '广州/深圳', '储能科研/自动化控制', '硕士18-28万', 'zhaopin.csg.cn'],
    ['4', '国家能源集团', '北京/全国', '氢能/储能研究员', '硕士15-25万', '国家能源集团招聘平台'],
    ['5', '国家管网储能', '上海/常州', '电气自控/新能源', '硕士14-22万', '国家管网招聘平台'],
    ['6', '东风汽车集团', '武汉', 'BMS算法/能量管理', '硕士20-35万', 'dfmc.com.cn'],
    ['7', '中国一汽', '长春/南京', '新能源研发/算法', '硕士15-25万', 'zhaopin.faw.com.cn'],
    ['8', '中石化石油工程设计', '武汉/成都/西安', '智能低碳研究岗', '硕士15-22万', '中石化人才招聘网'],
    ['9', '中石化江汉设计', '武汉/重庆', '智能低碳/智慧能源', '硕士15-22万', '中石化人才招聘网'],
    ['10', '中船集团', '上海/广州', '新能源/储能/控制', '硕士16-26万', '中船招聘平台'],
    ['11', '中核集团', '北京/全国', '储能/新能源控制', '硕士15-25万', 'cnnc.chinahr.com'],
    ['12', '中国化学工程', '北京/武汉', '新能源/储能技术', '硕士14-22万', 'cncec.chinahr.com'],
    ['13', '中国三峡集团', '武汉/北京/成都', '新能源/储能', '硕士16-26万', 'zhaopin.ctg.com.cn'],
    ['14', '中国华能集团', '北京/全国', '新能源/储能研发', '硕士15-25万', '华能校招系统'],
])
doc.add_paragraph()

# 二、央企研究院
doc.add_heading('二、央企研究院（7家）', level=2)
doc.add_paragraph('福利好、户口有优势、技术方向对口')
add_table(doc, [
    ['#', '单位', '城市', '对口岗位', '薪资福利', '投递渠道'],
    ['15', '华能清洁能源研究院', '北京', '算法工程师/自动化', '七险二金+北京落户', '华能校招系统'],
    ['16', '大唐科研总院', '北京/合肥/成都等', '控制工程/新能源', '六险二金+安家费', '大唐人才招聘'],
    ['17', '中石化工程设计研究院', '天津/成都/西安', '智能低碳研究', '央企编制', '中石化人才招聘网'],
    ['18', '中电科蓝天研究院(18所)', '天津', 'BMS/电源系统/储能', '博士40-100万', 'dklt.zhiye.com'],
    ['19', '南方电网储能科研院', '广州', '储能技术/AI', '高层次人才补贴', 'zhaopin.csg.cn'],
    ['20', '天津电气科学研究院', '天津', '电气控制/储能', '硕士15-22万', '官网招聘'],
    ['21', '湖州工控技术研究院', '浙江湖州', '算法开发(风光储氢)', '事业编制', '浙大就业网'],
])
doc.add_paragraph()

# 三、国企
doc.add_heading('三、国企·地方能源集团（6家）', level=2)
add_table(doc, [
    ['#', '单位', '城市', '对口岗位', '年薪参考', '投递渠道'],
    ['22', '潍柴动力', '潍坊/苏州', 'BMS/SOC/SOH算法', '硕士15-22万', 'xyzp@weichai.com'],
    ['23', '中创新航(CALB)', '常州', 'BMS算法工程师', '13-18万(13薪)', 'calb.com.cn'],
    ['24', '国轩高科', '合肥/上海/南京', 'BMS软硬件/算法', '18-26万', 'gotion.zhiye.com'],
    ['25', '中国一汽研发总院', '长春/南京', '新能源/智能控制', '15-25万', 'zhaopin.faw.com.cn'],
    ['26', '东方电气', '成都/广州', '氢能/储能控制', '15-24万', '东方电气招聘平台'],
    ['27', '上海电气', '上海', '储能/新能源控制', '16-26万', '上海电气校招'],
])
doc.add_paragraph()

# 四、事业单位
doc.add_heading('四、事业单位·高校/研究所（3家）', level=2)
add_table(doc, [
    ['#', '单位', '城市', '对口岗位', '薪资福利', '投递渠道'],
    ['28', '湖州工控技术研究院', '浙江湖州', '算法开发工程师', '事业编制', '浙大就业网'],
    ['29', '天津电气科学研究院', '天津', '电气/控制研发', '事业编制', '官网招聘'],
    ['30', '中国汽车技术研究中心', '天津/武汉', '新能源/算法', '硕士15-22万', '中汽中心招聘'],
])
doc.add_paragraph()

# 五、军工
doc.add_heading('五、特殊赛道·军工/航天（3家）', level=2)
add_table(doc, [
    ['#', '单位', '城市', '对口岗位', '说明'],
    ['31', '航天科工集团', '北京/武汉/南京', '电源/控制/算法', '关注各校就业网'],
    ['32', '中国电科集团(其他所)', '多地', '电源/BMS/控制', '各所独立招聘'],
    ['33', '中船重工712所', '武汉', '电池/储能/控制', '船用电池方向'],
])
doc.add_paragraph()

# 城市分布
doc.add_heading('城市分布（体制内）', level=2)
add_table(doc, [
    ['城市', '数量', '主要单位'],
    ['天津/北京', '约10家', '央企总部集中地'],
    ['武汉', '约6家', '东风+中石化+中船712'],
    ['成都/西安', '约4家', '中电科+中石化布局'],
    ['广州/深圳', '约4家', '南方电网+中广核'],
    ['合肥/南京', '约3家', ''],
    ['上海', '约3家', ''],
    ['常州/湖州', '约2家', ''],
])
doc.add_paragraph()

# 对比
doc.add_heading('体制内 vs 市场化对比', level=2)
add_table(doc, [
    ['对比项', '体制内（央企/研究所）', '市场化（比亚迪/汇川等）'],
    ['硕士薪资', '15-25万', '20-37万'],
    ['涨幅', '慢，每年5-10%', '快，跳槽可翻倍'],
    ['稳定性', '极高（5星）', '一般（3星）'],
    ['户口', '大概率解决', '看公司实力'],
    ['加班', '相对少', '普遍较多'],
    ['技术成长', '慢，流程多', '快，项目多'],
    ['福利', '六险二金+宿舍+食堂', '高薪但无额外福利'],
])
doc.add_paragraph()

# 策略建议
doc.add_heading('投递策略建议', level=2)
doc.add_paragraph('适合去体制内的情况：')
for item in ['求稳 > 求财', '想要户口/编制/福利房', '能接受前期薪资低一点、涨得慢一点', '想长期在一个城市定居']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('推荐折中策略：')
p = doc.add_paragraph()
run = p.add_run('主投市场化（比亚迪/汇川/豪鹏等20-30万档）+ 同时申2-3家央企研究院/国企保底（中电科/南方电网/东风）')
run.bold = True

# 平台
doc.add_paragraph()
doc.add_heading('体制内招聘关键平台', level=2)
add_table(doc, [
    ['平台', '网址', '说明'],
    ['国聘网', 'iguopin.com', '央企招聘官方平台'],
    ['国资小新', '公众号', '国资委官方发布'],
    ['央企校招专栏', '各央企官网', '最权威'],
    ['高校人才网', 'gaoxiaojob.com', '事业单位/高校'],
    ['北极星招聘', 'hr.bjx.com.cn', '能源电力行业垂直'],
    ['应届生求职网', 'yingjiesheng.com', '校招信息汇总'],
])

output_path = os.path.join('F:\\CLAUDE\\research\\ems-platform\\docs', 'EMS_BMS_体制内投递清单.docx')
doc.save(output_path)
print(f'[OK] 已生成: {output_path}')
