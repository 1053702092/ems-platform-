# -*- coding: utf-8 -*-
"""Generate DP vs ECMS formula comparison docx (readable version)
Uses content derived from formula-derivation skill, written in plain language."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DOCS_DIR = r'F:\CLAUDE\research\ems-platform\docs'

def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def p(doc, text, bold=False, size=10.5):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def tbl(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table

def formula(doc, text, size=10):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(size)
    run.font.italic = True
    return para


def main():
    os.makedirs(DOCS_DIR, exist_ok=True)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    title = doc.add_heading('DP 与 ECMS/Hamiltonian 公式关系详解', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('第5周补充材料 | 2026-06-12\n'
                       '答疑：DP 代码公式 vs ECMS 原理文档中的 Hamiltonian 公式为什么不一样？')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()

    # ================================================================
    # 1. 先摆出两个公式
    # ================================================================
    h(doc, '1. 先看两个公式', level=1)

    p(doc, '公式 A — 代码中的 DP Bellman 方程（day8_dp_ems.py）：', bold=True)
    formula(doc,
        'J[k][i] = min [ g(p_fc) + alpha * (SOC_next - SOC_ref)^2 + J[k+1][lookup(SOC_next)] ]')

    p(doc, '公式 B — ECMS 原理文档中的 Hamiltonian（连续最优控制）：', bold=True)
    formula(doc,
        'H(t) = H_fc(P_fc) + s(t) * H_bat(P_bat) + lambda(t) * d(SOC)/dt')

    p(doc,
        '一个看起来是递推公式，一个看起来是瞬时函数。结构完全不同。'
        '但它们优化的目标是同一个：最小化总氢耗。')

    # ================================================================
    # 2. 为什么 DP 没有 s * H_bat 项？
    # ================================================================
    h(doc, '2. 为什么 DP 里没有 s * H_bat？', level=1)

    p(doc,
        '简单说：DP 不需要它。\n\n'
        'DP 的思路是"把所有可能的路径都算一遍，然后选最好的"。'
        '它评估一个候选控制 p_fc 的好坏，看三项：')

    tbl(doc,
        ['DP 代价函数的三项', '含义', '大白话'],
        [
            ['g(p_fc)', 'FC 实际氢耗', '这一步烧了多少氢'],
            ['alpha * (SOC_next - SOC_ref)^2', 'SOC 偏离惩罚', '这一步跑完后 SOC 离目标有多远'],
            ['J[k+1][lookup(SOC_next)]', '未来最优代价', '从下一步到终点，最少还要烧多少氢'],
        ])

    p(doc,
        '注意第三项 J[k+1]——它已经包含了"电池现在用多了，后面要还"的全部信息。'
        'DP 是通过全局搜索知道这一点的，所以它不需要像 ECMS 那样，'
        '把电池的电能"换算"成等效氢耗。\n\n'
        'ECMS 之所以需要 s * H_bat，恰恰是因为它**每一步只能看当前**，看不到未来，'
        '所以它需要一个"价格标签"来预估"现在用电要付出多少代价"。')

    # ================================================================
    # 3. 两个公式的逐项对照
    # ================================================================
    h(doc, '3. 两个公式的逐项对照', level=1)

    p(doc,
        '把两个公式拆开来看，它们之间有清晰的对应关系：')
    doc.add_paragraph()

    tbl(doc,
        ['DP 里的项', 'Hamiltonian 里的项', '说明'],
        [
            ['g(p_fc)', 'H_fc(P_fc)',
             '同一回事，都是 FC 的实际氢耗。不同框架叫法不同而已。'],
            ['J[k+1]', 'lambda(t) * d(SOC)/dt',
             '最关键的一行！两者都编码了"未来信息"，但表达方式完全不同。见下文详解。'],
            ['alpha * (SOC-SOC_ref)^2', '(隐含在 lambda 里)',
             'alpha 是 DP 对 SOC 偏离的"罚金"。lambda 是 Hamiltonian 中 SOC 的"边际价格"。'
             '两者的作用类似，但数值上不能直接换算。'],
            ['(没有 s*H_bat)', 's(t) * H_bat(P_bat)',
             'DP 不需要这个项，因为 J[k+1] 已经管了电池的未来了。'],
            ['min_{p_fc}', 'min_{P_fc}',
             '都是"选最优控制"，一个离散一个连续。'],
        ])

    p(doc,
        '重点解释 J[k+1] 和 lambda(t) 的关系：', bold=True)

    p(doc,
        'J[k+1] 是 DP 的"未来记忆"——它存着"从下一步到终点最少要烧多少氢"这个信息。\n'
        'lambda(t) 是 Hamiltonian 的"未来记忆"——它表示"SOC 每变一点，对终点总代价的影响有多大"。\n\n'
        '物理意义完全相同：都是未来状态的**影子价格**。\n'
        '• DP 用查表 J[k+1] 直接读取\n'
        '• Hamiltonian 用协态变量 lambda 的微分方程来"传播"这个信息\n'
        '• 在连续极限下（时间步长趋近于 0），两者收敛到同一结果')

    # ================================================================
    # 4. 从 Hamiltonian 推到 ECMS
    # ================================================================
    h(doc, '4. 从 Hamiltonian 推到 ECMS', level=1)

    p(doc,
        'ECMS 不是凭空冒出来的，它是 Hamiltonian 在特殊情况下的简化。')

    h(doc, '4.1 一般情况：协态变量 lambda(t) 是变化的', level=2)
    p(doc,
        'Hamiltonian 的完整形式：')
    formula(doc, 'H(t) = H_fc(P_fc(t)) + lambda(t) * d(SOC)/dt')
    p(doc,
        '其中 lambda(t) 不是常数，它随时间变化。变化规律由协态方程给出：')
    formula(doc, 'd(lambda)/dt = -dH/d(SOC)')
    p(doc,
        '这个方程的意思是：lambda 的变化率等于 Hamiltonian 对 SOC 的偏导数的负值。'
        '简单理解：SOC 的"价格"会随着 SOC 本身的变化而调整。')

    h(doc, '4.2 特殊情况：lambda 是常数 -> ECMS', level=2)
    p(doc,
        '如果 lambda(t) 不变化（即 d(lambda)/dt = 0），协态方程就简化了。'
        '这要求 FC 效率曲线在工作点附近近似平坦——也就是说，FC 效率不随功率剧烈变化。\n\n'
        '在这个近似下，Hamiltonian 变成：')
    formula(doc, 'H = H_fc(P_fc) + lambda * (-I / (Q_bat * 3600))')
    p(doc,
        '再把电池电流 I 用功率 P_bat 表示（当内阻 R_int 远小于开路电压 V_oc 时，I ≈ P_bat / V_oc）：')
    formula(doc, 'H ≈ H_fc(P_fc) + lambda * (-P_bat / (Q_bat * 3600 * V_oc))')
    p(doc,
        '整理一下，把 lambda / (Q_bat * V_oc) 打包成一个新参数 s：')
    formula(doc, 's = lambda / (Q_bat * V_oc)')
    p(doc, '代入后得到：')
    formula(doc, 'H_eq = H_fc(P_fc) + s * P_bat / 3600')
    p(doc,
        '这就是 ECMS 的标准形式！s 就是等效因子，单位 g/kWh。'
        '它表示"电池每用 1kWh 电，等效于烧 s 克氢"。\n\n'
        '代入典型值：lambda ≈ 180 * 50 * 350 ≈ 3,150,000 g')

    h(doc, '4.3 更一般的情况：lambda 随 SOC 变化 -> A-ECMS', level=2)
    p(doc,
        '现实中 lambda 不是常数。当 SOC 偏离目标值时，横截条件会驱动 lambda 变化：')
    formula(doc, 'lambda(tf) = beta * (SOC(tf) - SOC_ref)')
    p(doc,
        '如果终点 SOC 不够（SOC(tf) < SOC_ref），则 lambda(tf) > 0\n'
        '-> lambda 在整个路径上为正 -> 电池放电"变贵"\n'
        '-> ECMS 倾向于多用 FC -> SOC 回升\n\n'
        'A-ECMS 的 SOC 反馈律就是对这个现象的简化表达：')
    formula(doc, 'lambda(t) ≈ lambda_0 * (1 + Kp * (SOC_ref - SOC(t)))')
    p(doc,
        '这个公式的意思是：SOC 比目标低时，lambda 增大（电池变贵，保护电池）；'
        'SOC 比目标高时，lambda 减小（电池变便宜，多用电池）。')

    # ================================================================
    # 5. 两个公式为什么看起来不一样？
    # ================================================================
    h(doc, '5. 为什么两个公式看起来完全不一样？', level=1)

    p(doc,
        '因为它们来自不同的数学"门派"，解题方向相反：')

    tbl(doc,
        ['', 'DP（Bellman 动态规划）', 'Hamiltonian（Pontryagin 最小原理）'],
        [
            ['解题方向', '从终点往回推（后向递推）', '从起点往前进（前向积分）'],
            ['怎么看到未来', 'J[k+1] 表——把每个状态的未来最优代价都存下来', 'lambda(t) 微分方程——通过方程传播未来信息'],
            ['SOC 怎么管', 'alpha 惩罚 + beta 终端惩罚', 'lambda 的边际价格 + 横截条件'],
            ['电池成本', '不需要单独算（J[k+1] 已包含）', '需要 s * H_bat 来预估'],
            ['能不能在线', '不能（需要已知全程）', '能（lambda 已知就能实时算）'],
        ])

    p(doc,
        '打个比方：\n\n'
        'DP 像一个有上帝视角的规划师——出发前就把整条路的所有加油站、油价、'
        '每个路口怎么走都算好了。\n\n'
        'Hamiltonian 像一个精明的商人——他不知道前面的路况，'
        '但他给 SOC 标了一个"价格" lambda。SOC 越少，价格越高，'
        '他就会越舍不得用电池。这个价格 lambda 会随着他的行程动态调整。\n\n'
        'ECMS 就是那个商人用的"快速算法"——假设 lambda 不变（或按简单规则调整），'
        '每一步只看当前的等效成本。')

    # ================================================================
    # 6. 参数之间的"换算"关系
    # ================================================================
    h(doc, '6. 参数之间的"换算"关系', level=1)

    p(doc,
        '很多人想知道：alpha 和 s 能不能直接换算？\n'
        '答案是：**不能精确换算，但有近似关系**。')

    p(doc, '通过协态变量 lambda 作为桥梁：', bold=True)
    formula(doc, 's = lambda / (Q_bat * V_oc)')
    formula(doc, 'alpha ≈ lambda * dt / (Q_bat * 3600)')

    p(doc,
        '代入典型值（Q_bat=50Ah, V_oc=350V, dt=1s, lambda 对应 s=180）：')
    formula(doc, 'lambda = s * Q_bat * V_oc = 180 * 50 * 350 = 3,150,000 g')
    formula(doc, 'alpha ≈ 3,150,000 * 1 / (50 * 3600) ≈ 17.5 g/SOC^2')

    p(doc,
        '但这个 alpha=17.5 只是理论估计。实际中 alpha 的最优值（100）和这个估算差了很多，'
        '原因：\n'
        '1. 上面的推导用了 I ≈ P_bat/V_oc 近似（忽略了内阻损耗）\n'
        '2. alpha 的实际最优值受离散化方式和工况影响\n'
        '3. DP 的 alpha 和 ECMS 的 s 本来就属于不同的优化框架\n\n'
        '所以 alpha 和 s 的"对应关系"更多是概念上的，不能直接数值换算。')

    # ================================================================
    # 7. 完整对照表
    # ================================================================
    h(doc, '7. 完整对照表', level=1)
    tbl(doc,
        ['概念', 'DP（离散）', 'Hamiltonian（连续）', 'ECMS（实时）', '大白话'],
        [
            ['目标函数', 'J[k][i] = min[...]', 'min integral H_fc dt', 'min H_eq', '都是算总氢耗'],
            ['FC 成本', 'g(p_fc)', 'H_fc(P_fc)', 'H_fc(P_fc)', '同一个东西'],
            ['电池成本', 'alpha * (SOC-SOC_ref)^2', 'lambda * d(SOC)/dt', 's * H_bat', '都是电池使用的代价'],
            ['未来信息', 'J[k+1] 查表', 'lambda 微分方程', 's 预估', '都带着未来的影子'],
            ['终端约束', 'beta * (SOC_N-SOC_ref)^2', 'lambda(tf) = beta * dSOC', 's0 + SOC 反馈', '保证终点 SOC 达标'],
            ['最优控制', 'pi[k][i] = argmin', 'dH/dP_fc = 0', 'argmin H_eq', '选最优 FC 功率'],
            ['求解方式', '后向递推（离线）', '两点边值问题', '瞬时最小化（在线）', '计算方式不同'],
        ])

    # ================================================================
    # 8. 面试回答
    # ================================================================
    h(doc, '8. 面试回答模板', level=1)

    p(doc, 'Q: "DP 公式和 ECMS 的 Hamiltonian 不一样，你用的哪个？"', bold=True)
    p(doc,
        'A: "DP 用的是 Bellman 方程：J[k][i] = min[g(p_fc) + alpha*(SOC_next-SOC_ref)^2 + J[k+1]]。'
        'ECMS 用的是 Hamiltonian：H = H_fc + s*H_bat + lambda*d(SOC)/dt。'
        '两个公式最小化的是同一个目标——总氢耗，约束也相同——SOC 动力学。'
        '它们来自不同的数学框架：DP 从 Bellman 的动态规划出发，Hamiltonian 从 Pontryagin 最小原理出发。'
        'DP 的 J[k+1] 和 Hamiltonian 的 lambda 都代表了未来 SOC 的边际成本，'
        '但具体数值上不能直接换算。"')

    p(doc, 'Q: "DP 有 s*H_bat 吗？ECMS 有 J[k+1] 吗？"', bold=True)
    p(doc,
        'A: "DP 不需要 s*H_bat，因为它通过 J[k+1] 已经全局看到了未来——'
        'J[k+1] 就是 DP 版的电池未来成本。ECMS 不能看未来，所以需要 s*H_bat 来预估。'
        '反过来，ECMS 也不需要 J[k+1]，因为它是瞬时决策，不管未来。"')

    p(doc, 'Q: "alpha 和 s 到底什么关系？"', bold=True)
    p(doc,
        'A: "物理意义上，alpha 是 SOC 偏离 1 单位的罚金，s 是电池用 1kWh 的等效氢耗。'
        '通过协态变量 lambda 可以建立联系：s = lambda / (Q_bat * V_oc)。'
        '但 lambda 本身没法直接测，需要通过 DP 反推或者数值求解才能得到。'
        '所以 alpha 和 s 之间没有简单的换算公式。'
        '工程上更可靠的做法是用 DP 反推最优 s*(t) = -H_fc/P_bat 来标定 ECMS。"')

    # ================================================================
    # Save
    # ================================================================
    out_path = os.path.join(DOCS_DIR, 'DP_ECMS_Hamiltonian_公式关系详解.docx')
    doc.save(out_path)
    print(f'[OK] 已生成: {out_path}')
    print(f'     大小: {os.path.getsize(out_path) / 1024:.0f} KB')

if __name__ == '__main__':
    main()
