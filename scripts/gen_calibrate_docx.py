#!/usr/bin/env python3
"""批量将 docs/ 下的 .md 转换为 .docx，然后删除 .md 原文件"""
import os, re, sys, glob
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')


def parse_markdown(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    elements = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        # 代码块
        if line.startswith('```'):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip('\n').startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1
            elements.append({'type': 'code', 'lang': lang, 'text': '\n'.join(code_lines)})
            continue
        # 分隔线
        if re.match(r'^-{3,}\s*$', line.strip()):
            elements.append({'type': 'hr'})
            i += 1
            continue
        # 块引用
        if line.startswith('>'):
            quote_parts = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                q = lines[i].strip()
                q = re.sub(r'^>\s?', '', q).replace('**', '')
                quote_parts.append(q)
                i += 1
            elements.append({'type': 'quote', 'text': '\n'.join(quote_parts)})
            continue
        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            title_text = m.group(2).strip().replace('**', '')
            elements.append({'type': 'heading', 'level': level, 'text': title_text})
            i += 1
            continue
        # 表格
        if '|' in line and line.strip().startswith('|'):
            rows = []
            rows.append([c.strip() for c in line.strip('| \n').split('|')])
            i += 1
            if i < len(lines) and re.match(r'^[\|\s\-:]+$', lines[i].strip()):
                i += 1
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip('| \n').split('|')]
                rows.append(row)
                i += 1
            elements.append({'type': 'table', 'rows': rows})
            continue
        # 空行
        if line.strip() == '':
            i += 1
            continue
        # 列表
        if line.startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            elements.append({'type': 'list', 'items': items})
            continue
        # 普通段落
        text = line.strip()
        if text:
            elements.append({'type': 'para', 'text': text})
        i += 1
    return elements


def format_inline(text):
    parts = re.split(r'\*\*(.+?)\*\*', text)
    runs = []
    for j, part in enumerate(parts):
        if not part:
            continue
        runs.append({'text': part, 'bold': j % 2 == 1})
    return runs


def add_para(doc, text, size=10, spacing=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing)
    runs = format_inline(text)
    if not runs:
        runs = [{'text': text, 'bold': False}]
    for ri in runs:
        r = p.add_run(ri['text'])
        r.font.size = Pt(size)
        r.font.name = 'Arial'
        r.bold = ri['bold']


def clean_cell(text):
    return text.replace('`', '').replace('**', '')


def add_table(doc, rows):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    for i, rd in enumerate(rows):
        rd = rd + [''] * (num_cols - len(rd))
        for j, ct in enumerate(rd):
            cell = table.cell(i, j)
            cell.text = clean_cell(str(ct))
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
                    if i == 0:
                        run.bold = True


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    for line in text.split('\n'):
        r = p.add_run(line + '\n')
        r.font.size = Pt(8.5)
        r.font.name = 'Consolas'
        r.font.color.rgb = RGBColor(30, 30, 30)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '0066CC')
    pBdr.append(left)
    pPr.append(pBdr)
    for line in text.split('\n'):
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        runs = format_inline(line)
        for ri in runs:
            r = p2.add_run(ri['text'])
            r.font.size = Pt(9)
            r.font.name = 'Arial'
            r.bold = ri['bold']
            r.italic = True


HEADING_SIZES = {1: 18, 2: 15, 3: 12, 4: 11, 5: 10, 6: 10}


def convert_one(md_path, docx_path):
    """转换单个 md → docx"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    elements = parse_markdown(md_path)

    for elem in elements:
        t = elem['type']
        if t == 'heading':
            h = doc.add_heading(elem['text'], level=elem['level'])
            for run in h.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(HEADING_SIZES.get(elem['level'], 11))
        elif t == 'hr':
            add_hr(doc)
        elif t == 'quote':
            add_quote(doc, elem['text'])
        elif t == 'code':
            add_code(doc, elem['text'])
        elif t == 'table':
            doc.add_paragraph()
            add_table(doc, elem['rows'])
            doc.add_paragraph()
        elif t == 'list':
            for item in elem['items']:
                p = doc.add_paragraph(item, style='List Bullet')
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                p.paragraph_format.space_after = Pt(2)
        elif t == 'para':
            add_para(doc, elem['text'], size=10, spacing=4)

    doc.save(docx_path)
    return len(elements)


def main():
    # 找到 docs/ 下所有 .md（不含子目录太深的，最多 2 层）
    md_files = sorted(glob.glob(os.path.join(DOCS_DIR, '*.md'))) + \
               sorted(glob.glob(os.path.join(DOCS_DIR, '*', '*.md')))

    if not md_files:
        print('[!] 未找到任何 .md 文件')
        return

    print(f'找到 {len(md_files)} 个 .md 文件\n')

    ok, fail = 0, 0
    for md_path in md_files:
        rel = os.path.relpath(md_path, PROJECT_ROOT)
        docx_path = os.path.splitext(md_path)[0] + '.docx'

        try:
            n = convert_one(md_path, docx_path)
            ok += 1
            size_kb = os.path.getsize(docx_path) / 1024
            print(f'  [OK] {rel} -> .docx ({n} 元素, {size_kb:.0f}KB)')

            # 删除 md 原文件
            os.remove(md_path)
            print(f'       已删除 {os.path.basename(md_path)}')
        except Exception as e:
            fail += 1
            print(f'  [FAIL] {rel}: {e}')

    print(f'\n完成: {ok} 成功, {fail} 失败')


if __name__ == '__main__':
    main()
