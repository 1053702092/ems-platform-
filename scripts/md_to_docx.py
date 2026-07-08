#!/usr/bin/env python3
"""将 tune_aecms_代码逐行分析.md 转换为 tune_aecms_代码逐行分析.docx"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── 配置 ──
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')
MD_FILE = os.path.join(DOCS_DIR, 'tune_aecms_代码逐行分析.md')
DOCX_FILE = os.path.join(DOCS_DIR, 'tune_aecms_代码逐行分析.docx')

def parse_markdown(md_path):
    """极简 Markdown 解析器：支持标题(##)、列表(>、-、|)、代码块( ``` )、粗体(**)、分隔线(---)"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    elements = []  # list of dict: {type, ...}
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 代码块
        if line.startswith('```'):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip('\n').startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```
            elements.append({'type': 'code', 'lang': lang, 'text': '\n'.join(code_lines)})
            continue

        # 分隔线
        if re.match(r'^-{3,}\s*$', line.strip()):
            elements.append({'type': 'hr'})
            i += 1
            continue

        # 块引用 >
        if line.startswith('>'):
            quote_parts = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                q = lines[i].strip()
                q = re.sub(r'^>\s?', '', q)
                # 去除内部粗体标记，保留文字
                q = q.replace('**', '')
                quote_parts.append(q)
                i += 1
            elements.append({'type': 'quote', 'text': '\n'.join(quote_parts)})
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            title_text = m.group(2).strip()
            # 解析内部粗体
            title_text = title_text.replace('**', '')
            elements.append({'type': 'heading', 'level': level, 'text': title_text})
            i += 1
            continue

        # 表格
        if '|' in line and line.strip().startswith('|'):
            rows = []
            # 读首行
            rows.append([c.strip() for c in line.strip('| \n').split('|')])
            i += 1
            # 跳过分隔行 ---|---|
            if i < len(lines) and re.match(r'^[\|\s\-:]+$', lines[i].strip()):
                i += 1
            # 读数据行
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip('| \n').split('|')]
                rows.append(row)
                i += 1
            elements.append({'type': 'table', 'rows': rows})
            continue

        # 空行
        if line.strip() == '':
            i += 1
            continue

        # 普通段落（可能是列表项或正文）
        if line.startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            elements.append({'type': 'list', 'items': items})
            continue

        # 普通文本段落
        text = line.strip()
        if text:
            elements.append({'type': 'para', 'text': text})
        i += 1

    return elements


def format_inline(text):
    """处理内联粗体 **text** → (text, bold=True) 对"""
    parts = re.split(r'\*\*(.+?)\*\*', text)
    runs = []
    for j, part in enumerate(parts):
        if not part:
            continue
        bold = (j % 2 == 1)  # 偶数索引是普通文本，奇数是 **/** 包裹的
        runs.append({'text': part, 'bold': bold})
    return runs


def add_paragraph_with_runs(doc, text, bold=False, size=10, spacing=4, alignment=None):
    """添加带内联格式的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing)
    runs = format_inline(text)
    if not runs:
        runs = [{'text': text, 'bold': bold}]
    for run_info in runs:
        r = p.add_run(run_info['text'])
        r.font.size = Pt(size)
        r.font.name = 'Arial'
        r.bold = run_info['bold']
    if alignment is not None:
        p.alignment = alignment
    return p


def add_multiline_runs(doc, text, size=9, spacing=2):
    """多行文本，每行一个 paragraph（用于引用块中的多行）"""
    lines = text.split('\n')
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(spacing if idx < len(lines) - 1 else 0)
        runs = format_inline(line)
        for run_info in runs:
            r = p.add_run(run_info['text'])
            r.font.size = Pt(size)
            r.font.name = 'Arial'
            r.bold = run_info['bold']
            r.italic = True


def clean_cell(text):
    """清理单元格文本：去掉 Markdown 代码标记"""
    return text.replace('`', '').replace('**', '')

def add_table_to_doc(doc, rows):
    """添加表格"""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        # 补齐列数
        row_data = row_data + [''] * (num_cols - len(row_data))
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = clean_cell(str(cell_text))
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
                    if i == 0:
                        run.bold = True


def add_code_block(doc, text, lang=''):
    """添加代码块（使用等宽字体 + 浅灰背景）"""
    # 代码块用表格模拟浅灰背景
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 设表格宽度占满
    for section in table.rows[0].cells[0].paragraphs:
        section.clear()

    cell = table.rows[0].cells[0]
    # 设置单元格底色
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    for line in text.split('\n'):
        r = p.add_run(line + '\n')
        r.font.size = Pt(8.5)
        r.font.name = 'Consolas'
        r.font.color.rgb = RGBColor(30, 30, 30)

    return table


def elements_to_docx(elements, doc):
    """将解析后的元素列表写入 doc"""
    heading_sizes = {1: 18, 2: 15, 3: 12, 4: 11, 5: 10, 6: 10}
    heading_spacings = {1: 12, 2: 8, 3: 6, 4: 4, 5: 4, 6: 2}

    for elem in elements:
        t = elem['type']

        if t == 'heading':
            level = elem['level']
            text = elem['text']
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(heading_sizes.get(level, 11))

        elif t == 'hr':
            # 添加一条水平线
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif t == 'quote':
            text = elem['text']
            # 添加左侧竖线样式
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '0066CC')
            pBdr.append(left)
            pPr.append(pBdr)

            add_multiline_runs(doc, text, size=9, spacing=2)

        elif t == 'code':
            add_code_block(doc, elem['text'], elem.get('lang', ''))

        elif t == 'table':
            doc.add_paragraph()  # 空行间隔
            add_table_to_doc(doc, elem['rows'])
            doc.add_paragraph()  # 空行间隔

        elif t == 'list':
            for item in elem['items']:
                p = doc.add_paragraph(item, style='List Bullet')
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                p.paragraph_format.space_after = Pt(2)

        elif t == 'para':
            # 检查是否包含表格行分隔符样式（如 `|---|---|` 已跳过）
            add_paragraph_with_runs(doc, elem['text'], size=10, spacing=4)


def main():
    doc = Document()

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 解析 Markdown
    elements = parse_markdown(MD_FILE)
    print(f'[OK] 解析了 {len(elements)} 个元素')

    # 写入 docx
    elements_to_docx(elements, doc)

    # 保存
    doc.save(DOCX_FILE)
    print(f'[OK] 已生成: {DOCX_FILE}')


if __name__ == '__main__':
    main()
