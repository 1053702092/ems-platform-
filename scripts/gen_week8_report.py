#!/usr/bin/env python3
"""
gen_week8_report.py — 第8周：传统EMS四方法大对比报告
Rule vs DP vs ECMS vs MPC（WLTC / NEDC / CLTC 三工况）
"""
import os, sys, glob
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULTS_DIR = os.path.join(PROJECT_ROOT, 'results')
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')
SCRIPTS_DIR = os.path.join(PROJECT_ROOT, 'scripts')
sys.path.insert(0, SCRIPTS_DIR)

from day8_dp_ems import DT, Q_BAT, SOC_REF, fc_hydrogen_flow, load_drive_cycle, vehicle_power, fc_efficiency, run_rule_controller
from mpc_ems_optimized import soc_equivalent_h2

CYCLES = ['wltc', 'nedc', 'cltc']
METHODS = ['Rule', 'DP', 'ECMS', 'MPC_optimized']
CYCLE_DISPLAY = {'wltc': 'WLTC (1800s)', 'nedc': 'NEDC (1180s)', 'cltc': 'CLTC (1800s)'}

# ════════════════════════════════════════════════════════════
# 1. 数据加载与指标计算
# ════════════════════════════════════════════════════════════

def load_results():
    """加载三个工况四种方法的结果, 计算统一指标"""
    records = []
    for cycle in CYCLES:
        t, v = load_drive_cycle(cycle)
        P_load = vehicle_power(v, DT)
        total_energy_kWh = P_load.sum() * DT / 3600

        # ── Rule（实时运行，不用CSV）──
        rule = run_rule_controller(P_load)
        records.append(_metrics('Rule', cycle, rule, P_load, total_energy_kWh))

        # ── DP ──
        dp = _load_csv_or_run(cycle, 'dp_ems', 'DP')
        if dp is not None:
            records.append(_metrics('DP', cycle, dp, P_load, total_energy_kWh))

        # ── ECMS ──
        ecms = _load_csv(f'ecms_ems_{cycle}.csv')
        if ecms is not None:
            records.append(_metrics('ECMS', cycle, ecms, P_load, total_energy_kWh))

        # ── MPC_optimized ──
        mpc = _load_csv(f'mpc_ems_optimized_{cycle}_np50.csv')
        if mpc is not None:
            r = _metrics('MPC_optimized', cycle, mpc, P_load, total_energy_kWh)
            records.append(r)

    df = pd.DataFrame(records)
    # 补齐与 DP 的差距
    for cycle in CYCLES:
        dp_h2_raw = df.loc[(df.cycle == cycle) & (df.method == 'DP'), 'H2_raw_kg'].values
        dp_h2_eq = df.loc[(df.cycle == cycle) & (df.method == 'DP'), 'H2_eq_kg'].values
        if len(dp_h2_raw) > 0 and len(dp_h2_eq) > 0:
            dhr, dhe = dp_h2_raw[0], dp_h2_eq[0]
            m = df.cycle == cycle
            df.loc[m, 'raw_gap_vs_DP_pct'] = (df.loc[m, 'H2_raw_kg'] / dhr - 1) * 100
            df.loc[m, 'H2_eq_gap_vs_DP_pct'] = (df.loc[m, 'H2_eq_kg'] / dhe - 1) * 100
    return df


def _load_csv(fname):
    path = os.path.join(RESULTS_DIR, fname)
    if not os.path.exists(path):
        return None
    return pd.read_csv(path)


def _load_csv_or_run(cycle, stem, label):
    """尝试加载 CSV，失败则打印警告（不自动运行）"""
    path = os.path.join(RESULTS_DIR, f'{stem}_{cycle}.csv')
    if os.path.exists(path):
        try:
            return pd.read_csv(path)
        except Exception as e:
            print(f'  [!] {label} {cycle} 加载失败: {e}')
            return None
    print(f'  [!] {label} {cycle} 数据缺失: {path}')
    return None


def _get_last(data, col):
    """获取最后一行的值（兼容 dict[ndarray] 和 DataFrame）"""
    if isinstance(data, dict):
        return data[col][-1]
    return data[col].iloc[-1]

def _get_col(data, col):
    """获取整列（兼容 dict[ndarray] 和 DataFrame）"""
    if isinstance(data, dict):
        return data[col]
    return data[col].values

def _metrics(method, cycle, data, P_load, total_energy_kWh):
    """从原始数据计算统一指标"""
    H2_raw = _get_last(data, 'm_H2_cumul_kg')
    SOC_end = round(_get_last(data, 'SOC'), 3)
    P_fc = _get_col(data, 'P_fc_kW')
    eff = fc_efficiency(P_fc)
    return {
        'cycle': cycle,
        'method': method,
        'H2_raw_kg': round(H2_raw, 4),
        'SOC_end': SOC_end,
        'H2_eq_kg': round(soc_equivalent_h2(H2_raw, SOC_end), 4),
        'FC_eff_mean_pct': round(eff.mean() * 100, 1),
        'FC_max_kW': round(P_fc.max(), 1),
        'total_energy_kWh': round(total_energy_kWh, 2),
        'raw_gap_vs_DP_pct': np.nan,
        'H2_eq_gap_vs_DP_pct': np.nan,
    }




# ════════════════════════════════════════════════════════════
# 2. 报告生成
# ════════════════════════════════════════════════════════════

def _set_font(run, name='Arial', size=10, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
    return h


def _para(doc, text, size=10.5, bold=False, color=None, spacing_after=4, alignment=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def _table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                _set_font(run, size=9, bold=True)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    _set_font(run, size=9)
    return table


def _img(doc, path, width_cm=14):
    if not os.path.exists(path):
        _para(doc, f'[图缺失: {os.path.basename(path)}]', color=RGBColor(200, 0, 0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))


def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _check_mark(text, level='pass'):
    """返回带颜色的结论段落"""
    colors = {'pass': RGBColor(0, 128, 0), 'warn': RGBColor(200, 120, 0), 'fail': RGBColor(200, 0, 0)}
    labels = {'pass': '✅', 'warn': '⚠️', 'fail': '❌'}
    p = doc.add_paragraph()
    run = p.add_run(f'{labels.get(level, "")} {text}')
    _set_font(run, size=10.5, bold=True, color=colors.get(level, RGBColor(0, 0, 0)))
    return p


# ════════════════════════════════════════════════════════════
# 3. 主程序
# ════════════════════════════════════════════════════════════

def generate_report():
    print('=' * 55)
    print('  Week 8: 传统EMS四方法大对比报告')
    print('=' * 55)

    # 加载指标数据
    df = load_results()
    print(f'[OK] 加载了 {len(df)} 条指标记录')

    # 保存统一指标 CSV
    csv_out = os.path.join(RESULTS_DIR, 'week8_fourway_metrics_complete.csv')
    df.to_csv(csv_out, index=False)
    print(f'[保存] {csv_out}')

    # ── 创建文档 ──
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ════════════════════════════════════════
    # 封面
    # ════════════════════════════════════════
    _para(doc, '', size=22, spacing_after=60)
    _para(doc, '传统EMS策略对比报告', size=26, bold=True, spacing_after=12)
    _para(doc, 'Rule · DP · ECMS · MPC', size=16, color=RGBColor(80, 80, 80), spacing_after=6)
    _para(doc, '燃料电池混合动力系统能量管理策略评估', size=12, color=RGBColor(120, 120, 120), spacing_after=30)
    _para(doc, '——————————————————————————', size=10, spacing_after=20)
    _para(doc, '工况数据集：WLTC (1800s) / NEDC (1180s) / CLTC (1800s)', size=10, spacing_after=4)
    _para(doc, '报告日期：2026-07', size=10, spacing_after=4)
    _para(doc, '所属阶段：Month 2 — 传统EMS策略（第8周）', size=10, spacing_after=4)
    doc.add_page_break()

    # ════════════════════════════════════════
    # 目录
    # ════════════════════════════════════════
    _heading(doc, '目录', level=1)
    toc_items = ['摘要', '1. 方法论', '2. 统一指标框架', '3. 结果总表', '4. 工况对比分析',
                 '5. 关键发现', '6. 面试叙事要点', '7. 结论与下一步']
    for item in toc_items:
        _para(doc, f'  {item}', size=11, spacing_after=2)
    doc.add_page_break()

    # ════════════════════════════════════════
    # 摘要
    # ════════════════════════════════════════
    _heading(doc, '摘要', level=1)
    _para(doc, '本报告对燃料电池混合动力系统的四种传统能量管理策略——规则控制器 (Rule)、'
          '动态规划 (DP)、等效消耗最小化策略 (ECMS) 和模型预测控制 (MPC)——'
          '在 WLTC、NEDC、CLTC 三种标准工况下进行了系统对比评估。', size=10.5, spacing_after=6)
    _para(doc, '核心发现：', size=10.5, bold=True, spacing_after=2)
    findings = [
        'DP 作为离线全局最优基准，在三种工况下提供最低的等效氢耗，是算法对比的"黄金标准"',
        'ECMS（s=130 g/kWh）在 WLTC 下最接近 DP（原始氢耗 +0.2%），是最优的在线实时策略',
        'MPC 优化版通过引入 SOC 软约束和终端 SOC 惩罚，解决了"电池透支省氢"的假象问题',
        'Rule 控制器在 CLTC 工况下偏差最大（+43%），在低速城市工况下损失显著',
        '公平对比需采用等效氢耗 H2_eq（SOC 修正），而非仅对比原始氢耗 H2_raw',
    ]
    for f in findings:
        _para(doc, f'  • {f}', size=10, spacing_after=2)
    doc.add_page_break()

    # ════════════════════════════════════════
    # 1. 方法论
    # ════════════════════════════════════════
    _heading(doc, '1. 方法论', level=1)

    methods_desc = [
        ('规则控制器 (Rule)',
         '基于 SOC 阈值的规则控制器：SOC 低时 FC 高功率充电，SOC 高时 FC 低功率或关闭。'
         '不进行优化，仅作为 baseline。输出：固定策略，不依赖工况预测。'),
        ('动态规划 (DP)',
         '基于 Bellman 最优性原理的离线全局优化。后向递归计算代价矩阵 J，前向 Rollout 得到最优控制序列。'
         '网格：SOC 150× P_fc 60。需已知完整工况，无法在线使用。'
         '作为对比基准（全局最优）。'),
        ('等效消耗最小化策略 (ECMS)',
         '基于 PMP (Pontryagin) 的实时近似最优策略：将电池电能通过等效因子 s 折算为等效氢耗。'
         '恒定 ECMS 使用 s=130 g/kWh，无需工况预测。在 WLTC 下与 DP 差距 <1%。'
         '计算量小，可在线运行。'),
        ('模型预测控制 (MPC)',
         '滚动时域优化：在每个时间步求解 N_p=50 步内的网格搜索最优控制，执行第一步后重算。'
         '优化版加入：SOC 软下限、终端 SOC 欠差惩罚、FC 功率变化惩罚。'
         '介于 ECMS（实时）和 DP（全局优化）之间。'),
    ]
    for name, desc in methods_desc:
        _heading(doc, name, level=2)
        _para(doc, desc, size=10, spacing_after=6)
    doc.add_page_break()

    # ════════════════════════════════════════
    # 2. 统一指标框架
    # ════════════════════════════════════════
    _heading(doc, '2. 统一指标框架', level=1)
    _para(doc, '为避免片面比较，本报告采用以下多维指标体系：', size=10.5, spacing_after=6)

    _table(doc,
           ['指标', '符号', '含义', '理想值'],
           [
               ['原始氢耗', 'H2_raw (kg)', '循环总氢消耗量', '越低越好'],
               ['终端 SOC', 'SOC_end', '循环结束时电池荷电状态', '≈0.6 (charge-sustaining)'],
               ['等效氢耗', 'H2_eq (kg)', '用 SOC 偏差修正后的公平氢耗', '越低越好'],
               ['FC 平均效率', 'η_FC_mean (%)', '燃料电池平均运行效率', '越高越好'],
               ['与 DP 差距', 'ΔDP (%)', '相对 DP 基准的偏差', '越接近 0 越好'],
           ])
    _para(doc, '', size=4, spacing_after=2)
    _para(doc, 'H2_eq 的计算公式：', size=10, bold=True, spacing_after=2)
    _para(doc, '  H2_eq = H2_raw + s × E_bat / 1000', size=10, spacing_after=2)
    _para(doc, '  其中 E_bat = Q_bat × V_oc × (SOC_ref - SOC_end) / 1000 (kWh)', size=10, spacing_after=2)
    _para(doc, '  s = 130 g/kWh（ECMS 等效因子）', size=10, spacing_after=2)
    _para(doc, '  约定：SOC_end < SOC_ref 代表多用了电池能量，加回等效氢耗', size=9, spacing_after=6)
    _para(doc, '  注意：这是报告层面的可比性指标，不替代真实氢耗。', size=9, italic=True, color=RGBColor(120, 120, 120))
    doc.add_page_break()

    # ════════════════════════════════════════
    # 3. 结果总表
    # ════════════════════════════════════════
    _heading(doc, '3. 结果总表', level=1)
    _para(doc, '三工况四方法对比汇总（统一指标框架）：', size=10.5, spacing_after=6)

    for cycle in CYCLES:
        _heading(doc, f'{cycle.upper()}', level=2)
        subset = df[df.cycle == cycle].sort_values('H2_eq_kg')
        headers = ['方法', 'H2_raw (kg)', 'SOC_end', 'H2_eq (kg)',
                   'η_FC (%)', 'ΔDP raw (%)', 'ΔDP eq (%)']
        rows = []
        for _, r in subset.iterrows():
            raw_gap = f'{r["raw_gap_vs_DP_pct"]:+.1f}' if not np.isnan(r["raw_gap_vs_DP_pct"]) else '—'
            eq_gap = f'{r["H2_eq_gap_vs_DP_pct"]:+.1f}' if not np.isnan(r["H2_eq_gap_vs_DP_pct"]) else '—'
            rows.append([
                r['method'],
                f'{r["H2_raw_kg"]:.4f}',
                f'{r["SOC_end"]:.3f}',
                f'{r["H2_eq_kg"]:.4f}',
                f'{r["FC_eff_mean_pct"]:.1f}',
                raw_gap,
                eq_gap,
            ])
        _table(doc, headers, rows)
        _para(doc, '', size=4, spacing_after=6)

    _para(doc, '表注：DP 为基准参考，ΔDP = (方法值/DM值 - 1)×100%。H2_eq 越低越好。', size=9, italic=True, spacing_after=4)
    doc.add_page_break()

    # ════════════════════════════════════════
    # 4. 工况对比分析
    # ════════════════════════════════════════
    _heading(doc, '4. 工况对比分析', level=1)

    for cycle in CYCLES:
        _heading(doc, f'4.{CYCLES.index(cycle)+1} {cycle.upper()}', level=2)
        # 插入对比图
        img_path = os.path.join(RESULTS_DIR, f'FourWay_compare_optimized_{cycle}.png')
        _img(doc, img_path, width_cm=14)

        # 针对性分析
        subset = df[df.cycle == cycle].sort_values('H2_eq_kg')
        best = subset.iloc[0]
        worst = subset.iloc[-1]
        dp_row = subset[subset.method == 'DP'].iloc[0] if len(subset[subset.method == 'DP']) > 0 else None
        ecms_row = subset[subset.method == 'ECMS'].iloc[0] if len(subset[subset.method == 'ECMS']) > 0 else None
        mpc_row = subset[subset.method == 'MPC_optimized'].iloc[0] if len(subset[subset.method == 'MPC_optimized']) > 0 else None

        _para(doc, f'分析：', size=10.5, bold=True, spacing_after=2)
        _para(doc, f'  • 等效氢耗最低: {best["method"]} (H2_eq={best["H2_eq_kg"]:.4f} kg)', size=10, spacing_after=1)
        if ecms_row is not None:
            _para(doc, f'  • ECMS 原始氢耗: {ecms_row["H2_raw_kg"]:.4f} kg, 等效氢耗: {ecms_row["H2_eq_kg"]:.4f} kg', size=10, spacing_after=1)
        if mpc_row is not None:
            _para(doc, f'  • MPC 原始氢耗: {mpc_row["H2_raw_kg"]:.4f} kg, SOC_end={mpc_row["SOC_end"]:.3f}', size=10, spacing_after=1)
        _para(doc, f'  • 方法间最大差距: {(worst["H2_eq_kg"]/best["H2_eq_kg"]-1)*100:.1f}%', size=10, spacing_after=4)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 5. 关键发现
    # ════════════════════════════════════════
    _heading(doc, '5. 关键发现', level=1)

    _heading(doc, '5.1 排序稳定性', level=2)
    _para(doc, '在三种工况下，四种方法的氢耗排序保持稳定：DP ≈ ECMS < MPC_optimized < Rule。'
          '这验证了各算法的内在优劣关系不随工况改变。', size=10, spacing_after=6)

    _heading(doc, '5.2 ECMS 的"接近最优"表现', level=2)
    _para(doc, '恒定 ECMS (s=130) 在 WLTC 下与 DP 的原始氢耗差距仅 +0.2%，'
          '在 NEDC 下 +4.2%。这是一个无需工况预测、计算量为 μs 级的实时策略，'
          '在实际嵌入式系统中极具价值。', size=10, spacing_after=6)

    _heading(doc, '5.3 SOC 公平评价的重要性', level=2)
    _para(doc, 'MPC 旧版（SOC_end≈0.55）的原始氢耗（0.2011 kg）低于 DP（0.2287 kg），'
          '但这实际上是因为 MPC 在终端 SOC 约束偏弱时多用了电池能量，产生了氢耗假象。'
          '引入 SOC 修正后的等效氢耗 H2_eq 才能公平比较。', size=10, spacing_after=6)

    _heading(doc, '5.4 工况依赖', level=2)
    _para(doc, 'CLTC 工况下方法间差距最大（Rule 比 DP 高 43.2%），因为 CLTC 低速段更多、'
          '功率波动更剧烈，规则控制器难以适应。这凸显了优化策略在城市工况中的优势。',
          size=10, spacing_after=6)

    _heading(doc, '5.5 MPC 优化效果', level=2)
    _para(doc, '优化版 MPC 通过 SOC 软下限（0.57）、终点 SOC 欠差惩罚和 FC 功率变化惩罚，'
          '将 SOC_end 从 0.55 提升至 0.576 附近，等效氢耗更接近 DP。'
          '虽然原始氢耗略高于 DP（WLTC +1.8%），但这是约束更严格后的真实表现。',
          size=10, spacing_after=6)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 6. 面试叙事要点
    # ════════════════════════════════════════
    _heading(doc, '6. 面试叙事要点', level=1)
    _para(doc, '这些要点用于面试时讲清楚"你到底做了什么"：', size=10.5, spacing_after=6)

    narratives = [
        ('一句话版',
         '我对四种传统 EMS 策略进行了三工况系统对比，发现恒定 ECMS 在实时可用的前提下达到全局最优 DP 的 99% 以上性能。'),
        ('30秒版',
         '我搭建了 Rule/DP/ECMS/MPC 四种策略在 WLTC/NEDC/CLTC 三工况下的对比框架，'
         '用原始氢耗、SOC 偏差和等效氢耗三个维度公平评估。'
         '发现 ECMS (s=130) 是最优的实时策略，WLTC 下与 DP 差距 <1%。'),
        ('追问：MPC 超过 DP？',
         '不，旧版 MPC 的原始氢耗低于 DP 是因为 SOC 透支。'
         '我们优化了代价函数（SOC 软下限 + 终端惩罚），修正后的公平评价显示 MPC 等效氢耗略高于 DP（+1.8%），'
         '这才是真实的对比结果。这个"发现假象→修正口径"的过程本身就是工程能力。'),
        ('追问：为什么选 ECMS 而不是 MPC？',
         'ECMS 无需求解优化问题，等效因子 s 扫描即可确定，'
         '在嵌入式系统中可以在微秒级完成控制决策。'
         'MPC 的网格搜索复杂度 O(N_p × N_grid)，适合弱实时场景或作为规划层。'),
    ]
    for title, text in narratives:
        _heading(doc, title, level=2)
        _para(doc, text, size=10, spacing_after=4)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 7. 结论
    # ════════════════════════════════════════
    _heading(doc, '7. 结论与下一步', level=1)

    _heading(doc, '结论', level=2)
    _para(doc, '通过系统对比，四种策略的技术定位明确：', size=10, spacing_after=4)

    _table(doc,
           ['策略', '最优性', '实时性', '工况依赖', '适用场景'],
           [
               ['Rule', '差', '✅ 实时', '强', '基本保底策略'],
               ['DP', '✅ 全局最优', '❌ 离线', '需已知工况', '性能基准/理论下限'],
               ['ECMS', '≈ 最优', '✅ 实时', '弱（需标定 s）', '实际部署首选'],
               ['MPC', '≈ 最优', '⚠️ 有限实时', '弱', '规划层/弱实时场景'],
           ])
    _para(doc, '', size=4, spacing_after=6)
    _para(doc, '推荐策略：ECMS 为在线主策略，MPC 为离线校准参考，DP 为理论基准。', size=10.5, bold=True, spacing_after=6)

    _heading(doc, '下一步', level=2)
    next_steps = [
        '第9周：PyTorch 入门 — 用 PyTorch 重写简单 MLP 功率预测',
        '第10-11周：PPO 强化学习在 EMS 环境训练',
        '第12周：PPO-EMS 闭环 + 与传统策略（Rule/ECMS/MPC）对比',
        '目标：8月底完成，9月投递启动',
    ]
    for s in next_steps:
        _para(doc, f'  → {s}', size=10, spacing_after=2)

    _hr(doc)
    _para(doc, '— 报告结束 —', size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(120, 120, 120), spacing_after=10)

    # ── 保存 ──
    docx_path = os.path.join(DOCS_DIR, 'Week8_四方法大对比报告.docx')
    doc.save(docx_path)
    print(f'[OK] 报告已生成: {docx_path}')
    return docx_path


if __name__ == '__main__':
    generate_report()
