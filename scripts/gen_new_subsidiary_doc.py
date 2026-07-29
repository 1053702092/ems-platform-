#!/usr/bin/env python3
"""Generate updated DOCX for newly established SOE subsidiaries."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_title(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_heading2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)


def add_heading3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)


def add_para(text, bold=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style=f"List {'Bullet' if level == 0 else f'Bullet {level+1}'}")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return p


def make_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Arial"

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"
    doc.add_paragraph()
    return tbl


# ============================================================
# DOCUMENT CONTENT
# ============================================================

add_title("2025-2026 年新设/扩产央企国企及子公司名单")

add_para(
    "筛选范围：福建、广东、浙江、江苏、湖北 | 整理日期：2026-07-29 | 版本：V2 更新版\n"
    "背景：211本 + 双非硕 + 化学/材料 + EMS算法项目经验 | 目标城市：南方为主",
    size=9,
)

add_para(
    "说明：以下企业均为 2025-2026 年已官宣新设子公司/新基地/新合资公司，预计 2026-2027 年投产或大规模招聘。"
    "新设公司的特点是刚成立需要招人、HC多、竞争相对成熟企业小——是应届生进央企的黄金窗口。",
    bold=False,
    size=9.5,
)

# ──────────────────────────────────────────────
# Section 1: 产能扩产类新基地
# ──────────────────────────────────────────────
add_heading2("一、产能扩产类新基地（2025开工 → 2026投产招聘）")
add_para("这些是成熟企业的新建基地，扩产期人员需求大，招聘确定性高。", size=9.5)

make_table(
    ["#", "企业名称", "地点", "项目方向", "状态", "备注"],
    [
        ["1", "宁德时代 — 惠州基地", "广东惠州", "储能电池产线", "待核验", "成熟企业扩产"],
        ["2", "比亚迪储能 — 揭阳基地", "广东揭阳", "储能系统", "待核验", "比亚迪系，需求量大"],
        ["3", "国轩高科 — 福州基地", "福建福州", "储能电池", "待核验", "BMS算法岗高度对口"],
        ["4", "瑞浦兰钧 — 太仓基地", "江苏太仓", "储能电池", "待核验", "青山集团系"],
        ["5", "海辰储能 — 漳州基地", "福建漳州", "储能电芯", "待核验", "快速扩产中"],
        ["6", "远景动力 — 湛江基地", "广东湛江", "储能电池", "待核验", "远景系"],
        ["7", "三花储能 — 杭州基地", "浙江杭州", "储能热管理/集成", "待核验", "三花系"],
        ["8", "中创新航 — 惠州基地", "广东惠州", "储能电池", "待核验", "你最稳的选择之一"],
        ["9", "海博思创 — 佛山基地", "广东佛山", "储能系统集成", "待核验", "系统集成龙头"],
    ],
)

# ──────────────────────────────────────────────
# Section 2: 新设子公司/合资公司
# ──────────────────────────────────────────────
add_heading2("二、新设子公司 & 合资公司（央企国企为主）")
add_para("这些是真正的新设主体——刚注册或刚成立不久，招聘窗口期最宝贵。", size=9.5)

make_table(
    ["#", "企业名称", "地点", "方向/岗位", "背景", "你的匹配度"],
    [
        ["1", "国宁新储（福建）科技", "宁德/福州", "储能系统集成工程师 / 仿真工程师", "央企新设储能子公司", "⭐⭐⭐⭐⭐"],
        ["2", "中国电气装备储能科技", "上海", "EMS软件工程师（30万/年）", "央企新设", "⭐⭐⭐⭐⭐"],
        ["3", "新能时代储能", "福建多地", "储能全产业链", "宁德时代系新设", "⭐⭐⭐"],
        ["4", "广东储能产业发展", "广东", "储能产业平台", "省级储能平台", "⭐⭐⭐⭐"],
        ["5", "广东新型储能国家研究院", "广东", "新型储能研发", "省级研究院", "⭐⭐⭐⭐"],
        ["6", "广州储能集团", "广州", "EMS研发（八险二金）", "广州市属储能集团", "⭐⭐⭐⭐"],
        ["7", "中国资源循环集团新能源科技", "深圳", "新能源/储能系统集成", "央企新设", "⭐⭐⭐⭐"],
        ["8", "中国资源循环集团电池公司", "深圳", "电池管理系统/材料研发", "央企新设", "⭐⭐⭐⭐"],
        ["9", "中资环新能源循环利用科技（深圳）", "深圳", "电池回收/新能源循环利用", "央企新设", "⭐⭐⭐"],
        ["10", "湖北省新能源有限公司", "武汉", "开发经理/储能方向/虚拟电厂", "省属新能源平台", "⭐⭐⭐⭐"],
        ["11", "华电福新（华电福建）", "福建", "水电/风电/光伏/储能", "华电集团新能源子公司", "⭐⭐⭐"],
        ["12", "广东华电储能", "广东汕尾", "储能", "华电系新设", "⭐⭐⭐"],
        ["13", "广州星翼智慧能源", "广州南沙", "BMS/EMS（清华系创业）", "清华系创业公司", "⭐⭐⭐⭐"],
    ],
)

# ──────────────────────────────────────────────
# Section 3: 新设重点企业详解
# ──────────────────────────────────────────────
add_heading2("三、重点企业详解 — 你最值得关注的6家")

companies = [
    {
        "name": "1️⃣ 国宁新储（福建）科技有限公司",
        "tag": "🏆 首推",
        "why": "央企（国家能源集团/国网系）在福建新设的储能子公司，储能系统集成工程师/仿真工程师岗位直接对口你的EMS项目。新公司意味着：① HC充足 ② 竞争较小 ③ 晋升空间大。地点宁德/福州，福建人首选。",
        "action": "重点关注其2026秋招公告，预计9-10月启动。可提前搜索'国宁新储 招聘'关注动态。",
    },
    {
        "name": "2️⃣ 中国电气装备集团储能科技有限公司",
        "tag": "💰 高薪",
        "why": "央企直接下场做储能，EMS软件工程师岗标价30万/年（上海），和你做的DP/ECMS/MPC/PPO完全对口。央企+高薪+技术岗，三者兼得。",
        "action": "上海岗位可接受的话，这是你薪资天花板级的央企选择。留意中国电气装备集团校招官网。",
    },
    {
        "name": "3️⃣ 广州储能集团",
        "tag": "🏛️ 稳定",
        "why": "广州市属储能国企，八险二金，EMS研发岗。项目和岗位高度对口，一线城市+国企编制，性价比极高。",
        "action": "广州是目标城市之一，可优先投递。关注广州市国资委招聘公告。",
    },
    {
        "name": "4️⃣ 广东新型储能国家研究院",
        "tag": "🔬 研发",
        "why": "研究院性质，偏研发而非工程。适合喜欢技术深挖的你。且研究院对专业限制通常比企业宽松。",
        "action": "关注其招聘公告，简历突出算法实现和论文调研能力。",
    },
    {
        "name": "5️⃣ 中国资源循环集团（新能源科技 + 电池公司）",
        "tag": "♻️ 央企新赛道",
        "why": "2025年新成立的央企（国务院国资委直接监管），资源循环利用是政策强推赛道。新能源科技子公司做储能系统，电池子公司做BMS，两个方向都匹配。",
        "action": "央企新设，HC多且竞争少。留意中国资源循环集团官网。",
    },
    {
        "name": "6️⃣ 新能时代储能（宁德时代系）",
        "tag": "🔋 CATL系平替",
        "why": "CATL系的独立储能子公司，门槛比CATL总部低，但技术和品牌关联度高。如果CATL总部投不进，这是很好的替代选项。",
        "action": "关注其2026秋招，和CATL同步投递。",
    },
]

for c in companies:
    add_heading3(c["name"])
    # Tag
    tag_p = doc.add_paragraph()
    run = tag_p.add_run(f"【{c['tag']}】")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0xD4, 0x6A, 0x00)

    # Why
    p = doc.add_paragraph()
    run = p.add_run("为什么值得去：")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run2 = p.add_run(c["why"])
    run2.font.size = Pt(10)
    run2.font.name = "Arial"

    # Action
    p2 = doc.add_paragraph()
    run = p2.add_run("建议行动：")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run2 = p2.add_run(c["action"])
    run2.font.size = Pt(10)
    run2.font.name = "Arial"

# ──────────────────────────────────────────────
# Section 4: 完整N档名单纯表
# ──────────────────────────────────────────────
add_heading2("四、完整 N档·新设扩产待核验 名单（按省份）")
add_para("以下为 Excel 中 'N档·新设扩产待核验' 分类的全部条目，按省份分组。", size=9.5)

province_groups = [
    ("福建", [
        ("Ampace新能安", "厦门", "BMS/EMS算法"),
        ("北辰星储能", "厦门", "BMS/EMS技术"),
        ("厦门艾能星储能", "厦门", "储能电气(BMS/EMS)"),
        ("拓普菲斯新能源", "厦门", "储能系统(BMS)"),
        ("索克曼能源", "厦门", "软件(UPS/BMS)"),
        ("储能电池EMS工程师（跨国公司）", "宁德", "25-40k·14薪"),
        ("东南汽车", "福州", "助理电控(BMS)"),
        ("智狐能源", "福州", "EMS研发"),
        ("飞毛腿动力科技", "福州", "BMS算法/SOX"),
        ("华电福新（华电福建）", "福建", "水电/风电/光伏/储能"),
        ("国宁新储(福建)", "福建", "储能系统集成"),
        ("新能时代储能（宁德时代系）", "福建多地", "储能全产业链"),
        ("海辰储能 — 漳州基地", "福建漳州", "储能电芯"),
        ("国轩高科 — 福州基地", "福建福州", "储能电池"),
    ]),
    ("广东", [
        ("ATL(新能源科技)", "东莞", "BMS算法"),
        ("新能德(NVT)", "东莞", "BMS算法"),
        ("广州东信软件", "云浮", "EMS/调度高级工程师"),
        ("广东储能产业发展", "广东", "储能产业平台"),
        ("广东新型储能国家研究院", "广东", "新型储能研发"),
        ("海博思创 — 佛山基地", "广东佛山", "储能系统集成"),
        ("中创新航 — 惠州基地", "广东惠州", "储能电池"),
        ("宁德时代 — 惠州基地", "广东惠州", "储能电池产线"),
        ("比亚迪储能 — 揭阳基地", "广东揭阳", "储能系统"),
        ("广东华电储能", "广东汕尾", "储能"),
        ("远景动力 — 湛江基地", "广东湛江", "储能电池"),
        ("小鹏汽车", "广州", "能量管理/控制"),
        ("广州储能集团", "广州", "EMS研发"),
        ("广州奥鹏能源", "广州", "BMS软件"),
        ("广汽埃安", "广州", "VCU/BMS控制"),
        ("广汽能源科技", "广州", "BMS"),
        ("智光电气", "广州", "储能EMS"),
        ("菲利斯太阳能", "广州", "BMS软件"),
        ("广州星翼智慧能源", "广州南沙", "BMS/EMS(清华系创业)"),
        ("华阳集团", "惠州", "BMS/控制算法"),
        ("德赛电池", "惠州/深圳", "BMS/EMS"),
        ("中国资源循环集团新能源科技", "深圳", "新能源/储能系统集成"),
        ("中国资源循环集团电池公司", "深圳", "电池管理系统"),
        ("中资环新能源循环利用", "深圳", "电池回收"),
        ("华宝新能", "深圳", "BMS/EMS"),
        ("古瑞瓦特", "深圳", "储能控制"),
        ("安克创新", "深圳", "储能EMS算法工程师"),
        ("德兰明海(Bluetti)", "深圳", "储能BMS/嵌入式(独角兽)"),
        ("未蓝新能源(Vilion)", "深圳", "储能BMS"),
        ("深圳协能科技", "深圳", "BMS/储能系统"),
        ("深圳天邦达新能源", "深圳", "BMS/锂电池管理"),
        ("深圳易储数智能源", "深圳", "储能BMS软件"),
        ("深圳科华恒盛", "深圳", "储能EMS"),
        ("科士达", "深圳", "储能EMS"),
        ("科陆电子", "深圳", "储能EMS"),
        ("英威腾", "深圳", "储能控制"),
        ("英飞源", "深圳", "EMS/BMS/算法"),
        ("豪鹏科技", "深圳", "BMS软件/AI仿真"),
        ("首航新能源", "深圳", "储能BMS/EMS"),
        ("长园能源", "珠海/深圳", "储能EMS/BMS"),
        ("中国广核 CGN", "阳江/防城港/台山/惠州", "核电/电气"),
    ]),
    ("浙江", [
        ("昱能科技", "嘉兴", "储能EMS"),
        ("华塑科技", "杭州", "储能BMS"),
        ("协能科技(杭州)", "杭州", "BMS算法工程师"),
        ("奥能电源(杭州)", "杭州", "储能BMS算法"),
        ("杭州德创电子", "杭州", "微网EMS/BMS"),
        ("科工电子", "杭州", "储能BMS/EMS"),
        ("高泰昊能", "杭州", "动力BMS/EMS"),
        ("高特电子", "杭州", "储能BMS算法"),
        ("正泰电源", "杭州/上海", "储能EMS"),
        ("三花储能 — 杭州基地", "浙江杭州", "储能热管理/集成"),
        ("艾罗能源", "杭州/苏州/深圳", "控制算法/BMS软件"),
    ]),
    ("湖北", [
        ("中步擎天新能源", "武汉", "储能系统"),
        ("岚图汽车(东风旗下)", "武汉", "新能源BMS/VCU控制"),
        ("楚能新能源", "武汉", "BMS/EMS系统开发"),
        ("武汉大全能源技术", "武汉", "电力电子软件"),
        ("武汉珠和辰仁", "武汉", "储能系统"),
        ("武汉菱电电控", "武汉", "EMS/电控"),
        ("骆驼集团", "武汉", "BMS工程师"),
        ("巨安储能武汉科技", "武汉光谷", "BMS开发(C++/Linux)"),
        ("武汉吉兆储能科技", "武汉光谷", "BMS软件(液态金属电池)"),
        ("君安储能(武汉)", "武汉洪山", "BMS/EMS研发(院士团队)"),
        ("湖北省新能源有限公司", "武汉", "储能/虚拟电厂"),
    ]),
]

for prov, companies in province_groups:
    add_heading3(f"📍 {prov}（{len(companies)}家）")
    make_table(
        ["企业名称", "地点", "方向/岗位"],
        [[c[0], c[1], c[2]] for c in companies],
    )

# ──────────────────────────────────────────────
# Section 5: 投递策略建议
# ──────────────────────────────────────────────
add_heading2("五、针对新设公司的投递策略")

strategies = [
    ("优先级排序",
     "第一优先：央企新设子公司（国宁新储、中国电气装备储能、中国资源循环集团）\n"
     "第二优先：地方储能集团/研究院（广州储能集团、广东新型储能研究院）\n"
     "第三优先：成熟企业新基地（中创新航惠州、国轩高科福州、海博思创佛山）\n"
     "第四优先：CATL/ATL系（新能时代、ATL）"),
    ("优势总结",
     "• 新公司HC多、竞争小：新设企业第一年招人通常招不满，门槛比成熟企业低\n"
     "• 央企新设=应届窗口：错过这波，以后社招进去要难得多\n"
     "• 你的项目在新公司是加分项：新公司更需要能直接上手的人\n"
     "• 专业限制相对宽松：新公司招人更务实，更看重你能做什么而不是你专业叫什么"),
    ("行动建议",
     "① 8月：整理这些公司的招聘官网/公众号，提前关注\n"
     "② 9月：第一批投递时把这些新设公司放在优先位置\n"
     "③ 投递时注意区分'校招'和'社招'——你是应届生，走校招通道\n"
     "④ 对于没有明确校招公告的，直接打HR电话问应届生招聘计划\n"
     "⑤ 有福建/广东关系的（校友、老乡），优先走内推"),
]

for title, content in strategies:
    add_heading3(title)
    for line in content.split("\n"):
        if line.strip():
            add_para(line.strip(), size=10)

# ──────────────────────────────────────────────
# Footer
# ──────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— END —")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = "Arial"

out_path = "docs/岗位/2025-2026新设央企国企子公司名单_V2.docx"
doc.save(out_path)
print(f"Done -> {out_path}")
