#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成：各模式下多目标权重 + 硬约束变化详表"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Style ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_table(doc, headers, rows, col_widths=None, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    return table

# ══════════════════════════════════════════════════════
#   COVER
# ══════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(60)
r = t.add_run('多堆燃料电池系统\n模式—多目标—约束映射矩阵')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('各模式下优化目标权重与硬约束的对应关系 · 2026-07').font.size = Pt(12)
doc.add_page_break()

# ══════════════════════════════════════════════════════
#  SECTION 1: 整体架构回顾
# ══════════════════════════════════════════════════════
doc.add_heading('一、模式架构总览', level=1)
doc.add_paragraph('控制模式分为三层，各层职责分离：')

add_table(doc,
    ['层级', '名称', '职责', '触发方式', '与优化的关系'],
    [
        ['L1', '正常作业模式', '根据工况选择最优权重分配', '工况自动切换', '四个目标做 Pareto 优化'],
        ['L2', '安全保底', '覆盖 L1，以安全为唯一目标', '安全阈值触发', '目标函数退化为 J = J_safety'],
        ['ASync', '故障降级', '切出故障堆，剩余堆降级运行', '故障检测异步触发', '目标函数结构不变，但参数/约束重标定'],
    ],
    col_widths=[1.5, 3, 4, 3, 4]
)

doc.add_paragraph()
doc.add_heading('二、L1 正常作业模式——四模式定义', level=1)
doc.add_paragraph('四种模式覆盖 95% 以上运行工况，模式间通过工况特征自动切换：')

add_table(doc,
    ['模式', '典型场景', '占运行\n时间', '负载\n范围', '核心意图'],
    [
        ['① 巡航经济',    '长距离 A 到 B 航行', '~60%', '30–70%', '同储氢量跑最远'],
        ['② 精细作业',    'ROV 收放/机械臂/科考', '~10%', '10–30%', '功率绝对稳定，不干扰作业'],
        ['③ 高功率作业',  '高速航行/上浮/避障', '~15%', '60–100%', '输出最大可用功率，保动态响应'],
        ['④ 停泊待命',    '靠港/漂流/待机', '~15%', '< 20%', '减少启停，低功耗保温'],
    ],
    col_widths=[2.5, 4, 1.5, 2, 5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  SECTION 2: 目标权重矩阵 (核心)
# ══════════════════════════════════════════════════════
doc.add_heading('三、四目标权重在各模式下的分布（核心表）', level=1)

doc.add_paragraph(
    '每个目标 J₁～J₄ 在不同模式下赋予不同权重，权重和为 1。'
    '★ 代表该模式下该目标的优先级：★=低 ★★=中 ★★★=高 ★★★★=最高'
)

add_table(doc,
    ['模式', 'J₁ 等效氢耗', 'J₂ 功率波动', 'J₃ 启停次数', 'J₄ 效率偏差', '权重分配依据'],
    [
        ['① 巡航经济', '0.55 ★★★★', '0.10 ★', '0.10 ★', '0.25 ★★',
         '长时稳定 → 省氢立竿见影；\n波动/启停天生少；偏差长时累积不可忽视'],
        ['② 精细作业', '0.10 ★', '0.55 ★★★★', '0.05 ★', '0.30 ★★★',
         '功率抖动直接影响作业精度 → 稳定压倒一切；\n低负载下偏差容易被放大'],
        ['③ 高功率作业', '0.25 ★★', '0.40 ★★★', '0.05 ★', '0.30 ★★★',
         '大功率下供气滞后危险 → 波动控制关键；\n偏差控制防止某堆过载'],
        ['④ 停泊待命', '0.05 ★', '0.05 ★', '0.70 ★★★★', '0.20 ★★',
         '氢耗/波动几乎不重要；\n启停是寿命第一杀手 → 宁可 idle 不 shutdown；\n低负载偏差也要关注'],
    ],
    col_widths=[2, 2.2, 2.2, 2, 2.2, 4]
)

doc.add_paragraph()
doc.add_paragraph(
    '说明：上述权重为推荐起始值，实际工程中可通过层次分析法（AHP）结合专家打分标定，'
    '或通过仿真敏感性分析微调。权重切换时建议加 0.5–2s 线性过渡（防突变震荡）。'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  SECTION 3: 硬约束变化矩阵
# ══════════════════════════════════════════════════════
doc.add_heading('四、硬约束在各模式下的松紧变化（核心表）', level=1)

doc.add_paragraph(
    '硬约束不参与 Pareto 优化，但不同模式下部分约束的边界值可以适当放宽或收紧。'
    'L2 安全保底模式下所有约束取最严值。'
)

add_table(doc,
    ['约束项', '基准值\n(无模式区分)', '① 巡航经济', '② 精细作业', '③ 高功率作业', '④ 停泊待命', 'L2 安全保底'],
    [
        ['功率变化率\n|dP/dt|', '≤ 5 kW/s', '≤ 5', '≤ 2 ← 更严\n(防抖动)', '≤ 8 ← 放宽\n(响应需求)', '≤ 5', '≤ 3'],
        ['负载跟随\nt_response', '≤ 2 s', '≤ 2', '≤ 1 ← 更严', '≤ 0.5 ← 最严\n(紧急避障)', '≤ 5 ← 放宽\n(无急迫)', '不考核'],
        ['温度窗口', '60–80°C', '60–80', '60–80', '60–80', '55–75 ← 放宽下限\n(省保温能耗)', '55–80'],
        ['堆间温差\nΔT_max', '≤ 5°C', '≤ 5', '≤ 3 ← 更严\n(防止偏差干\n扰作业)', '≤ 5', '≤ 5', '≤ 3'],
        ['单池最低电压\nV_cell_min', '≥ 0.55 V', '≥ 0.55', '≥ 0.60 ← 更严\n(高可靠性)', '≥ 0.50 ← 放宽\n(短时可忍)', '≥ 0.55', '≥ 0.60'],
        ['母线电压\n波动 ΔV_bus', '≤ ±5%', '≤ ±5%', '≤ ±2% ← 最严\n(作业设备)\n', '≤ ±5%', '≤ ±5%', '≤ ±5%'],
        ['过载能力', '120%/30s\n(紧急)', '不启用', '不启用', '150%/60s ← 增强\n避障冗余)', '不启用', '已激活\n(默认)'],
        ['冷启动能耗', '≤ 5% SOC', '不考核', '不考核', '不考核', '≤ 2% SOC ← 更严\n(闲时少耗电)', '不考核'],
        ['H₂ 浓度', '≤ 0.4% LFL', '≤ 0.4%', '≤ 0.2% ← 更严\n(作业舱有人)', '≤ 0.4%', '≤ 0.4%', '≤ 0.1% ← 最严\n(预泄漏)'],
    ],
    col_widths=[2.8, 2.2, 2, 2, 2, 2, 2.2],
    font_size=8.5
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('表格解读规则：')
r.bold = True
p.add_run(' ← 更严 = 约束收紧（安全性或品质要求更高，括号内写原因）；')
p.add_run(' ← 放宽 = 约束放松（牺牲一定安全性换取性能）；')
p.add_run(' 不考核 = 该模式下此项不构成限制。')

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  SECTION 4: L2 安全保底 + 故障降级
# ══════════════════════════════════════════════════════
doc.add_heading('五、L2 安全保底模式详解', level=1)
doc.add_paragraph(
    'L2 不是"一个模式"，而是一组安全规则的集合。当任意安全阈值被触发时，'
    'L2 立即覆盖 L1 的正常优化。目标函数退化为：'
)
p = doc.add_paragraph()
r = p.add_run('  J_L2 = 0 · J₁ + 0 · J₂ + 0 · J₃ + 0 · J₄ + 1.0 · J_safety')
r.bold = True

doc.add_paragraph()

add_table(doc,
    ['触发条件', '阈值', 'L2 动作', '恢复正常条件'],
    [
        ['单电池电压越限', 'V_cell < 0.55 V', '降功率 10%/s 直到 V_cell 恢复', 'V_cell > 0.60 V 持续 10s'],
        ['堆间温差越限', 'ΔT > 5°C', '重新分配功率，降低高温堆负载', 'ΔT < 3°C 持续 30s'],
        ['H₂ 浓度越限', 'C_H₂ > 0.4% LFL', '切断氢源 + 强制通风 + 停机', 'H₂ < 0.1% 且排查完泄漏源'],
        ['绝缘电阻越限', 'R_ins < 1 MΩ', '逐堆切出定位泄漏堆 → 切出该堆', '替换故障堆后人工复位'],
        ['功率变化率超硬件极限', '|dP/dt| > 15 kW/s', '强制限幅到 5 kW/s', '负载变化率回落持续 5s'],
    ],
    col_widths=[3.5, 2.5, 4.5, 4]
)

doc.add_paragraph()
doc.add_heading('六、故障降级模式详解', level=1)
doc.add_paragraph(
    '故障降级是异步触发的，与 L1/L2 正交。当某堆故障被诊断确认后执行：'
)

doc.add_heading('6.1 故障降级下的目标权重重新标定', level=2)

add_table(doc,
    ['参数', '正常 (4 堆)', '故障降级 (N-1 或 N-2)'],
    [
        ['J₁ 氢耗权重', '按 L1 模式', '↑ 上浮 10–20%（剩余堆跑更多功率，效率更关键）'],
        ['J₂ 波动权重', '按 L1 模式', '↑ 上浮 10%（剩余堆负担重，波动危害更大）'],
        ['J₃ 启停权重', '按 L1 模式', '→ 不变（降级途中不应再启停）'],
        ['J₄ 偏差权重', '按 L1 模式', '↓ 下调 10–15%（故障堆已切出，剩余堆偏差意义降低）'],
        ['安全约束', '正常值', '全部收紧到 L2 级别（系统容错余量减少）'],
        ['功率上限', '4 × 100%', '(N−1) × 100%（自动降额）'],
    ],
    col_widths=[3, 4, 7]
)

doc.add_heading('6.2 故障恢复路径', level=2)
doc.add_paragraph('故障堆修复/更换后重新并入系统的流程：')
add_table(doc,
    ['步骤', '动作', '条件', '耗时估计'],
    [
        ['1', '故障堆隔离确认', '故障已定位，已切出', '—'],
        ['2', '修复/更换故障堆', '备件到位', '1–4 h (深海换堆)'],
        ['3', '气密性测试', '各腔保压 5 min 压降 < 1%', '10 min'],
        ['4', '绝缘测试', 'R_ins ≥ 1 MΩ', '5 min'],
        ['5', '低功率预热', '5–10% 额定功率运行至温度稳定', '15–30 min'],
        ['6', '逐步加载并入', '以 2 kW/s 加载到目标功率', '1–2 min'],
        ['7', '偏差监控', '确认 σ_eff < 初始阈值后切回 L1', '5 min'],
    ],
    col_widths=[1, 3.5, 4, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  SECTION 7: 汇总速查表
# ══════════════════════════════════════════════════════
doc.add_heading('七、完整速查表——所有模式的目标 + 约束一览', level=1)
doc.add_paragraph('一张表看清楚所有模式下四个目标的权重走向和约束松紧趋势：')

add_table(doc,
    ['模式', 'J₁\n氢耗', 'J₂\n波动', 'J₃\n启停', 'J₄\n偏差', '关键收紧约束', '关键放宽约束'],
    [
        ['① 巡航经济', '▲▲▲▲', '▶', '▶', '▲▲', '堆间温差', '—'],
        ['② 精细作业', '▶', '▲▲▲▲', '▶', '▲▲▲', 'ΔV_bus ≤ ±2%\n|dP/dt| ≤ 2 kW/s\nH₂ ≤ 0.2% LFL', '—'],
        ['③ 高功率作业', '▲▲', '▲▲▲', '▶', '▲▲▲', '响应 ≤ 0.5s\n过载 150%/60s', 'V_cell ≥ 0.50V\n|dP/dt| ≤ 8 kW/s'],
        ['④ 停泊待命', '▶', '▶', '▲▲▲▲', '▲▲', '冷启动 ≤ 2% SOC', '响应 ≤ 5s\n温度 55–75°C'],
        ['L2 安全保底', '×', '×', '×', '×', '全部最严', '—'],
        ['故障降级 (N-1)', '▲▲▲', '▲▲', '▶', '▲', '全部 → L2 级别', '功率上限降额\n偏差权重下调'],
    ],
    col_widths=[2.5, 1.5, 1.5, 1.5, 1.5, 4.5, 4],
    font_size=8.5
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('图例：').bold = True
p.add_run('▲▲▲▲ = 最高优先级  ▲▲▲ = 高优先级  ▲▲ = 中优先级  ▶ = 低优先级  × = 不参与优化')

doc.add_paragraph()
doc.add_paragraph('— 文档结束 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──
out = r'F:\CLAUDE\research\ems-platform\docs\模式_目标_约束映射矩阵.docx'
doc.save(out)
print(f'OK: {out}')
print(f'Size: {os.path.getsize(out)/1024:.1f} KB')
