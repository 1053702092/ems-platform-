# -*- coding: utf-8 -*-
"""生成敏感性分析代码逐行解释 (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')

def add_line(doc, line_num, code, explanation):
    """添加一行代码+解释"""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    c0 = table.rows[0].cells[0]; c0.text = ''
    p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(line_num)); run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x99,0x99,0x99)
    c1 = table.rows[0].cells[1]; c1.text = ''
    run = c1.paragraphs[0].add_run(code); run.font.name = 'Consolas'; run.font.size = Pt(9)
    c2 = table.rows[0].cells[2]; c2.text = ''
    run = c2.paragraphs[0].add_run(explanation); run.font.size = Pt(9)
    c0.width = Inches(0.4); c1.width = Inches(4.0); c2.width = Inches(3.6)

def add_h(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x22,0x55,0x88)
    p.paragraph_format.space_before = Pt(12)

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'; style.font.size = Pt(10); style.paragraph_format.line_spacing = 1.3

    title = doc.add_heading('week4_sensitivity_analysis.py 逐行解释', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('这个脚本遍历 Alpha、Beta、网格密度，观察参数变化对 DP 结果的影响。')

    # ═══════════════ 1. 导入与路径 ═══════════════
    add_h(doc, '1. 导入与路径设置 (L1-26)')
    doc.add_paragraph()
    add_line(doc, '1-5', '# -*- coding: utf-8 -*- ... 文件头',
             '文件编码声明 + 文件说明注释。告诉 Python 这个文件用 UTF-8 编码。')
    add_line(doc, '6-11', 'import sys, os, time / numpy / pandas / matplotlib',
             '导入依赖库：sys/os 路径操作，time 计时，numpy 数值计算，pandas 保存 CSV，matplotlib 画图。')
    add_line(doc, '10', "matplotlib.use('Agg')",
             '设置 matplotlib 后端为 Agg（不弹出图形窗口，直接保存为图片文件）。在服务器上跑脚本时必须加这行。')
    add_line(doc, '13', "PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))",
             "获取项目根目录路径。__file__ 是当前脚本的路径，两次 dirname 得到 ems-platform/。")
    add_line(doc, '14', "RESULTS_DIR = os.path.join(PROJECT_ROOT, 'results')",
             '结果目录：ems-platform/results/')
    add_line(doc, '15', "SCRIPTS_DIR = os.path.join(PROJECT_ROOT, 'scripts')",
             '脚本目录：ems-platform/scripts/')
    add_line(doc, '16', 'sys.path.insert(0, SCRIPTS_DIR)',
             '把 scripts/ 加入 Python 导入路径，这样才能 from day8_dp_ems import ...')
    add_line(doc, '19', 'import day8_dp_ems',
             '导入 DP 模块。注意是 import 模块本身（不是 from ... import 函数），因为后面要修改模块的全局变量。')
    add_line(doc, '20-23', 'from day8_dp_ems import load_drive_cycle, vehicle_power, ...',
             '同时导入一些不需要修改的函数，方便直接调用。')
    add_line(doc, '26', 'DT = 1.0',
             '时间步长 1 秒，和 DP 脚本保持一致。')

    # ═══════════════ 2. run_dp 函数 ═══════════════
    add_h(doc, '2. run_dp() — 修改参数跑 DP (L28-44)')
    doc.add_paragraph()
    add_line(doc, '28', "def run_dp(P_load, N_SOC=150, N_PFC=60, ALPHA=100.0, BETA=10000.0):",
             '定义 run_dp 函数。输入：功率需求 + 4个可调参数（默认值就是 day8_dp_ems.py 里的默认值）。')
    add_line(doc, '30-33', "day8_dp_ems.N_SOC = N_SOC / day8_dp_ems.ALPHA = ALPHA / ...",
             '★ 关键技巧：直接修改 day8_dp_ems 模块的全局变量。'
             '因为 backward_dp 函数内部会读取 day8_dp_ems.ALPHA 等全局变量，'
             '改了模块变量后，函数就会用新值运行。不修改原文件。')
    add_line(doc, '35', 't0 = time.time()',
             '记录开始时间，用于计算单次 DP 运行耗时。')
    add_line(doc, '36', 'J, pi = day8_dp_ems.backward_dp(P_load)',
             '通过模块对象调用 backward_dp。因为刚修改了 day8_dp_ems 的全局变量，函数会使用新参数。')
    add_line(doc, '37', "dp = day8_dp_ems.forward_rollout(P_load, pi)",
             '前向 Rollout，得到仿真结果。')
    add_line(doc, '38', 't_cost = time.time() - t0',
             '计算耗时（秒）。')
    add_line(doc, '39-44', "return {'H2_kg': ..., 'SOC_end': ..., 'time_s': ..., 'improve_pct': 0}",
             '返回结果字典：氢耗、SOC 终值、计算时间。improve_pct 初始为 0，外面再算。')

    # ═══════════════ 3. main 整体结构 ═══════════════
    add_h(doc, '3. main() 整体结构 (L47-264)')
    doc.add_paragraph('main() 分 6 步：加载工况 → 预计算规则结果 → Alpha 分析 → Beta 分析 → 网格分析 → 画图 → 输出 CSV → 打印总结。')

    # ═══════════════ 4. 加载工况 ═══════════════
    add_h(doc, '4. 加载工况 (L48-66)')
    doc.add_paragraph()
    add_line(doc, '49', 'cycles = {}',
             '创建空字典，存工况数据。格式：{名字: {t, v, P_load}}')
    add_line(doc, '50', "for name in ['wltc', 'nedc', 'cltc']:",
             '遍历三种工况。')
    add_line(doc, '51-54', 'try: t, v = load_drive_cycle(name) / P_load = vehicle_power(v) / cycles[name] = {...}',
             '加载工况数据 → 算功率需求 → 存入字典。')
    add_line(doc, '55-56', 'except FileNotFoundError: print(f"跳过")',
             '如果工况数据不存在，跳过不报错。')
    add_line(doc, '62-66', 'rule_results = {} / for ... rule = run_rule_controller / rule_results[name] = ...',
             '预计算规则控制器的氢耗，作为后续对比的基线。只算一次，不用每次重复。')

    # ═══════════════ 5. Alpha 敏感性 ═══════════════
    add_h(doc, '5. Alpha (SOC维持惩罚) 敏感性 (L68-86)')
    doc.add_paragraph()
    add_line(doc, '74', 'alpha_vals = [10, 50, 100, 200, 500]',
             '要测试的 5 个 α 值。从 10（松）到 500（严），覆盖一个数量级。')
    add_line(doc, '75', 'alpha_data = []',
             '存结果，后面转 DataFrame 用。')
    add_line(doc, '77', 'for name, data in cycles.items():',
             '外层循环：遍历三种工况。')
    add_line(doc, '78', 'for a in alpha_vals:',
             '内层循环：遍历 5 个 α 值。共 3×5=15 次。')
    add_line(doc, '79', "r = run_dp(data['P_load'], ALPHA=a)",
             '调用 run_dp，只改 ALPHA，其他参数用默认值。')
    add_line(doc, '80-83', "r['param']=a / r['cycle']=name / r['improve_pct'] = (规则氢耗 - DP氢耗) / 规则氢耗 × 100",
             '补充结果信息：参数值、工况名、相对于规则控制器的改善率。')
    add_line(doc, '84', 'alpha_data.append(r)',
             '存到列表，后面统一转 CSV。')
    add_line(doc, '85-86', "print(f'  {name:5s}  Alpha={a:4d}  H2={r[...]:.4f}  ...')",
             '打印进度，格式化输出，方便实时观察。')

    # ═══════════════ 6. Beta 敏感性 ═══════════════
    add_h(doc, '6. Beta (终端惩罚) 敏感性 (L88-106)')
    doc.add_paragraph()
    add_line(doc, '94', 'beta_vals = [1000, 5000, 10000, 50000, 100000]',
             '5 个 β 值，从 1000（终端SOC约束松）到 100000（强制回0.6）。')
    add_line(doc, '95', 'beta_data = []',
             '存 Beta 结果。')
    add_line(doc, '97-106', '遍历循环，和 Alpha 部分结构完全一样，只是改 BETA 参数。',
             '代码复用模式——同样的循环结构，换参数名。')

    # ═══════════════ 7. 网格密度 ═══════════════
    add_h(doc, '7. 网格密度敏感性 (L108-131)')
    doc.add_paragraph()
    add_line(doc, '114-119', "grid_configs = [(50,20,'50x20'), (100,40,'100x40'), (150,60,'150x60'), (200,80,'200x80')]",
             '四种网格配置：(N_SOC, N_PFC, 标签)。50×20 最粗最快，200×80 最密最慢。')
    add_line(doc, '122-131', '遍历循环，每次改 N_SOC 和 N_PFC 两个参数。',
             '和之前结构一样，但一次改两个参数。')

    # ═══════════════ 8. 画图 ═══════════════
    add_h(doc, '8. 画图 — 3×3 子图 (L133-225)')
    doc.add_paragraph()
    add_line(doc, '137', 'fig, axes = plt.subplots(3, 3, figsize=(15, 12))',
             '创建 3 行 3 列共 9 张子图。figsize=(15,12) 设置画布宽15英寸、高12英寸。')
    add_line(doc, '138', "colors = {'wltc': '#2196F3', 'nedc': '#FF9800', 'cltc': '#4CAF50'}",
             '三工况的颜色：WLTC蓝色、NEDC橙色、CLTC绿色。')
    add_line(doc, '139', "markers = {'wltc': 'o', 'nedc': 's', 'cltc': '^'}",
             '三工况的标记形状：圆点、方块、三角。')

    add_h(doc, '8.1 辅助画图函数 plot_data()')
    add_line(doc, '141-150', 'def plot_data(ax, x_vals_groups, y_vals_groups, xlabel, ylabel):',
             '定义本地函数封装重复的绘图逻辑。简化后面的代码。')
    add_line(doc, '142-145', "for name in cycles: ax.plot(x, y, color=colors[name], marker=markers[name], label=name.upper())",
             '在每个子图上画三条线（三工况各一条），不同颜色和标记区分。')
    add_line(doc, '146-149', 'ax.set_xlabel/set_ylabel/legend/grid/tick_params',
             '设置坐标轴标签、图例、网格、字体大小。')

    add_h(doc, '8.2 Alpha 行 — 3 张子图')
    add_line(doc, '153-156', "axes[0,0] = Alpha vs H2 / axes[0,1] = Alpha vs SOC / axes[0,2] = Alpha vs Improvement",
             '第一行三张图：α 对氢耗、对 SOC 终值、对改善率的影响。')
    add_line(doc, '154', "groups = {n: [d['H2_kg'] for d in alpha_data if d['cycle']==n] for n in cycles}",
             '用列表推导式从 alpha_data 中按工况筛选数据。这是一个技巧：一次遍历构建三个工况的数据组。')

    add_h(doc, '8.3 Beta 行 — 3 张子图')
    add_line(doc, '170-184', '同样结构，数据来自 beta_data。',
             '第二行三张图：β 对氢耗、SOC、改善率的影响。')

    add_h(doc, '8.4 Grid 行 — 3 张子图')
    add_line(doc, '187-218', "第三行略有不同：grid 的 x 轴是分类标签('50x20'等)，用 for 循环逐条画线",
             '第三行三张图：网格对氢耗、SOC、计算时间的影响。'
             '第三张图（右下）最特别——画的是计算时间，能看出网格密度和速度的 trade-off。')

    add_line(doc, '220', "plt.suptitle('DP Parameter Sensitivity Analysis', fontsize=14, fontweight='bold')",
             '设置整张图的标题。')
    add_line(doc, '221', 'plt.tight_layout()',
             '自动调整子图间距，防止标签重叠。')
    add_line(doc, '222-223', "plt.savefig(png_path, dpi=150, bbox_inches='tight')",
             '保存图片。dpi=150 控制分辨率，bbox_inches="tight" 裁剪白边。')

    # ═══════════════ 9. 保存 CSV ═══════════════
    add_h(doc, '9. 保存 CSV (L227-238)')
    doc.add_paragraph()
    add_line(doc, '230-234', "for name, dlist, fname in [('Alpha', alpha_data, ...), ('Beta', ...), ('Grid', ...)]:",
             '循环保存三组数据为 CSV 文件，方便后续用 Excel 或 Python 进一步分析。')
    add_line(doc, '235', 'df = pd.DataFrame(dlist)',
             '把字典列表转为 DataFrame（表格）。Pandas 会自动识别字典的键为列名。')
    add_line(doc, '237', "df.to_csv(csv_path, index=False)",
             '保存为 CSV，index=False 表示不写入行号。')

    # ═══════════════ 10. 总结打印 ═══════════════
    add_h(doc, '10. 总结打印 (L240-264)')
    doc.add_paragraph()
    add_line(doc, '246', 'for name in cycles:',
             '遍历工况，每个打印一段总结。')
    add_line(doc, '248', "base = [d for d in alpha_data if d['cycle']==name and d['param']==100][0]",
             '找到默认参数（α=100）的那个结果。列表推导式+[0] 取第一个匹配项。')
    add_line(doc, '253-255', '计算 Alpha 敏感性范围 = 最大氢耗 - 最小氢耗',
             '敏感性范围越大，说明该参数对结果影响越大。如果范围接近 0，说明参数不敏感。')
    add_line(doc, '263-264', "if __name__ == '__main__': main()",
             'Python 标准写法：只有直接运行本脚本时才执行 main()，被 import 时不执行。')

    # ═══════════════ 重点总结 ═══════════════
    add_h(doc, '重点总结')
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('这个脚本最核心的技巧是第 30-33 行：')
    run.bold = True
    doc.add_paragraph("day8_dp_ems.N_SOC = N_SOC  # 直接修改模块的全局变量", style='List Bullet')
    doc.add_paragraph("day8_dp_ems.backward_dp(P_load)  # 函数会自动读取修改后的变量", style='List Bullet')
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    run = p2.add_run('原理：')
    run.bold = True
    p2.add_run(' Python 函数在运行时从模块的 __dict__ 中查找全局变量。修改模块变量后，函数下次调用就会用新值。'
               '这样就不用 rewrite 原文件，也不用写复杂的配置文件。')
    doc.add_paragraph()
    doc.add_paragraph('整个脚本就是三层嵌套循环：', style='List Bullet')
    doc.add_paragraph('  外层：遍历 WLTC/NEDC/CLTC 三种工况', style='List Bullet')
    doc.add_paragraph('  中层：遍历参数值（5个Alpha / 5个Beta / 4种网格）', style='List Bullet')
    doc.add_paragraph('  内层：调用 run_dp() 算一次 DP → 存结果', style='List Bullet')

    # ── 保存 ──
    os.makedirs(DOCS_DIR, exist_ok=True)
    out_path = os.path.join(DOCS_DIR, 'Day8_敏感性分析代码逐行解释.docx')
    doc.save(out_path)
    print(f'[OK] {out_path}')

if __name__ == '__main__':
    main()
