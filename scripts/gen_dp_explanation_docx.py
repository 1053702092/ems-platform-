# -*- coding: utf-8 -*-
"""后向DP核心逻辑通俗解释 (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Inches(0.3)

def add_h(doc, text, level=1):
    return doc.add_heading(text, level=level)

def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # ── 标题 ──
    title = doc.add_heading('后向DP核心逻辑 —— 通俗解释', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        '本文档用"北京到上海开车"的比喻，解释后向动态规划（Backward DP）的核心思想。'
        '适合在理解代码前先读。'
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 1. 问题是什么
    # ══════════════════════════════════════════
    add_h(doc, '问题是什么？', level=1)

    doc.add_paragraph('你要从北京开车到上海，路上有 1800 个路口（每个路口=1秒）。')
    doc.add_paragraph('在每个路口，你都要决定：')
    doc.add_paragraph('• FC（燃料电池）发多少电？', style='List Bullet')
    doc.add_paragraph('• 电池出多少力？', style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('核心矛盾：')
    run.bold = True
    p.add_run('FC 发电越多，电池越省，但氢耗越高。反过来电池用太多，到终点时 SOC 会掉太低。每一步都不知道前面路况怎么样，怎么选才能让全程总氢耗最低、同时 SOC 不掉到底？')

    # ══════════════════════════════════════════
    # 2. 普通人的想法 vs DP 的想法
    # ══════════════════════════════════════════
    add_h(doc, '普通人的想法 vs DP 的想法', level=1)

    add_h(doc, '普通人（贪心算法）：', level=2)
    doc.add_paragraph('"到每个路口了再看。如果电池还多，就多用电池省油。如果电池少了，就多用发动机充电。"')
    doc.add_paragraph('问题是——你不知道前面是上坡还是下坡。现在多用电池，万一前面是连续上坡需要大功率呢？')

    add_h(doc, 'DP 的想法：', level=2)
    doc.add_paragraph('"我先从终点倒着想，一路倒推回起点。"')
    doc.add_paragraph('这听起来反直觉，但它的好处是：**倒着走的时候，前面的路你已经"看过"了。**')

    # ══════════════════════════════════════════
    # 3. 比喻
    # ══════════════════════════════════════════
    add_h(doc, '用一个小例子理解', level=1)

    doc.add_paragraph('假设只有 3 个路口，SOC 只有 2 种可能（0.5 或 0.7），FC 功率只有 2 个选择（5kW 或 10kW）。')

    add_h(doc, '第1步：想终点', level=2)
    doc.add_paragraph('你站在最后（路口3），发现 SOC=0.5 要罚 1600 块，SOC=0.7 要罚 100 块。最好的是 SOC=0.6，罚 0 块。')
    p = doc.add_paragraph()
    run = p.add_run('这是代码里的：')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p2 = doc.add_paragraph()
    run = p2.add_run('J[N, :] = BETA * (SOC_GRID - SOC_0) ** 2')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_h(doc, '第2步：想倒数第2个路口', level=2)
    doc.add_paragraph('站在路口2，SOC=0.5。你考虑两种选择：')

    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'
    headers = ['选择', '这步氢耗', '下一步SOC', '未来代价（路口3）']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = 'FC=5kW'
    table.rows[1].cells[1].text = '10g'
    table.rows[1].cells[2].text = '0.55'
    table.rows[1].cells[3].text = '查J[3]得800'
    table.rows[2].cells[0].text = 'FC=10kW'
    table.rows[2].cells[1].text = '18g'
    table.rows[2].cells[2].text = '0.65'
    table.rows[2].cells[3].text = '查J[3]得50'

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('总代价 = 这步氢耗 + 未来代价')
    run.bold = True
    doc.add_paragraph('FC=5kW: 10 + 800 = 810')
    doc.add_paragraph('FC=10kW: 18 + 50 = 68')
    doc.add_paragraph()
    doc.add_paragraph('选 FC=10kW！记录下来：J[2] = 68, pi[2] = 10kW')

    p = doc.add_paragraph()
    run = p.add_run('注意：')
    run.bold = True
    p.add_run('这里的关键是——你在算"这步代价"的时候，"未来代价"已经通过第1步算好了，直接查表就行。这就是"倒着走"的优势：前面（其实是你后面的时间步）的答案已经知道了。')

    add_h(doc, '第3步：往前推到第1个路口', level=2)
    doc.add_paragraph('站在路口1，SOC=0.5。同样的逻辑：')

    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = 'FC=5kW'
    table.rows[1].cells[1].text = '10g'
    table.rows[1].cells[2].text = '0.52'
    table.rows[1].cells[3].text = '查J[2]得200'
    table.rows[2].cells[0].text = 'FC=10kW'
    table.rows[2].cells[1].text = '18g'
    table.rows[2].cells[2].text = '0.62'
    table.rows[2].cells[3].text = '查J[2]得80'

    doc.add_paragraph()
    doc.add_paragraph('选最小的 → 记录')

    # ══════════════════════════════════════════
    # 4. 最终你得到什么
    # ══════════════════════════════════════════
    add_h(doc, '最终你得到了什么？', level=1)

    doc.add_paragraph('你得到了一张表：')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = '路口'
    table.rows[0].cells[1].text = 'SOC'
    table.rows[0].cells[2].text = '最优FC功率'
    table.rows[1].cells[0].text = '0'
    table.rows[1].cells[1].text = '0.5'
    table.rows[1].cells[2].text = '10kW'
    table.rows[2].cells[0].text = '0'
    table.rows[2].cells[1].text = '0.7'
    table.rows[2].cells[2].text = '5kW'
    table.rows[3].cells[0].text = '...'
    table.rows[3].cells[1].text = '...'
    table.rows[3].cells[2].text = '...'

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('这就是 pi 表（策略表）。')
    run.bold = True
    p.add_run('有了它，开车时就不用想了——看看现在 SOC 多少，查表选对应的 FC 功率就行。')

    # ══════════════════════════════════════════
    # 5. 对应到代码
    # ══════════════════════════════════════════
    add_h(doc, '对应到代码', level=1)

    add_h(doc, '两个表', level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = '变量'
    table.rows[0].cells[1].text = '对应比喻'
    table.rows[1].cells[0].text = 'J[k][i]'
    table.rows[1].cells[1].text = '"站在路口k、SOC状态i，到终点最少还要花多少代价"'
    table.rows[2].cells[0].text = 'pi[k][i]'
    table.rows[2].cells[1].text = '"站在路口k、SOC状态i，应该选哪个FC功率"'

    add_h(doc, '三步流程', level=2)

    doc.add_paragraph('1. 初始化终点惩罚', style='List Number')
    add_code(doc, 'J[N, :] = BETA * (SOC_GRID - SOC_0) ** 2')
    doc.add_paragraph('"到上海了，SOC 不在 0.6 就要被罚。"')

    doc.add_paragraph('2. 从后往前推', style='List Number')
    add_code(doc, 'for k in range(N-1, -1, -1):')
    doc.add_paragraph('"从最后一个路口倒着走到起点。"')

    doc.add_paragraph('3. 每个状态选最优', style='List Number')
    doc.add_paragraph('对每个 SOC 状态，遍历所有 FC 功率，算：', style='List Bullet')
    add_code(doc, '总代价 = 这一步氢耗 + SOC惩罚 + 未来代价')
    doc.add_paragraph('选总代价最小的 FC 功率，存下。', style='List Bullet')

    add_h(doc, '前向 Rollout（查表跑仿真）', level=2)
    doc.add_paragraph('后向 DP 算完 pi 表后，真正跑仿真时是**从前向后**的：')
    add_code(doc,
        'for k in range(N):\n'
        '    pfc = 查 pi[k, SOC[k]]   # 看表就知道怎么选\n'
        '    SOC[k+1] = 更新SOC       # 执行控制，往前走'
    )
    doc.add_paragraph('像个自动驾驶——路书（pi表）已经规划好了，你照着开就行。')

    # ══════════════════════════════════════════
    # 6. 向量化
    # ══════════════════════════════════════════
    add_h(doc, '向量化（为什么代码里没有第三层循环？）', level=1)

    doc.add_paragraph('按照上面讲的，应该有三层循环：')
    add_code(doc,
        'for 每个路口:               # 1800次\n'
        '    for 每个SOC状态:         # 150次\n'
        '        for 每个FC功率:      # 60次\n'
        '            算一下总代价'
    )
    doc.add_paragraph('总共 1800×150×60 = 1620 万次。Python 的 for 循环跑这个会非常慢。')
    doc.add_paragraph()
    doc.add_paragraph('所以代码把最内层循环（遍历60个FC功率）改成了一次算完：')
    add_code(doc,
        '# 一次算所有60个FC功率的SOC_next\n'
        'SOC_next_all = state_transition(soc, PFC_GRID, P_load_k)\n'
        '\n'
        '# 一次算所有60个的未来代价\n'
        'J_future = np.interp(SOC_next_all, SOC_GRID, J_next_k)\n'
        '\n'
        '# 选最小的\n'
        'pi[k, i] = PFC_GRID[np.argmin(total)]'
    )

    # ══════════════════════════════════════════
    # 总结
    # ══════════════════════════════════════════
    add_h(doc, '总结', level=1)

    p = doc.add_paragraph()
    run = p.add_run('一句话记住后向 DP：')
    run.bold = True
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(
        '"从终点倒着走，每个路口记下"如果 SOC 是 X，应该选 Y"。\n'
        '走回起点时，最优路径就知道了。"'
    )
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph(
        '面试被问到"后向DP的原理"时，用这个比喻讲清楚，比背公式好得多。'
    )

    # ── 保存 ──
    os.makedirs(DOCS_DIR, exist_ok=True)
    out_path = os.path.join(DOCS_DIR, 'Day8_后向DP核心逻辑通俗解释.docx')
    doc.save(out_path)
    print(f'[OK] 已保存: {out_path}')

if __name__ == '__main__':
    main()
