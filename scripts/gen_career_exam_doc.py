#!/usr/bin/env python3
"""Generate DOCX from 央企笔试备考计划 markdown."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# Style defaults
style = doc.styles["Normal"]
font = style.font
font.name = "Arial"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

with open("docs/interview/央企笔试_行测备考计划.md", "r", encoding="utf-8") as f:
    md = f.read()

lines = md.split("\n")
code_block = False
table_headers = []
table_rows = []


def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        colors = {
            1: RGBColor(0x1A, 0x56, 0xDB),
            2: RGBColor(0x2C, 0x3E, 0x50),
            3: RGBColor(0x34, 0x49, 0x5E),
            4: RGBColor(0x5D, 0x6D, 0x7E),
        }
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
        sizes = {1: 22, 2: 16, 3: 13, 4: 11}
        run.font.size = Pt(sizes.get(level, 11))


def add_table(headers, rows):
    if not rows:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Arial"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"


def add_normal(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_bold_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_rich_line(line):
    """Add a line with **bold** support."""
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = "Arial"
        run.font.size = Pt(11)


for line in lines:
    if line.startswith("```"):
        code_block = not code_block
        continue
    if code_block:
        continue

    stripped = line.strip()

    # --- Headings ---
    if line.startswith("# ") and not line.startswith("## "):
        add_heading(line[2:].strip(), 1)
    elif line.startswith("## ") and not line.startswith("### "):
        add_heading(line[2:].strip(), 2)
    elif line.startswith("### ") and not line.startswith("#### "):
        add_heading(line[3:].strip(), 3)
    elif line.startswith("#### "):
        add_heading(line[4:].strip(), 4)

    # --- Separator ---
    elif stripped in ("---", "___", "___") and len(stripped) >= 3:
        doc.add_paragraph().add_run("─" * 50)

    # --- Tables ---
    elif line.startswith("|") and "|" in line:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if line.startswith("|---"):
            continue
        if not table_headers:
            table_headers = cells
        else:
            table_rows.append(cells)

    # --- Empty line -> flush table if any ---
    elif stripped == "":
        if table_headers:
            add_table(table_headers, table_rows)
            doc.add_paragraph()
            table_headers, table_rows = [], []
        else:
            doc.add_paragraph()

    # --- Lists ---
    elif stripped.startswith("- ") or stripped.startswith("* "):
        content = stripped[2:]
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(content)
        run.font.name = "Arial"
        run.font.size = Pt(11)

    elif stripped.startswith("  - ") or stripped.startswith("  * "):
        content = stripped[4:]
        p = doc.add_paragraph(style="List Bullet 2")
        run = p.add_run(content)
        run.font.name = "Arial"
        run.font.size = Pt(10)

    # --- Rich text with **bold** ---
    elif "**" in line:
        add_rich_line(line)

    # --- Plain text ---
    elif stripped:
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.name = "Arial"
        run.font.size = Pt(11)


# Flush remaining table
if table_headers:
    add_table(table_headers, table_rows)

for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

out_path = "docs/interview/央企笔试_行测备考计划.docx"
doc.save(out_path)
print(f"Done -> {out_path}")
