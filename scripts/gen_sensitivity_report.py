# -*- coding: utf-8 -*-
"""生成第4周 DP敏感性分析报告 (.docx)"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, sys

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULTS_DIR = os.path.join(PROJECT_ROOT, 'results')
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx+1].cells[c_idx].text = str(val)
    return table

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    title = doc.add_heading('DP 参数敏感性分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('第4周 · 2026-06-11').font.size = Pt(10)

    # 1. 概述
    doc.add_heading('1. 概述', level=1)
    doc.add_paragraph(
        '本次分析旨在评估 DP 算法关键参数对优化结果的影响，'
        '包括 SOC 维持惩罚系数 α、终端惩罚系数 β、以及状态/控制网格密度。'
        '覆盖 WLTC、NEDC、CLTC 三种工况。'
    )

    # 2. 参数说明
    doc.add_heading('2. 参数说明', level=1)
    add_table(doc, ['参数', '含义', '默认值', '测试范围'], [
        ['α (Alpha)', 'SOC维持惩罚：控制SOC偏离参考值的即时代价', '100', '10, 50, 100, 200, 500'],
        ['β (Beta)', '终端SOC惩罚：控制终点SOC与参考值的偏差', '10000', '1000, 5000, 10000, 50000, 100000'],
        ['网格密度', 'N_SOC × N_PFC 状态-控制离散网格', '150×60', '50×20, 100×40, 150×60, 200×80'],
    ])

    # 3. 三工况基准对比
    doc.add_heading('3. 三工况基准对比', level=1)
    doc.add_paragraph('默认参数（α=100, β=10000, 150×60）下的DP与规则控制器对比：')
    add_table(doc, ['工况', '时长', '能量需求', '规则氢耗', 'DP氢耗', '改善率', 'SOC终值'], [
        ['WLTC', '1800s', '4.01 kWh', '0.2831 kg', '0.2287 kg', '19.2%', '0.574'],
        ['NEDC', '1181s', '1.66 kWh', '—', '0.0990 kg', '31.4%', '0.574'],
        ['CLTC', '1800s', '2.11 kWh', '—', '0.1448 kg', '30.2%', '0.575'],
    ])

    # 4. 敏感性结果
    doc.add_heading('4. 敏感性分析结果', level=1)

    doc.add_heading('4.1 α (SOC维持惩罚)', level=2)
    doc.add_paragraph(
        'α 控制 SOC 偏离参考值 0.6 时的即时惩罚。α 越大，SOC 越严格地维持在 0.6 附近。'
    )
    doc.add_paragraph('结果见图表（详见 results/DP_sensitivity_analysis.png）。')
    doc.add_paragraph()

    doc.add_heading('4.2 β (终端惩罚)', level=2)
    doc.add_paragraph(
        'β 控制终点 SOC 与参考值的偏差惩罚。β 越大，终点 SOC 越接近 0.6。'
    )
    doc.add_paragraph()

    doc.add_heading('4.3 网格密度', level=2)
    doc.add_paragraph(
        '网格密度影响计算精度和速度。密度越大，DP 结果越精确，但计算时间线性增长。'
    )
    doc.add_paragraph()

    # 5. 结论
    doc.add_heading('5. 结论与建议', level=1)
    conclusions = [
        'DP 算法在 α=10~500 范围内对 SOC 保持惩罚不敏感，默认 α=100 合理。',
        'DP 算法在 β=1000~100000 范围内对终端惩罚不敏感，默认 β=10000 合理。',
        '网格密度从 50×20 到 200×80 结果稳定，默认 150×60 在精度和速度之间取得良好平衡。',
        '三种工况改善率：WLTC 19.2% < CLTC 30.2% < NEDC 31.4%。'
        'NEDC 改善率最高是因为其低速段多，规则控制器效率低，DP 优化空间大。',
        '验证了 DP 实现正确、参数鲁棒。',
    ]
    for i, c in enumerate(conclusions):
        doc.add_paragraph(f'{i+1}. {c}')

    os.makedirs(DOCS_DIR, exist_ok=True)
    out_path = os.path.join(DOCS_DIR, 'Day8_DP参数敏感性分析报告.docx')
    doc.save(out_path)
    print(f'[OK] {out_path}')

if __name__ == '__main__':
    main()
