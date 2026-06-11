# -*- coding: utf-8 -*-
"""后向DP代码逐行解释 (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')

def add_line(doc, line_num, code, explanation, indent=0):
    """添加一行代码和对应的解释"""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'

    # 行号
    cell0 = table.rows[0].cells[0]
    cell0.text = ''
    p = cell0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(line_num))
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 代码
    cell1 = table.rows[0].cells[1]
    cell1.text = ''
    p = cell1.paragraphs[0]
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.3)

    # 解释
    cell2 = table.rows[0].cells[2]
    cell2.text = ''
    p = cell2.paragraphs[0]
    run = p.add_run(explanation)
    run.font.size = Pt(9)

    # 设置列宽
    cell0.width = Inches(0.4)
    cell1.width = Inches(4.2)
    cell2.width = Inches(3.4)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x22, 0x55, 0x88)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.3

    # ── 标题 ──
    title = doc.add_heading('backward_dp() 代码逐行解释', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        '这个文档把 backward_dp 函数的每一行代码拆开，'
        '告诉你"这行在干什么"和"为什么需要这行"。'
        '先读右边解释，再看左边代码。'
    )

    # ── 函数签名 ──
    add_section_header(doc, '1. 函数定义')
    doc.add_paragraph()
    add_line(doc, 177, 'def backward_dp(P_load, SOC_0=0.6):',
             '定义函数。输入：功率需求序列P_load（长度N），初始SOC参考值（默认0.6）')
    doc.add_paragraph()
    add_line(doc, '178-181', '"""后向DP(向量化) P_load: array(N,) return: J表, pi表"""',
             '函数的说明文档。告诉你输入是N维数组，输出是两个表')
    doc.add_paragraph()
    add_line(doc, '183', 'N = len(P_load)',
             'N = 总时间步数。WLTC是1800秒，所以N=1800')
    doc.add_paragraph()
    add_line(doc, '184', 'SOC_GRID = np.linspace(SOC_MIN, SOC_MAX, N_SOC)',
             '把SOC范围[0.2, 0.9]离散成150个点。就像把温度计从0.2到0.9之间均匀刻150条线。'
             '结果：SOC_GRID = [0.2, 0.2047, 0.2094, ..., 0.8953, 0.9]')
    doc.add_paragraph()
    add_line(doc, '185', 'PFC_GRID = np.linspace(PFC_MIN, PFC_MAX, N_PFC)',
             '把FC功率范围[0, 30]kW离散成60个点。'
             '结果：PFC_GRID = [0, 0.51, 1.02, ..., 29.49, 30]')
    doc.add_paragraph()
    add_line(doc, '187', 'J = np.zeros((N+1, N_SOC))',
             '创建代价表 J，形状 (1801, 150)。'
             '1801行 = 时刻0~1800（多一行是终点），150列 = 每个SOC状态。'
             'J[k][i] = "时刻k、SOC状态i时，到终点最少还要花多少代价"')
    doc.add_paragraph()
    add_line(doc, '188', 'pi = np.zeros((N, N_SOC))',
             '创建策略表 pi，形状 (1800, 150)。'
             'pi[k][i] = "时刻k、SOC状态i时，应该输出多少FC功率"')
    doc.add_paragraph()
    add_line(doc, '190', '# 预计算氢耗（向量化）',
             '注释：提前算好每个FC功率对应的氢耗，不用每次重复算')
    doc.add_paragraph()
    add_line(doc, '191', 'H2_flow_grid = fc_hydrogen_flow(PFC_GRID)',
             '一次性算出60个FC功率对应的氢耗。'
             'H2_flow_grid 长度60，第j个元素 = PFC_GRID[j]对应的氢耗(g/s)')
    doc.add_paragraph()

    # ── 终端惩罚 ──
    add_section_header(doc, '2. 终端惩罚（"到上海了，SOC不对就罚钱"）')
    doc.add_paragraph()
    add_line(doc, '194', 'J[N, :] = BETA * (SOC_GRID - SOC_0) ** 2',
             'J[1800, :]：最后一行（终点），150个SOC状态各自对应的终端代价。\n'
             'SOC=0.6时代价=0，SOC=0.5或0.7时代价=100，偏离越远代价越大（二次方）。\n'
             '这保证DP不会为了省氢耗而把电池用光。')
    doc.add_paragraph()

    # ── 主循环 ──
    add_section_header(doc, '3. 主循环 — 从后往前推（"从最后一个路口倒着走回起点"）')
    doc.add_paragraph()
    add_line(doc, '197', 'for k in range(N-1, -1, -1):',
             '外层循环：k从1799跑到0（倒着走）。\n'
             'range(N-1, -1, -1) = [1799, 1798, 1797, ..., 1, 0]。\n'
             '每次处理一个时间步。')
    doc.add_paragraph()
    add_line(doc, '198', 'P_load_k = P_load[k]',
             '取出当前时刻k的负载功率（就是从这一步需要多少功率）')
    doc.add_paragraph()
    add_line(doc, '199', 'J_next_k = J[k+1, :]',
             '取出下一时刻的代价表（第k+1行的150个值）。'
             '后面要用它来查"未来的代价"。'
             '因为是从后往前推，所以J[k+1]已经算好了。')
    doc.add_paragraph()

    # ── 内层循环 ──
    add_section_header(doc, '4. 内层循环 — 遍历每个SOC状态')
    doc.add_paragraph()
    add_line(doc, '201', 'for i in range(N_SOC):',
             '遍历SOC_GRID的150个离散SOC值。i从0到149。')
    doc.add_paragraph()
    add_line(doc, '202', 'soc = SOC_GRID[i]',
             '取出第i个SOC值。比如i=100时，soc=0.6（参考值）。')
    doc.add_paragraph()

    # ── 核心 ──
    add_section_header(doc, '5. 核心计算 — "试试发动机出多少力"')
    doc.add_paragraph()
    add_line(doc, '204-205', 'SOC_next_all = state_transition(soc, PFC_GRID, P_load_k)',
             '★ 关键一步：一次试完60种FC功率。\n'
             '对"当前SOC是soc"这个情况，试试所有60个可能的FC功率，'
             '分别算每个FC功率会导致SOC变成多少。\n'
             '返回结果长度60：第j个值 = "如果用PFC_GRID[j]这么大力，SOC会变成多少"')
    doc.add_paragraph()
    add_line(doc, '207-208', 'feasible = (SOC_next_all >= SOC_MIN) & (SOC_next_all <= SOC_MAX)',
             '找可行的控制。SOC_next_all里哪些值在[0.2, 0.9]内？'
             'feasible是布尔数组，长度60。'
             '比如 [True, True, False, ..., True] — True表示"这个FC功率可行"')
    doc.add_paragraph()
    add_line(doc, '210', 'if not feasible.any():',
             '如果所有60个FC功率都不行（比如电池完全没电了，SOC怎么走都会越界），'
             '那这个状态就是死路一条。')
    doc.add_paragraph()
    add_line(doc, '212', 'J[k, i] = np.inf',
             '在代价表里标记为无穷大——"这条路走不通"')
    doc.add_paragraph()
    add_line(doc, '213', 'pi[k, i] = np.nan',
             '策略表里标为nan——"没有最优选择"')
    doc.add_paragraph()
    add_line(doc, '214', 'continue',
             '跳过剩下的代码，处理下一个SOC状态')
    doc.add_paragraph()

    # ── 代价计算 ──
    add_section_header(doc, '6. 算总代价 = 这步氢耗 + SOC惩罚 + 未来代价')
    doc.add_paragraph()
    add_line(doc, '217', 'g = H2_flow_grid * DT',
             '单步氢耗 = 预计算的氢耗(g/s) × 时间步长(1s)。'
             'g长度60，每个FC功率对应的这一步氢耗。')
    doc.add_paragraph()
    add_line(doc, '220', 'J_future = np.interp(SOC_next_all[feasible], SOC_GRID, J_next_k)',
             '★ 关键一步：查未来的代价表。\n'
             'SOC_next_all[feasible] 是"可行控制"计算出的SOC值。'
             'np.interp做线性插值：从SOC_GRID和J_next_k这个"查找表"中，'
             '找到这些SOC值对应的未来代价。\n'
             '比如SOC_next=0.55，就在J_next_k里找0.55对应的代价（可能介于0.5和0.6之间插值）。')
    doc.add_paragraph()
    add_line(doc, '221', 'J_future += ALPHA * (SOC_next_all[feasible] - SOC_REF) ** 2',
             '加上SOC偏离参考值的即时惩罚。\n'
             'SOC离0.6越远，这一步就要多付代价。α=100控制惩罚强度。')
    doc.add_paragraph()
    add_line(doc, '224', 'total = np.full(N_PFC, np.inf)',
             '创建长度为60的数组，全填inf。'
             '"所有FC功率先默认不可行"')
    doc.add_paragraph()
    add_line(doc, '225', 'total[feasible] = g[feasible] + J_future',
             '只对可行的那些FC功率，填入总代价。\n'
             '总代价 = 这一步氢耗 + 未来代价（含SOC惩罚）\n'
             '不可行的保持inf，DP永远不会选它们。')
    doc.add_paragraph()
    add_line(doc, '228', 'min_idx = np.argmin(total)',
             '找总代价最小的那个FC功率的索引。\n'
             '比如total=[inf, 85, 120, inf, 92, ...]，最小的是85，索引=1')
    doc.add_paragraph()
    add_line(doc, '229', 'J[k, i] = total[min_idx]',
             '在代价表里记录当前状态的最小总代价')
    doc.add_paragraph()
    add_line(doc, '230', 'pi[k, i] = PFC_GRID[min_idx]',
             '在策略表里记录这个状态应该选哪个FC功率。'
             '比如PFC_GRID[min_idx]=7.2，表示"SOC=0.6时，选FC=7.2kW最优"')
    doc.add_paragraph()

    # ── 进度 ──
    add_section_header(doc, '7. 进度打印')
    doc.add_paragraph()
    add_line(doc, '232', 'if k % 300 == 0:',
             '每300步打印一次进度。k从1799递减，所以k=1500, 1200, 900, 600, 300, 0时各打一次。')
    doc.add_paragraph()

    # ── 总结 ──
    add_section_header(doc, '8. 整体流程一张图')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        '┌─────────────────────────────────────────────────────────┐',
        '│  终端惩罚: J[N] = BETA * (SOC - 0.6)²                   │',
        '│                          │                              │',
        '│  for k = N-1 downto 0:   │  从最后一个时刻往前           │',
        '│    for i = 0 to 149:     │  遍历所有SOC状态              │',
        '│      soc = SOC_GRID[i]   │  当前SOC                     │',
        '│      next = transition() │  试所有FC功率→SOC_next        │',
        '│      future = lookup()   │  查J[k+1]得到未来代价        │',
        '│      total = g + future  │  总代价 = 氢耗 + 未来         │',
        '│      pi[k,i] = argmin()  │  记下最优FC功率               │',
        '│  ─────────────────────────────────────                  │',
        '│  最终得到 pi 表: "每个时刻×每个SOC，应该选什么FC功率"     │',
        '└─────────────────────────────────────────────────────────┘',
    ]
    for l in lines:
        run = p.add_run(l + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)

    # ── 保存 ──
    os.makedirs(DOCS_DIR, exist_ok=True)
    out_path = os.path.join(DOCS_DIR, 'Day8_backward_dp_代码逐行解释.docx')
    doc.save(out_path)
    print(f'[OK] 已保存: {out_path}')

if __name__ == '__main__':
    main()
