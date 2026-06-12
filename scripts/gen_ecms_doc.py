# -*- coding: utf-8 -*-
"""生成 ECMS 原理与实现学习文档 V2（更新版，含详细推导）"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DOCS_DIR = r'F:\CLAUDE\research\ems-platform\docs'

def h(doc, text, level=1):
    """Add Chinese heading with YaHei font"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def p(doc, text, bold=False, size=10.5):
    """Add Chinese paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def tbl(doc, headers, rows):
    """Add formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table

def code(doc, text, size=9):
    """Add monospace code block"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    return para

def formula(doc, text, size=10):
    """Add formula-style text"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(size)
    run.font.italic = True
    return para

def main():
    os.makedirs(DOCS_DIR, exist_ok=True)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ================================================================
    # COVER
    # ================================================================
    title = doc.add_heading('ECMS 等效消耗最小化策略 — 原理、推导与实现', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('第5周深度学习文档 | 2026-06-12\n'
                       '项目：EMS 平台 — 燃料电池混合动力能量管理\n'
                       '参考来源：MathWorks 官方文档 / Stanford Onori 论文 / MDPI 综述 / 国内核心期刊')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()

    # ================================================================
    # 1. ECMS 概述
    # ================================================================
    h(doc, '1. ECMS 概述', level=1)
    p(doc, 'ECMS（Equivalent Consumption Minimization Strategy，等效消耗最小化策略）'
         '是混合动力汽车（HEV/FCHEV）能量管理中最经典、应用最广泛的实时优化策略。'
         '它通过"等效因子"（Equivalence Factor）将电池的电能消耗折算为等效燃料消耗，'
         '把多目标优化问题（最小化氢耗 + 维持SOC）转化为单目标瞬时优化问题，'
         '实现在线实时求解。')

    h(doc, '1.1 核心定位', level=2)
    p(doc, 'ECMS 是 DP（动态规划）的实时近似：')
    tbl(doc,
        ['对比维度', 'DP（动态规划）', 'ECMS（等效消耗最小化）'],
        [
            ['优化性质', '全局最优（offline）', '局部最优（online）'],
            ['是否需要已知工况', '必须知道完整行驶循环', '仅需当前时刻信息'],
            ['计算复杂度', 'O(N × S × A)，N=1800步', 'O(S × A) 每步，约 150×60=9000 次/秒'],
            ['实时性', '不可在线（需几十秒到几分钟）', '可在线（毫秒级）'],
            ['SOC 维持', '显式惩罚（α + β）', '靠等效因子间接维持'],
            ['工程实现', '简单（Python/Matlab 均可）', '简单（适合嵌入式移植）'],
            ['实际应用', 'Benchmark 对比基准', '已广泛商用（Hybrid 车辆量产）'],
        ])

    h(doc, '1.2 发展历史', level=2)
    p(doc,
        '• 1999年：Paganelli G. 等首次提出 ECMS 概念，用于并联混合动力汽车（SAE 论文）\n'
        '• 2010年：Onori S. 等引入 Pontryagin 最小原理框架，建立严格的 Hamiltonian 推导\n'
        '• 2016年：Onori 专著《Hybrid Electric Vehicles Energy Management Systems》出版\n'
        '• 2019-2023年：自适应 ECMS（A-ECMS）大发展——SOC反馈、机器学习、预测型 ECMS\n'
        '• 2023-2025年：ECMS-MPC 融合、考虑 FC 寿命退化、多维度自适应因子')

    h(doc, '1.3 分类', level=2)
    tbl(doc,
        ['类型', '英文缩写', '核心特点', '适用场景'],
        [
            ['标准 ECMS', 'ECMS', '恒定等效因子，离线标定', '已知工况的基准测试'],
            ['自适应 ECMS', 'A-ECMS', '等效因子随 SOC/工况动态调整', '实际驾驶（工况不确定）'],
            ['预测型 ECMS', 'P-ECMS', '结合短时速度预测优化因子', '有 CAN 总线/导航数据的车辆'],
            ['ECMS-MPC 融合', 'ECMS-MPC', 'MPC 预测窗口 + ECMS 瞬时优化', '复杂工况/多目标权衡'],
        ])

    # ================================================================
    # 2. 核心原理
    # ================================================================
    h(doc, '2. ECMS 核心原理', level=1)

    h(doc, '2.1 物理直觉', level=2)
    p(doc,
        '想象一个燃料电池混合动力车：电池里存着电，氢气罐里存着氢。\n'
        '问题来了：某个时刻，车辆需要 10kW 功率，这 10kW 应该让 FC 发，还是让电池放？\n\n'
        'ECMS 的回答很直观：\n'
        '• 如果 FC 发这 10kW 很便宜（效率高，氢耗低），就让 FC 多发\n'
        '• 如果 FC 发这 10kW 很贵（效率低），就让电池放电来替代\n'
        '• "替代"的比例由等效因子 s 决定：s = "用电池放电替代FC发电"的代价（g H₂/kWh）\n\n'
        '关键洞见：电池不是免费的。你用了电池的电，相当于"透支"了未来需要用 FC 充回去的电量。'
        '这个"未来充电成本"就是等效因子的物理本质。')

    h(doc, '2.2 Hamiltonian 最小化框架', level=2)
    p(doc, 'ECMS 的严格数学基础来自最优控制理论中的 Pontryagin 最小原理。')
    p(doc, '定义 Hamiltonian 函数：', bold=True)
    formula(doc, 'H(t) = H_fc(P_fc(t)) + s(t) · H_bat(P_bat(t)) + λ(t) · ṡoc(t)')

    p(doc,
        '其中各部分含义：\n'
        '• H(t) — Hamiltonian（总代价函数），在每个时刻取最小\n'
        '• H_fc(P_fc) — FC 实际氢耗率 [g/s]\n'
        '• s(t) — 等效因子 [g/kWh]，是电池消耗电能对应的等效氢耗\n'
        '• H_bat(P_bat) — 电池电功率 [kW]，放电为正，充电为负\n'
        '• λ(t) —协态变量（costate），在标准 ECMS 中省略，在 Pontryagin 框架中保留')

    p(doc, '在 FCHEV 场景下的具体展开：', bold=True)
    formula(doc, 'H_fc(P_fc) = P_fc / (η_fc(P_fc) · LHV_H2) × 1000  [g/s]')
    formula(doc, 'H_bat(P_bat) = P_bat = P_load - P_fc  [kW]')

    p(doc,
        'ECMS 在每个时刻 k 做瞬时优化：\n'
        '  选择 P_fc* = argmin [H_fc(P_fc) + s · H_bat(P_load - P_fc)]\n'
        '  约束：SOC_next ∈ [SOC_min, SOC_max]')

    h(doc, '2.3 等效因子 s 的物理意义', level=2)
    p(doc, '等效因子 s 是 ECMS 最核心、最敏感的参数。它的单位是 [g H₂ / kWh]：')
    tbl(doc,
        ['s 的物理含义', '典型数值范围', '效果'],
        [
            ['"用电池1kWh等效于烧多少克氢"', '100~250 g/kWh', '越大 → 电池越"贵"→ 多用 FC 保护电池'],
            ['由 FC 效率和电池充放电效率共同决定', '—', 's 小 → 电池"便宜"→ 多用电池追求低氢耗'],
            ['标准 ECMS 用恒定值', '~180 g/kWh', '需离线标定'],
            ['自适应 ECMS 动态调整', '120~250 g/kWh', '随 SOC/工况变化'],
        ])

    h(doc, '2.4 等效因子的三种确定方法', level=2)

    h(doc, '方法 A：理论公式法', level=3)
    p(doc,
        '基于热力学第一定律，考虑电池充放电效率和 FC 效率：')
    formula(doc, 's_theory = (η_bat_chg · η_bat_dis · LHV_H2) / (η_fc · 3600) × 1000')
    p(doc,
        '代入典型值：η_bat_chg≈0.92, η_bat_dis≈0.92, LHV_H2=120×10⁶ J/kg, η_fc≈0.50\n'
        's_theory ≈ (0.92 × 0.92 × 120×10⁶) / (0.50 × 3600) × 1000 ≈ 188 g/kWh\n'
        '→ 理论值约 180~190 g/kWh，可作为初始搜索点')

    h(doc, '方法 B：DP 离线反推法（推荐 ⭐）', level=3)
    p(doc,
        '利用 DP 得到的全局最优解，反推出最优等效因子序列 s*(t)：')
    formula(doc, 's*(t) = - H_fc(P_fc*(t)) / H_bat(P_bat*(t))   （当电池放电，H_bat > 0）')
    formula(doc, 's*(t) = - H_fc(P_fc*(t)) / H_bat(P_bat*(t))   （当电池充电，H_bat < 0）')
    p(doc,
        '步骤：\n'
        '1. 用 DP 求解 WLTC/NEDC 工况的最优轨迹 → 得到 P_fc*(t), P_bat*(t)\n'
        '2. 在每个时刻 t 计算 s*(t) = -H_fc(P_fc*) / P_bat*\n'
        '3. 对 s*(t) 做均值/分段均值 → 作为 ECMS 的等效因子\n'
        '4. 在不同工况下重复 → 取平均值作为通用标定值')

    h(doc, '方法 C：自适应 ECMS（A-ECMS）', level=3)
    p(doc,
        '恒定 s 无法适应变化的 SOC 和工况，引入自适应机制。'
        '最常用的是 SOC 反馈型自适应：')
    formula(doc, 's(t) = s₀ · [1 + Kp · (SOC_ref - SOC(t))]')
    p(doc,
        '• 当 SOC > SOC_ref 时：s 减小 → 电池放电"变便宜" → 多用电池 → SOC 回落\n'
        '• 当 SOC < SOC_ref 时：s 增大 → 电池放电"变贵" → 少用电多用 FC → SOC 回升\n\n'
        '进阶：MathWorks 官方实现使用 PI 控制器自适应（参考 Adaptive ECMS 文档）：')
    formula(doc, 'Δs(t) = Kp · e(t) + Ki · ∫e(τ)dτ')
    formula(doc, 'e(t) = SOC_ref - SOC(t)')
    p(doc,
        'MathWorks 默认参数：Kp=0（可调）, Ki=0（可调），SOCTrgt=60%, SOCmin=40%, SOCmax=80%')

    # ================================================================
    # 3. 电池充放电等效成本
    # ================================================================
    h(doc, '3. 电池充放电等效成本（核心细节）', level=1)
    p(doc,
        'ECMS 中电池的成本不是对称的。放电和充电的等效氢耗完全不同：')

    tbl(doc,
        ['电池模式', 'P_bat 符号', '等效电能', '等效氢耗', '直觉理解'],
        [
            ['放电（Discharge）', 'P_bat > 0', '正值', 's × P_bat [g/s]', '用了电池1kWh = 多烧 s 克氢'],
            ['充电（Charge）', 'P_bat < 0', '负值', 's × P_bat [g/s]（负值→减少总氢耗）', '回收能量 = 减少 FC 输出 = "赚"氢'],
            ['制动回收', 'P_bat < 0（大）', '强负值', '大幅降低 H_eq', '再生制动是最大节能点'],
        ])

    p(doc,
        '重要推论：\n'
        '• 电池放电时 H_bat = +P_bat → 增加 H_eq → ECMS 会"犹豫"要不要用电池\n'
        '• 电池充电时 H_bat = -|P_bat| → 减少 H_eq → ECMS "欢迎"回收能量\n'
        '• 再生制动（P_bat < 0）是 ECMS 天然的节能机制——制动时电池充电等效于减少 FC 氢耗\n'
        '• 充放电效率不对称（η_chg ≠ η_dis）会影响 s 的最优值')

    h(doc, '3.1 充放电效率不对称处理', level=2)
    p(doc,
        '实际电池充电效率 η_chg < 放电效率 η_dis。这意味着：\n'
        '• 电池放 1kWh 出去 → 要充回 1/η_chg > 1kWh 才能回到相同 SOC\n'
        '• 所以"放电成本"应该略高于"充电收益"\n'
        '• 修正公式：s_chg = s_dis / (η_chg × η_dis)')
    formula(doc, 's_charge = s_discharge / (η_chg · η_dis)')
    p(doc,
        '例如 η_chg=0.92, η_dis=0.92 → s_charge = s_discharge / 0.85 ≈ 1.18 × s_discharge\n'
        '→ 充电等效因子比放电大 18%，反映"充回去更费电"的事实。')

    # ================================================================
    # 4. 约束处理
    # ================================================================
    h(doc, '4. ECMS 约束处理', level=1)

    h(doc, '4.1 SOC 约束', level=2)
    p(doc,
        'ECMS 没有显式的 SOC 维持惩罚项（不像 DP 有 α 和 β），'
        '但通过两种机制间接维持 SOC：')
    p(doc, '机制1：SOC 范围硬约束', bold=True)
    p(doc, '  SOC_min ≤ SOC(k+1) ≤ SOC_max\n'
           '  → 超出范围的控制直接排除（不可行）\n'
           '  → MathWorks 默认：SOCmin=40%, SOCmax=80%')
    p(doc, '机制2：等效因子自适应（软约束）', bold=True)
    p(doc, '  SOC 低 → s 增大 → 少用电池 → SOC 回升\n'
           '  SOC 高 → s 减小 → 多用电池 → SOC 回落')

    h(doc, '4.2 FC 功率约束', level=2)
    p(doc,
        '• P_fc_min ≤ P_fc ≤ P_fc_max\n'
        '• 本项目：PFC_MIN=0, PFC_MAX=30 kW\n'
        '• FC 有最小稳定功率（约 3kW），低于此值 FC 效率骤降甚至停机\n'
        '→ 工程实践中常加启停惩罚或最小运行时间约束')

    h(doc, '4.3 约束惩罚方法', level=2)
    p(doc,
        '对于超出范围的控制，ECMS 采用惩罚函数法（Penalty Method）：\n'
        '  如果 SOC_next ∉ [SOC_min, SOC_max] 或 P_fc ∉ [P_min, P_max]\n'
        '  → 该控制的总代价 H_eq 加上一个极大的惩罚值\n'
        '  → 使其永远不被选为最优\n\n'
        'MathWorks 默认惩罚因子：PenaltyFctr = 10,000,000（远大于任何合理的 H_eq 值）')

    # ================================================================
    # 5. 三种 ECMS 实现详解
    # ================================================================
    h(doc, '5. 三种 ECMS 实现详解', level=1)

    h(doc, '5.1 标准 ECMS（Non-adaptive）', level=2)
    p(doc, '最简单的实现——恒定等效因子：')
    code(doc,
        'def ecms_standard(P_load, SOC_0=0.6, s_factor=180.0):\n'
        '    """标准 ECMS：恒定等效因子"""\n'
        '    N = len(P_load)\n'
        '    SOC = np.zeros(N + 1)\n'
        '    P_fc = np.zeros(N)\n'
        '    SOC[0] = SOC_0\n'
        '    for k in range(N):\n'
        '        PFC_GRID = np.linspace(PFC_MIN, PFC_MAX, N_PFC)\n'
        '        H_fc = fc_hydrogen_flow(PFC_GRID)          # 实际氢耗 [g/s]\n'
        '        P_bat = P_load[k] - PFC_GRID               # 电池功率 [kW]\n'
        '        H_eq = H_fc + s_factor * P_bat             # 等效总氢耗\n'
        '        # SOC 约束筛选\n'
        '        feasible = []\n'
        '        for j, pf in enumerate(PFC_GRID):\n'
        '            soc_next = state_transition(SOC[k], pf, P_load[k])\n'
        '            if SOC_MIN <= soc_next <= SOC_MAX:\n'
        '                feasible.append(j)\n'
        '        if not feasible:\n'
        '            P_fc[k] = P_load[k] * 0.5              # fallback\n'
        '        else:\n'
        '            best = min(feasible, key=lambda j: H_eq[j])\n'
        '            P_fc[k] = PFC_GRID[best]\n'
        '        SOC[k+1] = state_transition(SOC[k], P_fc[k], P_load[k])\n'
        '    return SOC, P_fc')

    h(doc, '5.2 自适应 ECMS（A-ECMS）', level=2)
    p(doc, '等效因子随 SOC 动态调整：')
    code(doc,
        'def ecms_adaptive(P_load, SOC_0=0.6, s_0=180.0, Kp=30.0, SOC_ref=0.6):\n'
        '    """自适应 ECMS：SOC 反馈调整等效因子"""\n'
        '    N = len(P_load)\n'
        '    SOC = np.zeros(N + 1)\n'
        '    P_fc = np.zeros(N)\n'
        '    s_history = np.zeros(N)\n'
        '    SOC[0] = SOC_0\n'
        '    for k in range(N):\n'
        '        # SOC 反馈调整等效因子\n'
        '        s_k = s_0 * (1 + Kp * (SOC_ref - SOC[k]))\n'
        '        s_history[k] = s_k\n'
        '        # 同标准 ECMS 的瞬时优化\n'
        '        ...  # 同上，使用 s_k 替代固定 s\n'
        '    return SOC, P_fc, s_history')

    p(doc,
        '参数说明：\n'
        '• s_0 = 180 g/kWh（基准等效因子）\n'
        '• Kp = 30（比例增益，控制自适应强度）\n'
        '• SOC_ref = 0.6（目标 SOC）\n'
        '• 当 SOC=0.6 时 s_k = 180；SOC=0.5 时 s_k ≈ 210；SOC=0.7 时 s_k ≈ 150')

    h(doc, '5.3 PI 自适应 ECMS（MathWorks 方法）', level=2)
    p(doc, 'MathWorks 官方实现采用 PI 控制器调整等效因子：')
    formula(doc, 's(k) = s(k-1) + Kp · [e(k) - e(k-1)] + Ki · e(k)')
    formula(doc, 'e(k) = SOC_ref - SOC(k)')
    p(doc,
        '特点：\n'
        '• Kp 控制响应速度（比例项），Ki 消除稳态误差（积分项）\n'
        '• s 的变化率被限制（防止突变）\n'
        '• PI 输出经过限幅（s ∈ [s_min, s_max]）\n'
        '• 适合长周期行驶（多次 trip 累积 SOC 漂移校正）')

    # ================================================================
    # 6. ECMS 参数调优指南
    # ================================================================
    h(doc, '6. ECMS 参数调优指南', level=1)

    h(doc, '6.1 标准 ECMS 调优', level=2)
    p(doc, '核心参数：等效因子 s。调优目标：最小化总氢耗 + SOC 终值接近 SOC_ref。')
    tbl(doc,
        ['步骤', '操作', '预期结果'],
        [
            ['1. 确定搜索范围', 's ∈ [120, 250] g/kWh', '覆盖理论值 ±30%'],
            ['2. 粗扫描', '步长 10，跑 WLTC', '找到氢耗最低的 3~5 个 s 值'],
            ['3. 精细搜索', '步长 2~5，在粗扫描最优附近', '精确定位最优 s'],
            ['4. 多工况验证', 'WLTC + NEDC + CLTC', '验证泛化性'],
            ['5. SOC 终值检查', '目标 SOC_end ≈ 0.6', '偏差 >0.05 需要调 Kp'],
        ])

    h(doc, '6.2 A-ECMS 调优', level=2)
    tbl(doc,
        ['参数', '作用', '调优范围', '默认值'],
        [
            ['s_0', '基准等效因子', '150~200', '180'],
            ['Kp', '自适应强度', '10~60', '30'],
            ['SOC_ref', '目标 SOC', '0.5~0.7', '0.6'],
            ['s_min', '等效因子下限', '100~150', '120'],
            ['s_max', '等效因子上限', '200~300', '250'],
        ])

    # ================================================================
    # 7. 实现代码框架
    # ================================================================
    h(doc, '7. 实现代码框架（Python）', level=1)
    p(doc,
        '基于项目现有的 day8_dp_ems.py 框架，只需替换核心决策逻辑。'
        '以下是可以复用的组件：')
    tbl(doc,
        ['函数', '功能', '是否复用'],
        [
            ['fc_hydrogen_flow(P_fc)', 'FC 功率 → 氢耗率 [g/s]', '✅ 直接复用'],
            ['fc_efficiency(P_fc)', 'FC 效率曲线查表', '✅ 直接复用'],
            ['vehicle_power(v_kmh)', '车速 → 功率需求 [kW]', '✅ 直接复用'],
            ['state_transition(SOC_k, P_fc, P_load)', 'SOC 单步状态转移', '✅ 直接复用'],
            ['load_drive_cycle(name)', '加载 WLTC/NEDC/CLTC', '✅ 直接复用'],
            ['backward_dp(P_load)', '后向 DP', '❌ 替换为 ECMS'],
            ['forward_rollout()', '前向 Rollout', '❌ 替换为 ECMS 仿真'],
            ['run_rule_controller()', '规则控制器', '✅ 作为 Baseline'],
        ])

    code(doc,
        '# ecms_ems.py — 最小可用 ECMS 实现\n'
        'import numpy as np\n'
        'from day8_dp_ems import (\n'
        '    fc_hydrogen_flow, state_transition, vehicle_power,\n'
        '    load_drive_cycle, SOC_MIN, SOC_MAX, PFC_MIN, PFC_MAX,\n'
        '    N_PFC, DT, RESULTS_DIR\n'
        ')\n\n'
        'def ecms_sim(P_load, SOC_0=0.6, s_factor=180.0, adaptive=False,\n'
        '             Kp=30.0, SOC_ref=0.6):\n'
        '    """ECMS 仿真主循环\n'
        '    P_load : 功率需求 [kW], shape(N,)\n'
        '    SOC_0  : 初始 SOC\n'
        '    s_factor: 等效因子 [g/kWh]\n'
        '    adaptive: 是否启用 SOC 反馈自适应\n'
        '    """\n'
        '    N = len(P_load)\n'
        '    PFC_GRID = np.linspace(PFC_MIN, PFC_MAX, N_PFC)\n'
        '    H2_grid = fc_hydrogen_flow(PFC_GRID)  # 预计算\n\n'
        '    SOC = np.zeros(N + 1)\n'
        '    P_fc = np.zeros(N)\n'
        '    P_bat = np.zeros(N)\n'
        '    m_H2 = np.zeros(N)\n'
        '    s_hist = np.zeros(N) if adaptive else None\n\n'
        '    SOC[0] = SOC_0\n\n'
        '    for k in range(N):\n'
        '        # 自适应等效因子\n'
        '        if adaptive:\n'
        '            s_k = s_factor * (1 + Kp * (SOC_ref - SOC[k]))\n'
        '            s_hist[k] = s_k\n'
        '        else:\n'
        '            s_k = s_factor\n\n'
        '        # 瞬时优化：找使 H_eq 最小的 P_fc\n'
        '        H_fc = H2_grid                          # 实际氢耗\n'
        '        P_bat_candidates = P_load[k] - PFC_GRID  # 候选电池功率\n'
        '        H_eq = H_fc + s_k * P_bat_candidates     # 等效总氢耗\n\n'
        '        # SOC 约束筛选\n'
        '        feasible = []\n'
        '        for j, pf in enumerate(PFC_GRID):\n'
        '            soc_next = state_transition(SOC[k], pf, P_load[k], DT)\n'
        '            if SOC_MIN + 0.01 <= soc_next <= SOC_MAX - 0.01:\n'
        '                feasible.append(j)\n\n'
        '        if feasible:\n'
        '            best_j = min(feasible, key=lambda j: H_eq[j])\n'
        '            P_fc[k] = PFC_GRID[best_j]\n'
        '        else:\n'
        '            # 无可行解：强制中点\n'
        '            P_fc[k] = np.clip(P_load[k] * 0.5, PFC_MIN, PFC_MAX)\n\n'
        '        P_bat[k] = P_load[k] - P_fc[k]\n'
        '        m_H2[k] = fc_hydrogen_flow(P_fc[k]) * DT\n'
        '        SOC[k+1] = state_transition(SOC[k], P_fc[k], P_load[k], DT)\n\n'
        '    return {\n'
        '        "SOC": SOC[:N], "P_fc_kW": P_fc, "P_bat_kW": P_bat,\n'
        '        "m_H2_g": m_H2, "m_H2_cumul_kg": np.cumsum(m_H2) / 1000,\n'
        '        "s_history": s_hist\n'
        '    }')

    # ================================================================
    # 8. 与 DP 的对比预期
    # ================================================================
    h(doc, '8. ECMS vs DP 对比预期', level=1)
    p(doc,
        '基于文献（MDPI 2023、MathWorks 案例、多篇 FCHEV 论文）和项目自身 DP 基准：')
    tbl(doc,
        ['指标', 'DP（全局最优）', 'ECMS 预期', 'ECMS+自适应预期', '差距来源'],
        [
            ['WLTC 氢耗', '0.2287 kg', '0.235~0.245 kg', '0.232~0.240 kg', '局部最优 + s 标定精度'],
            ['SOC 终值', '0.574', '0.50~0.65', '0.55~0.63', '自适应 Kp 调优'],
            ['FC 平均效率', '45.7%', '43~46%', '44~46%', '功率分配质量'],
            ['FC>50% 占比', '40.5%', '35~42%', '38~43%', '等效因子精度'],
            ['计算时间', '~30s（Python）', '<0.5s', '<0.5s', '单步 O(9000) vs O(135000)'],
            ['实现难度', '中', '低', '中', '自适应逻辑简单'],
        ])
    p(doc,
        '文献参考值：\n'
        '• MDPI WEVJ (2025): ECMS vs Rule-based 改善 15-25%，ECMS vs DP 差距 2-5%\n'
        '• Stanford Onori (2010): A-ECMS 可达 DP 的 97-99%\n'
        '• MathWorks 案例：ECMS 在 HEV 上已量产验证\n'
        '• SAE 2023 (P-ECMS): 预测型 ECMS 比标准 A-ECMS 额外降低 3% 氢耗')

    # ================================================================
    # 9. Week 5 任务清单
    # ================================================================
    h(doc, '9. Week 5 任务清单', level=1)

    h(doc, 'Day 1-2：ECMS 理论学习', level=2)
    p(doc, '□ 理解 Hamiltonian 最小化框架\n'
           '□ 掌握等效因子的物理意义和三种确定方法\n'
           '□ 理解 SOC 反馈自适应的数学原理\n'
           '□ 阅读本文档 + MathWorks ECMS 文档（extract 已获取）\n'
           '□ 手写推导 H_eq 公式和自适应律公式')

    h(doc, 'Day 3-4：Python 实现 ECMS', level=2)
    p(doc, '□ 创建 scripts/ecms_ems.py\n'
           '□ 实现 ecms_sim() — 标准 ECMS（恒定 s）\n'
           '□ 实现 ecms_adaptive() — SOC 反馈自适应\n'
           '□ 复用 day8_dp_ems.py 的 vehicle_power / state_transition / fc_hydrogen_flow\n'
           '□ 参数扫描：s ∈ [120, 250]，步长 5\n'
           '□ 输出：sensitivity_s_factor.csv + 最优 s 值')

    h(doc, 'Day 5：验证与对比', level=2)
    p(doc, '□ ECMS vs DP vs Rule 三方法对比图（五合一）\n'
           '□ 等效因子敏感性分析图\n'
           '□ A-ECMS vs 标准 ECMS 对比\n'
           '□ 输出 Week 5 成果报告')

    h(doc, '并行：C++ 基础', level=2)
    p(doc, '□ 变量 / 函数 / STL 容器（vector, map, set）\n'
           '□ LeetCode Easy 1-2 题\n'
           '□ 目标：看懂 C++ 语法，能写简单算法')

    # ================================================================
    # 10. 面试八股文
    # ================================================================
    h(doc, '10. ECMS 面试八股文', level=1)

    qa = [
        ('Q1: ECMS 和 DP 的本质区别？为什么 ECMS 能实时而 DP 不能？',
         'A: 本质区别在于"已知信息量"。DP 利用 Bellman 最优性原理做全局优化，需要知道未来所有时刻的负载信息，计算复杂度 O(N×S×A)，N 是工况时长（如 1800s），所以无法在线。\n'
         'ECMS 在每个时刻只做局部瞬时优化，通过等效因子 s 把电池的"未来成本"折叠到当前代价中，每步复杂度 O(S×A)，与工况长度无关，所以可在线。'
         '性能上 ECMS 通常能达到 DP 的 95~99%，取决于 s 的标定质量。'),

        ('Q2: 等效因子 s 怎么确定？三种方法的优缺点？',
         'A: ①理论公式法：s_theory = η_bat_chg·η_bat_dis·LHV/(η_fc·3600)，约 180 g/kWh。优点是物理意义明确，缺点是没有考虑工况特性，只是一个粗略估计。\n'
         '②DP 反推法：用 DP 得到全局最优解，反推 s*(t) = -H_fc/P_bat，取均值。优点是利用了真实最优轨迹，最接近实际最优，缺点是需要先跑一遍 DP。\n'
         '③自适应 ECMS：s(t) = s₀·[1+Kp·(SOC_ref-SOC(t))]。优点是能适应 SOC 漂移和不同工况，缺点是调参复杂（Kp 需要 tune）。\n'
         '工程实践推荐：DP 反推得到基准 s₀ + 自适应微调 Kp。'),

        ('Q3: ECMS 如何维持 SOC 平衡？和规则控制器比有什么优势？',
         'A: ECMS 通过"等效因子间接维持 SOC"：SOC 低 → s 增大 → 电池放电变贵 → ECMS 倾向于多用 FC → SOC 回升。这和规则控制器的"硬分段"（SOC<0.4 强制 FC 多发电）完全不同——ECMS 是优化驱动，规则是启发式驱动。\n'
         '优势：①优化目标统一（都是最小化 H_eq），不存在规则控制器的分段跳跃；②可以精确权衡氢耗和 SOC 维持；③SOC 不会出现规则控制器常见的"过充/过放"问题。'),

        ('Q4: 电池充放电不对称对 ECMS 有什么影响？怎么处理？',
         'A: 因为 η_chg < η_dis，电池放 1kWh 再充回相同 SOC 需要充 1/η_chg > 1kWh，'
         '这意味着"放电成本"和"充电收益"不对称。如果不处理，ECMS 会过度使用电池（因为回收看起来太"划算"）。\n'
         '处理方法：对充电侧乘以效率修正系数 s_chg = s_dis / (η_chg × η_dis)。'
         '例如 η_chg=η_dis=0.92 → s_chg ≈ 1.18 × s_dis。'),

        ('Q5: 说说 A-ECMS 和 P-ECMS 的区别？',
         'A: A-ECMS（Adaptive ECMS）的自适应依据是当前状态（SOC、功率），'
         '用的是反馈控制——"SOC 低了就调大 s"。\n'
         'P-ECMS（Predictive ECMS）在此基础上增加了短时速度预测（如基于 CAN 总线车速或导航数据），'
         '预测未来几秒到几十秒的功率需求，提前调整 s——"前面要爬坡了就提前多充点电"。\n'
         'SAE 2023 论文显示 P-ECMS 比 A-ECMS 额外降低 3% 氢耗，同时减少 FC 启停次数。'),

        ('Q6: ECMS 的局限性？怎么改进？',
         'A: ①局部最优 gap：ECMS 每步贪心，可能陷入局部最优。改进：ECMS-MPC 融合，用 MPC 的预测窗口弥补。\n'
         '②s 调优依赖经验：不同车辆参数/工况需要重新标定。改进：机器学习离线学习 s*(t) 映射。\n'
         '③极端工况鲁棒性差：急加减速时 ECMS 决策可能振荡。改进：加入规则控制器 fallback 或启停惩罚。\n'
         '④未考虑 FC 寿命：频繁启停和负载波动加速 FC 退化。改进：Degradation-aware ECMS，在 H_eq 中加入退化惩罚项（MDPI 2023 有相关研究）。'),
    ]

    for q, a in qa:
        p(doc, q, bold=True)
        p(doc, a)
        doc.add_paragraph()  # spacer

    # ================================================================
    # 11. 参考资料
    # ================================================================
    h(doc, '11. 参考资料', level=1)
    refs = [
        '[1] Onori S., Serrao L., Rizzoni G., "Hybrid Electric Vehicles Energy Management Systems", Springer, 2016.\n'
        '    — ECMS 领域最权威的专著，Stanford Onori 教授著。',

        '[2] Paganelli G. et al., "Equivalent Consumption Minimization Strategy for Parallel Hybrid Vehicles", '
        'JSAE Review, 1999.\n'
        '    — ECMS 原始论文，首次提出等效因子概念。',

        '[3] MathWorks, "Equivalent Consumption Minimization Strategy" (ECMS Block Documentation).\n'
        '    mathworks.com/help/autoblks — 含自适应/非自适应两种实现、参数说明、PI 控制器设计。',

        '[4] MDPI Sustainability, "Adaptive Equivalent Fuel Consumption Minimization Based Energy Management...", 2023.\n'
        '    — A-ECMS 最新综述，涵盖等效因子优化、机器学习辅助调参。',

        '[5] Stanford ERE, "Adaptive Equivalent Consumption Minimization Strategy for Hybrid Electric Vehicles" (Onori PDF).\n'
        '    pangea.stanford.edu — SOC 反馈自适应律的经典推导。',

        '[6] 华南理工大学, "插电式混合动力汽车等效因子的实时优化".\n'
        '    期刊：zrb.bjb.scut.edu.cn — 线性规划法 + shooting 算法。',

        '[7] 汽车工程, "燃料电池汽车行驶里程自适应ECMS策略", 2019.\n'
        '    — 基于行驶里程的自适应 ECMS，含参考 SOC 修正。',

        '[8] 机械工程学报, "等效能量因子自适应调整的瞬时优化新方法", 2023.\n'
        '    cjmenet.com.cn — 三维转移概率矩阵 + 蒙特卡洛采样 + ECMS。',

        '[9] MDPI Energies, "Degradation-Conscious ECMS for Fuel Cell Hybrid System", 2021.\n'
        '    — 考虑 FC 寿命退化的 ECMS，在 H_eq 中加入退化惩罚。',

        '[10] SAE 2023, "Development of a Predictive ECMS Based on Short-term Velocity Forecast".\n'
        '     — P-ECMS 预测型 ECMS，比 A-ECMS 额外降低 3% 氢耗。',
    ]
    for ref in refs:
        p(doc, ref)

    # ================================================================
    # 12. 附录：与 DP 代码的关键差异速查
    # ================================================================
    h(doc, '12. 附录：DP vs ECMS 代码差异速查', level=1)
    tbl(doc,
        ['环节', 'DP（day8_dp_ems.py）', 'ECMS（ecms_ems.py）'],
        [
            ['求解方式', '后向递推 J[N→0]', '每步瞬时最小化'],
            ['核心数据结构', 'J 表 (N+1)×N_SOC + π 表 N×N_SOC', '无需表格，每步在线计算'],
            ['需要未来信息', '是（J[k+1] 依赖未来代价）', '否（只依赖当前时刻）'],
            ['计算量', 'N×N_SOC×N_PFC ≈ 6.5M', 'N×N_PFC ≈ 108K'],
            ['SOC 维持', 'α惩罚 + β终端惩罚', '等效因子 s + SOC 范围硬约束'],
            ['输出', '策略表 → 前向 Rollout', '直接得到 P_fc 时间序列'],
            ['代码行数', '~512行', '预计 ~150行'],
            ['可在线运行', '否', '是'],
        ])

    # Save
    out_path = os.path.join(DOCS_DIR, 'ECMS_原理与实现_Week5学习文档.docx')
    doc.save(out_path)
    print(f'[OK] 已生成: {out_path}')
    print(f'     路径: {out_path}')
    print(f'     大小: {os.path.getsize(out_path) / 1024:.0f} KB')

if __name__ == '__main__':
    main()
