#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 Week9 Part 5-8 逐行解释文档（一次生成4份）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

# ── 辅助函数 ──
style = None
doc = None

def new_doc():
    global doc, style
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return d

def add_title_page(doc, title, subtitle):
    for _ in range(6):
        doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run(subtitle + f'\n生成日期：{datetime.date.today().isoformat()}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(11)

def add_bold(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(11)
    run.bold = True

def add_line(lineno, code, explanation):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{lineno:4d}  ')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    if code.strip().startswith('#'):
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run = p2.add_run(f'└─ {explanation}')
    run.font.name = '微软雅黑'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_section_break():
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def add_qa(q, a):
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.name = '微软雅黑'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    doc.add_paragraph(a)

def save(doc, filename):
    path = os.path.join(r'F:\CLAUDE\research\ems-platform\docs\notes', filename)
    doc.save(path)
    print(f'OK: {path}')

# ═══════════════════════════════════════════════════════════════
# PART 5: MDP GridWorld
# ═══════════════════════════════════════════════════════════════
doc = new_doc()
add_title_page(doc, 'Week 9 · Part 5\nMDP 五元组 — GridWorld',
               'scripts/week9_complete.py  Lines 272–340')

add_h1('本节在学什么')
add_text(
    'Part 5 的目标是建立一个完整的 MDP（马尔可夫决策过程）实例。\n'
    '具体任务：在 4×4 网格世界中，定义状态、动作、转移概率、奖励函数和折扣因子，'
    '为后续的 Bellman 方程、策略迭代、值迭代提供环境。\n\n'
    'GridWorld 是 RL 领域的 "Hello World" —— 足够简单可视化，又能展示 MDP 的所有要素。'
)

add_h1('MDP 五元组速查')
add_text('<S, A, P, R, γ> 分别对应：')
add_text('  S — 状态集合：16 个网格位置（4×4）')
add_text('  A — 动作集合：↑ ↓ ← →')
add_text('  P — 转移概率：80% 往目标方向，20% 随机')
add_text('  R — 奖励函数：终点 +1，陷阱 -1，其他 0')
add_text('  γ — 折扣因子：0.9')

doc.add_page_break()
add_h1('逐行代码解析')

add_line(275, 'def part5_mdp_gridworld():',
         '函数定义。返回一个包含 MDP 所有数据的字典，供 Part 6/7/8 使用。')
add_line(276, '    """', '文档字符串开始。')
add_line(277, '    GridWorld MDP:', '一句话说明。')
add_line(278, '      4×4 网格, 起点 (0,0), 终点 (3,3) +1, 陷阱 (1,1) -1',
         '环境描述。4×4 共 16 个格子。')
add_line(279, '      动作: ↑ ↓ ← →   (随机移动: 80% 目标方向, 20% 随机)',
         '动作空间说明。加随机性模拟真实环境的"不完美控制"。')
add_line(280, '      折扣因子 γ = 0.9', '折扣因子。越接近 1 越看重长期收益。')
add_line(281, '    """', '文档字符串结束。')

add_section_break()
add_line(282, '    SIZE = 4', '网格边长。4×4 = 16 个状态。')
add_line(283, '    n_states = SIZE * SIZE', '状态总数 = 16。每个格子是一个状态。')
add_line(284, '    actions = [\'↑\', \'↓\', \'←\', \'→\']', '动作集合。共 4 个离散动作。')
add_line(285, '    n_actions = len(actions)', '动作数 = 4。')
add_line(286, '    gamma = 0.9', '折扣因子 γ = 0.9。')

add_line(288, '    GOAL = (3, 3)', '终点坐标：右下角。')
add_line(289, '    TRAP = (1, 1)', '陷阱坐标：靠近中心的位置。')
add_line(290, '    GOAL_IDX = GOAL[0] * SIZE + GOAL[1]',
         '将二维坐标转为一维索引。公式：行×宽度 + 列。(3,3) → 3×4+3 = 15。')
add_line(291, '    TRAP_IDX = TRAP[0] * SIZE + TRAP[1]',
         '陷阱的索引。(1,1) → 1×4+1 = 5。')

add_section_break()
add_line(293, '    action_delta = {', '定义动作对应的坐标偏移量。')
add_line(294, '        \'↑\': (-1, 0), \'↓\': (1, 0),', '↑: 行-1（向上），↓: 行+1（向下）。')
add_line(295, '        \'←\': (0, -1), \'→\': (0, 1),', '←: 列-1（向左），→: 列+1（向右）。')
add_line(296, '    }', '')

add_line(298, '    def pos_to_idx(r, c):',
         '辅助函数：将 (行, 列) 坐标转为一维索引。')
add_line(299, '        return r * SIZE + c', '索引 = 行 × 每行格数 + 列。')

add_line(301, '    def is_valid(r, c):',
         '辅助函数：判断坐标是否在网格范围内。')
add_line(302, '        return 0 <= r < SIZE and 0 <= c < SIZE',
         '行和列都在 [0, SIZE) 范围内才合法。如果出界则卡在原位（后面处理）。')

doc.add_page_break()
add_h2('构建奖励函数 R 和转移概率 P')

add_line(304, '    R = {s: {a: 0.0 for a in range(n_actions)} for s in range(n_states)}',
         '初始化奖励函数 R[s][a]。'
         '这是一个嵌套字典：外层键是状态 s（0-15），内层键是动作 a（0-3），'
         '值是该状态下执行该动作的即时奖励。初始全部设为 0.0。')

add_line(305, '    P = {s: {a: {} for a in range(n_actions)} for s in range(n_states)}',
         '初始化转移概率 P[s][a][s_next]。'
         '三层嵌套字典：P[s][a] 返回一个字典，'
         '该字典的键是下一个状态 s_next，值是概率。'
         '例如 P[0][0] = {4: 0.8, 1: 0.0667, 2: 0.0667} '
         '表示在状态 0 执行 ↑，有 80% 到状态 4，'
         '各 6.67% 偏向其他方向。')

add_section_break()
add_line(307, '    for r, c in itertools.product(range(SIZE), range(SIZE)):',
         '双重循环遍历所有 16 个格子（(0,0) 到 (3,3)）。')
add_line(308, '        s = pos_to_idx(r, c)', '将当前格子坐标转为一维索引。')

add_line(309, '        if s == GOAL_IDX or s == TRAP_IDX:',
         '终点和陷阱是终止状态：一旦到达，游戏结束。')
add_line(310, '            for a in range(n_actions):', '遍历所有动作。')
add_line(311, '                R[s][a] = 1.0 if s == GOAL_IDX else -1.0',
         '到达终点给 +1 奖励，陷阱给 -1 惩罚。')
add_line(312, '                P[s][a][s] = 1.0',
         '终止状态的转移概率：无论执行什么动作，100% 留在原地（不离开）。')
add_line(313, '            continue', '跳过下面的非终止状态处理逻辑。')

doc.add_page_break()
add_text('——以下是非终止状态的处理（普通格子）——')

add_line(314, '        for a_idx, (action_name, (dr, dc)) in enumerate(action_delta.items()):',
         '遍历 4 个动作。enumerate 同时取出动作索引和坐标偏移。')

add_line(315, '            nr, nc = r + dr, c + dc',
         '计算目标方向的坐标：当前位置 + 动作偏移量。')
add_line(316, '            if not is_valid(nr, nc):', '如果目标坐标出界...')
add_line(317, '                nr, nc = r, c', '则卡在原地不动。')
add_line(318, '            target_s = pos_to_idx(nr, nc)', '将目标坐标转为一维索引。')

add_line(319, '            R[s][a_idx] = 0.0',
         '普通格子的即时奖励为 0（无惩罚也无奖励）。')

add_line(320, '            P[s][a_idx][target_s] = P[s][a_idx].get(target_s, 0) + 0.8',
         '核心：80% 的概率转移到目标方向。'
         '.get(target_s, 0) 先取出已有的概率值（如果没有则默认为 0），加上 0.8。')

add_line(321, '            for other_dr, other_dc in action_delta.values():',
         '遍历所有 4 个方向（包括目标方向本身）。')
add_line(322, '                if (other_dr, other_dc) == (dr, dc):',
         '如果是目标方向本身...')
add_line(323, '                    continue', '跳过——已经在上面处理过了。')
add_line(324, '                nr2, nc2 = r + other_dr, c + other_dc',
         '计算其他 3 个方向的坐标。')
add_line(325, '                if not is_valid(nr2, nc2):', '如果出界...')
add_line(326, '                    nr2, nc2 = r, c', '卡在原地。')
add_line(327, '                other_s = pos_to_idx(nr2, nc2)', '转索引。')
add_line(328, '                P[s][a_idx][other_s] = P[s][a_idx].get(other_s, 0) + 0.2 / 3',
         '剩余 20% 的概率均分给其他 3 个方向，每个方向 6.67%。'
         '这就是"80% 目标方向，20% 随机"的具体实现。')

doc.add_page_break()
add_text('——打印输出——')
add_line(330, '    print(f\'[Part 5] GridWorld MDP: {n_states} 状态 × {n_actions} 动作\')',
         '打印环境信息。')
add_line(331, '    print(f\'  起点 (0,0), 终点 {GOAL} (+1), 陷阱 {TRAP} (-1)\')', '')
add_line(332, '    print(f\'  折扣因子 γ = {gamma}\')', '')
add_line(333, '    print(f\'  随机转移: 80% 目标方向, 20% 随机方向\')', '')

add_line(335, '    return {', '返回 MDP 数据字典，供后续 Part 使用。')
add_line(336, '        \'SIZE\': SIZE, \'n_states\': n_states, \'n_actions\': n_actions,', '')
add_line(337, '        \'actions\': actions, \'gamma\': gamma,', '')
add_line(338, '        \'GOAL_IDX\': GOAL_IDX, \'TRAP_IDX\': TRAP_IDX,', '')
add_line(339, '        \'R\': R, \'P\': P, \'pos_to_idx\': pos_to_idx,', '')
add_line(340, '    }', '')

doc.add_page_break()
add_h1('核心知识点')

add_bold('1. MDP 是 RL 的基础建模框架')
add_text(
    '任何 RL 问题都可以抽象为 MDP：状态 = 环境的情况，'
    '动作 = 你做的决策，奖励 = 你做得好不好。')

add_bold('2. 随机转移模拟现实的不确定性')
add_text(
    '在真实环境中，你"想"往上走，但因为有风、打滑、执行器误差，'
    '不一定真的走到上面。80% 目标 + 20% 随机 是对这种不确定性的简化建模。')

add_bold('3. 字典比矩阵更适合稀疏结构')
add_text(
    'P 用三层嵌套字典而不是 16×4×16 的矩阵，'
    '因为每个状态下只有少数几个可能的下一状态，用字典节省内存且更易读。')

doc.add_page_break()
add_h1('面试追问')
add_qa('Q: MDP 的"马尔可夫"性质是什么意思？',
       '"未来只和现在有关，和过去无关。" '
       '在 GridWorld 中，不管你是从 (0,0) 直接走下来的，'
       '还是在其他地方转了 10 圈才来的，只要当前在同一个格子，'
       '下一步的转移概率完全一样。')

add_qa('Q: 为什么终止状态的转移是 P[s][a][s] = 1.0？',
       '因为游戏已经结束。在终点或陷阱，'
       '不管执行什么动作都不会改变状态。'
       '如果没有这个设置，agent 可能会"走出"终止状态，逻辑上不合理。')

add_qa('Q: γ=0.9 是什么意思？',
       '明年的 1 块钱相当于现在的 9 毛钱。'
       'agent 越看重短期收益（γ 越小），越倾向于尽快拿到奖励。'
       'γ=0.9 是一个合理的中间值，让 agent 愿意接受短期绕路换取更大长期收益。')

save(doc, 'Week9_Part5_MDP_GridWorld_逐行解释.docx')

# ═══════════════════════════════════════════════════════════════
# PART 6: Bellman 方程
# ═══════════════════════════════════════════════════════════════
doc = new_doc()
add_title_page(doc, 'Week 9 · Part 6\nBellman 方程',
               'scripts/week9_complete.py  Lines 343–420')

add_h1('本节在学什么')
add_text(
    'Part 6 的目标是用 Bellman 方程计算两个核心价值函数：\n\n'
    '1) V^π(s) — 在随机策略下（每个动作等概率），每个状态的价值\n'
    '2) V*(s)  — 最优价值函数，即选择最好动作时的状态价值\n\n'
    'Bellman 方程是 RL 的数学基础——它把"长期回报"分解为'
    '"立即奖励 + 下一状态的折扣价值"。'
)

add_h1('核心公式')
add_text('Bellman 期望方程（用于策略评估）：')
add_text('  V^π(s) = Σ_a π(a|s) · [R(s,a) + γ · Σ_s\' P(s\'|s,a) · V^π(s\')]')
add_text('')
add_text('Bellman 最优方程（用于求最优策略）：')
add_text('  V*(s) = max_a [R(s,a) + γ · Σ_s\' P(s\'|s,a) · V*(s\')]')
add_text('')
add_text('两者区别：前者对所有动作加权平均，后者取最大值。')

doc.add_page_break()
add_h1('逐行代码解析')

add_line(346, 'def part6_bellman(mdp):',
         '函数定义。接收 Part 5 返回的 MDP 字典。')
add_line(347, '    """用 Bellman 方程计算 V(s) 和 Q(s,a)"""', '')
add_line(348, '    n_states = mdp[\'n_states\']', '从 MDP 字典中解出状态数。')
add_line(349, '    n_actions = mdp[\'n_actions\']', '解出动作数。')
add_line(350, '    gamma = mdp[\'gamma\']', '解出折扣因子。')
add_line(351, '    R = mdp[\'R\']', '解出奖励函数字典。')
add_line(352, '    P = mdp[\'P\']', '解出转移概率字典。')

add_h2('随机策略的 V(s) — Bellman 期望方程')
add_line(355, '    policy = np.ones((n_states, n_actions)) / n_actions',
         '创建随机策略矩阵。形状 (16, 4)，每个元素 = 1/4 = 0.25。'
         '表示在每个状态下，4 个动作被选择的概率相同。')

add_line(358, '    V = np.zeros(n_states)',
         '初始化状态价值函数。初始全部为 0，后续通过迭代更新。')
add_line(359, '    theta = 1e-6', '收敛阈值。当 ΔV < 1e-6 时认为已收敛。')
add_line(360, '    max_iter = 1000', '最大迭代次数，防止死循环。')

add_line(361, '    for i in range(max_iter):', '主迭代循环。')
add_line(362, '        delta = 0', '记录本轮最大的 V 值变化量。')
add_line(363, '        for s in range(n_states):', '遍历所有 16 个状态。')
add_line(364, '            v_old = V[s]', '记住当前的 V[s] 值，用于计算变化量。')
add_line(365, '            v_new = 0', '新的 V[s] 值，初始为 0。')

add_line(366, '            for a in range(n_actions):',
         '遍历 4 个动作。这里在做 Bellman 期望公式中的 Σ_a π(a|s)。')
add_line(367, '                p_a = policy[s, a]',
         '策略的概率 π(a|s)：随机策略下永远是 0.25。')
add_line(368, '                if p_a == 0:', '如果策略概率为 0 则跳过（但这里不会触发）。')
add_line(369, '                    continue', '')
add_line(370, '                bellman_sum = R[s][a]',
         '从即时奖励 R(s,a) 开始。这就是公式里的 R(s,a) 项。')

add_line(371, '                for s_next, prob in P[s][a].items():',
         '遍历可能的下一状态及其概率。这就是公式里的 Σ P(s\'|s,a)。')
add_line(372, '                    bellman_sum += gamma * prob * V[s_next]',
         '加上 γ · P · V(s\')。这就是公式里的 γ · Σ P·V(s\') 项。')

add_line(373, '                v_new += p_a * bellman_sum',
         '加权求和：π(a|s) × [R + γ·Σ P·V(s\')]。'
         '这就是完整的 Bellman 期望方程的一步。')

add_line(374, '            V[s] = v_new',
         '更新 V[s] 为计算出的新值。')
add_line(375, '            delta = max(delta, abs(v_old - v_new))',
         '更新最大变化量。用于判断是否收敛。')

add_line(376, '        if delta < theta:',
         '如果所有状态中最大的 V 值变化都小于阈值...')
add_line(377, '            break', '则认为已收敛，退出迭代。')

doc.add_page_break()
add_h2('最优 V*(s) — Bellman 最优方程')
add_line(379, '    V_opt = np.zeros(n_states)',
         '最优价值函数，初始全部为 0。')
add_line(381, '    for i in range(max_iter):', '同样迭代直到收敛。')
add_line(382, '        delta = 0', '记录最大变化。')
add_line(383, '        for s in range(n_states):', '遍历状态。')
add_line(384, '            v_old = V_opt[s]', '记住旧值。')

add_line(385, '            q_values = []', '存所有动作的 Q 值。')
add_line(386, '            for a in range(n_actions):',
         '遍历动作，计算 Q(s,a)。')
add_line(387, '                q = R[s][a]', '从即时奖励开始。')
add_line(388, '                for s_next, prob in P[s][a].items():',
         '遍历下一状态。')
add_line(389, '                    q += gamma * prob * V_opt[s_next]',
         '加上 γ·Σ P·V*(s\')。注意这里用的是 V_opt（正在更新的）。')

add_line(390, '                q_values.append(q)',
         '把 Q(s,a) 加入列表。')
add_line(391, '            V_opt[s] = max(q_values)',
         '关键区别：取最大值而非加权平均！'
         'V*(s) = max_a Q(s,a) — 假设你总是选最好的动作。')

add_line(392, '            delta = max(delta, abs(v_old - V_opt[s]))',
         '更新变化量。')
add_line(393, '        if delta < theta:', '收敛判断。')
add_line(394, '            break', '')

doc.add_page_break()
add_h2('打印输出 + 可视化')
add_line(396, '    print(f\'\\n[Part 6] Bellman 方程\')', '')
add_line(397, '    print(f\'  随机策略 V(s) 收敛于 {i+1} 次迭代\')', '')
add_line(398, '    print(f\'  起点 V(0)     = {V[0]:.4f}\')',
         '起点 (0,0) 在随机策略下的价值。')
add_line(399, '    print(f\'  最优 V*(0)    = {V_opt[0]:.4f}\')',
         '起点在最优策略下的价值。V_opt[0] > V[0] 是必然的——有策略总比随机好。')

add_line(402, '    fig, axes = plt.subplots(1, 2, figsize=(10, 4))',
         '1 行 2 列子图，左边随机策略，右边最优策略。')
add_line(403, '    SIZE = mdp[\'SIZE\']', '网格大小 4。')
add_line(404, '    im0 = axes[0].imshow(V.reshape(SIZE, SIZE), cmap=\'RdYlBu_r\', vmin=-1, vmax=1)',
         '将 V 从 16 维向量重塑为 4×4 网格，用红蓝颜色映射显示。'
         '红色 = 高价值，蓝色 = 低价值。vmin/vmax 固定为 [-1,1] 以便两张图颜色一致可比。')
add_line(405, '    axes[0].set_title(\'V^π(s) — 随机策略\')', '')
add_line(406, '    for r in range(SIZE):', '')
add_line(407, '        for c in range(SIZE):', '')
add_line(408, '            axes[0].text(c, r, f\'{V[r*SIZE+c]:.2f}\', ha=\'center\', va=\'center\', fontsize=8)',
         '在每个格子上标注 V 值，保留 2 位小数。')

add_line(411, '    for r in range(SIZE):', '')
add_line(412, '        for c in range(SIZE):', '')
add_line(413, '            axes[1].text(c, r, f\'{V_opt[r*SIZE+c]:.2f}\', ha=\'center\', va=\'center\', fontsize=8)',
         '右侧子图同样标注 V_opt 值。')

add_line(414, '    plt.tight_layout()', '')
add_line(415, '    path = os.path.join(RESULTS_DIR, \'week9_complete_bellman_value.png\')', '')
add_line(416, '    plt.savefig(path, dpi=150)', '保存结果图。')
add_line(417, '    plt.close()', '')
add_line(418, '    print(f\'  图: {path}\')', '')

add_line(420, '    return {\'V_pi\': V, \'V_opt\': V_opt}', '返回随机策略和最优策略的价值函数。')

doc.add_page_break()
add_h1('核心知识点')
add_bold('1. 策略评估 vs 策略优化的关系')
add_text(
    'Part 6 的前半部分（V^π）是策略评估：给定一个策略，算它的价值。'
    '后半部分（V*）是策略优化：在所有策略中找出最优的。'
    '策略迭代 = 评估 → 改进 → 评估 → 改进 ...')

add_bold('2. 加权平均 vs 取最大')
add_text(
    'V^π 对所有动作加权平均（因为随机策略下每个动作都可能被执行）。'
    'V* 取 Q 值的最大值（因为你假设自己每次都会选最好的动作）。'
    '这个区别是 Bellman 期望方程和最优方程的唯一不同。')

add_bold('3. 自举（Bootstrapping）')
add_text(
    '用 V(s_next) 来更新 V(s) 称为"bootstrap" — 用猜测来更新猜测。'
    '这是 DP/RL 的核心思想，也是它高效的原因——不需要完整轨迹就能学习。')

doc.add_page_break()
add_h1('面试追问')
add_qa('Q: 为什么 V* 一定比 V^π 大？',
       'V* 在每一步都选最好的动作，而 V^π 随机选。'
       '最好的动作的 Q 值 ≥ 平均的 Q 值，所以 V* ≥ V^π。')

add_qa('Q: 迭代到收敛为什么需要那么多步？',
       '因为信息在网格中"扩散"需要时间。终点 +1 的价值需要一步步向起点传播。'
       '虽然终点值固定，但远离终点的格子需要通过多次迭代"感受到"终点的存在。')

add_qa('Q: 这和强化学习中的 Q-learning 有什么关系？',
       'Q-learning 就是 Part 6 的 Bellman 最优方程 + 采样。'
       '这里我们已知 P 和 R（有模型），直接迭代。'
       'Q-learning 不知道 P 和 R，通过实际交互的样本来更新。')

save(doc, 'Week9_Part6_Bellman_方程_逐行解释.docx')

# ═══════════════════════════════════════════════════════════════
# PART 7: 策略迭代
# ═══════════════════════════════════════════════════════════════
doc = new_doc()
add_title_page(doc, 'Week 9 · Part 7\n策略迭代',
               'scripts/week9_complete.py  Lines 423–512')

add_h1('本节在学什么')
add_text(
    'Part 7 实现策略迭代（Policy Iteration）算法。\n\n'
    '策略迭代 = 重复两个步骤直到收敛：\n'
    '  ① 策略评估（Policy Evaluation）：计算当前策略的 V(s)\n'
    '  ② 策略改进（Policy Improvement）：根据 V(s) 更新策略\n\n'
    '和 Part 6 的区别：Part 6 是"给定策略算价值"，Part 7 是"边改进策略边算价值"。'
)

add_h1('算法流程图')
add_text('  随机初始化策略 π₀')
add_text('          ↓')
add_text('  策略评估：算 V^π(s)')
add_text('          ↓')
add_text('  策略改进：π\' = argmax_a Q^π(s,a)')
add_text('          ↓')
add_text('  判断 π\' == π ？')
add_text('    ├─ 是 → 收敛，输出 π*')
add_text('    └─ 否 → π = π\'，回到策略评估')

doc.add_page_break()
add_h1('逐行代码解析')

add_line(426, 'def part7_policy_iteration(mdp):',
         '策略迭代主函数。接收 MDP 字典，返回最优策略和对应的 V。')
add_line(427, '    """策略迭代: 策略评估 → 策略改进 → 直到收敛"""', '')
add_line(428, '    n_states = mdp[\'n_states\']', '')
add_line(429, '    n_actions = mdp[\'n_actions\']', '')
add_line(430, '    gamma = mdp[\'gamma\']', '')
add_line(431, '    R = mdp[\'R\']', '')
add_line(432, '    P = mdp[\'P\']', '')

add_line(434, '    policy = np.random.randint(0, n_actions, size=n_states)',
         '随机初始化策略。对每个状态随机选一个动作（0-3）。'
         '和 P5 的随机策略不同——那里是概率分布，这里是确定性策略。')
add_line(435, '    V = np.zeros(n_states)',
         '价值函数初始化为 0。')

doc.add_page_break()
add_h2('策略评估函数')
add_line(437, '    def policy_evaluation(policy, V, theta=1e-6):',
         '策略评估：计算当前策略 π 下的 V(s)。本质 = Part 6 的 Bellman 期望方程。')
add_line(438, '        for _ in range(1000):', '')
add_line(439, '            delta = 0', '')
add_line(440, '            for s in range(n_states):',
         '遍历所有状态。')
add_line(441, '                v_old = V[s]', '')
add_line(442, '                a = policy[s]',
         '关键区别：这里不是对所有动作加权平均，'
         '而是只取策略指定的那个动作（因为策略是确定性的）。')
add_line(443, '                v_new = R[s][a]', '立即奖励。')
add_line(444, '                for s_next, prob in P[s][a].items():',
         '遍历可能的下一状态。')
add_line(445, '                    v_new += gamma * prob * V[s_next]',
         '加上 γ·Σ P·V(s\')。')
add_line(446, '                V[s] = v_new', '')
add_line(447, '                delta = max(delta, abs(v_old - v_new))', '')
add_line(448, '            if delta < theta:', '')
add_line(449, '                break', '')
add_line(450, '        return V', '返回收敛后的 V(s)。')

doc.add_page_break()
add_h2('策略改进函数')
add_line(452, '    def policy_improvement(policy, V):',
         '策略改进：根据当前的 V(s)，在每个状态下选 Q 值最大的动作。')
add_line(453, '        policy_stable = True',
         '标记策略是否稳定（即是否还有任何改进）。')
add_line(454, '        for s in range(n_states):', '遍历所有状态。')
add_line(455, '            old_action = policy[s]', '记住旧动作。')

add_line(456, '            q_values = []',
         '存所有动作的 Q(s,a)。')
add_line(457, '            for a in range(n_actions):',
         '遍历动作计算 Q 值。')
add_line(458, '                q = R[s][a]', '即时奖励。')
add_line(459, '                for s_next, prob in P[s][a].items():', '')
add_line(460, '                    q += gamma * prob * V[s_next]',
         'Q(s,a) = R + γ·Σ P·V(s\')。')
add_line(461, '                q_values.append(q)', '')
add_line(462, '            policy[s] = int(np.argmax(q_values))',
         '选 Q 值最大的动作。这就是策略改进的核心：greedy improvement。')
add_line(463, '            if old_action != policy[s]:',
         '如果动作变了...')
add_line(464, '                policy_stable = False',
         '说明策略还在变化，没收敛。')

add_line(465, '        return policy, policy_stable',
         '返回更新后的策略和稳定标志。')

doc.add_page_break()
add_h2('主循环')
add_line(467, '    print(f\'\\n[Part 7] 策略迭代\')', '')
add_line(468, '    for iteration in range(50):',
         '主循环：最多 50 轮策略评估 + 改进。')
add_line(469, '        V = policy_evaluation(policy, V)',
         '① 策略评估：从当前策略算 V(s)。'
         '注意传入 V 作为初始值，利用之前迭代的结果加速收敛。')
add_line(470, '        policy, stable = policy_improvement(policy, V)',
         '② 策略改进：根据 V(s) 更新策略。')
add_line(471, '        if stable:', '如果策略没变...')
add_line(472, '            print(f\'  第 {iteration+1} 轮: 收敛 [OK]\')', '')
add_line(473, '            break', '收敛，退出。')
add_line(474, '        else:', '')
add_line(475, '            print(f\'  第 {iteration+1} 轮: 策略改进中...\')', '')

doc.add_page_break()
add_h2('输出最优策略')
add_line(477, '    action_symbols = {0: \'↑\', 1: \'↓\', 2: \'←\', 3: \'→\'}', '')
add_line(478, '    SIZE = mdp[\'SIZE\']', '')
add_line(479, '    print(f\'\\n  最优策略:\')', '')
add_line(480, '    for r in range(SIZE):', '')
add_line(481, '        row_str = \'  \'', '')
add_line(482, '        for c in range(SIZE):', '')
add_line(483, '            s = r * SIZE + c', '')
add_line(484, '            if s == mdp[\'GOAL_IDX\']:', '')
add_line(485, '                row_str += \' G  \'', '')
add_line(486, '            elif s == mdp[\'TRAP_IDX\']:', '')
add_line(487, '                row_str += \' X  \'', '')
add_line(488, '            else:', '')
add_line(489, '                row_str += f\' {action_symbols[policy[s]]}  \'',
         '打印箭头表示最优动作。')
add_line(490, '        print(row_str)', '')

doc.add_page_break()
add_h2('可视化')
add_line(493, '    fig, ax = plt.subplots(figsize=(5, 5))', '')
add_line(494, '    grid = V.reshape(SIZE, SIZE)', 'V 重塑为网格。')
add_line(495, '    ax.imshow(grid, cmap=\'RdYlBu_r\', vmin=-1, vmax=1)',
         '颜色背景表示 V 值。')
add_line(496, '    for r in range(SIZE):', '')
add_line(497, '        for c in range(SIZE):', '')
add_line(498, '            s = r * SIZE + c', '')
add_line(499, '            if s == mdp[\'GOAL_IDX\']:', '')
add_line(500, '                ax.text(c, r, \'GOAL\', ha=\'center\', va=\'center\', fontsize=12, fontweight=\'bold\')', '')
add_line(501, '            elif s == mdp[\'TRAP_IDX\']:', '')
add_line(502, '                ax.text(c, r, \'TRAP\', ha=\'center\', va=\'center\', fontsize=12, fontweight=\'bold\')', '')
add_line(503, '            else:', '')
add_line(504, '                ax.text(c, r, action_symbols[policy[s]], ha=\'center\', va=\'center\', fontsize=16)',
         '在格子上画箭头。')
add_line(505, '    ax.set_title(\'Optimal Policy (Policy Iteration)\')', '')
add_line(506, '    plt.tight_layout()', '')
add_line(507, '    path = os.path.join(RESULTS_DIR, \'week9_complete_optimal_policy.png\')', '')
add_line(508, '    plt.savefig(path, dpi=150)', '')
add_line(509, '    plt.close()', '')
add_line(511, '    return {\'policy\': policy, \'V\': V}', '返回最优策略和价值函数。')

doc.add_page_break()
add_h1('核心知识点')
add_bold('1. 策略迭代保证收敛到最优策略')
add_text(
    '数学上已经证明：策略迭代在有限 MDP 中一定收敛到全局最优策略，'
    '而且通常比值迭代更快（需要的迭代次数更少）。')

add_bold('2. 策略评估是"内循环"')
add_text(
    '每轮策略改进前，都要先把当前策略的 V(s) 算准。'
    '这个"内循环"可能要迭代很多次，但可以利用上一轮的 V 作为初始值加速。')

add_bold('3. 确定性策略就够了')
add_text(
    '在确定性的 MDP 中，最优策略可以证明是确定性的（每个状态只有一个最优动作）。'
    '随机性只在不确定的环境中才有优势。')

doc.add_page_break()
add_h1('面试追问')
add_qa('Q: 策略迭代和值迭代的区别？',
       '策略迭代：每次改进策略前先把 V(s) 算准（内循环收敛），再改进策略。\n'
       '值迭代：不需要显式存储策略，直接迭代 V*，迭代结束后一次性提取策略。\n'
       '策略迭代收敛更快（轮数少），但每轮更慢（要算到收敛）；'
       '值迭代每轮快但需要更多轮。')

add_qa('Q: 为什么策略改进取 argmax 就能保证收敛？',
       '策略改进定理：如果 Q^π(s, π\'(s)) ≥ V^π(s) 对所有 s 成立，'
       '那么 π\' 不比 π 差。'
       '取 argmax 自然满足这个条件，所以每轮策略迭代都在改善。')

save(doc, 'Week9_Part7_策略迭代_逐行解释.docx')

# ═══════════════════════════════════════════════════════════════
# PART 8: 值迭代
# ═══════════════════════════════════════════════════════════════
doc = new_doc()
add_title_page(doc, 'Week 9 · Part 8\n值迭代 + 收敛可视化',
               'scripts/week9_complete.py  Lines 515–627')

add_h1('本节在学什么')
add_text(
    'Part 8 实现值迭代（Value Iteration）算法，是 DP 求解 MDP 的第二种方法。\n\n'
    '值迭代 = 直接迭代 Bellman 最优方程，不显式存储策略。\n'
    '策略是算完 V* 后一次性提取出来的。\n\n'
    'Extra 部分可视化了值迭代的收敛过程。'
)

add_h1('和策略迭代的对比')
add_text('策略迭代：评估（内循环收敛）→ 改进 → 评估 → 改进 → ...')
add_text('值迭代：一步 Bellman 更新 → 一步 Bellman 更新 → ... → 最后提取策略')
add_text('')
add_text('值迭代本质上是把策略迭代的"内循环"压缩成一步。'
         '不需要完全收敛到 V^π，只要移动一步，就开始下一轮。')

doc.add_page_break()
add_h1('逐行代码解析')

add_line(518, 'def part8_value_iteration(mdp):',
         '值迭代主函数。')
add_line(519, '    """值迭代: 直接迭代 Bellman 最优方程"""', '')
add_line(520, '    n_states = mdp[\'n_states\']', '')
add_line(521, '    n_actions = mdp[\'n_actions\']', '')
add_line(522, '    gamma = mdp[\'gamma\']', '')
add_line(523, '    R = mdp[\'R\']', '')
add_line(524, '    P = mdp[\'P\']', '')

add_line(526, '    V = np.zeros(n_states)', '初始全 0。')
add_line(527, '    theta = 1e-6', '收敛阈值。')

add_line(529, '    print(f\'\\n[Part 8] 值迭代\')', '')
add_line(530, '    for iteration in range(1000):', '主循环。')

add_line(531, '        delta = 0', '记录最大变化。')
add_line(532, '        for s in range(n_states):', '遍历状态。')
add_line(533, '            v_old = V[s]', '')

add_line(534, '            q_max = -np.inf',
         '初始化 q_max 为负无穷。'
         '这样无论 Q 值是多少，第一次比较时一定会更新。')
add_line(535, '            for a in range(n_actions):', '遍历动作，算 Q(s,a)。')
add_line(536, '                q = R[s][a]', '')
add_line(537, '                for s_next, prob in P[s][a].items():', '')
add_line(538, '                    q += gamma * prob * V[s_next]', '')

add_line(539, '                if q > q_max:',
         '取最大值。这是值迭代和策略评估的关键区别！')
add_line(540, '                    q_max = q', '')

add_line(541, '            V[s] = q_max',
         'V(s) = max_a [R + γ·Σ P·V(s\')]。'
         '这就是 Bellman 最优算子的一步应用。')
add_line(542, '            delta = max(delta, abs(v_old - V[s]))', '')

add_line(543, '        if delta < theta:', '如果变化很小...')
add_line(544, '            print(f\'  收敛于第 {iteration+1} 次迭代\')', '')
add_line(545, '            break', '')

doc.add_page_break()
add_h2('提取最优策略')
add_line(547, '    policy = np.zeros(n_states, dtype=int)',
         '初始化策略数组。')
add_line(548, '    for s in range(n_states):', '')
add_line(549, '        q_values = []', '')
add_line(550, '        for a in range(n_actions):',
         '和策略评估一样，计算 Q(s,a)。')
add_line(551, '            q = R[s][a]', '')
add_line(552, '            for s_next, prob in P[s][a].items():', '')
add_line(553, '                q += gamma * prob * V[s_next]', '')
add_line(554, '            q_values.append(q)', '')
add_line(555, '        policy[s] = int(np.argmax(q_values))',
         '选 Q 值最大的动作。'
         '值迭代在迭代过程中不存策略，迭代结束后一次提取。')

add_line(557, '    action_symbols = {0: \'↑\', 1: \'↓\', 2: \'←\', 3: \'→\'}', '')
add_line(559, '    print(f\'  最优策略:\')', '')
add_line(560, '    for r in range(SIZE):', '')
add_line(561, '        row_str = \'  \'', '')
add_line(562, '        for c in range(SIZE):', '')
add_line(563, '            s = r * SIZE + c', '')
add_line(564, '            if s == mdp[\'GOAL_IDX\']:', '')
add_line(565, '                row_str += \' G  \'', '')
add_line(566, '            elif s == mdp[\'TRAP_IDX\']:', '')
add_line(567, '                row_str += \' X  \'', '')
add_line(568, '            else:', '')
add_line(569, '                row_str += f\' {action_symbols[policy[s]]}  \'', '')
add_line(570, '        print(row_str)', '')

add_line(572, '    print(f\'\\n  值迭代 V*(0) = {V[0]:.4f}\')',
         '打印收敛后的起点价值。')
add_line(574, '    return {\'V\': V, \'policy\': policy}', '')

doc.add_page_break()
add_h1('Extra: 收敛过程可视化')
add_text('Part 8 之后还有一个 extra 函数，专门展示值迭代的收敛动态。')

add_line(580, 'def extra_convergence_plot():',
         '收敛可视化函数。')
add_line(581, '    """演示值迭代收敛速度"""', '')
add_line(582, '    mdp = part5_mdp_gridworld()',
         '重新创建 MDP 环境。这里的 part5_mdp_gridworld() 之前已经定义过。')
add_line(584, '    n_states, n_actions = mdp[\'n_states\'], mdp[\'n_actions\']', '')
add_line(585, '    gamma = mdp[\'gamma\']', '')
add_line(586, '    R = mdp[\'R\']', '')
add_line(587, '    P = mdp[\'P\']', '')

add_line(589, '    V_track = []',
         '记录每一轮的 V(s)，用于绘制收敛轨迹。')
add_line(590, '    V = np.zeros(n_states)', '')
add_line(591, '    theta = 1e-6', '')

add_line(593, '    for iteration in range(100):', '最多 100 轮（实际 < 30 轮就收敛）。')
add_line(606, '        V_track.append(V.copy())',
         '记录当前 V 的快照。.copy() 很重要，否则存的是引用，后面改了就变了。')

add_line(610, '    V_track = np.array(V_track)',
         '形状：(n_iterations, 16)。每一行是在某一轮迭代时 16 个状态的 V 值。')

doc.add_page_break()
add_h2('双图可视化')
add_line(612, '    fig, axes = plt.subplots(2, 1, figsize=(10, 8))', '')

add_line(613, '    axes[0].plot(V_track[:, 0], \'b-\', linewidth=1.5, label=\'Start (0,0)\')',
         '绘制起点 V(0) 随迭代次数的变化。')
add_line(614, '    axes[0].plot(V_track[:, mdp[\'GOAL_IDX\']], \'g-\', linewidth=1.5, label=f\'Goal {mdp["GOAL_IDX"]}\')',
         '终点 V(15) 的变化。')
add_line(615, '    axes[0].plot(V_track[:, mdp[\'TRAP_IDX\']], \'r-\', linewidth=1.5, label=f\'Trap {mdp["TRAP_IDX"]}\')',
         '陷阱 V(5) 的变化。')
add_line(616, '    axes[0].set_xlabel(\'Iteration\')', '')
add_line(617, '    axes[0].set_ylabel(\'V(s)\')', '')
add_line(618, '    axes[0].set_title(\'Value Iteration Convergence\')', '')
add_line(619, '    axes[0].legend()', '')
add_line(620, '    axes[0].grid(True, alpha=0.3)', '')

add_line(622, '    deltas = [np.max(np.abs(V_track[i+1] - V_track[i])) for i in range(len(V_track)-1)]',
         '计算每轮之间的变化量（ΔV 的最大值）。')
add_line(623, '    axes[1].plot(deltas, \'k-\', linewidth=1.0)',
         '绘制 ΔV 随迭代次数的变化。')
add_line(624, '    axes[1].set_xlabel(\'Iteration\')', '')
add_line(625, '    axes[1].set_ylabel(\'Max ΔV\')', '')
add_line(626, '    axes[1].set_title(\'Convergence: Max Change per Iteration\')', '')
add_line(627, '    axes[1].set_yscale(\'log\')',
         'y 轴用对数尺度，因为 ΔV 是指数衰减。'
         '如果画出来是一条直线（在 log 尺度下），说明收敛速度是线性的。')

doc.add_page_break()
add_h1('核心知识点')
add_bold('1. 值迭代是在迭代 Bellman 最优算子')
add_text(
    'V_{k+1}(s) = max_a [R(s,a) + γ·Σ P·V_k(s\')]\n'
    '这个算子是压缩映射（contraction mapping），保证收敛到唯一不动点 V*。')

add_bold('2. 策略是副产品')
add_text(
    '值迭代在整个迭代过程中都不显式存策略。'
    '迭代结束后，从 V* 提取策略：π(s) = argmax_a Q*(s,a)。')

add_bold('3. 收敛速度指数级')
add_text(
    '从收敛图中可以看到 ΔV 呈指数衰减（log 尺度下是直线）。'
    '这是因为 Bellman 算子的压缩系数 = γ < 1。')

doc.add_page_break()
add_h1('面试追问')
add_qa('Q: 值迭代和策略迭代哪个更好？',
       '取决于问题规模。策略迭代收敛轮数少但每轮贵（内循环到收敛）；'
       '值迭代每轮便宜但需要更多轮。'
       '对于 GridWorld 这样的小问题，两者差别不大；'
       '大状态空间下策略迭代通常更优。')

add_qa('Q: 这个 DP 求解和 EMS 中的 DP 是什么关系？',
       '完全相同的数学原理！EMS 的 DP 求解器也是用 Bellman 方程：'
       '状态 = SOC × 功率，动作 = FC 功率分配，'
       '奖励 = -氢耗（负值，因为要最小化）。'
       '区别只是状态空间从 16 变成了 9000（150 SOC × 60 功率）。')

save(doc, 'Week9_Part8_值迭代_收敛可视化_逐行解释.docx')

print('\n=== 全部生成完成 ===')
