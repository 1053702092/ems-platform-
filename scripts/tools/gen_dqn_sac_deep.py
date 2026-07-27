#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 DQN / SAC 深度学习原理精讲 (v2)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def p(text, bold=False, sz=11, color=None, indent=0):
    pa = doc.add_paragraph()
    if indent: pa.paragraph_format.left_indent = Cm(indent)
    run = pa.add_run(text)
    run.font.name = '微软雅黑'; run.font.size = Pt(sz); run.bold = bold
    if color: run.font.color.rgb = color

def b(text, sz=11, color=None): p(text, bold=True, sz=sz, color=color)

def bullet(text, lv=0):
    pa = doc.add_paragraph(text, style='List Bullet')
    pa.paragraph_format.left_indent = Cm(1.5 + lv * 0.8)

def code(lines, label=None):
    if label:
        pa = doc.add_paragraph()
        run = pa.add_run(label); run.bold = True; run.font.size = Pt(10)
    for line in lines.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before = Pt(0)
        pa.paragraph_format.space_after = Pt(0)
        pa.paragraph_format.left_indent = Cm(1)
        run = pa.add_run(line)
        run.font.name = 'Consolas'; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, hd in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(hd)
        run.bold = True; run.font.size = Pt(10); run.font.name = '微软雅黑'
    for rd in rows:
        row = t.add_row()
        for c, txt in enumerate(rd):
            run = row.cells[c].paragraphs[0].add_run(txt)
            run.font.size = Pt(10); run.font.name = '微软雅黑'

def brk(): doc.add_page_break()

# ======================== 封面 ========================
for _ in range(5): doc.add_paragraph('')
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('DQN / SAC 深度学习原理解析'); run.font.size = Pt(26); run.bold = True; run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t2.add_run('从监督学习视角理解强化学习算法'); run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t3.add_run(f'生成日期：{datetime.date.today().isoformat()}'); run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
brk()

# ======================== 预备知识 ========================
h('预备知识：监督学习 vs 强化学习', 1)
p('在深入 DQN 和 SAC 之前，先建立两个框架的对应关系。这样你从监督学习切过来就不会觉得完全是另一套东西。')

tbl(
    ['维度', '监督学习', '强化学习'],
    [
        ['数据', '(x, y) 特征-标签对', '(s, a, r, sp) 状态-动作-奖励-下一状态'],
        ['目标', '最小化预测误差', '最大化累积奖励'],
        ['训练数据来源', '固定的数据集', '智能体自己采集（探索 vs 利用）'],
        ['标签', '人工标注，固定不变', '奖励函数 + 自举（动态变化）'],
        ['输出', '分类/回归值', '动作（离散或连续）'],
        ['评估指标', '准确率 / MSE', '累积奖励 / 成功率'],
    ]
)
p('')
b('核心差异一句话：', sz=12, color=RGBColor(0xC0, 0x39, 0x2B))
b('监督学习 = 给定数据学映射；强化学习 = 边试边学，数据本身由当前策略产生。', sz=11)
p('')
p('这个差异导致了 RL 的三个核心挑战：')
bullet('探索-利用困境：是选已知的好动作，还是试新动作？')
bullet('信用分配：走了100步才得到奖励，中间哪一步该负责？')
bullet('非平稳目标：训练目标（Q 值）本身在训练过程中不断变化')
brk()

# ================================================================
# 第一部分：DQN
# ================================================================
h('第一部分：DQN -- 深度 Q 网络', 1)
p('-- Value-Based 深度强化学习的开山之作（DeepMind, 2013 / Nature 2015）', color=RGBColor(0x66,0x66,0x66))

h('1.1 从 Q-Learning 到 DQN：一步一步推导', 2)

h('1.1.1 先理解 Q-Learning（表格版）', 3)
p('Q-Learning 是 DQN 的前身，用一个 Q 表存储每个 (状态, 动作) 的价值：')

code('''Q 表：每个格子存一个数值
        a1 (左)   a2 (右)   a3 (上)   a4 (下)
s1       0.5       0.8       0.3       0.6
s2      -0.2       0.1       0.4       0.7''')

p('更新公式（贝尔曼最优方程的自举形式）：', bold=True)
code('''Q(s,a) <- Q(s,a) + alpha * [ r + gamma * max_a' Q(s',a') - Q(s,a) ]
                  |               TD target (目标值)           |
                                                  TD error (差距)

如果目标值 > 当前估计：增大 Q(s,a)
如果目标值 < 当前估计：减小 Q(s,a)''')

p('逐项拆解：')
bullet('alpha (学习率)：每次更新迈多大步。太大震荡，太小收敛慢。')
bullet('r + gamma * max_a\' Q(s\', a\')：这是"实际观察到的奖励 + 对未来最优价值的估计"，比当前 Q(s,a) 更接近真实值，所以把它当目标')
bullet('Q-Learning 是 off-policy 的：更新时用的 max 对应的是"最优策略"，而不是当前正在执行的策略')
p('')
b('表格 Q-Learning 的局限：', color=RGBColor(0xC0,0x39,0x2B))
bullet('围棋状态数约 10^170，存 Q 表需要比宇宙原子还多的内存')
bullet('表格方法无法泛化：没见过的新状态不知道该怎么处理')

h('1.1.2 DQN 的解决方案：用神经网络代替 Q 表', 3)
p('DQN 把 Q(s,a) 替换为一个神经网络 Q_theta(s,a)，参数为 theta：')
bullet('输入：状态 s（可以是图像、向量等）')
bullet('输出：每个动作的 Q 值（离散动作）')
bullet('参数 theta：通过梯度下降从数据中学习')
p('')
b('监督学习视角理解 DQN：', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
p('把 DQN 的每一步训练看作一个监督学习问题：')
tbl(
    ['', '标准监督学习', 'DQN'],
    [
        ['输入 x', '图片', '状态 s'],
        ['输入 x', '图片', '状态 s'],
        ['标签 y', '人工标注（固定）', 'r + gamma * max Q(s\', a\')（动态生成）'],
        ['预测', 'f_theta(x)', 'Q_theta(s, a)'],
        ['损失', '(y - f_theta(x))^2', '(y - Q_theta(s,a))^2'],
        ['本质', '拟合固定的标签', '拟合自举生成的目标'],
    ]
)
p('')
b('关键区别：监督学习的标签是固定的，DQN 的目标是动态生成的——这就是 RL 训练不稳定的根本原因。')
brk()

# 1.2 两大创新
h('1.2 DQN 的两大核心创新', 2)

h('1.2.1 创新一：经验回放（Experience Replay）', 3)
b('标准监督学习假设数据是独立同分布的。但 RL 的数据是序列相关的：')
bullet('时间步 t 的状态 s_t 和 t+1 的状态 s_{t+1} 高度相关')
bullet('如果用顺序数据训练神经网络，网络会"过拟合"到最近的经验上')
bullet('这就像考试前只复习最后几页书——前面的全忘了')
p('')
b('经验回放的做法：')
p('1) 用一个缓冲区 D（通常容量 10^5 到 10^6）存储智能体的经验 (s, a, r, sp)')
p('2) 训练时从 D 中随机均匀采样一个 mini-batch')
p('3) 用采样的数据计算梯度，更新网络参数')
p('')
b('为什么有效：')
bullet('打破相关性：随机采样去除了时间上的序列相关性 -> 更像监督学习的独立同分布假设')
bullet('数据复用：一条经验可以被多次使用 -> 样本效率提高')
bullet('避免灾难性遗忘：网络不会只记住最近的行为')
p('')
b('类比理解：')
p('  学生 A（无经验回放）：考完一题扔一题，只记得最后几题')
p('  学生 B（有经验回放）：把所有做过的题记在本子上，复习时随机抽着看 -> 成绩更好')
p('')
p('伪代码：', bold=True)
code('''buffer = deque(maxlen=100000)   # 环形缓冲区

for step in range(T):
    action = agent.act(state)           # 选动作
    next_state, reward, done = env.step(action)  # 执行
    buffer.append((state, action, reward, next_state, done))
    state = next_state

    if len(buffer) >= batch_size:
        batch = random.sample(buffer, batch_size)  # 随机采样！
        loss = train_step(batch)         # 更新网络''')

h('1.2.2 创新二：目标网络（Target Network）', 3)
b('DQN 的损失函数是：')
code('''L(theta) = ( r + gamma * max_a' Q_theta(s',a')  -  Q_theta(s,a)  )^2
                |______ target y ______|   |_ prediction _|

问题：target y 和 prediction 用的是同一个 Q_theta！
如果 theta 更新了，target y 也跟着变 -> 靶子在动 -> 射不准。''')
p('')
b('目标网络的解决方案：')
p('维护两个 Q 网络：')
tbl(
    ['', '在线网络 Q_theta', '目标网络 Q_theta_target'],
    [
        ['作用', '输出预测值，计算梯度更新', '输出目标值中的 max Q(sp,ap)'],
        ['更新频率', '每步更新', '每 C 步复制，或软更新'],
        ['是否需要梯度', '是（被训练的主体）', '否（冻结参数）'],
    ]
)
p('')
p('硬更新（DQN 原文）：', bold=True)
code('''每 C 步执行一次:    theta_target <- theta
其余时间:           theta_target 冻结不变''')
p('')
p('软更新（更平滑）：', bold=True)
code('''每一步执行:  theta_target <- tau * theta + (1-tau) * theta_target
              tau 很小，通常 0.001 ~ 0.01''')
p('')
b('为什么目标网络有效：')
bullet('C 步内目标 y = r + gamma * max Q_target(s\', a\') 不变 -> 目标固定 -> 训练稳定')
bullet('目标网络的参数始终落后于在线网络 -> 避免"追自己尾巴"')
bullet('这是 RL 特有的技巧。监督学习中标签固定，不需要目标网络。')
brk()

# 1.3 完整算法
h('1.3 DQN 完整算法流程', 2)
code('''算法：DQN (Deep Q-Network with Experience Replay)

初始化：
  Q_theta           <- 随机初始化（在线网络）
  Q_theta_target    <- Q_theta（目标网络，参数相同）
  D                 <- 空缓冲区（容量 N）
  epsilon           <- 1.0（初始探索率）

for episode = 1 to M:
    s <- env.reset()
    for t = 1 to T:
        # 1. epsilon-贪婪选动作
        if random() < epsilon:
            a <- random_action()           # 探索
        else:
            a <- argmax Q_theta(s)         # 利用

        # 2. 与环境交互
        sp, r, done <- env.step(a)
        D.push(s, a, r, sp, done)          # 存储经验

        # 3. 训练（当 D 足够大时）
        if len(D) >= batch_size:
            batch <- D.sample(batch_size)  # 经验回放
            对每个 (s, a, r, sp, done) in batch:
                if done:
                    y = r                  # 终止状态无未来
                else:
                    y = r + gamma * max Q_target(sp)  # 目标网络
            loss = MSE(Q_theta(s,a), y)
            loss.backward()                # 计算梯度
            optimizer.step()               # 更新 Q_theta

        # 4. 更新目标网络（硬更新）
        if t % C == 0:
            Q_target.load_state_dict(Q_theta.state_dict())

        # 5. epsilon 衰减
        epsilon <- max(epsilon_min, epsilon * decay)

        s <- sp
        if done: break''')
brk()

# 1.4 网络结构
h('1.4 DQN 的网络结构（理解 DL 视角的关键）', 2)
p('DQN 的网络结构取决于输入类型：')
p('')
b('情况一：图像输入（如 Atari 游戏）')
code('''输入：4 帧堆叠的灰度图像 -> shape = [84, 84, 4]
  ↓
CNN 层：3 层卷积 -> 提取空间特征
  ↓
全连接层：512 个神经元 -> 特征融合
  ↓
输出：每个动作的 Q 值 -> shape = [n_actions]  （无 softmax！是原始分值）
      例如 Atari 有 4-18 个离散动作''')
p('')
b('情况二：向量输入（如 EMS 状态）')
code('''输入：[SOC, P_load, V_bat, T_fc, ...] -> shape = [n_features]
  ↓
全连接层：256 -> 128 -> 64 -> 逐层提取特征
  ↓
输出：每个动作的 Q 值 -> shape = [n_actions]''')
p('')
b('注意：DQN 输出层没有激活函数（回归头）。Q 值可以是任意实数。', color=RGBColor(0xC0,0x39,0x2B))
brk()

# 1.5 DQN 的改进
h('1.5 DQN 的三大改进', 2)
p('原始 DQN 有三个明显缺陷，后来有针对性改进：')

h('缺陷 1：Q 值过估计 -> Double DQN', 3)
b('问题：max Q 操作导致估计值系统性偏高。')
p('Q_theta(s\', a\') 本身有估计误差，max 会取到被高估的动作 -> Q 值越算越高')
p('')
b('修复：选动作和算 Q 值用不同的网络')
code('''原始 DQN:    y = r + gamma * max_a' Q_target(s', a')
              <- 同一个目标网络既选动作又算 Q

Double DQN:  y = r + gamma * Q_target(s', argmax_a' Q_online(s', a'))
              |                          |
          目标网络算 Q 值          在线网络选动作''')
p('用在线网络选"哪个动作最好"，用目标网络算"这个动作值多少"——双重校验，降低高估。')

h('缺陷 2：所有动作共享网络 -> Dueling DQN', 3)
p('有些状态下，动作选择不重要（比如直线公路，左右转都一样）。把 Q 拆成两部分：')
code('''Q(s,a) = V(s) + A(s,a)
         |         |
   状态价值   动作优势（相对 V(s) 的优劣）''')
p('V(s) 告诉你在状态 s 能拿多少分，A(s,a) 告诉你选 a 比平均好多少。')

h('缺陷 3：采样不高效 -> PER (Prioritized Experience Replay)', 3)
p('所有经验同等对待不是最优的。有些经验更"意外"，应该更常学。')
bullet('每条经验计算 TD error = |y - Q(s,a)|')
bullet('TD error 越大 -> 采样概率越高 -> 学到更多')
bullet('同时保留一定的随机采样，避免过拟合到少数极端经验')
brk()

# 1.6 DQN 局限
h('1.6 DQN 的固有局限', 2)
p('DQN 在 Atari 上取得了突破性成果，但它有三条根深蒂固的局限：')
tbl(
    ['局限', '根本原因', '后果', '解决方向'],
    [
        ['只能处理离散动作', '输出层每个神经元对应一个动作\nargmax需要遍历所有动作', '无法处理连续控制\n（发动机功率、机器人关节）', '-> 策略梯度方法\n-> Actor-Critic'],
        ['Q 值过估计', 'max 操作放大估计误差', '策略次优\nQ 值发散', '-> Double DQN'],
        ['训练不稳定', '自举目标 + 函数近似\n（双重不稳定源）', '收敛慢\n需要大量调参', '-> 目标网络\n-> 梯度裁剪'],
    ]
)
p('')
b('这些局限促使研究者开辟了另一条路线——"基于策略的强化学习"，也就是 SAC 所在的流派。')
brk()

# ================================================================
# 第二部分：SAC
# ================================================================
h('第二部分：SAC -- Soft Actor-Critic', 1)
p('-- 基于最大熵的连续控制算法（Haarnoja et al., 2018）', color=RGBColor(0x66,0x66,0x66))

h('2.0 背景：为什么需要 Policy-Based 方法？', 2)
p('DQN 的"先算 Q 值再 argmax"范式有两个根本问题：')
p('')
bullet('问题 1：连续动作空间无法 argmax。假设 P_fc 属于 [0, 30] kW，你怎么遍历所有实数求最大值？即使离散化成 100 档，100 个输出神经元 + 精度损失也不理想。')
bullet('问题 2：DQN 是确定性策略。给同一个状态永远输出同一个动作。但现实中最优策略往往是随机的——比如石头剪刀布。')
p('')
b('Policy-Based 方法直接学一个策略网络 pi_theta(a|s) -> 输出动作的概率分布：')
bullet('连续动作：输出高斯分布的均值 mu 和标准差 sigma -> 从分布中采样')
bullet('随机策略：可以输出动作的概率 -> 在不确定时保持随机')
brk()

# 2.1 从 REINFORCE 到 Actor-Critic
h('2.1 从 REINFORCE 到 Actor-Critic', 2)

h('阶段 1：REINFORCE（蒙特卡洛策略梯度）', 3)
p('最朴素的策略梯度方法：')
code('''策略网络 pi_theta(a|s)：输出动作的概率

更新公式：grad_J(theta) = E[ grad_log_pi_theta(a|s) * G ]
            提升方向         对动作概率求导    整条轨迹的累积奖励

直觉：如果整条轨迹拿到了高奖励 G -> 增大每个动作的概率
      如果拿到了低奖励 G    -> 减小每个动作的概率''')
p('')
b('REINFORCE 的问题：', color=RGBColor(0xC0,0x39,0x2B))
bullet('G 直到轨迹结束才能算出来 -> 更新慢（需要整条轨迹跑完）')
bullet('G 的方差极大 -> 训练不稳定')
bullet('类比：看完整部电影才评价演员"你第三分钟演得好不好"——说不清楚')

h('阶段 2：Actor-Critic', 3)
p('改进思路：不等到轨迹结束，每步都做评估。引入 Critic 网络。')
p('')
b('Actor（演员）：策略网络 pi_theta(a|s) -> 决定怎么做')
b('Critic（评委）：价值网络 V_phi(s) 或 Q_phi(s,a) -> 评价做得好不好')
p('')
code('''Actor 的更新（基于 Critic 的反馈）：
  grad_J(theta) = E[ grad_log_pi_theta(a|s) * Q_phi(s,a) ]
  如果 Critic 说 Q 值高 -> 增大这个动作的概率

Critic 的更新（基于实际奖励）：
  L(phi) = ( r + gamma * V_phi(sp) - V_phi(s) )^2
  让价值估计更接近实际观测到的回报''')
p('')
b('类比理解：')
p('  REINFORCE = 演员自己看完整部票房结果来琢磨哪演得好（方差大）')
p('  Actor-Critic = 演员演一幕，导演当场点评（方差小，更新快）')
brk()

# 2.2 最大熵
h('2.2 SAC 的核心创新：最大熵框架', 2)

b('标准 RL 的目标')
code('max E [ sum gamma^t * r(s_t, a_t) ]')

b('最大熵 RL 的目标')
code('max E [ sum gamma^t * ( r(s_t, a_t) + alpha * H(pi(·|s_t)) ) ]')
p('')
b('多出来的一项：alpha * H(pi(·|s_t))', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
p('')
b('什么是熵（Entropy）？')
p('H(pi(·|s)) = -sum pi(a|s) * log pi(a|s)')
p('')
bullet('熵 -> 0：动作分布集中在某一个动作上 -> 确定性策略')
bullet('熵 -> 大：动作分布均匀 -> 随机策略')
bullet('例子：{左=1.0, 右=0.0} -> H=0  （完全确定）')
bullet('例子：{左=0.5, 右=0.5} -> H=0.693  （随机）')
bullet('例子：{四个方向各 0.25} -> H=1.386  （完全均匀）')
p('')
b('为什么加熵项？三个层面理解：')
p('')
b('层面一：鼓励探索', color=RGBColor(0x1F,0x3A,0x5F))
bullet('标准 RL 一旦发现某个动作好，就会死锁在这个动作上（过早收敛）')
bullet('熵奖励让策略保持一定的随机性，继续探索其他可能性')
bullet('类比：一个厨师学会做宫保鸡丁后就不再学新菜 -> 熵太低。保持探索才可能发现更好的菜。')
p('')
b('层面二：鲁棒性', color=RGBColor(0x1F,0x3A,0x5F))
bullet('学过"多种做法"的策略，遇到环境变化不容易崩溃')
bullet('类比：只会做一道菜的厨师，缺了关键材料就完了。会十道菜的厨师，换材料也能应变。')
bullet('在 EMS 中：训练时学的工况和实际工况不可能完全一样 -> 鲁棒策略更重要')
p('')
b('层面三：多模态', color=RGBColor(0x1F,0x3A,0x5F))
bullet('当多个动作同样优秀时，标准 RL 会随机选一个，但 SAC 会均匀分配概率')
bullet('保持多种可能 -> 未来信息到来时可以灵活切换')
brk()

# 2.3 alpha
h('2.3 自动温度调节（Auto Temperature）', 2)
p('alpha（温度系数）控制"探索"和"利用"的平衡：')
bullet('alpha = 0：退化为标准 RL，熵项消失')
bullet('alpha 很大：熵奖励主导，智能体几乎随机行动')
p('')
b('手动调 alpha 很麻烦——不同任务最优 alpha 不同。SAC 自动调 alpha：')
p('')
code('''把 alpha 当作一个可训练的参数（而不是手动设定的超参数）
目标熵 H_target = -dim(A)  （比如动作空间 4 维 -> H_target = -4）

loss_alpha = -alpha * ( log_pi_theta(a|s) + H_target )
            <- 梯度下降更新 alpha

当实际熵 < H_target（策略太确定）：loss 增大 -> alpha 增大 -> 鼓励探索
当实际熵 > H_target（策略太随机）：loss 减小 -> alpha 减小 -> 专注利用''')
p('')
b('这意味着你不需要手动设 epsilon 或探索率——SAC 自动调节探索程度。')
brk()

# 2.4 5网络架构
h('2.4 SAC 的五网络架构', 2)
p('SAC 同时维护 5 个网络。听起来多，但分工清晰：')
tbl(
    ['网络', '符号', '输入', '输出', '更新方式'],
    [
        ['Actor\n（策略网络）', 'pi_phi', 's', '动作分布\n(mu, sigma)', '最大化 Q + 熵\n（策略梯度）'],
        ['Critic Q1', 'Q_theta1', '(s, a)', 'Q 值', '最小化 Bellman 误差'],
        ['Critic Q2', 'Q_theta2', '(s, a)', 'Q 值', '最小化 Bellman 误差\n（独立于 Q1）'],
        ['目标 Q1_target', 'Q_theta1_t', '(s, a)', '目标 Q 值', '软更新自 Q_theta1'],
        ['目标 Q2_target', 'Q_theta2_t', '(s, a)', '目标 Q 值', '软更新自 Q_theta2'],
    ]
)
p('')
b('为什么两个 Critic（双 Q 技巧）？')
p('DQN 的过估计问题同样影响 SAC。解法：取 min(Q1, Q2) 作为 Q 值的估计。')
p('两个 Critic 独立训练，同一个 (s,a) 给出两个独立估计 -> 取最小值 -> 防止高估。')
p('这是 SAC 训练稳定的关键设计之一。')
brk()

# 2.5 完整算法
h('2.5 SAC 完整算法流程（伪代码）', 2)
code('''算法: SAC (Soft Actor-Critic)

初始化:
  pi_phi             <- 随机初始化（Actor）
  Q_theta1, Q_theta2 <- 随机初始化（两个 Critic）
  Q_t1, Q_t2         <- Q_theta1, Q_theta2（目标网络）
  D                  <- 空缓冲区
  alpha              <- 0.2（可训练的探索温度）

for each step:
    # ====== 采集 ======
    s -> a ~ pi_phi(a|s)           # 从策略分布采样连续动作
    sp, r, done <- env.step(a)
    D.push(s, a, r, sp, done)

    if len(D) < batch_size: continue

    # ====== 训练 ======
    batch <- D.sample(batch_size)

    # 1. 更新 Critic（最小化 Bellman 误差）
    ap ~ pi_phi(ap|sp)            # 下一状态的动作
    Q_target = min(Q_t1(sp,ap), Q_t2(sp,ap)) - alpha * log_pi_phi(ap|sp)
    y = r + gamma * Q_target       # 目标值（含熵奖励）
    loss_Q1 = MSE(Q_theta1(s,a), y)
    loss_Q2 = MSE(Q_theta2(s,a), y)
    梯度下降更新 theta1, theta2

    # 2. 更新 Actor（最大化 Q + 熵）
    a_new ~ pi_phi(a|s)
    Q_value = min(Q_theta1(s,a_new), Q_theta2(s,a_new))
    loss_pi = (alpha * log_pi_phi(a_new|s) - Q_value).mean()
    梯度下降更新 phi

    # 3. 更新 alpha（自动温度调节）
    loss_alpha = -alpha * (log_pi_phi(a|s) + H_target).mean()
    梯度下降更新 alpha

    # 4. 软更新目标网络
    Q_t1 <- tau * Q_theta1 + (1-tau) * Q_t1
    Q_t2 <- tau * Q_theta2 + (1-tau) * Q_t2''')
brk()

# 2.6 重参数化
h('2.6 重参数化技巧（Reparameterization Trick）', 2)
p('Actor 输出的是"动作分布"（均值 mu 和标准差 sigma），但采样动作 a ~ N(mu, sigma) 这个操作不可导——梯度无法回传到 pi_phi。')
p('')
b('SAC 的解法：')
code('''采样过程：a = mu + sigma * epsilon,  其中 epsilon ~ N(0, 1)
                               |
                       噪声从标准正态采样，与网络无关

这样梯度可以流过 mu 和 sigma 直接更新 Actor 网络参数 phi''')
p('')
b('类比：')
p('  不可导版本：拿一个骰子掷出结果，改骰子的形状和结果无关（梯度断掉）')
p('  重参数化：做一个"确定性部分 + 随机噪声部分"的骰子，改形状能影响结果分布（梯度可通过）')
brk()

# 2.7 SAC vs PPO
h('2.7 SAC vs PPO：两大主流算法对比', 2)
p('SAC 和 PPO 是目前连续控制领域最主流的两个算法。它们的核心差异：')
tbl(
    ['维度', 'SAC', 'PPO'],
    [
        ['算法类型', 'Off-policy', 'On-policy'],
        ['样本效率', '高（复用历史数据）', '低（每个样本只能用一次）'],
        ['网络数量', '5 个', '2 个（Actor + Critic）'],
        ['实现复杂度', '高', '中'],
        ['超参数敏感度', '较低', '中（clip 参数敏感）'],
        ['训练稳定', '双 Q + 软更新', 'clip 限制单步更新幅度'],
        ['收敛速度（样本数）', '快（off-policy）', '慢（on-policy）'],
        ['收敛速度（壁钟时间）', '慢（5 网络计算量大）', '快（2 网络计算量小）'],
    ]
)
p('')
b('如何选择：', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
bullet('数据获取代价高（如真实机器人）-> SAC（样本效率高）')
bullet('计算资源有限/需要快速迭代 -> PPO（实现简单，训练快）')
bullet('工业落地 -> PPO（超参数鲁棒，调参成本低）')
bullet('SOTA 追求 -> SAC（通常性能略高于 PPO）')
brk()

# 2.8 应用案例
h('2.8 SAC 在深度学习中的典型应用', 2)
p('SAC 不只为 EMS 而生，它在整个深度强化学习领域广泛使用：')
p('')
b('应用一：机器人控制（MuJoCo 基准）')
bullet('Ant, HalfCheetah, Humanoid 等连续控制任务')
bullet('SAC 在这些基准上取得了 SOTA 成绩')
bullet('输出：关节力矩（连续值）')
p('')
b('应用二：自动驾驶决策')
bullet('状态：车辆位置、速度、周围车辆信息')
bullet('动作：方向盘角度、油门开度、刹车力度（全连续）')
bullet('奖励：到达目标 + 安全约束 + 舒适度')
p('')
b('应用三：能源管理')
bullet('电池储能调度、建筑 HVAC 控制、数据中心冷却')
bullet('连续控制 + 长期回报 + 安全约束 -> SAC 很适合')
p('')
b('应用四：游戏（连续动作类）')
bullet('虽然游戏通常用 PPO 或 DQN，SAC 也可用于连续动作的游戏')
brk()

# ================================================================
# 第三部分：总对比
# ================================================================
h('第三部分：全方位对比', 1)
tbl(
    ['维度', 'DQN', 'SAC', 'PPO'],
    [
        ['提出年份', '2013 (DeepMind)', '2018 (UC Berkeley)', '2017 (OpenAI)'],
        ['流派', 'Value-Based', 'Actor-Critic', 'Policy-Based'],
        ['动作空间', '离散', '连续', '连续'],
        ['策略类型', '确定性（argmax）', '随机性（分布采样）', '随机性（分布采样）'],
        ['是否 off-policy', '是', '是', '否（on-policy）'],
        ['网络数量', '2', '5', '2'],
        ['核心机制', '经验回放 + 目标网络', '最大熵 + 双Q + 自动alpha', 'Clipped Surrogate'],
        ['样本效率', '中', '高', '低'],
        ['训练稳定性', '较不稳定', '稳定', '稳定'],
        ['实现难度', '中等', '高', '中等'],
        ['主要局限', '离散动作\nQ值过估计', '实现复杂\n计算量大', '样本效率低\n交互需求大'],
    ]
)
brk()

# ================================================================
# 第四部分：面试
# ================================================================
h('第四部分：面试准备（从 DL 视角回答）', 1)

b('Q1: "DQN 为什么需要经验回放？从监督学习角度解释。"')
p('监督学习假设数据是独立同分布的。但 RL 的数据是序列相关的——连续采样的状态之间高度相关。如果用顺序数据训练神经网络，网络会过拟合到最近的经验上，产生灾难性遗忘。经验回放通过随机采样打破数据相关性，使训练数据更接近独立同分布假设。同时，一条经验可以被多次使用，提高了样本效率。')
p('')
b('Q2: "DQN 的目标网络解决了什么问题？"')
p('DQN 的更新目标 y = r + gamma * max Q_theta(s\', a\') 依赖于同一个网络 Q_theta。这意味着每次 theta 更新时 y 也随之变化——训练目标在"追着自己的尾巴跑"。目标网络把 y 固定一段时间，让在线网络去追赶一个相对稳定的目标，大幅提高训练的稳定性。这在监督学习中不会出现，因为监督学习的标签是固定的。')
p('')
b('Q3: "SAC 的最大熵是什么意思？为什么有用？"')
p('最大熵是指在标准 RL 目标上增加一项策略熵的奖励。熵衡量随机性，熵越大表示动作分布越均匀。加熵奖励的三个好处：1) 鼓励探索，避免过早收敛到次优策略；2) 提高鲁棒性，学到多种应对方式；3) 处理多模态情况，多个等优动作时均匀分配概率。alpha 自动调节探索程度，不需要手动设 epsilon。')
p('')
b('Q4: "SAC 为什么需要两个 Critic 网络？"')
p('这是为了解决 Q 值的过估计问题。DQN 中 max Q 操作会系统性高估 Q 值。SAC 训练两个独立的 Critic 网络，在计算目标值时取 min(Q1, Q2) 作为 Q 值的估计——这相当于一个"悲观的估计"，可以有效防止高估。这是 SAC 训练稳定的关键设计之一。')
p('')
b('Q5: "SAC 和 PPO 各适合什么场景？"')
p('SAC 是 off-policy 算法，样本效率高，适合数据获取成本高的场景（如真实机器人、硬件在环）。PPO 是 on-policy 算法，实现简单、超参数鲁棒、壁钟时间收敛快，适合计算资源有限或需要快速迭代的场景。工业界 PPO 更常见（实现简单意味着维护成本低），学术界 SAC 更常见（追求 SOTA 性能）。')

brk()

# ================================================================
# 附录
# ================================================================
h('附录：关键公式速查', 1)

b('DQN')
p('Q-learning 更新：', bold=True)
code('Q(s,a) <- Q(s,a) + alpha * [r + gamma * max Q(s\',a\') - Q(s,a)]')
p('DQN 损失函数：', bold=True)
code('L(theta) = E[( r + gamma * max Q_target(s\',a\') - Q_theta(s,a) )^2]')
p('Epsilon-greedy 探索：', bold=True)
code('a = random_action  if p < epsilon;  else argmax Q(s)')
p('')
b('SAC')
p('最大熵目标：', bold=True)
code('J = sum E[ r + alpha * H(pi(·|s)) ]')
p('Critic 损失：', bold=True)
code('y = r + gamma * ( min(Q_target(sp,ap)) - alpha * log_pi(ap|sp) )')
code('L_Q = E[(Q1(s,a) - y)^2 + (Q2(s,a) - y)^2]')
p('Actor 损失：', bold=True)
code('L_pi = E[ alpha * log_pi(a|s) - min(Q1(s,a), Q2(s,a)) ]')
p('Alpha 更新：', bold=True)
code('L_alpha = -E[ alpha * (log_pi(a|s) + H_target) ]')
p('软更新目标网络：', bold=True)
code('Q_target <- tau * Q + (1-tau) * Q_target')

# ======================== 保存 ========================
path = os.path.join(OUT_DIR, 'DQN_SAC_深度学习原理解析.docx')
doc.save(path)
print(f'OK: {path}')
