#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 week10_dqn_gridworld.py 逐行精讲 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def p(text, bold=False, sz=11, color=None, indent=0):
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(3)
    if indent: pa.paragraph_format.left_indent = Cm(indent)
    run = pa.add_run(text); run.font.name='微软雅黑'; run.font.size=Pt(sz); run.bold=bold
    if color: run.font.color.rgb = color

def b(text, sz=11, color=None): p(text, bold=True, sz=sz, color=color)

def bullet(text, lv=0):
    pa = doc.add_paragraph(text, style='List Bullet')
    pa.paragraph_format.left_indent = Cm(1.5 + lv*0.8)

def code(lines, label=None):
    if label:
        pa = doc.add_paragraph(); run=pa.add_run(label); run.bold=True; run.font.size=Pt(10)
    for line in lines.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before=Pt(0); pa.paragraph_format.space_after=Pt(1)
        pa.paragraph_format.left_indent=Cm(1)
        run=pa.add_run(line); run.font.name='Consolas'; run.font.size=Pt(9.5); run.font.color.rgb=RGBColor(0x33,0x33,0x33)

def tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i,hd in enumerate(headers):
        run=t.rows[0].cells[i].paragraphs[0].add_run(hd); run.bold=True; run.font.size=Pt(10); run.font.name='微软雅黑'
    for rd in rows:
        row=t.add_row()
        for c,txt in enumerate(rd):
            run=row.cells[c].paragraphs[0].add_run(txt); run.font.size=Pt(10); run.font.name='微软雅黑'

def brk(): doc.add_page_break()
def note(text):
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    run = pa.add_run(f'  {text}')
    run.font.name = '微软雅黑'; run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

def explain(line_num, code_line, explanation):
    """一行代码 + 解释"""
    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(6)
    pa.paragraph_format.space_after = Pt(2)
    run = pa.add_run(f'第 {line_num} 行  ')
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x99,0x99,0x99)
    run = pa.add_run(code_line)
    run.font.name = 'Consolas'; run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = RGBColor(0x1F,0x3A,0x5F)
    p(explanation, indent=1, sz=10)

# ======================== 封面 ========================
for _ in range(4): doc.add_paragraph('')
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t.add_run('Week 10 — DQN GridWorld\n逐行精讲'); run.font.size=Pt(26); run.bold=True; run.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t2.add_run('用神经网络代替 Q 表，每一行都讲清楚"为什么这么写"'); run.font.size=Pt(13); run.font.color.rgb=RGBColor(0x66,0x66,0x66)
t3=doc.add_paragraph(); t3.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t3.add_run(f'生成日期：{datetime.date.today().isoformat()}'); run.font.size=Pt(10); run.font.color.rgb=RGBColor(0x99,0x99,0x99)
brk()

# ================================================================
h('文件概览', 1)
p('文件名: week10_dqn_gridworld.py')
p('功能: 在 4×4 GridWorld 上用 DQN 训练一个智能体从起点走到终点')
p('对照: 和 week10_qlearning_gridworld.py 共用同一个环境，唯一区别是 Q 表换成了神经网络')
p('')
tbl(
    ['部分', '行号', '内容'],
    [
        ['环境定义', '22-65', '4×4 GridWorld（和 Q-learning 完全一样）'],
        ['DQN 神经网络', '52-74', '用 PyTorch 定义网络结构'],
        ['DQN 算法', '78-182', '经验回放 + 目标网络 + 训练循环'],
        ['结果展示', '186-194', '保存训练数据到 CSV'],
        ['主程序', '197-250', '跑 5000 局 + 打印策略'],
    ]
)

brk()

# ================================================================
h('第一部分：导入和常量（1-49 行）', 1)

explain(1, '#!/usr/bin/env python3', 'Shebang 行。告诉系统用 Python 3 解释器运行这个文件。')

explain(2, '# -*- coding: utf-8 -*-', '文件编码声明。支持中文注释和字符串。')

explain(3, '"""', '多行文档字符串开始。描述这个文件是干什么的。')

p('从第 3 行到第 11 行的注释块说明了这个文件的核心卖点：', indent=1, sz=10)
bullet('Q-learning: Q[s][a] ← 一张 16×4 的表格', lv=1)
bullet('DQN:        Q_theta(s, a) ← 一个神经网络', lv=1)
bullet('环境：同一个 4x4 GridWorld', lv=1)
p('所以当你读这个文件时，心里要一直想着：和 Q-learning 比，到底改了哪？答案就一个——存 Q 值的方式变了。', indent=1, sz=10)

explain(13, 'import numpy as np', 'NumPy——处理数组、算均值、argmax。和 Q-learning 一样。')

explain(14, 'import random', 'Python 内置随机库。用来生成随机数、随机选动作。和 Q-learning 一样。')

explain(15, 'import os', '路径和文件操作。用来创建 results 目录、拼接文件路径。')

explain(16, 'import torch', 'PyTorch 主库。这是 DQN 和 Q-learning 最根本的区别——QL 不需要深度学习框架，DQN 需要。')

explain(17, 'import torch.nn as nn', 'PyTorch 的神经网络模块。包含 Linear（全连接层）、ReLU 激活函数、MSELoss（损失函数）。')

explain(18, 'import torch.optim as optim', 'PyTorch 的优化器模块。包含 Adam、SGD 等优化算法。DQN 用 Adam 代替了 QL 的手动 lr 更新。')

explain(20, 'RESULTS_DIR = ...', '保存结果文件的目录路径。所有生成的 CSV 文件都会放在这里。')

brk()

h('1.1 环境常量', 2)

explain(22, 'SIZE = 4', '网格大小。4×4 = 16 个格子。和 Q-learning 一模一样。')

explain(23, 'N_STATES = SIZE * SIZE', '状态总数 = 16。每个格子是一个状态，编号 0-15。')

explain(24, 'N_ACTIONS = 4', '动作数量 = 4 个：上(0)、下(1)、左(2)、右(3)。')

explain(25, 'GOAL_IDX = 15', '终点状态编号。右下角格子 (3,3)，编号 15。走到这里得 +1 分。')

explain(26, 'TRAP_IDX = 5', '陷阱状态编号。位置 (1,1)，编号 5。走到这里得 -1 分。')

explain(27, 'GAMMA = 0.9', '折扣因子 γ。未来奖励的折扣率。和 Q-learning 完全一样。')

explain(28, 'ACTION_DELTA = [(-1, 0), (1, 0), (0, -1), (0, 1)]', '动作对应的坐标变化。用行、列的偏移量来表示：↑(-1,0)、↓(1,0)、←(0,-1)、→(0,1)。')

explain(29, 'ACTION_SYMBOLS = {0: "↑", 1: "↓", 2: "←", 3: "→"}', '给动作编号加上箭头符号，打印策略时更方便看清。')

brk()

h('1.2 环境函数', 2)

explain(31, 'def is_valid(r, c):', '检查坐标 (r, c) 是否在网格内。如果越界返回 False。')
p('比如 is_valid(-1, 0) → False（出了上边界）', sz=10, indent=1, color=RGBColor(0x66,0x66,0x66))
p('比如 is_valid(0, 0) → True（左上角，有效）', sz=10, indent=1, color=RGBColor(0x66,0x66,0x66))

explain(35, 'def step(s, a):', '执行动作 a，返回 (s_next, reward, done)。这是环境和智能体交互的唯一接口。')
b('和 Q-learning 完全一样。', color=RGBColor(0x1F,0x3A,0x5F))

explain(37, 'r, c = divmod(s, SIZE)', '把状态编号 s 转成网格坐标。divmod(5, 4) → (1, 1)，因为 5 = 1×4 + 1。')

explain(38-42, 'if random.random() < 0.8: ... else: ...', '80% 概率走选的方向，20% 概率滑到其他方向。模拟环境的不确定性，让问题更有挑战性。如果没有这个随机性，问题就太简单了。')

explain(44-46, 'nr, nc = r + dr, c + dc', '计算新坐标。如果撞墙（超出边界），就留在原地。')

explain(49-65, 'if s_next == GOAL_IDX: ...', '根据新位置决定奖励和是否结束：')
bullet('走到终点 → reward=1.0, done=True', lv=1)
bullet('走进陷阱 → reward=-1.0, done=True', lv=1)
bullet('其他位置 → reward=0.0, done=False', lv=1)

brk()

# ================================================================
h('第二部分：DQN 神经网络（52-74 行）', 1)

p('这是 DQN 和 Q-learning 的第一个核心区别：Q 表 → 神经网络。', bold=True, sz=12, color=RGBColor(0x1F,0x3A,0x5F))

explain(53, 'class DQN(nn.Module):', '定义一个神经网络类，继承自 PyTorch 的 nn.Module。这个网络的作用和 Q 表一模一样：输入状态 s，输出 4 个 Q 值。')

explain(58, 'def __init__(self, state_dim=16, hidden=32):', '构造函数。state_dim=16 输入维度（one-hot 编码），hidden=32 隐藏层神经元数。')

explain(59, 'super().__init__()', '调用父类 nn.Module 的构造函数。PyTorch 要求必须调用。')

explain(60-64, 'self.net = nn.Sequential(...)', '定义网络结构，按顺序堆叠：')
bullet('nn.Linear(16, 32)：全连接层，16 维输入 → 32 维隐藏。权重矩阵形状 [32×16]', lv=1)
bullet('nn.ReLU()：激活函数，把负数变成 0，引入非线性。没有它，多层就和一层没区别。', lv=1)
bullet('nn.Linear(32, 4)：全连接层，32 维隐藏 → 4 维输出（4 个 Q 值）。权重矩阵形状 [4×32]', lv=1)
p('总参数 = 16×32 + 32 + 32×4 + 4 = 512 + 32 + 128 + 4 = 644 个 (加上bias是676)', indent=1, sz=10)

explain(66, 'def forward(self, x):', '前向传播函数。输入 x（状态向量），输出 4 个 Q 值。')
p('PyTorch 网络的核心：把输入"流"过各层，最终输出结果。', indent=1, sz=10)

explain(67, 'return self.net(x)', '把输入 x 送进 net 序列，依次通过 Linear → ReLU → Linear，返回 4 个 Q 值。')
p('这和 Q 表的关系：net(格子0) → [Q_↑, Q_↓, Q_←, Q_→]，和 Q[0] → [Q_↑, Q_↓, Q_←, Q_→] 输出格式完全一样。', indent=1, sz=10, color=RGBColor(0x1F,0x3A,0x5F))

brk()

h('2.1 状态编码函数', 2)

explain(70, 'def state_to_tensor(s):', '把状态编号（0-15）转成 PyTorch 能用的 one-hot 向量。')

explain(71, '"""把状态 s（0-15）转成 one-hot 向量"""', '为什么不用数字 0-15 直接输入网络？因为格子 3 不是格子 2 的"一半"，格子 7 也不是格子 8 的"0.875 倍"。数字编码隐含了大小关系，但 GridWorld 的状态之间没有大小关系——它们只是不同的格子。')

explain(72, 'x = torch.zeros(16)', '创建一个长度为 16 的全零张量（数组）。比如 x = [0, 0, 0, ..., 0]。')

explain(73, 'x[s] = 1.0', '把第 s 个位置设为 1.0。比如 s=3 → x = [0, 0, 0, 1, 0, 0, ..., 0]。这个向量只在第 3 个位置是 1，其他都是 0——所以叫 one-hot（一个热）。')

explain(74, 'return x.unsqueeze(0)', '在 0 维上加一个"batch"维度。[16] → [1, 16]。因为 PyTorch 网络处理数据时要求一个 batch 一起输入，即使只有一条数据也要加这个维度。')

brk()

# ================================================================
h('第三部分：DQN 算法（78-182 行）', 1)

explain(78, 'def dqn(episodes=5000, lr=0.01, epsilon_start=1.0, epsilon_end=0.01, epsilon_decay=0.998):', '主训练函数。参数和 Q-learning 基本一样：')
bullet('episodes=5000：训练 5000 局', lv=1)
bullet('lr=0.01：学习率（QL 用 0.1，DQN 用更小的 0.01，因为梯度下降比直接 Q 表更新更敏感）', lv=1)
bullet('epsilon_start=1.0：初始探索率，从完全随机开始', lv=1)
bullet('epsilon_end=0.01：最低探索率，保持 1% 的随机性', lv=1)
bullet('epsilon_decay=0.998：探索率衰减系数', lv=1)

brk()

h('3.1 初始化（86-98 行）', 2)

explain(87, 'q_network = DQN()', '创建在线网络（Online Network）。这个网络负责：1) 选动作时算 Q 值；2) 训练时实时更新参数。')

explain(88, 'target_network = DQN()', '创建目标网络（Target Network）。这个网络只负责算 target（训练目标），参数不频繁更新。')
b('这就是 DQN 和 Q-learning 的第二大区别：QL 没有目标网络，DQN 有。', color=RGBColor(0x1F,0x3A,0x5F))

explain(89, 'target_network.load_state_dict(q_network.state_dict())', '把在线网络的参数复制给目标网络。初始时两个网络一模一样。')

explain(90, 'optimizer = optim.Adam(q_network.parameters(), lr=lr)', '创建 Adam 优化器。它负责根据 loss 计算出梯度，然后更新 q_network 的参数。')
p('和 Q-learning 的对比：', sz=10, indent=1)
p('  QL: Q[s][a] += lr × td_error  ← 手动更新一个格子', sz=10, indent=1)
p("  DQN: optimizer.step()            ← 优化器自动更新全部 644 个参数", sz=10, indent=1)

explain(91, "loss_fn = nn.MSELoss()", '创建损失函数：均方误差 MSE。用来衡量当前 Q 值和目标 Q 值的差距。')

brk()

h('3.2 经验回放缓冲区（93-96 行）', 2)

explain(94, 'replay_buffer = []', '经验回放缓冲区。一个空列表，用来存交互数据 (s, a, r, sp, done)。')
b('这是 DQN 和 Q-learning 的第三个区别：QL 不需要缓冲区，DQN 需要。', color=RGBColor(0x1F,0x3A,0x5F))
p('为什么？Q-learning 每次只改 Q 表的一个格子，顺序数据不影响。但神经网络一次更新影响所有参数，如果数据是顺序相关的（相邻几步高度相关），网络会过拟合到最近的经验。', indent=1, sz=10)

explain(95, 'BUFFER_SIZE = 10000', '缓冲区最大容量。存满 10000 条后，新数据进来就挤掉最旧的数据。')

explain(96, 'BATCH_SIZE = 32', '每次训练从缓冲区随机抽 32 条经验。这叫一个 batch。')
p('为什么要凑一批（batch）而不是一条一条学？单条数据噪声大，32 条平均一下，更新方向更准。', indent=1, sz=10)

explain(101, 'for ep in range(1, episodes + 1):', '主训练循环。从第 1 局跑到第 episodes 局（5000 局）。和 Q-learning 一样。')

brk()

h('3.3 ε-贪心选动作（108-114 行）', 2)

explain(108, 'if random.random() < epsilon:', '以 epsilon 概率随机探索。和 Q-learning 完全一样。')

explain(109, 'a = random.randint(0, N_ACTIONS - 1)', '探索：从 4 个动作中随机选一个。')

explain(110, 'else:', '以 1-epsilon 概率利用已有知识。')

explain(111-114, 'with torch.no_grad():', '利用：用在线网络算 Q 值，选最大的。')
p('torch.no_grad() 告诉 PyTorch："我只是查一下值，不需要计算梯度。"如果不加这个，PyTorch 会记录所有计算过程以备求导，浪费内存。', indent=1, sz=10)
p('和 Q-learning 的对比：', sz=10, indent=1)
p("  QL:  a = int(np.argmax(Q[s]))     ← 查 Q 表第 s 行就行", sz=10, indent=1)
p("  DQN: q_values = q_network(x); a = argmax(q_values)  ← 前向传播算一遍", sz=10, indent=1)

explain(116, 's_next, reward, done = step(s, a)', '执行动作，和环境交互。和 Q-learning 完全一样。')

brk()

h('3.4 存储经验 + 训练（119-155 行）', 2)

explain(119, 'replay_buffer.append((s, a, reward, s_next, done))', '把这一步的经验存到缓冲区。存的是五元组 (当前状态, 动作, 奖励, 下个状态, 是否结束)。')

explain(120-121, 'if len(replay_buffer) > BUFFER_SIZE:', '如果缓冲区满了（超过 10000 条），就删掉最旧的一条（pop(0)）。保持缓冲区大小稳定。')

explain(124, 'if len(replay_buffer) >= BATCH_SIZE:', '缓冲区里的经验够 32 条才开始训练。刚开始时 buffer 是空的，先攒经验。')

explain(125, 'batch = random.sample(replay_buffer, BATCH_SIZE)', '核心：经验回放！从缓冲区随机抽 32 条经验，打乱顺序。')
b('这就是解决"数据相关性"的关键。', color=RGBColor(0x1F,0x3A,0x5F))

explain(127-131, 'states = torch.zeros(BATCH_SIZE, 16)', '创建 5 个张量来装这个 batch 的数据：')
bullet('states: [32, 16] → 32 条状态，每条 one-hot 16 维', lv=1)
bullet("next_states: [32, 16] → 32 条下一个状态", lv=1)
bullet('actions: [32] → 32 个动作编号', lv=1)
bullet('rewards: [32] → 32 个奖励值', lv=1)
bullet('dones: [32] → 32 个结束标志', lv=1)

explain(133-138, 'for i, (s, a, r, ns, d) in enumerate(batch):', '遍历 batch 里的 32 条经验，把数据填到刚刚创建的张量里。')

brk()

h('3.5 核心：算 target + loss + 梯度下降（140-155 行）', 1)

p('这里就是 DQN 最核心的 3 步，对照 Q-learning 来看：', bold=True, sz=12, color=RGBColor(0x1F,0x3A,0x5F))

tbl(
    ['步骤', 'Q-learning (查 Q 表)', 'DQN (用神经网络)'],
    [
        ['算 target', 'r + γ·max Q[sp]', 'r + γ·max Q_target(sp)'],
        ['算误差', 'td_error = target - Q[s][a]', 'loss = MSE(Q_net(s,a), target)'],
        ['更新', "Q[s][a] += lr × td_error", 'loss.backward(); optimizer.step()'],
    ]
)
p('公式几乎一模一样，区别只在于：', indent=1, sz=10)
bullet('QL 用 Q 表查一下，DQN 用目标网络跑一次前向传播', lv=1)
bullet('QL 改一个格子，DQN 用梯度下降改所有参数', lv=1)

explain(141, 'with torch.no_grad():', '计算 target 时不需要梯度。目标网络只负责"查值"，不训练。')

explain(142, 'next_q = target_net(next_states)', '把 32 个下个状态送进目标网络，得到 [32, 4] 的 Q 值矩阵。')
b('注意：用的是目标网络（target_net），不是在线网络（q_network）！', color=RGBColor(0xC0,0x39,0x2B))

explain(143, 'max_next_q = torch.max(next_q, dim=1).values', '从每个状态的 4 个 Q 值中取最大值。dim=1 表示沿着动作维度取 max。')
p('结果是 [32] 的张量，每个值代表"下个状态的最好 Q 值"。', indent=1)
p('  Q-learning 中这一步是 np.max(Q[sp])', sz=10, indent=1)

explain(144, 'td_targets = rewards + GAMMA * max_next_q * (1 - dones)', '计算 TD target：r + γ·max Q(sp)。(1 - dones) 表示如果这一步是终点，就不加未来奖励了。')
p('这和 Q-learning 的 target 公式一模一样！', bold=True, sz=10, indent=1, color=RGBColor(0x1F,0x3A,0x5F))

explain(146-148, 'current_q = q_network(states)', '用在线网络算当前状态 s 的 Q 值。[32, 16] → 网络 → [32, 4]。')
p('然后用 gather 从 4 个 Q 值中挑出实际执行的那个动作 a 对应的 Q 值。', indent=1, sz=10)
p('  Q-learning 中这一步是 Q[s][a]——直接从 Q 表读取', sz=10, indent=1)

explain(150, 'loss = loss_fn(current_q_a, td_targets)', '计算均方误差 loss = mean((current_Q - target)²)。衡量"当前预测"和"目标"的差距。')
p('  Q-learning 中这一步是 td_error = target - Q[s][a]——只是一个差值', sz=10, indent=1)

explain(152, 'optimizer.zero_grad()', '清空上一轮计算的梯度。PyTorch 的梯度会累积，不清空的话这轮的梯度会加上轮的。')

explain(153, 'loss.backward()', '反向传播！计算 loss 对每个参数的偏导数（梯度）。')
p('一句话：计算出"每个参数应该往哪个方向调、调多少"才能让 loss 变小。', indent=1, sz=10, color=RGBColor(0x1F,0x3A,0x5F))
p('  这行代码就是 DQN 和 QL 最根本的区别：', sz=10, indent=1)
p("  QL: Q[s][a] += lr * td_error  ← 手动改一个方格", sz=10, indent=1)
p("  DQN: loss.backward() + optimizer.step()  ← 自动调全部 644 个参数", sz=10, indent=1)

explain(155, 'optimizer.step()', '用优化器（Adam）根据梯度更新 q_network 的所有参数。参数朝着"让 loss 变小"的方向微调。')

brk()

h('3.6 目标网络更新（168-169 行）', 2)

explain(168, 'if ep % 200 == 0:', '每 200 局更新一次目标网络。')

explain(169, 'target_network.load_state_dict(q_network.state_dict())', '把在线网络 q_network 的当前参数完整复制给目标网络 target_network。')
p('为什么不是每局更新？如果每局都更新，目标网络跟着在线网络一起变，target 就不稳定了（移动靶）。每 200 局更新一次，target 相对稳定，在线网络有一个"固定的目标"去追赶。', indent=1, sz=10)

brk()

# ================================================================
h('第四部分：结果展示（186-194 行）', 1)

explain(186, 'def save_results(policy, rewards, steps):', '把训练结果保存到 CSV 文件，方便以后画图分析。')

explain(189, "path = os.path.join(RESULTS_DIR, 'week10_dqn_training.csv')", "拼接结果文件路径。结果保存在 F:\\CLAUDE\\research\\ems-platform\\results\\week10_dqn_training.csv。")

explain(190-193, "with open(path, 'w', encoding='utf-8') as f:", "写入 CSV 文件。每行一条记录：局数, 奖励, 步数。可以用 Excel 打开画训练曲线。")

brk()

# ================================================================
h('第五部分：主程序（197-250 行）', 1)

explain(197, "if __name__ == '__main__':", 'Python 标准写法。只有直接运行这个文件时才会执行下面的代码。如果被别的文件 import，不会跑。')

explain(198-205, 'print(...)', '打印 DQN 的介绍信息，标明和 Q-learning 的区别。')

explain(208-214, 'policy, rewards, steps = dqn(...)', '调用 dqn 函数开始训练。参数：5000 局、lr=0.01、ε 从 1.0 衰减到 0.01。')

explain(220-233, '打印策略网格', '把训练得到的策略打印成箭头网格，和 Q-learning 的输出格式完全一样，方便对比。')

explain(238-250, '三个方法对比', '打印值迭代、Q-learning、DQN 三个方法的对比总结。')
p('核心信息：DQN 在这个小问题上没有优势（甚至更慢），但真实问题状态空间巨大，Q 表存不下，只能用神经网络。', bold=True, indent=1, sz=10, color=RGBColor(0x1F,0x3A,0x5F))

brk()

# ================================================================
h('附录：QL vs DQN 代码行对照表', 1)
p('同一个功能，在 Q-learning 和 DQN 中分别怎么写：', sz=10)
brk()

tbl(
    ['功能', 'Q-learning (week10_qlearning)', 'DQN (week10_dqn)'],
    [
        ['存储 Q 值', 'Q = np.zeros((16, 4))\n第 80 行', 'q_network = DQN()\ntarget_network = DQN()\n第 87-88 行'],
        ['状态编码', '直接用数字 s (0-15)', 'one-hot 向量 [16 维]\n第 70-74 行'],
        ['选动作(利用)', 'a = int(np.argmax(Q[s]))\n第 103 行', 'q_values = q_network(x)\na = argmax(q_values)\n第 112-114 行'],
        ['算 target', "r + GAMMA * np.max(Q[s_next])\n第 110 行", 'r + GAMMA * max(Q_target(sp))\n第 142-144 行'],
        ['更新', "Q[s, a] += lr * td_error\n第 112 行", 'loss.backward()\noptimizer.step()\n第 153-155 行'],
        ['目标网络', '不需要 ❌', '每 200 局同步一次\n第 168-169 行'],
        ['经验回放', '不需要 ❌', 'buffer + 随机采样\n第 119-125 行'],
    ]
)

brk()

h('核心记忆点', 1)
p('1. DQN 和 Q-learning 的框架一模一样：ε-贪心 → 执行 → 更新', sz=12)
p('2. 唯一区别：Q 表 → 神经网络。QL 改 1 个格子，DQN 改 644 个参数', sz=12, bold=True)
p('3. 加了两个辅助组件：经验回放（打破数据相关性）+ 目标网络（稳定训练目标）', sz=12)
p('4. 这两组件都是因为"用了神经网络"才需要的——QL 不需要', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
note(f'生成日期: {datetime.date.today().isoformat()}')

# 保存
os.makedirs(OUT_DIR, exist_ok=True)
path = os.path.join(OUT_DIR, 'Week10_DQN_GridWorld_逐行精讲.docx')
doc.save(path)
print(f'OK: {path}')
