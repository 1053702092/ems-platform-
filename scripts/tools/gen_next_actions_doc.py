#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成「下一步行动计划」docx
==========================
汇总当前所有待办事项：LG HG2 数据下载（用户）、方案 A 后续、方案 B、
以及贯穿主线（Week12 简历/投递）。供用户对照执行。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\soc-estimation\doc'

BLUE = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x66, 0x66, 0x66)
DGRAY = RGBColor(0x99, 0x99, 0x99)
RED = RGBColor(0xC0, 0x39, 0x2B)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GREEN = RGBColor(0x27, 0xAE, 0x60)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.line_spacing = 1.35

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.color.rgb = BLUE

def h2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.color.rgb = ORANGE

def h3(t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.color.rgb = ORANGE

def p(t, size=11, bold=False, color=None):
    pa = doc.add_paragraph()
    r = pa.add_run(t)
    r.font.name = 'Microsoft YaHei'; r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return pa

def bullet(t, indent=0, done=None):
    pa = doc.add_paragraph(style='List Bullet')
    pa.paragraph_format.left_indent = Cm(1.5 + indent * 0.8)
    if done is True:
        r = pa.add_run('✅ '); r.font.color.rgb = GREEN
    elif done is False:
        r = pa.add_run('☐ '); r.font.color.rgb = ORANGE
    else:
        r = pa.add_run('▶ '); r.font.color.rgb = ORANGE
    r2 = pa.add_run(t)
    r2.font.name = 'Microsoft YaHei'; r2.font.size = Pt(11)

def code(t):
    for line in t.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Cm(1)
        pa.paragraph_format.space_after = Pt(1)
        r = pa.add_run(line)
        r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def tbl(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers))
    tb.style = 'Table Grid'
    for i, hh in enumerate(headers):
        r = tb.rows[0].cells[i].paragraphs[0].add_run(hh)
        r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'
    for rd in rows:
        row = tb.add_row()
        for c, txt in enumerate(rd):
            r = row.cells[c].paragraphs[0].add_run(txt)
            r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'

# ── 封面 ──
for _ in range(3):
    doc.add_paragraph('')
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('下一步行动计划\n'); r.font.size = Pt(24); r.bold = True; r.font.color.rgb = BLUE
r = t.add_run('EMS-PLATFORM 简历加分项 · 待办总览')
r.font.size = Pt(13); r.font.color.rgb = GRAY
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(f'生成日期：{datetime.date.today().isoformat()} · 目标：8 月底完成技术积累，9 月全力投递')
r.font.size = Pt(10); r.font.color.rgb = DGRAY
doc.add_page_break()

# ── 一、总览 ──
h1('一、待办总览')
tbl(['#', '任务', '负责人', '优先级', '预计耗时'], [
    ['1', '下载 LG HG2 数据集（Mendeley 2GB）', '你（另一台电脑）', 'P0 阻塞项', '下载+解压 ~1h'],
    ['2', '方案 A 后续：HPPC 辨识 → 真实数据对比 → 报告', '我', 'P0', '1 周（等数据）'],
    ['3', '方案 B：PPO 跑 OPF-Gym benchmark', '我', 'P1', '2-3 天'],
    ['4', 'Week12：简历打磨 + 投递准备', '你', 'P0', '贯穿'],
])

p('依赖关系：任务 1（下载数据）阻塞任务 2（方案 A 真实数据部分）。数据到位前，我已完成合成数据流程，数据一到即可直接替换。', bold=True, color=RED)

# ── 二、任务 1：下载数据（你来） ──
h1('二、任务 1：下载 LG HG2 数据集（你来，阻塞项）')
p('详见已生成的《SOC估计器_LG_HG2_数据集下载指南.docx》，要点：', bold=True)
bullet('打开数据集页：https://data.mendeley.com/datasets/cp3473x7xv/3', done=False)
bullet('注册/登录 Mendeley（免费邮箱即可），下载整个 zip（约 2 GB）', done=False)
bullet('解压后放到：ems-platform/datasets/LG_HG2/（按温度分文件夹，不要改结构）', done=False)
bullet('核对：208 个 CSV、约 490 万行', done=False)
bullet('下载完成后告诉我"数据放好了"，我继续任务 2', done=False)
p('提示：如果本机之前下载过（你跑过 LG HG2 验证 RMSE=0.0465），先检查 datasets/ 是否残留，能省一次下载。', size=10)

# ── 三、任务 2：方案 A 后续（我来，等数据） ──
h1('三、任务 2：方案 A 后续（我来做，等数据到位）')
p('已完成部分（合成数据流程已跑通）：', bold=True)
bullet('pybamm 26.7.1 安装 ✅', done=True)
bullet('run_pybamm_synth.py：DFN 合成数据 + EKF 对比 ✅（RMSE=0.113）', done=True)
bullet('发现 3 个工程问题：电流符号 / OCV 标定 / 边界 nan ✅', done=True)
p('数据到位后的待办：', bold=True)
bullet('解析 LG HG2 CSV → 提取电流/电压/温度序列（25°C 优先）', done=False)
bullet('用 HPPC 文件辨识 R0/R1/C1（STATUS.md 里的"下一步"）', done=False)
bullet('真实数据跑现有 EKF/AEKF 估计器', done=False)
bullet('用 PyBaMM 的 DFN 物理模型（NMC 参数接近 LG HG2）建立对比', done=False)
bullet('低温（-10°C/0°C）验证模型失配：等效电路 vs 物理模型 RMSE 对比', done=False)
bullet('生成对比报告 docx：方法 / 结果表 / 面试叙事要点', done=False)

# ── 四、任务 3：方案 B ──
h1('四、任务 3：方案 B —— PPO 跑 OPF-Gym benchmark（我来做）')
p('背景：让 Week11 的 PPO 在行业标准 RL 环境上验证，增加项目公信力。', bold=True)
p('待办：', bold=True)
bullet('pip install gymnasium opfgym（opfgym 基于 pandapower，依赖较轻）', done=False)
bullet('跑通 opfgym 自带示例（EcoDispatch / VoltageControl）', done=False)
bullet('把 Week11 PPO 适配成 gymnasium API（reset/step 已是标准接口，改动很小）', done=False)
bullet('在 benchmark 上训练 + 对比官方 baseline', done=False)
bullet('产出：训练曲线 + 结果表 + 简历叙事', done=False)
p('依赖：无需外部数据，随时可做。可在等数据下载时并行启动。', size=10)

# ── 五、任务 4：Week12 主线 ──
h1('五、任务 4：Week12 主线（你来，贯穿）')
p('技术积累 8 月底截止，9 月全力投递。主线事项：', bold=True)
bullet('EMS-PLATFORM 项目叙事定稿（已生成 docs/interview/EMS项目叙事_简历面试版.docx）', done=False)
bullet('中石化/中海油/发电集团笔试准备（综合知识 60% + 专业 40%）', done=False)
bullet('LeetCode Easy 刷题（DP/数组/字符串 30 道）', done=False)
bullet('八股文积累（控制理论 / EMS 算法 / 燃料电池基础）', done=False)
bullet('目标企业清单更新（20-30 家）', done=False)
bullet('第一批投递启动（8 月下旬）', done=False)

# ── 六、执行顺序建议 ──
h1('六、执行顺序建议')
p('现在（数据下载前）：', bold=True)
bullet('你把下载指南拷走，找另一台电脑下载数据（后台进行）')
bullet('我同时启动任务 3（方案 B），不浪费等待时间')
p('数据到位后：', bold=True)
bullet('我接着做任务 2（方案 A 真实数据 + 物理模型对比）')
bullet('方案 A 出报告后，方案 B 也基本完成')
p('最终目标（8 月底前）：简历新增两项成果 —— ① SOC 估计器升级到电化学物理模型 ② PPO 通过行业标准 benchmark。', bold=True, color=GREEN)

path = os.path.join(OUT_DIR, '下一步行动计划_简历加分项.docx')
doc.save(path)
print('OK: ' + path)
