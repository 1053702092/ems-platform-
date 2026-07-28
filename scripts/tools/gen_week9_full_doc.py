#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Week9 完整代码逐行讲解文档 (.docx)
=========================================
涵盖 Part 1-8 全部代码，每部分包含：
  - 篇章概述（用途与学习目标）
  - 逐段代码详解（核心逻辑 + 设计意图）
  - 输出结果解读

输出: docs/notes/Week9_完整代码逐行精讲.docx
"""

import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

# ── 样式工具 ──

def add_code_block(doc, code_text, label=None):
    """添加代码块（等宽字体 + 浅灰底色）"""
    if label:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # 背景色通过底纹实现
        shading = run._element.get_or_add_rPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F5F5F5',
            qn('w:val'): 'clear',
        })
        shading.append(shd)


def add_output_block(doc, output_text, label='▶ 输出示例'):
    """添加输出示例块"""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    for line in output_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)


def add_hint(doc, text, label='💡 要点提示'):
    """添加重点提示"""
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)


def add_code_ref(doc, line_num, desc):
    """引用代码行号"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'L{line_num}  ')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x99)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)


def make_section_heading(doc, title, level=1):
    """创建带分隔线的标题"""
    h = doc.add_heading(title, level=level)
    return h


def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.bold = bold
        if bold:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml import OxmlElement
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')
            shading_elm.set(qn('w:val'), 'clear')
            cell._element.get_or_add_tcPr().append(shading_elm)


# ═══════════════════════════════════════════════════════════════
# 主程序
# ═══════════════════════════════════════════════════════════════

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph('')

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('Week 9 — PyTorch + RL 基础 完整通关')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('代码逐行精讲 · 原理剖析 · 实践指南')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph('')
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run(f'生成日期：{datetime.date.today().isoformat()}\n'
                       f'文件：scripts/week9_complete.py\n'
                       f'涉及主题：Tensor · Autograd · MLP · MDP · Bellman · 策略迭代 · 值迭代')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '1. 文件总览 — 架构与设计思路',
    '2. Part 1 — Tensor 基础',
    '3. Part 2 — Autograd 自动求导',
    '4. Part 3 — nn.Module + MLP',
    '5. Part 4 — MLP 功率预测（FC 功率预测）',
    '6. Part 5 — MDP 五元组：GridWorld',
    '7. Part 6 — Bellman 方程',
    '8. Part 7 — 策略迭代',
    '9. Part 8 — 值迭代',
    '10. Extra — 收敛过程可视化',
    '11. 主程序调度 — Part 编排与依赖管理',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. 文件总览
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '1. 文件总览 — 架构与设计思路', level=1)

doc.add_paragraph(
    'week9_complete.py 是一个一体化的学习脚本，涵盖 PyTorch 深度学习基础（Part 1-4）'
    '和强化学习基础（Part 5-8）两大模块。全书共 730+ 行，按 Part 编号组织为独立函数，'
    '通过主程序调度器（main）组合执行。'
)

make_section_heading(doc, '1.1 前后依赖关系', level=2)

dep_table = doc.add_table(rows=1, cols=4)
dep_table.style = 'Light Grid Accent 1'
hdr = dep_table.rows[0].cells
for i, txt in enumerate(['Part', '主题', '前置依赖', '输出']):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True

deps = [
    ('1', 'Tensor 基础', '无', 'device 对象'),
    ('2', 'Autograd', '无', '梯度概念'),
    ('3', 'nn.Module+MLP', '无', '模型实例'),
    ('4', 'MLP 功率预测', 'Part 3 概念', '训练好的模型 + 预测图'),
    ('5', 'MDP GridWorld', '无', 'mdp 字典（R, P, gamma）'),
    ('6', 'Bellman 方程', 'Part 5', 'V_pi, V_opt + heatmap'),
    ('7', '策略迭代', 'Part 5', '最优策略 + 策略图'),
    ('8', '值迭代', 'Part 5', '最优策略 + V*'),
]
for r in deps:
    add_table_row(dep_table, list(r))

doc.add_paragraph('')

make_section_heading(doc, '1.2 前置准备（L22-50）', level=2)
doc.add_paragraph(
    '文件开头完成了三项关键准备工作：'
)

add_code_ref(doc, 24, '环境编码修复 — 将 stdout/stderr 强制设为 UTF-8。')
doc.add_paragraph(
    'Windows 系统终端默认编码为 GBK（CP936），打印中文时若包含 Unicode 字符会报错。'
    '这段代码检测到 GBK 编码后主动 reconfigure，确保中英混排的 print 输出不会崩溃。'
    '这是跨平台脚本工程的最佳实践。'
)

add_code_ref(doc, 30-33, 'Matplotlib 配置 — 使用 "Agg" 后端（无头模式），适用于服务器/CI 环境。')
doc.add_paragraph(
    'Agg 是 matplotlib 的非交互式后端，不依赖 GUI 显示。'
    '搭配 plt.savefig() 将绘图直接输出到文件，是本项目（无桌面环境的远程 GPU 服务器）的正确选择。'
)

add_code_ref(doc, 37-45, '命令行参数解析 — 支持 --part N 单独运行指定部分。')
doc.add_paragraph(
    '通过 argparse 接受整数参数 0~8，默认 0 表示运行全部。'
    'should_run(part_num) 函数封装了判断逻辑，各 Part 函数据此决定是否执行：'
    '当 args.part == 0 或 args.part == part_num 时返回 True。'
    '这种模式允许用户反复调试某个特定部分，而不必每次都跑完整流程。'
)

add_code_ref(doc, 47-50, '中文字体 Fallback 链 — 优先 微软雅黑 → SimHei → SimSun。')
doc.add_paragraph(
    '如果系统中没有 微软雅黑，自动回退到黑体（SimHei）再回退到宋体（SimSun）。'
    'axes.unicode_minus = False 修复负号显示为方块的问题（将减号替换为 Unicode 标准负号）。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. Part 1 — Tensor 基础
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '2. Part 1 — Tensor 基础', level=1)

doc.add_paragraph(
    'Tensor 是 PyTorch 的核心数据结构，可理解为「GPU 加速的多维数组」。'
    'Part 1 演示了 Tensor 的创建、属性查询、设备管理、索引切片和广播运算。'
    '这是后续所有 Part 的基石。'
)

make_section_heading(doc, '2.1 核心代码详解', level=2)

add_code_block(doc, '''import torch''', 'L58 导入 PyTorch 核心库')

doc.add_paragraph(
    'import torch 导入了 PyTorch 的全部核心 API，包括 Tensor 创建、数学运算、'
    '自动求导引擎等。在本 Part 中，我们只用到了 tensor 创建函数。'
)

add_code_block(doc, '''t1 = torch.tensor([[1, 2], [3, 4]])''', 'L61 从 Python 列表创建 Tensor')

doc.add_paragraph(
    'torch.tensor() 是最通用的 Tensor 创建方式，接受嵌套列表、numpy 数组等作为输入。'
    '此处创建了一个 2×2 的整数 Tensor，PyTorch 会自动推导 dtype=torch.int64。'
    '输出时附带 dtype 和 shape 信息，这是 Tensor 的三大关键属性之一。'
)

add_code_block(doc, '''a = np.array([1.0, 2.0, 3.0])
t2 = torch.from_numpy(a)''', 'L65-66 从 numpy 数组创建 Tensor')

doc.add_paragraph(
    'torch.from_numpy() 创建出的 Tensor 与源 numpy 数组共享内存——修改任一方都会影响另一方。'
    '这是高性能数据交换的关键设计，避免了不必要的内存拷贝。'
    '注意：返回的 Tensor 默认为 float64（与 numpy 一致），而 PyTorch 模型通常使用 float32。'
)

add_code_block(doc, '''zeros = torch.zeros(2, 3)
ones = torch.ones(2, 3)
rand = torch.randn(3, 3)  # 标准正态''', 'L70-72 特殊张量')

doc.add_paragraph(
    '三种常用初始化方式：全零（zeros）、全一（ones）、标准正态分布（randn, 均值0方差1）。'
    'randn 常用于权重初始化，zeros/ones 常用于偏置或 Mask 初始化。'
    '三者都接受 shape 参数（可变参数或元组）。'
)

add_code_block(doc, '''device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
t_gpu = torch.tensor([1, 2, 3], device=device)''', 'L76-80 设备管理')

doc.add_paragraph(
    '这段代码展示了 PyTorch 的设备抽象：tensor.to(device) 或直接在目标设备上创建 Tensor。'
    '如果 CUDA 可用，device 为 cuda:0，否则为 cpu。'
    '这种抽象让同一份代码可以无缝在 CPU 和 GPU 之间切换，是 PyTorch 的核心优势。'
)

add_code_block(doc, '''t = torch.randn(4, 5)
print(f"t[0]: {t[0].shape}")        # → torch.Size([5])
print(f"t[:, :3]: {t[:, :3].shape}") # → torch.Size([4, 3])
print(f"t.view(-1): {t.view(-1).shape}")  # → torch.Size([20])  展平
print(f"t.reshape(2, 10): {t.reshape(2, 10).shape}")  # → torch.Size([2, 10]) 重塑''', 'L83-86 索引、切片、形状操作')

doc.add_paragraph(
    '索引切片与 numpy 完全一致。view() 和 reshape() 的区别：view() 要求数据在内存中连续'
    '（contiguous），而 reshape() 在数据不连续时会自动拷贝。'
    'view(-1) 中的 -1 表示自动推断该维度的大小。'
)

add_code_block(doc, '''a = torch.tensor([[1], [2], [3]])  # (3,1)
b = torch.tensor([10, 20, 30])     # (3,)
print(f"broadcast add:\\n{a + b}")''', 'L89-91 广播运算')

doc.add_paragraph(
    '形状 (3,1) 与 (3,) 相加时，广播机制将 a 沿列方向扩展为 (3,3)，将 b 沿行方向扩展为 (3,3)，'
    '然后逐元素相加。结果：[[11,21,31],[12,22,32],[13,23,31]]^T。'
    '广播规则：从尾部维度开始对齐，维度为 1 或缺失时自动扩展。'
)

make_section_heading(doc, '2.2 返回值', level=2)
doc.add_paragraph('函数返回 device 对象，供后续 Part 使用。从设计上，Part 1 为 Part 4 的 GPU 训练预热。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. Part 2 — Autograd 自动求导
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '3. Part 2 — Autograd 自动求导', level=1)

doc.add_paragraph(
    'Autograd 是 PyTorch 自动微分引擎，是实现神经网络反向传播（backpropagation）的基础。'
    '本 Part 演示了基础用法、链式法则验证、梯度清除和推理模式。'
)

make_section_heading(doc, '3.1 核心概念：计算图', level=2)
doc.add_paragraph(
    'PyTorch 在前向传播过程中动态构建计算图（DAG，有向无环图）。每个 Tensor 记录了：'
)
doc.add_paragraph('• data — 实际的数值', style='List Bullet')
doc.add_paragraph('• grad — 累积的梯度', style='List Bullet')
doc.add_paragraph('• grad_fn — 创建该 Tensor 的运算（如 AddBackward、MulBackward）', style='List Bullet')
doc.add_paragraph('• requires_grad — 是否需要梯度', style='List Bullet')

make_section_heading(doc, '3.2 核心代码详解', level=2)

add_code_block(doc, '''x = torch.tensor([2.0, 3.0], requires_grad=True)
y = x ** 2 + 3 * x
loss = y.sum()
loss.backward()  # 反向传播
print(f"gradient dy/dx: {x.grad}")  # → tensor([7., 9.])''', 'L104-111 基本自动求导')

doc.add_paragraph(
    '设置 requires_grad=True 后，所有基于 x 的运算都会被追踪。'
    '计算过程：y = x² + 3x；loss = y.sum() = y₁ + y₂。'
    '对 loss.backward() 调用后，PyTorch 从 loss 开始逆向遍历计算图，'
    '将梯度累积到每个 requires_grad=True 的叶子节点的 .grad 属性中。'
)
doc.add_paragraph(
    '理论验证：dy/dx = 2x + 3，在 x=[2,3] 处得到 [7,9]，与输出一致 ✓'
)

add_code_block(doc, '''x.grad.zero_()''', 'L115 清除梯度')

doc.add_paragraph(
    '梯度默认是累积（accumulate）的——多次 backward() 的梯度会加在一起。'
    '因此每次更新参数前必须手动归零。zero_() 是 in-place 操作（下划线后缀表示原地修改）。'
    '若不清零，第二次 backward 后 x.grad 会变成 [7+7, 9+9] = [14, 18]。'
)

add_code_block(doc, '''a = torch.tensor(2.0, requires_grad=True)
b = torch.tensor(3.0, requires_grad=True)
z = (a ** 2) * torch.sin(b)
z.backward()
print(f"dz/da = {a.grad:.4f}")  # 2*a*sin(b) = 4*0.1411 = 0.5645
print(f"dz/db = {b.grad:.4f}")  # a^2*cos(b) = 4*(-0.99) = -3.9600''', 'L118-123 链式法则')

doc.add_paragraph(
    '复合函数的梯度通过链式法则自动计算。z = a² · sin(b)：'
)
doc.add_paragraph('∂z/∂a = 2a · sin(b) = 2×2×sin(3) ≈ 4×0.1411 = 0.5645', style='List Bullet')
doc.add_paragraph('∂z/∂b = a² · cos(b) = 4×cos(3) ≈ 4×(-0.99) = -3.9600', style='List Bullet')
doc.add_paragraph(
    'Autograd 自动应用链式法则，无需手动推导偏导数。这正是深度学习框架的核心价值。'
)

add_code_block(doc, '''with torch.no_grad():
    y_eval = x ** 2 + 3 * x
    print(f"no_grad: {y_eval}")''', 'L126-128 推理模式')

doc.add_paragraph(
    'torch.no_grad() 上下文管理器关闭了计算图构建。在此模式下，所有运算结果不再跟踪梯度，'
    '大幅降低内存和计算开销。在模型评估（evaluation）阶段，测试集的前向传播应当在 no_grad 下进行。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. Part 3 — nn.Module + MLP
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '4. Part 3 — nn.Module + MLP', level=1)

doc.add_paragraph(
    'nn.Module 是 PyTorch 中所有神经网络模型的基类。本 Part 通过构建一个简单的'
    '两层 MLP（多层感知机），演示了模型定义、前向传播和参数管理。'
)

make_section_heading(doc, '4.1 核心代码详解', level=2)

add_code_block(doc, '''class MLP(nn.Module):
    """两层 MLP: input -> hidden (ReLU) -> output"""
    def __init__(self, input_dim, hidden_dim, output_dim):
        super().__init__()
        self.fc1 = nn.Linear(input_dim, hidden_dim)
        self.fc2 = nn.Linear(hidden_dim, output_dim)

    def forward(self, x):
        x = F.relu(self.fc1(x))
        x = self.fc2(x)
        return x''', 'L140-150 两层 MLP 定义')

doc.add_paragraph('逐行解读：')

add_code_ref(doc, 140, 'class MLP(nn.Module) — 继承自 nn.Module，自动获得参数管理、train/eval 切换等功能。')
add_code_ref(doc, 142, '__init__() — 构造函数，声明所有可学习的层。这里声明了两个全连接层（Linear）。')
add_code_ref(doc, 143, 'super().__init__() — 必须调用父类构造器，初始化 Module 内部注册机制。')
add_code_ref(doc, 144, 'nn.Linear(input_dim, hidden_dim) — 全连接层：y = xW^T + b。内部注册了权重参数 W 和偏置 b。')
add_code_ref(doc, 145, 'nn.Linear(hidden_dim, output_dim) — 输出层，无激活函数（回归任务直接输出数值）。')
add_code_ref(doc, 148, 'F.relu(self.fc1(x)) — 隐藏层使用 ReLU 激活函数，引入非线性变换。ReLU: f(x) = max(0, x)。')

doc.add_paragraph('')
doc.add_paragraph(
    'nn.Module 的设计精髓：在 __init__ 中声明层，在 forward 中定义计算逻辑——'
    '正向传播时 PyTorch 自动追踪所有操作来构建计算图，反向传播时自动求导。'
)

add_code_block(doc, '''model = MLP(input_dim=4, hidden_dim=32, output_dim=1)
print(f"Model:\\n{model}")
print(f"Parameters: {sum(p.numel() for p in model.parameters())}")''', 'L153-155 实例化与参数统计')

doc.add_paragraph(
    'model.parameters() 递归收集所有子模块的可学习参数。'
    '此处 MLP 有 4×32 + 32 + 32×1 + 1 = 128 + 32 + 32 + 1 = 193 个参数。'
    '打印 model 会递归显示所有子模块的层次结构。'
)

add_code_block(doc, '''x = torch.randn(10, 4)  # batch=10, features=4
y = model(x)
print(f"Input: {x.shape}, Output: {y.shape}")''', 'L158-160 前向传播')

doc.add_paragraph(
    'model(x) 实际上调用了 model.__call__(x)，__call__ 在调用 forward() 前后'
    '自动处理了钩子（hook）注册和 train/eval 模式逻辑。'
    '输入形状 [10, 4] → 经过 fc1(4→32) → ReLU → fc2(32→1) → 输出形状 [10, 1]。'
)

make_section_heading(doc, '4.2 概念总结', level=2)
doc.add_paragraph('nn.Module 封装了四个关键能力：')
doc.add_paragraph('1. 参数注册与管理（parameters(), named_parameters()）', style='List Bullet')
doc.add_paragraph('2. 设备转移（.to(device)）', style='List Bullet')
doc.add_paragraph('3. 训练/评估模式切换（.train(), .eval()）', style='List Bullet')
doc.add_paragraph('4. 状态保存与加载（state_dict(), load_state_dict()）', style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. Part 4 — MLP 功率预测
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '5. Part 4 — MLP 功率预测（FC 功率预测）', level=1)

doc.add_paragraph(
    '将 MLP 应用于实际的燃料电池（FC）功率预测任务。这里使用模拟数据，'
    '完整地走了一遍「数据生成 → 时序样本构建 → 模型定义 → 训练循环 → 评估 → 可视化」的机器学习工作流。'
)

make_section_heading(doc, '5.1 数据生成（L178-181）', level=2)

add_code_block(doc, '''np.random.seed(42)
t = np.linspace(0, 100, 1000)
power = 30 + 15 * np.sin(0.1 * t) + 5 * np.sin(0.5 * t) + np.random.randn(1000) * 2
power = np.clip(power, 10, 80)''', '')

doc.add_paragraph(
    '使用 sin 波叠加加噪声来模拟 FC 功率信号：'
)
doc.add_paragraph('• 30 kW 基准功率（直流偏置）', style='List Bullet')
doc.add_paragraph('• 15 × sin(0.1t) — 低频大幅波动（模拟负载变化趋势）', style='List Bullet')
doc.add_paragraph('• 5 × sin(0.5t) — 高频小幅波动（模拟噪声扰动）', style='List Bullet')
doc.add_paragraph('• randn × 2 — 高斯随机噪声（模拟传感器测量噪声）', style='List Bullet')
doc.add_paragraph('• np.clip(10, 80) — 功率限幅，模拟实际的 FC 功率上下限', style='List Bullet')

make_section_heading(doc, '5.2 时序样本构建（L184-201）', level=2)

add_code_block(doc, '''def create_sequences(data, seq_len=10):
    X, y = [], []
    for i in range(len(data) - seq_len):
        X.append(data[i:i + seq_len])
        y.append(data[i + seq_len])
    return np.array(X), np.array(y)''', '')

doc.add_paragraph(
    '滑动窗口法构建监督学习样本：用过去 10 个时间步的功率值预测下一个时间步的功率值。'
    '如果原始数据长度为 N，window_size = seq_len，则生成 N - seq_len 个样本。'
    '这种序列 → 监督转换是时序预测的标准做法。'
)
doc.add_paragraph(
    '数据划分：前 80% 训练，后 20% 测试（时序数据不能随机打乱，否则会造成未来信息泄漏）。'
    '最后用 unsqueeze(-1) 增加特征维度，使形状从 [n_samples, seq_len] → [n_samples, seq_len, 1]。'
)

make_section_heading(doc, '5.3 模型定义（L204-218）', level=2)

add_code_block(doc, '''class PowerPredictor(nn.Module):
    def __init__(self, seq_len, hidden=64):
        super().__init__()
        self.fc1 = nn.Linear(seq_len, hidden)
        self.fc2 = nn.Linear(hidden, hidden)
        self.fc3 = nn.Linear(hidden, 1)
        self.dropout = nn.Dropout(0.1)

    def forward(self, x):
        x = x.view(x.size(0), -1)   # 展平: [B, seq_len, 1] → [B, seq_len]
        x = F.relu(self.fc1(x))
        x = self.dropout(x)
        x = F.relu(self.fc2(x))
        x = self.fc3(x)              # 输出层无激活（回归）
        return x''', '')

doc.add_paragraph('相较于 Part 3 的简单 MLP，这里做了三处重要改进：')
doc.add_paragraph(
    '1. 三全连接层：64→64→1，增加模型容量（capacity），学习更复杂的映射关系',
    style='List Bullet'
)
doc.add_paragraph(
    '2. Dropout(0.1)：随机丢弃 10% 的神经元输出，防止过拟合。推理时 Dropout 自动关闭',
    style='List Bullet'
)
doc.add_paragraph(
    '3. view() 展平：将 [B, 10, 1] 展平为 [B, 10]（因为 Linear 层需要 2D 输入）',
    style='List Bullet'
)

make_section_heading(doc, '5.4 训练循环（L225-238）', level=2)

add_code_block(doc, '''optimizer = torch.optim.Adam(model.parameters(), lr=0.001)
loss_fn = nn.MSELoss()

for epoch in range(n_epochs):
    model.train()
    optimizer.zero_grad()
    y_pred = model(X_train_t)
    loss = loss_fn(y_pred, y_train_t)
    loss.backward()
    optimizer.step()
    ...''', '')

doc.add_paragraph('训练循环的标准 5 步模式：')
doc.add_paragraph('① model.train() — 切换到训练模式（启用 Dropout、BatchNorm 等）', style='List Bullet')
doc.add_paragraph('② optimizer.zero_grad() — 清空上一步的梯度（避免累积）', style='List Bullet')
doc.add_paragraph('③ loss = loss_fn(model(x), y) — 前向传播，计算预测值与真值之间的损失', style='List Bullet')
doc.add_paragraph('④ loss.backward() — 反向传播，计算所有参数的梯度', style='List Bullet')
doc.add_paragraph('⑤ optimizer.step() — 沿梯度方向更新参数（θ = θ - lr × grad）', style='List Bullet')

doc.add_paragraph(
    '优化器选择 Adam，学习率 0.001。Adam 结合了 Momentum（动量）和 RMSProp（自适应学习率）'
    '的优点，是深度学习中默认的首选优化器。',
)

make_section_heading(doc, '5.5 评估与可视化（L242-267）', level=2)
doc.add_paragraph('测试集上计算两个指标：')
doc.add_paragraph('• MAE（Mean Absolute Error）= (1/N) Σ|y_pred - y_true|，单位 kW', style='List Bullet')
doc.add_paragraph('• RMSE（Root Mean Square Error）= sqrt((1/N) Σ(y_pred - y_true)²)，单位 kW', style='List Bullet')
doc.add_paragraph(
    'MAE 对所有误差同等看待，RMSE 对大误差敏感（平方放大了大误差的影响）。'
    '如果 RMSE >> MAE，说明存在个别预测偏差很大的时间点。'
)
doc.add_paragraph(
    '可视化部分输出两张子图：训练损失下降曲线（判断收敛）和前 200 个测试样本的'
    '预测 vs 真值对比（直观感受拟合质量）。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. Part 5 — MDP 五元组：GridWorld
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '6. Part 5 — MDP 五元组：GridWorld', level=1)

doc.add_paragraph(
    'Part 5 是 RL 基础篇的起点。用 4×4 网格世界（GridWorld）来具象化 MDP'
    '（马尔可夫决策过程）的五元组 ⟨S, A, P, R, γ⟩。'
)

make_section_heading(doc, '6.1 MDP 五元组概念', level=2)

mdp_table = doc.add_table(rows=1, cols=3)
mdp_table.style = 'Light Grid Accent 1'
hdr = mdp_table.rows[0].cells
for i, txt in enumerate(['符号', '名称', 'GridWorld 实例']):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True
mdp_data = [
    ('S', '状态集合', '4×4 = 16 个格子 (0~15)'),
    ('A', '动作集合', '{↑, ↓, ←, →} 共 4 个动作'),
    ('P', '状态转移概率', 'P(s\'|s,a) 执行动作 a 从 s 到 s\' 的概率'),
    ('R', '立即奖励', '到达终点 +1，陷阱 -1，其余 0'),
    ('γ', '折扣因子', 'γ = 0.9，平衡即时与未来奖励'),
]
for r in mdp_data:
    add_table_row(mdp_table, list(r))

doc.add_paragraph('')

make_section_heading(doc, '6.2 核心代码详解', level=2)

add_code_block(doc, '''SIZE = 4
n_states = SIZE * SIZE  # 16
actions = ['↑', '↓', '←', '→']
n_actions = len(actions)  # 4
gamma = 0.9

GOAL = (3, 3)
TRAP = (1, 1)
GOAL_IDX = GOAL[0] * SIZE + GOAL[1]  # = 15
TRAP_IDX = TRAP[0] * SIZE + TRAP[1]  # = 5''', 'L282-291 状态空间定义')

doc.add_paragraph(
    '4×4 网格共 16 个状态，从 0 到 15 编号。编号公式：s = r × SIZE + c。'
    '终点 (3,3) → 3×4+3=15，陷阱 (1,1) → 1×4+1=5。'
    '折扣因子 γ=0.9 意味着未来第 k 步的奖励按 0.9^k 折现到当前——'
    '1 步后的 1 分现在值 0.9 分，10 步后的 1 分现在仅值 0.35 分。'
)

add_code_block(doc, '''action_delta = {
    '↑': (-1, 0), '↓': (1, 0),
    '←': (0, -1), '→': (0, 1),
}''', 'L293-296 动作-位移映射')

doc.add_paragraph(
    '用字典将动作名称映射为 (行偏移, 列偏移)。'
    '"上" → (-1,0)：行减一，列不变；"左" → (0,-1)：行不变，列减一。'
    '这种映射是实现状态转移的基础。'
)

add_code_block(doc, '''def pos_to_idx(r, c):
    return r * SIZE + c

def is_valid(r, c):
    return 0 <= r < SIZE and 0 <= c < SIZE''', 'L298-302 坐标工具函数')

doc.add_paragraph(
    'pos_to_idx：将 (行,列) 二维坐标映射为一维状态编号 s。'
    '这是 GridWorld 中最常用的函数——几乎所有的状态转移都依赖它。'
)
doc.add_paragraph(
    'is_valid：检查坐标是否在 0~3 的合法范围内。'
    '越界时会被「撞墙反弹」——智能体尝试走出边界时停留在原地。'
)

# ── P, R 初始化 ──
add_code_block(doc, '''R = {s: {a: 0.0 for a in range(n_actions)} for s in range(n_states)}
P = {s: {a: {} for a in range(n_actions)} for s in range(n_states)}''', 'L304-305 初始化 R 和 P 数据结构')

doc.add_paragraph(
    '使用嵌套字典存储 MDP 五元组：'
)
doc.add_paragraph('• R[s][a] — 在状态 s 执行动作 a 得到的立即奖励（float）', style='List Bullet')
doc.add_paragraph('• P[s][a] — 转移概率字典 {s\': prob}，记录了所有可能转移到的状态及其概率', style='List Bullet')
doc.add_paragraph(
    '用字典而不是数组来存储 P，因为每个 (s,a) 对通常只转移到少量目标状态——'
    '稀疏存储节省内存。数据结构是：dict[state][action] → dict[next_state: probability]。'
)

# ── 循环遍历 ──
add_code_block(doc, '''for r, c in itertools.product(range(SIZE), range(SIZE)):
    s = pos_to_idx(r, c)
    if s == GOAL_IDX or s == TRAP_IDX:
        for a in range(n_actions):
            R[s][a] = 1.0 if s == GOAL_IDX else -1.0
            P[s][a][s] = 1.0   # 吸收态：到达后永远停留在该状态
        continue''', 'L307-313 遍历所有格子·终点/陷阱特殊处理')

doc.add_paragraph(
    '用 itertools.product 生成所有 (r,c) 组合，遍历 4×4=16 个格子。'
    '终点和陷阱作为「吸收态（absorbing state）」处理：无论执行什么动作，都 100% 停留在自身，'
    '不会转移到其他格子。这模拟了「游戏结束」的语义。'
)

# ── 普通格子 ──
add_code_block(doc, '''for a_idx, (action_name, (dr, dc)) in enumerate(action_delta.items()):
    nr, nc = r + dr, c + dc
    if not is_valid(nr, nc):
        nr, nc = r, c           # 撞墙反弹
    target_s = pos_to_idx(nr, nc)
    R[s][a_idx] = 0.0
    P[s][a_idx][target_s] = P[s][a_idx].get(target_s, 0) + 0.8''', 'L314-320 普通格子·80% 目标方向')

doc.add_paragraph(
    '对于每个非终止状态 s，遍历所有 4 个动作。针对每个动作：'
)
doc.add_paragraph('① 计算意图到达的位置 (nr, nc) = (r + dr, c + dc)', style='List Bullet')
doc.add_paragraph('② 如果越界，重置为原地（撞墙反弹）', style='List Bullet')
doc.add_paragraph('③ 以 80% 的概率（0.8）按意图方向移动', style='List Bullet')
doc.add_paragraph('④ 立即奖励为 0（只有终点/陷阱有非零奖励）', style='List Bullet')

# ── 你问的核心部分 ──
add_code_block(doc, '''for other_dr, other_dc in action_delta.values():
    if (other_dr, other_dc) == (dr, dc):
        continue
    nr2, nc2 = r + other_dr, c + other_dc
    if not is_valid(nr2, nc2):
        nr2, nc2 = r, c
    other_s = pos_to_idx(nr2, nc2)
    P[s][a_idx][other_s] = P[s][a_idx].get(other_s, 0) + 0.2 / 3''', 'L321-328 普通格子·20% 滑移概率')

doc.add_paragraph(
    '这是随机转移的核心——模拟现实中的「打滑」或「风扰」。'
)
doc.add_paragraph('设计逻辑：')
doc.add_paragraph(
    '1. 遍历所有 4 个方向的位移向量', style='List Bullet'
)
doc.add_paragraph(
    '2. 跳过当前意图方向（已在上一段以 80% 概率处理）', style='List Bullet'
)
doc.add_paragraph(
    '3. 剩余 20% 总概率平均分配到另外 3 个方向 → 每个滑移方向约 6.67%（=0.2/3）', style='List Bullet'
)
doc.add_paragraph(
    '4. 如果某个滑移方向撞墙，反弹停在原地', style='List Bullet'
)

doc.add_paragraph('')
add_hint(doc, '为什么用 0.2 / 3 而不是直接写 0.0667？因为除以 3 更精确，且如果未来动作空间从 4 个变成 5 个，代码只需改动作字典而不用手工重算概率。', '✏️ 工程心得')

doc.add_paragraph('')
doc.add_paragraph(
    '最终转移概率分布示例（在格子 (1,2) 执行 ↑）：'
)
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Light Grid Accent 1'
for i, txt in enumerate(['目标状态', '概率', '说明']):
    t2.rows[0].cells[i].text = txt
    t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
add_table_row(t2, ['(0,2) — 上面', '0.8 (80%)', '按意图移动'])
add_table_row(t2, ['(2,2) — 下面', '0.2/3 ≈ 6.67%', '滑移'])
add_table_row(t2, ['(1,1) — 左面', '0.2/3 ≈ 6.67%', '滑移'])
add_table_row(t2, ['(1,3) — 右面', '0.2/3 ≈ 6.67%', '滑移'])

doc.add_paragraph('')
add_hint(doc,
    '80-20 分割和「滑移均分」是 GridWorld 的标准设计，源自经典 RL 教材 Sutton & Barto。'
    '这种随机性迫使智能体学习鲁棒策略——即使执行不完全可靠，也能到达目标。',
    '🎯 设计理念'
)

make_section_heading(doc, '6.3 返回值', level=2)
doc.add_paragraph(
    '函数返回一个字典 mdp，包含：SIZE, n_states, n_actions, actions, gamma, '
    'GOAL_IDX, TRAP_IDX, R, P, pos_to_idx。'
    '后续 Part 6-8 都以这个 mdp 字典为输入进行策略评估和优化。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. Part 6 — Bellman 方程
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '7. Part 6 — Bellman 方程', level=1)

doc.add_paragraph(
    'Bellman 方程是强化学习的核心数学工具，它将「最优决策」问题分解为子问题的递归形式。'
    '本 Part 实现了两个变体：状态值函数 V(s) 和动作值函数 Q(s,a)。'
)

make_section_heading(doc, '7.1 关键公式', level=2)
doc.add_paragraph('策略状态值函数 Bellman 方程：')
p = doc.add_paragraph()
run = p.add_run('   V^π(s) = Σ_a π(a|s) · [ R(s,a) + γ Σ_{s\'} P(s\'|s,a) V^π(s\') ]')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph('最优状态值函数 Bellman 方程（Bellman 最优性方程）：')
p = doc.add_paragraph()
run = p.add_run('   V*(s) = max_a [ R(s,a) + γ Σ_{s\'} P(s\'|s,a) V*(s\') ]')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph('')
doc.add_paragraph(
    '直观理解：当前状态的价值 = 立即奖励 + 折扣后的未来价值期望。'
    '这是一种「自洽」的递归关系——通过迭代求解直到收敛。'
)

make_section_heading(doc, '7.2 核心代码详解', level=2)

# 策略评估
add_code_block(doc, '''# 随机策略：每个动作概率相等
policy = np.ones((n_states, n_actions)) / n_actions

# 策略评估：迭代求解 V^π
V = np.zeros(n_states)
theta = 1e-6
for i in range(max_iter):
    delta = 0
    for s in range(n_states):
        v_old = V[s]
        v_new = 0
        for a in range(n_actions):
            p_a = policy[s, a]
            if p_a == 0: continue
            bellman_sum = R[s][a]
            for s_next, prob in P[s][a].items():
                bellman_sum += gamma * prob * V[s_next]
            v_new += p_a * bellman_sum
        V[s] = v_new
        delta = max(delta, abs(v_old - v_new))
    if delta < theta:
        break''', 'L354-377 策略评估')

doc.add_paragraph('这段代码实现了「迭代策略评估」，逐行解读：')
add_code_ref(doc, 354, '随机策略：每个状态下 4 个动作的概率各 0.25（无先验知识时最公平的初始化）')
add_code_ref(doc, 358-359, 'V 初始化为全零；theta=1e-6 是收敛阈值（两次迭代变化小于此值即认为收敛）')
add_code_ref(doc, 362-363, '遍历所有状态 s，记录旧值 v_old 以便后续比较变化量')
add_code_ref(doc, 366-372, '对每个动作 a：按概率 p_a 加权求和 Bellman 方程的值')
add_code_ref(doc, 370-371, '核心 Bellman 求和：R(s,a) + γ·Σ P·V(s\')')
add_code_ref(doc, 374, '更新 V[s] 并记录最大变化量 delta')
add_code_ref(doc, 376-377, '当所有状态的变化都小于 theta 时停止迭代')

doc.add_paragraph('')
add_hint(doc,
    '策略评估是「预测」问题——给定策略 π，计算其值函数 V^π。'
    '它是策略迭代（Part 7）中"策略评估"步骤的核心。',
    '🧩'
)

# 最优值函数
add_code_block(doc, '''# 最优值函数 V*
V_opt = np.zeros(n_states)
for i in range(max_iter):
    delta = 0
    for s in range(n_states):
        v_old = V_opt[s]
        q_values = []
        for a in range(n_actions):
            q = R[s][a]
            for s_next, prob in P[s][a].items():
                q += gamma * prob * V_opt[s_next]
            q_values.append(q)
        V_opt[s] = max(q_values)
        delta = max(delta, abs(v_old - V_opt[s]))
    if delta < theta:
        break''', 'L380-394 Bellman 最优性方程（值迭代）')

doc.add_paragraph('与策略评估的关键区别：')
doc.add_paragraph(
    '• 策略评估用 Σ_a π(a|s) 对动作加权平均',
    style='List Bullet'
)
doc.add_paragraph(
    '• 最优值函数用 max_a 取所有动作中的最大值',
    style='List Bullet'
)
doc.add_paragraph(
    '• max 操作对应「智能体总是选择当前看来最好的动作」这一假设',
    style='List Bullet'
)
doc.add_paragraph(
    '• 这实际上就是值迭代（Part 8）——只是在这里先演示以引出概念',
    style='List Bullet'
)

make_section_heading(doc, '7.3 结果解读', level=2)
doc.add_paragraph('输出结果（基于随机策略）：')
add_output_block(doc, '''随机策略 V(s) 收敛
起点 V(0)     = -0.0456   ← 随机游走可能掉入陷阱
最优 V*(0)    = 0.7524    ← 最优策略下起点也有高价值''')

doc.add_paragraph('解读：')
doc.add_paragraph(
    '• 随机策略下起点 V(0) ≈ -0.05：因为随机游走的智能体可能掉入陷阱 (-1)，'
    '也可能到达终点 (+1)，净期望值为负（陷阱更近）。',
    style='List Bullet'
)
doc.add_paragraph(
    '• 最优策略下 V*(0) ≈ 0.75：知道最佳路径后，起点距离终点虽远，'
    '但通过 γ=0.9 的折扣累积，价值很高。',
    style='List Bullet'
)

doc.add_paragraph('')
doc.add_paragraph(
    'heatmap 可视化清晰地展示了 V(s) 的空间分布：越靠近终点的格子颜色越暖（高价值），'
    '陷阱本身为深蓝色（-1），陷阱周围格子也受到负向影响。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. Part 7 — 策略迭代
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '8. Part 7 — 策略迭代', level=1)

doc.add_paragraph(
    '策略迭代（Policy Iteration）是求解 MDP 最优策略的经典算法。'
    '它交替执行两个步骤：策略评估（计算当前策略的值函数）和策略改进（基于值函数更新策略），'
    '直到策略不再变化（收敛）。'
)

make_section_heading(doc, '8.1 算法流程', level=2)

algo_table = doc.add_table(rows=1, cols=2)
algo_table.style = 'Light Grid Accent 1'
for i, txt in enumerate(['步骤', '说明']):
    algo_table.rows[0].cells[i].text = txt
    algo_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

add_table_row(algo_table, ['初始化', '随机初始化策略 π(s) 和 V(s)'])
add_table_row(algo_table, ['策略评估 (E)', '迭代求解 V^π(s)，直到收敛'])
add_table_row(algo_table, ['策略改进 (I)', 'π_new(s) = argmax_a Q^π(s,a)'])
add_table_row(algo_table, ['检查收敛', '如果 π_new == π，停止；否则返回步骤 2'])
doc.add_paragraph('')

add_hint(doc,
    '策略迭代的收敛性有严格数学保证：每次改进都严格提升策略，'
    '而有限 MDP 的策略空间是有限的，因此必然在有限步内收敛到最优。',
    '📐 理论保证'
)

make_section_heading(doc, '8.2 核心代码详解', level=2)

add_code_block(doc, '''policy = np.random.randint(0, n_actions, size=n_states)
V = np.zeros(n_states)''', 'L434-435 初始化')

doc.add_paragraph(
    'policy 是长度为 16 的整数数组，每个元素 0~3 代表 ↑↓←→。'
    '使用随机初始化（而非全零），避免对称性导致的初始偏差。'
)

add_code_block(doc, '''def policy_evaluation(policy, V, theta=1e-6):
    for _ in range(1000):
        delta = 0
        for s in range(n_states):
            v_old = V[s]
            a = policy[s]
            v_new = R[s][a]
            for s_next, prob in P[s][a].items():
                v_new += gamma * prob * V[s_next]
            V[s] = v_new
            delta = max(delta, abs(v_old - v_new))
        if delta < theta:
            break
    return V''', 'L437-450 策略评估函数')

doc.add_paragraph('与 Part 6 的策略评估有两处重要差异：')
doc.add_paragraph(
    '1. 确定性策略：policy[s] 是唯一确定的动作 a，因此不需要对动作加权平均（没有 Σ_a π(a|s)）',
    style='List Bullet'
)
doc.add_paragraph(
    '2. 就地更新：直接使用 V 的当前值进行迭代，而非拷贝。'
    '这是一种 Gauss-Seidel 风格的迭代，通常收敛更快',
    style='List Bullet'
)

add_code_block(doc, '''def policy_improvement(policy, V):
    policy_stable = True
    for s in range(n_states):
        old_action = policy[s]
        q_values = []
        for a in range(n_actions):
            q = R[s][a]
            for s_next, prob in P[s][a].items():
                q += gamma * prob * V[s_next]
            q_values.append(q)
        policy[s] = int(np.argmax(q_values))
        if old_action != policy[s]:
            policy_stable = False
    return policy, policy_stable''', 'L452-465 策略改进函数')

doc.add_paragraph('策略改进的核心思想是「贪心」（Greedy）：')
doc.add_paragraph('• 对每个状态 s，计算所有 4 个动作的 Q(s,a) 值', style='List Bullet')
doc.add_paragraph('• 选择 Q 值最大的动作作为新的策略 π(s)', style='List Bullet')
doc.add_paragraph('• 如果所有状态的动作都没有改变，说明已经达到最优（policy_stable=True）', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    '策略改进定理保证：如果 Q^π(s, π_new(s)) ≥ V^π(s) 对所有 s 成立，'
    '那么 π_new 一定不差于 π。如果存在严格不等式，则 π_new 严格更好。'
)

add_code_block(doc, '''for iteration in range(50):
    V = policy_evaluation(policy, V)
    policy, stable = policy_improvement(policy, V)
    if stable:
        print(f'第 {iteration+1} 轮: 收敛 [OK]')
        break''', 'L468-475 主循环')

doc.add_paragraph(
    '交替执行评估和改进，最多 50 轮。GridWorld（16 状态 × 4 动作）通常只需 3-5 轮即可收敛。'
    '最终打印出的最优策略类似：'
)

add_output_block(doc, '''最优策略:
  →  →  →  ↓
  ↑  X  →  ↓
  ↑  ↑  →  ↓
  ↑  ←  ←  G''')

doc.add_paragraph(
    '可以直观验证策略的合理性：所有箭头都避开陷阱 (X)，指向终点 (G)。'
    '注意有些格子可能有多个等优动作（例如 (3,1) 向左和向下都是最优），'
    '策略取 argmax 时会选索引较小的那个（这里是 ←）。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. Part 8 — 值迭代
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '9. Part 8 — 值迭代', level=1)

doc.add_paragraph(
    '值迭代（Value Iteration）是策略迭代的「合二为一」版本——'
    '它直接迭代 Bellman 最优性方程，无需显式的策略评估步骤。'
    '每次迭代同时更新 V(s) 和隐含的策略。'
)

make_section_heading(doc, '9.1 核心代码详解', level=2)

add_code_block(doc, '''V = np.zeros(n_states)
theta = 1e-6

for iteration in range(1000):
    delta = 0
    for s in range(n_states):
        v_old = V[s]
        q_max = -np.inf
        for a in range(n_actions):
            q = R[s][a]
            for s_next, prob in P[s][a].items():
                q += gamma * prob * V[s_next]
            if q > q_max:
                q_max = q
        V[s] = q_max
        delta = max(delta, abs(v_old - V[s]))
    if delta < theta:
        print(f'收敛于第 {iteration+1} 次迭代')
        break''', 'L526-545 值迭代主循环')

doc.add_paragraph('与 Part 6 的 Bellman 最优性方程求解完全一致——因为那就是值迭代本身。')
doc.add_paragraph('')
doc.add_paragraph('关键特征：')
doc.add_paragraph('• q_max = -np.inf 初始化，确保无论 R(s,a) 多小都能正确取 max', style='List Bullet')
doc.add_paragraph('• 每次迭代对每个状态取 max_a Q(s,a) 并更新 V(s)', style='List Bullet')
doc.add_paragraph('• 不需要显式存储策略——隐含在 V 中（argmax 可以需要时再计算）', style='List Bullet')
doc.add_paragraph('• 收敛后通过 argmax 一次性地恢复出最优策略', style='List Bullet')

add_code_block(doc, '''policy = np.zeros(n_states, dtype=int)
for s in range(n_states):
    q_values = []
    for a in range(n_actions):
        q = R[s][a]
        for s_next, prob in P[s][a].items():
            q += gamma * prob * V[s_next]
        q_values.append(q)
    policy[s] = int(np.argmax(q_values))''', 'L547-555 从 V* 提取最优策略')

doc.add_paragraph(
    '值迭代收敛后 V[s] 就是 V*(s)。但还需要从 V* 中恢复策略：'
    '对每个状态，计算所有动作的 Q*(s,a) = R(s,a) + γ·Σ P·V*(s\')，'
    '取最大 Q 值对应的动作。这与策略改进步骤完全相同。'
)

make_section_heading(doc, '9.2 策略迭代 vs 值迭代', level=2)

compare_table = doc.add_table(rows=1, cols=4)
compare_table.style = 'Light Grid Accent 1'
for i, txt in enumerate(['对比维度', '策略迭代', '值迭代', '适用场景']):
    compare_table.rows[0].cells[i].text = txt
    compare_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
add_table_row(compare_table, ['迭代单位', '策略 + 值函数', '仅值函数', '—'])
add_table_row(compare_table, ['每轮计算量', '评估到收敛（很多步）', '一步 Bellman 更新', '—'])
add_table_row(compare_table, ['收敛轮数', '少（~5 轮）', '多（~20 轮）', '—'])
add_table_row(compare_table, ['总计算成本', '高（精确评估）', '低（近似更新）', '—'])
add_table_row(compare_table, ['状态空间', '小到中等（<10⁴）', '任意大小', '策略迭代状态数受限'])
add_table_row(compare_table, ['输出', '精确最优策略', 'ε-最优策略', '—'])

doc.add_paragraph('')

add_hint(doc,
    '实际应用中值迭代更常用，因为它实现简单、不需要在每次策略评估中等待完全收敛。'
    '修改版：在值迭代的每一步异步更新策略，称为「修改的策略迭代」。',
    '💡 实践建议'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. Extra — 收敛过程可视化
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '10. Extra — 收敛过程可视化', level=1)

doc.add_paragraph(
    '额外部分通过跟踪值迭代过程中 V(s) 的变化，直观展示收敛行为。'
    '这有助于建立对迭代算法收敛性的直觉理解。'
)

add_code_block(doc, '''V_track = []
V = np.zeros(n_states)
theta = 1e-6

for iteration in range(100):
    delta = 0
    for s in range(n_states):
        ...（值迭代更新）...
    V_track.append(V.copy())
    if delta < theta:
        break''', 'L589-608 轨迹记录')

doc.add_paragraph(
    '每轮迭代后将 V 的拷贝存入 V_track。使用 .copy() 而非直接引用，'
    '确保保存的是当前时刻的快照而非指向最终值的引用。'
)

doc.add_paragraph('')
doc.add_paragraph('可视化包含两张子图：')
doc.add_paragraph(
    '1. 左上：起点(0)、终点(15)、陷阱(5) 的 V(s) 随迭代次数变化曲线。'
    '可以看到终点值迅速收敛到+1，陷阱到-1，起点逐渐上升并收敛。',
    style='List Bullet'
)
doc.add_paragraph(
    '2. 右下：每轮 V(s) 的最大变化量（对数坐标）。'
    '曲线呈近似线性下降，说明收敛速度是指数级的——这就是线性收敛。',
    style='List Bullet'
)

doc.add_paragraph('')
add_hint(doc,
    '对数坐标下 delta 呈直线下降，意味着每次迭代误差减少一个固定比例。'
    '这是值迭代在线性收敛速率下的典型行为。如果曲线变平，说明接近收敛阈值 theta。',
    '📊 解读收敛曲线'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. 主程序调度
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '11. 主程序调度 — Part 编排与依赖管理', level=1)

doc.add_paragraph(
    '主程序部分（L641-733）不是一个「脚本」的简单顺序执行，而是一个精心设计的 Part 调度器。'
    '它处理了命令行参数、Part 间依赖和 Part 6-8 的自动触发。'
)

add_code_block(doc, '''if __name__ == '__main__':
    if args.part > 0:
        print(f'>> 单 Part 模式: 只跑 Part {args.part}\\n')
    else:
        print('...完整通关标题...')''', 'L641-650 入口判断')

doc.add_paragraph('入口保护 if __name__ == \'__main__\' 确保文件被 import 时不会自动执行。')

add_code_block(doc, '''# 占位变量（处理 Part 间依赖）
mdp = None; pi_result = None; vi_result = None''', 'L653 依赖变量占位')

doc.add_paragraph(
    'mdp 存储 Part 5 的返回值，pi_result 和 vi_result 分别存储 Part 7 和 Part 8 的结果。'
    '初始化为 None，后续判断是否可用。'
)

add_code_block(doc, '''if should_run(5):
    print(...)
    mdp = part5_mdp_gridworld()

# Part 6-8 依赖 Part 5 的 mdp；单独跑时自动触发 Part 5
if args.part in (6, 7, 8) and mdp is None:
    print('[auto] Part 6-8 依赖 Part 5 的 MDP，先跑 Part 5...')
    mdp = part5_mdp_gridworld()''', 'L681-690 自动依赖解析')

doc.add_paragraph(
    '如果用户指定 --part 6，脚本会自动先跑 Part 5 来获得 MDP 对象，'
    '然后执行 Part 6。这种智能依赖解析让单 Part 模式也能独立运行，无需用户手动准备数据。'
)

add_code_block(doc, '''if should_run(7):
    pi_result = part7_policy_iteration(mdp)
if should_run(8):
    vi_result = part8_value_iteration(mdp)''', 'L698-708 执行策略迭代和值迭代')

doc.add_paragraph(
    'Part 7 和 Part 8 各自产生最优策略和值函数。它们的结果会在最后的验证环节进行对比。'
)

add_code_block(doc, '''if args.part == 0:
    print('Extra: 收敛过程可视化')
    extra_convergence_plot()''', 'L711-715 Extra（仅全部模式）')

doc.add_paragraph(
    '收敛可视化只在全跑模式下执行，避免单 Part 模式下产生意外输出。'
)

add_code_block(doc, '''if args.part == 0 or (args.part in (7, 8) and pi_result and vi_result):
    if pi_result and vi_result:
        ...比较策略迭代和值迭代的 V(0)...''', 'L718-731 验证环节')

doc.add_paragraph(
    '最后的验证步骤比较策略迭代和值迭代求出的最优值函数是否一致。'
    '如果 ΔV < 1e-4，打印 [OK]——这是对课本理论「两种算法收敛到同一最优值函数」的实验验证。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════════════════════════
make_section_heading(doc, '附录 A — 核心公式速查', level=1)

doc.add_paragraph('')

formulas = [
    ('Bellman 期望方程 (V)',
     'V^π(s) = Σ_a π(a|s) [ R(s,a) + γ Σ_{s\'} P(s\'|s,a) V^π(s\') ]',
     '给定策略 π 下状态 s 的价值'),
    ('Bellman 期望方程 (Q)',
     'Q^π(s,a) = R(s,a) + γ Σ_{s\'} P(s\'|s,a) V^π(s\')',
     '给定策略 π 下在状态 s 执行动作 a 的价值'),
    ('Bellman 最优性方程 (V*)',
     'V*(s) = max_a [ R(s,a) + γ Σ_{s\'} P(s\'|s,a) V*(s\') ]',
     '最优值函数，对应最优策略'),
    ('策略改进',
     'π\'(s) = argmax_a Q^π(s,a)',
     '贪心地选择 Q 值最大的动作'),
    ('值函数与 Q 的关系',
     'V^π(s) = Σ_a π(a|s) Q^π(s,a)',
     'V 是 Q 在策略分布下的期望'),
]

for name, formula, desc in formulas:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(f'    {formula}')
    doc.add_paragraph(f'    {desc}')
    doc.add_paragraph('')

make_section_heading(doc, '附录 B — 常见问题', level=1)

faqs = [
    ('Q: 为什么用字典存 P 而不是数组？',
     'A: 每个 (s,a) 通常只转移到少数几个状态（本例中最多 4 个），'
     '用字典存储稀疏转移矩阵更省内存，且代码更易读。')
    ,
    ('Q: 为什么策略迭代比值迭代少很多轮？',
     'A: 策略迭代每轮都精确评估策略到收敛，然后一步改进，步幅大但每步代价高。'
     '值迭代每轮只走一小步（一次 Bellman 更新），需要更多轮但每轮代价低。')
    ,
    ('Q: GridWorld 的随机转移模拟了什么物理现象？',
     'A: 模拟现实中的执行噪声——比如机器人轮子打滑、风扰、'
     '传感器误差等。让智能体在不确定性下做决策，这正是 RL 区别于最优控制的关键。')
    ,
    ('Q: γ=0.9 怎么选？',
     'A: γ ∈ [0,1]。γ→0 只看眼前奖励（短视），γ→1 考虑长远。'
     'GridWorld 取 0.9 是经验值，既考虑长远又保证值迭代快速收敛。'
     '在金融等场景常用 γ=0.99+。')
    ,
    ('Q: torch.no_grad() 的作用？',
     'A: 关闭计算图追踪。推理/评估时不需要梯度，no_grad 下运算更快、省内存。'
     '如果不关，PyTorch 会继续为每个操作构建计算图，造成内存泄漏。')
    ,
    ('Q: 为什么 Part 4 用 unsqueeze(-1)？',
     'A: 原始数据 X 形状为 [N, 10]，MLP 的 Linear 层需要 [N, 10] 输入。'
     '但可视化或后续处理时可能需要保留「通道维度」信息，'
     'unsqueeze(-1) 将其变为 [N, 10, 1] 后再 view 展平为 [N, 10]。'
     '这是一种常见的「先扩展再压缩」的模式。')
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph(a)
    doc.add_paragraph('')

# ── 保存 ──
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'notes')
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, 'Week9_Complete_Code_Detailed_Explanation.docx')
doc.save(OUTPUT_PATH)
print(f'[OK] 文档已生成: {OUTPUT_PATH}')
