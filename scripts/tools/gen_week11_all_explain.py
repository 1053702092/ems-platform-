#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 Week 11 四个 py 文件的逐行分析 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'
BLUE = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x66, 0x66, 0x66)
LRED = RGBColor(0xC0, 0x39, 0x2B)
DGRAY = RGBColor(0x99, 0x99, 0x99)
CODE_C = RGBColor(0x33, 0x33, 0x33)

def new_doc(title_str, subtitle_str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.paragraph_format.line_spacing = 1.35

    def add_heading(text, level=1):
        hd = doc.add_heading(text, level=level)
        for r in hd.runs:
            r.font.color.rgb = BLUE
        return hd

    def add_para(text, bold=False, sz=11, color=None, indent=0):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_after = Pt(3)
        if indent:
            pa.paragraph_format.left_indent = Cm(indent)
        run = pa.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(sz)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        return pa

    def add_bullet(text, lv=0):
        pa = doc.add_paragraph(text, style='List Bullet')
        pa.paragraph_format.left_indent = Cm(1.5 + lv * 0.8)
        return pa

    def add_code(lines_s, label=None):
        if label:
            pa = doc.add_paragraph()
            run = pa.add_run(label)
            run.bold = True
            run.font.size = Pt(10)
        for line in lines_s.split('\n'):
            pa = doc.add_paragraph()
            pa.paragraph_format.space_before = Pt(0)
            pa.paragraph_format.space_after = Pt(1)
            pa.paragraph_format.left_indent = Cm(1)
            run = pa.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_C

    def add_tbl(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        for i, hd in enumerate(headers):
            run = t.rows[0].cells[i].paragraphs[0].add_run(hd)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'
        for rd in rows:
            row = t.add_row()
            for c, txt in enumerate(rd):
                run = row.cells[c].paragraphs[0].add_run(txt)
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'

    def add_brk():
        doc.add_page_break()

    def add_xplain(line_num, code_line, explanation):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before = Pt(4)
        pa.paragraph_format.space_after = Pt(1)
        run = pa.add_run(f'L{line_num}  ')
        run.font.size = Pt(8)
        run.font.color.rgb = DGRAY
        run = pa.add_run(code_line)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = BLUE
        add_para(explanation, indent=1, sz=10)

    # 封面
    for _ in range(4):
        doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f'{title_str}\n')
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = BLUE
    run = t.add_run(subtitle_str)
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t2.add_run(f'Generated: {datetime.date.today().isoformat()}')
    run.font.size = Pt(10)
    run.font.color.rgb = DGRAY
    add_brk()

    return doc, add_heading, add_para, add_bullet, add_code, add_tbl, add_brk, add_xplain


def add_common(doc, add_heading, add_para, add_bullet, add_code, add_tbl, add_brk, add_xplain):
    """Placeholder for shared content"""
    pass


# ================================================================
# 1. continuous_env.py
# ================================================================
print("Generating Step 1 doc...")
doc, h, p, b, c, t, brk, x = new_doc(
    'Step 1: Continuous Action Environment',
    'week11_continuous_env.py - EMS Simplified + Why DQN Fails'
)

h('File Overview')
p('Creates a continuous-action EMS environment (2D state, 1D continuous action) and demonstrates why DQN cannot handle continuous actions.')

t(['Section', 'Lines', 'Content'], [
    ['EMSEnv class', '26-111', 'Continuous state + continuous action + reward'],
    ['DQN failure demo', '115-203', 'Shows why DQN fails on continuous actions'],
    ['Environment test', '207-235', 'Random policy test'],
    ['Main', '238-247', 'Run test + demo'],
])

brk()
h('EMSEnv Class (Lines 26-111)')

x(26, 'class EMSEnv:', 'Defines the simplified EMS environment. Key difference from GridWorld: states are continuous (not discrete grids), actions are continuous (not 4 directions).')
x(35, 'def __init__(self):', 'Init env params. SOC range 0.2-0.9, state_dim=2, action_dim=1.')
x(39, 'self.state_dim = 2', 'State = [SOC, P_load], 2D continuous vector. Compare to GridWorld\'s 16 discrete states.')
x(40, 'self.action_dim = 1', 'Action = P_fc, 1D continuous value (normalized [0,1]). Compare to GridWorld\'s 4 discrete actions.')
x(49, 'def reset(self):', 'Reset env. SOC=0.6, P_load=0.5 kW. Returns initial state.')
x(57, 'def _get_state(self):', 'Returns current state [SOC, P_load] as numpy float32 array.')
x(61, 'def step(self, action):', 'Core: execute action. Action is normalized [0,1] P_fc.')
x(68, 'p_fc = float(np.clip(action, 0, 1)) * 30.0', 'Denormalize: map [0,1] to [0,30] kW. E.g., action=0.5 -> P_fc=15kW.')
x(71, 'self.p_load = 0.3 + 0.4 * (0.5 + 0.5 * np.sin(...))', 'Simulate load variation using sine wave, range ~0.3-0.7.')
x(77, 'soc_change = power_diff / self.battery_capacity', 'SOC change = net power / capacity. P_fc > P_load -> charging, SOC up.')
x(82, 'fuel_cost = -0.01 * p_fc', 'Fuel cost. Higher P_fc = more fuel = more negative reward.')
x(85, 'soc_penalty = -0.5 * (self.soc - 0.6) ** 2', 'SOC penalty. Deviation from 0.6 is penalized.')
x(89, 'if self.soc <= self.soc_min or self.soc >= self.soc_max:', 'Bound penalty. SOC outside [0.2, 0.9] gets -1.0.')

brk()
h('DQN Failure Demo (Lines 115-203)')
p('This is the key section showing why DQN fundamentally cannot handle continuous actions.', bold=True, sz=12)

x(115, 'class DQN_Continuous(nn.Module):', 'Fake DQN - output layer has only 1 neuron (not 4). It directly outputs action values instead of Q values. This is supervised learning, not RL.')
x(126, 'self.net = nn.Sequential(...)', 'Network: 2D input -> 64 hidden -> 1D output. Standard DQN outputs Q-values for each action; this forcibly outputs a single continuous value.')
x(136, 'def demo_dqn_failure():', 'Main demo function. Shows two fundamental problems.')
x(157, "print('Problem 1: Cannot argmax')", 'Problem 1: DQN selects actions by argmax over all actions. Continuous actions have infinite possibilities - cannot enumerate.')
x(163, "print('Problem 2: Update formula also breaks')", 'Problem 2: Q-learning formula target = r + gamma * max Q(s\', a\') requires max over a\' - impossible for continuous actions.')
x(181, 'a = float(torch.sigmoid(q_net(s_tensor)).item())', 'Forcibly outputs an action value directly. sigmoid ensures [0,1] range. But this is not RL - no Q-value concept.')
x(199, "print('DQN cannot handle continuous actions.')", 'Core conclusion. This is a mathematical limitation (argmax requires enumeration), not a tuning problem.')
x(202, "print('Must switch to Policy Gradient methods.')", 'Leads to next sections: REINFORCE, Actor-Critic, PPO.')

path1 = os.path.join(OUT_DIR, 'Week11_Step1_ContinuousEnv_逐行精讲.docx')
doc.save(path1)
print(f'  OK: {path1}')


# ================================================================
# 2. reinforce.py
# ================================================================
print("Generating Step 2 doc...")
doc, h, p, b, c, t, brk, x = new_doc(
    'Step 2: REINFORCE (Policy Gradient)',
    'week11_reinforce.py - First continuous-action RL algorithm'
)

h('File Overview')
p('Implements REINFORCE (also called Vanilla Policy Gradient or Monte Carlo Policy Gradient). This is the simplest algorithm that can handle continuous actions.')

t(['Section', 'Lines', 'Content'], [
    ['Environment', '28-62', 'Same EMS env as Step 1'],
    ['PolicyNet', '66-104', 'pi(s) -> [mu, sigma] -> sample action'],
    ['REINFORCE algorithm', '108-189', 'Run episode -> compute G_t -> loss = -logPi x G'],
    ['Test + Plot', '193-250', 'Test learned policy + plot training curve'],
])

brk()
h('Policy Network - PolicyNet (Lines 66-104)')
p('Core difference from DQN. DQN outputs Q-values, policy network outputs action distribution parameters (mean and std).', bold=True, sz=12)

x(66, 'class PolicyNet(nn.Module):', 'Policy network. Input state, output action distribution [mu, sigma]. Fundamentally different from DQN\'s Q(s) -> [Qup, Qdown, Qleft, Qright].')
x(73, 'def __init__(self, state_dim=2, hidden=64, action_dim=1):', 'Network: 2D input -> 64 hidden -> 64 hidden -> two outputs (mean + log_std).')
x(78, 'self.mean_head = nn.Linear(hidden, action_dim)', 'Mean output head: 64 -> 1 (action mu). Final action will be near mu.')
x(79, 'self.log_std = nn.Parameter(torch.zeros(action_dim))', 'Log standard deviation, trainable. Init log_std=0 -> std=1. Controls exploration range.')
x(84, 'mean = torch.tanh(self.mean_head(x))', 'tanh maps output to [-1, 1], then rescale to [0, 1] as action mean.')
x(86, 'std = torch.exp(self.log_std.clamp(-5, 2))', 'std = e^(log_std), guaranteed positive. Clamp prevents numerical overflow.')
x(89, 'def get_action(self, state):', 'Select action: forward -> [mu, sigma] -> sample from Normal distribution -> continuous action.')
x(94, 'm = dist.Normal(mean, std)', 'Create Normal distribution N(mu, sigma). Key to continuous actions - not picking discrete actions but sampling from a distribution.')
x(95, 'a = m.sample()', 'Sample a continuous action value. Different each time - this is "policy stochasticity."')
x(99, 'def evaluate(self, state, action):', 'Compute log_prob for a given state-action pair (used in training with gradients).')

brk()
h('REINFORCE Algorithm (Lines 108-189)')
p('Core formula: del J = E[ del log pi(a|s) x G ]', bold=True, sz=12)
p('G > 0 -> increase action probability; G < 0 -> decrease action probability', sz=10)

x(108, 'def reinforce(episodes=500, lr=0.001):', 'REINFORCE main function. 500 episodes, lr=0.001.')
x(120, 'policy = PolicyNet()', 'Create policy network. This is the ONLY network - REINFORCE has no Critic.')
x(121, 'optimizer = optim.Adam(policy.parameters(), lr=lr)', 'Adam optimizer. Same as DQN.')
x(140, 'while not done:', 'Run one episode. Same as DQN.')
x(141, 'a, trace = policy.get_action(s)', 'Select action via policy (continuous value!). Key difference from DQN\'s argmax.')
x(143, 'sp, reward, done, _ = env.step(a)', 'Execute action. Different environment from DQN\'s GridWorld, but same interface.')
x(144, 'transitions.append((s_a, a_val, reward))', 'Store (state, action, reward). Note: no Q-values stored - REINFORCE does not learn Q-values.')
x(148, '# Compute G_t (backward cumulative)', 'Core step: Monte Carlo return. Accumulate rewards backward from last step.')
x(151, 'for r in reversed(rewards):', 'Traverse rewards backward from last step.')
x(152, 'G = r + 0.99 * G', 'G_t = r_t + gamma * G_{t+1}. Accumulated from the end. This is the "episode return."')
x(153, 'returns.insert(0, G)', 'Insert computed G at front of list to maintain time order.')
x(158, 'returns_t = (returns_t - returns_t.mean()) / (returns_t.std() + 1e-8)', 'Normalize returns. Positive G -> increase prob, negative G -> decrease prob.')
x(164, 'for (s_i, a_i, _), G_i in zip(transitions, returns_t):', 'Loop over each step to compute loss.')
x(167, 'log_prob = policy.evaluate(s_t, a_t)', 'Recompute log_prob with gradients (get_action was gradient-free).')
x(168, 'loss = loss + (-log_prob * G_i)', 'Core: loss = -sum log_prob x G. G>0 -> minimize -logProb G -> increase logProb -> action more likely.')
x(174, 'optimizer.zero_grad()', 'Clear gradients.')
x(175, 'loss.backward()', 'Backpropagation. Same as DQN.')
x(176, 'optimizer.step()', 'Update parameters. Same as DQN.')

brk()
h('REINFORCE vs DQN Comparison')
t(['Dimension', 'DQN', 'REINFORCE'], [
    ['Network output', 'Q-values (one per action)', 'Action distribution [mu, sigma]'],
    ['Action selection', 'argmax (enumerate all)', 'Sample from distribution'],
    ['Action type', 'Discrete (finite)', 'Continuous (infinite)'],
    ['Value estimate', 'Q(s,a) via TD', 'G_t = episode return (MC)'],
    ['Update timing', 'Every step', 'After episode ends'],
    ['Loss formula', 'MSE(Q, r+gamma Q\')', '-logPi x G'],
    ['Num params', '644 (16-dim input)', '~2500 (2-dim input)'],
])
p('REINFORCE problem: waits for entire episode, high variance, slow learning. Actor-Critic solves this.', bold=True, sz=11)

path2 = os.path.join(OUT_DIR, 'Week11_Step2_REINFORCE_逐行精讲.docx')
doc.save(path2)
print(f'  OK: {path2}')


# ================================================================
# 3. actor_critic.py
# ================================================================
print("Generating Step 3 doc...")
doc, h, p, b, c, t, brk, x = new_doc(
    'Step 3: Actor-Critic',
    'week11_actor_critic.py - Update every step, no need to wait'
)

h('File Overview')
p('Adds a Critic network V(s) on top of REINFORCE, enabling per-step updates instead of waiting for episode end.')

t(['Section', 'Lines', 'Content'], [
    ['Environment', '33-65', 'Same EMS env'],
    ['Actor network', '69-101', 'pi(s) -> [mu, sigma] (same as REINFORCE)'],
    ['Critic network', '104-122', 'V(s) -> scalar (NEW!)'],
    ['Actor-Critic algorithm', '126-224', 'Per-step Advantage -> update Actor + Critic'],
])

brk()
h('Actor and Critic Networks (Lines 69-122)')

h('Actor (Lines 69-101)', 2)
x(69, 'class Actor(nn.Module):', 'Actor network. Same as REINFORCE PolicyNet: state -> [mu, sigma] -> action distribution.')
x(86, 'def get_action(self, state):', 'Select action. Same as REINFORCE: forward -> sample -> return continuous action.')
x(96, 'def evaluate(self, state, action):', 'Compute log_prob with gradients. Same as REINFORCE.')

h('Critic (Lines 104-122)', 2)
p('Critic is the ONLY new addition compared to REINFORCE, but it solves the core problem.', bold=True, sz=11)
x(104, 'class Critic(nn.Module):', 'Critic network. REINFORCE does NOT have this. The Critic estimates "how good is the current state."')
x(111, 'def __init__(self, state_dim=2, hidden=64):', 'Two hidden layers MLP: 2D state -> 64 hidden -> 64 hidden -> 1 scalar V(s).')
x(118, 'nn.Linear(hidden, 1)  # Output a scalar V(s)', 'Output layer: 1 neuron. V(s) is a single number representing "state value."')

brk()
h('Actor-Critic Algorithm (Lines 126-224)')
p('Fundamental difference from REINFORCE: update every step instead of waiting for entire episode.', bold=True, sz=12)

x(144, 'env = EMSEnv()', 'Create environment. Same as REINFORCE.')
x(145, 'actor = Actor()', 'Create Actor network. Same as REINFORCE PolicyNet.')
x(146, 'critic = Critic()', 'Create Critic network. NEW! This is the difference between AC and REINFORCE.')
x(147, 'actor_opt = optim.Adam(actor.parameters(), lr=lr)', 'Actor optimizer.')
x(148, 'critic_opt = optim.Adam(critic.parameters(), lr=lr * 2)', 'Critic optimizer. Faster learning rate for Critic.')
x(162, 'for ep in range(1, episodes + 1):', 'Main loop. Same as REINFORCE.')
x(170, 'a = actor.get_action(s)', 'Actor selects action (continuous). Same as REINFORCE.')
x(173, 'sp, reward, done, _ = env.step(a)', 'Execute action. Same as REINFORCE.')
x(185, 'V_s = critic(s_t)', 'Critic estimates current state value V(s). REINFORCE does NOT have this.')
x(188, 'advantage = r_t + gamma * V_sp * (not done) - V_s', 'KEY: Advantage = r + gamma*V(s\') - V(s). If > 0, step was better than expected; < 0, worse.')
p('Advantage replaces G_t from REINFORCE. G_t needs entire episode; Advantage can be computed every step.', sz=10, indent=1)
x(193, 'actor_loss = -(log_prob * advantage.detach()).mean()', 'Update Actor: same formula as REINFORCE! But G_t is replaced by Advantage.')
x(196, 'actor_opt.zero_grad()', 'Clear gradients.')
x(197, 'actor_loss.backward()', 'Backprop for Actor.')
x(198, 'actor_opt.step()', 'Update Actor params.')
x(203, 'td_target = r_t + gamma * V_sp * (not done)', 'Critic target: r + gamma*V(s\'). Same formula as DQN target.')
x(204, 'critic_loss = loss_fn(V_s, td_target)', 'Update Critic: make V(s) close to r + gamma*V(s\'). This is TD learning.')
x(207, 'critic_opt.step()', 'Update Critic params.')

brk()
h('REINFORCE vs Actor-Critic Comparison')
t(['Dimension', 'REINFORCE', 'Actor-Critic'], [
    ['Policy network', 'PolicyNet pi(s)->[mu,sigma]', 'Actor pi(s)->[mu,sigma]'],
    ['Value network', 'None', 'Critic V(s)'],
    ['Evaluation', 'G_t = sum(r) (MC return)', 'Advantage = r+gammaV-V\' (TD)'],
    ['Update timing', 'After episode', 'Every step'],
    ['Variance', 'High', 'Low (due to Critic)'],
    ['Speed', 'Slow (wait for episode)', 'Fast (per-step update)'],
])
p('AC advantage: per-step update, lower variance. But Critic can be biased, causing wrong update directions. PPO\'s clip solves this.', bold=True, sz=11)

path3 = os.path.join(OUT_DIR, 'Week11_Step3_ActorCritic_逐行精讲.docx')
doc.save(path3)
print(f'  OK: {path3}')


# ================================================================
# 4. ppo.py
# ================================================================
print("Generating Step 4 doc...")
doc, h, p, b, c, t, brk, x = new_doc(
    'Step 4: PPO (Proximal Policy Optimization)',
    'week11_ppo.py - Interview focus, final algorithm chosen for EMS'
)

h('File Overview')
p('Adds the clip mechanism on top of Actor-Critic - limits policy update magnitude to prevent catastrophic updates.')
p('PPO is the algorithm ultimately chosen for the EMS project. Interview essential.', bold=True, sz=12)

t(['Section', 'Lines', 'Content'], [
    ['Environment', '36-74', 'Slightly easier EMS env'],
    ['Actor + Critic', '78-125', 'Same network structure as AC'],
    ['PPO algorithm', '129-273', 'GAE + importance ratio + clip + multi-epoch'],
    ['Test + Summary', '277-359', 'Test + 3-method comparison'],
])

brk()
h('PPO Core Innovation - clip (Lines 240-244)')
p('PPO adds 3 things to Actor-Critic: (1) importance ratio, (2) clip, (3) multi-epoch updates. Clip is the most important.', bold=True, sz=12)

x(129, 'def ppo(episodes=500, lr=0.0003, clip_eps=0.2, epochs=10, ...):', 'PPO main function. New params: clip_eps (clip range), epochs (data reuse count).')
x(159, 'gamma = 0.99', 'Discount factor. Same as AC.')
x(160, 'gae_lambda = 0.95', 'GAE smoothing coefficient. More stable advantage estimation than AC\'s simple TD error.')
x(161, 'entropy_coef = 0.01', 'Entropy bonus coefficient. Encourages exploration (prevents premature convergence).')

h('Step 1: Collect Data (Lines 164-183)', 2)
x(164, 's = env.reset()', 'Reset environment.')
x(166, 'states, actions, rewards, dones, log_probs_old = [], [], [], [], []', 'Note: log_probs_old stores the OLD policy\'s log_prob. Core to PPO: must compare with old policy.')
x(169, 'a, lp = actor.get_action(s)', 'Actor selects action, records log_prob. Unlike AC, PPO stores the old log_prob.')

h('Step 2: Compute GAE (Lines 185-213)', 2)
x(185, '# Compute GAE (Generalized Advantage Estimation)', 'GAE is smoother than AC\'s simple TD error, reducing training variance.')
x(193, 'with torch.no_grad():', 'Critic inference only, no gradients.')
x(194, 'values = critic(states_t).squeeze()', 'Critic computes V(s) for each state.')
x(197, 'gae = 0', 'Accumulate GAE backward. Smoother than AC\'s simple Advantage.')
x(204, 'delta = rewards[t] + gamma * next_value * (1 - dones[t]) - values[t]', 'TD error = r + gamma*V(sp) - V(s). Same as AC\'s Advantage.')
x(205, 'gae = delta + gamma * gae_lambda * (1 - dones[t]) * gae', 'GAE core: current TD error + discounted * next GAE. Integrates multi-step info.')

h('Step 3: PPO Core - Clip Update (Lines 215-264)', 2)
p('This is the most important part of PPO. Essential for interviews.', bold=True, sz=12)

x(218, 'for _ in range(epochs):', 'Multi-epoch update. Same batch reused epochs times (default 10). AC uses each datum once.')
x(232, 'log_probs_new, entropy = actor.evaluate(batch_s, batch_a)', 'Recompute log_prob with CURRENT policy (after updates).')
x(238, 'ratio = torch.exp(log_probs_new - batch_old_lp)', 'Importance ratio = pi_new / pi_old. Ratio > 1 means action more likely now; < 1 means less likely.')
x(242, 'surr1 = ratio * batch_adv', 'Standard policy gradient objective: ratio x Advantage. Same as AC.')
x(243, 'surr2 = torch.clamp(ratio, 1 - clip_eps, 1 + clip_eps) * batch_adv', 'CLIP version: limit ratio to [0.8, 1.2]. If ratio exceeds bounds, clip it - prevents too-large updates.')
x(244, 'actor_loss = -torch.min(surr1, surr2).mean()', 'Take min(surr1, surr2): when ratio exceeds [0.8, 1.2], gradient goes to zero - stops update.')
x(248, 'entropy_loss = -entropy_coef * entropy.mean()', 'Entropy bonus. Encourages policy to maintain randomness (prevents premature convergence).')
x(255, 'torch.nn.utils.clip_grad_norm_(actor.parameters(), 0.5)', 'Gradient clipping. Another safety mechanism to prevent gradient explosion.')

brk()
h('3-Method Formula Comparison')
t(['', 'REINFORCE', 'Actor-Critic', 'PPO'], [
    ['Action', 'pi(s) -> sample', 'pi(s) -> sample', 'pi(s) -> sample'],
    ['Value', 'G_t = sum(r)', 'A = r+gammaV-V\'', 'GAE (smoothed)'],
    ['Update', 'Episode end', 'Every step', 'Episode -> multi-epoch'],
    ['Core loss', '-logPi x G', '-logPi x A', '-min(clip ratio x A)'],
    ['Safety', 'None', 'None', 'Clip + grad clipping'],
])
p('PPO is the algorithm chosen for the EMS project. Reasons: continuous actions + stable training + moderate complexity.', bold=True, sz=12)

path4 = os.path.join(OUT_DIR, 'Week11_Step4_PPO_逐行精讲.docx')
doc.save(path4)
print(f'  OK: {path4}')

print('\nAll 4 docx files generated successfully.')
