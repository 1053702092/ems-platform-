#!/usr/bin/env python3
"""
《2026年EMS相关岗位市场调研 + 学习计划对照分析报告》
V2 — 仅含应届生可投岗位（校招/硕士起点）
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 样式设定 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1a, 0x47, 0x8a)

def add_table_h(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def bold_before_colon(doc, text):
    """加粗冒号前的文字"""
    p = doc.add_paragraph()
    if '：' in text:
        idx = text.index('：')
        r1 = p.add_run(text[:idx])
        r1.bold = True
        p.add_run(text[idx:])
    elif ':' in text:
        idx = text.index(':')
        r1 = p.add_run(text[:idx])
        r1.bold = True
        p.add_run(text[idx:])
    else:
        p.add_run(text)
    return p

# ═══════════════════ 封面 ═══════════════════
title = doc.add_heading('2026年 EMS 相关岗位\n应届生校园招聘调研报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('附：与你学习计划的对照分析与调整建议')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('调研日期：2026年7月27日\n筛选条件：仅含应届/校招/硕士起点岗位\n对照基准：STATUS.md (2026-07-24 更新)\n').font.size = Pt(9)

doc.add_page_break()

# ═══════════════════ 目录 ═══════════════════
doc.add_heading('目录', level=1)
toc = [
    '1. 核心结论（先说结果）',
    '2. 六大方向应届岗位一览',
    '  2.1 纯正 EMS（能量管理系统）校招',
    '  2.2 BMS 算法工程师 校招',
    '  2.3 整车能量管理 / 混动控制 校招',
    '  2.4 RL + 控制策略（自动驾驶 / 机器人）校招',
    '  2.5 储能系统 / 数字化能源 校招',
    '  2.6 央企国企 校招',
    '3. 应届薪资天梯图',
    '4. 你的技能 vs 岗位要求对照',
    '5. 当前差距分析',
    '6. 学习计划调整建议',
    '7. 投递优先级 + 时间线',
    '8. 行动清单',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ═══════════════════ 1. 核心结论 ═══════════════════
doc.add_heading('1. 核心结论', level=1)
p = doc.add_paragraph()
p.add_run(
    '这次筛掉资深社招岗后，发现应届可投的 EMS 相关岗位比你想象的多。\n\n'
    '关键数字：\n'
    '  - 纯正EMS校招：约 5-8 家（华思、果下、东方电子、华昱欣等）\n'
    '  - BMS算法校招：约 10-15 家（中创新航、国轩高科、摩瓦、力高等）\n'
    '  - 整车能量管理/控制校招：约 8-12 家（吉利、明阳、三晶、英集动力等）\n'
    '  - RL+控制校招：约 5-8 家（地平线、蔚来、小米等）\n'
    '  - 央企国企：约 20-30 家（发电集团、石油、三峡、中广核等）\n\n'
    '总目标投递数：30-50 家完全足够覆盖。\n\n'
    '好消息: 你对口的岗位远比你感觉的多。'
    '关键是要把项目经验包装成两类（EMS/BMS方向 & RL/控制方向），'
    '分别投不同的公司。'
)

doc.add_page_break()

# ═══════════════════ 2. 六大方向 ═══════════════════
doc.add_heading('2. 六大方向应届岗位一览', level=1)

# ── 2.1 EMS ──
doc.add_heading('2.1 纯正 EMS（能量管理系统）校招', level=2)
doc.add_paragraph(
    '岗位数量不多但稳定存在，集中在储能系统集成商和能源数字化企业。'
)
headers = ['公司', '岗位', '薪资(月/年)', '地点', '专业要求']
rows = [
    ['华思系统', '嵌入式软件工程师(EMS方向)', '7-10k/月', '合肥', '电子/自动化/计算机'],
    ['果下科技', '算法/软件研发(储能EMS)', '15w+/年', '无锡', '控制/电气/软件工程'],
    ['东方电子(国有上市)', '储能EMS工程师', '10-15k/月', '烟台', '电气/计算机/控制'],
    ['上海某国企', 'EMS软件工程师', '10.5-15k/月', '上海闵行', '电气/控制/通信'],
    ['浙江华昱欣', '嵌入式软件(MCU/EMS)', '面议(校招)', '杭州', '电子/自动化/计算机'],
    ['浙江英集动力', '算法工程师(智慧能源)', '面议', '杭州/常州', '能动/控制/计算机'],
]
add_table_h(doc, headers, rows)

p = doc.add_paragraph()
p.add_run('\n小结: ').bold = True
p.add_run(
    '纯正EMS校招约5-8家。薪资一般10-15k/月或15w/年，不算高但胜在对口。'
    '这些岗位更偏嵌入式软件而非纯算法。'
)

# ── 2.2 BMS ──
doc.add_heading('2.2 BMS 算法工程师 校招', level=2)
doc.add_paragraph(
    '这是你最大的应届对口市场。岗位多、企业多、薪资覆盖面广。'
)
headers = ['公司', '岗位', '薪资', '地点', '备注']
rows = [
    ['中创新航 CALB', 'BMS算法工程师', '1-1.4万·13薪', '常州', '440人校招规模'],
    ['国轩高科', 'BMS软/硬件/算法工程师', '10.5-15k/月', '合肥/上海/南京/青岛', '全国+海外多地'],
    ['摩瓦新能源', 'BMS算法工程师', '20-40万/年', '上海/深圳/成都/西安', '硕士20-40万，最高一档'],
    ['吉利控股', '新能源电池BMS开发', '10.5-15k/月', '杭州/宁波/成都', '全球校招'],
    ['力高新能源', 'BMS/储能软件开发', '14-25万/年', '合肥/深圳/烟台', '专业BMS厂商'],
    ['华昱欣', '嵌入式软件(储能/BMS)', '面议', '杭州', '光伏+储能'],
    ['三晶电气', '电池测试/电机控制算法', '面议', '广州', '分布式光伏+储能'],
]
add_table_h(doc, headers, rows)

p = doc.add_paragraph()
p.add_run('\n小结: ').bold = True
p.add_run(
    'BMS算法校招约10-15家。薪资跨度大：国轩/中创新航给1-1.5万/月（中等），'
    '摩瓦新能源给20-40万/年（高配）。你的SOC估计(EKF)项目+BMS方向叙事可以直接用。'
)

# ── 2.3 整车能量管理 ──
doc.add_heading('2.3 整车能量管理 / 混动控制 校招', level=2)
doc.add_paragraph(
    '传统车企校招规模大，岗位名称可能不直接叫"EMS"而是"能量控制平台开发"等。'
)
headers = ['公司', '岗位', '薪资', '地点', '备注']
rows = [
    ['吉利控股', '能量控制平台开发/运动能量应用软件开发', '10.5-15k/月', '杭州/宁波/北京/上海', '全球校招，岗位类型多'],
    ['明阳集团', '控制系统/控制算法/DSP软件工程师', '面议', '中山/深圳/天津/西安', '风电+储能'],
    ['三晶电气', '电机控制算法工程师', '面议', '广州', '光伏+储能'],
    ['力高新能源', 'VCU整车控制器/BMS开发', '14-25万/年', '合肥/深圳/烟台', '新能源控制系统'],
    ['东方电子', '储能EMS工程师(研发)', '10-15k/月', '烟台', '国企上市'],
]
add_table_h(doc, headers, rows)

p = doc.add_paragraph()
p.add_run('\n小结: ').bold = True
p.add_run(
    '约5-8家。吉利的能量控制平台开发岗最对口（直接跟能量管理相关）。'
    '明阳/三晶偏风电和储能控制。'
)

# ── 2.4 RL+控制 ──
doc.add_heading('2.4 RL + 控制策略（自动驾驶/机器人）校招', level=2)
doc.add_paragraph(
    '⚠ 这部分岗位不叫"EMS"，但你的RL+控制+仿真背景可以平移过去，薪资天花板最高。'
)
headers = ['公司', '岗位', '薪资参考', '地点', '要求']
rows = [
    ['地平线', '智能驾驶端到端算法(RL方向)', '30-50万/年', '北京/上海/南京/杭州/深圳', '硕+，熟悉PPO/GRPO等'],
    ['蔚来', '数据算法/端到端模型研发', '28-45万/年', '北京/上海', '硕+，RL/扩散模型'],
    ['小米（校招）', '机器人RL算法/自动驾驶决策', '30-50万/年', '北京/上海/武汉', '具身智能+自动驾驶'],
    ['小鹏汽车（校招）', 'RL算法工程师', '面议', '广州/北京', '端到端+RL后训练'],
]
add_table_h(doc, headers, rows)

p = doc.add_paragraph()
p.add_run('\n小结: ').bold = True
p.add_run(
    '约5-8家。这些岗位对RL要求高，需要能讲清楚PPO原理+有PyTorch实现经验。'
    '你的PPO-EMS项目可以直接用（但面试时要能说出通用RL原理，不局限于EMS场景）。'
)

# ── 2.5 储能 ──
doc.add_heading('2.5 储能系统 / 数字化能源 校招', level=2)
doc.add_paragraph('储能是增量市场，对硕士应届相对友好。')
headers = ['公司', '岗位', '薪资', '地点']
rows = [
    ['华思系统', '嵌入式/硬件/测试(储能BMS/EMS)', '7-10k/月', '合肥'],
    ['果下科技', '算法研发(储能方向)', '15w+/年', '无锡'],
    ['力高新能源', '储能电气/软件开发', '14-25万/年', '合肥/深圳/烟台'],
    ['华昱欣', '储能嵌入式/智能算法', '面议', '杭州'],
    ['三晶电气', '储能系统/电池测试', '面议', '广州'],
    ['英集动力', '智慧能源算法/咨询', '面议', '杭州/常州/郑州'],
]
add_table_h(doc, headers, rows)

p = doc.add_paragraph()
p.add_run('\n小结: ').bold = True
p.add_run(
    '约6-8家。储能公司校招特点是专业包容度高，电气/控制/计算机/化工都能报。'
    '但岗位多为嵌入式或系统工程师，纯算法岗偏少。'
)

# ── 2.6 央企国企 ──
doc.add_heading('2.6 央企国企 校招', level=2)
doc.add_paragraph(
    '这是你的主赛道之一（STATUS.md 写明的 50% 精力）。央企国企对化学/材料背景更友好。'
)
headers = ['单位', '岗位方向', '薪资参考', '投递时间', '化学背景友好度']
rows = [
    ['中石化', '新能源/储能/控制相关', '18-28w/年', '9月启动', '★★★★★'],
    ['中海油', '电气/控制/新能源', '20-35w/年', '9-10月', '★★★★★'],
    ['中石油', '电气/控制/新能源', '18-28w/年', '春招', '★★★★☆'],
    ['国家能源集团', '新能源/电气/控制', '15-25w/年', '9-10月', '★★★★☆'],
    ['中国华能/华电/国电投/大唐', '新能源/储能/电气', '15-25w/年', '9-10月', '★★★★☆'],
    ['三峡集团', '新能源/抽水蓄能', '20-30w/年', '9-10月', '★★★☆☆'],
    ['中广核', '核能/新能源/自动化', '20-30w/年', '9-10月', '★★★☆☆'],
    ['华润电力', '新能源/储能', '15-25w/年', '9-10月', '★★★★☆'],
    ['国网（二批备份）', '其他工学（名额<20%）', '15-30w/年', '2027年2-3月', '★★☆☆☆'],
]
add_table_h(doc, headers, rows)

doc.add_page_break()

# ═══════════════════ 3. 薪资天梯 ═══════════════════
doc.add_heading('3. 应届薪资天梯图', level=1)
doc.add_paragraph('按年薪从低到高排列，标注代表公司：')

headers = ['薪资带', '方向', '代表公司']
rows = [
    ['10-15万/年', '嵌入式EMS/储能', '华思系统、东方电子（国企）'],
    ['15-20万/年', '央企国企基层', '发电集团、中石化基层'],
    ['18-25万/年', 'BMS算法（校招主流）', '国轩高科、中创新航、力高新能源'],
    ['20-30万/年', '央企国企中等', '三峡、中广核、中海油'],
    ['25-35万/年', '整车能量管理/储能算法', '吉利、明阳、果下科技'],
    ['28-45万/年', 'RL+自动驾驶', '蔚来、地平线'],
    ['30-50万/年', 'RL+机器人/端到端', '小米、地平线、小鹏'],
    ['20-40万/年（高配）', 'BMS算法（最高一档）', '摩瓦新能源（硕士20-40万）'],
]
add_table_h(doc, headers, rows)

p = doc.add_paragraph()
p.add_run('\n注意: ').bold = True
p.add_run(
    '摩瓦新能源给的20-40万/年是BMS方向的最高一档，但竞争也会最激烈。'
    '国轩/中创新航的1-1.5万/月是BMS校招的基准线。'
)

doc.add_page_break()

# ═══════════════════ 4. 技能对照 ═══════════════════
doc.add_heading('4. 你的技能 vs 岗位要求对照', level=1)

doc.add_heading('4.1 校招岗位看什么？', level=2)
doc.add_paragraph(
    '校招不要求"量产经验"和"5年+经验"。校招主要看：\n'
    '  1. 专业基础（控制理论/电池原理/RL算法 → 你已经很强）\n'
    '  2. 项目经验（DP/ECMS/MPC/RL-EMS → 远超同龄人）\n'
    '  3. 编程能力（Python 熟练 + C 基础 → 够用）\n'
    '  4. 潜力（学习能力、表达逻辑、STAR叙事 → STATUS.md 在练了）\n\n'
    '⚠ 校招最大的风险不是"技能不够"，而是"投递太少"和"简历不会包装"。'
)

doc.add_heading('4.2 分方向技能匹配表', level=2)

# BMS方向
bold_before_colon(doc, '版本A：EMS/BMS算法工程师方向')
headers = ['技能', '校招需求', '你的现状', '状态']
rows = [
    ['DP/动态规划', '中等', '手写实现，WLTC氢耗↓19.2%', '强项'],
    ['ECMS', '中等', '已实现+多工况验证', '强项'],
    ['MPC', '较高', '已实现+EKF集成', '强项'],
    ['EKF/状态估计', '高（BMS核心）', 'MPC+EKF已集成', '已覆盖'],
    ['锂电池建模/RC模型', '高', '仅有概念，无独立模块', '需补'],
    ['Matlab/Simulink', '高', '已安装R2024b，未深度使用', '需补'],
    ['Python/C', '高', 'Python熟练，C基础', '够用'],
    ['RL算法', '中等（BMS方向不硬性要求）', 'PPO进行中', '够用'],
]
add_table_h(doc, headers, rows)
doc.add_paragraph('')

# RL方向
bold_before_colon(doc, '版本B：RL/控制算法工程师方向')
headers = ['技能', '校招需求', '你的现状', '状态']
rows = [
    ['PPO/RL理论基础', '高', '第11周实现中', '需完成'],
    ['PyTorch', '高', '第9-10周完成基础', '需继续'],
    ['仿真环境搭建', '高', 'EMS仿真平台已搭建', '强项'],
    ['MPC+控制理论', '较高', '已实现', '强项'],
    ['Python', '高', '熟练', '强项'],
    ['C++', '中等', 'LeetCode Easy', '够用（RL岗不硬性要求C++深度）'],
    ['顶会论文', '加分非必须', '无', '不影响校招'],
]
add_table_h(doc, headers, rows)

doc.add_page_break()

# ═══════════════════ 5. 差距分析 ═══════════════════
doc.add_heading('5. 当前差距分析', level=1)

doc.add_heading('5.1 最需要补的：电池建模（3-4天可解决）', level=2)
doc.add_paragraph(
    '现状: 你的项目有 EKF SOC 估计，但你把它当作"MPC的一个子模块"来叙事，'
    '而不是独立的"电池状态估计"项目。\n'
    '后果: 面试BMS岗时，面试官追问"你的电池模型是什么？几阶RC？参数怎么辨识的？"'
    '你目前回答不了。\n'
    '方案: 花3-4天补一个独立的电池建模模块：\n'
    '  1. 二阶RC等效电路模型（Python实现）\n'
    '  2. RLS 参数辨识（在线辨识）\n'
    '  3. EKF SOC估计（你已经会了，独立出来）\n'
    '  4. 整理成独立文档/脚本 → 简历上写成独立项目\n\n'
    '这样你的简历就从"EMS算法"扩展到"EMS+BMS算法"，覆盖面翻倍。'
)

doc.add_heading('5.2 需要熟悉但没那么紧迫的：Simulink', level=2)
doc.add_paragraph(
    '问题: BMS岗位JD普遍要求Matlab/Simulink，你只用Python实现。\n'
    '方案: 不用从零学全部。8月有空就用Simulink搭一个简单的电池RC模型，'
    '让你在面试中至少能说"我会用Simulink做算法验证"。\n'
    '优先级: 低。如果时间来不及，"Python手写实现"也能顶住校招面试。'
)

doc.add_heading('5.3 你的绝对强项', level=2)
doc.add_paragraph(
    '✅ DP/ECMS/MPC 手写实现 — 同龄人几乎没人做到这个深度\n'
    '✅ SOC估计(EKF) + MPC集成 — 工程能力强\n'
    '✅ 四方法量化对比 — 能用数据说话的面试者极少\n'
    '✅ 完整的项目叙事体系 — STATUS.md已有系统化面试训练\n'
    '✅ 127家投递清单 — 比95%的同龄人准备充分\n\n'
    '说白了: 你的核心问题不是"技能不够"，是"不知道这些技能对应哪些公司"。'
    '你现在的水平在校招市场已经是中上水平了。'
)

doc.add_page_break()

# ═══════════════════ 6. 调整建议 ═══════════════════
doc.add_heading('6. 学习计划调整建议', level=1)
doc.add_paragraph(
    '基于此次应届岗位调研，建议对你的学习计划做以下微调：'
)

doc.add_heading('6.1 第11周（8月初）：PPO 实现 — 保持，1周跑通', level=2)
doc.add_paragraph(
    '原计划2周，建议压到1周。只跑通闭环即可，不调优。\n'
    '理由: 地平线/蔚来/小米的RL岗确实要求RL经验，但校招更看重"有没有RL实践经验"'
    '而非"调参好不好"。一个能跑通的PPO-EMS项目足够面试用了。'
)

doc.add_heading('6.2 第12周（8月中）：拆分成"电池专项+简历"', level=2)
doc.add_paragraph(
    '原计划12周是"简历打磨+笔面试准备"。建议：\n'
    '  - 前4天: 电池建模专项（二阶RC + RLS + EKF独立模块）\n'
    '  - 后3天: 两版简历定稿 + 投递清单更新\n\n'
    '理由: 这是性价比最高的3-4天投入。补完电池专项后，你的目标岗位从EMS方向'
    '直接扩展到BMS方向（后者岗位数量多一倍）。'
)

doc.add_heading('6.3 8月下旬：第一批投递启动', level=2)
doc.add_paragraph(
    '没问题，保持原计划。建议第一批发力央企国企 + 市场化BMS公司。'
)

doc.add_heading('6.4 学习计划不需要大改', level=2)
doc.add_paragraph(
    '总体来看，你的STATUS.md学习计划方向正确，节奏合理，不需要大改。\n'
    '唯一实质性变更: 第12周加入"电池建模专项"模块（3-4天）。\n'
    '其他保持原样即可。'
)

doc.add_page_break()

# ═══════════════════ 7. 投递优先级 ═══════════════════
doc.add_heading('7. 投递优先级 + 时间线', level=1)

doc.add_heading('7.1 按应届成功率排序', level=2)
headers = ['优先级', '方向', '精力占比', '薪资范围', '原因']
rows = [
    ['P0', '央企国企（中石化/中海油/发电/三峡）', '40%', '18-35w/年', '应届一次性窗口，化学友好'],
    ['P1', 'BMS算法（摩瓦/国轩/中创新航/力高）', '25%', '14-40万/年', '岗位最多，你的匹配度最高'],
    ['P2', 'RL+控制（地平线/蔚来/小米/小鹏）', '15%', '28-50万/年', '薪资天花板最高'],
    ['P3', 'EMS/能量管理（吉利/明阳/东方电子）', '10%', '12-25万/年', '纯正对口但岗位偏少'],
    ['P4', '储能公司（华思/果下/华昱欣）', '8%', '10-25万/年', '增量市场，专业包容'],
    ['P5', '国网二批（备份）', '2%', '15-30w/年', '渠道窄，仅当备份'],
]
add_table_h(doc, headers, rows)

doc.add_heading('7.2 时间线', level=2)
doc.add_paragraph(
    '7月底（现在）  第11周 — PPO跑通（1周）\n'
    '8月上旬        电池建模专项（3-4天）+ 简历定稿\n'
    '8月中旬        第一批投递：央企 + BMS公司\n'
    '9-10月         大规模秋招投递 + 笔面试\n'
    '11月           中石化/中海油笔试\n'
    '2027年2-3月    国网二批（备份）\n'
)

doc.add_page_break()

# ═══════════════════ 8. 行动清单 ═══════════════════
doc.add_heading('8. 行动清单（可直接写进STATUS.md）', level=1)

doc.add_heading('本周（7月27日-8月3日）', level=2)
items1 = [
    'PPO 轻量实现跑通训练闭环。不调优，跑通即可',
    '产出：PPO-EMS训练曲线 + 与Rule的简要对比',
]
for item in items1:
    doc.add_paragraph(f'  {item}')

doc.add_heading('8月上旬', level=2)
items2 = [
    '电池建模专项（3-4天）：',
    '  1. 二阶RC等效电路模型 Python实现',
    '  2. RLS参数辨识脚本',
    '  3. EKF SOC估计独立成单独模块',
    '  4. 整理成独立文档（面试用）',
    '简历版本A（EMS/BMS方向）定稿',
    '简历版本B（RL/规控方向）定稿',
    '投递清单更新（新增摩瓦新能源、地平线、吉利能量控制等）',
]
for item in items2:
    doc.add_paragraph(f'  {item}')

doc.add_heading('8月下旬', level=2)
items3 = [
    '第一批投递启动（央企+市场化同步）',
    '央企笔试准备（综合60%+专业40%）',
    'LeetCode Easy保持手感',
    '面试八股文积累（控制/电池/RL/燃料电池）',
]
for item in items3:
    doc.add_paragraph(f'  {item}')

# ── 附录 ──
doc.add_heading('附录：你的50家目标投递清单', level=1)
doc.add_paragraph('建议优先投递的应届公司（按方向分类）：')

bold_before_colon(doc, '央企国企（15家）')
items_soe = [
    '中国石化（9月启动秋招）',
    '中国海油（9-10月，海上/沿海城市）',
    '中国石油（春招，25/26届都可报）',
    '国家能源集团（统招6400+人）',
    '中国华能',
    '中国华电',
    '国家电投',
    '中国大唐',
    '三峡集团（抽水蓄能）',
    '中广核（核电站，低房价城市）',
    '华润电力',
    '中电建',
    '中能建',
    '中国通用技术集团',
    '中国船舶集团',
]
for item in items_soe:
    doc.add_paragraph(f'    {item}')

bold_before_colon(doc, 'BMS/EMS算法公司（15家）')
items_bms = [
    '摩瓦新能源（上海/深圳，20-40万/年，最高一档）',
    '国轩高科（合肥/上海/南京）',
    '中创新航（常州）',
    '力高新能源（合肥/深圳/烟台）',
    '阳光电源（合肥）',
    '东方电子（烟台，国有上市）',
    '华思系统（合肥）',
    '果下科技（无锡）',
    '华昱欣（杭州）',
    '英集动力（杭州/常州）',
    '三晶电气（广州）',
    'CET中电技术（深圳/武汉）',
    '江苏林洋能源',
    '派能科技（上海）',
    '宁德时代（冲刺）',
]
for item in items_bms:
    doc.add_paragraph(f'    {item}')

bold_before_colon(doc, 'RL/控制算法公司（8家）')
items_rl = [
    '地平线（北京/上海/南京/杭州/深圳）',
    '蔚来（上海/北京）',
    '小米（北京/上海/武汉）',
    '小鹏汽车（广州/北京）',
    '吉利控股（杭州/宁波）',
    '明阳集团（中山/深圳/西安）',
    '千里智驾（北京/上海）',
    '德赛西威（北京/上海）',
]
for item in items_rl:
    doc.add_paragraph(f'    {item}')

bold_before_colon(doc, '比亚迪（1家）')
doc.add_paragraph('    比亚迪（全球校招，专业包容度高）')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('\n合计: 约39家。再加其他中小公司，凑到50家不难。').bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('\n—— 报告完 ——').bold = True
doc.add_paragraph('数据来源: 各高校就业信息网、各公司校招官网、猎聘、前程无忧\n调研日期: 2026年7月27日 | 对照基准: STATUS.md (2026-07-24)')

# ── 保存 ──
out_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    'docs',
    '2026年EMS应届校招市场调研_学习计划对照分析报告.docx'
)
doc.save(out_path)
print('Report saved:', out_path)
