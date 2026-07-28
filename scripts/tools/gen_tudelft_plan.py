#!/usr/bin/env python3
"""生成《TU Delft 数据集验证 — 实施任务书》DOCX"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1a, 0x47, 0x8a)

def heading(s, level=1):
    return doc.add_heading(s, level=level)

def para(s, bold=False):
    p = doc.add_paragraph(s)
    if bold:
        for r in p.runs:
            r.bold = True
    return p

def bullet(s):
    doc.add_paragraph(s, style='List Bullet')

def numbered(s):
    doc.add_paragraph(s, style='List Number')

# ═══════════════ 封面 ═══════════════
title = doc.add_heading('TU Delft 数据集验证\n实施任务书', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('让项目从"算法实现"升级为"学术基准验证"')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('基准文件：STATUS.md (2026-07-24) | 编制日期：2026-07-27\n预期耗时：5-7天 | 前置条件：Python 3.13+ 环境已就绪').font.size = Pt(9)

doc.add_page_break()

# ═══════════════ 0. 背景 ═══════════════
heading('0. 背景：为什么做这件事', level=1)
para(
    '你现有的 EMS-PLATFORM 项目（DP/ECMS/MPC 四方法对比）内容扎实，但呈现上有两个短板：\n'
    '  1. 只在标准驾驶循环（WLTC/NEDC/CLTC）上跑过，这是"教程级"数据，说服力不够\n'
    '  2. 没有第三方基准可以对照，结果只能"自己跟自己比"\n\n'
    'TU Delft 发表的船舶燃料电池混动系统数据集提供了：\n'
    '  1. 真实航线负荷数据（荷兰-立陶宛航线，172小时实测）\n'
    '  2. 论文已发表的 LPF-EMS（低通滤波能量管理）结果，可以作为对比基准\n'
    '  3. 学术论文的背书效应（面试时提及"我在 TU Delft 的公开数据集上做了验证"）\n\n'
    '完成此验证后，你的项目叙事从：\n'
    '  "我实现了四种EMS算法并在标准工况上做了对比"\n'
    '  升级为：\n'
    '  "我在 TU Delft 发表的船用混动系统基准上验证了我的算法，并与论文结果做了对比"'
)

# ═══════════════ 1. 总览 ═══════════════
doc.add_page_break()
heading('1. 任务总览', level=1)
para('整个验证分为 5 个步骤，建议顺序执行：')

tasks_overview = [
    ['步骤', '任务', '耗时', '产出'],
    ['Step 1', '数据提取 + 缩放', '1天', 'tudelft_load_profile.csv + 可视化'],
    ['Step 2', '跑 DP（全局最优基准）', '半天', 'DP 在实船数据上的氢耗/SOC结果'],
    ['Step 3', '跑 ECMS + MPC', '1天', 'ECMS/MPC 在实船数据上的结果'],
    ['Step 4', '跑 MPC+EKF', '1天', '带 SOC 估计的 MPC 结果'],
    ['Step 5', '对比报告 + 叙事包装', '1-2天', '验证报告 DOCX + 对比图 + README 更新'],
]
table = doc.add_table(rows=len(tasks_overview), cols=4)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(tasks_overview):
    for j, val in enumerate(row_data):
        table.rows[i].cells[j].text = val
        for p in table.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.bold = True

para('')
para('总耗时：5-7天（可并行 Step 3 和 Step 4 各减半天）', bold=True)

# ═══════════════ 2. Step 1 ═══════════════
doc.add_page_break()
heading('2. Step 1 — 数据提取 + 缩放（1天）', level=1)

heading('2.1 目标', level=2)
para('从 TU Delft Results_ch3.xlsx 中提取实船负荷功率数据，缩放后生成你的求解器可以直接读入的 CSV 文件。')

heading('2.2 输入文件', level=2)
bullet('datasets/TU_Delft_Ch3_LPF_EMS/Results_and_Plots_ch3/Results_ch3.xlsx')
bullet('Chapter 3_readme.txt（确认数据含义）')

heading('2.3 输出文件', level=2)
bullet('datasets/tudelft_load_profile.csv（2063 行，2 列：time_sec, load_kW）')
bullet('datasets/tudelft_load_profile_scaled.csv（缩放后的版本，适配你的 FC 系统）')
bullet('results/tudelft_load_profile.png（负荷曲线可视化）')

heading('2.4 具体操作', level=2)
para(
    '1) 用 pandas 读 Results_ch3.xlsx，提取 Time (sec) 和 Load Profile (kW) 两列\n'
    '2) 剔除 NaN 行，检查数据连续性\n'
    '3) 缩放：原始数据 max=1809kW，你的系统 FC max=30kW\n'
    '   缩放系数 k = 30/1809 ≈ 0.0166（按 FC 最大功率比例）\n'
    '   或按电池容量比例缩放（需要先确定你的电池参数与原系统的比例）\n'
    '4) 保存缩放到后的 CSV\n'
    '5) 画负荷曲线图（确认缩放后的数据在物理合理范围内）'
)

heading('2.5 注意事项', level=2)
bullet('原始数据采样间隔 300s（5分钟），比你的 WLTC（1s）粗糙很多')
bullet('172 小时 / 300s = 2064 个时间步，计算量不大')
bullet('缩放时保持负荷曲线的相对形状不变')

# ═══════════════ 3. Step 2 ═══════════════
heading('3. Step 2 — DP 全局最优基准（半天）', level=1)

heading('3.1 目标', level=2)
para('在你的 DP 求解器上换用 TU Delft 负荷数据，计算理论最优氢耗，作为后续所有方法对比的基准。')

heading('3.2 输入', level=2)
bullet('tudelft_load_profile_scaled.csv（Step 1 产出）')
bullet('scripts/day8_dp_ems.py（不需要改算法逻辑）')

heading('3.3 输出', level=2)
bullet('results/tudelft_dp_results.csv（时间、SOC、FC功率、氢耗等）')
bullet('results/tudelft_dp_profile.png（DP 最优控制轨迹图）')
bullet('数值结果：总氢耗(kg)、SOC初值/终值、FC平均效率')

heading('3.4 注意事项', level=2)
bullet('DP 的状态空间网格需要适应新的负荷范围（你的车用30kW FC 可能偏小，需确认缩放后负荷是否匹配）')
bullet('如果缩放后的负荷峰值超过 FC+电池联合供电能力，需做限幅处理')
bullet('建议先缩放到 FC max=30kW 跑一版，再缩放到 FC max=50kW 跑第二版做敏感性分析')

# ═══════════════ 4. Step 3 ═══════════════
heading('4. Step 3 — ECMS + MPC 对比（1天）', level=1)

heading('4.1 目标', level=2)
para('在相同数据上跑 ECMS 和 MPC（优化版），与 DP 基准对比。')

heading('4.2 输入', level=2)
bullet('tudelft_load_profile_scaled.csv')
bullet('scripts/day9_ecms_ems.py（ECMS 求解器）')
bullet('scripts/mpc_ems_optimized.py（MPC 优化版求解器）')

heading('4.3 输出', level=2)
bullet('results/tudelft_ecms_results.csv')
bullet('results/tudelft_mpc_results.csv')
bullet('results/tudelft_threeway_compare.png（DP vs ECMS vs MPC 三方法对比图）')
bullet('results/tudelft_comparison_metrics.csv（统一指标表）')

heading('4.4 ECMS 参数注意事项', level=2)
bullet('等效因子 s 需要重新校准（你的 s≈130 是基于 WLTC 的，船用负荷特性不同）')
bullet('建议：s 参数扫描 50~200 找最优值')
bullet('或者用 DP 反推标定法（你的 calibrate_s_from_dp.py 可以直接用）')

heading('4.5 MPC 参数注意事项', level=2)
bullet('预测时域 N_p 需要重新选择（采样间隔 300s，N_p=10 就覆盖 3000s）')
bullet('建议扫 N_p=[5, 10, 20, 50] 看收敛趋势')
bullet('SOC 约束参数需根据缩放后的电池容量重新设定')

# ═══════════════ 5. Step 4 ═══════════════
heading('5. Step 4 — MPC+EKF SOC 估计（1天）', level=1)

heading('5.1 目标', level=2)
para('在 MPC 基础上开启 EKF SOC 估计，验证状态估计算法在实船数据上的鲁棒性。')

heading('5.2 输入', level=2)
bullet('tudelft_load_profile_scaled.csv')
bullet('scripts/mpc_ems_ekf.py（MPC+EKF 集成求解器，1017 行）')

heading('5.3 输出', level=2)
bullet('results/tudelft_mpc_ekf_results.csv（含 SOC 真实值/估计值/开环值三路追踪）')
bullet('results/tudelft_soc_estimation.png（SOC 估计精度对比图）')
bullet('results/tudelft_ekf_metrics.csv（RMSE、终点误差等）')

heading('5.4 预期结果', level=2)
bullet('在 WLTC 上 SOC_RMSE=0.0024，在实船数据上预计 RMSE 也在相似量级')
bullet('如果实船数据噪声特性不同，RMSE 可能略高，但 EKF 应显著优于开环')

# ═══════════════ 6. Step 5 ═══════════════
heading('6. Step 5 — 对比报告 + 叙事包装（1-2天）', level=1)

heading('6.1 目标', level=2)
para('将所有结果整理成可展示的验证报告，并更新项目 README。')

heading('6.2 产出物清单', level=2)

heading('6.2.1 验证报告 DOCX', level=3)
bullet('docs/TU_Delft_验证报告.docx')
bullet('含：背景、数据说明、方法、结果总表、工况图、关键发现')
bullet('参考 docs/Week8_四方法大对比报告.docx 的格式')

heading('6.2.2 更新项目 README', level=3)
bullet('在 README 的"结果"章节增加 TU Delft 验证部分')
bullet('对比表：DP / ECMS / MPC / MPC+EKF vs 论文 LPF-EMS 基准')

heading('6.2.3 更新简历项目叙事', level=3)
para('版本A（EMS/BMS方向）增添一句：')
para(
    '"在 TU Delft 公开的船舶 FC 混动系统数据集上验证，'
    'ECMS 氢耗仅比 DP 全局最优高 X%，MPC+EKF SOC 估计 RMSE<0.3%"',
    bold=False
)
para('版本B（RL/控制方向）增添一句：')
para(
    '"EMS 算法平台在学术基准数据上完成第三方验证，'
    '并扩展了 EKF 状态估计模块，适用于实际工程场景"',
    bold=False
)

# ═══════════════ 7. 脚本框架 ═══════════════
heading('7. 脚本框架', level=1)
para('生成的 validate_tudelft.py 脚本建议按以下结构组织：', bold=True)

heading('7.1 脚本结构', level=2)
para(
    'validate_tudelft.py\n'
    '├── main()                          # 入口，解析参数\n'
    '├── step1_extract_load_profile()    # 提取+缩放+保存\n'
    '├── step2_run_dp()                  # 跑 DP 基准\n'
    '├── step3_run_ecms()                # 跑 ECMS\n'
    '├── step3_run_mpc()                 # 跑 MPC\n'
    '├── step4_run_mpc_ekf()             # 跑 MPC+EKF\n'
    '├── step5_compare()                 # 汇总对比\n'
    '├── step5_plot_comparison()         # 出图\n'
    '└── step5_generate_report()         # 生成报告'
)

heading('7.2 参数化运行', level=2)
para(
    '建议让脚本支持命令行参数，方便逐步调试：\n\n'
    '  python scripts/validate_tudelft.py --step 1        # 只做数据提取\n'
    '  python scripts/validate_tudelft.py --step 2        # 只跑 DP\n'
    '  python scripts/validate_tudelft.py --step 3 --method ecms  # 只跑 ECMS\n'
    '  python scripts/validate_tudelft.py --all           # 全部跑一遍\n'
    '  python scripts/validate_tudelft.py --step 5        # 只汇总出报告'
)

heading('7.3 关键参数（可配置）', level=2)
para(
    'scale_factor = 30/1809  # 缩放系数，可按需调整\n'
    's_range = (50, 200)     # ECMS 等效因子扫描范围\n'
    'N_p_list = [5, 10, 20, 50]  # MPC 预测时域\n'
    'sim_dt = 300            # 采样间隔（秒），与原始数据一致'
)

# ═══════════════ 8. 时间线 ═══════════════
heading('8. 时间线', level=1)

table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Light Grid Accent 1'
headers2 = ['天次', '任务', '备注']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

data2 = [
    ['第1天', 'Step 1: 数据提取+缩放\n+ Step 2: DP 基准', '数据提取半天，DP半天，合在一天做'],
    ['第2天', 'Step 3: ECMS + MPC', 'ECMS 参数扫描 + MPC 对比'],
    ['第3天', 'Step 4: MPC+EKF', 'SOC 估计验证'],
    ['第4天', 'Step 5: 对比报告', '数据汇总+出图+报告撰写'],
    ['第5天', 'Step 5 续: 叙事包装', 'README 更新 + 简历叙事打磨'],
]
for i, row_data in enumerate(data2):
    for j, val in enumerate(row_data):
        table2.rows[i+1].cells[j].text = val
        for p in table2.rows[i+1].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

para('')
para('如果时间紧张，可以合并到第3天做完。最低可行方案（3天）：', bold=True)
bullet('第1天：数据提取+DP+ECMS（并行走）')
bullet('第2天：MPC+MPC/EKF（并行走）')
bullet('第3天：报告+叙事包装')

# ═══════════════ 9. 面试话术 ═══════════════
heading('9. 面试话术速查', level=1)

heading('9.1 一句话版（30秒）', level=2)
para(
    '"我的EMS算法平台在TU Delft发表的船舶混动系统数据集上做了验证，'
    '结果与论文基准吻合，DP最优氢耗验证了算法正确性。"'
)

heading('9.2 两分钟版（技术面）', level=2)
para(
    '"我的EMS算法对比平台之前是在标准驾驶循环上跑的。'
    '为了验证它的通用性，我从TU Delft的一篇论文里找到了他们的公开数据——'
    '那是一艘滚装船的燃料电池混动系统，有荷兰-立陶宛航线的实测功率负荷，'
    '总共172小时的运行数据。'
    '我把负荷曲线提取出来、缩放到我系统的功率等级，然后用我的DP、ECMS、MPC重新跑了一遍。'
    '结果：DP作为全局最优给出了氢耗下限，ECMS只比DP高X%，MPC在SOC约束下表现稳定。'
    '而且我加了EKF SOC估计后，即使在这个完全不同的工况上，SOC跟踪误差仍然在0.3%以内。'
    '这说明我的算法平台不局限于汽车工况，具有跨场景的通用性。"'
)

heading('9.3 追问准备', level=2)
bullet('问："为什么选TU Delft的数据？" → "他们的论文有公开的模型和数据，结果可复现、可对比，这比我自己造数据更有说服力"')
bullet('问："缩放会不会引入偏差？" → "保持了负荷的相对形状，只是幅度按功率等级调整。如果时间允许，可以在多个缩放系数下做敏感性分析"')
bullet('问："跟论文的LPF-EMS比差多少？" → "LPF是简单的低通滤波，DP理论上限肯定优于它，ECMS接近DP。这个对比正好说明更复杂的算法能带来多少收益"')

# ═══════════════ 10. 交付物清单 ═══════════════
heading('10. 完整交付物清单', level=1)

heading('10.1 数据文件', level=2)
bullet('datasets/tudelft_load_profile.csv')
bullet('datasets/tudelft_load_profile_scaled.csv')

heading('10.2 脚本', level=2)
bullet('scripts/validate_tudelft.py（主脚本）')

heading('10.3 结果文件', level=2)
bullet('results/tudelft_load_profile.png')
bullet('results/tudelft_dp_results.csv / .png')
bullet('results/tudelft_ecms_results.csv')
bullet('results/tudelft_mpc_results.csv')
bullet('results/tudelft_mpc_ekf_results.csv')
bullet('results/tudelft_threeway_compare.png')
bullet('results/tudelft_soc_estimation.png')
bullet('results/tudelft_comparison_metrics.csv')
bullet('results/tudelft_ekf_metrics.csv')

heading('10.4 文档', level=2)
bullet('docs/TU_Delft_验证报告.docx')

heading('10.5 项目文档更新', level=2)
bullet('README.md 增加 TU Delft 验证章节')
bullet('简历项目叙事更新（两个版本）')

# ── 结尾 ──
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('\n—— 任务书完 ——').bold = True
para('编制日期：2026-07-27 | 基于 STATUS.md (2026-07-24) 学习计划')

# ── 保存 ──
out_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    'docs',
    'TU_Delft_验证_实施任务书.docx'
)
doc.save(out_path)
print('Done:', out_path)
