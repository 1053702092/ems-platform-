#!/usr/bin/env python3
"""
《EMS 仿真器 — 三大实验详细设计书》
实验1: 参数敏感性分析
实验2: 传感器故障鲁棒性
实验3: 实时性基准测试
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1a, 0x47, 0x8a)

def heading(s, level=1):
    return doc.add_heading(s, level=level)

def para(s, bold=False):
    p = doc.add_paragraph(s)
    if bold:
        for r in p.runs:
            r.bold = True
    return p

def bullet(s):
    doc.add_paragraph(s, style='List Bullet')

def code(s):
    p = doc.add_paragraph()
    r = p.add_run(s)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
    for rd in rows:
        r = t.add_row()
        for i, v in enumerate(rd):
            r.cells[i].text = str(v)
            for p in r.cells[i].paragraphs:
                for rr in p.runs:
                    rr.font.size = Pt(9)
    return t

# ═══════════════ 封面 ═══════════════
doc.add_heading('EMS 仿真器\n三大实验详细设计书', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('参数敏感性 + 传感器故障鲁棒性 + 实时性基准测试')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('编制日期：2026-07-27\n前置条件：EMS 仿真器骨架已完成（见 EMS仿真器_实施任务书）').font.size = Pt(9)

doc.add_page_break()

# ═══════════════ 总览 ═══════════════
heading('总览', level=1)
para('本任务书包含三个独立的实验，在 EMS 仿真器骨架完成后依次执行。')

add_table(['实验', '耗时', '验证什么', '简历价值'],
    [
        ['实验1：参数敏感性', '2天', 'ECMS s因子/MPC N_p/DP网格\n对结果的影响程度', '找到最优参数区间\n确定工程安全边界'],
        ['实验2：传感器故障鲁棒性', '2天', 'EKF 在有传感器故障/噪声时\n能否维持SOC估计精度', '实际系统一定有噪声\n能处理才是工程能力'],
        ['实验3：实时性基准测试', '1天', '每种算法的单步耗时\n能否满足实时控制要求', '算法不只比精度\n还要比算得快'],
    ])
para('')
para('总耗时：5天（可并行实验1和实验2各减1天）', bold=True)

# ═══════════════════════════════════════
# 实验1
# ═══════════════════════════════════════
doc.add_page_break()
heading('实验1：参数敏感性分析', level=1)

heading('1.1 背景', level=2)
para(
    '任何EMS算法都有"调参"这一步——面试官一定会问："你这个参数怎么确定的？"\n'
    '参数敏感性分析回答的是："参数变了多少，结果会差多少？"\n'
    '这证明你不只是调出来了，而是理解参数的作用边界。'
)

heading('1.2 三个子实验', level=2)

# ── 1.2.1 ECMS ──
heading('1.2.1 ECMS 等效因子 s 扫描', level=3)
para('目标：', bold=True)
para('找出 s 的最优区间，以及 s 偏大/偏小分别有什么后果。')

para('输入：', bold=True)
bullet('算法：ECMS（恒等效因子版本）')
bullet('工况：WLTC（基准），NEDC/CLTC（验证）')
bullet('参数范围：s = 50 ~ 250，步长 10，共 21 组')

para('步骤：', bold=True)
code(
    '# 伪代码\n'
    'results = []\n'
    'for s in range(50, 260, 10):\n'
    '    metrics = run_simulation("ecms", cycle="wltc", params={"s": s})\n'
    '    results.append({"s": s, "H2": metrics.H2, "SOC_end": metrics.SOC_end})\n\n'
    '# 绘图\n'
    'plot(results.s, results.H2, xlabel="s", ylabel="H2 consumption (kg)")\n'
    'plot(results.s, results.SOC_end, xlabel="s", ylabel="SOC_end")\n'
    '→ 确认 s 与氢耗呈 U 型曲线，SOC_end 呈单调关系'
)

para('预期结果：', bold=True)
bullet('s 过小（<80）：SOC 持续下降 → 终点 SOC 低于下限 → 氢耗虚低（电池透支）')
bullet('s 最优（120~140）：SOC 持平 → 氢耗真实最低')
bullet('s 过大（>180）：优先用 FC → SOC 上升，但氢耗也上升（FC 在低效区运行多）')
bullet('s 偏差 ±10%，氢耗波动 < 2% → 算法对参数不敏感，说明工程鲁棒性好')

para('产出文件：', bold=True)
bullet('results/exp1_ecms_s_scan_wltc.csv（21行 × 5列）')
bullet('results/exp1_ecms_s_scan_wltc.png（双轴图：H2 + SOC_end vs s）')
bullet('results/exp1_ecms_s_scan_summary.txt（最优s区间 + 结论）')

heading('1.2.2 MPC 预测时域 N_p 扫描', level=3)
para('目标：', bold=True)
para('N_p 设多长够用？太短会怎样？太长值得吗？')

para('输入：', bold=True)
bullet('算法：MPC（优化版）')
bullet('工况：WLTC（基准），NEDC/CLTC（验证）')
bullet('参数范围：N_p = 5, 10, 20, 30, 50, 80, 120, 200，共 8 组')

para('步骤：', bold=True)
code(
    '# 伪代码\n'
    'results = []\n'
    'for Np in [5, 10, 20, 30, 50, 80, 120, 200]:\n'
    '    metrics = run_simulation("mpc", cycle="wltc", params={"N_p": Np})\n'
    '    results.append({"N_p": Np, "H2": metrics.H2, "SOC_end": metrics.SOC_end,\n'
    '                     "time_ms": metrics.compute_time_ms})\n\n'
    '# 绘图\n'
    'plot(results.N_p, results.H2, xlabel="N_p", ylabel="H2 (kg)")\n'
    'plot(results.N_p, results.time_ms, xlabel="N_p", ylabel="Compute time (ms)")\n'
    '→ 找"拐点"：N_p 超过多少后氢耗不再下降'
)

para('预期结果：', bold=True)
bullet('N_p < 20：氢耗波动大（看不远，决策近视）')
bullet('N_p = 30~50：氢耗收敛到接近 DP 的水平')
bullet('N_p > 50：边际收益递减，但计算时间线性增加')
bullet('拐点≈50 → "对于WLTC工况，N_p=50是性价比最优的选择"')

para('产出文件：', bold=True)
bullet('results/exp1_mpc_np_scan_wltc.csv（8行 × 4列）')
bullet('results/exp1_mpc_np_scan_wltc.png（双Y轴：H2 + compute_time vs N_p）')
bullet('results/exp1_mpc_np_scan_convergence.png（氢耗收敛曲线 + 标注最优区间）')

heading('1.2.3 DP 状态网格密度扫描', level=3)
para('目标：', bold=True)
para('DP 的 SOC 网格分多细才够？网格太粗精度差，太细算不动。')

para('输入：', bold=True)
bullet('算法：DP')
bullet('工况：WLTC')
bullet('参数范围：SOC 网格 = 30, 60, 100, 150, 200, 300，共 6 组')

para('步骤：', bold=True)
code(
    '# 伪代码\n'
    'results = []\n'
    'for n_grid in [30, 60, 100, 150, 200, 300]:\n'
    '    metrics = run_simulation("dp", cycle="wltc", params={"n_grid_soc": n_grid})\n'
    '    results.append({"n_grid": n_grid, "H2": metrics.H2,\n'
    '                     "SOC_end": metrics.SOC_end, "time_ms": metrics.compute_time_ms})\n\n'
    '# 绘图\n'
    'plot(results.n_grid, results.H2, marker="o")\n'
    '→ 找"收敛网格密度"：超过多少后氢耗不再变化'
)

para('预期结果：', bold=True)
bullet('n_grid < 60：SOC 插值误差大，氢耗不稳定')
bullet('n_grid = 100~150：氢耗收敛，与 300 的差异 < 0.5%')
bullet('n_grid = 300：精度没有明显提高，但计算时间成倍增加')
bullet('→ "150 格是精度和速度的最佳平衡点"')

para('产出文件：', bold=True)
bullet('results/exp1_dp_grid_scan_wltc.csv（6行 × 4列）')
bullet('results/exp1_dp_grid_scan_wltc.png')

heading('1.3 参数敏感性总结表', level=2)
add_table(['算法', '参数', '范围', '最优区间', '敏感性'],
    [
        ['ECMS', '等效因子 s', '50~250', '120~140', '低：±10%内氢耗波动<2%'],
        ['MPC', '预测时域 N_p', '5~200', '30~50', '中：N_p<20 时结果不稳定'],
        ['DP', 'SOC 网格密度', '30~300', '100~150', '低：>100 后氢耗已收敛'],
    ])

para('')
para('面试话术：', bold=True)
para(
    '"我对三种算法分别做了参数敏感性分析。ECMS的等效因子s在120~140之间时氢耗最低，'
    '而且s偏差±10%氢耗波动不到2%，说明工程上不需要精确标定。'
    'MPC的预测时域N_p到50之后氢耗就收敛了，再增加时域只增加计算量不改善效果。'
    'DP的SOC网格到150格也收敛了。'
    '这些分析帮我确定了每组参数的安全边界和最佳取值区间。"'
)

# ═══════════════════════════════════════
# 实验2
# ═══════════════════════════════════════
doc.add_page_break()
heading('实验2：传感器故障鲁棒性', level=1)

heading('2.1 背景', level=2)
para(
    '实际 EMS 系统中，传感器信号不可能完美——有噪声、有偏置、甚至可能完全丢失。\n'
    '面试官一定会问："你的EKF在真实场景下到底能用吗？"\n'
    '这个实验就是回答这个问题的。\n\n'
    '你的 mpc_ems_ekf.py 已经有传感器噪声模拟能力了，'
    '但这个实验把它**专门拎出来作为独立的验证场景**。'
)

heading('2.2 核心思路', level=2)
para('在不同类型的传感器故障下，对比"有EKF"和"无EKF"的SOC估计精度。')

add_table(['传感器故障场景', '严重程度', '实际对应情况'],
    [
        ['偏置故障', '0.5A / 2A / 5A', '电流传感器零点漂移'],
        ['噪声', 'σ=0.1A / 0.5A / 1.0A', '传感器本身的测量噪声'],
        ['混合场景', '偏置2A + 噪声σ=0.5A', '真实传感器最常见的状态'],
    ])

heading('2.3 子实验2.1：电流传感器偏置', level=3)
para('设置：', bold=True)
bullet('EKF 模式和 Open Loop 模式并列运行')
bullet('给电流测量值施加固定偏置：0A（基准）, 0.5A, 2A, 5A')
bullet('运行整个 WLTC 工况，记录每一步的 SOC 估计误差')

para('伪代码：', bold=True)
code(
    'faults = {"bias": [0, 0.5, 2, 5]}  # 单位：安培\n\n'
    'for bias in faults["bias"]:\n'
    '    # 开环 SOC（无EKF）\n'
    '    metrics_ol = run_simulation("mpc", cycle="wltc",\n'
    '                                params={"sensor_bias": bias, "use_ekf": False})\n'
    '    # EKF SOC\n'
    '    metrics_ekf = run_simulation("mpc", cycle="wltc",\n'
    '                                 params={"sensor_bias": bias, "use_ekf": True})\n\n'
    '    record(bias, "open_loop", metrics_ol.SOC_RMSE, metrics_ol.SOC_end_error)\n'
    '    record(bias, "ekf", metrics_ekf.SOC_RMSE, metrics_ekf.SOC_end_error)'
)

para('预期结果：', bold=True)

add_table(['偏置(A)', '开环 SOC_RMSE', 'EKF SOC_RMSE', '改进倍数'],
    [
        ['0（基准）', '~0.003', '~0.002', '1.5×'],
        ['0.5', '~0.005', '~0.002', '2.5×'],
        ['2', '~0.012', '~0.002', '6×'],
        ['5', '~0.030', '~0.003', '10×'],
    ])

para('')
bullet('EKF 在所有偏置场景下都能将 SOC_RMSE 维持在 ~0.003 以下')
bullet('开环 SOC 误差与偏置量成正比（积分累积效应）')
bullet('EKF 因为有电压观测值修正，偏置误差被校准掉')

para('产出文件：', bold=True)
bullet('results/exp2_bias_sensitivity.csv（8行 × 4列）')
bullet('results/exp2_bias_comparison.png（柱状图：开环 vs EKF，按偏置分组）')
bullet('results/exp2_bias_trajectory_bias2A.png（SOC轨迹对比图，偏置2A场景）')

heading('2.4 子实验2.2：电流传感器噪声', level=3)
para('设置：', bold=True)
bullet('给电流测量值叠加高斯噪声：σ=0（基准）, 0.1A, 0.5A, 1.0A')
bullet('运行 WLTC，对比开环和 EKF 的 SOC 估计')

para('伪代码：', bold=True)
code(
    'faults = {"noise_std": [0, 0.1, 0.5, 1.0]}\n\n'
    'for noise_std in faults["noise_std"]:\n'
    '    metrics_ol = run_simulation("mpc", cycle="wltc",\n'
    '                                params={"noise_std": noise_std, "use_ekf": False})\n'
    '    metrics_ekf = run_simulation("mpc", cycle="wltc",\n'
    '                                 params={"noise_std": noise_std, "use_ekf": True})\n\n'
    '    result = compare(metrics_ol, metrics_ekf)\n'
    '    # EKF 应该稳定，开环会波动但无累积误差'
)

para('预期结果：', bold=True)
bullet('开环 SOC：噪声导致轨迹波动，但无系统性偏差 → RMSE 与噪声 std 成正比')
bullet('EKF SOC：卡尔曼增益自适应调整 → RMSE 稳定，对噪声不敏感')
bullet('关键结论：EKF 在高噪声场景下的优势体现在"平滑性"而非"准确性"（开环没累积误差）')

para('产出文件：', bold=True)
bullet('results/exp2_noise_sensitivity.csv')
bullet('results/exp2_noise_trajectory_std1.png（σ=1.0 场景的轨迹对比）')

heading('2.5 子实验2.3：混合场景（最真实的测试）', level=3)
para('设置：', bold=True)
bullet('同时施加偏置 2A + 噪声 σ=0.5A → 模拟真实传感器的状态')
bullet('跑三个工况（WLTC / NEDC / CLTC）验证一致性')

para('伪代码：', bold=True)
code(
    'for cycle in ["wltc", "nedc", "cltc"]:\n'
    '    metrics_ol = run_simulation("mpc", cycle=cycle,\n'
    '                                params={"sensor_bias": 2, "noise_std": 0.5,\n'
    '                                        "use_ekf": False})\n'
    '    metrics_ekf = run_simulation("mpc_ekf", cycle=cycle,\n'
    '                                 params={"sensor_bias": 2, "noise_std": 0.5})\n'
    '    ol_rmse = metrics_ol.SOC_RMSE\n'
    '    ekf_rmse = metrics_ekf.SOC_RMSE\n'
    '    impr = ol_rmse / ekf_rmse\n'
    '    print(f"{cycle}: OL_RMSE={ol_rmse:.4f}, EKF_RMSE={ekf_rmse:.4f}, {impr:.1f}x")'
)

para('预期结果：', bold=True)
bullet('WLTC/NEDC/CLTC 三工况下，EKF 的 SOC_RMSE 均稳定在 0.003 左右')
bullet('改进倍数：3-5 倍（取决于工况的动态程度）')
bullet('证明 EKF 具有跨工况鲁棒性，不是只在某个工况下好用')

para('产出文件：', bold=True)
bullet('results/exp2_hybrid_scenario_summary.csv（3行（三工况）× 5列）')
bullet('results/exp2_hybrid_comparison.png（分群的对比柱状图）')

heading('2.6 实验2总结表', level=2)
add_table(['场景', '开环 SOC_RMSE', 'EKF SOC_RMSE', '改进倍数', '说明'],
    [
        ['偏置2A', '~0.012', '~0.002', '6×', 'EKF 用电压观测修正电流积分误差'],
        ['噪声σ=1.0', '~0.008', '~0.003', '2.7×', '开环无累积误差，但轨迹抖动'],
        ['偏置2A+噪声0.5A', '~0.015', '~0.003', '5×', '最接近真实传感器情况'],
    ])

para('')
para('面试话术：', bold=True)
para(
    '"我专门测了EKF在传感器故障下的表现。电流偏置2A时，开环SOC误差会累积到1.2%，'
    '但EKF因为有端电压观测来修正，误差不到0.3%。'
    '在偏置+噪声的混合场景下，EKF在三工况上SOC_RMSE都稳定在0.3%以内，'
    '比开环提高了5倍。'
    '这说明即使传感器不是完美的，EMS系统也能可靠运行——这对实际工程很有意义。"'
)

# ═══════════════════════════════════════
# 实验3
# ═══════════════════════════════════════
doc.add_page_break()
heading('实验3：实时性基准测试', level=1)

heading('3.1 背景', level=2)
para(
    '所有 EMS 算法最终都要跑在真实的控制器上，控制器的采样间隔通常是 0.1s~1s。\n'
    '如果算法每步决策时间 > 采样间隔，就来不及算出结果，控制会失效。\n'
    '这个实验回答：你的算法在实际控制器的时间约束下跑得动吗？'
)

heading('3.2 测试方法', level=2)
para('工具：', bold=True)
bullet('Python time.perf_counter() 或 timeit 模块')
bullet('测的是"单次控制决策"的时间，不是整个仿真时间')
bullet('每种算法跑 100 次取中位数（排除单次抖动）')

para('伪代码：', bold=True)
code(
    'import time\n\n'
    'def benchmark_one_step(engine, P_load, SoC):\n'
    '    """测单步决策时间，重复100次取中位数"""\n'
    '    times = []\n'
    '    for _ in range(100):\n'
    '        t0 = time.perf_counter()\n'
    '        action = engine.step(t=0, P_load=P_load, SoC=SoC)\n'
    '        t1 = time.perf_counter()\n'
    '        times.append((t1 - t0) * 1000)  # 转为毫秒\n'
    '    return {"median_ms": np.median(times),\n'
    '            "p95_ms": np.percentile(times, 95),\n'
    '            "p99_ms": np.percentile(times, 99)}'
)

heading('3.3 测试场景设计', level=2)
para('只在一种状态下测不够全面，要覆盖几种典型情况：')

add_table(['场景', 'P_load', 'SoC', '模拟什么情况'],
    [
        ['低负荷', '5 kW', '0.60', '怠速/滑行'],
        ['中负荷', '15 kW', '0.50', '匀速巡航'],
        ['高负荷', '28 kW', '0.40', '急加速/爬坡'],
        ['边界（SoC低）', '15 kW', '0.25', '电池快没电，算法需做保护决策'],
    ])

para('')
para('对每个场景，每种算法都跑 100 次取统计值。')

heading('3.4 四种算法的预期结果', level=2)

add_table(['算法', '单步时间(中位数)', 'p99(最大)', '瓶颈在哪'],
    [
        ['DP', 'NA（离线批量）', 'NA', 'DP 不是在线算法，不参与单步比较'],
        ['ECMS（网格搜索）', '~5-20 ms', '~30 ms', '暴力搜索所有 FC 功率候选值'],
        ['MPC（优化版）', '~20-100 ms', '~200 ms', '滚动优化 + SOC 软约束求解'],
        ['MPC+EKF', '~21-101 ms', '~201 ms', 'MPC + EKF 传播+更新'],
        ['PPO（完成后测）', '<1 ms', '<2 ms', '前向推理，无优化循环'],
    ])

para('')
para('判断标准：', bold=True)

add_table(['采样间隔', '典型场景', '哪些算法满足'],
    [
        ['1.0 s', '整车控制器 VCU', '全部满足'],
        ['0.1 s', 'BMS 内部循环', 'ECMS / PPO 满足，MPC 可能过载'],
        ['0.01 s', '底层电流环', '只有 PPO 满足'],
    ])

heading('3.5 测试执行计划', level=2)

heading('第1步：编写基准测试脚本', level=3)
bullet('新建 scripts/simulator/benchmark.py')
bullet('对每种算法、每个场景，跑100步取统计值')
bullet('结果输出到 results/exp3_benchmark_results.csv')

heading('第2步：绘制对比图', level=3)
bullet('results/exp3_benchmark_bar.png — 柱状图：不同场景下各算法单步时间')
bullet('results/exp3_benchmark_cdf.png — CDF 图：各算法的延时分布')
bullet('标注 0.1s 和 1s 的实时性红线')

heading('第3步：三工况重复验证', level=3)
bullet('在 WLTC/NEDC/CLTC 全工况跑一遍，测平均单步时间和总仿真时间')
bullet('验证：即使工况不同，单步时间基本一致（不受负荷值影响）')

heading('3.6 产出文件', level=2)
bullet('scripts/simulator/benchmark.py（基准测试脚本）')
bullet('results/exp3_benchmark_results.csv（4算法 × 4场景 × 3个统计值 = 48行）')
bullet('results/exp3_benchmark_bar.png')
bullet('results/exp3_benchmark_cdf.png')
bullet('results/exp3_full_cycle_time.csv（三工况 × 4算法 × 总仿真时间）')
bullet('results/exp3_realtime_summary.txt（判断结论）')

heading('3.7 面试话术', level=2)
para('一句话版：', bold=True)
para(
    '"ECMS单步5ms、MPC单步50ms，都满足1秒采样间隔的实时性要求。PPO单步不到1ms，'
    '甚至可以跑在更快的控制环上。"'
)
para('')
para('深入版：', bold=True)
para(
    '"实时性测试是很多人忽略的。我专门测了四种算法在四个典型工况点的计算耗时，'
    '每种跑100次取统计值。结果：ECMS中位数5ms、MPC 50ms、PPO<1ms，'
    '全部满足1秒采样间隔的要求。PPO甚至可以在0.1秒的BMS内环运行。'
    '同时我发现MPC的p99时间接近200ms，说明偶尔会出现较长的求解延时，'
    '如果要部署到实车，需要增加看门狗机制。——这个问题就是做实验才发现的。"'
)

# ═══════════════════════════════════════
# 执行计划
# ═══════════════════════════════════════
doc.add_page_break()
heading('执行计划总表', level=1)

add_table(['周次', '任务', '具体内容'],
    [
        ['第1-2天', '实验1：参数敏感性', 'ECMS s扫描 + MPC N_p扫描 + DP网格扫描'],
        ['第3-4天', '实验2：传感器故障鲁棒性', '偏置测试 + 噪声测试 + 混合场景测试'],
        ['第5天', '实验3：实时性基准测试', '编写benchmark.py + 四种算法逐一测试'],
        ['并行', '所有实验的结果图', '每次运行自动保存 CSV + PNG 到 results/exp*'],
        ['最后', 'README 更新', '三个实验各写一段结论 + 关键图'],
    ])

para('')
para('产出文件总清单：', bold=True)
bullet('results/exp1_ecms_s_scan_wltc.csv + .png')
bullet('results/exp1_mpc_np_scan_wltc.csv + .png')
bullet('results/exp1_dp_grid_scan_wltc.csv + .png')
bullet('results/exp2_bias_sensitivity.csv + .png')
bullet('results/exp2_noise_sensitivity.csv + .png')
bullet('results/exp2_hybrid_scenario_summary.csv + .png')
bullet('results/exp3_benchmark_results.csv + bar.png + cdf.png')
bullet('results/exp3_full_cycle_time.csv')
bullet('合计约 18 个文件（含 CSV 和 PNG）')

# ── 结尾 ──
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('\n—— 任务书完 ——').bold = True
para('编制日期：2026-07-27')

out_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    'docs',
    'EMS三大实验详细设计书.docx'
)
doc.save(out_path)
print('Done:', out_path)
