#!/usr/bin/env python3
"""生成《公开实船功率/能耗数据集调研报告》.docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 全局样式 ──────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 封面 / 标题 ──────────────────────────────────────
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('公开实船功率／能耗数据集调研报告')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('——面向船舶能源管理系统（EMS）算法训练与验证')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'\n生成日期：{datetime.date.today().isoformat()}\n'
                     f'研究项目：深海多堆FC系统 · 能源管理平台')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ── 目录占位 ─────────────────────────────────────────
toc_heading = doc.add_heading('目  录', level=1)
toc_items = [
    '一、实测船舶功率时间序列数据集',
    '二、燃料电池／混合动力船舶专题数据集',
    '三、基于 AIS 的功率重构数据集',
    '四、学术论文附带实验数据',
    '五、船级社／行业工具',
    '六、其他学术论文引用数据集一览',
    '七、推荐优先级与使用建议',
    '附录：数据源链接汇总',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.space_after = Pt(4)
doc.add_page_break()

# ── 辅助函数 ─────────────────────────────────────────
def add_table(doc, headers, rows):
    """添加带格式的表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def add_star_rating(n):
    return '⭐' * n + '☆' * (5 - n)

def severity_badge(text, color):
    """彩色标签"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)
    return p

# ══════════════════════════════════════════════════════
# 1. 实测船舶功率时间序列
# ══════════════════════════════════════════════════════
doc.add_heading('一、实测船舶功率时间序列数据集', level=1)

doc.add_heading('1.1  Shifts Marine Cargo Vessel 功率预测数据集', level=2)
p = doc.add_paragraph()
run = p.add_run('推荐指数：')
run.bold = True
run = p.add_run(add_star_rating(5))
run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)

add_table(doc,
    ['项目', '内容'],
    [
        ['平台', 'Zenodo (DOI: 10.5281/zenodo.7057666)'],
        ['数据规模', '营运货船连续 4 年传感器数据，每分钟采样'],
        ['关键特征', '主机轴功率(kW)、航速、吃水、风速、波高、距上次进坞时间'],
        ['格式', '表格回归数据，含分布偏移 train/val/eval 划分'],
        ['许可', 'CC BY-NC-SA 4.0'],
        ['关联竞赛', 'Shifts Challenge 2022/2023，175 支队伍参赛'],
        ['论文', 'arXiv:2206.15407 (https://arxiv.org/pdf/2206.15407)'],
        ['备注', '含真实数据 + 物理模型合成数据两部分'],
    ])

doc.add_paragraph('')
doc.add_heading('1.2  中国国家基础科学数据中心 — 船舶能耗评估模型数据集', level=2)
add_table(doc,
    ['项目', '内容'],
    [
        ['标识', 'CSTR: 16666.11.nbsdc.6LlX5R4z'],
        ['链接', 'https://60.245.194.22/general/dataDetail?id=69b6d7a5195d2623ac632d4c&type=1'],
        ['船型', 'VLOC（巴西–中国航线）+ 散货船（1–14 kn 航速范围）'],
        ['参数', '主机功率、航速、油耗率连续时间序列'],
        ['精度', '模型验证误差 < 5%'],
    ])

doc.add_paragraph('')
doc.add_heading('1.3  Kaggle 平台相关数据集', level=2)

p = doc.add_paragraph()
run = p.add_run('(a) Ship Performance Analysis 数据集')
run.bold = True
doc.add_paragraph('2,736 条 × 18 列，含 Engine_Power_kW、航速、吃水、载重吨、效率等。', style='List Bullet')
p = doc.add_paragraph('链接：')
run = p.add_run('https://www.kaggle.com/code/hazemsayedabdullah/ship-performance-analysis/clustering/log')
run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

p = doc.add_paragraph()
run = p.add_run('(b) Preventive Maintenance for Marine Engines')
run.bold = True
doc.add_paragraph('船用发动机仿真传感器数据（温度、油压、油耗、转速、振动），含维护标签。', style='List Bullet')
p = doc.add_paragraph('GitHub：')
run = p.add_run('https://github.com/mlatinov/Preventive-Maintenance-for-Marine-Engines')
run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 2. 燃料电池/混合动力船舶专题
# ══════════════════════════════════════════════════════
doc.add_heading('二、燃料电池／混合动力船舶专题数据集', level=1)

doc.add_heading('2.1  NAUTILUS SOFC + 电池发电机组实测数据', level=2)
p = doc.add_paragraph()
run = p.add_run('推荐指数：')
run.bold = True
run = p.add_run(add_star_rating(5))
run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)

p = doc.add_paragraph()
run = p.add_run('⭐ 这是目前唯一公开的实船燃料电池+电池混合系统实测数据集，对 EMS 验证极有价值。')
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

add_table(doc,
    ['项目', '内容'],
    [
        ['链接', 'https://zenodo.org/records/14643552 (DOI: 10.5281/zenodo.14643552)'],
        ['系统构成', '60 kWe SOFC（SolydEra）+ 40 kWh 锂电池（MAN ES），总额定 80 kWe'],
        ['测试周期', '2024 年 9 月 – 12 月'],
        ['测试负荷', '真实游轮负荷剖面，向电网输送 62.6 MWh'],
        ['文件内容', 'CSV 时间序列 + 索引文件 + Python 分析脚本 dem_analysis.py + 调试报告'],
        ['来源机构', 'DLR、MAN ES、SolydEra、RWTH Aachen、Lund University'],
        ['项目背景', 'EU Horizon 2020 NAUTILUS 项目，目标 5–60 MW 船载系统'],
    ])

doc.add_paragraph('')
doc.add_heading('2.2  TU Delft — SH2IPDRIVE 氢燃料电池船舶系列', level=2)
p = doc.add_paragraph()
run = p.add_run('推荐指数：')
run.bold = True
run = p.add_run(add_star_rating(5))
run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)

add_table(doc,
    ['数据集', 'DOI', '核心内容'],
    [
        ['Ch.3 LPF-EMS', '10.4121/589dc384-b5bf-450b-bba2-c6cd1d33e378',
         'MATLAB/Simulink 低通滤波EMS；FC-电池电动船功率分配；退化评估'],
        ['Ch.4 系统优化', '10.4121/c1f91c19-b975-4dcc-acb5-58316b11f9a9',
         'Python 寿命成本优化；功率剖约简；稳性约束下的系统布置'],
        ['分层控制策略', '10.4121/e9fd5f83-35bc-4745-a2a6-682500b2646b',
         'FC-电池 DC 船舶电力系统分层分散式EMS (IEEE ITEC 2023)'],
    ])

doc.add_paragraph('')
doc.add_heading('2.3  退化感知预测性 EMS（MPC / ECMS / 滤波）', level=2)
add_table(doc,
    ['项目', '内容'],
    [
        ['DOI', '10.4121/2cc57bd5-a76d-4c15-9abe-ed814c70c607'],
        ['链接', 'https://data.4tu.nl/datasets/2cc57bd5-a76d-4c15-9abe-ed814c70c607/1'],
        ['方法', 'MPC / ECMS / 滤波三种 EMS 对比，数据驱动负荷预测 + FC 退化感知'],
        ['备注', '负荷预测数据因保密未公开，但框架和代码完全开放，可替换用户数据'],
    ])

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 3. 基于 AIS 的功率重构
# ══════════════════════════════════════════════════════
doc.add_heading('三、基于 AIS 的功率重构数据集', level=1)

doc.add_heading('3.1  AIS-TSH 数据集', level=2)
add_table(doc,
    ['项目', '内容'],
    [
        ['来源', 'IEEE Data Descriptions (2026)'],
        ['范围', '法罗群岛 Tórshavn 港及周边，547 艘船，780 万条 AIS 消息'],
        ['格式', 'newline-delimited JSON，含射频元数据'],
        ['链接', 'https://zendy.io/title/10.1109%2Fieeedata.2026.3676933'],
    ])

doc.add_heading('3.2  AIS + 实测功率 + 天气 → ML 功率预测', level=2)
p = doc.add_paragraph()
run = p.add_run('论文：')
run.bold = True
run = p.add_run('"Prediction of vessel propulsion power using machine learning on AIS data, '
                 'ship performance measurements and weather data"')
run.italic = True
doc.add_paragraph('AIS + 船上传感器功率 + ECMWF 气象数据，ML 模型 R²=0.78（vs 物理模型 0.48）', style='List Bullet')
p = doc.add_paragraph('链接：')
run = p.add_run('https://iopscience.iop.org/article/10.1088/1742-6596/1357/1/012038')
run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

doc.add_heading('3.3  功率重构方法论', level=2)
doc.add_paragraph('IMO / UMAS 方法：基于船长、船宽、吃水、DWT 的多重回归填补缺失的功率数据', style='List Bullet')
doc.add_paragraph('Admiralty 公式法：W(t) = δ_w / (η_w η_f) · 航速³，结合天气修正和污底修正', style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 4. 学术论文附带实验数据
# ══════════════════════════════════════════════════════
doc.add_heading('四、学术论文附带实验数据', level=1)

doc.add_heading('4.1  NTUA HIPPO-2 + Marine Loading Cycles', level=2)
p = doc.add_paragraph()
run = p.add_run('推荐指数：')
run.bold = True
run = p.add_run(add_star_rating(4))
run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)

add_table(doc,
    ['项目', '内容'],
    [
        ['论文', 'Applied Energy, Vol.307, 2022 (被引 68 次)'],
        ['实验台架', 'HIPPO-2：260 kW 柴油机 + 90 kW 电机 + 28 kWh 电池，并联混合动力'],
        ['核心产出', '从实船数据提取 20 条典型 Marine Loading Cycles（类似汽车行驶工况）'],
        ['EMS 验证', 'NMPC 功率分配 → 节油 6%，NOx 减排 8.5%'],
        ['链接', 'https://www.sciencedirect.com/science/article/abs/pii/S0306261921013702'],
    ])

doc.add_heading('4.2  Polish Maritime Research 2022 — 2400TEU 集装箱船', level=2)
doc.add_paragraph('主机转速、轴功率、油耗率、航速、吃水、滑失比 —— 2020/2021 年各 7 个月数据', style='List Bullet')
doc.add_paragraph('聚类分析识别出三个典型工况中心：低速(~9.5kn, 1295kW) / 中速(~12.6kn, 2970kW) / 高速(~16.1kn, 6096kW)', style='List Bullet')

doc.add_heading('4.3  韩国渔业巡逻船 — 5 秒间隔实测', level=2)
doc.add_paragraph('MDPI JMSE 2025：韩国渔业巡逻船真实运行剖面，5 秒采样间隔，用于验证混合动力标准负荷分析模型', style='List Bullet')

doc.add_paragraph('')
doc.add_heading('4.4  MDPI Sustainability 2025 — 集装箱港口数据', level=2)
doc.add_paragraph('440 艘船在港一年：主机功率、辅机功率、GRT、TEU、停泊时间 —— 港口能耗预测', style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 5. 船级社／行业工具
# ══════════════════════════════════════════════════════
doc.add_heading('五、船级社／行业工具', level=1)

doc.add_heading('5.1  DNV EETA（Energy Efficiency Performance Indicator）', level=2)
doc.add_paragraph('基于全船队 AIS 数据的能效基准测试工具，覆盖散货船/油轮/集装箱船', style='List Bullet')
doc.add_paragraph('商业服务，非开放数据集，但其方法论（工况剖面构建）值得参考', style='List Bullet')
p = doc.add_paragraph('链接：')
run = p.add_run('https://www.dnv.us/maritime/advisory/eeta/faq/')
run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

doc.add_heading('5.2  DeepSea Technologies AI Benchmarking', level=2)
doc.add_paragraph('免费 AI 基准测试入口，需航运公司注册后可用', style='List Bullet')
p = doc.add_paragraph('链接：')
run = p.add_run('https://elnavi.gr/en/shippings-news/871-deepsea-creates-free-entry-point-for-shipping-companies-to-explore-ai-benchmarking-tools')
run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 6. 其他学术论文引用数据集一览
# ══════════════════════════════════════════════════════
doc.add_heading('六、其他学术论文引用数据集一览', level=1)

add_table(doc,
    ['来源机构', '船型', '关键参数', '采样', '备注'],
    [
        ['Aalborg University', '油/化学品船', '主机 7000kW, 3×580kW 发电机', '—', '含转速、油耗、功率'],
        ['Chalmers University', '化学品船', 'MCR 7200kW', '1 min', '3 年连续数据'],
        ['Chalmers University', '双头渡轮', '2×709kW', '1 min', '1 年数据'],
        ['Simon Fraser Univ.', '多发动机船', 'Power_1/2, SFC, 扭矩, 推力', '—', '含工况/风/货载标签'],
        ['Ch. University (2025)', '化学品船', '主机功率、航速等', '10 min', '3 年数据降采样'],
    ])

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 7. 推荐优先级与使用建议
# ══════════════════════════════════════════════════════
doc.add_heading('七、推荐优先级与使用建议', level=1)

doc.add_heading('针对深海多堆 FC 系统 · EMS 研究的推荐路径', level=2)

recommendations = [
    ('🥇 第一优先 — NAUTILUS 数据集',
     '唯一公开的实船 SOFC+电池实测数据，含完整 CSV 时间序列。'
     '可直接分析真实功率波动特性，提取 FC 与电池的动态响应特征。'),
    ('🥈 第二优先 — TU Delft SH2IPDRIVE 系列',
     '最贴近 FC+电池 EMS 研究场景。MPC/ECMS/滤波三种 EMS 框架代码全开放，'
     '可直接替换负荷数据跑对比实验。Ch.4 含寿命成本优化代码。'),
    ('🥉 第三优先 — Shifts 数据集',
     '4 年实船分钟级功率数据，数据量大、含分布偏移。'
     '最适合做负荷预测模型训练和鲁棒性验证。'),
    ('④ NTUA Marine Loading Cycles',
     '20 条标准工况，可直接作为 EMS 仿真输入，便于与其他研究横向对比。'),
    ('⑤ 中国国家基础科学数据中心',
     'VLOC 和散货船能耗数据，含主机功率和油耗时间序列，适合补充船型多样性。'),
]

for title, desc in recommendations:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(desc)

doc.add_paragraph('')
doc.add_heading('数据获取与使用注意事项', level=2)
doc.add_paragraph('NAUTILUS 数据和 TU Delft 数据集均为 CC BY 4.0 许可，可直接下载使用。', style='List Bullet')
doc.add_paragraph('Shifts 数据集为 CC BY-NC-SA 4.0（非商业），学术研究可用，商业使用需授权。', style='List Bullet')
doc.add_paragraph('Kaggle 数据集需注册 Kaggle 账号后下载。', style='List Bullet')
doc.add_paragraph('中国基础科学数据中心可能需要国内 IP 访问。', style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 附录
# ══════════════════════════════════════════════════════
doc.add_heading('附录：数据源链接汇总', level=1)

links = [
    ('Shifts 功率预测数据集', 'https://zenodo.org/record/7057666'),
    ('Shifts 论文', 'https://arxiv.org/pdf/2206.15407'),
    ('NAUTILUS SOFC+电池实测', 'https://zenodo.org/records/14643552'),
    ('TU Delft Ch.3 LPF-EMS', 'https://data.4tu.nl/datasets/589dc384-b5bf-450b-bba2-c6cd1d33e378/1'),
    ('TU Delft Ch.4 系统优化', 'https://data.4tu.nl/datasets/c1f91c19-b975-4dcc-acb5-58316b11f9a9/1'),
    ('TU Delft 分层控制策略', 'https://data.4tu.nl/datasets/e9fd5f83-35bc-4745-a2a6-682500b2646b/1'),
    ('TU Delft 退化感知 MPC-EMS', 'https://data.4tu.nl/datasets/2cc57bd5-a76d-4c15-9abe-ed814c70c607/1'),
    ('中国国家基础科学数据中心', 'https://60.245.194.22/general/dataDetail?id=69b6d7a5195d2623ac632d4c&type=1'),
    ('Kaggle Ship Performance', 'https://www.kaggle.com/code/hazemsayedabdullah/ship-performance-analysis/clustering/log'),
    ('GitHub 船舶发动机维护', 'https://github.com/mlatinov/Preventive-Maintenance-for-Marine-Engines'),
    ('AIS-TSH 数据集', 'https://zendy.io/title/10.1109%2Fieeedata.2026.3676933'),
    ('AIS+ML 功率预测论文', 'https://iopscience.iop.org/article/10.1088/1742-6596/1357/1/012038'),
    ('NTUA HIPPO-2 Marine Loading Cycles', 'https://www.sciencedirect.com/science/article/abs/pii/S0306261921013702'),
    ('DNV EETA', 'https://www.dnv.us/maritime/advisory/eeta/faq/'),
    ('DeepSea AI Benchmarking', 'https://elnavi.gr/en/shippings-news/871-deepsea-creates-free-entry-point-for-shipping-companies-to-explore-ai-benchmarking-tools'),
]

add_table(doc,
    ['名称', '链接'],
    links,
)

# ── 保存 ──────────────────────────────────────────────
output_path = r'F:\CLAUDE\research\ems-platform\docs\公开实船功率能耗数据集调研报告.docx'
doc.save(output_path)
print(f'OK: {output_path}')
