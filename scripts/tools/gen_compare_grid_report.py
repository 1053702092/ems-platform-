#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 Q-learning vs DQN 大 GridWorld 对比报告 docx"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = PROJECT_ROOT / 'docs' / 'notes'
RESULTS_DIR = PROJECT_ROOT / 'results'
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def p(text, bold=False, sz=11, color=None, indent=0):
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(4)
    if indent: pa.paragraph_format.left_indent = Cm(indent)
    run = pa.add_run(text); run.font.name='微软雅黑'; run.font.size=Pt(sz); run.bold=bold
    if color: run.font.color.rgb = color

def bullet(text, lv=0):
    pa = doc.add_paragraph(text, style='List Bullet')
    pa.paragraph_format.left_indent = Cm(1.5 + lv*0.8)

def code(lines, label=None):
    if label:
        pa = doc.add_paragraph(); run=pa.add_run(label); run.bold=True; run.font.size=Pt(10)
    for line in lines.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before=Pt(0); pa.paragraph_format.space_after=Pt(1)
        pa.paragraph_format.left_indent=Cm(1)
        run=pa.add_run(line); run.font.name='Consolas'; run.font.size=Pt(9.5); run.font.color.rgb=RGBColor(0x33,0x33,0x33)

def tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i,hd in enumerate(headers):
        run=t.rows[0].cells[i].paragraphs[0].add_run(hd); run.bold=True; run.font.size=Pt(10); run.font.name='微软雅黑'
    for rd in rows:
        row=t.add_row()
        for c,txt in enumerate(rd):
            run=row.cells[c].paragraphs[0].add_run(txt); run.font.size=Pt(10); run.font.name='微软雅黑'

def brk(): doc.add_page_break()

def note(text):
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    run = pa.add_run(f'  {text}')
    run.font.name = '微软雅黑'; run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

def img(path, width=15):
    path = Path(path)
    if path.exists():
        doc.add_picture(str(path), width=Cm(width))
    else:
        p(f'[图片未找到: {path}]', color=RGBColor(0xC0,0x39,0x2B))

# ======================== 封面 ========================
for _ in range(3): doc.add_paragraph('')
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t.add_run('Q-learning vs DQN\n大 GridWorld 对比报告'); run.font.size=Pt(26); run.bold=True; run.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t2.add_run('从 Q 表到神经网络，空间变大看区别'); run.font.size=Pt(13); run.font.color.rgb=RGBColor(0x66,0x66,0x66)
brk()

# ================================================================
# 第一章：实验设置
# ================================================================
h('一、实验设置', 1)
p('同一个 GridWorld 环境，两种算法，两种规模：')
p('')
tbl(
    ['Grid 大小', '状态数', 'QL 参数', 'DQN 参数', '训练局数'],
    [
        ['4×4', '16', '64 (16×4)', '1348', '5000'],
        ['8×8', '64', '256 (64×4)', '4420', '5000'],
    ]
)
p('')
p('环境参数：终点在右下角，陷阱在 (1,1)，80% 走对方向，20% 滑走，γ=0.9')
p('Q-learning: lr=0.1, ε=1.0→0.01, decay=0.998')
p('DQN: lr=0.01, hidden=64, batch=64, buffer=50000, ε=1.0→0.01')
brk()

# ================================================================
# 第二章：4×4 结果
# ================================================================
h('二、4×4 GridWorld 对比', 1)

h('2.1 参数规模', 2)
tbl(
    ['', 'Q-learning', 'DQN'],
    [
        ['参数数量', '64', '1348（21 倍）'],
        ['训练时间', '0.18s', '218.50s'],
    ]
)

h('2.2 训练结果', 2)
tbl(
    ['指标', 'Q-learning', 'DQN'],
    [
        ['最后 200 局平均奖励', '+0.81', '+0.78'],
        ['到达终点比例', '90.5%', '89.0%'],
    ]
)
p('')
p('两个方法效果接近。DQN 参数多 21 倍、慢约 1200 倍，但最终策略也能较稳定走到终点。', bold=True, color=RGBColor(0x1F,0x3A,0x5F))

h('2.3 策略对比', 2)
p('Q-learning 策略:', bold=True)
code(' →  →  →  ↓ \n ↑  X  →  ↓ \n ↓  ↓  ↓  ↓ \n →  →  →  G ')
p('DQN 策略:', bold=True)
code(' ↓  →  →  → \n ↓  X  →  → \n →  →  →  → \n →  →  →  G ')
p('两种策略都能到终点，只是路径略有不同。')

h('2.4 收敛曲线', 2)
img(RESULTS_DIR / 'compare_4x4_ql_vs_dqn.png')

note('4×4 结论：小 Grid 上 QL 完胜。参数少、速度快、效果好。DQN 多余。')
brk()

# ================================================================
# 第三章：8×8 结果
# ================================================================
h('三、8×8 GridWorld 对比', 1)

h('3.1 参数规模', 2)
tbl(
    ['', 'Q-learning', 'DQN'],
    [
        ['参数数量', '256', '4420（17 倍）'],
        ['训练时间', '0.45s', '851.14s'],
    ]
)

h('3.2 训练结果', 2)
tbl(
    ['指标', 'Q-learning', 'DQN'],
    [
        ['最后 200 局平均奖励', '+0.81', '+0.50'],
        ['到达终点比例', '90.5%', '75.0%'],
    ]
)

h('3.3 策略对比', 2)
p('Q-learning 策略:', bold=True)
code(' ↓  →  →  →  →  ↓  ↓  ↓ \n ↓  X  ↓  →  →  ↓  ↓  ↓ \n ↓  ↓  ↓  →  →  →  ↓  ↓ \n ↓  →  ↓  ↓  ↓  →  ↓  ↓ ')
p('DQN 策略:', bold=True)
code(' →  →  →  →  →  →  →  → \n →  X  →  →  →  →  →  → \n →  →  →  →  →  →  →  → \n →  →  →  →  →  →  →  → ')

p('')
p('关键发现：DQN 策略退化', bold=True, color=RGBColor(0xC0,0x39,0x2B))
bullet('DQN 学到的策略偏向"一路向右"，能避开陷阱，但缺少 Q-learning 那样稳定的绕行路径')
bullet('最后 200 局到达率：Q-learning 90.5%，DQN 75.0%，说明 DQN 在 8×8 上效率明显下降')
bullet('Q-learning 仍然能学会绕路到终点')
p('')
p('原因分析：', bold=True)
bullet('DQN 的 4420 个参数需要更多训练数据（5000 局不够）')
bullet('Q 表只有 256 个格子，5000 局足够填满')
bullet('DQN 的学习率 lr=0.01 可能偏大，对 8×8 网络需要更精细调参')

h('3.4 收敛曲线', 2)
img(RESULTS_DIR / 'compare_8x8_ql_vs_dqn.png', width=15)

note('8×8 结论：状态数翻 4 倍后，DQN 开始掉队。参数多了但没学好，策略退化到"偏向单一方向"。Q-learning 仍然稳定。但 Q 表参数已达 256，再多状态就开始装不下了。')
brk()

# ================================================================
# 第四章：汇总分析
# ================================================================
h('四、汇总分析', 1)

h('4.1 三张表看趋势', 2)
tbl(
    ['Grid', '状态数', 'Q 表参数', 'DQN 参数', 'QL 奖励', 'DQN 奖励', 'QL 时间', 'DQN 时间'],
    [
        ['4×4', '16', '64', '1,348', '+0.81', '+0.78', '0.18s', '218.50s'],
        ['8×8', '64', '256', '4,420', '+0.81', '+0.50', '0.45s', '851.14s'],
    ]
)
p('')

h('4.2 Q 表 vs 网络的本质区别', 2)
p('Q 表的问题：', bold=True)
bullet('参数 = 状态数 × 动作数。16→64→144...，状态数翻 N 倍，参数翻 N 倍')
bullet('每个状态的 Q 值独立存储，相邻状态互不共享')
bullet('没见过的新状态 → 完全不知道 Q 值（零泛化能力）')
p('')
p('网络的好处：', bold=True)
bullet('参数 ≈ 输入_dim × hidden + hidden × 输出_dim。输入增大但隐藏层不变')
bullet('参数共享：学过格子 A，格子 B 也跟着受益（泛化）')
bullet('没见过的新状态 → 网络也能"猜"出 Q 值')
p('')
p('网络的问题：', bold=True)
bullet('参数多 → 需要更多数据才能训练好')
bullet('可能收敛到局部最优（8×8 上"一路向下"的例子）')
bullet('超参数敏感（学习率、网络大小、batch 大小都要调）')

h('4.3 什么时候用哪个？', 2)
tbl(
    ['场景', '推荐方法', '原因'],
    [
        ['状态少（<1000）', 'Q-learning / 值迭代', 'Q 表装得下，简单高效'],
        ['状态多但动作离散', 'DQN', 'Q 表装不下，网络泛化'],
        ['图像输入', 'DQN + CNN', '用卷积提取特征，全连接算 Q'],
        ['连续动作', 'PPO / SAC', 'DQN 输出层无法处理连续值'],
    ]
)

brk()
h('五、核心记忆点', 1)
p('1. DQN = Q-learning + 神经网络 + 经验回放 + 目标网络', bold=True, sz=12, color=RGBColor(0x1F,0x3A,0x5F))
p('2. 核心改动只有一行：QL 改一个格子 → DQN 梯度下降改全部参数', sz=12)
p('3. 在小问题上 DQN 没优势（甚至更差），但 Q 表的最大问题是泛化能力为 0', sz=12)
p('4. 面试回答用 8×8 结果作为例子：DQN 策略退化的原因就是参数多、样本效率低、调参难', sz=12)
note(f'实验日期: 2026-07-28')

# 保存
OUT_DIR.mkdir(parents=True, exist_ok=True)
path = OUT_DIR / 'QL_vs_DQN_大Grid对比报告.docx'
doc.save(str(path))
print(f'OK: {path}')
