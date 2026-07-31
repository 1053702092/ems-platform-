#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Week 11 四个 py 文件的详细中文逐行分析 docx (v2)
=====================================================
相比 v1 的改进：
  1. 同步最新代码行号（脚本已加入 argparse / week11_common / CLI 参数）
  2. 覆盖文件中每一行代码，不再只挑关键行
  3. 每个函数/方法给出【功能】【输入】【输出】【调用关系】
  4. 补充：数学推导、训练流程走查、数值示例、面试要点、常见疑问 FAQ
  5. 新增 callout 高亮结论 / flow 文字流程图 两种排版
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'

BLUE = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x66, 0x66, 0x66)
DGRAY = RGBColor(0x99, 0x99, 0x99)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x27, 0xAE, 0x60)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)


def build_doc(title, subtitle, sections, filename):
    """sections = list of (type, args...)
       type: 'h1','h2','h3','p','bullet','code','tbl','brk',
             'x'(line,code,explain), 'callout'(text), 'flow'(lines)
    """
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.paragraph_format.line_spacing = 1.35
    style.paragraph_format.space_after = Pt(2)

    # 封面
    for _ in range(3):
        doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title + '\n')
    r.font.size = Pt(24); r.bold = True; r.font.color.rgb = BLUE
    r = t.add_run(subtitle)
    r.font.size = Pt(13); r.font.color.rgb = GRAY
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f'生成日期：{datetime.date.today().isoformat()}')
    r.font.size = Pt(10); r.font.color.rgb = DGRAY
    doc.add_page_break()

    for sec in sections:
        ty = sec[0]
        if ty == 'h1':
            hd = doc.add_heading(sec[1], level=1)
            for rr in hd.runs: rr.font.color.rgb = BLUE
        elif ty == 'h2':
            hd = doc.add_heading(sec[1], level=2)
            for rr in hd.runs: rr.font.color.rgb = BLUE
        elif ty == 'h3':
            hd = doc.add_heading(sec[1], level=3)
            for rr in hd.runs: rr.font.color.rgb = ORANGE
        elif ty == 'p':
            pa = doc.add_paragraph()
            pa.paragraph_format.space_after = Pt(3)
            rr = pa.add_run(sec[1])
            rr.font.name = 'Microsoft YaHei'
            rr.font.size = Pt(sec[2] if len(sec) > 2 else 11)
            rr.bold = sec[3] if len(sec) > 3 else False
            color = sec[4] if len(sec) > 4 else None
            if color: rr.font.color.rgb = color
        elif ty == 'bullet':
            pa = doc.add_paragraph(sec[1], style='List Bullet')
            pa.paragraph_format.left_indent = Cm(1.5 + (sec[2] if len(sec) > 2 else 0) * 0.8)
        elif ty == 'code':
            for line in sec[1].split('\n'):
                pa = doc.add_paragraph()
                pa.paragraph_format.space_before = Pt(0)
                pa.paragraph_format.space_after = Pt(1)
                pa.paragraph_format.left_indent = Cm(1)
                rr = pa.add_run(line)
                rr.font.name = 'Consolas'; rr.font.size = Pt(9.5)
                rr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif ty == 'flow':
            for line in sec[1].split('\n'):
                pa = doc.add_paragraph()
                pa.paragraph_format.space_before = Pt(0)
                pa.paragraph_format.space_after = Pt(1)
                pa.paragraph_format.left_indent = Cm(1)
                rr = pa.add_run(line)
                rr.font.name = 'Consolas'; rr.font.size = Pt(9)
                rr.font.color.rgb = GREEN
        elif ty == 'tbl':
            hdrs, rows = sec[1], sec[2]
            tb = doc.add_table(rows=1, cols=len(hdrs))
            tb.style = 'Table Grid'
            for i, h in enumerate(hdrs):
                rr = tb.rows[0].cells[i].paragraphs[0].add_run(h)
                rr.bold = True; rr.font.size = Pt(9.5); rr.font.name = 'Microsoft YaHei'
            for rd in rows:
                row = tb.add_row()
                for c, txt in enumerate(rd):
                    rr = row.cells[c].paragraphs[0].add_run(txt)
                    rr.font.size = Pt(9.5); rr.font.name = 'Microsoft YaHei'
        elif ty == 'brk':
            doc.add_page_break()
        elif ty == 'callout':
            pa = doc.add_paragraph()
            pa.paragraph_format.space_before = Pt(6)
            pa.paragraph_format.space_after = Pt(6)
            pa.paragraph_format.left_indent = Cm(0.5)
            pa.paragraph_format.right_indent = Cm(0.5)
            rr = pa.add_run('💡 ' + sec[1])
            rr.font.name = 'Microsoft YaHei'; rr.font.size = Pt(11)
            rr.bold = True; rr.font.color.rgb = ORANGE
        elif ty == 'x':
            pa = doc.add_paragraph()
            pa.paragraph_format.space_before = Pt(6)
            pa.paragraph_format.space_after = Pt(2)
            rr = pa.add_run(f'  L{sec[1]}  ')
            rr.font.size = Pt(8); rr.font.color.rgb = DGRAY
            rr = pa.add_run(sec[2])
            rr.font.name = 'Consolas'; rr.font.size = Pt(10)
            rr.bold = True; rr.font.color.rgb = BLUE
            pa2 = doc.add_paragraph()
            pa2.paragraph_format.left_indent = Cm(0.8)
            pa2.paragraph_format.space_before = Pt(0)
            pa2.paragraph_format.space_after = Pt(6)
            rr = pa2.add_run(sec[3])
            rr.font.name = 'Microsoft YaHei'; rr.font.size = Pt(10.5)

    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    print(f'  OK: {path}')


# =====================================================================
# File 1: week11_continuous_env.py —— 连续动作环境 + DQN 为什么不行
# =====================================================================
build_doc(
    'Step 1: 连续动作环境',
    'week11_continuous_env.py — 从 GridWorld 到 EMS 连续控制的第一步',
    [
        # ============ 一、文件概览 ============
        ('h1', '一、文件概览'),
        ('p', '这个文件是 Week 11 的第一课，目标只有一个：**造出一个"连续动作"的强化学习环境，并亲眼验证 DQN 为什么搞不定它**。上一周（Week 10）的 GridWorld 环境里，状态是 16 个格子编号、动作是 4 个方向键，都是离散的；本文件把两者都换成连续值，这是从"离散 RL 入门"跨入"连续控制 RL"的分水岭。', 11),
        ('callout', '一句话记忆：本文件 = ① 一个连续动作的 EMS 简化环境（EMSEnv）+ ② 一个证明 DQN 用不了的演示（demo_dqn_failure）。'),
        ('h2', '1.1 与 Week10 GridWorld 的对比'),
        ('tbl', ['维度', 'Week10 GridWorld', '本文件 EMSEnv', '关键差异'], [
            ['状态', '16 个离散格子编号 0-15', '[SOC, P_load] 2 维连续值', '离散→连续，用 2 个浮点数直接描述'],
            ['动作', '4 个方向（上/下/左/右）', 'P_fc ∈ [0,1] 归一化连续值', '从"4 选 1"变成"区间内任取值"'],
            ['奖励', '走到终点 +1 / 陷阱 -1', '三项连续奖励（燃料成本+SOC 惩罚+越界惩罚）', '从离散奖励变成连续数值'],
            ['终止', '到达终点或陷阱', '步数超限或 SOC 越界', '物理约束（电池安全）驱动终止'],
            ['能否用 DQN', '可以（4 个 Q 值 argmax）', '不行（见 5 章演示）', 'DQN 的 argmax 需要离散动作'],
        ]),
        ('p', '注意：虽然环境变成了"连续"，但 RL 的基本接口没变——依然是 reset() 返回初始状态、step(action) 返回 (下一个状态, 奖励, 是否结束, 附加信息)。你调用的方式跟 GridWorld 一模一样，变的只是里面的"值"。', 11, True),

        ('h2', '1.2 与真实 EMS（能量管理系统）的对应关系'),
        ('tbl', ['代码里的元素', '物理含义', '真实系统对应'], [
            ['SOC（State of Charge）', '电池剩余电量比例 0~1', '锂电池荷电状态，如 0.6 = 60%'],
            ['P_load（负载功率）', '当前用电设备的总功率', '船舶/车辆的推进+辅助用电负荷'],
            ['P_fc（燃料电池功率）', '燃料电池输出的功率，0~30kW', '发动机/燃料电池的出力'],
            ['battery_capacity = 50 kWh', '电池总容量', '电池包的额定能量（度）'],
            ['reward', '每步的即时得分', '优化目标（经济性+安全性）'],
        ]),
        ('p', '物理关系：P_fc 是"发电"，P_load 是"用电"。发电 > 用电 → 多余电能给电池充电（SOC 上升）；发电 < 用电 → 电池放电补缺（SOC 下降）。我们要让燃料电池既别烧太多燃料（经济），又别让电池过充过放（安全）。', 11),

        # ============ 二、代码结构地图 ============
        ('brk',),
        ('h1', '二、代码结构地图'),
        ('tbl', ['部分', '行号', '内容', '作用'], [
            ['文件头注释', '1-14', '说明状态/动作/奖励的定义', '定位这个文件在 Week11 里的角色'],
            ['import 语句', '15-23', '导入 random/argparse/numpy/torch 等', '准备工具库 + week11_common'],
            ['EMSEnv 类', '26-111', '连续动作 EMS 环境', '本文件主角①：环境'],
            ['DQN_Continuous 类', '115-133', '强行套用 DQN 的网络', '演示用的"假 DQN"'],
            ['demo_dqn_failure()', '136-203', '证明 DQN 数学上不可用', '本文件主角②：论证'],
            ['test_env()', '207-235', '随机策略跑 10 步测环境', '冒烟测试：确认环境不报错'],
            ['main()', '238-259', '解析 CLI 参数并调度', '程序入口'],
        ]),
        ('p', '整个文件的执行顺序：main() → 先 test_env() 测环境 → 再 demo_dqn_failure() 演示失败。', 11, True),

        # ============ 三、import 精讲 ============
        ('brk',),
        ('h1', '三、import 语句逐行精讲（15-23 行）'),
        ('x', 1, '#!/usr/bin/env python3', 'Shebang 行。在 Unix/Linux 下让系统用 python3 解释器直接执行本文件（./week11_continuous_env.py）。Windows 上不起作用，但保留无妨。'),
        ('x', 2, '# -*- coding: utf-8 -*-', '声明文件编码为 UTF-8。保证文件里的中文注释和字符串在旧版本 Python（<3.0）下不乱码。Python 3 默认 UTF-8，这行是给极端兼容场景用的。'),
        ('x', 4, 'Week 11 Step 1: 连续动作环境 — EMS 简化版', '模块 docstring 标题行：标注这是 Week11 的第 1 步。'),
        ('x', 9, '状态: [SOC, P_load] ← 2 维连续值', '状态的定义。2 个浮点数：SOC（电池电量）和 P_load（负载功率）。注意它们是"连续"的——SOC 可以是 0.573 这种任意值，而不是"第 5 个格子"。'),
        ('x', 10, '动作: P_fc ∈ [0, 1] ← 1 维连续值（归一化，对应 0-30 kW）', '动作的定义。P_fc 归一化到 [0,1]，对应实际燃料电池功率 0~30kW。后面 step() 里乘以 30 就是反归一化。'),
        ('x', 11, '奖励: -fuel_cost - penalty ← 连续值', '奖励是连续值，由燃料成本（负）和惩罚项（负）叠加。'),
        ('x', 13, '对比 DQN 在这个环境上为什么会失败', '点明本文件第二个任务：论证 DQN 在此环境不可用。'),
        ('x', 15, 'import random', '导入 random 库。用于 test_env() 里产生随机动作（random.random()），模拟"随机策略"。'),
        ('x', 16, 'import argparse', '导入 argparse（命令行参数解析）。main() 里用它解析 --seed 和 --skip-dqn-demo 参数。这是 Week11 脚本统一加的 CLI 能力。'),
        ('x', 18, 'import numpy as np', '导入 NumPy。用于数组运算（状态向量）、np.clip（数值裁剪）、np.sin（正弦负载波形）、np.mean（求平均）。'),
        ('x', 19, 'import torch', '导入 PyTorch 主模块。用于神经网络（DQN_Continuous）、张量运算。'),
        ('x', 20, 'import torch.nn as nn', '导入神经网络子模块。DQN_Continuous 继承 nn.Module，用 nn.Sequential / nn.Linear / nn.ReLU 搭建网络。'),
        ('x', 21, 'import torch.optim as optim', '导入优化器。demo_dqn_failure() 里用 optim.Adam 创建优化器（虽然是"假的更新"）。'),
        ('x', 23, 'from week11_common import set_seed', '从共享模块 week11_common 导入 set_seed。这是 Week11 重构时新增的公共工具，统一设置 random/numpy/torch 的随机种子，保证实验可复现。注意：week11_common.py 与被导入脚本在同一目录（scripts/），Python 运行时靠该目录下的 __init__ 或直接脚本位置找到它。'),

        # ============ 四、EMSEnv 类 ============
        ('brk',),
        ('h1', '四、EMSEnv 环境类逐行精讲（26-111 行）'),
        ('p', '这是本文件的核心。EMSEnv 实现了 Gym 风格的强化学习环境接口：reset() 初始化、step(action) 推进一步。理解它的关键是抓住**物理模型**：SOC 怎么随功率变化、奖励三项各代表什么、终止条件怎么触发。', 12, True),

        ('h2', '4.1 类定义与 __init__（26-47 行）'),
        ('x', 26, 'class EMSEnv:', '定义 EMS 简化环境类。和 GridWorld 的关键区别：① 状态是连续向量 [SOC, P_load]，不是离散格子编号；② 动作是连续值 P_fc ∈ [0,1]，不是 4 个方向键。这两点直接导致 DQN 用不了。'),
        ('x', 27, '"""简化 EMS 环境', '类的 docstring，解释物理意义。'),
        ('x', 31, 'SOC（电池电量）会随着负载用电和燃料电池供电而变化。', '核心动力学：SOC 不是固定值，而是每步根据功率差变化的动态量。'),
        ('x', 32, 'P_fc（燃料电池功率）是我们要控制的变量。', '明确：P_fc 是控制量（动作），P_load 是扰动量（环境变化），SOC 是状态量（被影响）。这是 EMS 问题的标准三分法。'),
        ('x', 33, '目标是：用最少的燃料（最小化 P_fc），同时维持 SOC 在合理范围。', '目标函数的双重性：既要省钱（P_fc 小 = 燃料少），又要安全（SOC 不出界）。这两个目标有冲突，RL 就是来权衡的。'),
        ('x', 35, 'def __init__(self):', '构造函数：初始化环境参数。'),
        ('x', 37, 'self.soc_min = 0.2', 'SOC 下限 20%。锂电池放太干净会损伤电池，这是安全硬边界。'),
        ('x', 38, 'self.soc_max = 0.9', 'SOC 上限 90%。过充同样危险，也是硬边界。'),
        ('x', 39, 'self.state_dim = 2   # [SOC, P_load]', '状态维度 = 2。状态向量由 SOC 和 P_load 两个连续值组成，直接输入神经网络。对比 GridWorld 用 one-hot 编码 16 个格子。'),
        ('x', 40, 'self.action_dim = 1  # P_fc (归一化 0~1)', '动作维度 = 1。动作就是燃料电池功率一个数，归一化到 [0,1]。'),
        ('x', 43, 'self.battery_capacity = 50  # kWh', '电池容量 50 kWh。这个值决定了 SOC 变化的速度：容量越大，同样的功率差引起的 SOC 变化越小。'),
        ('x', 44, 'self.dt = 1/60  # 每步 = 1 分钟', '每步时长 = 1 分钟（1/60 小时）。dt 在动力学中用于把功率换算成能量（但本代码简化处理，见 4.3 说明）。'),
        ('x', 47, 'self.reset()', '构造时立即调用 reset()，把环境初始化到默认状态，这样对象创建后立刻就能用。'),

        ('h2', '4.2 reset 与 _get_state（49-59 行）'),
        ('x', 49, 'def reset(self):', '重置环境到初始状态，并返回初始状态向量。训练每局开始时都会调用它，清空上一局的残留。'),
        ('x', 51, 'self.soc = 0.6  # 初始 SOC = 60%', '初始电量 60%，取在安全区间 [0.2, 0.9] 的中偏上位置，给智能体留足上下操作空间。'),
        ('x', 52, 'self.p_load = 0.5  # 初始负载（KW）', '初始负载 0.5 kW。注意这里 p_load 单位写 KW（应为 kW），是作者的小笔误，不影响逻辑。'),
        ('x', 53, 'self.steps = 0', '步数清零。用于判断是否达到 max_steps。'),
        ('x', 54, 'self.max_steps = 200', '一局最多 200 步（约 3 小时 20 分钟）。防止环境无限运行下去，也定义了"一局"的长度。'),
        ('x', 55, 'return self._get_state()', '返回初始状态向量。RL 训练代码用这个返回值作为第一轮的观测。'),
        ('x', 57, 'def _get_state(self):', '内部辅助方法：把当前 SOC 和 P_load 组装成状态向量。'),
        ('x', 59, 'return np.array([self.soc, self.p_load], dtype=np.float32)', '返回 numpy float32 数组 [SOC, P_load]。用 np.float32 与 PyTorch 默认张量精度一致（PyTorch 默认 float32），避免类型转换开销和精度警告。'),
        ('callout', '记忆点：reset() = 把电量设回 60%、负载设回 0.5、步数清零，然后吐出初始状态。step() = 接收动作 → 更新 SOC → 算奖励 → 返回新状态。这就是"环境"的全部职责。'),

        ('h2', '4.3 step() 的核心动力学（61-104 行）'),
        ('x', 61, 'def step(self, action):', '环境最核心的方法：执行动作，推进一步，返回 (next_state, reward, done, info)。所有 RL 算法（REINFORCE/AC/PPO）都通过它和环境互动，接口和 Gym 完全一致。'),
        ('x', 62, '"""执行动作 action = P_fc (归一化 0~1)', 'docstring：说明入参是归一化的 P_fc。'),
        ('x', 65, '返回: (next_state, reward, done, info)', '标准四元组返回值：下一状态、即时奖励、是否终止、附加信息字典。'),
        ('x', 68, 'p_fc = float(np.clip(action, 0, 1)) * 30.0  # [0, 30] kW', '★ 反归一化：把网络输出的 [0,1] 归一化动作，映射回真实功率 [0,30] kW。np.clip 把动作硬夹在 [0,1]（防止网络输出 1.3 或 -0.2 这种越界值），float() 转成 Python 浮点，×30 得到 kW。例：action=0.5 → 15 kW，action=1.0 → 30 kW。'),
        ('x', 71, 'self.p_load = 0.3 + 0.4 * (0.5 + 0.5 * np.sin(self.steps * 0.1))', '★ 负载波形：生成随时间变化的正弦负载。展开计算：sin 部分在 [-1,1] 波动 → (0.5+0.5·sin) ∈ [0,1] → ×0.4 ∈ [0,0.4] → +0.3 ∈ [0.3,0.7] kW。周期为 2π/0.1 ≈ 63 步，也就是负载约 1 小时一循环。模拟真实工况里负荷的起伏。'),
        ('x', 74, '# P_fc > P_load → 充电，SOC ↑', '充电逻辑注释：发电 > 用电时，多余电能充进电池。'),
        ('x', 75, '# P_fc < P_load → 放电，SOC ↓', '放电逻辑注释：发电 < 用电时，电池放电补缺口。'),
        ('x', 76, 'power_diff = p_fc - self.p_load  # kW', '功率差 = 发电 - 用电。这是 SOC 变化的驱动量。正值 → 充电，负值 → 放电。'),
        ('x', 77, 'soc_change = power_diff / self.battery_capacity  # 归一化变化量', 'SOC 变化量 = 功率差 ÷ 电池容量。这是从"功率"到"电量比例"的换算。严格物理上应该还要乘以 dt（因为功率×时间=能量），但这里 dt=1 分钟，步长小，代码直接省略了，是简化处理。例：功率差 +5 kW，容量 50 kWh → SOC 每步上升 0.1。'),
        ('x', 78, 'self.soc = np.clip(self.soc + soc_change, self.soc_min, self.soc_max)', '更新 SOC 并夹到 [0.2, 0.9] 安全区间。clip 保证 SOC 永远不会真的越界存储（越界的惩罚在奖励里体现，见 4.4）。'),
        ('callout', '物理本质：这是一个离散时间的一阶惯性系统 SOC[t+1] = SOC[t] + (P_fc[t] - P_load[t])/C。它把"连续动作"和"连续状态"直接连起来了——这正是 DQN 棘手的地方。'),

        ('h2', '4.4 step() 的奖励函数（80-92 行）'),
        ('x', 80, '# 计算奖励', '奖励计算分三部分：燃料成本 + SOC 偏离惩罚 + SOC 越界惩罚。'),
        ('x', 82, 'fuel_cost = -0.01 * p_fc', '① 燃料成本：P_fc 越大 → 负奖励越大（惩罚越重）。0.01 是单位成本系数。这是"经济性"项——逼着智能体别让燃料电池乱烧。例：P_fc=30kW → 该项 = -0.3。'),
        ('x', 85, 'soc_penalty = -0.5 * (self.soc - 0.6) ** 2', '② SOC 偏离惩罚：SOC 偏离目标值 0.6 越远，惩罚越大。平方项保证：无论过高(>0.6)还是过低(<0.6)都受罚，且偏离越远惩罚增长越快（二次）。系数 0.5 控制惩罚力度。这是"安全性/电量保持"项。例：SOC=0.7 → -0.5×0.01 = -0.005；SOC=0.2 → -0.5×0.16 = -0.08。'),
        ('x', 88, 'soc_bound_penalty = 0.0  #0.2-0.9', '③ 越界惩罚：默认 0。注释标出安全范围是 0.2-0.9。'),
        ('x', 89, 'if self.soc <= self.soc_min or self.soc >= self.soc_max:', '判断 SOC 是否触碰硬边界（≤0.2 或 ≥0.9）。'),
        ('x', 90, 'soc_bound_penalty = -1.0', '触界 → 一次性罚 -1.0。相比燃料成本（~-0.1 量级），这是重罚，让智能体强烈避免触碰边界。'),
        ('x', 92, 'reward = fuel_cost + soc_penalty + soc_bound_penalty', '★ 总奖励 = 三项相加。注意三者量级差异：燃料成本常态 ~-0.1，SOC 偏离 ~-0.01~-0.1，越界一次性 -1。智能体要平衡：不能只图省油（导致 SOC 掉出界吃重罚），也不能只守 SOC（一直满功率烧油）。这就是"权衡"。'),

        ('h2', '4.5 step() 的终止与返回（94-104 行）'),
        ('x', 95, 'self.steps += 1', '步数自增，用于触发最大步数终止。'),
        ('x', 96, 'done = (self.steps >= self.max_steps or', '终止条件第 1 项：达到 200 步上限（一局自然结束）。'),
        ('x', 97, 'self.soc <= self.soc_min or', '终止条件第 2 项：SOC 掉到 0.2 以下（过放）。'),
        ('x', 98, 'self.soc >= self.soc_max)', '终止条件第 3 项：SOC 顶到 0.9 以上（过充）。后两项模拟真实电池保护系统——电量出界就停机。'),
        ('x', 100, 'return self._get_state(), reward, done, {', '返回标准四元组。info 字典携带诊断信息（调试/画图用）。'),
        ('x', 101, "'p_fc': p_fc,", 'info 里带上反归一化后的实际功率（kW），方便外部查看智能体到底输出了多少功率。'),
        ('x', 102, "'fuel_cost': fuel_cost,", 'info 里带上燃料成本分量，便于分析奖励构成。'),
        ('x', 103, "'soc': self.soc", 'info 里带上最新 SOC，便于监控电量轨迹。'),
        ('x', 104, '}', '结束 step()。注意 dict 最后一项没逗号，是合法 Python（省略了尾随逗号）。'),

        ('h2', '4.6 render()（106-111 行）'),
        ('x', 106, 'def render(self):', '文字渲染方法：把当前状态用 ASCII 条形图打印出来，方便肉眼观察。'),
        ('x', 108, 'bar_len = 20', '条形图总长度 20 格。'),
        ('x', 109, 'soc_bar = int((self.soc - self.soc_min) / (self.soc_max - self.soc_min) * bar_len)', '把 SOC 从 [0.2,0.9] 映射到 0-20 格的整数长度。公式本质是线性归一化后×20。例：SOC=0.6 → (0.6-0.2)/(0.9-0.2)×20 ≈ 11.4 → int=11。'),
        ('x', 110, "bar = '█' * soc_bar + '░' * (bar_len - soc_bar)", '拼出条形图：soc_bar 格实心块 + 剩余格空心块。'),
        ('x', 111, 'print(f"Step {self.steps:3d} | SOC [{bar}] {self.soc:.2f} | Load {self.p_load:.2f}")', '打印一行：步数、SOC 条形图、SOC 数值、负载。f-string 里 :3d 右对齐占 3 位，:.2f 保留 2 位小数。'),

        # ============ 五、DQN 演示 ============
        ('brk',),
        ('h1', '五、DQN 为什么不行 —— 逐行精讲（115-203 行）'),
        ('p', '这一章是本文件的教学重点。结论先说：**DQN 处理不了连续动作，不是调参的问题，是数学结构（argmax 和 max）决定的**。', 12, True),

        ('h2', '5.1 DQN_Continuous 网络（115-133 行）'),
        ('x', 115, 'class DQN_Continuous(nn.Module):', '定义一个"假 DQN"网络。为什么说"假"？因为标准 DQN 的输出层应该有 N 个神经元（N = 离散动作数），每个神经元对应一个动作的 Q 值，最后 argmax。而这个网络输出层只有 1 个神经元，直接输出一个连续值——它本质上是个回归网络，已经失去了 DQN 的灵魂。'),
        ('x', 117, 'DQN 强行用在连续动作上——', 'docstring：直接点明这是"强行"套用。'),
        ('x', 119, '但 Q-learning 需要 argmax，对连续值没法 argmax！', '★ 核心论证起点：Q-learning 选动作靠 argmax，连续空间里没法遍历求 argmax。'),
        ('x', 124, 'def __init__(self, state_dim=2, hidden=64):', '构造函数。输入维度 2（状态），隐藏层 64。'),
        ('x', 126, 'self.net = nn.Sequential(', '用 nn.Sequential 顺序堆叠网络层。'),
        ('x', 127, 'nn.Linear(state_dim, hidden),', '第 1 层：全连接层，2 维输入 → 64 维隐藏。'),
        ('x', 128, 'nn.ReLU(),', 'ReLU 激活函数：max(0,x)，引入非线性，让网络能学复杂映射。'),
        ('x', 129, 'nn.Linear(hidden, 1)  # 输出 1 个连续动作值', '★ 输出层：hidden → 1 个神经元。这个输出被"当成动作值直接使用"，而不是 N 个 Q 值。注意：如果这是真 DQN，这里应该是 nn.Linear(hidden, num_actions)，且没有激活——Q 值可以是任意实数。'),
        ('x', 132, 'def forward(self, x):', '前向传播：输入状态，输出网络预测。'),
        ('x', 133, 'return self.net(x)', '直接把 x 送进 Sequential 网络并返回结果。'),

        ('h2', '5.2 demo_dqn_failure() 论证（136-203 行）'),
        ('x', 136, 'def demo_dqn_failure():', '主论证函数：从数学原理到代码实验，逐步展示 DQN 的失败。'),
        ('x', 137, '"""演示 DQN 强行用在连续动作上会怎么样', 'docstring 概述。'),
        ('x', 141, 'DQN 的核心是：输出每个动作的 Q 值 → argmax 选最大的', '复盘 DQN 决策机制。'),
        ('x', 142, '连续动作有无限个可能值，没法 argmax！', '★ 核心矛盾：连续区间 [0,1] 有无穷多个点，argmax 需要"遍历所有候选取最大"，无穷集没法遍历。'),
        ('x', 147, 'print("=" * 65)', '打印分隔线，纯排版。'),
        ('x', 152, 'env = EMSEnv()', '创建环境，为后面的"伪训练"实验做准备。'),
        ('x', 153, 'q_net = DQN_Continuous()', '创建这个假 DQN 网络。'),
        ('x', 154, 'optimizer = optim.Adam(q_net.parameters(), lr=0.01)', '创建 Adam 优化器。注意：这个优化器在后面根本不会真正用来做 Q-learning 更新（因为公式不成立），它是为了演示"即便有优化器也学不了"的荒谬。'),
        ('x', 155, 'loss_fn = nn.MSELoss()', '创建均方误差损失函数（本文件里实际未使用，留着是惯例占位）。'),
        ('x', 157, 'print("问题 1: 没法 argmax")', '开始论证问题 1。'),
        ('x', 158, 'print("  DQN 选动作: a = argmax Q(s, a) ← 需要遍历所有动作")', 'DQN 的决策公式：遍历所有动作，找 Q 值最大的那个。'),
        ('x', 159, 'print("  连续动作:    a ∈ [0, 1]，有无穷多个值")', '连续动作空间 [0,1] 有无穷多个候选值。'),
        ('x', 160, 'print("  遍历不了 → 没法用 argmax → Q-learning 公式失效")', '结论：argmax 要求遍历，连续空间无法遍历 → 决策环节直接失效。'),
        ('x', 163, 'print("问题 2: 即使强行输出一个动作值，更新公式也不对")', '开始论证问题 2（更致命）。'),
        ('x', 164, "print(\"  Q-learning 的 target = r + γ·max Q(s', a')\")", 'Q-learning 的 TD 目标里含 max Q(s\', a\')——"下一状态的最优动作价值"。'),
        ('x', 165, "print(\"  连续空间里 max Q(s', a') 没法算\")", '★ 这个 max 同样要遍历下一动作空间。就算你解决了"选动作"（问题 1），"算目标"（问题 2）还要再遍历一次。两个环节都卡死在连续空间上。'),
        ('x', 168, 'print("做个实验：让网络直接输出动作（不是 Q 值）")', '好，那"绕过"这两个问题行不行？试试让网络直接输出动作值。'),
        ('x', 169, 'print("跑 50 局看看效果：")', '实验设置：跑 50 局。'),
        ('x', 172, 'for ep in range(1, 51):', '外层循环：50 局。'),
        ('x', 173, 's = env.reset()', '每局重置环境。'),
        ('x', 174, 'total_reward = 0', '累计本局总奖励。'),
        ('x', 176, 'for t in range(200):', '内层循环：每局最多 200 步。'),
        ('x', 177, 's_tensor = torch.FloatTensor(s).unsqueeze(0)', '把状态转成 PyTorch 张量，并 unsqueeze(0) 增加 batch 维（shape 从 [2] 变成 [1,2]），满足网络对 batch 输入的要求。'),
        ('x', 180, 'with torch.no_grad():', '关闭梯度追踪。因为这里只是"推理"（inference），不需要反向传播。'),
        ('x', 181, 'a = float(torch.sigmoid(q_net(s_tensor)).item())', '★ 绕过 argmax：让网络直接输出一个值，再用 sigmoid 压到 (0,1)。这就是"假装 DQN 直接给动作"。但注意——这是监督学习式映射（状态→动作），完全没有 Q 值的概念、没有贝尔曼公式，已经不是强化学习了。'),
        ('x', 183, 'sp, reward, done, _ = env.step(a)', '执行动作，环境返回新状态/奖励/是否结束。'),
        ('x', 185, '# 尝试用 Q-learning 公式更新——但这里 max Q(sp) 算不了', '注释说明：想用 Q-learning 更新，但 max Q(sp) 无从计算。'),
        ('x', 186, '# 因为连续动作没有"max"', '原因：连续动作空间没有 max 操作。'),
        ('x', 187, '# 所以这根本就不是 DQN 了，只是瞎更新', '结论：没有正确更新公式，这不是 DQN，是随机游走。'),
        ('x', 189, 'total_reward += reward', '累加奖励。'),
        ('x', 190, 's = sp', '状态推进。'),
        ('x', 191, 'if done:', '如果环境终止。'),
        ('x', 192, 'break', '提前结束本局。'),
        ('x', 194, 'if ep % 10 == 0:', '每 10 局打印一次进度。'),
        ('x', 195, 'print(f"  第{ep:3d}局 | 总奖励={total_reward:.3f} | "', '格式化打印局数（右对齐 3 位）和总奖励（3 位小数）。'),
        ('x', 196, 'f"无意义——因为没有正确的更新公式")', '打印关键注脚：这局成绩毫无意义，因为没有正确更新。'),
        ('x', 199, 'print("结论：")', '输出结论段。'),
        ('x', 200, 'print("  DQN 从数学结构上就无法处理连续动作。")', '★ 总结论一：数学结构（argmax + max）决定了 DQN 无法处理连续动作。'),
        ('x', 201, 'print(" 这不是调参能解决的，是 Q-learning 的 argmax 决定的。")', '★ 总结论二：增加网络宽度/深度、改学习率、加训练次数——都救不了，因为病根在决策和更新的数学公式。'),
        ('x', 202, 'print(" 要处理连续动作，必须换方法——策略梯度。")', '★ 总结论三：出路是策略梯度方法（REINFORCE/AC/PPO），它们不靠 argmax 选动作，而是直接输出动作分布、从分布采样。这是 Week11 后续三步的主题。'),

        # ============ 六、test_env 与 main ============
        ('brk',),
        ('h1', '六、环境测试与主程序（207-259 行）'),
        ('h2', '6.1 test_env()（207-235 行）'),
        ('x', 207, 'def test_env():', '冒烟测试：不关心学得好不好，只确认环境能正常跑、奖励边界合理。'),
        ('x', 213, 'env = EMSEnv()', '创建环境。'),
        ('x', 214, 's = env.reset()', '重置，拿初始状态。'),
        ('x', 215, 'print(f"初始状态: SOC={s[0]:.2f}, P_load={s[1]:.2f}")', '打印初始 SOC 和负载（都应为 0.60 / 0.50）。'),
        ('x', 221, 'for t in range(10):', '随机策略跑 10 步。'),
        ('x', 222, 'a = random.random()', '★ 随机动作：random.random() 产生 [0,1) 均匀随机数，正好对应归一化动作范围。这是"随机策略"——不看状态，随便输出。'),
        ('x', 223, 'sp, reward, done, info = env.step(a)', '推进环境。'),
        ('x', 224, 'total_reward += reward', '累计奖励。'),
        ('x', 225, 'print(f"  步{t+1:2d}: P_fc={info[\'p_fc\']:5.1f}kW | SOC={info[\'soc\']:.3f} | "', '打印本步：实际功率（info 里取，5 位宽 1 位小数）、SOC（3 位小数）。'),
        ('x', 226, 'f"奖励={reward:+.4f}")', '奖励带符号显示（+/- 4 位小数）。'),
        ('x', 227, 's = sp', '状态推进。'),
        ('x', 229, 'print(f"\\n10 步总奖励: {total_reward:+.4f}")', '汇总打印 10 步总奖励。'),
        ('x', 231, 'print("状态维度:", env.state_dim, "(连续)")', '打印状态维度，标注"连续"。'),
        ('x', 232, 'print("动作维度:", env.action_dim, "(连续)")', '打印动作维度，标注"连续"。'),
        ('x', 233, 'print("状态范围: SOC=[0.2, 0.9], P_load=[0.3, 0.7]")', '打印状态范围。P_load 范围 0.3-0.7 来自 step() 里的正弦波形。'),
        ('x', 234, 'print("动作范围: P_fc=[0, 1] 归一化 → [0, 30] kW")', '打印动作范围：归一化 [0,1] 对应实际 [0,30] kW。'),

        ('h2', '6.2 main()（238-259 行）'),
        ('x', 238, 'def main() -> None:', '程序入口函数。返回值类型标注为 None。'),
        ('x', 239, "parser = argparse.ArgumentParser(description='Week 11 simplified continuous-action EMS environment demo')", '创建参数解析器，description 会在 --help 时显示。'),
        ('x', 240, "parser.add_argument('--seed', type=int, default=42, help='Random seed')", '定义 --seed 参数（默认 42）。方便复现实验。'),
        ('x', 241, "parser.add_argument('--skip-dqn-demo', action='store_true', help='Only run the environment smoke test')", '定义 --skip-dqn-demo 开关：加上它就跑环境测试、跳过 DQN 失败演示（快速冒烟用）。action=\'store_true\' 表示"存在即 True"。'),
        ('x', 242, 'args = parser.parse_args()', '解析命令行参数。'),
        ('x', 244, 'set_seed(args.seed)', '★ 调用 week11_common 的 set_seed，统一设置 random/numpy/torch 随机种子，保证可复现。'),
        ('x', 247, 'print("╔══════════════════════════════════════════════════════╗")', '打印框线装饰（盒形字符），让程序输出有辨识度。'),
        ('x', 248, 'print("║  Week 11 Step 1: 连续动作环境                     ║")', '标题行。'),
        ('x', 249, 'print("║  从离散（GridWorld 4×4）到连续（EMS 简化版）      ║")', '副标题：点明学习跨度。'),
        ('x', 250, 'print("╚══════════════════════════════════════════════════════╝")', '盒形收尾。'),
        ('x', 253, 'test_env()', '先跑环境冒烟测试。'),
        ('x', 254, 'if not args.skip_dqn_demo:', '如果没加 --skip-dqn-demo。'),
        ('x', 255, 'demo_dqn_failure()', '再跑 DQN 失败演示。默认两条都跑。'),
        ('x', 258, "if __name__ == '__main__':", '★ Python 入口惯用法：只有直接执行本文件（而非被 import）时，才执行 main()。被 import 时只加载类定义，不自动跑演示。'),
        ('x', 259, 'main()', '调用主函数，启动程序。'),

        # ============ 七、机制专题 ============
        ('brk',),
        ('h1', '七、机制专题'),
        ('h2', '7.1 一个 step 的完整数值走查'),
        ('p', '假设当前 SOC=0.55，steps=30，动作 action=0.7，逐步推演：', 11),
        ('flow', '''1. 反归一化:    p_fc = clip(0.7,0,1)×30 = 21.0 kW
2. 负载更新:     sin(30×0.1)=sin(3)≈0.141
                p_load = 0.3 + 0.4×(0.5+0.5×0.141) = 0.3+0.4×0.5706 ≈ 0.528 kW
3. 功率差:       power_diff = 21.0 - 0.528 ≈ 20.47 kW
4. SOC 变化:     soc_change = 20.47 / 50 ≈ 0.409   ← 很大！因为满载充电
5. 新 SOC:       soc = clip(0.55+0.409, 0.2, 0.9) = 0.9  ← 直接顶到上界
6. 奖励:         fuel_cost = -0.01×21 = -0.21
                soc_penalty = -0.5×(0.9-0.6)² = -0.5×0.09 = -0.045
                soc_bound_penalty = -1.0  (SOC≥0.9 触界)
                reward = -0.21 - 0.045 - 1.0 = -1.255
7. 终止:         SOC=0.9 → done = True   ← 这一局提前结束'''),
        ('callout', '这个例子揭示了环境的设计张力：动作直接取 0.7（21kW）远超负载 0.53kW，导致 SOC 一步冲顶。智能体必须学会"跟随负载"——P_fc 大致等于 P_load，SOC 才平稳，奖励才高。这正是 PPO 环境里加入 tracking_bonus 的伏笔。'),

        ('h2', '7.2 奖励三项的权衡博弈'),
        ('tbl', ['项目', '公式', '量级', '主导倾向'], [
            ['燃料成本', '-0.01 × P_fc', '-0.3 ~ 0', '让 P_fc 越小越好（省油）'],
            ['SOC 偏离', '-0.5 × (SOC-0.6)²', '-0.08 ~ 0', '让 SOC 贴近 0.6'],
            ['越界惩罚', '触界时 -1.0', '0 或 -1', '绝不让 SOC 出 [0.2,0.9]'],
        ]),
        ('p', '三者角力：如果只省油（P_fc≈0），SOC 会一路掉到 0.2 以下触发 -1.0 重罚；如果只保电量（P_fc≈30），燃料成本 -0.3 且 SOC 冲到 0.9 又触发重罚。最优解就是 P_fc ≈ P_load，SOC 微幅波动。RL 的任务就是从这个连续空间中搜出这个平衡点。', 11),

        # ============ 八、面试要点与 FAQ ============
        ('brk',),
        ('h1', '八、面试要点'),
        ('bullet', '「为什么 EMS 的功率分配是连续控制问题？」→ 燃料电池功率可以在 0-30kW 任意取值，不是 4 个档位；强行离散化会损失精度并导致维度灾难（离散档数指数增长）。', 0),
        ('bullet', '「DQN 为什么不能直接处理连续动作？」→ 两个数学环节：决策 a=argmax Q(s,a) 和更新 target=r+γ·max Q(s\',a\')，都需要遍历动作空间；连续空间有无穷多个点，无法遍历。这是结构性问题，调参无法解决。', 0),
        ('bullet', '「连续动作该怎么解决？」→ 策略梯度方法：直接输出动作分布参数 [μ,σ]，从正态分布采样得动作，绕开 argmax。', 0),
        ('bullet', '「奖励函数怎么设计的？」→ 三项：燃料成本（经济性）、SOC 偏离惩罚（电量保持）、SOC 越界惩罚（安全性硬约束）。', 0),
        ('bullet', '「环境接口？」→ Gym 风格：reset() 返回初始状态，step(action) 返回 (state, reward, done, info)。', 0),

        ('h1', '九、常见疑问 FAQ'),
        ('bullet', 'Q：为什么 P_fc 归一化到 [0,1] 而不是直接用 kW？→ 神经网络对固定范围输入更稳定，且方便换不同功率等级的燃料电池（改乘数即可），解耦了"策略"和"物理量纲"。', 0),
        ('bullet', 'Q：SOC 更新为什么没乘 dt？→ 代码把每步时间隐含为 1 分钟（dt=1），功率差×1分钟≈能量增量，直接除以容量即得 SOC 变化，是教学环境的简化。真实系统需严格积分。', 0),
        ('bullet', 'Q：DQN_Continuous 输出的 sigmoid 值算什么？→ 什么都不算——它既不是 Q 值（没法 argmax），也不是策略采样（没有分布概念），只是"监督式硬映射"，这正是演示要揭示的荒谬。', 0),
        ('bullet', 'Q：200 步一局是什么量级？→ 每步 1 分钟，200 步≈3.3 小时。对应一次短途航行/行驶的负荷时长。', 0),
        ('bullet', 'Q：正弦负载曲线合理吗？→ 只是教学简化。真实负载是实测序列，但正弦能模拟周期性起伏（如加速/巡航交替），足以演示"负载在变"这一核心矛盾。', 0),
    ],
    'Week11_Step1_ContinuousEnv_逐行精讲.docx'
)


# =====================================================================
# File 2: week11_reinforce.py —— REINFORCE (策略梯度)
# =====================================================================
build_doc(
    'Step 2: REINFORCE (策略梯度)',
    'week11_reinforce.py — 第一个能处理连续动作的 RL 算法',
    [
        # ============ 一、文件概览 ============
        ('h1', '一、文件概览'),
        ('p', '本文件实现 REINFORCE——也叫"蒙特卡洛策略梯度"或"Vanilla Policy Gradient"。它是解决 Step 1 遗留问题（连续动作）的第一个算法。核心思想一句话：**直接学习一个策略 π(a|s)（输出动作分布），跑完一整局，用真实的整局回报 G 来评价每个动作好坏，好的增大概率、坏的减小概率**。', 11),
        ('callout', '一句话记忆：REINFORCE = ① 一个输出动作分布的策略网络 PolicyNet（π(s)→[μ,σ]）+ ② 跑完一局 → 从后往前算 G_t → loss = -Σ log π(a|s)·G_t → 更新策略。'),
        ('h2', '1.1 与 DQN 的本质区别'),
        ('tbl', ['维度', 'DQN（Week10）', 'REINFORCE（本文件）'], [
            ['网络输出', 'Q(s) → N 个 Q 值（每个动作一个）', 'π(s) → [μ, σ]（动作分布参数）'],
            ['选动作', 'argmax Q(s)，遍历所有离散动作', '从 Normal(μ,σ) 采样，无需遍历'],
            ['动作类型', '离散（GridWorld 的 4 个方向）', '连续（P_fc ∈ [0,1]）'],
            ['价值估计', 'Q(s,a) 用 TD 学习，每步可更新', 'G_t = 整局实际回报，须等整局结束'],
            ['更新时机', '每步更新', '整局跑完才更新一次'],
            ['核心公式', 'MSE(Q, r+γ·max Q\')', '-Σ log π(a|s)·G_t'],
        ]),

        ('h2', '1.2 本文件的结构'),
        ('tbl', ['部分', '行号', '内容'], [
            ['import 与配置', '15-29', '导入库 + matplotlib 中文字体配置'],
            ['EMSEnv 精简环境', '32-66', '复用 Step1 的环境（略精简）'],
            ['PolicyNet 策略网络', '70-108', 'π(s)→[μ,σ]，含 get_action/evaluate'],
            ['reinforce() 主算法', '112-193', '跑局→算 G→标准化→loss→更新'],
            ['test_policy()', '197-219', '用训练好的策略跑 10 局评估'],
            ['plot_results()', '223-252', '画训练曲线（原始+平滑+分块均值）'],
            ['main()', '256-294', 'CLI 入口'],
        ]),

        # ============ 二、import 精讲 ============
        ('brk',),
        ('h1', '二、import 与配置（15-29 行）'),
        ('x', 15, 'import argparse', '导入 argparse，用于 --episodes/--lr/--seed/--output-dir 命令行参数。'),
        ('x', 16, 'from pathlib import Path', '导入 Path，用于类型标注（output_dir 参数类型）。'),
        ('x', 18, 'import matplotlib', '导入 matplotlib 顶层模块。'),
        ('x', 19, "matplotlib.use('Agg')", '★ 选择 Agg 后端（无窗口渲染）。在服务器/无显示器环境（或自动化跑批）下，matplotlib 默认 GUI 后端会报错；Agg 直接输出图片文件，不弹窗口。'),
        ('x', 20, 'import matplotlib.pyplot as plt', '导入 pyplot，用于画训练曲线。'),
        ('x', 25, 'import torch.distributions as dist', '★ 导入概率分布库。策略采样和 log_prob 计算全靠它（dist.Normal）。这是本文件处理"连续动作"的关键工具。'),
        ('x', 27, 'from week11_common import configure_matplotlib, ensure_results_dir, set_seed', '从共享模块导入三个工具：configure_matplotlib（设置中文字体）、ensure_results_dir（确保输出目录存在）、set_seed（设随机种子）。'),
        ('x', 29, 'configure_matplotlib()', '模块加载时立即配置 matplotlib 中文字体（在 Week11 批量生成图片时避免中文变方框）。'),

        # ============ 三、环境 ============
        ('brk',),
        ('h1', '三、EMSEnv 精简环境（32-66 行）'),
        ('p', '和 Step 1 的 EMSEnv 几乎一样，只是做了精简：去掉了 dt、render()、max_steps 拆分逻辑等。逐行快读（相同处不再赘述，只讲差异）：', 11),
        ('x', 37, 'self.state_dim = 2', '状态维度 2：同为 [SOC, P_load]。'),
        ('x', 46, 'return np.array([self.soc, self.p_load], dtype=np.float32)', 'reset() 直接返回状态数组（少了 _get_state 的二次封装调用，代码更紧凑）。'),
        ('x', 49, 'p_fc = float(np.clip(action, 0, 1)) * 30.0', '反归一化，与 Step1 完全一致。'),
        ('x', 50, 'self.p_load = 0.3 + 0.4 * (0.5 + 0.5 * np.sin(self.steps * 0.1))', '负载正弦波，与 Step1 完全一致。'),
        ('x', 57, 'soc_bound_penalty = -1.0 if (self.soc <= self.soc_min or self.soc >= self.soc_max) else 0.0', '★ 越界惩罚改成了单行三元表达式写法（Step1 是 if 语句多行），逻辑完全相同，只是代码更紧凑。'),
        ('x', 63, "return self._get_state(), reward, done, {'p_fc': p_fc}", '返回四元组，info 只保留 p_fc（去掉了 fuel_cost/soc，简化）。'),
        ('x', 65, 'def _get_state(self):', '内部方法，把 SOC 和负载打包成状态数组。'),
        ('callout', '环境被两个文件重复写了一遍——这是刻意为之（每步教学文件自包含，方便单独阅读/运行），不是好工程实践。真实项目应 import 共享。但作为教学脚本，自包含更容易跑通。'),

        # ============ 四、PolicyNet ============
        ('brk',),
        ('h1', '四、策略网络 PolicyNet 逐行精讲（70-108 行）—— 本文件核心'),
        ('p', '这是 REINFORCE 和 DQN 的根本分水岭。DQN 输出"每个动作的 Q 值"；PolicyNet 输出"动作分布的参数 [μ, σ]"。理解了它，就理解了策略梯度方法的大半。', 12, True),

        ('h2', '4.1 网络结构（70-91 行）'),
        ('x', 70, 'class PolicyNet(nn.Module):', '定义策略网络。与 DQN 的 TinyDQN 本质区别：TinyDQN 输出 4 个 Q 值（离散动作各一个），PolicyNet 输出动作分布参数。'),
        ('x', 71, '"""策略网络 π_θ(s) → [μ, σ]', 'docstring 点明输出。'),
        ('x', 73, 'DQN 输出 Q 值（每个动作一个值）', '对比 DQN。'),
        ('x', 74, '策略网络输出动作分布的参数（μ 和 σ）', '策略网络输出：均值 μ + 标准差 σ，共同定义一个高斯分布。'),
        ('x', 77, 'def __init__(self, state_dim=2, hidden=64, action_dim=1):', '构造函数。输入状态维度 2，隐藏层 64，动作维度 1。'),
        ('x', 79, 'self.fc1 = nn.Linear(state_dim, hidden)', '第 1 层全连接：2 维状态 → 64 维隐藏。'),
        ('x', 80, 'self.fc2 = nn.Linear(hidden, hidden)', '第 2 层全连接：64 → 64。相比 DQN 的 TinyDQN（一层隐藏），这里多一层——连续控制任务的映射更复杂。'),
        ('x', 81, '# 输出层分成两路：均值 μ + 对数标准差 log_std', '注释说明输出层两路结构。'),
        ('x', 82, 'self.mean_head = nn.Linear(hidden, action_dim)', '★ 均值输出头：64 维 → 1 维，输出动作均值 μ。μ 表示"策略当前倾向于输出的动作值"。'),
        ('x', 83, 'self.log_std = nn.Parameter(torch.zeros(action_dim))  # 可训练的标准差', '★ 对数标准差 log_std：这是**可训练的参数**（nn.Parameter），不是网络层输出。初始化全 0 → std = e⁰ = 1。为什么存 log_std 而不是 std？因为 std 必须 > 0，直接梯度下降可能把它推到负值；而对 log_std 做无约束优化，e^(log_std) 自动保证正数。std 控制探索范围：std 大 → 动作随机（多探索），std 小 → 动作集中（多利用）。训练中会自动调整。'),

        ('h2', '4.2 forward()（85-91 行）'),
        ('x', 85, 'def forward(self, x):', '前向传播。'),
        ('x', 86, 'x = torch.relu(self.fc1(x))', '第 1 层 ReLU 激活。'),
        ('x', 87, 'x = torch.relu(self.fc2(x))', '第 2 层 ReLU 激活。'),
        ('x', 88, 'mean = torch.tanh(self.mean_head(x))  # tanh → [-1, 1]', '★ tanh 把均值压到 [-1,1]。tanh 以 0 为中心、有梯度饱和区，适合作为输出激活（比 ReLU 输出有界）。'),
        ('x', 89, 'mean = (mean + 1) / 2  # 映射到 [0, 1]', '★ 把 [-1,1] 平移缩放到动作范围 [0,1]：(m+1)/2。这样均值 μ 天然落在有效动作区间。'),
        ('x', 90, 'std = torch.exp(self.log_std.clamp(-5, 2))  # 保证正数', '★ 标准差 = e^clamp(log_std, -5, 2)。指数运算保证正数；clamp 限制 log_std ∈ [-5,2]，避免 e^5≈148（过大）或 e^-5≈0.007（过小）的数值溢出/退化。'),
        ('x', 91, 'return mean, std', '返回 (μ, σ) 元组，构成高斯分布。'),

        ('h2', '4.3 get_action()（93-101 行）'),
        ('x', 93, 'def get_action(self, state):', '★ 选动作——和 DQN 差异最大的一步。DQN：前向传播 → 4 个 Q 值 → argmax。这里：前向传播 → 高斯分布 → 采样。不用 argmax，自然能处理连续动作。'),
        ('x', 94, '"""选动作：从策略分布中采样"""', 'docstring。'),
        ('x', 95, 'with torch.no_grad():', '★ 关闭梯度：采样过程不需要反向传播（这是"行为"不是"训练"）。'),
        ('x', 96, 's = torch.FloatTensor(state).unsqueeze(0)', '状态转张量并加 batch 维，shape [1,2]。'),
        ('x', 97, 'mean, std = self.forward(s)', '前向得到分布参数。'),
        ('x', 98, 'm = dist.Normal(mean, std)', '★ 用 PyTorch 概率库构建高斯分布 N(μ,σ)。'),
        ('x', 99, 'a = m.sample()', '★ 从分布中采样一个动作值。即便输入同一状态，每次采样都可能不同（策略的随机性 = 探索）。'),
        ('x', 100, 'a = a.clamp(0, 1)', '把采样值夹到 [0,1] 有效动作范围（高斯分布理论上可能采到负数或 >1）。'),
        ('x', 101, 'return a.item(), (state.copy(), a.item())', '返回两样东西：① a.item()——动作的 Python float；② (state.copy(), a.item())——"trace"（轨迹记录），供训练时重放 (s,a)。state.copy() 用浅拷贝，防止后续状态被覆盖影响记录。'),

        ('h2', '4.4 evaluate()（103-108 行）'),
        ('x', 103, 'def evaluate(self, state, action):', '★ 给定状态和动作，计算"在当前策略下这个动作的对数概率 log π(a|s)"。这是训练用的（带梯度）。'),
        ('x', 104, '"""给定状态和动作，算 log_prob（用于训练时计算 loss，有梯度）"""', 'docstring 强调"有梯度"。'),
        ('x', 105, 'mean, std = self.forward(state)', '前向得到分布参数。'),
        ('x', 106, 'm = dist.Normal(mean, std)', '构建高斯分布。'),
        ('x', 107, 'log_prob = m.log_prob(action)', '计算该动作的对数概率密度 log π(a|s)。'),
        ('x', 108, 'return log_prob', '返回 log_prob。它在 loss 里作为"这个动作有多可能"的梯度载体。'),

        # ============ 五、reinforce 主算法 ============
        ('brk',),
        ('h1', '五、reinforce() 主算法逐行精讲（112-193 行）'),
        ('p', '整个算法的骨架在 docstring 里写得很清楚，我们把它拆成 5 步走查：跑局 → 算 G → 标准化 → 算 loss → 更新。', 12, True),

        ('h2', '5.1 函数签名与初始化（112-133 行）'),
        ('x', 112, 'def reinforce(episodes=500, lr=0.001):', '主函数。500 局（比 DQN 的 5000 局少——环境简单、每局信息量大），学习率 0.001（比 DQN 的 0.01 小——策略梯度对步长更敏感，大了容易崩）。'),
        ('x', 123, 'env = EMSEnv()', '创建环境。'),
        ('x', 124, 'policy = PolicyNet()', '★ 创建策略网络——这是**唯一的**网络。REINFORCE 没有价值网络（Critic），不学 Q 值，只学"给定状态该输出什么动作"。对比 DQN 有 q_network + target_network 两个网络。'),
        ('x', 125, 'optimizer = optim.Adam(policy.parameters(), lr=lr)', 'Adam 优化器。优化目标与 DQN 不同：DQN 最小化 Q 值预测误差（MSE），REINFORCE 最大化期望回报（等价于最小化 -logπ·G）。'),
        ('x', 127, 'episode_rewards = []', '记录每局总奖励，画图用。'),
        ('x', 128, 'episode_lengths = []', '记录每局步数，画图/观察用。'),
        ('x', 131, 'print(f"  策略网络: {2}维状态 → {64}隐藏 → [μ, σ] → 采样动作")', '打印网络结构概览（用硬编码的 2/64 展示，教学脚本追求可读性）。'),

        ('h2', '5.2 第 1 步：跑一整局（135-150 行）'),
        ('x', 135, 'for ep in range(1, episodes + 1):', '外层训练循环：逐局训练。'),
        ('x', 136, 's = env.reset()', '每局重置环境。'),
        ('x', 137, 'log_probs = []', '声明用于存 log_prob 的列表（本文件实际用 trace 替代，保留声明）。'),
        ('x', 138, 'rewards = []', '存本局每一步的奖励。'),
        ('x', 139, 'done = False', '初始化终止标志。'),
        ('x', 141, 'transitions = []  # 存 (状态, 动作, 奖励)', '★ 记录本局所有 (状态, 动作, 奖励) 三元组。这是"蒙特卡洛"的载体——等整局跑完，用这些数据回溯。'),
        ('x', 143, '# ---- 1. 跑一局，记录 (s, a, r) ----', '注释标记步骤 1。'),
        ('x', 144, 'while not done:', '内层循环：一直走到环境终止（200 步或 SOC 越界）。'),
        ('x', 145, 'a, trace = policy.get_action(s)', '★ 用当前策略选动作。返回连续值（如 0.63）+ 轨迹记录。这就是 REINFORCE 能处理连续动作的原因——输出是连续值。'),
        ('x', 146, 's_a, a_val = trace', '解包轨迹：s_a 是采动作时的状态拷贝，a_val 是动作值。'),
        ('x', 147, 'sp, reward, done, _ = env.step(a)', '执行动作，环境返回新状态/奖励/终止标志。'),
        ('x', 148, 'transitions.append((s_a, a_val, reward))', '★ 把 (状态, 动作, 奖励) 存入轨迹。注意存的是采样时的状态 s_a（拷贝），而不是执行后的 sp——因为 loss 要用"当时做决策的状态"。'),
        ('x', 149, 'rewards.append(reward)', '把奖励也单独存一份（算 G_t 用）。'),
        ('x', 150, 's = sp', '状态推进，进入下一步循环。'),

        ('h2', '5.3 第 2 步：算 G_t（152-159 行）'),
        ('x', 152, '# ---- 2. 算 G_t（从后往前累加）----', '注释标记步骤 2。'),
        ('x', 153, 'G = 0', '初始化累计回报 G。'),
        ('x', 154, 'returns = []', '存每步的 G_t。'),
        ('x', 155, 'for r in reversed(rewards):', '★ 从最后一步往前遍历。为什么反着来？因为 G_t = r_t + γ·r_{t+1} + γ²·r_{t+2} + ...，后面的回报要先算出来才能叠加给前面的。'),
        ('x', 156, 'G = r + 0.99 * G', '★ 核心递推：G ← r + γ·G。从终点开始：最后一步 G = r_last；倒数第二步 G = r_{prev} + 0.99·r_last；依此类推。0.99 是折扣因子 γ——越远的未来回报折得越狠。'),
        ('x', 157, 'returns.insert(0, G)', 'insert(0, G) 把结果插到列表头部，保证 returns 的索引和 rewards 一一对应（因为遍历是反的，需要逆序放回）。'),
        ('x', 159, 'returns_t = torch.FloatTensor(returns)', '转成 PyTorch 张量，便于后面做张量运算和标准化。'),

        ('h2', '5.4 第 3 步：标准化 G_t（161-163 行）'),
        ('x', 161, '# ---- 3. 标准化 G_t（减均值÷标准差，稳定训练）----', '注释标记步骤 3。'),
        ('x', 162, 'if len(returns_t) > 1:', '只有一局 >1 步时才标准化（防除以 0）。'),
        ('x', 163, 'returns_t = (returns_t - returns_t.mean()) / (returns_t.std() + 1e-8)', '★ 把所有 G_t 变成均值 0、标准差 1。+1e-8 防除零。为什么要标准化？不同局的 G_t 绝对值差异很大（有的全负、有的全正）。若不做，可能某局所有 G_t 全为正 → 所有动作都增大概率，等于没学。标准化后正负分界清晰：G>0 好动作（增大）、G<0 坏动作（减小），尺度统一，训练更稳。'),

        ('h2', '5.5 第 4 步：算 loss（165-176 行）'),
        ('x', 165, '# ---- 4. loss = -Σ log_prob(a|s) × G_t ----', '注释标记步骤 4。'),
        ('x', 166, '#  关键：把 (s, a) 重新送进网络算 log_prob（带梯度）', '★ 强调：虽然 get_action 时也算过概率，但那是 no_grad 的。这里重新前向一次（带梯度），让梯度能流回网络。'),
        ('x', 167, 'loss = 0', '初始化 loss 为 0。'),
        ('x', 168, 'for (s_i, a_i, _), G_i in zip(transitions, returns_t):', '遍历每一步：(状态, 动作, 奖励) 与对应的 G_t 配对。'),
        ('x', 169, 's_t = torch.FloatTensor(s_i).unsqueeze(0)', '状态转张量 + batch 维。'),
        ('x', 170, 'a_t = torch.FloatTensor([a_i]).unsqueeze(0)', '动作转张量 + batch 维。'),
        ('x', 171, 'log_prob = policy.evaluate(s_t, a_t)', '★ 带梯度计算 log π(a|s)。'),
        ('x', 172, 'loss = loss + (-log_prob * G_i)  # G>0 → 增大概率，G<0 → 减小概率', '★ 核心公式累加：loss += -log π(a|s) × G。展开解释：若 G>0（好回报），最小化 -logπ·G 意味着让 -logπ 变小即 logπ 变大 → 该动作概率上升；若 G<0，则概率下降。这就是"好事多干、坏事少干"。'),
        ('x', 174, 'loss = loss / len(transitions)', '对整局的步数取平均（否则长局 loss 天然大，局间不可比）。'),
        ('x', 175, '#    ↑ 如果 G>0（好回报），增大 log_prob（增大这个动作的概率）', '注释重申核心语义。'),
        ('x', 176, '#      如果 G<0（坏回报），减小 log_prob（减小这个动作的概率）', '注释重申核心语义。'),

        ('h2', '5.6 第 5 步：梯度更新（178-193 行）'),
        ('x', 178, 'optimizer.zero_grad()', '清空上一轮梯度。PyTorch 梯度默认累积，不清会叠加错误。'),
        ('x', 179, 'loss.backward()', '反向传播，计算每个参数的梯度。语法和 DQN 一样，含义不同：DQN 的梯度让 Q 值更准，REINFORCE 的梯度让策略更好。'),
        ('x', 180, 'optimizer.step()', '用梯度更新参数。'),
        ('x', 182, 'episode_rewards.append(sum(rewards))', '记录本局总奖励。'),
        ('x', 183, 'episode_lengths.append(len(rewards))', '记录本局步数。'),
        ('x', 186, 'if ep % 50 == 0:', '每 50 局打印一次进度。'),
        ('x', 187, 'avg_r = np.mean(episode_rewards[-50:])', '最近 50 局平均奖励。'),
        ('x', 188, 'avg_len = np.mean(episode_lengths[-50:])', '最近 50 局平均步数。'),
        ('x', 189, 'print(f"  第{ep:4d}/{episodes}局 | 平均奖励={avg_r:+.3f} | "', '打印训练进度：局数、平均奖励（带符号）。'),
        ('x', 190, 'f"平均步数={avg_len:.0f} | "', '打印平均步数。'),
        ('x', 191, 'f"log_std={policy.log_std.item():.3f}")', '★ 打印 log_std 训练值——观察探索节奏。log_std 下降说明策略在"收敛确定性"（从随机探索走向利用）。'),
        ('x', 193, 'return policy, episode_rewards, episode_lengths', '返回训练好的策略 + 训练曲线数据。'),

        # ============ 六、测试与画图 ============
        ('brk',),
        ('h1', '六、测试与画图（197-252 行）'),
        ('h2', '6.1 test_policy()（197-219 行）'),
        ('x', 197, 'def test_policy(policy, episodes=10):', '评估：用训练好的策略跑 10 局（注意：不在线更新，只测表现）。'),
        ('x', 202, 'env = EMSEnv()', '新建环境（独立于训练环境，干净评估）。'),
        ('x', 206, 'for t in range(200):', '每局最多 200 步。'),
        ('x', 207, 'a, _ = policy.get_action(s)', '用训练好的策略选动作（这里丢弃 trace，只取动作）。'),
        ('x', 208, 'sp, r, done, info = env.step(a)', '执行。'),
        ('x', 211, 'if done:', '若环境终止。'),
        ('x', 212, 'break', '提前结束。'),
        ('x', 218, 'print(f"  平均总奖励: {np.mean(total_rewards):+.3f}")', '打印 10 局平均总奖励——这是衡量策略好坏的最终指标。'),
        ('x', 219, 'return np.mean(total_rewards)', '返回平均奖励，供调用方比较。'),

        ('h2', '6.2 plot_results()（223-252 行）'),
        ('x', 223, 'def plot_results(rewards, label=\'REINFORCE\', output_dir: str | Path | None = None):', '画训练曲线函数。output_dir 支持传路径（来自 CLI），None 则用默认 results/。'),
        ('x', 224, 'plt.figure(figsize=(10, 4))', '画布大小 10×4 英寸（宽图，左右两个子图）。'),
        ('x', 227, 'plt.subplot(1, 2, 1)', '左子图：训练曲线。'),
        ('x', 228, "plt.plot(rewards, alpha=0.3, color='blue')", '原始曲线（半透明），展示噪声。'),
        ('x', 230, 'window = 20', '平滑窗口 20。'),
        ('x', 231, 'smoothed = np.convolve(rewards, np.ones(window)/window, mode=\'valid\')', '★ 滑动平均：用长度 20 的均值核做卷积，得到平滑曲线。np.ones(20)/20 是均值为 1/20 的核，mode=\'valid\' 只保留完整窗口的输出。'),
        ('x', 232, "plt.plot(smoothed, color='blue', linewidth=2)", '画平滑曲线（粗线）。'),
        ('x', 235, "plt.title(f'{label} 训练曲线')", '标题。'),
        ('x', 236, 'plt.grid(alpha=0.3)', '浅色网格，辅助读数。'),
        ('x', 239, 'plt.subplot(1, 2, 2)', '右子图：每 50 局均值柱状图。'),
        ('x', 240, 'chunk = 50', '分块大小 50。'),
        ('x', 241, 'means = [np.mean(rewards[i:i+chunk]) for i in range(0, len(rewards), chunk)]', '把 rewards 按每 50 局切成块，每块取均值。'),
        ('x', 242, "plt.bar(range(len(means)), means, color='blue', alpha=0.7)", '画柱状图。'),
        ('x', 248, 'plt.tight_layout()', '自动调整子图间距，避免标签重叠。'),
        ('x', 249, "path = ensure_results_dir(output_dir) / 'week11_reinforce_training.png'", '★ 用 week11_common 的 ensure_results_dir 确保目录存在，输出固定文件名。'),
        ('x', 250, 'plt.savefig(path, dpi=150)', '保存图片，dpi 150（较高清）。'),
        ('x', 252, 'plt.close()', '关闭画布，释放内存（防止 GUI 环境下重复弹窗/内存泄漏）。'),

        # ============ 七、主程序 ============
        ('brk',),
        ('h1', '七、main() 主程序（256-294 行）'),
        ('x', 257, "parser = argparse.ArgumentParser(description='Week 11 REINFORCE demo on a simplified continuous EMS environment')", '创建参数解析器。'),
        ('x', 258, "parser.add_argument('--episodes', type=int, default=500, help='Training episodes')", '训练局数（默认 500）。'),
        ('x', 259, "parser.add_argument('--lr', type=float, default=0.001, help='Learning rate')", '学习率（默认 0.001）。'),
        ('x', 260, "parser.add_argument('--seed', type=int, default=42, help='Random seed')", '随机种子（默认 42）。'),
        ('x', 261, "parser.add_argument('--output-dir', type=Path, default=None, help='Directory for generated figures')", '输出目录（默认 None → 用 results/）。'),
        ('x', 262, 'args = parser.parse_args()', '解析参数。'),
        ('x', 264, 'set_seed(args.seed)', '设置随机种子，保证可复现。'),
        ('x', 273, 'policy, rewards, lengths = reinforce(episodes=args.episodes, lr=args.lr)', '训练。'),
        ('x', 274, 'test_policy(policy)', '评估训练好的策略。'),
        ('x', 275, 'plot_results(rewards, output_dir=args.output_dir)', '画训练曲线。'),
        ('x', 281, 'print("""', '打印 DQN vs REINFORCE 的对照总结（多行字符串）。'),
        ('x', 282, 'DQN:        Q(s) → [Q↑, Q↓, Q←, Q→] → argmax → 选动作', 'DQN 路径回顾。'),
        ('x', 283, '↑ 只能处理离散动作', 'DQN 局限。'),
        ('x', 285, 'REINFORCE:  π(s) → [μ, σ] → Normal(μ,σ) → 采样得动作', 'REINFORCE 路径。'),
        ('x', 286, '↑ 输出的动作分布，可以取任意连续值', 'REINFORCE 优势。'),
        ('x', 288, 'print("  REINFORCE 的问题：要等整局跑完才能更新")', '★ 预告下一步：REINFORCE 的最大短板是"蒙特卡洛式——整局跑完才更新一次，方差大、样本利用率低"。'),
        ('x', 289, 'print("  → 下一节 Actor-Critic 解决这个问题")', '引出 Step 3。'),
        ('x', 293, "if __name__ == '__main__':", '入口惯用法。'),
        ('x', 294, 'main()', '启动。'),

        # ============ 八、数学推导 ============
        ('brk',),
        ('h1', '八、数学推导专题'),
        ('h2', '8.1 策略梯度定理（从直觉到公式）'),
        ('p', '强化学习的目标是最大化期望回报 J(θ) = E[Σ γ^t r_t]。策略梯度定理指出：', 11),
        ('flow', '''∇θ J(θ) = E[ Σ_t  ∇θ log πθ(a_t|s_t) · G_t ]

含义：
  · ∇θ log πθ(a|s)：策略输出该动作的"概率变化方向"
  · G_t：这个动作带来的"实际回报"（标量，作为权重）
  · 乘积后求期望：好回报(大G)的动作，沿增大概率方向更新；
    坏回报(小/负G)的动作，沿减小概率方向更新'''),
        ('p', '代码对应：第 172 行 loss = -log_prob * G_i 正是上式取负号后的逐样本形式（取负是因为 PyTorch 做梯度下降，而我们要最大化回报）。', 11, True),

        ('h2', '8.2 G_t 递推的正确读法'),
        ('p', 'G_t = r_t + γ·r_{t+1} + γ²·r_{t+2} + ...，代码用反向遍历高效实现：', 11),
        ('flow', '''G = 0
for r in reversed(rewards):   # 从最后一刻往前
    G = r + 0.99 * G          # G_new = r + γ·G_old
    returns.insert(0, G)

例（γ=0.99，rewards=[r0,r1,r2]）：
  第一步（i=2，反向第一个）：G = r2
  第二步（i=1）：            G = r1 + 0.99·r2
  第三步（i=0）：            G = r0 + 0.99·r1 + 0.9801·r2
  → returns = [G_t0, G_t1, G_t2] 与原索引对齐'''),

        ('h2', '8.3 为什么标准化 G_t 是"锦上添花"但几乎必需'),
        ('p', '若某一局所有奖励都为负（早期策略常如此），所有 G_t 都为负 → loss = -Σ logπ·(负数) 会让所有 logπ 同时增大 → 策略盲目增大概率，学习信号退化。标准化把均值归 0：正 G 增大概率、负 G 减小概率，训练方向立刻清晰。实践中策略梯度几乎都会加这一步。', 11),

        # ============ 九、训练走查 ============
        ('brk',),
        ('h1', '九、一局完整训练走查'),
        ('flow', '''【开始第 ep 局】
  reset() → s = [0.60, 0.50]

【循环跑局，直到 done】
  ① a, trace = policy.get_action(s)
       → 网络输出 μ≈0.5, σ≈1 → 采样 a≈0.63
  ② sp, r, done = env.step(a)
       → p_fc=18.9kW, p_load≈0.5kW → 充电 → r≈-0.19-0.005
  ③ transitions 记 (s_0.6, 0.63, -0.19)； rewards 记 -0.19
  ④ s = sp；重复直到 200 步或 SOC 出界

【算 G_t（反向递推）】→ returns = [G_0, ..., G_T]
【标准化】→ 均值为 0，方差 1
【算 loss】loss = -Σ log π(a_t|s_t)·G_t（取平均）
【更新】zero_grad → backward → step
【记录】episode_rewards / episode_lengths'''),

        # ============ 十、面试要点与 FAQ ============
        ('brk',),
        ('h1', '十、面试要点'),
        ('bullet', '「REINFORCE 和 DQN 的核心区别？」→ DQN 输出 Q 值 + argmax（只能离散动作）；REINFORCE 输出动作分布 [μ,σ] + 采样（可连续动作）。', 0),
        ('bullet', '「为什么用 log_prob 而不是概率？」→ ① 计算方便（softmax/gaussian 下 log 抵消指数）；② 数值稳定；③ 梯度 ∇logπ = ∇π/π 天然消除了"概率低导致梯度小"的偏置。', 0),
        ('bullet', '「REINFORCE 的缺点？」→ 蒙特卡洛：必须等整局结束，样本利用率低、方差大；学得慢。', 0),
        ('bullet', '「log_std 为什么初始化为 0？」→ log_std=0 → std=e⁰=1，初始分布较宽，探索充分。训练中自动收缩。', 0),
        ('bullet', '「为什么用 tanh 压缩动作均值？」→ 输出有界 [-1,1]，避免动作均值跑到 [0,1] 之外；配合 (m+1)/2 映射到有效区间。', 0),

        ('h1', '十一、常见疑问 FAQ'),
        ('bullet', 'Q：为什么采样时用 no_grad，训练时又算一次？→ 采样只是"做动作"（行为），不需要梯度；训练要回传梯度，所以重新前向一次。两次的目的不同。', 0),
        ('bullet', 'Q：标准化会不会改变策略梯度无偏性？→ 理论上引入轻微偏差（因为统计量依赖样本），但方差下降显著，实践中公认收益更大。', 0),
        ('bullet', 'Q：episodes=500 够吗？→ 环境简单，500 局足够观察到收敛趋势。调参可加 --episodes 实验。', 0),
        ('bullet', 'Q：为什么 trace 里要存 state.copy()？→ 后续 s 会被覆盖（s=sp），如果不拷贝，记录里存的就变成最新状态，loss 会算错。', 0),
        ('bullet', 'Q：REINFORCE 是 on-policy 吗？→ 是。策略更新后数据就"过期"了，所以每局数据只用一次（与 PPO 的多次重用不同）。', 0),
    ],
    'Week11_Step2_REINFORCE_逐行精讲.docx'
)


# =====================================================================
# File 3: week11_actor_critic.py —— Actor-Critic (演员-评委)
# =====================================================================
build_doc(
    'Step 3: Actor-Critic (演员-评委)',
    'week11_actor_critic.py — 每步更新，不再等整局结束',
    [
        # ============ 一、文件概览 ============
        ('h1', '一、文件概览'),
        ('p', 'Actor-Critic（AC）解决 REINFORCE 的最大痛点：**REINFORCE 要等一整局跑完、用实际的 G_t 才能评价动作好坏；AC 引入一个 Critic（价值网络），每走一步就能估计"这个状态值多少钱"，从而即时计算 Advantage 并更新**。名字很形象：Actor 是演员（负责出招），Critic 是评委（负责打分），评委实时点评，演员边演边改。', 11),
        ('callout', '一句话记忆：AC = REINFORCE（Actor 选动作） + 新增 Critic 网络 V(s) 打分期中成绩 → Advantage = r + γV(s\') - V(s) → 每步更新 Actor 和 Critic。'),
        ('h2', '1.1 与 REINFORCE 的核心对比'),
        ('tbl', ['维度', 'REINFORCE', 'Actor-Critic'], [
            ['策略网络', 'PolicyNet π(s)→[μ,σ]', 'Actor π(s)→[μ,σ]（一样）'],
            ['价值网络', '无', 'Critic V(s)（新增）'],
            ['评价标准', 'G_t = 整局真实回报（MC）', 'Advantage = r+γV(s\')-V(s)（TD）'],
            ['更新时机', '整局结束后才更新', '每一步都更新'],
            ['方差', '很大（G_t 覆盖整局波动大）', '较小（每步计算更平滑）'],
            ['偏差', '无偏（G_t 是真实值）', '有偏（V(s) 是估计值，可能不准）'],
            ['样本效率', '低（一局才学一次）', '高（每步都在学）'],
        ]),

        ('h2', '1.2 本文件的结构'),
        ('tbl', ['部分', '行号', '内容'], [
            ['import 与配置', '20-34', '导入库 + 中文字体配置'],
            ['EMSEnv 环境', '37-69', '同 Step2 的精简环境'],
            ['Actor 网络', '73-105', '策略网络，和 REINFORCE 的 PolicyNet 相同'],
            ['Critic 网络', '108-126', '价值网络 V(s)——AC 的新增核心'],
            ['actor_critic() 算法', '130-228', '每步算 Advantage → 双网络更新'],
            ['test_policy()', '232-252', '评估'],
            ['plot_results()', '256-281', '画训练曲线'],
            ['main()', '285-325', 'CLI 入口'],
        ]),

        # ============ 二、import ============
        ('brk',),
        ('h1', '二、import 与配置（20-34 行）'),
        ('x', 20, 'import argparse', 'CLI 参数解析。'),
        ('x', 24, "matplotlib.use('Agg')", '无窗口后端（批量出图）。'),
        ('x', 30, 'import torch.distributions as dist', '概率分布库（Actor 采样 / log_prob 用）。'),
        ('x', 32, 'from week11_common import configure_matplotlib, ensure_results_dir, set_seed', '共享工具：中文字体 / 输出目录 / 随机种子。'),
        ('x', 34, 'configure_matplotlib()', '模块加载即配置中文字体。'),

        # ============ 三、环境 ============
        ('brk',),
        ('h1', '三、EMSEnv 环境（37-69 行）'),
        ('p', '与 Step 2 完全相同的环境（SOC 动力学 + 三项奖励 + 200 步终止），逐行含义参见 Step1/Step2 文档。此处只点出关键行，不再重复展开：', 11),
        ('x', 54, 'p_fc = float(np.clip(action, 0, 1)) * 30.0', '动作反归一化 → 实际 kW。'),
        ('x', 55, 'self.p_load = 0.3 + 0.4 * (0.5 + 0.5 * np.sin(self.steps * 0.1))', '正弦负载 0.3-0.7 kW。'),
        ('x', 58, 'self.soc = np.clip(self.soc + soc_change, self.soc_min, self.soc_max)', 'SOC 积分 + 夹取。'),
        ('x', 61, 'soc_bound_penalty = -1.0 if (self.soc <= self.soc_min or self.soc >= self.soc_max) else 0.0', '越界重罚（三元表达式）。'),
        ('x', 66, "return self._get_state(), reward, done, {'p_fc': p_fc}", '标准四元组返回。'),

        # ============ 四、Actor 网络 ============
        ('brk',),
        ('h1', '四、Actor 网络（73-105 行）'),
        ('p', 'Actor 和 REINFORCE 的 PolicyNet 结构完全一致（输入状态 → 输出动作分布 [μ,σ] → 采样得连续动作）。快速走读：', 11),
        ('x', 73, 'class Actor(nn.Module):', '★ Actor（演员）网络。名字的由来：负责"出招"。与 REINFORCE 的 PolicyNet 结构完全一致。'),
        ('x', 75, 'def __init__(self, state_dim=2, hidden=64, action_dim=1):', '构造函数。'),
        ('x', 77, 'self.fc1 = nn.Linear(state_dim, hidden)', '输入层：2 → 64。'),
        ('x', 78, 'self.fc2 = nn.Linear(hidden, hidden)', '隐藏层：64 → 64。'),
        ('x', 79, 'self.mean_head = nn.Linear(hidden, action_dim)', '均值输出头：64 → 1。'),
        ('x', 80, 'self.log_std = nn.Parameter(torch.zeros(action_dim))', '对数标准差（可训练参数），初始 0 → std=1。'),
        ('x', 82, 'def forward(self, x):', '前向传播。'),
        ('x', 83, 'x = torch.relu(self.fc1(x))', 'ReLU 激活。'),
        ('x', 84, 'x = torch.relu(self.fc2(x))', 'ReLU 激活。'),
        ('x', 85, 'mean = torch.tanh(self.mean_head(x))', 'tanh → [-1,1]。'),
        ('x', 86, 'mean = (mean + 1) / 2  # [0, 1]', '映射到动作范围 [0,1]。'),
        ('x', 87, 'std = torch.exp(self.log_std.clamp(-5, 2))', '标准差（指数 + clamp 保证正数且不溢出）。'),
        ('x', 88, 'return mean, std', '返回分布参数。'),
        ('x', 90, 'def get_action(self, state):', '选动作（无梯度）。'),
        ('x', 92, 'with torch.no_grad():', '关闭梯度（纯行为采样）。'),
        ('x', 93, 's = torch.FloatTensor(state).unsqueeze(0)', '状态转张量 + batch 维。'),
        ('x', 94, 'mean, std = self.forward(s)', '前向。'),
        ('x', 95, 'm = dist.Normal(mean, std)', '构建高斯分布。'),
        ('x', 96, 'a = m.sample()', '采样动作。'),
        ('x', 97, 'a = a.clamp(0, 1)', '夹到 [0,1]。'),
        ('x', 98, 'return a.item()', '★ 注意：只返回动作值（相比 REINFORCE 的 get_action 少返回了 trace）。因为 AC 的训练数据不需要重放整局——每步即时用 V(s) 算 Advantage，无需回溯。'),
        ('x', 100, 'def evaluate(self, state, action):', '计算 log_prob（带梯度，训练用）。'),
        ('x', 102, 'mean, std = self.forward(state)', '前向。'),
        ('x', 103, 'm = dist.Normal(mean, std)', '构建分布。'),
        ('x', 104, 'log_prob = m.log_prob(action)', '对数概率。'),
        ('x', 105, 'return log_prob', '返回 log_prob。'),

        # ============ 五、Critic 网络 ============
        ('brk',),
        ('h1', '五、Critic 网络（108-126 行）—— 本文件核心新增'),
        ('p', 'Critic 是 AC 相比 REINFORCE 唯一新增的组件，但它一举解决了"等整局"的问题。类比：REINFORCE 是考完试看总分才知道学得怎么样，AC 是边做边有老师告诉你"这步做得对不对"。', 12, True),
        ('x', 108, 'class Critic(nn.Module):', '★ Critic（评委）网络。作用：估计状态值 V(s)="当前状态值多少钱"（期望未来累计奖励）。REINFORCE 没有它——REINFORCE 要等整局跑完用真实 G_t 评价；Critic 给出"预估值"，不用等结局。'),
        ('x', 111, 'REINFORCE 没有这个——它要等整局跑完才知道好坏。', 'docstring 再次强调对比。'),
        ('x', 112, 'Critic 的作用是"预估"当前状态值多少钱，不用等到整局结束。', '★ 核心定位：V(s) 是"预估价"，不是真实值。'),
        ('x', 115, 'def __init__(self, state_dim=2, hidden=64):', '构造函数。输入状态维度 2，隐藏 64。'),
        ('x', 117, 'self.net = nn.Sequential(', '顺序网络。'),
        ('x', 118, 'nn.Linear(state_dim, hidden),', '输入层：2 → 64。'),
        ('x', 119, 'nn.ReLU(),', '激活。'),
        ('x', 120, 'nn.Linear(hidden, hidden),', '隐藏层：64 → 64。'),
        ('x', 121, 'nn.ReLU(),', '激活。'),
        ('x', 122, 'nn.Linear(hidden, 1)  # 输出一个标量 V(s)', '★ 输出层：64 → 1，单个标量 V(s)。注意结构差异：Actor 输出 2 个值（μ 和 σ），Critic 只输出 1 个标量（状态价值）。'),
        ('x', 125, 'def forward(self, x):', '前向传播。'),
        ('x', 126, 'return self.net(x)', '直接返回 V(s) 预测值。'),

        # ============ 六、actor_critic 算法 ============
        ('brk',),
        ('h1', '六、actor_critic() 算法逐行精讲（130-228 行）'),
        ('p', '算法骨架见 docstring。核心是"每步循环里同时更新两个网络"。我们按每步流程走查。', 12, True),

        ('h2', '6.1 初始化（148-164 行）'),
        ('x', 148, 'env = EMSEnv()', '创建环境。'),
        ('x', 149, 'actor = Actor()', '创建 Actor 网络。'),
        ('x', 150, 'critic = Critic()', '★ 创建 Critic 网络——REINFORCE 没有的东西。'),
        ('x', 151, 'actor_opt = optim.Adam(actor.parameters(), lr=lr)', 'Actor 优化器（lr=0.001）。'),
        ('x', 152, 'critic_opt = optim.Adam(critic.parameters(), lr=lr * 2)  # Critic 学快一点', '★ Critic 优化器，学习率是 Actor 的两倍。为什么？Critic 要先"学好"才能给出可靠评价，评价准了 Actor 才学得好——所以让它跑快点。'),
        ('x', 153, 'loss_fn = nn.MSELoss()', 'MSE 损失：Critic 的训练指标（V(s) 与 target 的平方差）。'),
        ('x', 164, 'gamma = 0.99', '折扣因子，用于 Advantage 和 TD target。'),

        ('h2', '6.2 每步循环（166-216 行）'),
        ('x', 166, 'for ep in range(1, episodes + 1):', '训练局循环。'),
        ('x', 167, 's = env.reset()', '重置环境。'),
        ('x', 172, 'while not done:', '内层：一直走到终止。'),
        ('x', 174, 'a = actor.get_action(s)', '★ Actor 选动作（和 REINFORCE 一样采样得连续值）。'),
        ('x', 177, 'sp, reward, done, _ = env.step(a)', '执行动作。'),
        ('x', 180, 's_t = torch.FloatTensor(s).unsqueeze(0)', '当前状态 → 张量 + batch 维。'),
        ('x', 181, 'sp_t = torch.FloatTensor(sp).unsqueeze(0)', '下一状态 → 张量 + batch 维。'),
        ('x', 182, 'a_t = torch.FloatTensor([a]).unsqueeze(0)', '动作 → 张量 + batch 维。'),
        ('x', 183, 'r_t = torch.FloatTensor([reward])', '奖励 → 张量（标量包装）。'),

        ('h2', '6.3 计算 Advantage（185-193 行）—— 核心公式'),
        ('x', 185, '# ---- 4. 算 Advantage ----', '注释标记。'),
        ('x', 186, '#  A = r + γ·V(s\') - V(s)', '★ Advantage 公式。'),
        ('x', 187, '#  如果 A > 0：这一步比预期好 → 增大这个动作的概率', '语义：实际结果优于预期 → 鼓励。'),
        ('x', 188, '#  如果 A < 0：这一步比预期差 → 减小这个动作的概率', '语义：实际结果差于预期 → 抑制。'),
        ('x', 189, 'V_s = critic(s_t)', '★ Critic 估当前状态值 V(s)。'),
        ('x', 190, 'with torch.no_grad():', '关闭梯度：Advantage 是"评价"不是"训练目标"，不希望梯度流经 Critic 参与 Advantage 的误判（详见 8.3 注释）。'),
        ('x', 191, 'V_sp = critic(sp_t)', '估下一状态值 V(s\')。'),
        ('x', 192, 'advantage = r_t + gamma * V_sp * (not done) - V_s', '★ 核心公式：A = r + γ·V(s\')·(1-done) - V(s)。(not done) 处理终止步：若 done=True，则 (not done)=False → 0，V(s\') 被清零（终止后没有未来回报）。展开解释：r + γV(s\') 是"实际奖励 + 未来估值"（TD target），减去当前估值 V(s) 就是"这一步比预期好多少"。'),
        ('x', 193, '#   ↑ TD error = 实际奖励 + 未来估值 - 当前估值', '注释点明：Advantage 本质就是 TD error（时序差分误差）。'),

        ('h2', '6.4 更新 Actor（195-202 行）'),
        ('x', 196, 'log_prob = actor.evaluate(s_t, a_t)', '带梯度计算 log π(a|s)。'),
        ('x', 197, 'actor_loss = -(log_prob * advantage.detach()).mean()', '★ Actor loss = -log π(a|s) × Advantage。和 REINFORCE 公式形态完全一样，只是把 G_t 换成了 Advantage。.detach() 切断 Advantage 的梯度（Advantage 是"老师打分"，不该反向影响 Critic）。mean() 对 batch 平均（此处 batch=1，等于没平均，但保持写法一致）。'),
        ('x', 200, 'actor_opt.zero_grad()', '清梯度。'),
        ('x', 201, 'actor_loss.backward()', '反向传播。'),
        ('x', 202, 'actor_opt.step()', '更新 Actor 参数。'),

        ('h2', '6.5 更新 Critic（204-212 行）'),
        ('x', 204, '# ---- 6. 更新 Critic ----', '注释标记。'),
        ('x', 205, '#  让 V(s) 接近 r + γ·V(s\')', '★ Critic 的目标：学会预测"状态价值"，用 TD 目标逼近。'),
        ('x', 206, 'with torch.no_grad():', '关闭梯度：target 本身不作为可训练量（它依赖 V(s\')，若带梯度会形成循环依赖）。'),
        ('x', 207, 'td_target = r_t + gamma * V_sp * (not done)', '★ TD target = r + γV(s\')·(1-done)。这和 DQN 的 target 公式几乎一样！DQN 用 max Q(s\',a\')，这里用 Critic 的 V(s\')（少了一层 max，因为不看动作）。'),
        ('x', 208, 'critic_loss = loss_fn(V_s, td_target)', '★ MSE(V(s), td_target)：让 Critic 预测逼近 TD target。这就是时序差分（TD）学习——用"一步的真实奖励 + 下一步的估值"来校准"当前估值"。'),
        ('x', 210, 'critic_opt.zero_grad()', '清梯度。'),
        ('x', 211, 'critic_loss.backward()', '反向传播。'),
        ('x', 212, 'critic_opt.step()', '更新 Critic 参数。'),
        ('x', 214, 'total_reward += reward', '累计本局总奖励（记录用）。'),
        ('x', 215, 'steps += 1', '累计步数。'),
        ('x', 216, 's = sp', '状态推进。'),

        ('h2', '6.6 每局收尾（218-228 行）'),
        ('x', 218, 'episode_rewards.append(total_reward)', '记录本局奖励。'),
        ('x', 219, 'episode_lengths.append(steps)', '记录本局步数。'),
        ('x', 221, 'if ep % 50 == 0:', '每 50 局打印。'),
        ('x', 222, 'avg_r = np.mean(episode_rewards[-50:])', '近 50 局平均奖励。'),
        ('x', 226, 'f"log_std={actor.log_std.item():.3f}")', '打印 Actor 的探索参数。'),
        ('x', 228, 'return actor, episode_rewards, episode_lengths', '返回 Actor + 曲线数据。'),

        # ============ 七、测试与画图 ============
        ('brk',),
        ('h1', '七、测试与画图（232-281 行）'),
        ('x', 232, 'def test_policy(actor, episodes=10):', '评估：用训练好的 Actor 跑 10 局（不更新，纯测）。'),
        ('x', 240, 'for t in range(200):', '每局最多 200 步。'),
        ('x', 241, 'a = actor.get_action(s)', '用 Actor 选动作。'),
        ('x', 242, 'sp, r, done, info = env.step(a)', '执行。'),
        ('x', 251, 'print(f"  平均总奖励: {np.mean(total_rewards):+.3f}")', '输出平均奖励。'),
        ('x', 256, 'def plot_results(rewards, label=\'Actor-Critic\', output_dir: str | Path | None = None):', '画图函数（绿色系，与 REINFORCE 蓝色区分）。'),
        ('x', 261, 'smoothed = np.convolve(rewards, np.ones(window)/window, mode=\'valid\')', '滑动平均平滑。'),
        ('x', 278, "path = ensure_results_dir(output_dir) / 'week11_ac_training.png'", '输出文件名 week11_ac_training.png。'),
        ('x', 280, 'plt.savefig(path, dpi=150)', '保存。'),

        # ============ 八、机制专题 ============
        ('brk',),
        ('h1', '八、机制专题'),
        ('h2', '8.1 Advantage 为什么比 G_t 好？'),
        ('p', 'REINFORCE 用 G_t（整局累计回报）评价动作，G_t 含很多噪声（远未来的奖励波动大）。AC 用 A = r + γV(s\') - V(s)：V(s) 和 V(s\') 都是 Critic 的期望估计，差值只反映"这一步偏离预期的部分"，天然去掉了"这个状态本来平均就好/差"的基准。方差显著降低，学习更快。', 11),
        ('h2', '8.2 为什么更新 Critic 用 TD 而不是 MC？'),
        ('p', 'TD（时序差分）只用"一步奖励 + 下一步估值"，可以每步都更新（快、方差小）；代价是估值有偏差（因为用估计值校准估计值）。MC 用整局真实回报（无偏但慢、方差大）。AC 的哲学是"宁要快的近似，不要慢的精确"——因为 Critic 会随着训练越来越准。', 11),
        ('h2', '8.3 为什么 Advantage 和 td_target 都要 detach/no_grad？'),
        ('flow', '''关键点：这两个量都依赖 Critic 的输出 V(s\')。
· 若 Advantage 不带 detach：梯度会从 actor_loss 流经 advantage → 再流回 critic 网络，
  导致"更新 Actor 时顺便改 Critic"，两个网络互相牵扯，训练混乱。
· detach()/no_grad() 的本质：
    把 V(s') 当成"固定的打分"，Critic 的更新只走 critic_loss 这一条路，
    Actor 的更新只走 actor_loss 这一条路。职责分离，各学各的。'''),

        # ============ 九、训练走查 ============
        ('brk',),
        ('h1', '九、一局完整训练走查'),
        ('flow', '''【开始第 ep 局】reset() → s = [0.60, 0.50]

【每步循环】
  ① a = actor.get_action(s)          → 采样 a≈0.55
  ② sp, r, done = env.step(a)
  ③ s_t/sp_t/a_t/r_t = 转张量
  ④ V_s  = critic(s_t)               → 当前状态价值
     V_sp = critic(sp_t)（no_grad）  → 下一状态价值
     A = r + 0.99·V_sp·(not done) - V_s   ← 核心 Advantage
  ⑤ actor_loss = -log π(a|s)·A(detach) → 更新 Actor
  ⑥ td_target = r + 0.99·V_sp·(not done)
     critic_loss = MSE(V_s, td_target)   → 更新 Critic
  ⑦ s = sp；循环直到 done
【记录】episode_rewards / episode_lengths'''),

        # ============ 十、面试要点与 FAQ ============
        ('brk',),
        ('h1', '十、面试要点'),
        ('bullet', '「AC 和 REINFORCE 的区别？」→ AC 多了一个 Critic 网络 V(s)，用每步的 TD Advantage 取代整局 G_t，从"整局更新一次"变成"每步更新"，方差更小、学得更快；代价是引入了 Critic 估计的偏差。', 0),
        ('bullet', '「Advantage 公式及含义？」→ A = r + γV(s\') - V(s)。表示"这一步的实际结果比预期（V(s)）好多少"。A>0 鼓励，A<0 抑制。', 0),
        ('bullet', '「为什么 Critic 学习率是 Actor 的两倍？」→ Critic 要快速变准才能当好"评委"；评委准了，Actor 才有可靠信号。', 0),
        ('bullet', '「AC 的缺点？」→ Critic 是估计值，可能不准（有偏）；且仍可能"一步更新太多"导致策略突变崩溃——这正是 PPO 要修的问题。', 0),

        ('h1', '十一、常见疑问 FAQ'),
        ('bullet', 'Q：Critic 和 DQN 的 Q 网络有什么区别？→ DQN 输出 Q(s,a)（依赖动作），Critic 输出 V(s)（不看动作，只评估状态）。DQN 用 max 选最优动作，Critic 直接被 Actor 的动作所依赖。', 0),
        ('bullet', 'Q：为什么 (not done) 要乘在 V(s\') 上？→ 终止步之后没有未来回报，V(s\') 应为 0；否则会把"结束后的幻想价值"算进 Advantage。', 0),
        ('bullet', 'Q：AC 是 on-policy 还是 off-policy？→ on-policy（数据来自当前 Actor 采样，更新后数据即过期）。', 0),
        ('bullet', 'Q：Advantage 和 GAE 什么关系？→ GAE（PPO 用）是 Advantage 的推广：λ=0 时退化为单步 TD（即 AC 的 Advantage），λ=1 时退化为 MC（即 G_t）。PPO 取 λ=0.95 在两者间折中。', 0),
    ],
    'Week11_Step3_ActorCritic_逐行精讲.docx'
)


# =====================================================================
# File 4: week11_ppo.py —— PPO (Proximal Policy Optimization)
# =====================================================================
build_doc(
    'Step 4: PPO (Proximal Policy Optimization)',
    'week11_ppo.py — 面试重点，EMS 项目最终选用',
    [
        # ============ 一、文件概览 ============
        ('h1', '一、文件概览'),
        ('p', 'PPO（近端策略优化）在 AC 基础上加了最关键的一项——**clip（裁剪）**。AC 的问题是"一步更新可能让策略变化太大，直接崩掉"；PPO 用一个一行代码的 clamp 限制每次更新幅度，让策略"慢慢走"。它是当前工业界最主流的 on-policy 算法，也是本 EMS 项目 RL 部分的最终选型，**面试必问**。', 11),
        ('callout', '一句话记忆：PPO = AC（Actor+Critic） + ① 攒一局数据多轮更新 ② importance ratio（新旧策略概率比） ③ clip 裁剪防突变 ④ 熵奖励鼓励探索 ⑤ 梯度裁剪防爆炸。'),
        ('h2', '1.1 三种方法演进图'),
        ('flow', '''REINFORCE:  π 走完整局 → G_t → 更新一次
                 [等整局跑完才知道好坏，方差大]

Actor-Critic: π 走一步 → Critic 当场打分 → 每步更新
                 [学得快，但可能一步改太多搞崩策略]

PPO:          π 走一步 → Critic 打分 → clip(ratio) → 每步更新
                 [和 AC 一样快，但加了"保险丝"不让策略突变]'''),
        ('h2', '1.2 PPO 的 3+2 项升级'),
        ('tbl', ['机制', '作用', '代码位置'], [
            ['批量收集 + 多轮更新', '攒一局数据，反复用 epochs 次，提高样本效率', '222-268'],
            ['importance ratio', '对比新旧策略下同一动作的概率变化，为多轮复用校正偏差', '242'],
            ['clip 裁剪', '把 ratio 夹到 [0.8,1.2]，策略只能"微调"不能"突变"', '247'],
            ['熵奖励', '鼓励策略保持随机，避免过早收敛到确定性', '252'],
            ['梯度裁剪', '限制梯度范数 ≤0.5，防梯度爆炸', '259, 267'],
        ]),
        ('h2', '1.3 本文件的结构'),
        ('tbl', ['部分', '行号', '内容'], [
            ['import 与配置', '20-34', '库导入 + 中文字体'],
            ['EMSEnv（稍宽松版）', '40-78', '电池容量更大、负载更缓、奖励加了跟踪项'],
            ['Actor 网络', '82-114', '同 AC，但 get_action 额外返回 log_prob'],
            ['Critic 网络', '117-129', '同 AC'],
            ['ppo() 算法', '133-277', 'GAE + importance ratio + clip + 多轮更新'],
            ['test_policy()', '281-301', '评估'],
            ['plot_results()', '305-330', '画训练曲线'],
            ['main()', '334-376', 'CLI 入口'],
        ]),

        # ============ 二、环境 ============
        ('brk',),
        ('h1', '二、EMSEnv（稍宽松版）环境（40-78 行）'),
        ('p', '★ 注意：PPO 的环境**不是**直接复用 AC 的，而是做了三处"放宽"，让训练更容易成功。为什么？因为要演示"PPO 能学到稳定策略"这个卖点。注释（37-39 行）说明了背景：REINFORCE 和 AC 在这个（原版）环境上只撑了 2-3 步。', 11),
        ('x', 41, 'def __init__(self, hard=False):', '构造函数带 hard 参数：hard=False（默认宽松版），hard=True 时退回原版参数。'),
        ('x', 42, 'self.soc_min = 0.1', '① SOC 下限从 0.2 放宽到 0.1（给电池更宽的缓冲，少触发越界终止）。'),
        ('x', 43, 'self.soc_max = 0.95', 'SOC 上限从 0.9 放宽到 0.95。'),
        ('x', 47, 'self.battery_capacity = 100 if not hard else 50', '★ ② 电池容量从 50 翻倍到 100 kWh。容量越大，同样的功率差引起的 SOC 变化越慢——"惯性"更大，智能体不容易一步就把电量弄出界，训练友好很多。'),
        ('x', 50, 'def reset(self):', '重置。'),
        ('x', 51, 'self.soc = 0.5', '初始 SOC 改 0.5（目标值也从 0.6 改成 0.5，见 67 行）。'),
        ('x', 54, 'self.max_steps = 300', '③ 最大步数从 200 放宽到 300（局更长，学得信息更多）。'),
        ('x', 58, 'p_fc = float(np.clip(action, 0, 1)) * 30.0', '动作反归一化，一致。'),
        ('x', 59, 'self.p_load = 0.35 + 0.3 * (0.5 + 0.5 * np.sin(self.steps * 0.05))', '★ 负载波形调缓：振幅从 0.4 降到 0.3、频率从 0.1 降到 0.05。负载变化更平缓 → 任务更容易。范围 0.35~0.65 kW。'),
        ('x', 60, 'soc_change = (p_fc - self.p_load) / self.battery_capacity', 'SOC 变化量（容量改用 100）。'),
        ('x', 61, 'self.soc = np.clip(self.soc + soc_change, self.soc_min, self.soc_max)', 'SOC 更新 + 夹取。'),
        ('x', 63, '# 奖励设计：鼓励 P_fc 跟踪 P_load', '★ 奖励设计改动说明：加了"跟踪"项。'),
        ('x', 65, 'fuel_cost = -0.01 * p_fc', '燃料成本（同前）。'),
        ('x', 66, 'tracking_bonus = -0.5 * (p_fc - self.p_load) ** 2 / 900  # 鼓励跟踪负载', '★ 新增：跟踪惩罚（实为"负奖励"）。(p_fc-p_load)² 惩罚功率偏差，除以 900 是把量级压到合理范围（(30)²=900 是最大偏差的平方，归一化用）。这行让环境把"跟随负载"作为显式目标——P_fc 接近 P_load 时该项≈0。'),
        ('x', 67, 'soc_penalty = -1.0 * (self.soc - 0.5) ** 2', '★ SOC 惩罚系数从 0.5 加大到 1.0，目标点从 0.6 改到 0.5（对齐初始 SOC）。'),
        ('x', 69, 'reward = fuel_cost + tracking_bonus + soc_penalty', '总奖励 = 三项。注意：没有越界硬惩罚了（用更强的 SOC 偏离惩罚替代，边界靠终止兜底）。'),
        ('x', 72, 'done = (self.steps >= self.max_steps or', '终止条件。'),
        ('x', 73, 'self.soc <= self.soc_min or self.soc >= self.soc_max)', '仍保留 SOC 越界终止。'),
        ('x', 78, 'return np.array([self.soc, self.p_load], dtype=np.float32)', '返回状态。'),
        ('callout', '环境改动的教学意义：RL 环境设计本身就是调参的重要维度。PPO 选择"放宽物理边界 + 显式跟踪奖励"，让策略梯度更容易收敛——这也是真实项目里"奖励塑形（reward shaping）"和"任务简化"思想的体现。'),

        # ============ 三、Actor / Critic ============
        ('brk',),
        ('h1', '三、Actor 与 Critic 网络（82-129 行）'),
        ('h2', '3.1 Actor（82-114 行）'),
        ('x', 82, 'class Actor(nn.Module):', 'Actor 网络，结构与 AC 相同。'),
        ('x', 83, 'def __init__(self, state_dim=2, hidden=64, action_dim=1):', '构造函数。'),
        ('x', 85, 'self.fc1 = nn.Linear(state_dim, hidden)', '输入层 2→64。'),
        ('x', 86, 'self.fc2 = nn.Linear(hidden, hidden)', '隐藏层 64→64。'),
        ('x', 87, 'self.mean_head = nn.Linear(hidden, action_dim)', '均值头 64→1。'),
        ('x', 88, 'self.log_std = nn.Parameter(torch.zeros(action_dim))', '对数标准差（可训练）。'),
        ('x', 90, 'def forward(self, x):', '前向：输出 [μ,σ]。'),
        ('x', 91, 'x = torch.relu(self.fc1(x))', '激活。'),
        ('x', 92, 'x = torch.relu(self.fc2(x))', '激活。'),
        ('x', 93, 'mean = torch.tanh(self.mean_head(x))', 'tanh。'),
        ('x', 94, 'mean = (mean + 1) / 2  # [0, 1]', '映射到 [0,1]。'),
        ('x', 95, 'std = torch.exp(self.log_std.clamp(-5, 2))', '标准差。'),
        ('x', 96, 'return mean, std', '返回分布参数。'),
        ('x', 98, 'def get_action(self, state):', '★ 选动作。注意：相比 AC，**额外返回了 log_prob**——因为 PPO 要存"旧策略的 log_prob"，供后面算 importance ratio。'),
        ('x', 100, 'mean, std = self.forward(s)', '前向。'),
        ('x', 102, 'a = m.sample()', '采样动作。'),
        ('x', 104, 'log_prob = m.log_prob(a)', '★ 记录该动作在"当前（旧）策略"下的对数概率——这是 PPO 独有的需求。'),
        ('x', 105, 'a = a.clamp(0, 1)', '夹到 [0,1]。'),
        ('x', 106, 'return a.item(), log_prob.item()', '★ 返回 (动作值, 旧 log_prob) 二元组。log_prob 将被存入 log_probs_old 列表。'),
        ('x', 108, 'def evaluate(self, state, action):', '★ 训练用：返回 log_prob **和熵**（带梯度）。比 AC 多返回熵——PPO 用熵做奖励。'),
        ('x', 109, '"""返回 log_prob 和熵（带梯度）"""', 'docstring。'),
        ('x', 112, 'log_prob = m.log_prob(action)', '对数概率。'),
        ('x', 113, 'entropy = m.entropy()', '★ 计算高斯分布的熵。熵衡量分布"有多乱"：熵高=分布均匀（探索强），熵低=分布集中（利用强）。'),
        ('x', 114, 'return log_prob, entropy', '返回二元组。'),

        ('h2', '3.2 Critic（117-129 行）'),
        ('x', 117, 'class Critic(nn.Module):', 'Critic 网络，与 AC 完全相同。'),
        ('x', 118, 'def __init__(self, state_dim=2, hidden=64):', '构造函数。'),
        ('x', 120, 'self.net = nn.Sequential(', '顺序网络。'),
        ('x', 121, 'nn.Linear(state_dim, hidden),', '输入层。'),
        ('x', 122, 'nn.ReLU(),', '激活。'),
        ('x', 123, 'nn.Linear(hidden, hidden),', '隐藏层。'),
        ('x', 124, 'nn.ReLU(),', '激活。'),
        ('x', 125, 'nn.Linear(hidden, 1)', '输出标量 V(s)。'),
        ('x', 128, 'def forward(self, x):', '前向。'),
        ('x', 129, 'return self.net(x)', '返回 V(s)。'),

        # ============ 四、ppo 算法 ============
        ('brk',),
        ('h1', '四、ppo() 算法逐行精讲（133-277 行）'),
        ('h2', '4.1 参数与初始化（133-165 行）'),
        ('x', 133, 'def ppo(episodes=500, lr=0.0003, clip_eps=0.2, epochs=10, batch_size=64):', '★ 主函数签名，5 个参数：episodes（局数）、lr=0.0003（★ 学习率比 AC 的 0.001 还小——PPO 更新更保守，宁慢勿崩）、clip_eps=0.2（裁剪范围 ±0.2）、epochs=10（★ 同一批数据重用 10 次）、batch_size=64（mini-batch 大小）。'),
        ('x', 147, 'env = EMSEnv()', '创建宽松版环境。'),
        ('x', 148, 'actor = Actor()', 'Actor 网络。'),
        ('x', 149, 'critic = Critic()', 'Critic 网络。'),
        ('x', 150, 'actor_opt = optim.Adam(actor.parameters(), lr=lr)', 'Actor 优化器。'),
        ('x', 151, 'critic_opt = optim.Adam(critic.parameters(), lr=lr)', '★ Critic 优化器：和 Actor 相同学习率（AC 里是 lr×2，这里不设差异，说明学习率策略不是铁律，可灵活调）。'),
        ('x', 156, 'print(f"  clip_eps = {clip_eps}")', '打印裁剪参数。'),
        ('x', 158, 'print(f"  batch_size = {batch_size}")', '打印 batch 大小。'),
        ('x', 163, 'gamma = 0.99', '折扣因子。'),
        ('x', 164, 'gae_lambda = 0.95  # GAE 平滑系数', '★ GAE 参数 λ=0.95。AC 用单步 TD（λ=0），PPO 用 GAE（λ=0.95）在"单步"和"整局 MC"之间折中——既压低方差又少丢远期信息。'),
        ('x', 165, 'entropy_coef = 0.01  # 熵奖励系数', '★ 熵奖励系数 0.01。控制"探索奖励"的强度。'),

        ('h2', '4.2 第 1 步：跑一局收集数据（167-187 行）'),
        ('x', 167, 'for ep in range(1, episodes + 1):', '训练局循环。'),
        ('x', 169, 's = env.reset()', '重置环境。'),
        ('x', 170, 'states, actions, rewards, dones, log_probs_old = [], [], [], [], []', '★ 五个列表收集一局的全部数据。相比 AC 每步立即更新，PPO 是**攒一局再批量更新**（on-policy 的 batch 训练）。log_probs_old 是 PPO 独有——旧策略下的 log_prob。'),
        ('x', 172, 'while True:', '内层循环。'),
        ('x', 173, 'a, lp = actor.get_action(s)', '★ Actor 选动作 + 旧 log_prob（这是本文件 get_action 多返回的）。'),
        ('x', 174, 'sp, r, done, _ = env.step(a)', '执行动作。'),
        ('x', 176, 'states.append(s)', '收集状态。'),
        ('x', 177, 'actions.append(a)', '收集动作。'),
        ('x', 178, 'rewards.append(r)', '收集奖励。'),
        ('x', 179, 'dones.append(done)', '收集终止标志。'),
        ('x', 180, 'log_probs_old.append(lp)', '★ 收集旧 log_prob——后面 ratio 的"分母"。'),
        ('x', 182, 's = sp', '状态推进。'),
        ('x', 183, 'if done:', '若终止。'),
        ('x', 184, 'break', '退出本局。'),
        ('x', 186, 'episode_rewards.append(sum(rewards))', '记录本局总奖励。'),
        ('x', 187, 'episode_lengths.append(len(rewards))', '记录本局步数。'),

        ('h2', '4.3 第 2 步：算 GAE（189-217 行）—— PPO 相比 AC 的升级点'),
        ('p', 'GAE（Generalized Advantage Estimation）用递推把多步 TD error 加权融合，比 AC 的单步 Advantage 更平滑。', 12, True),
        ('x', 191, 'states_t = torch.FloatTensor(np.array(states))', '状态列表 → 张量，shape [n, 2]（n=本局步数）。'),
        ('x', 192, 'actions_t = torch.FloatTensor(actions).unsqueeze(1)', '动作 → 张量 [n] → unsqueeze 成 [n,1]。'),
        ('x', 193, 'rewards_t = torch.FloatTensor(rewards)', '奖励 → 张量 [n]。'),
        ('x', 194, 'dones_t = torch.FloatTensor(dones)', '终止标志 → 张量 [n]。'),
        ('x', 195, 'old_log_probs_t = torch.FloatTensor(log_probs_old).unsqueeze(1)', '旧 log_prob → 张量 [n,1]。'),
        ('x', 197, 'with torch.no_grad():', '★ GAE 计算全程不更新 Critic，只是"估值"，故关闭梯度。'),
        ('x', 198, 'values = critic(states_t).squeeze()', '★ 一次前向得到所有状态的 V(s)。squeeze() 把 [n,1] 压成 [n]，方便逐元素操作。'),
        ('x', 200, 'advantages = []', '存 GAE 结果。'),
        ('x', 201, 'gae = 0', '递推累加器。'),
        ('x', 202, 'next_value = 0', '下一状态价值（初始化 0）。'),
        ('x', 203, 'for t in reversed(range(len(rewards))):', '★ 从最后一步反向遍历（和 REINFORCE 算 G_t 同理，反向递推）。'),
        ('x', 204, 'if t == len(rewards) - 1:', '若 t 是最后一步。'),
        ('x', 205, 'next_value = 0  # 终点后的 V=0', '终点之后没有未来回报，V=0（等价于 AC 的 (not done) 处理）。'),
        ('x', 206, 'else:', '否则。'),
        ('x', 207, 'next_value = values[t + 1].item()', '取下一步的 Critic 估值（.item() 转 Python 数）。'),
        ('x', 208, 'delta = rewards[t] + gamma * next_value * (1 - dones[t]) - values[t].item()', '★ 单步 TD error：δ = r + γ·V(s\')·(1-done) - V(s)。和 AC 的 Advantage 一样。'),
        ('x', 209, 'gae = delta + gamma * gae_lambda * (1 - dones[t]) * gae', '★★ GAE 核心递推：GAE_t = δ_t + γλ·(1-done)·GAE_{t+1}。当前步的 GAE 包含当前 δ 加"折扣后的后续 GAE"。λ=0.95 意味着约 1/(1-0.95)≈20 步后的信息权重衰减到约 e⁻¹。这样 GAE 既用上了多步信息（比单步准），又不会像 MC 那样方差爆表。'),
        ('x', 210, 'advantages.insert(0, gae)', '逆序插入，让索引对齐。'),
        ('x', 212, 'advantages_t = torch.FloatTensor(advantages)', 'GAE → 张量。'),
        ('x', 213, 'returns_t = advantages_t + values  # returns = advantage + V(s)', '★ returns = A + V(s)。为什么？Critic 的训练目标是"状态价值"V(s)，其 TD target 正是"回报"= Advantage + V(s)（因为 A = G - V(s)，所以 G = A + V(s)）。这行直接构造出 Critic 的训练标签。'),
        ('x', 216, 'if len(advantages_t) > 1:', '标准化前检查。'),
        ('x', 217, 'advantages_t = (advantages_t - advantages_t.mean()) / (advantages_t.std() + 1e-8)', '★ 标准化 Advantage（同 REINFORCE 标准化 G 的理由，稳定训练）。'),

        ('h2', '4.4 第 3 步：PPO 核心——多轮 clip 更新（219-268 行）'),
        ('p', '这是 PPO 的招牌部分，面试高频。核心是"同一批数据多轮复用 + importance ratio 校正 + clip 防突变"。', 12, True),
        ('x', 220, 'n = len(states)', '本局步数（数据条数）。'),
        ('x', 222, 'for _ in range(epochs):', '★ 多轮更新循环：同一局数据重复用 epochs=10 次。AC 是"来一条学一条、用完即弃"；PPO 批量攒好反复学，样本效率高。这就是 on-policy 也能复用的关键——靠 importance ratio 校正。'),
        ('x', 224, 'indices = np.random.permutation(n)', '★ 打乱索引（shuffle），打破数据顺序相关性，让 mini-batch 更随机。'),
        ('x', 226, 'for start in range(0, n, batch_size):', '把 n 条数据切成长度 batch_size 的 mini-batch 依次处理。'),
        ('x', 227, 'idx = indices[start:start + batch_size]', '取本批索引。'),
        ('x', 229, 'batch_s = states_t[idx]', '本批状态。'),
        ('x', 230, 'batch_a = actions_t[idx]', '本批动作。'),
        ('x', 231, 'batch_adv = advantages_t[idx]', '本批 Advantage。'),
        ('x', 232, 'batch_ret = returns_t[idx]', '本批 Critic 目标。'),
        ('x', 233, 'batch_old_lp = old_log_probs_t[idx]', '本批旧 log_prob。'),
        ('x', 236, 'log_probs_new, entropy = actor.evaluate(batch_s, batch_a)', '★ 用**当前**策略（已经更新了几轮）重新算 log_prob。注意：这与 batch_old_lp（训练开始时存的）可能不同了——差值正是"策略变化了多少"的度量。'),
        ('x', 242, 'ratio = torch.exp(log_probs_new - batch_old_lp)', '★★ importance ratio = π_new(a|s)/π_old(a|s) = e^(logπ_new - logπ_old)。含义：**同一个动作，在新策略下比旧策略下"可能"了多少倍**。ratio>1 概率增大了，ratio<1 概率减小了，ratio=1 没变。例：ratio=1.5 → 概率增大 50%；ratio=0.5 → 概率减半。'),
        ('x', 244, '# 3c. PPO clip 核心公式', '注释：进入 clip。'),
        ('x', 245, '# L_clip = min(ratio × A, clip(ratio, 1-ε, 1+ε) × A)', '★ 公式原型：对两个候选目标取 min。'),
        ('x', 246, 'surr1 = ratio * batch_adv', '★ 未裁剪目标（surrogate 1）：和普通策略梯度一样 ratio×A。'),
        ('x', 247, 'surr2 = torch.clamp(ratio, 1 - clip_eps, 1 + clip_eps) * batch_adv', '★★ 裁剪目标（surrogate 2）：把 ratio 夹到 [1-0.2, 1+0.2]=[0.8,1.2]，再乘 A。这行就是 PPO 的全部创新——**一行 clamp 阻止策略突变**。若 ratio=2.0（概率翻倍），被砍到 1.2；ratio=0.3 被抬到 0.8。'),
        ('x', 248, 'actor_loss = -torch.min(surr1, surr2).mean()', '★★★ 最终目标：-min(surr1, surr2) 取平均、取负（PyTorch 做梯度下降）。展开解释：当 ratio 在安全区间 [0.8,1.2] 内时，surr1=surr2，正常更新；当 ratio 超界时，min 选中被裁剪的 surr2，梯度不再随 ratio 增长（被打平为 0），策略被"按住"——这就是"近端（Proximal）"的含义：**不让策略走太远**。'),
        ('x', 249, '#   ↑ 负号是因为我们要最大化，但 PyTorch 是做梯度下降', '注释：取负号的原因。'),
        ('x', 252, 'entropy_loss = -entropy_coef * entropy.mean()', '★ 熵奖励：-0.01 × 平均熵。最大化熵 = 保持随机性/探索。策略若过早变确定（std→0），熵→0，这个奖励会阻止它——相当于一种"探索保险"，防过早陷入局部最优。'),
        ('x', 254, 'actor_total = actor_loss + entropy_loss', '★ 总 Actor 损失 = clip 损失 + 熵奖励。'),
        ('x', 256, 'actor_opt.zero_grad()', '清梯度。'),
        ('x', 257, 'actor_total.backward()', '反向传播。'),
        ('x', 258, '# 梯度裁剪：防止梯度爆炸', '注释。'),
        ('x', 259, 'torch.nn.utils.clip_grad_norm_(actor.parameters(), 0.5)', '★ 梯度裁剪：把所有参数梯度的范数限制在 ≤0.5。若梯度范数超过 0.5 就按比例缩放。又一个"安全保险"——PPO 有多重保险机制，这是它训练稳定的原因。'),
        ('x', 260, 'actor_opt.step()', '更新 Actor。'),
        ('x', 263, 'V_pred = critic(batch_s).squeeze()', '★ Critic 对这批状态的预测 V(s)，squeeze 掉 [n,1]→[n]。'),
        ('x', 264, 'critic_loss = nn.MSELoss()(V_pred, batch_ret)', '★ Critic loss = MSE(V(s), returns)。returns 是在 4.3 步算好的"优势+价值"标签（等价于 TD 目标）。'),
        ('x', 265, 'critic_opt.zero_grad()', '清梯度。'),
        ('x', 266, 'critic_loss.backward()', '反向传播。'),
        ('x', 267, 'torch.nn.utils.clip_grad_norm_(critic.parameters(), 0.5)', 'Critic 也做梯度裁剪。'),
        ('x', 268, 'critic_opt.step()', '更新 Critic。'),
        ('x', 270, 'if ep % 50 == 0:', '每 50 局打印。'),
        ('x', 273, 'avg_r = np.mean(episode_rewards[-50:])', '近 50 局平均奖励。'),
        ('x', 275, 'f"log_std={actor.log_std.item():.3f}")', '打印探索参数。'),
        ('x', 277, 'return actor, episode_rewards, episode_lengths', '返回结果。'),

        # ============ 五、测试与画图 ============
        ('brk',),
        ('h1', '五、测试与画图（281-330 行）'),
        ('x', 281, 'def test_policy(actor, episodes=10):', '评估函数。'),
        ('x', 289, 'for t in range(300):', '每局最多 300 步（对齐 max_steps=300）。'),
        ('x', 290, 'a, _ = actor.get_action(s)', '取动作（丢弃 log_prob）。'),
        ('x', 297, 'if ep < 3:', '前 3 局打印详细。'),
        ('x', 298, 'print(f"  第{ep+1}局: 总奖励={total_r:+.3f} ({t+1}步)")', '★ 多打印了步数——观察策略是否"活"得久（学到稳定策略 → 走满 300 步而不出界）。'),
        ('x', 300, 'print(f"  平均总奖励: {np.mean(total_rewards):+.3f}")', '平均奖励。'),
        ('x', 305, 'def plot_results(rewards, label=\'PPO\', output_dir: str | Path | None = None):', '画图函数（红色系）。'),
        ('x', 311, 'smoothed = np.convolve(rewards, np.ones(window)/window, mode=\'valid\')', '滑动平均。'),
        ('x', 327, "path = ensure_results_dir(output_dir) / 'week11_ppo_training.png'", '输出 week11_ppo_training.png。'),

        # ============ 六、主程序 ============
        ('brk',),
        ('h1', '六、main() 主程序（334-376 行）'),
        ('x', 336, "parser = argparse.ArgumentParser(description='Week 11 PPO demo on a simplified continuous EMS environment')", '参数解析器。'),
        ('x', 337, "parser.add_argument('--episodes', type=int, default=500, help='Training episodes')", '局数。'),
        ('x', 338, "parser.add_argument('--lr', type=float, default=0.0003, help='Learning rate')", '学习率（默认 0.0003）。'),
        ('x', 339, "parser.add_argument('--seed', type=int, default=42, help='Random seed')", '种子。'),
        ('x', 340, "parser.add_argument('--output-dir', type=Path, default=None, help='Directory for generated figures')", '输出目录。'),
        ('x', 351, 'actor, rewards, lengths = ppo(episodes=args.episodes, lr=args.lr)', '训练。'),
        ('x', 353, 'test_policy(actor)', '评估。'),
        ('x', 354, 'plot_results(rewards, output_dir=args.output_dir)', '画图。'),
        ('x', 360, 'print(""""', '打印三种方法对比总结。'),
        ('x', 361, 'REINFORCE:    π 直接走完 → Σr → 更新', 'REINFORCE 路径。'),
        ('x', 364, 'Actor-Critic: π 走一步 → Critic 当场评价 → 更新', 'AC 路径。'),
        ('x', 367, 'PPO:          π 走一步 → Critic 评价 → clip(ratio) → 更新', 'PPO 路径。'),
        ('x', 370, 'print("  PPO 是 EMS 项目最终选用的算法。")', '★ 项目结论。'),
        ('x', 371, 'print("  原因：连续动作 + 训练稳定 + 实现复杂度适中")', '★ 选型三大理由：连续动作适配、clip 保证稳定、实现比 SAC（5 个网络）简单。'),
        ('x', 375, "if __name__ == '__main__':", '入口。'),
        ('x', 376, 'main()', '启动。'),

        # ============ 七、机制专题 ============
        ('brk',),
        ('h1', '七、机制专题'),
        ('h2', '7.1 clip 的三种情形'),
        ('flow', '''设 clip_eps=0.2，Advantage A>0（好动作）：
  情形一  ratio=1.1（在[0.8,1.2]内）：
         surr1=1.1A, surr2=1.1A → min 取 1.1A → 正常增大该动作概率 ✓
  情形二  ratio=1.5（超上界）：
         surr1=1.5A, surr2=1.2A → min 取 1.2A → 概率增长被"锁死"在 1.2 倍
          → 即使再更新，奖励也不会进一步上升（梯度≈0）→ 防突变 ✓
  情形三  ratio=0.9，但 A<0（坏动作）：
         surr1=0.9A, surr2=0.9A → 都在区间内，正常减小概率 ✓'''),
        ('callout', '直观类比：clip 像"安全带"或"防呆设计"——策略可以进步，但每一步只准进步一点。这避免了"一步跨太大掉下悬崖（策略崩溃）"的悲剧。'),

        ('h2', '7.2 importance ratio 为什么能"复用旧数据"'),
        ('p', 'on-policy 算法本应"数据用完即弃"。但 PPO 用重要性采样把"旧策略收集的数据"重新加权到"新策略"下使用：E_old[f(x)] ≈ E_new[f(x)·π_new/π_old]。ratio 就是这个权重。多轮更新时，虽然策略在变，但每次都用最新 ratio 校正，数据就不会"过时失真"——从而可以把同一批数据学 10 次。', 11),
        ('h2', '7.3 PPO vs AC vs REINFORCE 汇总'),
        ('tbl', ['维度', 'REINFORCE', 'Actor-Critic', 'PPO'], [
            ['策略网络', 'PolicyNet', 'Actor', 'Actor'],
            ['价值网络', '无', 'Critic V(s)', 'Critic V(s)'],
            ['评价信号', 'G_t（MC 整局回报）', 'A=r+γV\'-V（单步 TD）', 'GAE（多步加权，λ=0.95）'],
            ['更新时机', '整局结束一次', '每步一次', '整局攒批，多轮×mini-batch'],
            ['安全机制', '无', '无', 'clip + 梯度裁剪 + 熵奖励'],
            ['样本利用率', '低', '中', '高（10 次复用）'],
            ['方差', '很高', '中', '低'],
            ['训练稳定性', '差（易崩）', '中（可能一步突变）', '好（多重保险）'],
        ]),

        # ============ 八、训练走查 ============
        ('brk',),
        ('h1', '八、一局完整训练走查'),
        ('flow', '''【第 ep 局】
  ① 跑局收集：states/actions/rewards/dones/log_probs_old 各一列
       while True: a, lp = actor.get_action(s); 收集; s=sp; until done
  ② 算 GAE（反向递推）：
       values = critic(所有 states)
       for t in reversed: delta = r+γV'-V; gae = δ+γλ·gae
       returns = advantage + V(s)      ← Critic 的训练标签
       标准化 advantage
  ③ 多轮 clip 更新（epochs=10 × mini-batch=64）：
       ratio = exp(logπ_new - logπ_old)
       L = -min(ratio·A, clip(ratio,0.8,1.2)·A) + 0.01·H
       反向传播 + 梯度裁剪(≤0.5) → 更新 Actor
       MSE(V(s), returns) → 更新 Critic
【记录】episode_rewards / episode_lengths'''),

        # ============ 九、面试要点 ============
        ('brk',),
        ('h1', '九、面试要点（PPO 必背）'),
        ('bullet', '「PPO 相比 AC 的核心改进？」→ 用 importance ratio 衡量策略变化，再用 clip 把 ratio 限制在 [1-ε, 1+ε]，防止一步更新过多导致策略崩溃；同时用 GAE 估计 Advantage、熵奖励维持探索、梯度裁剪防爆炸。', 0),
        ('bullet', '「PPO 的核心公式？」→ L_clip = -min(ratio·A, clip(ratio, 1-ε, 1+ε)·A)，ratio = π_new/π_old。', 0),
        ('bullet', '「为什么 clip 能防崩溃？」→ ratio 超出 [0.8,1.2] 时 min 取裁剪项，梯度被"打平"，策略不再激进更新——训练稳定。', 0),
        ('bullet', '「PPO 为什么能多次使用一批数据（on-policy 却高效）？」→ importance sampling 用 ratio 校正新旧策略差异，同一批数据可安全复用 epochs 次。', 0),
        ('bullet', '「为什么 EMS 选 PPO？」→ 连续动作适配（输出高斯分布采样）+ 训练稳定（clip+梯度裁剪+熵）+ 实现复杂度适中（对比 SAC 需 5 个网络）+ 工业界主流、面试常考。', 0),
        ('bullet', '「GAE 是什么？」→ Generalized Advantage Estimation，用 λ 平滑多步 TD error，λ=0 退化单步 TD，λ=1 退化 MC，PPO 取 0.95 折中。', 0),

        ('h1', '十、常见疑问 FAQ'),
        ('bullet', 'Q：PPO 是 on-policy 吗？怎么又复用数据？→ 是 on-policy，但用 importance sampling 校正后可以把"当前策略采的数据"多学几次；数据本质还是旧策略的，所以每局收集一次、更新完就弃。', 0),
        ('bullet', 'Q：为什么 lr 只有 0.0003？→ 更新更保守，配合 clip 双重保险，宁慢勿崩。lr 太大策略梯度容易发散。', 0),
        ('bullet', 'Q：熵奖励会不会让策略永远不收敛？→ 系数只有 0.01，很弱；它只防止"过早确定"，不影响最终收敛到近确定性策略。', 0),
        ('bullet', 'Q：环境为什么改成宽松版？→ 原版环境对 REINFORCE/AC 太苛刻（几步就出界）；放宽边界+加跟踪奖励是为了展示 PPO 的稳定学习能力。真实项目中"先简化任务再逐步加难度"也是常见做法。', 0),
        ('bullet', 'Q：PPO 还有什么超参可以调？→ batch_size、epochs、clip_eps、gae_lambda、entropy_coef、网络宽度等。调参原则：clip_eps 太大→失去保护，太小→学得慢；gae_lambda 越大越接近 MC。', 0),
    ],
    'Week11_Step4_PPO_逐行精讲.docx'
)

# =====================================================================
# File 0: week11_common.py —— Week11 公共底座
# =====================================================================
build_doc(
    'Step 0: week11_common 公共底座',
    'week11_common.py — 6 个脚本共用的路径/编码/种子/字体基础设施',
    [
        # ============ 一、文件概览 ============
        ('h1', '一、文件概览'),
        ('p', '这个文件是 Week 11 所有 RL 脚本的"公共底座"——它本身不实现任何算法，而是解决 6 个脚本都会遇到的 4 类环境问题：**① 输出路径不统一 ② Windows 控制台中文/符号乱码崩溃 ③ 实验不可复现 ④ matplotlib 中文显示成方框**。把公共逻辑抽到这里，每个算法脚本只专注算法本身，这是 Week 11 工程化重构（2026-07-28）的核心成果。', 11),
        ('callout', '一句话记忆：week11_common = 2 个路径常量（PROJECT_ROOT / RESULTS_DIR）+ 4 个函数（configure_console / ensure_results_dir / set_seed / configure_matplotlib）+ 1 次模块级调用（configure_console()）。'),
        ('h2', '1.1 被哪些脚本引用'),
        ('tbl', ['脚本', 'import 行', '用了哪些函数'], [
            ['week11_continuous_env.py', '23', 'set_seed'],
            ['week11_reinforce.py', '27', 'configure_matplotlib, ensure_results_dir, set_seed'],
            ['week11_actor_critic.py', '32', 'configure_matplotlib, ensure_results_dir, set_seed'],
            ['week11_ppo.py', '32', 'configure_matplotlib, ensure_results_dir, set_seed'],
            ['week11_compare.py', '27', 'configure_matplotlib, ensure_results_dir, set_seed'],
            ['compare_large_grid.py', '27', 'configure_matplotlib, ensure_results_dir, set_seed'],
        ]),
        ('p', '可以看到：set_seed 是 6 个脚本都用的（保证实验可复现）；configure_matplotlib 和 ensure_results_dir 被除 continuous_env 外的 5 个脚本使用（它们都要画图）；configure_console 则不直接 import，而是靠本文件模块级自动执行（见 6 章）。', 11, True),

        ('h2', '1.2 解决了哪 4 类问题'),
        ('tbl', ['问题', '后果', '解决函数'], [
            ['输出目录硬编码 F:\\CLAUDE\\research\\...', '换机器/换目录就跑崩', 'PROJECT_ROOT + ensure_results_dir'],
            ['Windows 控制台打印 μ/√/中文报错', '程序中途崩溃，UnicodeEncodeError', 'configure_console'],
            ['每次运行结果不同', '实验无法复现、无法比较', 'set_seed'],
            ['matplotlib 中文变方框 □□', '图表不可读', 'configure_matplotlib'],
        ]),

        # ============ 二、结构地图 ============
        ('brk',),
        ('h1', '二、代码结构地图'),
        ('tbl', ['部分', '行号', '内容'], [
            ['文件头 + import', '1-16', '模块说明 + 容错导入'],
            ['路径常量', '19-20', 'PROJECT_ROOT / RESULTS_DIR'],
            ['configure_console()', '23-28', '控制台 UTF-8 编码'],
            ['ensure_results_dir()', '31-36', '确保输出目录存在'],
            ['set_seed()', '39-47', '设随机种子'],
            ['configure_matplotlib()', '50-68', '中文字体配置'],
            ['模块级调用', '71', 'configure_console() 自动执行'],
        ]),

        # ============ 三、import 精讲 ============
        ('brk',),
        ('h1', '三、文件头与 import 逐行精讲（1-16 行）'),
        ('x', 1, '#!/usr/bin/env python3', 'Shebang。Unix/Linux 下允许 ./week11_common.py 直接执行（尽管本文件主要被 import，不是独立入口）。'),
        ('x', 2, '# -*- coding: utf-8 -*-', '声明 UTF-8 编码，兼容旧 Python 的中文注释。'),
        ('x', 3, '"""Shared helpers for Week 11 RL scripts."""', '模块 docstring：点明定位——"Week 11 RL 脚本的共享辅助工具"。注意用英文写，因为这是给所有脚本共用的基础设施，保持中性。'),
        ('x', 5, 'from __future__ import annotations', '★ 延迟求值类型注解。作用：让 `str | Path | None` 这类 Python 3.10+ 的联合类型语法在 Python 3.9 及以下也能用于**注解**（不会在 import 时求值报错）。为什么需要？代码里大量用了 `output_dir: str | Path | None`、`-> Path`、`-> None`。不加这行，3.10 以下会因 `str | Path` 在函数定义时求值而抛 TypeError。'),
        ('x', 7, 'import random', '导入标准库 random——set_seed 要设置它的种子。'),
        ('x', 8, 'import sys', '导入标准库 sys——configure_console 要访问 sys.stdout / sys.stderr。'),
        ('x', 9, 'from pathlib import Path', '导入 pathlib.Path——路径常量构建和 ensure_results_dir 的类型标注都用它。'),
        ('x', 11, 'import numpy as np', '导入 NumPy——set_seed 设置 np.random 的种子。'),
        ('x', 13, 'try:', '开始容错导入块。'),
        ('x', 14, '    import torch', '尝试导入 PyTorch。'),
        ('x', 15, 'except ImportError:  # pragma: no cover - only for environments without torch', '★ 捕获导入失败。# pragma: no cover 是给 coverage 工具看的：这段异常路径不参与覆盖率统计。原因：如果机器没装 torch，set_seed 里关于 torch 的部分要能优雅跳过，而不是整个模块崩掉。'),
        ('x', 16, '    torch = None', '★ 关键兜底：torch 没装上时赋为 None。这样后续 `if torch is not None:` 就能安全跳过 torch 相关代码。注意：Week 11 算法脚本本身必须要有 torch，所以这个兜底主要保护"仅想用 common 画图/设种子的场景"。'),

        # ============ 四、路径常量 ============
        ('brk',),
        ('h1', '四、路径常量（19-20 行）—— 消灭硬编码'),
        ('x', 19, 'PROJECT_ROOT = Path(__file__).resolve().parents[1]', '★★ 计算项目根目录。逐层拆解：① __file__ = 本文件的绝对路径（如 scripts/week11_common.py）；② .resolve() 解析所有符号链接/相对路径为绝对路径；③ .parents[1] 取父目录的父目录 = scripts/ 的上一级 = 项目根目录。为什么不用硬编码 F:\\CLAUDE\\research\\ems-platform？因为换机器、换目录、被拷贝到别处，这行代码永远自适应。这是本次重构消除硬编码路径的关键一行。'),
        ('x', 20, 'RESULTS_DIR = PROJECT_ROOT / "results"', '★ 在项目根下拼接 results 子目录。所有输出图片/CSV 都默认进这里。用 pathlib 的 / 运算符拼接，比 os.path.join 更简洁且跨平台。'),
        ('callout', '对比重构前的痛点：各脚本硬编码 `F:\\CLAUDE\\research\\ems-platform\\results`。一旦仓库被 clone 到别的盘/目录，输出就会跑丢或报错。现在改成相对项目根定位，仓库搬到哪都能跑。'),

        # ============ 五、四个函数 ============
        ('brk',),
        ('h1', '五、四个函数逐行精讲'),

        ('h2', '5.1 configure_console()（23-28 行）—— 解决 Windows 控制台乱码崩溃'),
        ('x', 23, 'def configure_console() -> None:', '定义函数。返回 None。'),
        ('x', 24, '"""Avoid Windows console encoding crashes for math symbols and Chinese text."""', 'docstring：避免 Windows 控制台在打印数学符号和中文时报编码错误。'),
        ('x', 26, 'for stream in (sys.stdout, sys.stderr):', '遍历标准输出和标准错误两个流。'),
        ('x', 27, 'if hasattr(stream, "reconfigure"):', '★ 检测流对象是否有 reconfigure 方法。Python 3.7+ 的 TextIOWrapper 有这个方法，旧式流对象没有——hasattr 保证兼容性（没有就跳过）。'),
        ('x', 28, 'stream.reconfigure(encoding="utf-8", errors="replace")', '★★ 核心操作：把流编码改成 UTF-8，errors="replace" 表示遇到无法编码的字符时用替换符（?）而不是抛异常。不这么做，打印 μ、√、≈ 或中文时在 GBK 控制台会直接 UnicodeEncodeError 崩溃。'),
        ('callout', '为什么必须做？Windows 默认控制台编码是 GBK/cp936，而 Python 源码是 UTF-8。直接 print("μ=0.63") 在旧环境下会抛 UnicodeEncodeError，程序中途死亡。这一行把整个 Week11 脚本的打印稳定住了。'),

        ('h2', '5.2 ensure_results_dir()（31-36 行）—— 输出目录兜底'),
        ('x', 31, 'def ensure_results_dir(output_dir: str | Path | None = None) -> Path:', '函数签名。注意类型注解 `str | Path | None`：允许传字符串、Path 对象或不传。这是 CLI 参数（args.output_dir）和默认值（None→用 RESULTS_DIR）统一入口。'),
        ('x', 34, 'path = Path(output_dir) if output_dir is not None else RESULTS_DIR', '★ 三目表达式：传了就用它（转成 Path），没传就用 RESULTS_DIR 默认值。这样"用户指定目录"和"默认目录"一条逻辑搞定。'),
        ('x', 35, 'path.mkdir(parents=True, exist_ok=True)', '★★ 确保目录存在。mkdir(parents=True) 递归创建所有缺失的父目录；exist_ok=True 表示目录已存在也不报错。这是"幂等"操作——跑多少次都安全，目录一定在。'),
        ('x', 36, 'return path', '返回可写的输出目录 Path 对象。调用方直接拼文件名：`ensure_results_dir() / "week11_ppo_training.png"`。'),

        ('h2', '5.3 set_seed()（39-47 行）—— 实验可复现'),
        ('x', 39, 'def set_seed(seed: int) -> None:', '函数签名：接收整数种子。'),
        ('x', 40, '"""Make numpy/random/torch experiments easier to reproduce."""', 'docstring：让实验更易复现。'),
        ('x', 42, 'random.seed(seed)', '★ 设置 Python 内置 random 的种子（影响 random.random() 等，连续环境测试的随机策略用到）。'),
        ('x', 43, 'np.random.seed(seed)', '★ 设置 NumPy 随机种子（影响 np.random.* 和 np.random.permutation，PPO 的 mini-batch 打乱用到）。'),
        ('x', 44, 'if torch is not None:', '★ 容错：只有 torch 真的装上了才设置其种子（呼应第 13-16 行的 torch=None 兜底）。'),
        ('x', 45, 'torch.manual_seed(seed)', '★ 设置 PyTorch 全局种子（CPU 端）。影响 torch 张量的随机初始化（网络权重）和随机采样。'),
        ('x', 46, 'if torch.cuda.is_available():', '★ 判断是否有可用 GPU（CUDA）。'),
        ('x', 47, 'torch.cuda.manual_seed_all(seed)', '设置所有 CUDA 设备（含多卡）的种子。注意：即便只设了 CPU 种子，网络权重、Normal 采样等在 CPU 上已经可复现；GPU 端需要额外设置（且 GPU 上存在非确定性算子，完全复现需要额外配置）。'),
        ('callout', '为什么每个脚本的 main() 都要在开头调用 set_seed(args.seed)？因为不设种子，每次运行网络权重初始化和采样都不同，训练曲线五花八门——你无法判断算法差异是"方法本身"还是"运气"。设种子后，同一 seed 每次运行结果完全一致，对比实验才公平。'),

        ('h2', '5.4 configure_matplotlib()（50-68 行）—— 中文图表'),
        ('x', 50, 'def configure_matplotlib() -> None:', '函数签名。'),
        ('x', 51, '"""Configure matplotlib for Chinese labels on Windows and headless export."""', 'docstring：配置 matplotlib 以支持中文标签和无头导出。'),
        ('x', 53, 'import matplotlib.pyplot as plt', '★ 函数内 import。为什么不放顶部？因为有些脚本（如 continuous_env）不需要 matplotlib，函数内导入实现"用到才加载"，避免无谓开销。'),
        ('x', 54, 'from matplotlib import font_manager', '导入字体管理器，用于查询系统已安装字体。'),
        ('x', 56, 'installed = {font.name for font in font_manager.fontManager.ttflist}', '★ 构建"已安装字体名"集合：ttflist 是 matplotlib 扫描到的所有字体对象列表，取每个的 .name 组成集合（去重、快速查找）。'),
        ('x', 57, 'for candidate in (', '开始遍历候选字体列表。'),
        ('x', 58, '"Microsoft YaHei",', '候选① 微软雅黑（Windows 系统自带，最常见）。'),
        ('x', 59, '"SimHei",', '候选② 黑体（Windows 自带，无衬线）。'),
        ('x', 60, '"Noto Sans CJK SC",', '候选③ 思源黑体（Google 开源，跨平台）。'),
        ('x', 61, '"Source Han Sans SC",', '候选④ 思源黑体另一名称。'),
        ('x', 62, '"Arial Unicode MS",', '候选⑤ Mac 系统字体。'),
        ('x', 63, '"DejaVu Sans",', '候选⑥ 最后兜底——注意 DejaVu 并不含中文字形，是"最后选择"（至少不崩，中文可能仍变方框）。'),
        ('x', 65, 'if candidate in installed:', '★ 按优先级找第一个已安装的字体。顺序即优先级：微软雅黑最优先（本机 Windows）。'),
        ('x', 66, 'plt.rcParams["font.family"] = candidate', '★ 命中则把 matplotlib 全局字体族设为该字体。这样所有标题/坐标轴/图例的中文都能正常渲染。'),
        ('x', 67, 'break', '找到即停，不再往后找。'),
        ('x', 68, 'plt.rcParams["axes.unicode_minus"] = False', '★ 修复负号显示成方框：matplotlib 默认用 Unicode 负号 −（U+2212），但很多中文字体没有这个字形，导致轴上的负数变成 □。设为 False 后改用 ASCII 连字符 "-"，负号正常显示。'),
        ('callout', '为什么必须做？不配置字体，matplotlib 在 Windows 上画中文标题/标签会渲染成空心方框 □□，图表完全不可读。这次重构后，Week 11 所有图表的中文已复查不再显示方框（见 STATUS.md）。'),

        # ============ 六、模块级调用 ============
        ('brk',),
        ('h1', '六、模块级调用（71 行）'),
        ('x', 71, 'configure_console()', '★ 模块加载时立即执行 configure_console()。效果：**只要 import week11_common，控制台编码就被修好**，无需各脚本显式调用。这就是为什么算法脚本只 import 了 configure_matplotlib/ensure_results_dir/set_seed 却也能享受 UTF-8 控制台的原因——configure_console 在 import 时自动跑掉了。'),
        ('callout', '这是"副作用式 import"的典型例子：import 一个模块时自动执行模块级代码。优点：调用方零负担；注意点：副作用应只做"环境配置"这类安全操作，不能有破坏性行为。'),

        # ============ 七、调用关系图 ============
        ('brk',),
        ('h1', '七、调用关系总览'),
        ('flow', '''                   week11_common.py
   ┌──────────────────┼───────────────────┬───────────────┐
   │                  │                   │               │
 configure_console  ensure_results_dir  set_seed      configure_matplotlib
   │(import 自动)       │(画图输出)          │(main 开头)    │(模块加载时)
   ▼                   ▼                   ▼               ▼
 所有 6 个脚本        reinforce/ac/ppo/  6 个脚本         reinforce/ac/ppo/
 (靠 import 触发)     compare/large_grid  main() 首行     compare/large_grid
                      的 plot_results      set_seed()     模块加载即配置

调用顺序（以 ppo.py 为例）：
  import week11_common
    → 模块级 configure_console() 立即生效（编码修好）
    → configure_matplotlib() 立即生效（中文字体设好）
  main()
    → set_seed(args.seed)  （可复现）
    → ppo() → plot_results()
      → ensure_results_dir(args.output_dir) / "week11_ppo_training.png"'''),

        # ============ 八、面试要点 ============
        ('brk',),
        ('h1', '八、面试要点'),
        ('bullet', '「为什么把公共逻辑抽成模块？」→ 6 个脚本共享 4 类环境问题（路径/编码/种子/字体），抽取后算法脚本只关心算法，工程边界清晰；这也是"DRY 原则"和"基础设施与业务分离"的体现。', 0),
        ('bullet', '「PROJECT_ROOT 怎么动态定位项目根？」→ Path(__file__).resolve().parents[1]：取本文件绝对路径后向上两级，保证仓库可移植，避免硬编码绝对路径。', 0),
        ('bullet', '「torch 导入失败怎么兜底？」→ try/except ImportError 里 torch=None，后续 if torch is not None 判断，保证画图/设种子场景不崩。', 0),
        ('bullet', '「怎么保证实验可复现？」→ set_seed 统一设置 random/numpy/torch 三层种子；注意 GPU 还需 manual_seed_all。', 0),
        ('bullet', '「matplotlib 中文方框怎么修？」→ 按优先级扫描已装中文字体（微软雅黑等）设为 font.family，再关掉 Unicode 负号。', 0),

        ('h1', '九、常见疑问 FAQ'),
        ('bullet', 'Q：configure_console 用 errors="replace" 会不会丢字符？→ 会，但只是控制台显示层面用 ? 替代；文件和数据不受影响。比起程序崩溃，显示替代符是可接受的。', 0),
        ('bullet', 'Q：from __future__ import annotations 有什么用？→ 让 str | Path | None 这类 3.10+ 联合类型注解在旧版本也不报错（延迟求值）。本文件需要它在多个签名里放心写类型注解。', 0),
        ('bullet', 'Q：为什么不把 configure_matplotlib 的 import 放顶部？→ 函数内导入实现"用到才加载"，避免不需要 matplotlib 的脚本（如 continuous_env）白白加载重库。', 0),
        ('bullet', 'Q：DejaVu Sans 不含中文字形，放候选列表干嘛？→ 作为最后兜底——至少保证有字体可用不崩；若连它都没有说明环境极简。正常机器都会命中微软雅黑。', 0),
        ('bullet', 'Q：模块级副作用（configure_console()）安全吗？→ 安全。它只改 stdout/stderr 的编码，是无破坏性的环境配置；import 即生效，调用方零成本。', 0),
    ],
    'Week11_Step0_Common_逐行精讲.docx'
)

print('\nAll 5 detailed docs generated successfully! (v2 + Step0)')

