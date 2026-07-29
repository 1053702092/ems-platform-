#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 Week 11 — 连续动作 RL 对比报告 docx"""
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = PROJECT_ROOT / 'docs' / 'notes'
RESULTS_DIR = PROJECT_ROOT / 'results'
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def p(text, bold=False, sz=11, color=None, indent=0):
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(3)
    if indent: pa.paragraph_format.left_indent = Cm(indent)
    run = pa.add_run(text); run.font.name='微软雅黑'; run.font.size=Pt(sz); run.bold=bold
    if color: run.font.color.rgb = color

def b(text, sz=11, color=None): p(text, bold=True, sz=sz, color=color)

def bullet(text, lv=0):
    pa = doc.add_paragraph(text, style='List Bullet')
    pa.paragraph_format.left_indent = Cm(1.5 + lv*0.8)

def code(lines, label=None):
    if label:
        pa = doc.add_paragraph(); run=pa.add_run(label); run.bold=True; run.font.size=Pt(10)
    for line in lines.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before=Pt(0); pa.paragraph_format.space_after=Pt(1)
        pa.paragraph_format.left_indent=Cm(1)
        run=pa.add_run(line); run.font.name='Consolas'; run.font.size=Pt(9.5); run.font.color.rgb=RGBColor(0x33,0x33,0x33)

def tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i,hd in enumerate(headers):
        run=t.rows[0].cells[i].paragraphs[0].add_run(hd); run.bold=True; run.font.size=Pt(10); run.font.name='微软雅黑'
    for rd in rows:
        row=t.add_row()
        for c,txt in enumerate(rd):
            run=row.cells[c].paragraphs[0].add_run(txt); run.font.size=Pt(10); run.font.name='微软雅黑'

def brk(): doc.add_page_break()

def note(text):
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    run = pa.add_run(f'  {text}')
    run.font.name = '微软雅黑'; run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

def img(path, width=15):
    path = Path(path)
    if path.exists():
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pa.add_run()
        run.add_picture(str(path), width=Cm(width))
    else:
        p(f'[图片未找到: {path}]', color=RGBColor(0xC0,0x39,0x2B))

# ======================== 封面 ========================
for _ in range(3): doc.add_paragraph('')
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t.add_run('Week 11 — 连续动作 RL\nREINFORCE vs Actor-Critic vs PPO'); run.font.size=Pt(24); run.bold=True; run.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t2.add_run('从离散 DQN 到连续动作，三种方法的递进关系'); run.font.size=Pt(13); run.font.color.rgb=RGBColor(0x66,0x66,0x66)
t3=doc.add_paragraph(); t3.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t3.add_run(f'生成日期：{datetime.date.today().isoformat()}'); run.font.size=Pt(10); run.font.color.rgb=RGBColor(0x99,0x99,0x99)
brk()

# ================================================================
h('一、为什么 DQN 不行', 1)
p('DQN 的核心是输出每个动作的 Q 值，然后 argmax。在 GridWorld 的 4 个离散动作上没问题，但 EMS 的 P_fc 是 0-30 kW 连续值：', indent=0)
bullet('连续动作有无穷多个可能值，没法遍历 argmax')
bullet('如果强行离散化（切成 30 档），精度损失大，输出层爆炸')
bullet('这不是调参能解决的，是 Q-learning 数学结构决定的')
b('结论：连续动作必须换方法 → 策略梯度', color=RGBColor(0x1F,0x3A,0x5F))

brk()

# ================================================================
h('二、三种方法的递进关系', 1)

h('2.1 REINFORCE（策略梯度）', 2)
p('核心思想：不经过 Q 值，直接输出动作分布 π(s) → [μ, σ]', indent=0)
code('''# DQN:      Q(s)  → [Q↑, Q↓, Q←, Q→] → argmax → 离散动作
# REINFORCE: π(s)  → [μ, σ]  → Normal(μ,σ) → 采样 → 连续动作''')
p('')
p('更新公式:', bold=True)
code('''∇J = E[ ∇log π(a|s) × G ]   ← G 是整局回报
如果 G > 0（好结果）→ 增大这个动作的概率
如果 G < 0（差结果）→ 减小这个动作的概率''')
p('')
b('问题：等整局跑完才知道好坏，方差大、学得慢。', color=RGBColor(0xC0,0x39,0x2B))

h('2.2 Actor-Critic', 2)
p('加一个 Critic 网络 V(s) 来"当场评价"：', indent=0)
code('''Actor（演员）:  π(s)  → [μ, σ] → 采样动作（和 REINFORCE 一样）
Critic（评委）:  V(s)  → 标量值（新加的）''')
p('')
p('Advantage = r + γ·V(s\') - V(s)  → "这一步比预期好多少？"', bold=True)
p('更新：每步都能更新，不用等到整局结束。', indent=0)
p('')
b('优点：每步更新，方差比 REINFORCE 小。', color=RGBColor(0x1F,0x3A,0x5F))
b('缺点：Critic 估计不准，可能导致策略一步改太多搞崩。', color=RGBColor(0xC0,0x39,0x2B))

h('2.3 PPO（Proximal Policy Optimization）', 2)
p('在 Actor-Critic 基础上加了一个"保险"——clip 机制：', indent=0)
code('''ratio = π_new(a|s) / π_old(a|s)    ← 这个动作的概率变了多少倍？
L_clip = min(ratio × A, clip(ratio, 1-ε, 1+ε) × A)
                        ↑ 砍掉太大的改动''')
p('')
p('clip 的作用：', bold=True)
bullet('ratio > 1.2：动作概率增加了 1.2 倍以上 → 砍掉，不奖励更多')
bullet('ratio < 0.8：动作概率减少了 0.8 倍以下 → 砍掉，不惩罚更多')
bullet('策略每步只能"微调"，不能"突变"')
p('')
b('效果：训练更稳定，不容易崩。这是 EMS 选 PPO 的第一原因。', color=RGBColor(0x1F,0x3A,0x5F))

brk()

# ================================================================
h('三、实验对比', 1)
p('在同一个 EMS 简化环境上跑 500 局：', indent=0)
bullet('状态：[SOC, P_load]（2 维连续）')
bullet('动作：P_fc ∈ [0, 1]（归一化，对应 0-30 kW）')
bullet('奖励：-fuel_cost - tracking_penalty - SOC_penalty')
bullet('电池容量 5000 kWh，每局最多 200 步')

tbl(
    ['方法', '训练时间', '最后 50 局平均奖励', '特点'],
    [
        ['REINFORCE', '112s', '-38.3', '等整局结束才更新，方差大'],
        ['Actor-Critic', '269s', '-31.3', '每步更新，接近-30的理论最优'],
        ['PPO', '71s', '-39.5', 'clip+熵奖励探索多，收敛最慢'],
    ]
)
p('')
p('结果分析：', bold=True)
bullet("AC 学得最好（-31.3），已经接近理论最优（-30，即纯燃料成本）")
bullet('PPO 目前最低，因为熵奖励鼓励探索，需要更多局数收敛')
bullet('但 PPO 的优势在更复杂的问题上才体现：稳定、安全、不会突然崩')

p('实验图：三种连续动作 RL 方法训练对比', bold=True)
img(RESULTS_DIR / 'week11_comparison.png', width=16)

p('单算法训练曲线：', bold=True)
img(RESULTS_DIR / 'week11_reinforce_training.png', width=14.5)
img(RESULTS_DIR / 'week11_ac_training.png', width=14.5)
img(RESULTS_DIR / 'week11_ppo_training.png', width=14.5)

brk()

h('四、三种方法核心公式对比', 1)
tbl(
    ['', 'REINFORCE', 'Actor-Critic', 'PPO'],
    [
        ['策略', 'π(s) → μ,σ', 'π(s) → μ,σ', 'π(s) → μ,σ'],
        ['价值估计', 'G_t = Σr(MC)', 'A = r+γV-V\' (TD)', 'GAE (平滑版 TD)'],
        ['更新时机', '整局结束', '每步', '每局结束（多轮）'],
        ['核心loss', '-logπ × G', '-logπ × A', '-min(clip ratio×A)'],
        ['防崩机制', '无', '无', 'clip + 梯度裁剪'],
        ['方差', '高', '中', '低'],
        ['训练速度', '慢', '中', '中（但更稳）'],
    ]
)

brk()

h('五、为什么 EMS 选 PPO', 1)
p('面试回答（30 秒版）：', bold=True, sz=12)
p('"EMS 的核心变量 P_fc 是 0-30 kW 连续值，DQN 没法处理连续动作。PPO 通过输出正态分布参数 [μ, σ] 来采样连续动作，天生适合连续控制。同时 PPO 的 clip 机制限制了每步的更新幅度，训练比 Actor-Critic 稳定得多，不会出现一步更新把策略搞崩的情况。在实车部署时，这种稳定性比训练速度更重要。"', indent=0.5)
p('')
p('选型对比：', bold=True)
tbl(
    ['因素', 'DQN', 'PPO', 'SAC'],
    [
        ['连续动作', '❌ 不行', '✅ 可以', '✅ 可以'],
        ['训练稳定性', '⚠️ 一般', '✅ 最稳（clip）', '✅ 较稳'],
        ['实现复杂度', '简单', '中等', '复杂（5个网络）'],
        ['采样效率', '高（off-policy）', '低（on-policy）', '高（off-policy）'],
        ['调参难度', '低', '中', '较高'],
        ['EMS 适用性', '❌', '✅ 推荐', '✅ 备选'],
    ]
)

brk()

h('六、核心记忆点', 1)
p('1. DQN → REINFORCE → AC → PPO 是递进关系，每一步解决上一步的问题', sz=12)
p('2. 连续动作 = 输出 [μ, σ] 采样，不能再用 argmax', sz=12)
p('3. PPO 的核心就是一行 clamp：clip(ratio, 0.8, 1.2) × A', sz=12, bold=True)
p('4. EMS 选 PPO 原因：连续动作 + 训练稳定 + 实现适中', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
note(f'生成日期: {datetime.date.today().isoformat()}')

# 保存
OUT_DIR.mkdir(parents=True, exist_ok=True)
path = OUT_DIR / 'Week11_连续动作RL对比报告.docx'
doc.save(str(path))
print(f'OK: {path}')
