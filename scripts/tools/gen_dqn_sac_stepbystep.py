#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""DQN / SAC 逐级详解 — 从你会的 Q-learning 出发，每一步都讲清楚为什么"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(2)

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
def p(text, bold=False, sz=11, color=None, indent=0):
    pa = doc.add_paragraph()
    if indent: pa.paragraph_format.left_indent = Cm(indent)
    run = pa.add_run(text); run.font.name='微软雅黑'; run.font.size=Pt(sz); run.bold=bold
    if color: run.font.color.rgb = color
def b(text, sz=11, color=None): p(text, bold=True, sz=sz, color=color)
def bullet(text, lv=0):
    pa = doc.add_paragraph(text, style='List Bullet')
    pa.paragraph_format.left_indent = Cm(1.5 + lv * 0.8)
def code(lines, label=None):
    if label:
        pa = doc.add_paragraph(); run=pa.add_run(label); run.bold=True; run.font.size=Pt(10)
    for line in lines.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before=Pt(0); pa.paragraph_format.space_after=Pt(0)
        pa.paragraph_format.left_indent=Cm(1)
        run=pa.add_run(line); run.font.name='Consolas'; run.font.size=Pt(9.5); run.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i,hd in enumerate(headers):
        run=t.rows[0].cells[i].paragraphs[0].add_run(hd); run.bold=True; run.font.size=Pt(10); run.font.name='微软雅黑'
    for rd in rows:
        row=t.add_row()
        for c,txt in enumerate(rd):
            run=row.cells[c].paragraphs[0].add_run(txt); run.font.size=Pt(10); run.font.name='微软雅黑'
def brk(): doc.add_page_break()
def note(text):
    """高亮提示框"""
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    run = pa.add_run(f'   {text}')
    run.font.name = '微软雅黑'; run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

# ======================== 封面 ========================
for _ in range(4): doc.add_paragraph('')
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t.add_run('DQN / SAC\n逐级详解'); run.font.size=Pt(26); run.bold=True; run.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t2.add_run('从 Q-learning 出发，每一步讲清楚"为什么"'); run.font.size=Pt(13); run.font.color.rgb=RGBColor(0x66,0x66,0x66)
t3=doc.add_paragraph(); t3.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t3.add_run(f'生成日期：{datetime.date.today().isoformat()}'); run.font.size=Pt(10); run.font.color.rgb=RGBColor(0x99,0x99,0x99)
brk()

# ================================================================
# 第一章：你已经会了什么（复习）
# ================================================================
h('第一章：你已经会了什么', 1)
p('在学 DQN 之前，你已经掌握了两个关键基础：')
p('')
b('1. 值迭代（Week 9 Part 8）')
p('  V(s) = max_a [ R(s,a) + gamma * sum P * V(sp) ]')
p('  特点：知道 P，直接算。不需要真去走。')
p('')
b('2. Q-learning（你刚学的）')
p('  Q[s][a] += lr * ( r + gamma * max Q[sp] - Q[s][a] )')
p('  特点：不知道 P，真去走一步。用一张 Q 表记经验。')
p('')
p('DQN = Q-learning + 神经网络（把 Q 表换成网络）', bold=True, sz=12, color=RGBColor(0x1F,0x3A,0x5F))
p('SAC = Q-learning 解决不了连续动作问题 → 换成另一种思路（策略网络）')
p('')
p('这个文档就从这两个基础出发，一步步走向 DQN 和 SAC。', bold=True)
brk()

# ================================================================
# 第二章：从 Q 表到神经网络
# ================================================================
h('第二章：从 Q 表到神经网络（DQN 的核心）', 1)

h('2.1 先回忆：你刚写的 Q-learning 是怎么工作的', 2)
p('你刚写完的 Q-learning 代码核心就两样东西：')
p('')
b('1. Q 表：一个 16×4 的数组')
code('''Q = np.zeros((16, 4))   # 16 个格子 × 4 个动作

Q[0] = [0.19, 0.17, 0.25, 0.41]   ← 格子0 的 4 个 Q 值: ↑ ↓ ← →
Q[1] = [0.10, 0.30, 0.20, 0.50]   ← 格子1 的 4 个 Q 值
...''')
p('')
b('2. 更新代码：查 Q 表 → 执行 → 更新 Q 表')
code('''# 选动作
a = argmax Q[s]        # 查 Q 表第 s 行，取最大的列

# 执行
sp, r = step(s, a)     # 真走一步

# 更新
Q[s][a] += lr * (r + gamma * max(Q[sp]) - Q[s][a])   # 改 Q 表的一个格子''')

p('')
p('这个流程在 GridWorld 上工作得很好。因为 16 个格子，64 个 Q 值，一张小表格就够了。')
p('但想象一下：如果状态不是 16 个格子，而是：')
bullet('一部手机在充电时的电压、电流、温度、SOC、老化程度... → 连续值，无限个状态')
bullet('Atari 游戏的画面 → 128×128 像素 × 3 通道 = 49152 个值组合成无限种状态')
p('')
b('Q 表的致命问题：状态太多，表格装不下。', color=RGBColor(0xC0,0x39,0x2B))
p('GridWorld 16 个格子，刚好装下。')
p('Atari 游戏画面，全世界所有的硬盘加起来也存不下一张 Q 表。')

h('2.2 神经网络怎么代替 Q 表', 2)
b('核心思路：不"记"每个状态的 Q 值，而是"猜"没见过的状态的 Q 值。')
p('')
p('类比两种学习方法：')
bullet('Q 表 = 背答案。考试出了原题就会，换个数字就不会了。')
bullet('神经网络 = 理解了公式。题目怎么变都能算。')
p('')

b('具体做法：')

p('第 1 步：把 Q 表扔了，换成神经网络 Q_theta', bold=True, indent=0.5)
code('''# Q-learning:      Q = np.zeros((16, 4))        ← 表格
# DQN:             Q_theta = NeuralNetwork()    ← 网络

输入：状态 s（比如 one-hot 向量 [0,0,1,0,...,0] 表示"格子2"）
输出：4 个 Q 值   [q_up, q_down, q_left, q_right]''')

p('')
p('第 2 步：网络内部长这样', bold=True, indent=0.5)
code('''输入层(16个神经元)       ← 状态 s（16个格子，one-hot）
    ↓ 全连接 W1(16×32) + ReLU
隐藏层(32个神经元)
    ↓ 全连接 W2(32×4)
输出层(4个神经元)        ← 4 个 Q 值''')

p('')
p('第 3 步：选动作时，不是查 Q 表，而是跑一次网络', bold=True, indent=0.5)
code('''# Q-learning:
a = argmax Q[s]              # 查表，O(1) 时间复杂度

# DQN:
q_values = Q_theta(s)        # 网络前向传播，O(参数数) 时间
a = argmax(q_values)         # 选最大的 Q 值''')

p('')
p('第 4 步：更新时，不是改表格的一个格子，而是梯度下降', bold=True, indent=0.5)
code('''# Q-learning:
Q[s][a] += lr * (target - Q[s][a])        # 改一个格子

# DQN:
loss = MSE(Q_theta(s,a), target)           # 算误差
loss.backward()                            # 算梯度
optimizer.step()                           # 更新所有参数''')

note('关键区别：改一个格子只影响那一个状态。但梯度下降更新网络参数，会影响类似状态的输出。'
     '这就是"泛化"——没见过格子8的 Q 值，但格子7和格子9的 Q 值会影响它。')

h('2.3 但直接这么用会出问题（两个大坑）', 2)
p('如果你直接把 Q 表换成神经网络，按 Q-learning 的方式训练，会发现两个问题：')
p('')

b('问题 1：数据相关性 — 网络学不会', color=RGBColor(0xC0,0x39,0x2B))
p('Q-learning 用 Q 表时，数据是按顺序来的：')
code('''格子0 → 动作↓ → 格子4 → 动作→ → 格子5(Trap!) → 重置
        ↑                      ↑                       ↑
     时刻 t                  时刻 t+1                时刻 t+2''')
p('按顺序训练神经网络：相邻样本高度相关，网络会"过拟合"到最近几步。')
p('')
b('类比：', color=RGBColor(0xC0,0x39,0x2B))
p('  你连续做 10 道"鸡兔同笼"的题，第 11 道改成"相遇问题"，你的脑子还转不过来。')
p('  但如果 10 道题是随机混着的，你的脑子就能适应。')
p('')
b('解决方案：经验回放（Experience Replay）', color=RGBColor(0x1F,0x3A,0x5F))
p('把经验存到一个大缓冲区，训练时从中随机采样，打乱顺序。')
code('''buffer = [(s0,a0,r0,sp0), (s1,a1,r1,sp1), ..., (s9999,a9999,r9999,sp9999)]

# 不按顺序学，而是随机抽 32 条出来学
batch = random.sample(buffer, 32)
for s,a,r,sp in batch:
    target = r + gamma * max(Q_theta(sp))
    loss = MSE(Q_theta(s,a), target)
    ...''')
p('')
p('经验回放其实就是告诉神经网络："别死盯着当前这一步，回头看看以前的经验。"', bold=True)

brk()

p('')
b('问题 2：目标不稳定 — 网络追着自己的尾巴跑', color=RGBColor(0xC0,0x39,0x2B))
p('Q-learning 的更新公式：')
code('''target = r + gamma * max Q(sp)
                   ↑
              这个 Q 和 Q(s,a) 是同一个 Q 表！
              更新 Q(s,a) 时，max Q(sp) 也变了''')
p('')
p('用 Q 表时这个问题不明显，因为每次只改一个格子，对其他格子影响小。')
p('但用神经网络时，一次更新可能改变所有状态的 Q 值输出！')
p('')
b('具体来说：', color=RGBColor(0xC0,0x39,0x2B))
code('''第 1 步：Q_theta 说格子3的 Q 值是 0.5，我们算 target = r + 0.9*0.5 = ...
第 2 步：梯度下降更新 theta，Q_theta 变了
第 3 步：再用 Q_theta 算格子3的 Q 值，发现变成 0.7 了
          → 之前算的 target 现在不对了！
第 4 步：再用新 target 更新，Q_theta 又变了...
          → 永远追不上！''')
p('')
b('类比：', color=RGBColor(0xC0,0x39,0x2B))
p('  你射箭，靶子绑在你的箭上。箭飞出去，靶子也跟着飞。永远射不中。')
p('')
b('解决方案：目标网络（Target Network）', color=RGBColor(0x1F,0x3A,0x5F))
p('维护两个神经网络：')
bullet('在线网络 Q_theta：负责预测，每步都更新')
bullet('目标网络 Q_target：负责计算 target，不频繁更新')
code('''# 用目标网络算 target（它不动，靶子就稳了）
target = r + gamma * max Q_target(sp)     # Q_target 不更新

# 训练在线网络
loss = MSE(Q_theta(s,a), target)           # 只有 Q_theta 更新
loss.backward()
optimizer.step()

# 每隔 200 轮，把 Q_theta 复制给 Q_target
Q_target.load_state_dict(Q_theta.state_dict())''')
p('')
b('这样 target 相对稳定，在线网络可以去追赶一个"不那么快变"的目标。', color=RGBColor(0x1F,0x3A,0x5F))

h('2.4 现在把 DQN 的完整流程串起来', 2)
code('''# 初始化
Q_theta = DQN()              # 在线网络（负责预测）
Q_target = DQN()             # 目标网络（负责算 target）
buffer = []                  # 经验回放缓冲区
epsilon = 1.0                # 探索率

for episode in range(5000):
    s = 0
    while not done:
        # 1. ε-贪心选动作（和 Q-learning 一样！）
        if random() < epsilon: a = random_action()
        else: a = argmax Q_theta(s)           # 用在线网络

        # 2. 执行动作（和 Q-learning 一样！）
        sp, r, done = step(s, a)
        buffer.append((s, a, r, sp, done))    # 存起来

        # 3. 训练（和 Q-learning 不一样：用网络，不是 Q 表）
        if len(buffer) > 32:
            batch = random.sample(buffer, 32) # 经验回放！
            for s,a,r,sp,done in batch:
                target = r + gamma * max Q_target(sp)  # 目标网络！
                loss = MSE(Q_theta(s,a), target)
                loss.backward()
                optimizer.step()               # 更新 Q_theta

        # 4. 更新目标网络（每 200 局）
        if episode % 200 == 0:
            Q_target.load_state_dict(Q_theta.state_dict())

        s = sp''')
p('')
b('对比 Q-learning 和 DQN，你会发现框架完全一样：', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
tbl(
    ['步骤', 'Q-learning', 'DQN'],
    [
        ['存储 Q 值的地方', 'Q 表 [16×4]', '神经网络 Q_theta'],
        ['选动作', 'argmax Q[s]', 'argmax Q_theta(s)'],
        ['算 target', 'r + gamma * max Q[sp]', 'r + gamma * max Q_target(sp)'],
        ['更新', '改 Q 表的一个格子', '梯度下降更新所有参数'],
        ['处理数据相关性', '不需要（Q 表不受顺序影响）', '经验回放（随机采样）'],
        ['处理目标不稳定', '不需要（改一个格子影响小）', '目标网络（冻结 C 步）'],
    ]
)
p('')
b('DQN = Q-learning + 神经网络 + 经验回放 + 目标网络', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
p('前面两个（神经网络、经验回放、目标网络）都是因为"用神经网络代替 Q 表"引发的新问题。', bold=True)
brk()

# ================================================================
# 第三章：从离散到连续
# ================================================================
h('第三章：DQN 的局限 — 为什么还需要 SAC', 1)

h('3.1 DQN 的先天缺陷：只能处理离散动作', 2)
p('DQN 的输出层是"每个动作一个神经元"：')
code('''输出层: [Q(↑), Q(↓), Q(←), Q(→)]  ← 4 个神经元，对应 4 个离散动作
选动作: argmax → 选 Q 值最大的那个''')
p('')
b('如果是连续动作（比如"功率调到 0-30 kW 之间的任意值"），怎么办？', color=RGBColor(0xC0,0x39,0x2B))
p('')
p('方案 1：离散化（把连续值切成几档）', bold=True)
code('''P_fc 范围 [0, 30] kW，切 7 档：
{0, 5, 10, 15, 20, 25, 30} kW
→ 7 个输出神经元

问题：最优功率可能是 17.3 kW，但只能选 15 或 20，精度丢失！''')
p('')
p('方案 2：增加离散粒度', bold=True)
code('''切 100 档：每档 0.3 kW
→ 100 个输出神经元
→ 动作越多，输出层越大，训练越慢''')
p('')
p('方案 3：输出层做文章', bold=True)
code('''有人试过让 DQN 直接输出一个连续值（比如输出功率的数值）
但 Q-learning 的核心 argmax 要求"遍历所有动作取最大值"
连续值没法 argmax！''')
p('')
b('结论：DQN 从根本上无法处理连续动作。这不是调参能解决的，是数学结构决定的。', color=RGBColor(0xC0,0x39,0x2B))

h('3.2 换个思路：不经过 Q 值，直接输出动作', 2)
p('如果不用 Q 值绕圈子，直接让神经网络输出"应该做什么动作"呢？')
p('')
p('这就是策略网络（Policy Network）的思路：', bold=True)
code('''# DQN（Value-Based）
Q_theta(s) → [Q↑, Q↓, Q←, Q→]   → argmax → 选动作 a
  中间多了一层 Q 值

# Policy-Based
pi_theta(s) → 直接输出动作 a
  或者输出动作分布（比如均值 μ + 标准差 σ）
  不经过 Q 值！''')
p('')
b('这样连续动作就自然了：', color=RGBColor(0x1F,0x3A,0x5F))
code('''pi_theta(s) 输出：μ = 20.5, σ = 3.2
→ 从 N(20.5, 3.2) 分布中采样
→ 得到 a = 21.3 kW  （连续值！）''')

brk()

h('3.3 策略梯度思想：让好动作概率变大，坏动作概率变小', 2)
p('Policy-Based 的核心思想很直观：')
p('')
code('''# 智能体走完一局，得到总奖励 G
如果 G 大 → 这局的动作都不错 → 增大这些动作的概率
如果 G 小 → 这局的动作不好   → 减小这些动作的概率

grad_J = E[ grad_log_pi(a|s) * G ]
  提升方向 = 动作概率的梯度 × 整局奖励''')
p('')
p('但这里有个问题：')
code('''REINFORCE：等整局走完才更新
问题：走了 50 步，第 10 步走错了，第 40 步走对了
      但整局奖励 = -10，是所有的步子一起挨批还是只批第 10 步？
      说不清楚！→ 方差极大''')
p('')
b('解决方案：Actor-Critic（演员-评委）', color=RGBColor(0x1F,0x3A,0x5F))

brk()

h('3.4 Actor-Critic 架构', 2)
p('不等到整局结束，每走一步就做一次评价：')
p('')
code('''Actor（演员）：pi_theta(s) → 输出动作 a
Critic（评委）：Q_phi(s,a) → 评价"这一步走得好不好"

更新时：
  Actor 听 Critic 的：  如果 Critic 说 Q 值高 → 增大这个动作的概率
  Critic 听实际奖励的： 如果实际奖励和预期不符 → 调整 Critic 的评价标准''')
p('')
b('类比：', color=RGBColor(0x1F,0x3A,0x5F))
p('  REINFORCE = 演员演完整部电影，看票房（G）才知道自己演得怎么样')
p('  Actor-Critic = 演员演完一幕，导演（Critic）当场说"这幕好/这幕不好"')
p('  导演也在不断学习，他的评价标准越来越准')
p('')
p('SAC 就属于 Actor-Critic 流派。但 SAC 还有一个关键创新——最大熵。')

brk()

# ================================================================
# 第四章：最大熵
# ================================================================
h('第四章：SAC 的最大熵 — 为什么策略要保持随机', 1)

h('4.1 先看一个标准 RL 会出问题的例子', 2)
p('假设一个 GridWorld，有两个格子同样值得去：')
code('''起点格子0：往右 = +0.8，往下 = +0.8   ← 两个动作一样好

标准 RL（无熵）：
  学到最终 → Q(→) = 0.81, Q(↓) = 0.79
  argmax → 永远选"→"
  哪怕"↓"其实同样好，它也死认"→"''')
p('')
p('这有什么问题？', bold=True)
bullet('如果环境变了，"→"的路被堵死了，智能体不会自动切到"↓"')
bullet('在真实 EMS 中：一个 FC 衰退了，另一个 FC 状况良好，但策略死认第一个')

h('4.2 SAC 的解法：给"随机性"发奖励', 2)
p('SAC 在标准 RL 目标上加了一项——熵（Entropy）：')
code('''标准 RL 的目标： max E[ r_total ]        ← 只看奖励

SAC 的目标：       max E[ r_total + α × H ]  ← 奖励 + 熵
                                              ↑
                                    也给"保持随机"加分''')
p('')
b('熵 H 是什么？', color=RGBColor(0x1F,0x3A,0x5F))
p('H = -sum( p * log p )  衡量不确定性')

code('''策略 A：{↑=0.0, ↓=0.0, ←=0.0, →=1.0}  → H = 0    ← 完全确定
策略 B：{↑=0.25, ↓=0.25, ←=0.25, →=0.25} → H = 1.39 ← 完全随机

策略 A 死认"→"，策略 B 四个方向都试试''')

p('')
b('加熵奖励的效果：', color=RGBColor(0x1F,0x3A,0x5F))
code('''两个动作同样好时（Q=0.8 和 Q=0.8）：
  标准 RL：argmax 随机选一个，然后死认这个
  SAC：两个动作都保持概率 0.5/0.5 → 下次还能选另一个

一个动作更好时（Q=0.9 和 Q=0.6）：
  SAC 会偏向好的动作，但不会完全抛弃差的
  → 大概率选好的，小概率试差的''')

h('4.3 三个层面理解最大熵的好处', 2)

b('好处 1：鼓励探索', color=RGBColor(0x1F,0x3A,0x5F))
p('标准 RL 一旦发现"→"不错，就再也不试其他方向了。万一"↓"更好呢？')
p('SAC 保持一定的随机性，继续探索——"老走→，偶尔也走走↓"。')
p('')

b('好处 2：鲁棒性', color=RGBColor(0x1F,0x3A,0x5F))
p('如果环境变了（比如路堵了），确定性策略要重新学。')
p('SAC 因为一直保持随机尝试，可能自动切换到另一个可行路径。')
p('在 EMS 中：燃料电池会老化，训练时的特性和实际运行时的特性不同。鲁棒策略更重要。')
p('')

b('好处 3：多模态（Multiple Modes）', color=RGBColor(0x1F,0x3A,0x5F))
p('当多个策略同样好时，SAC 会学会"分头下注"，而不是赌一个。')

h('4.4 α 自动调节：不手动设探索率', 2)
p('α（温度系数）控制"探索"和"利用"的平衡：')
bullet('α 大 → 熵奖励权重大 → 策略更随机 → 多探索')
bullet('α 小 → 熵奖励权重小 → 策略更确定 → 多利用')
p('')
b('SAC 把 α 当做一个可训练的参数，自动调节：', color=RGBColor(0x1F,0x3A,0x5F))
code('''目标熵 H_target = -动作空间维数（比如 4 维动作 → -4）

如果当前策略太确定（H < H_target）→ 增大 α → 多探索
如果当前策略太随机（H > H_target）→ 减小 α → 多利用

你不需要手动设 ε，SAC 自动找到合适的探索程度。''')

brk()

# ================================================================
# 第五章：SAC 的网络结构
# ================================================================
h('第五章：SAC 为什么需要 5 个网络', 1)

h('5.1 逐个拆解：每个网络负责什么', 2)

b('网络 1：Actor π_phi（策略网络）', color=RGBColor(0x1F,0x3A,0x5F))
code('''输入：状态 s
输出：动作分布（均值 μ + 标准差 σ）
      比如 μ = 20.5 kW, σ = 3.2 kW
更新目标：让 Q 值 + 熵 最大化 → 找到"高 Q 值但仍保持随机"的动作''')
p('')
b('网络 2 & 3：Critic Q_theta1, Q_theta2（两个价值网络）', color=RGBColor(0x1F,0x3A,0x5F))
code('''输入：状态 s + 动作 a
输出：Q 值（这个动作好不好的评分）
更新目标：让 Q 值更接近实际观察到的回报
为什么两个？取 min(Q1, Q2) 作为最终 Q 值 → 防止高估''')
p('')
b('网络 4 & 5：目标网络 Q_target1, Q_target2', color=RGBColor(0x1F,0x3A,0x5F))
code('''和 DQN 的目标网络同理：冻结参数，提供稳定的训练目标
软更新：theta_target = τ * theta + (1-τ) * theta_target  (τ=0.005)''')

p('')
b('5 个网络的分工一句话：', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
p('  Actor 决定"做什么"，Critic 评价"做得多好"，两个 Critic 互相纠错防止高估，')
p('  目标网络让训练目标稳定。')

h('5.2 为什么两个 Critic（而不是一个）', 2)
p('DQN 有个著名问题：Q 值被高估。原因：')
code('''Q(s,a) 本身有估计误差（可能偏高 0.1，也可能偏低 0.1）
max_a Q(s,a) → 取所有动作中最大的 → 系统性地取到被高估的那个

例：真实 Q 值 = [0.7, 0.8, 0.6, 0.7]
    估计 Q 值 = [0.75, 0.85, 0.55, 0.65]
    max = 0.85 > 真实 max 0.8 → 高估了！''')
p('')
p('SAC 的解法：训练两个独立的 Critic，取最小值')
code('''Q1 估计 = [0.75, 0.85, 0.55, 0.65]
Q2 估计 = [0.68, 0.82, 0.62, 0.72]
min(Q1, Q2) = [0.68, 0.82, 0.55, 0.65]
max = 0.82 ≈ 真实值 0.8 ✓
''')
p('两个独立网络同时高估同一个动作的概率很低。取 min 相当于"悲观估计"。')

h('5.3 SAC 的更新流程（不写公式，用自然语言）', 2)
code('''每一步 SAC 做 4 件事：

1. 更新 Critic（让评分更准）
   用目标网络算出"未来回报的估计"
   让 Q1 和 Q2 的输出接近这个估计
   取 min(Q1, Q2) 作为最终评分

2. 更新 Actor（让动作更好）
   看 Critic 的评分 → 哪个动作评分高就增大哪个动作的概率
   同时看熵 → 动作太集中了就罚（鼓励保持随机）

3. 更新 α（自动调探索程度）
   当前策略太确定？→ 增大 α（多探索）
   当前策略太随机？→ 减小 α（多利用）

4. 软更新目标网络
   Q_target ← 0.005 * Q + 0.995 * Q_target
   每次只移动一点点，保持稳定''')

brk()

# ================================================================
# 第六章：面试
# ================================================================
h('第六章：面试 Q&A', 1)

b('Q1: "DQN 相比 Q-learning 改了哪些东西？"')
p('核心改了三点：1) 用神经网络代替 Q 表，能够处理连续状态输入和泛化到未见过状态。2) 加了经验回放，解决数据相关性问题。3) 加了目标网络，解决训练目标不稳定问题。后面两点都是因为"用神经网络"之后才出现的新问题。')

p('')
b('Q2: "经验回放解决了什么问题？为什么 Q-learning 不需要？"')
p('Q-learning 的 Q 表每次只修改一个格子，数据顺序不影响。但神经网络一次更新影响所有参数，如果数据是顺序相关的（相邻几步高度相关），网络会过拟合到最近的经验。经验回放把经验存起来随机采样，打乱顺序，让数据更像独立同分布的。')

p('')
b('Q3: "为什么 DQN 不能处理连续动作？"')
p('DQN 输出层每个神经元对应一个离散动作的 Q 值，选动作要 argmax 遍历所有动作。连续动作（比如 0-30 kW）有无穷多个可能取值，无法遍历。即使离散化，精度和动作数量的矛盾也很难避免。SAC 用策略网络直接输出动作分布（均值+标准差），从分布中采样得到连续动作，绕过了 argmax 问题。')

p('')
b('Q4: "SAC 的最大熵加的是什么？为什么有用？"')
p('在标准 RL 目标上加了一项"熵奖励"。熵衡量随机性，熵越大动作越均匀。加熵奖励鼓励策略保持一定随机性：1) 继续探索，不锁死在次优解；2) 学到多种应对方式，环境变化时更鲁棒；3) 多个动作同样好时均匀分配概率。α 自动调节探索程度，不需要手动设 ε。')

p('')
b('Q5: "SAC 为什么需要两个 Critic？"')
p('DQN 的 max 操作容易高估 Q 值（取到被噪声放大的动作）。SAC 训练两个独立的 Critic 网络，取 min(Q1, Q2) 作为评分。两个网络同时高估同一个动作的概率很低，相当于一个"悲观的估计"，防止 Q 值虚高。')

brk()

# ================================================================
# 附录
# ================================================================
h('附录：三方法代码骨架对比', 1)
p('同一件事（GridWorld 找终点）的三种解法：')
p('')

b('值迭代（DP，知道 P）', color=RGBColor(0x1F,0x3A,0x5F))
code('''V = zeros(16)
for _ in range(1000):
    for s in range(16):
        q_max = -inf
        for a in range(4):
            q = R[s][a] + gamma * sum(P[s][a][sp] * V[sp])
            q_max = max(q_max, q)
        V[s] = q_max''')
note('不需要和环境交互。拿到 P 就能算。')

p('')
b('Q-learning（无模型，有 Q 表）', color=RGBColor(0x1F,0x3A,0x5F))
code('''Q = zeros((16, 4))
for episode in range(5000):
    s = 0
    while not done:
        a = argmax Q[s] if random() > epsilon else random_action()
        sp, r, done = step(s, a)         # 真走！
        Q[s][a] += lr * (r + gamma * max(Q[sp]) - Q[s][a])
        s = sp''')
note('不知道 P。自己试，用 Q 表记经验。')

p('')
b('DQN（无模型，有神经网络）', color=RGBColor(0x1F,0x3A,0x5F))
code('''Q_theta = DQN()
Q_target = DQN()
buffer = []
for episode in range(5000):
    s = 0
    while not done:
        a = argmax Q_theta(s)
        sp, r, done = step(s, a)
        buffer.append((s,a,r,sp,done))
        if len(buffer) > 32:
            batch = random.sample(buffer, 32)
            target = r + gamma * max(Q_target(sp))
            loss = MSE(Q_theta(s,a), target)
            loss.backward()
            optimizer.step()
        s = sp''')
note('和 Q-learning 一样，不知道 P。只是把 Q 表换成了网络。')

# ======================== 保存 ========================
path = os.path.join(OUT_DIR, 'DQN_SAC_逐级详解.docx')
doc.save(path)
print(f'OK: {path}')
