#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 LG 18650HG2 数据集下载指南 docx
====================================
供用户在另一台电脑上下载 McMaster University 的 LG HG2 公开数据集，
用于 SOC 估计器升级（方案 A：PyBaMM 物理模型 vs 等效电路 EKF 对比）。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\soc-estimation\doc'

BLUE = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x66, 0x66, 0x66)
DGRAY = RGBColor(0x99, 0x99, 0x99)
RED = RGBColor(0xC0, 0x39, 0x2B)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GREEN = RGBColor(0x27, 0xAE, 0x60)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.line_spacing = 1.35

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.color.rgb = BLUE
def h2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.color.rgb = ORANGE
def p(t, size=11, bold=False, color=None):
    pa = doc.add_paragraph()
    r = pa.add_run(t)
    r.font.name = 'Microsoft YaHei'; r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return pa
def bullet(t, indent=0):
    pa = doc.add_paragraph(t, style='List Bullet')
    pa.paragraph_format.left_indent = Cm(1.5 + indent*0.8)
def code(t):
    for line in t.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Cm(1)
        pa.paragraph_format.space_after = Pt(1)
        r = pa.add_run(line)
        r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
def tbl(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers))
    tb.style = 'Table Grid'
    for i, hh in enumerate(headers):
        r = tb.rows[0].cells[i].paragraphs[0].add_run(hh)
        r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'
    for rd in rows:
        row = tb.add_row()
        for c, txt in enumerate(rd):
            r = row.cells[c].paragraphs[0].add_run(txt)
            r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'

# ── 封面 ──
for _ in range(3): doc.add_paragraph('')
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('LG 18650HG2 数据集下载指南\n'); r.font.size = Pt(24); r.bold = True; r.font.color.rgb = BLUE
r = t.add_run('用于 SOC 估计器升级（方案 A：PyBaMM 物理模型 vs 等效电路 EKF）')
r.font.size = Pt(13); r.font.color.rgb = GRAY
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(f'生成日期：{datetime.date.today().isoformat()}')
r.font.size = Pt(10); r.font.color.rgb = DGRAY
doc.add_page_break()

# ── 一、数据集是什么 ──
h1('一、数据集是什么')
p('这是 SOC 估计领域最常用的公开数据集之一，由加拿大 McMaster 大学 Philip Kollmeyer 团队发布，用全新的 3Ah LG 18650HG2 电芯在恒温箱中测试得到。')
tbl(['项目', '内容'], [
    ['数据集全名', 'LG 18650HG2 Li-ion Battery Data and Example Deep Neural Network xEV SOC Estimator Script'],
    ['发布机构', 'McMaster University（Hamilton, Ontario, Canada）'],
    ['发布平台', 'Mendeley Data（爱思唯尔旗下科研数据平台）'],
    ['DOI', '10.17632/cp3473x7xv.3（V3 版本）'],
    ['发布时间', '2020 年 3 月 5 日'],
    ['数据规模', '208 个 CSV 文件，约 490 万行，约 2 GB'],
    ['测试设备', '75A / 5V Digatron 电池测试仪，电流电压精度 0.1%'],
    ['电芯', 'LG 18650HG2，3 Ah（和你现有的 SOC 标签验证报告同型号！）'],
])

# ── 二、为什么需要它 ──
h1('二、为什么升级 SOC 项目需要它')
p('你现有的 SOC 估计器（docs/soc-estimation/code/）用的是 1RC 等效电路 + EKF/AEKF，在常温下表现好，但低温 / 大倍率下 OCV-only 模型会失配。')
p('方案 A 的目标：用 PyBaMM 的电化学物理模型（DFN/P2D）替代/对比等效电路，理解参数辨识与模型失配问题。', bold=True)
p('这份数据集提供：')
bullet('HPPC（Hybrid Pulse Power Characterization）测试数据 → 辨识电池 R0 / R1 / C1 参数，这是你 STATUS.md 里"下一步"要做的')
bullet('多个温度（-10°C / 0°C / 10°C / 25°C / 40°C）→ 验证低温模型失配问题（方案 A 的核心卖点）')
bullet('真实驾驶循环（UDDS / US06 / HWFET / LA92 / 混合循环）→ 贴近实际工况验证')
bullet('C/20 小倍率放电 → 近似真实 OCV 曲线')

# ── 三、下载步骤 ──
h1('三、下载步骤（在另一台电脑操作）')
h2('3.1 打开数据集页面')
code('https://data.mendeley.com/datasets/cp3473x7xv/3')
p('（如果打开是英文页面，用浏览器自带的翻译即可。）')

h2('3.2 下载方式：推荐用 Mendeley Desktop 同步（不易断流）')
p('数据集有 2 GB，浏览器直接下载容易断。Mendeley 官方推荐方式：', bold=True)
bullet('注册/登录 Mendeley Data（免费，用邮箱注册即可，不需要机构权限）')
bullet('在数据集页面右上角点 "Download" 会下载 zip（简单方式）')
bullet('更稳的方式：页面提供 Mendeley Desktop / 数据集同步工具，用它可以断点续传')

h2('3.3 必须下载的文件')
p('整个数据集是一份 zip，解压后按温度分文件夹。关键内容：', bold=True)
tbl(['文件夹', '内容', '对方案 A 的用途'], [
    ['25degC/', '36 个文件（含 HPPC、C/20 放电、各倍率放电、驾驶循环）', '★ 主用：HPPC 辨识参数 + 常温对比'],
    ['10degC/、0degC/、-10degC/', '低温测试数据', '验证低温模型失配（简历卖点）'],
    ['40degC/', '高温测试数据', '高温工况补充'],
    ['（根目录）', 'FNN 估计脚本 .mlx 等', '参考作者的神经网络 SOC 估计方法'],
])
p('HPPC 示例文件名：549_HPPC.csv（25°C）。HPPC 文件共 5 个，占数据集约 2.4%。', size=10)

h2('3.4 解压与存放')
p('解压后不要改动文件夹结构。建议放回本项目后按以下结构存放：')
code('ems-platform/datasets/\n└── LG_HG2/\n    ├── n10degC/    # -10°C（38 个 CSV）\n    ├── 0degC/      # 0°C（32 个 CSV）\n    ├── 10degC/     # 10°C（32 个 CSV）\n    ├── 25degC/     # 25°C（36 个 CSV）← 主用\n    └── 40degC/     # 40°C（35 个 CSV）')

# ── 四、文件格式说明 ──
h1('四、CSV 文件格式说明（拿到数据后对照）')
p('每个 CSV 是多级表头（第 25-26 行），列含：Time Stamp（时间戳）、Step、Status（状态）、Prog Time、Step Time、Cycle、Procedure（工况名）、Voltage（电压 V）、Current（电流 A）、Temperature（温度 °C）、Capacity（容量）、WhAccu 等。')
p('Status 列的取值含义：', bold=True)
tbl(['取值', '含义', '用途'], [
    ['CHA', '充电', 'SOC 上升段'],
    ['DCH', '放电', 'SOC 下降段'],
    ['TABLE', '静置/查表', '静置段'],
    ['PAU', '暂停', '一般过滤掉'],
])

# ── 五、注意事项 ──
h1('五、注意事项')
bullet('引用规范：使用该数据需引用原作者 —— Kollmeyer, Philip; Vidal, Carlos; Naguib, Mina; Skells, Michael (2020), "LG 18650HG2 Li-ion Battery Data and Example Deep Neural Network xEV SOC Estimator Script", Mendeley Data, V3, doi: 10.17632/cp3473x7xv.3')
bullet('相关论文：C. Vidal et al., "Robust xEV Battery State-of-Charge Estimator Design using Deep Neural Networks," SAE World Congress, Apr 2020.（读这篇能理解作者怎么用这份数据做 SOC 估计）')
bullet('网络提示：Mendeley 在某些网络下可能访问慢或被限制，如果页面打不开，可尝试科学上网或换网络环境。')
bullet('已有数据提醒：你之前跑过 LG HG2 SOC 标签验证（RMSE=0.0465），说明这台机器上曾有数据，检查一下 datasets/ 下是否残留旧数据，能省一次下载。')
bullet('下载大小约 2 GB，请预留磁盘空间。下载完成后核对：208 个 CSV、约 490 万行。')

# ── 六、下载完成后在本项目怎么用 ──
h1('六、下载完成后在项目里的使用路径')
p('拿到数据后，直接告诉我数据放好了，我会继续推进方案 A 的后续步骤：')
flow = ['1. 解析 CSV → 提取电流/电压/温度序列',
        '2. 用 HPPC 数据辨识 R0/R1/C1（你 STATUS.md 里的下一步）',
        '3. 在现有 EKF/AEKF 估计器上跑真实数据',
        '4. 用 PyBaMM 的 LG 参数库建立 DFN 物理模型对比',
        '5. 生成"等效电路 vs 物理模型"对比报告']
for s in flow:
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.8)
    r = pa.add_run(s)
    r.font.name = 'Consolas'; r.font.size = Pt(10); r.font.color.rgb = GREEN

p('')
p('—— 数据下载好之前，我会先用 PyBaMM 内置的 LG 参数做合成数据仿真，把代码流程跑通，等你数据到位直接替换。', 11, True, ORANGE)

path = os.path.join(OUT_DIR, 'SOC估计器_LG_HG2_数据集下载指南.docx')
doc.save(path)
print(f'OK: {path}')
