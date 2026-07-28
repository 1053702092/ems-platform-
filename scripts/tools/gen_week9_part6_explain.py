#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Week9 Part6 Bellman 方程 逐行精讲文档 (.docx)
==================================================
涵盖：
  - 直觉理解：格子打分法
  - 两段代码逐行对照
  - 手算演示
  - 结果解读
"""

import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.4

def add_code_block(doc, code_text, label=None):
    """添加代码块"""
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        shading = run._element.get_or_add_rPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F2F2F2',
            qn('w:val'): 'clear',
        })
        shading.append(shd)

def add_output_block(doc, output_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('▶ 输出')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    for line in output_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.bold = bold
        if bold:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml import OxmlElement
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')
            shading_elm.set(qn('w:val'), 'clear')
            cell._element.get_or_add_tcPr().append(shading_elm)

def add_hint(doc, text, label='💡'):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}  ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

def add_line_ref(doc, line_num, desc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'L{line_num}  ')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x99)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10.5)


# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph('')

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('Week 9 · Part 6\nBellman 方程 逐行精讲')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph('')
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('从直觉到代码 · 从公式到收敛')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph('')
mp = doc.add_paragraph()
mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = mp.add_run(f'生成日期：{datetime.date.today().isoformat()}\n'
               f'对应代码：week9_complete.py  L346-L420\n'
               f'前置知识：Part 5 MDP GridWorld（建议先看懂 P 和 R 的数据结构）')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 0. 到底在干嘛 — 一句话版
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第零章：Bellman 方程到底在干嘛？', level=1)

doc.add_paragraph(
    '一句话回答：给 GridWorld 的 16 个格子打分。'
)

doc.add_paragraph(
    '每个格子 s 有一个价值 V(s)。分数高的格子 = 「从这个格子出发，最终能拿到的总奖励多」。'
)
doc.add_paragraph(
    '比如终点 (3,3) 旁边格子分数很高，因为走两步就能拿 +1 分；'
    '陷阱 (1,1) 旁边格子分数低，因为一不小心就掉坑拿 -1 分。'
)

doc.add_paragraph('')
doc.add_heading('打分的核心思想 — 一个类比', level=2)

doc.add_paragraph('想象你是高中班主任，要给 16 个学生（格子）评「未来前途分」：')
add_bullet(doc, '好学生（靠近终点）= 将来能考上好大学（高分）')
add_bullet(doc, '差学生（靠近陷阱）= 将来可能学坏（低分）')
add_bullet(doc, '终点学生 = 已经功成名就了（+1 分）')
add_bullet(doc, '陷阱学生 = 已经辍学了（-1 分）')

doc.add_paragraph('')
doc.add_paragraph(
    '怎么给分？规则是：一个学生的前途 = 他现在的表现 + 他周围朋友的前途 × 0.9。'
)
doc.add_paragraph(
    '这就产生了一个循环——为了算 A 的前途，得先知道 B 的前途；'
    '但 B 的前途又依赖 A。怎么办？答案是：反复算，每次都基于上一轮的结果，'
    '直到结果不再变化。这就是「迭代」。'
)

doc.add_paragraph('')
add_hint(doc, '核心心法：每个格子的分数 = 当前奖励 + 下一格分数的折扣和', '🎯')
doc.add_paragraph(
    '所以 Bellman 方程其实就这一句话。后面那么长的代码，都是在重复算这句话。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. Part 6 整体结构
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第一章：Part 6 整体在干什么', level=1)

doc.add_paragraph('Part 6 的核心是「用 Bellman 方程求解值函数」，它做了两件事：')

t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
for i, txt in enumerate(['', '段1: 随机策略', '段2: 最优策略']):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True

cells_data = [
    ('输入策略', '瞎走（↑25%, ↓25%, ←25%, →25%）', '智能体自己选最好的方向'),
    ('公式', '对 4 个动作求平均（加权和）', '对 4 个动作取最大值'),
    ('结果', 'V^π(s) — 瞎走情况下的格子分数', 'V*(s) — 最聪明情况下的格子分数'),
    ('用途', '评估一个策略的好坏', '找出最优策略的理论上限'),
]
for row_data in cells_data:
    row = t.add_row()
    for i, txt in enumerate(row_data):
        row.cells[i].text = txt

doc.add_paragraph('')
add_hint(doc,
    '两段代码几乎一样，只有第 11 行不同：一段用 Σ (求和)，一段用 max (取最大)。'
    '理解了这一点就理解了 Part 6 的一半。',
    '✨')

doc.add_paragraph('')

doc.add_paragraph('代码的整体数据流：')
add_bullet(doc, '输入：从 Part 5 传来的 mdp 字典（含 R, P, gamma, n_states, n_actions）')
add_bullet(doc, '处理：用 Bellman 方程迭代更新 V(s) 直到收敛')
add_bullet(doc, '输出：V_pi（随机策略分值）和 V_opt（最优策略分值），以及 heatmap 图')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. 准备工作
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第二章：准备工作 — 从 MDP 取数据（L346-352）', level=1)

add_code_block(doc, '''def part6_bellman(mdp):
    """用 Bellman 方程计算 V(s) 和 Q(s,a)"""
    n_states = mdp['n_states']
    n_actions = mdp['n_actions']
    gamma = mdp['gamma']
    R = mdp['R']
    P = mdp['P']''', 'L346-352')

doc.add_paragraph('')

doc.add_paragraph('从 Part 5 的 mdp 字典中取出五元组：')

data_table = doc.add_table(rows=1, cols=3)
data_table.style = 'Light Grid Accent 1'
for i, txt in enumerate(['变量', '值', '含义']):
    data_table.rows[0].cells[i].text = txt
    data_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
add_table_row(data_table, ['n_states', '16', '4×4 网格 = 16 个格子'])
add_table_row(data_table, ['n_actions', '4', '↑↓←→'])
add_table_row(data_table, ['gamma', '0.9', '折扣因子，未来奖励每步打 9 折'])
add_table_row(data_table, ['R', 'dict', 'R[s][a] = 奖励，终点 +1，陷阱 -1，其他 0'])
add_table_row(data_table, ['P', 'dict', 'P[s][a][s\'] = 从 s 走到 s\' 的概率'])

doc.add_paragraph('')

add_hint(doc,
    '如果对 R 和 P 的数据结构还不清楚，回到 Part 5 看一下：'
    'R = {s: {a: 0.0}} 和 P = {s: {a: {s\': prob}}} 的嵌套字典结构。'
    '这是理解 Part 6 的前提。',
    '⚠️')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. 随机策略评估
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第三章：随机策略评估 — V^π(s)（L354-377）', level=1)

doc.add_paragraph(
    '假设智能体在每个格子都「瞎走」——四个方向各 25% 概率。这种情况下每个格子值多少分？'
)

doc.add_heading('3.1 建立直觉：用手算理解迭代过程', level=2)

doc.add_paragraph('为了看懂代码，我们先手算两个格子。假设初始所有 V = 0，只看第一轮：')

doc.add_paragraph('')
doc.add_paragraph('格子 0 — 起点 (0,0)：')
p = doc.add_paragraph()
r = p.add_run(
    '  四个动作都有效：↑(撞墙=原地), ↓(到(1,0)), ←(撞墙=原地), →(到(0,1))\n'
    '  R(s,a) 全部 = 0（只有终点/陷阱有奖励）\n'
    '  V(0) = 0.25×(0+0.9×V(0)) + 0.25×(0+0.9×V(4)) \n'
    '       + 0.25×(0+0.9×V(0)) + 0.25×(0+0.9×V(1))\n'
    '       = 0.25×0.9×0 + 0.25×0.9×0 + 0.25×0.9×0 + 0.25×0.9×0\n'
    '       = 0'
)
r.font.name = 'Consolas'
r.font.size = Pt(9.5)

doc.add_paragraph('第一轮全是 0，因为周围格子也都是 0。')

doc.add_paragraph('')
doc.add_paragraph('格子 15 — 终点 (3,3)：')
p = doc.add_paragraph()
r = p.add_run(
    '  终点是吸收态：无论执行什么动作，100% 停留在自己\n'
    '  V(15) = 0.25×(1 + 0.9×1×V(15)) × 4\n'
    '         = 0.25×4 × (1 + 0.9×V(15))\n'
    '         = 1 + 0.9×V(15)\n'
    '  解方程：V(15) - 0.9·V(15) = 1  →  0.1·V(15) = 1  →  V(15) = 10\n'
    '  但代码是通过迭代逼近的，不是直接解方程：\n'
    '  第1轮: V=1, 第2轮: V=1+0.9×1=1.9, 第3轮: V=1+0.9×1.9=2.71, ...'
)
r.font.name = 'Consolas'
r.font.size = Pt(9.5)

doc.add_paragraph('')
doc.add_paragraph('终点最终收敛到 V(15) ≈ 10，为什么？')
doc.add_paragraph(
    '因为到终点得了 1 分，但终点本身还值 10 分（因为 γ=0.9，是无穷级数 1+0.9+0.81+...=10）。'
    '智能体一旦进入终点，就永远留在那里，每步都拿 +1——所以终点像个「永续年金」，价值 = 1/(1-0.9) = 10。'
)

doc.add_heading('3.2 逐行解读代码', level=2)

add_code_block(doc, '''# 随机策略（4 个动作概率各 0.25）
policy = np.ones((n_states, n_actions)) / n_actions''', 'L354-355')

doc.add_paragraph(
    'policy 是一个 16×4 的矩阵，每行 = 一个格子的动作概率。'
    'np.ones((16,4)) / 4 → 每个格子都是 [0.25, 0.25, 0.25, 0.25]。'
    '意思就是「瞎走」。'
)

add_code_block(doc, '''V = np.zeros(n_states)
theta = 1e-6
max_iter = 1000''', 'L358-360')

add_line_ref(doc, 358, 'V：长度 16 的数组，存储每个格子的分数。初始全部猜 0。')
doc.add_paragraph(
    '    初始值很重要：猜 0 = 「我什么都不知道，先假设所有格子都没价值」。'
    '    通过迭代，分数会从终点和陷阱向四周传播。'
)
add_line_ref(doc, 359, 'theta = 1e-6：收敛判断阈值。如果两轮之间所有格子的变化 < 0.000001，就认为是够了。')
add_line_ref(doc, 360, 'max_iter = 1000：兜底，死循环保护。1000 轮还没收敛也强行退出。')

add_code_block(doc, '''for i in range(max_iter):
    delta = 0''', 'L361-362')

doc.add_paragraph(
    '外层循环：反复更新 V，每轮都基于上一轮的结果算新的。'
    'delta 用来记录本轮中「变化最大的那个格子的变化量」。'
    '相当于老师每次改分后问：「这次改分最大的变化是多少？如果很小，就停吧。」'
)

add_code_block(doc, '''for s in range(n_states):
    v_old = V[s]
    v_new = 0''', 'L363-365')

doc.add_paragraph(
    '遍历 16 个格子 (s=0 到 15)，给每个格子重新算分。'
    'v_old 记住旧分，v_new 从 0 开始累加新分。'
    '循环结束时，如果 |v_new - v_old| 比之前的 delta 大，就更新 delta。'
)

add_code_block(doc, '''for a in range(n_actions):
    p_a = policy[s, a]
    if p_a == 0:
        continue''', 'L366-369')

doc.add_paragraph(
    '遍历 4 个动作 (a=0↑, 1↓, 2←, 3→)。'
    'p_a 是当前策略下执行这个动作的概率。在随机策略下永远是 0.25，所以 if 永远不会触发。'
    '但保留这个 if 可以应对特殊的确定性策略——某些动作概率为 0 时跳过，省点计算。'
)

add_code_block(doc, '''bellman_sum = R[s][a]
for s_next, prob in P[s][a].items():
    bellman_sum += gamma * prob * V[s_next]''', 'L370-372')

doc.add_paragraph('这是整段代码最核心的三行，理解它们就理解了 Bellman 方程：')

doc.add_paragraph('')
doc.add_paragraph('第一行：bellman_sum = R[s][a]')
doc.add_paragraph(
    '  从格子 s 执行动作 a，能立刻拿到多少奖励？'
    '  大部分格子是 0，终点是 +1，陷阱是 -1。'
)
doc.add_paragraph('')
doc.add_paragraph('第二-三行：遍历 P[s][a] 字典')
doc.add_paragraph(
    '  P[s][a] 是一个字典 {目标格子: 概率}。比如 P[5][0] 可能 = {5: 0.8, 1: 0.0667, ...}。\n'
    '  对每个可能的下一格 s\'，加上：γ × 概率 × V[s\']\n'
    '  翻译成人话：「未来能得多少分 = 去每个方向的可能性 × 那边的分数，打折之后加起来」'
)

doc.add_paragraph('')
doc.add_paragraph('用一个实际例子跑一遍——格子 1（紧挨终点左上方），选择「→」这个动作：')

p = doc.add_paragraph()
r = p.add_run(
    '  R[1][3] = 0  （→ 对应索引 3；从格子1去格子2，没有立即奖励）\n'
    '  P[1][3] = {2: 0.8, 1: 0.0667, 5: 0.0667, 0: 0.0667}\n'
    '    → 80% 去格子2（右边），6.67% 留在原地，6.67% 掉到下面，6.67% 回到左边\n'
    '  bellman_sum = 0 + 0.9 × (0.8×V[2] + 0.0667×V[1] + 0.0667×V[5] + 0.0667×V[0])\n'
    '  = 0.9 × (0.8×V[2] + 0.0667×(V[1]+V[5]+V[0]))'
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

doc.add_paragraph('')
doc.add_paragraph('注意这里 V[2] 有 0.8 的权重（因为这是意图方向），其他三个共享 0.2 的权重。')

add_code_block(doc, '''v_new += p_a * bellman_sum''', 'L373')

doc.add_paragraph(
    '将当前动作的「折扣后总价值」按概率加权加到 v_new。'
    '随机策略下 p_a=0.25，所以：'
)
doc.add_paragraph(
    '  v_new = 0.25 × bellman_sum(↑) + 0.25 × bellman_sum(↓)\n'
    '        + 0.25 × bellman_sum(←) + 0.25 × bellman_sum(→)'
)

doc.add_paragraph('')
doc.add_paragraph('这对应于公式：')
p = doc.add_paragraph()
r = p.add_run('  V^π(s) = Σ_a π(a|s) · [ R(s,a) + γ Σ_{s\'} P(s\'|s,a) V^π(s\') ]')
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_code_block(doc, '''V[s] = v_new
delta = max(delta, abs(v_old - v_new))''', 'L374-375')

doc.add_paragraph(
    '把算好的新分写进 V[s]，并更新 delta——记录这轮中变化最大的格子的变化幅度。'
    '到后期，所有格子的变化都会越来越小，直到被 theta 截停。'
)

add_code_block(doc, '''if delta < theta:
    break''', 'L376-377')

doc.add_paragraph(
    '如果 16 个格子中变化最大的那个，变化量也不到百万分之一（1e-6），'
    '就认为 V 已经收敛了——即使再算 1000 轮结果也基本一样。'
)

doc.add_paragraph('')
add_hint(doc,
    '收敛后 V[s] 就是随机策略下每个格子的「稳定分数」。'
    '这个分数是自洽的——如果你拿着这份分数去算下一轮，结果几乎一模一样。',
    '✅')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. 最优值函数
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第四章：最优值函数 — V*(s)（L379-394）', level=1)

doc.add_paragraph(
    '和第三章的步骤几乎一样，唯一的区别在第 11 行——不是加权平均，而是取最大值。'
)

add_code_block(doc, '''# 最优值函数 V*
V_opt = np.zeros(n_states)
for i in range(max_iter):
    delta = 0
    for s in range(n_states):
        v_old = V_opt[s]
        q_values = []
        for a in range(n_actions):
            q = R[s][a]
            for s_next, prob in P[s][a].items():
                q += gamma * prob * V_opt[s_next]
            q_values.append(q)
        V_opt[s] = max(q_values)     # ← 唯一区别：取最大
        delta = max(delta, abs(v_old - V_opt[s]))
    if delta < theta:
        break''', '')

doc.add_paragraph('')

doc.add_heading('4.1 两段代码的逐行对比', level=2)

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Light Grid Accent 1'
for i, txt in enumerate(['', '随机策略 (V^π)', '最优策略 (V*)']):
    t2.rows[0].cells[i].text = txt
    t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

compare_data = [
    ('初始化', 'V = zeros(16)', 'V_opt = zeros(16)'),
    ('动作循环', 'for a in range(4):\n    p_a = policy[s,a]\n    …\n    v_new += p_a × q', 'for a in range(4):\n    q = R + γ·ΣP·V\n    q_values.append(q)'),
    ('合并方式', 'v_new += p_a × q  (加权求和)', 'V_opt[s] = max(q_values)  (取最大)'),
    ('含义', '对所有动作的结果求平均', '只取最好的那个动作的结果'),
]
for row_data in compare_data:
    row = t2.add_row()
    for i, txt in enumerate(row_data):
        row.cells[i].text = txt

doc.add_paragraph('')

add_hint(doc,
    '取 max 意味着智能体可以「选择」最好的动作——'
    '所以 V*(s) ≥ V^π(s) 对所有 s 成立。起点 V*(0) ≈ 0.75 >> V^π(0) ≈ -0.05。',
    '🎯')

doc.add_paragraph('')
doc.add_paragraph(
    '具体到格子 0（起点）：最优策略下，它会选择四个动作中 Q 值最大的那个（可能是 → 或 ↓）。'
    '因此不会像随机策略那样「25% 概率掉陷阱」，得分自然高得多。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. 输出与可视化
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第五章：输出与可视化（L396-420）', level=1)

add_code_block(doc, '''print(f'[Part 6] Bellman 方程')
print(f'  随机策略 V(s) 收敛于 {i+1} 次迭代')
print(f'  起点 V(0)     = {V[0]:.4f}')
print(f'  最优 V*(0)    = {V_opt[0]:.4f}')''', 'L396-399')

doc.add_paragraph('')
doc.add_paragraph('典型输出：')

add_output_block(doc, '''[Part 6] Bellman 方程
  随机策略 V(s) 收敛于 148 次迭代
  起点 V(0)     = -0.0457
  最优 V*(0)    = 0.7524''')

doc.add_paragraph('')
doc.add_paragraph('解读：')
add_bullet(doc, '随机策略下起点是负分（-0.0457）：瞎走的话大概率掉陷阱，所以起点 = 危险区域')
add_bullet(doc, '最优策略下起点是正分（0.7524）：虽然离终点远，但绕开陷阱走，累计折扣奖励可观')
add_bullet(doc, '差距 ≈ 0.8 分：这就是「一个好策略」带来的价值增益')

doc.add_heading('5.2 Heatmap 可视化', level=2)

doc.add_paragraph(
    '两张热力图并排对比 V^π(s) 和 V*(s)。颜色越暖（红）分数越高，越冷（蓝）分数越低。'
)

doc.add_paragraph('')
doc.add_paragraph('左图：V^π(s) — 随机策略')
add_bullet(doc, '陷阱(1,1) 处深蓝 = -1.00')
add_bullet(doc, '陷阱周围几个格子也是蓝色或灰色，因为 20% 概率滑进陷阱')
add_bullet(doc, '终点(3,3) 处深红 ≈ 10.00')
add_bullet(doc, '起点(0,0) 附近灰色偏蓝 ≈ -0.05')

doc.add_paragraph('')
doc.add_paragraph('右图：V*(s) — 最优策略')
add_bullet(doc, '同样陷阱处 -1.00，终点处 ≈ 10.00')
add_bullet(doc, '但中间区域整体偏暖（更高分），因为智能体知道避开陷阱')
add_bullet(doc, '起点(0,0) ≈ 0.75，明显暖色调')

doc.add_paragraph('')
add_hint(doc,
    '两幅图最直观的差异就在起点附近——随机策略下起点周围偏冷，'
    '最优策略下起点周围偏暖。这就是「会做选择」的优势。',
    '📊')

add_code_block(doc, '''return {'V_pi': V, 'V_opt': V_opt}''', 'L420')

doc.add_paragraph(
    '返回两个值函数字典。V_pi 供你分析「某个策略有多好」，V_opt 供你作为后续 Part 7/8 '
    '策略优化的验证基准。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. 常见困惑解答
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第六章：常见困惑解答', level=1)

faqs = [
    ('Q1: 为什么 V 要迭代？不能一步算出来吗？',
     '理论上可以解线性方程组 V = R + γPV，但 16 个格子就需要解 16 元方程组。'
     '如果 GridWorld 扩大到 100×100=10000 个格子，解方程组就太慢了。'
     '迭代法每轮 O(n)，简单且可扩展。'
     '而且迭代过程能直观看到分数从终点/陷阱向四周「扩散」。'),

    ('Q2: 终点的 V 为什么是 10？不是 +1 吗？',
     '终点是吸收态——进去了就出不来，每步都拿 +1。'
     '价值是未来所有奖励的折扣和：1 + 0.9 + 0.81 + 0.729 + ... = 1/(1-0.9) = 10。'
     '如果你把 γ 改成 0.5，终点 V 就变成 1/(1-0.5)=2。'
     'γ 越大，终点分值越高。'),

    ('Q3: 随机策略那段在 Part 7 还有用吗？',
     '有！Part 7 的「策略评估」步骤做的就是这件事——'
     '给一个确定性策略（不是随机 0.25）打分。'
     'Part 6 的随机策略评估是 Part 7 策略评估的一个特例（策略全部 0.25）。'),

    ('Q4: 两段代码哪个算得快？',
     '随机策略那段更慢，因为它需要 × 4 个动作再求平均。'
     '最优策略那段直接取 max，少一层乘法。'
     '但真正的时间都花在 P[s][a] 的字典遍历上。'),

    ('Q5: theta = 1e-6 会不会太小？',
     '对于 16 个格子的简单问题，1e-6 是合适的。'
     '更大的问题可以用 1e-4 或 1e-3（早停），因为误差不会灾难性地传播。'
     '你可以试试把 theta 改成 0.01 或 0.1，看看 V(0) 差多少。'),

    ('Q6: 为什么 Part 6 和 Part 8 代码这么像？',
     '因为 Part 6 的最优值函数代码就是 Part 8 的值迭代。'
     'Part 6 先展示「最优 Bellman 方程算出来长什么样」，'
     'Part 8 正式命名并深入讨论。'
     'Part 6 独特之处在于它算了两组值（随机 vs 最优）做对比。'),

    ('Q7: 代码里 L372 的 V[s_next] 用的是更新前的值还是更新后的？',
     '用的是当前 V 数组中的值。同一个循环内，格子 0 先更新，格子 1 后更新——'
     '格子 1 更新时看到的是格子 0 的「新值」。'
     '这叫 Gauss-Seidel 风格更新（而非 Jacobi 风格），收敛更快。'),

    ('Q8: 到底什么是「自洽」？',
     '如果你拿收敛后的 V 代入 Bellman 方程右边的 Σ P·V(s\')，'
     '算出来的新 V 和旧 V 几乎一样。'
     '这意味着「这些分数不再自我矛盾」——就像解方程 x = 0.9x + 1 一样，找到了平衡点。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    doc.add_paragraph(a)
    doc.add_paragraph('')

# ── 保存 ──
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'docs', 'notes')
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, 'Week9_Part6_Bellman_逐行精讲.docx')
doc.save(OUTPUT_PATH)
print(f'[OK] {OUTPUT_PATH}')
