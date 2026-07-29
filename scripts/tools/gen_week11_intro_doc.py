#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate a Week 11 overview and study handoff DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = PROJECT_ROOT / "docs" / "notes"
RESULTS_DIR = PROJECT_ROOT / "results"
OUT_PATH = OUT_DIR / "Week11_完整介绍与学习路线.docx"

BLUE = RGBColor(0x1F, 0x3A, 0x5F)
HEADING = RGBColor(0x2E, 0x74, 0xB5)
DARK = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xC0, 0x39, 0x2B)
FILL = "E8EEF5"


def set_run_font(run, name="微软雅黑", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, HEADING, 18, 10),
        ("Heading 2", 13, HEADING, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        st = styles[style_name]
        st.font.name = "微软雅黑"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25


def paragraph(doc, text="", *, bold=False, color=None, size=11, italic=False, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), color=HEADING if level <= 2 else DARK, bold=True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths_dxa)
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(str(text))
            set_run_font(r, size=10)
    set_table_geometry(table, widths_dxa)
    paragraph(doc)
    return table


def callout(doc, title, body, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    paragraph(doc)


def add_image(doc, path, caption, width_cm=15.3):
    path = Path(path)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        cap = paragraph(doc, caption, size=9.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cap.paragraph_format.space_after = Pt(10)
    else:
        paragraph(doc, f"[图片未找到: {path}]", color=RED, bold=True)


def add_cover(doc):
    for _ in range(3):
        paragraph(doc)
    p = paragraph(doc, "Week 11 完整介绍与学习路线", bold=True, color=BLUE, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(12)
    paragraph(doc, "连续动作强化学习：REINFORCE、Actor-Critic、PPO 与 Q-learning/DQN 对比", color=MUTED, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "适用：自学复盘 / ChatGPT 接力 / 面试项目包装", color=MUTED, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(8):
        paragraph(doc)
    callout(
        doc,
        "一句话结论",
        "Week11 主线已经完整：从离散 DQN 的边界讲到连续动作 RL，再用 REINFORCE -> Actor-Critic -> PPO 形成递进实现，并补齐了可复现脚本、结果图、报告和接力说明。",
    )
    doc.add_page_break()


def main():
    doc = Document()
    style_doc(doc)
    add_cover(doc)

    heading(doc, "一、Week11 在 EMS-PLATFORM 中的位置", 1)
    paragraph(doc, "Week11 的作用不是把强化学习做到论文级调参，而是把 EMS 项目从“传统优化策略”自然过渡到“学习型控制策略”。前面已经有 DP、ECMS、MPC 与 SOC 估计底座，Week11 要补的是：为什么离散动作方法不够、连续动作策略怎么建模、PPO 为什么适合拿来做轻量落地。")
    callout(
        doc,
        "完成度判断",
        "学习计划意义上：完成。项目包装意义上：可以作为“强化学习 EMS 原型验证”材料。下一步不建议继续深挖更多 RL 算法，而是把 PPO-EMS 奖励函数、SOC 约束和简历叙事打磨清楚。",
    )

    heading(doc, "二、文件地图：先看哪些", 1)
    add_table(
        doc,
        ["优先级", "文件", "用途"],
        [
            ["1", "docs/notes/Week11_连续动作RL对比报告.docx", "主报告：先看整体逻辑、三算法对比和核心记忆点"],
            ["2", "docs/notes/QL_vs_DQN_大Grid对比报告.docx", "辅助报告：理解为什么 DQN 不适合直接处理连续动作"],
            ["3", "scripts/week11_continuous_env.py", "连续动作 EMS 环境，理解状态、动作、奖励"],
            ["4", "scripts/week11_compare.py", "三算法统一对比入口，建议优先读这个脚本"],
            ["5", "scripts/week11_reinforce.py / week11_actor_critic.py / week11_ppo.py", "逐个看算法实现差异"],
            ["6", "scripts/compare_large_grid.py", "Q-learning vs DQN 大 Grid 对比，属于概念辅助材料"],
        ],
        [1100, 3900, 4360],
    )

    heading(doc, "三、学习主线：从离散动作到 PPO", 1)
    add_table(
        doc,
        ["阶段", "你要理解什么", "一句话记忆"],
        [
            ["Q-learning", "用 Q 表记录每个状态-动作价值", "小状态空间好用，但没有泛化能力"],
            ["DQN", "用神经网络近似 Q(s,a)", "能泛化，但仍然偏离散动作"],
            ["REINFORCE", "直接优化策略概率", "能处理连续动作，但方差大"],
            ["Actor-Critic", "策略网络 + 价值网络", "用价值估计降低方差，训练更稳"],
            ["PPO", "限制策略更新幅度", "用 clip 防止一步改太猛，适合工程原型"],
        ],
        [1500, 5000, 2860],
    )

    heading(doc, "四、实验结果怎么读", 1)
    heading(doc, "4.1 连续动作 RL 三算法对比", 2)
    add_table(
        doc,
        ["方法", "最终表现", "训练时间", "解释"],
        [
            ["REINFORCE", "最后50局平均奖励约 -38.26", "115.6s", "能跑通，但曲线波动较大"],
            ["Actor-Critic", "最后50局平均奖励约 -31.31", "273.1s", "本次综合奖励最好，但耗时最长"],
            ["PPO", "最后50局平均奖励约 -38.20", "75.4s", "速度快、实现稳定，适合作为后续 EMS-RL 基线"],
        ],
        [1700, 2300, 1800, 3560],
    )
    add_image(doc, RESULTS_DIR / "week11_comparison.png", "图 1：Week11 连续动作 RL 三算法统一对比。")

    heading(doc, "4.2 Q-learning vs DQN 大 Grid 对比", 2)
    add_table(
        doc,
        ["Grid", "Q-learning", "DQN", "要点"],
        [
            ["4x4", "奖励 +0.81，时间 0.18s", "奖励 +0.78，时间 218.50s", "小问题上 Q 表更划算"],
            ["8x8", "奖励 +0.81，时间 0.45s", "奖励 +0.50，时间 851.14s", "DQN 参数多、样本效率低，短训练下容易策略退化"],
        ],
        [1200, 2600, 2600, 2960],
    )
    doc.add_page_break()
    add_image(doc, RESULTS_DIR / "compare_8x8_ql_vs_dqn.png", "图 2：8x8 GridWorld 中 DQN 与 Q-learning 的收敛差异。")

    heading(doc, "五、你应该怎么学这部分", 1)
    add_table(
        doc,
        ["顺序", "任务", "完成标准"],
        [
            ["第1步", "先读本介绍文档和 Week11 主报告", "能说清楚 Week11 为什么从 DQN 过渡到 PPO"],
            ["第2步", "运行 week11_compare.py", "知道三种算法输出了哪些曲线和指标"],
            ["第3步", "读 week11_continuous_env.py", "能解释状态、动作、奖励函数分别代表什么"],
            ["第4步", "对照读 REINFORCE / Actor-Critic / PPO", "能说出三者更新方式的差异"],
            ["第5步", "整理成面试话术", "能用 1 分钟讲清 PPO 为什么适合 EMS 连续控制"],
        ],
        [1100, 4300, 3960],
    )

    doc.add_page_break()
    heading(doc, "六、面试怎么讲", 1)
    paragraph(doc, "推荐用下面这段作为 Week11 的项目叙事基础：")
    callout(
        doc,
        "面试口径",
        "我在 EMS-PLATFORM 中先实现了 DP、ECMS、MPC 等传统能量管理方法，之后用 Week11 做强化学习扩展。由于 EMS 中燃料电池功率、电池功率分配更接近连续控制，DQN 这类离散动作方法需要强行离散化，动作精度和搜索效率都会受限。因此我构建了一个连续动作 EMS 环境，并递进实现 REINFORCE、Actor-Critic 和 PPO。最后选 PPO 作为后续基线，是因为它通过 clipped surrogate objective 限制策略更新幅度，训练稳定性更适合工程原型。",
    )

    heading(doc, "七、给 ChatGPT 的接力提示", 1)
    paragraph(doc, "如果要让 ChatGPT 做简单接力，可以直接把下面这段发过去：")
    callout(
        doc,
        "接力摘要",
        "当前项目位于 D:\\CHAT\\_projects\\EMS-paltform，分支 feat/soc-estimation-research。Week11 已完成连续动作 RL 学习闭环：包含 REINFORCE、Actor-Critic、PPO 三个脚本、统一对比脚本 week11_compare.py、连续动作环境 week11_continuous_env.py，以及 Q-learning vs DQN 大 Grid 辅助实验。请优先阅读 STATUS.md、docs/notes/Week11_连续动作RL对比报告.docx 和 docs/notes/QL_vs_DQN_大Grid对比报告.docx；简单任务可以帮我整理学习笔记、提炼面试话术、检查表述是否清楚，复杂代码修改交给 Codex。",
    )

    doc.add_page_break()
    heading(doc, "八、下一步计划", 1)
    bullet(doc, "进入 Week12：简历打磨 + 笔面试准备，不再把时间花在扩展更多 RL 算法。")
    bullet(doc, "把 PPO-EMS 讲成“传统 EMS + RL 原型扩展”，避免夸大成成熟工业控制器。")
    bullet(doc, "SOC 项目继续往 1RC/HPPC 参数标定走，这是 BMS 岗位更硬的项目亮点。")
    bullet(doc, "如果需要继续包装，可生成一份“简历项目版 Week11 话术”和一页项目卡片。")

    section = doc.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("EMS-PLATFORM Week11 学习说明 | 2026-07-28")
    set_run_font(r, size=9, color=MUTED)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"OK: {OUT_PATH}")


if __name__ == "__main__":
    main()
