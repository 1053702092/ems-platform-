#!/usr/bin/env python3
"""生成《SOC估计器独立项目 — 实施任务书》DOCX"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
ST = doc.styles['Normal']
ST.font.name = 'Microsoft YaHei'
ST.font.size = Pt(10.5)
ST.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Microsoft YaHei'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h.font.color.rgb = RGBColor(0x1a, 0x47, 0x8a)

def H(s, lv=1): return doc.add_heading(s, level=lv)
def P(s, bold=False):
    p = doc.add_paragraph(s)
    if bold:
        for r in p.runs:
            r.bold = True
    return p
def B(s):
    p = doc.add_paragraph()
    p.add_run(s).bold = True
    return p
def L(s): doc.add_paragraph(s, style='List Bullet')
def add_table(hd, rows):
    t = doc.add_table(rows=1, cols=len(hd))
    t.style = 'Light Grid Accent 1'
    for i, hd1 in enumerate(hd):
        t.rows[0].cells[i].text = hd1
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
    for rd in rows:
        r = t.add_row()
        for i, v in enumerate(rd):
            r.cells[i].text = str(v)
            for p in r.cells[i].paragraphs:
                for rr in p.runs:
                    rr.font.size = Pt(9)
    return t

doc.add_heading('电池 SOC 估计器\n独立项目 — 实施任务书', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('将散落的 SOC 估计代码包装成一个可演示的独立项目')
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
P('编制日期：2026-07-27 | 预期耗时：1-2天')
doc.add_page_break()

H('项目目标', 1)
P('把现有的 EKF/AEKF SOC 估计代码从一个"MPC里的子模块"包装成一个"独立可运行的项目"。')
B('最终效果：一行命令跑完，出对比图，面试时能现场演示。')

H('目录结构（已在 scripts/soc_estimator/ 下建好骨架）', 1)
C = """scripts/soc_estimator/
├── run.py              ← 主入口（骨架已搭好，需填空）
├── estimator.py        ← 三个估计器类（骨架已搭好，需填算法）
├── battery_model.py    ← OCV-SOC表+数据加载（骨架已搭好）
├── requirements.txt    ← numpy, matplotlib
└── README.md           ← 需要你写"""
p = doc.add_paragraph()
r = p.add_run(C)
r.font.name = 'Consolas'; r.font.size = Pt(9)

doc.add_page_break()
H('需要你完成的步骤', 1)

H('第1步：填 estimator.py 中的算法逻辑（关键，2-3小时）', 2)
P('打开 estimator.py，文件里有 TODO 标注。具体要填的：')
add_table(['类', '方法', '要填什么'],
    [
        ['OpenLoopEstimator', 'step()', 'SOC_{k+1} = SOC_k - I/Q*dt，限幅到 [0.2, 0.9]'],
        ['EKFEstimator', 'step()\nPredict阶段', 'soc_pred = x - I/Q*dt\nP_pred = P + Q'],
        ['', 'Update阶段', 'v_pred = lookup_ocv(soc_pred)\ny = v_t_meas - v_pred\nH = lookup_docv_dsoc(soc_pred)\nK = P_pred*H / (H*P_pred*H + R)\nx_est = soc_pred + K*y\nP_est = (1-K*H)*P_pred'],
        ['AEKFEstimator', 'step()', '同EKF + 用innov_buffer自适应调整Q/R'],
    ])
P('')
P('提示：可以从 mpc_ems_ekf.py 中复制对应的实现逻辑，但建议自己手打一遍理解。', bold=True)

H('第2步：填 battery_model.py（1小时）', 2)
P('两个函数需要填：')
L('lookup_ocv(soc) → np.interp 查表')
L('lookup_docv_dsoc(soc) → np.gradient 求导')
L('load_cycle_data() → 从已有数据加载或生成模拟数据')

H('第3步：填 run.py 仿真循环（2小时）', 2)
P('run.py 的 main() 里需要：')
L('循环时间步，逐时间步调用三种 estimator.step()')
L('记录真实SOC、开环估计、EKF估计、AEKF估计')
L('计算 SOC_RMSE 和终点误差')
L('画出三条对比曲线')
P('')
B('建议从 mpc_ems_ekf.py 的 main() 函数中提取仿真循环逻辑，简化成独立版本。')

H('第4步：加故障场景（1-2小时）', 2)
P('数据准备：在加载数据后，对电流加偏置或噪声：')
C2 = """# run.py 中支持:
python run.py                      # 标准模式
python run.py --fault bias --value 2   # 电流偏置 2A
python run.py --fault noise --value 1  # 电流噪声 σ=1.0"""
p = doc.add_paragraph()
r = p.add_run(C2)
r.font.name = 'Consolas'; r.font.size = Pt(9)
P('')
P('在仿真循环开始时，对电流 I 施加偏置或噪声：')
L('偏置: i_meas = i_true + fault_value')
L('噪声: i_meas = i_true + normal(0, fault_value)')

H('第5步：写 README.md 并跑通验证（1小时）', 2)
P('README.md 应包含：')
L('项目一句话介绍')
L('使用方法（3条命令）')
L('结果图（放一张对比图）')
L('量化指标（SOC_RMSE 0.0024等）')
L('参考 docs/soc-estimation/md/ 里的文档')

doc.add_page_break()
H('完成前后对比', 1)
add_table(['维度', '现在', '完成后'],
    [
        ['项目形态', 'MPC里的一个子模块\n或散落的独立脚本', '独立项目，有统一入口\n一行命令出结果'],
        ['可演示性', '要打开代码解释', '终端直接跑+出图'],
        ['简历描述', '"在MPC中集成了EKF"', '"电池SOC估计器\nEKF/AEKF/故障测试"'],
        ['面试追问', '可能要翻代码', 'README 写清楚了\n自己打过一遍理解了'],
    ])

H('预期的简历描述', 1)
B('电池SOC估计器 / 锂电池状态估计算法')
L('实现 EKF/AEKF 锂电池 SOC 在线估计算法，在传感器偏置 2A 故障下 SOC_RMSE 仅 0.3%（比开环安时积分改进 6×）')
L('支持标准模式与故障模式（偏置/噪声）测试，自动生成对比报告')

H('时间预估', 1)
add_table(['步骤', '时间', '说明'],
    [
        ['第1步 填estimator.py', '2-3小时', '核心算法，对照现有代码填'],
        ['第2步 填battery_model.py', '1小时', '查表+导数+数据加载'],
        ['第3步 填run.py', '2小时', '仿真循环+出图'],
        ['第4步 加故障场景', '1-2小时', '偏置+噪声参数支持'],
        ['第5步 README+验证', '1小时', '写文档+跑通全流程'],
        ['合计', '7-9小时', '约1-2天'],
    ])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('\n—— 任务书完 ——').bold = True
P('tips: 所有 TODO 都有提示，填完后从 mpc_ems_ekf.py 对比验证结果是否一致')

out_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    'docs',
    'SOC估计器独立项目_实施任务书.docx'
)
doc.save(out_path)
print('Done:', out_path)
