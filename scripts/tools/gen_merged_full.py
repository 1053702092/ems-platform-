#!/usr/bin/env python3
"""week9_complete.py 完整逐行精讲 — 全篇合并（保留全部细节）"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
font = style.font; font.name = 'Microsoft YaHei'; font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# ===== helpers =====
def _code(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.8)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name='Consolas'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x2D,0x2D,0x2D)
    s=OxmlElement('w:shd'); s.set(qn('w:fill'),'F0F0F0'); s.set(qn('w:val'),'clear')
    p._element.get_or_add_pPr().append(s); return p

def _h1(t):
    h=doc.add_heading(t,1)
    for r in h.runs: r.font.size=Pt(20); r.font.color.rgb=RGBColor(0x0D,0x2B,0x5E)
def _h2(t):
    h=doc.add_heading(t,2)
    for r in h.runs: r.font.size=Pt(17); r.font.color.rgb=RGBColor(0x1A,0x3C,0x6E)
def _h3(t):
    h=doc.add_heading(t,3)
    for r in h.runs: r.font.size=Pt(15); r.font.color.rgb=RGBColor(0x2D,0x5A,0x8C)

def _b(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(14); return p
def _n(t):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(1.2)
    r=p.add_run('> '+t); r.font.size=Pt(13); r.font.color.rgb=RGBColor(0x00,0x70,0xC0); r.italic=True
def _w(t):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(1.2)
    r=p.add_run('! '+t); r.font.size=Pt(13); r.font.color.rgb=RGBColor(0xCC,0x33,0x00); r.bold=True
def _bx(t,c='b'):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(t); r.font.size=Pt(14); f='E8F0FE' if c=='b' else 'FFF0F0'
    r.font.color.rgb=RGBColor(0x00,0x55,0x99) if c=='b' else RGBColor(0xCC,0x33,0x00)
    if c=='r': r.bold=True
    s=OxmlElement('w:shd'); s.set(qn('w:fill'),f); s.set(qn('w:val'),'clear')
    p._element.get_or_add_pPr().append(s)
def _bl(b,n=''):
    p=doc.add_paragraph(style='List Bullet'); r1=p.add_run(b); r1.font.size=Pt(14); r1.bold=True
    if n: r2=p.add_run(n); r2.font.size=Pt(14)
def _line(): _b('─'*60)

LINE = lambda n: f'（第 {n} 行）'

# ================================================================
_h1('week9_complete.py 完整逐行精讲')
_b('本文档对 week9_complete.py（698 行）做逐行分析。约定：灰底 = 代码，蓝字 > = 注释，蓝框 = 概念，红框 = 注意。')

# ===================== 1-34 =====================
_h1('一、代码头部（1-34 行）')

_h2('1-2 行：Shebang 和编码')
_code('#!/usr/bin/env python3')
_code('# -*- coding: utf-8 -*-')
_b('第 1 行：类 Unix 系统用 env 查找 Python3 解释器。第 2 行：声明 UTF-8。Windows 下忽略，但保留是好习惯。')

_h2('3-20 行：文件文档字符串')
_code('"""Week 9 — PyTorch + RL 基础 完整通关"""')
_b('多行 docstring，描述文件总目标、8 个 Part 说明、输出产物。IDE 会显示在文件预览中。')

_h2('22-26 行：库导入')
_code('import os, sys, itertools')
_code('import numpy as np')
_code('import matplotlib')
_code("matplotlib.use('Agg')")
_code('import matplotlib.pyplot as plt')
_bl('os：','文件和路径操作')
_bl('itertools：','Part 5 的 product() 生成网格坐标')
_bl('numpy：','数组运算')
_bl("matplotlib.use('Agg')：","无 GUI 后端，服务器脚本必须。一定在 import pyplot 之前调用。")
_w("matplotlib.use() 必须在 import pyplot 之前，否则无效！")

_h2('28-29 行：结果目录')
_code('RESULTS_DIR = os.path.join(..., "results")')
_code('os.makedirs(RESULTS_DIR, exist_ok=True)')
_b('取项目根目录 → 拼接 results/ → 创建目录。exist_ok=True 避免重复创建报错。')

_h2('32-34 行：中文字体')
_code("plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']")
_code("plt.rcParams['axes.unicode_minus'] = False")
_b('配置 matplotlib 使用微软雅黑显示中文。第三行防止负号显示为方块。')

# ===================== PART 1 =====================
_h1('二、Part 1: Tensor 基础（40-77 行）')

_h2('40 行：函数定义')
_code('def part1_tensor_basics():')
_b('返回 device 对象供后续使用。')

_h2('43 行：导入 torch')
_code('import torch')
_b('函数内导入（非全局），各 Part 可独立运行。')

_h2('45-46 行：列表 → Tensor')
_code('t1 = torch.tensor([[1, 2], [3, 4]])')
_b('从嵌套列表创建 2×2 整数张量。自动推断 dtype=torch.int64。')
_n('输出: tensor([[1,2],[3,4]]), dtype=int64, shape=[2,2]')

_h2('48-51 行：NumPy → Tensor')
_code('a = np.array([1.0, 2.0, 3.0])')
_code('t2 = torch.from_numpy(a)')
_b('from_numpy 和 NumPy 共享内存！修改 a 会影响 t2，反之亦然。')

_h2('53-57 行：特殊张量')
_code('zeros = torch.zeros(2, 3)')
_code('ones = torch.ones(2, 3)')
_code('rand = torch.randn(3, 3)')
_b('全零、全一、标准正态分布 N(0,1)。randn 的 n 代表 normal。')

_h2('60 行：设备检测')
_code('device = torch.device("cuda" if torch.cuda.is_available() else "cpu")')
_b('三目运算符。有 GPU 用 GPU，否则 CPU。本环境输出 Device: cpu。')

_h2('64 行：指定设备')
_code('t_gpu = torch.tensor([1, 2, 3], device=device)')
_b('直接在目标设备上创建。GPU 上创建的在显存中。')

_h2('67-70 行：索引、切片、形状')
_code("print(f't[0]: {t[0].shape}')")
_code("print(f't[:, :3]: {t[:, :3].shape}')")
_code("print(f't.view(-1): {t.view(-1).shape}')")
_code("print(f't.reshape(2, 10): {t.reshape(2, 10).shape}')")
_b('t[0] = 第一行（形状[5]）；t[:,:3] = 前三列（形状[4,3]）。view(-1) 展平为 [20]；reshape(2,10) 重塑。')
_n('view 要求连续内存，reshape 不要求（不连续时自动拷贝）。')

_h2('73-75 行：广播')
_code("print(f'broadcast add:\\n{a + b}')")
_b('[3,1] + [3] → 广播为 [3,3] + [3,3] → [[11,21,31],[12,22,32],[13,23,33]]')
_b('广播自动扩展维度使形状匹配，省去手动复制。')

_h2('77 行：返回值')
_code('return device')

# ===================== PART 2 =====================
_h1('三、Part 2: Autograd 自动求导（88-113 行）')

_h2('88 行：标记求导')
_code('x = torch.tensor([2.0, 3.0], requires_grad=True)')
_b('requires_grad=True 让 PyTorch 追踪 x 的所有操作。只有浮点类型可标记，整数不行。')

_h2('89-95 行：前向 + 反向 + 梯度')
_code('y = x ** 2 + 3 * x')
_code('loss = y.sum()')
_code('loss.backward()')
_code("print(f'gradient dy/dx: {x.grad}')")
_b('y = x² + 3x。对 y 求和得到标量 loss（backward 必须从标量出发）。')
_b('backward() 自动计算梯度，存到 x.grad。')
_n('x.grad = [7., 9.] 验证: dy/dx=2x+3, 代入 [2,3] → [7,9] ✓')
_code("print(f'loss: {loss.item()}')")
_b('loss.item() 从单元素张量提取 Python float。')

_h2('98-99 行：清零')
_code('x.grad.zero_()')
_w('PyTorch 默认累积梯度！不清零会导致每次 backward 叠加，参数更新异常。')
_b('.zero_() 后缀下划线表示原地操作（in-place），直接修改 x.grad 本身。')

_h2('102-107 行：链式法则')
_code('a = torch.tensor(2.0, requires_grad=True)')
_code('b = torch.tensor(3.0, requires_grad=True)')
_code('z = (a ** 2) * torch.sin(b)')
_code('z.backward()')
_bl('dz/da = 2a·sin(b) = 4×0.1411 = 0.5645','验证 ✓')
_bl('dz/db = a²·cos(b) = 4×(-0.99) = -3.96','验证 ✓')
_n('你只写前向计算，PyTorch 自动算所有偏导——这就是"自动求导"的意义。')

_h2('110-112 行：no_grad')
_code('with torch.no_grad():')
_code('    y_eval = x ** 2 + 3 * x')
_b('推理时关闭梯度追踪，不构建计算图，节省内存和计算。')

# ===================== PART 3 =====================
_h1('四、Part 3: nn.Module + MLP（117-146 行）')

_h2('117-123 行：函数定义 + 导入')
_code("""# ══════════════════════════════════════════════════════
# Part 3: nn.Module + MLP
# ══════════════════════════════════════════════════════
def part3_mlp_module():
    \"\"\"用 nn.Module 构建一个多层感知机.\"\"\"
    import torch
    import torch.nn as nn
    import torch.nn.functional as F""")
_b('117-118 行：大片 # 分隔线，把文件分成清晰段落。多人协作时快速定位。')
_b('119 行：函数定义，返回训练好的 model。')
_b('120 行：函数级 docstring。')
_b('121-123 行：PyTorch 三个核心层级：')
_bl('torch：','基础库，Tensor 创建设备管理')
_bl('torch.nn as nn：','网络层（Linear、Conv2d），有可训练参数')
_bl('torch.nn.functional as F：','函数（ReLU、softmax），无参数')
_n('nn 有参数（需学习），F 无参数（只是计算）。两者平行导入是 PyTorch 标准写法。')

_h2('124-134 行：定义 MLP 类')
_code("""    class MLP(nn.Module):
        \"\"\"两层 MLP: input -> hidden (ReLU) -> output.\"\"\"
        def __init__(self, input_dim, hidden_dim, output_dim):
            super().__init__()
            self.fc1 = nn.Linear(input_dim, hidden_dim)
            self.fc2 = nn.Linear(hidden_dim, output_dim)

        def forward(self, x):
            x = F.relu(self.fc1(x))
            x = self.fc2(x)
            return x""")

_h3('124 行：class MLP(nn.Module)')
_b('定义 MLP 类，继承 nn.Module。nn.Module 是 PyTorch 所有网络的基类，提供：')
_bl('参数自动注册：','所有 nn.Parameter 自动被 model.parameters() 收集')
_bl('设备迁移：','model.to(device) 一键切换 CPU/GPU')
_bl('模式切换：','model.train() / .eval() 控制 Dropout/BN 行为')
_bl('序列化：','torch.save(model.state_dict(), ...) 保存/加载')
_bx('如果没有 nn.Module，你需要手动记录每个参数的位置、手动写 to(device)、手动实现序列化。50 层的 ResNet 就彻底失控了。')

_h3('125 行：docstring')
_code('"""两层 MLP: input -> hidden (ReLU) -> output."""')
_b('类文档字符串。不执行代码，但几周后回看时帮你省 5 分钟回忆时间。')

_h3('126 行：构造函数')
_code('def __init__(self, input_dim, hidden_dim, output_dim):')
_b('当你写 MLP(4, 32, 1) 时自动调用。三个参数：')
_bl('input_dim：','输入特征维度（如 4 个传感器读数）')
_bl('hidden_dim：','隐藏层神经元数量（越大容量越大，但易过拟合）')
_bl('output_dim：','输出维度（功率预测为 1）')

_h3('127 行：super().__init__()')
_code('super().__init__()')
_b('调用父类 nn.Module 构造函数。这是参数注册的入口——不调用的话 model.parameters() 返回空列表，梯度无法更新。')
_w('绝对不要忘记 super().__init__()！这是最常见的 PyTorch 新手 bug。')

_h3('128-129 行：全连接层')
_code('self.fc1 = nn.Linear(input_dim, hidden_dim)')
_code('self.fc2 = nn.Linear(hidden_dim, output_dim)')
_b('nn.Linear 执行 y = xW^T + b。自动创建两个参数 weight 和 bias。')
_code('''input_dim=4, hidden_dim=32, output_dim=1:
fc1: weight [32,4], bias [32]  → 128+32=160 参数
fc2: weight [1,32], bias [1]   →  32+1= 33 参数
总计:                                         193 参数''')
_b('带 self 才是实例属性，不加 self 在 __init__ 结束后变量就丢了。')

_h3('131-134 行：前向传播')
_code('def forward(self, x):')
_code('    x = F.relu(self.fc1(x))')
_code('    x = self.fc2(x)')
_code('    return x')
_b('131 行：forward = 前向传播。与 __init__ 分工：__init__ 说"有哪些层"，forward 说"数据怎么流过"。')
_b('132 行：从里往外拆解：')
_bl('① self.fc1(x)：','线性变换 [10,4]→[10,32]')
_bl('② F.relu(...)：','max(0,x) 负数清零，引入非线性')
_bl('③ x = ...：','结果存回变量 x')
_b('133 行：self.fc2(x): [10,32]→[10,1]。输出层无激活 = 回归头（输出任意实数）。')
_b('分类任务在损失函数中处理：CrossEntropyLoss 自带 softmax。')
_bx('如果没有 ReLU：两层线性可合并为 y=x(W1W2)+(b1W2+b2)，等价于单层！ReLU 打破线性，让网络真正"深"起来。')

_h2('137-144 行：实例化和使用')
_code('model = MLP(input_dim=4, hidden_dim=32, output_dim=1)')
_code("print(f'Model:\\n{model}')")
_code("print(f'Parameters: {sum(p.numel() for p in model.parameters())}')")
_code('x = torch.randn(10, 4)   # batch=10, features=4')
_code('y = model(x)')
_code("print(f'Input: {x.shape}, Output: {y.shape}')")

_b('137 行：MLP(4,32,1) 触发 __init__，分配内存 + 初始化父类 + 建层。')
_b('138 行：print(model) 输出结构树：')
_code('''MLP((fc1): Linear(4→32), (fc2): Linear(32→1))''')
_b('139 行：p.numel() = number of elements。193 个 float32 = 772 字节（不到 1KB！）。')
_b('142 行：10 个样本批处理，batch 利用 GPU 并行加速。')
_b('143 行：model(x) 而非 model.forward(x)——__call__ 自动注册 hook 和控制 Dropout。')
_n('输入 [10,4] → 输出 [10,1]。每个样本 4 维输入映射到 1 维预测。')

_h3('Part 3 总结')
_code('''① 继承 nn.Module        class MLP(nn.Module):
② __init__ 定义层       self.fc1 = nn.Linear(...)
③ forward 定义前向      x = F.relu(self.fc1(x))
④ 实例化               model = MLP(4, 32, 1)
⑤ 查看结构             print(model)
⑥ 前向传播             y = model(x)''')
_b('这个模式对任何 PyTorch 模型都适用——无论是 2 层 MLP 还是 100 层 ResNet。')

# ===================== PART 4 =====================
_h1('五、Part 4: MLP 功率预测（151-253 行）')

_h2('162-165 行：生成模拟数据')
_code('np.random.seed(42)')
_code('t = np.linspace(0, 100, 1000)')
_code('power = 30 + 15 * np.sin(0.1 * t) + 5 * np.sin(0.5 * t) + np.random.randn(1000) * 2')
_code('power = np.clip(power, 10, 80)')
_b('模拟燃料电池功率信号：基线 30kW + 低频波动 15kW + 高频波动 5kW + 高斯噪声 σ=2。')
_b('np.clip(10,80) 将功率限制在 10~80kW 之间（真实燃料电池的合理范围）。')

_h2('168-173 行：滑动窗口建样本')
_code('def create_sequences(data, seq_len=10):')
_code('    X, y = [], []')
_code('    for i in range(len(data) - seq_len):')
_code('        X.append(data[i:i+seq_len])')
_code('        y.append(data[i+seq_len])')
_code('    return np.array(X), np.array(y)')
_b('用前 10 步预测下一步。X=[样本数,10]，y=[样本数]。1000 个点 → 990 个样本。')

_h2('177-184 行：分割和预处理')
_code('split = int(0.8 * len(X))')
_code('X_train_t = torch.tensor(X_train).unsqueeze(-1)')
_b('80% 训练（792 样本），20% 测试（198 样本）。')
_b('unsqueeze(-1) 添加尾椎，形状 [792,10]→[792,10,1]，符合 Linear 输入约定。')

_h2('188-202 行：PowerPredictor 模型')
_code("""class PowerPredictor(nn.Module):
    def __init__(self, seq_len, hidden=64):
        super().__init__()
        self.fc1 = nn.Linear(seq_len, hidden)
        self.fc2 = nn.Linear(hidden, hidden)
        self.fc3 = nn.Linear(hidden, 1)
        self.dropout = nn.Dropout(0.1)
    def forward(self, x):
        x = x.view(x.size(0), -1)
        x = F.relu(self.fc1(x))
        x = self.dropout(x)
        x = F.relu(self.fc2(x))
        x = self.fc3(x)
        return x""")

_b('188 行：3 层 MLP，相比 Part 3 的 2 层 MLP 多了一个隐藏层。')
_b('189 行：seq_len=10（输入维度），hidden=64（默认值）。')

_h3('191-193 行：三层全连接')
_code('self.fc1 = nn.Linear(seq_len, hidden)    # 10→64')
_code('self.fc2 = nn.Linear(hidden, hidden)     # 64→64')
_code('self.fc3 = nn.Linear(hidden, 1)          # 64→1')
_b('参数量明细：')
_code('''层        权重形状     权重数    偏置    小计
fc1:      [64,10]      640       64      704
fc2:      [64,64]     4096       64     4160   ← 84% 参数
fc3:       [1,64]       64        1       65
总计:                                     4,929 参数 (~19KB)''')
_b('fc2 占绝大部分参数，承担主要特征提取任务。')

_h3('194 行：Dropout')
_code('self.dropout = nn.Dropout(0.1)')
_b('训练时 10% 概率随机置零神经元输出。类比"考试不能抄同桌"——每个神经元必须独立学好。')
_bl('防过拟合：','神经元不能依赖其他神经元的存在')
_bl('训练时开，测试时关：','.eval() 自动关闭 Dropout')
_bl('0.1 = 弱正则化：','4929 参数不算大，10% 丢弃足够')

_h3('197 行：view 展平')
_code('x = x.view(x.size(0), -1)')
_b('形状 [batch, 10, 1] → [batch, 10]。因为 nn.Linear 期望 2D 输入。')
_b('x.size(0) 保留 batch 大小，-1 自动计算剩余维度（10×1=10）。')
_n('等效写法: x = x.squeeze(-1) 或 x = x.reshape(x.shape[0], -1)')

_h3('198-201 行：前向计算流')
_code('x = F.relu(self.fc1(x))    # [batch,64]  线性→ReLU')
_code('x = self.dropout(x)        # [batch,64]  随机丢弃')
_code('x = F.relu(self.fc2(x))    # [batch,64]  二次特征提取')
_code('x = self.fc3(x)            # [batch,1]   输出层（无激活）')
_b('数据流总览：')
_code('''输入 [32,10,1] → view → [32,10] → fc1 → [32,64] → ReLU → [32,64]
→ Dropout → [32,64] → fc2 → [32,64] → ReLU → [32,64]
→ fc3 → [32,1] → return''')

_h2('204-206 行：实例化 + 优化器 + 损失')
_code('model = PowerPredictor(SEQ_LEN, hidden=64)')
_code('optimizer = torch.optim.Adam(model.parameters(), lr=0.001)')
_code('loss_fn = nn.MSELoss()')

_b('204 行：创建 4929 参数的模型。')
_b('205 行：Adam 优化器——当前最主流的默认优化器。')
_bx('Adam = Momentum(惯性)+RMSProp(自适应学习率)。内部 5 步：算梯度→更新动量 m_t→更新 RMS v_t→偏差校正→更新参数。')

_h3('Adam 的进化链')
_b('SGD（随机梯度下降）：θ = θ - lr·∇L。固定学习率，峡谷震荡。类比"蒙眼下山"。')
_b('Momentum：v = γv + lr·∇L, θ = θ - v。加惯性，平路加速震荡抵消。类比"滚球下山"。')
_b('RMSProp：E[g²]=β·E[g²]+(1-β)·g², θ=θ-lr/√E[g²]·g。每个参数自适应学习率。类比"登山杖"。')
_b('Adam = Momentum + RMSProp。既有惯性又自适应。lr=0.001 大多数任务直接可用。')

_b('206 行：MSELoss = (1/n)Σ(y_pred-y_true)²。平方放大误差惩罚，适合回归。')

_h2('208-223 行：训练循环')
_code('n_epochs = 200')
_code('train_losses = []')

_h3('211-217 行：训练四件套')
_code('for epoch in range(n_epochs):')
_code('    model.train()')
_code('    optimizer.zero_grad()')
_code('    y_pred = model(X_train_t)')
_code('    loss = loss_fn(y_pred, y_train_t)')
_code('    loss.backward()')
_code('    optimizer.step()')

_b('212 行：model.train() —— 开启 Dropout。如果忘记，Dropout 不生效。')
_b('213 行：optimizer.zero_grad() —— 清零梯度。PyTorch 默认累积梯度（设计如此——可做梯度累积模拟大 batch）。')
_b('214 行：前向传播。model(X_train_t) 调用 forward。X_train_t: [792,10,1]→y_pred: [792,1]。')
_b('215 行：算损失。比较预测值和真实值的差距。loss 是标量张量。')
_b('216 行：loss.backward() —— 自动求导！从 loss 沿计算图反向遍历，链式法则算每个参数梯度，存到 param.grad。')
_b('217 行：optimizer.step() —— 参数更新。θ = θ - lr·∇θ。')
_b('218 行：train_losses.append(loss.item()) 记录损失画图用。')

_h3('219-223 行：定期打印')
_code('if (epoch+1) % 50 == 0:')
_code('    model.eval()')
_code('    with torch.no_grad():')
_code('        test_loss = loss_fn(model(X_test_t), y_test_t)')
_code("    print(f'Epoch {epoch+1:3d}: ...')")
_b('model.eval() 关 Dropout，no_grad() 省内存。每 50 轮打印训练和测试损失。')
_code('''Epoch  50: train_loss=14.458, test_loss=5.418
Epoch 100: train_loss= 9.251, test_loss=4.946
Epoch 150: train_loss= 7.968, test_loss=5.443
Epoch 200: train_loss= 7.212, test_loss=5.635''')
_b('训练损失持续下降（模型在学习），测试损失在 100 轮后趋稳（已达最优）。')

_h2('226-232 行：评估')
_code('model.eval()')
_code('with torch.no_grad():')
_code('    y_pred = model(X_test_t).numpy().flatten()')
_code('    y_true = y_test')
_code('mae = np.mean(np.abs(y_pred - y_true))')
_code('rmse = np.sqrt(np.mean((y_pred - y_true) ** 2))')
_code("print(f'Test MAE: {mae:.3f} kW, RMSE: {rmse:.3f} kW')")
_b('model.eval() 关闭 Dropout，确保测试一致性。')
_b('.numpy() Tensor → NumPy，.flatten() [198,1]→[198]。')
_b('MAE = mean(|y_pred-y_true|) = 1.888kW —— 平均每个预测差 1.888kW。')
_b('RMSE = sqrt(mean((y_pred-y_true)^2)) = 2.374kW —— RMSE>MAE 说明存在大误差点。')

_h2('235-251 行：可视化')
_code('fig, axes = plt.subplots(2, 1, figsize=(12, 8))')
_code('axes[0].plot(train_losses)')
_code('axes[0].set_xlabel("Epoch")')
_code('axes[0].set_ylabel("MSE Loss")')
_code('axes[0].set_title("Training Loss")')
_code('axes[0].grid(True)')
_b('上子图：训练损失曲线。x=Epoch, y=MSE Loss。')

_code('axes[1].plot(y_true[:200], label="True", alpha=0.7)')
_code('axes[1].plot(y_pred[:200], label="Predicted", alpha=0.7)')
_code('axes[1].set_xlabel("Time Step")')
_code('axes[1].set_ylabel("FC Power (kW)")')
_code('axes[1].legend()')
_code('axes[1].grid(True)')
_b('下子图：前 200 个测试点的预测 vs 真实。alpha=0.7 透明度便于重叠观察。')
_code('plt.savefig(os.path.join(RESULTS_DIR, "..."), dpi=150)')
_code('plt.close()')
_b('dpi=150 比默认 100 清晰。保存到 results/。plt.close() 释放内存。')

_h2('253 行：返回值')
_code('return model')
_b('返回训练好的 PowerPredictor 模型。')

# ===================== PART 5 =====================
_h1('六、Part 5: MDP 五元组 — GridWorld（258-324 行）')
_b('Part 5 构建完整的 MDP 环境——4×4 网格世界。这是后续 Bellman 方程和策略求解的基础。')

_h2('259 行：函数定义')
_code('def part5_mdp_gridworld():')
_b('返回包含完整 MDP 数据的字典（SIZE, n_states, R, P 等），供 Part 6-8 使用。')

_h2('260-265 行：docstring')
_code('"""GridWorld MDP: ... """')
_b('完整描述环境设定：4×4 网格、起点终点、动作空间、随机转移、折扣因子。')

_h2('266-270 行：MDP 基本参数（S 和 A）')
_code('SIZE = 4')
_code('n_states = SIZE * SIZE')
_code("actions = ['↑', '↓', '←', '→']")
_code('n_actions = len(actions)')
_code('gamma = 0.9')
_bl('S = 状态空间：','4×4 = 16 个离散状态')
_bl('A = 动作空间：','4 个离散动作')
_bl('γ = 折扣因子：','0.9。未来第 k 步的奖励现在值 γ^k 倍')

_h3('折扣因子详解')
_b('γ ∈ [0,1] 控制对未来的重视程度：')
_bl('γ = 0：','只看眼前（即时奖励），完全近视')
_bl('γ = 0.9：','第 10 步奖励现在值 0.9^10 ≈ 0.35 倍')
_bl('γ → 1：','远视，未来奖励几乎和现在一样重要')
_n('γ=0.9 是常用值。智能体优先拿近奖励，但也愿意为远处更大奖励绕路。')

_h2('272-275 行：特殊位置')
_code('GOAL = (3, 3); TRAP = (1, 1)')
_code('GOAL_IDX = GOAL[0] * SIZE + GOAL[1]')
_code('TRAP_IDX = TRAP[0] * SIZE + TRAP[1]')
_b('坐标转线性索引 idx = r·4 + c：')
_code('''(0,0)=0  (0,1)=1  (0,2)=2  (0,3)=3
(1,0)=4  (1,1)=5  (1,2)=6  (1,3)=7    ← 5 陷阱
(2,0)=8  (2,1)=9  (2,2)=10 (2,3)=11
(3,0)=12 (3,1)=13 (3,2)=14 (3,3)=15   ← 15 终点''')

_h2('277-280 行：动作映射')
_code("action_delta = {'↑':(-1,0), '↓':(1,0), '←':(0,-1), '→':(0,1)}")
_b('每个动作对应 (行偏移, 列偏移)。')

_h2('282-286 行：辅助函数')
_code('def pos_to_idx(r, c): return r * SIZE + c')
_code('def is_valid(r, c): return 0 <= r < SIZE and 0 <= c < SIZE')
_b('pos_to_idx：坐标→索引。is_valid：检查是否出界。')

_h2('288-289 行：MDP 核心数据结构（R 和 P）')
_code('R = {s: {a: 0.0 for a in range(n_actions)} for s in range(n_states)}')
_code('P = {s: {a: {} for a in range(n_actions)} for s in range(n_states)}')

_h3('R[s][a] — 奖励函数')
_b('嵌套字典：R[状态索引][动作索引] = 即时奖励。16 状态 × 4 动作。')

_h3('P[s][a] — 转移概率')
_b('P[状态索引][动作索引] = {目标状态: 概率, ...}。字典形式只存非零项（稀疏存储）。')
_code('''例: P[0]["↑"] = {0:0.2667, 1:0.0667, 4:0.0667}
在左上角向上走 → 80% 撞墙留原地 → 20% 滑向其他方向''')

_h2('291-312 行：填充 R 和 P（MDP 构建核心）')
_code('for r, c in itertools.product(range(SIZE), range(SIZE)):')
_b('itertools.product 生成 4×4=16 个坐标组合，遍历所有格子。')

_h3('分支 1：终点和陷阱（293-297 行）')
_code('if s == GOAL_IDX or s == TRAP_IDX:')
_code('    for a in range(n_actions):')
_code('        R[s][a] = 1.0 if s == GOAL_IDX else -1.0')
_code('        P[s][a][s] = 1.0  # 100% 留原地')
_code('    continue')
_b('终止状态：所有动作奖励固定（+1/-1），100% 留在原地（游戏结束）。')

_h3('分支 2：普通格子（298-312 行）')
_code('for a_idx, (action_name, (dr, dc)) in enumerate(action_delta.items()):')
_code('    nr, nc = r + dr, c + dc')
_code('    if not is_valid(nr, nc): nr, nc = r, c   # 撞墙留原地')
_code('    P[s][a_idx][target_s] += 0.8  # 80% 目标方向')
_code('    for other_dr, other_dc in action_delta.values():')
_code('        if same direction: continue')
_code('        P[s][a_idx][other_s] += 0.2 / 3  # 20% 其他方向均分')
_b('80% 目标方向，20% 均匀分配到其他 3 个方向（每个 ≈ 6.67%）。撞墙则留在原地。')
_bx('随机转移模拟了现实不确定性：轮子打滑、风阻、路面不平。没有随机性的是玩具环境。')
_n('P[s][a].get(target_s, 0) 确保键不存在时从 0 开始加。')

_h3('转移概率示例')
_code('在 (0,0) 执行 ↓：')
_code('  80% → (1,0)        [目标]')
_code('  6.67% → (0,0)      [↑ 撞墙]')
_code('  6.67% → (0,0)      [← 撞墙]')
_code('  6.67% → (0,1)      [→ 滑过去]')
_code('  总计 = 100% ✓')

_h2('314-317 行：打印 MDP 信息')
_code('print(f\'GridWorld: {n_states} 状态 × {n_actions} 动作\')')
_code('print(f\'起点 (0,0), 终点 (3,3)+1, 陷阱 (1,1)-1\')')
_code('print(f\'γ={gamma}, 80%目标/20%随机\')')

_h2('319-324 行：返回值')
_code("return {'SIZE':SIZE, 'n_states':n_states, 'R':R, 'P':P, ...}")
_b('打包返回完整的 MDP 结构。后续 Part 通过 mdp["R"]、mdp["gamma"] 访问。')

_h2('Part 5 总结')
_bl('MDP 五元组：','⟨S,A,P,R,γ⟩ = ⟨16,4,dict,dict,0.9⟩')
_bl('奖励：','终点+1（目标），陷阱-1（惩罚），正常格子0')
_bl('随机转移：','80% 目标方向 + 20% 随机滑走')
_bl('终止状态：','终点和陷阱不再移动（episode 结束）')

# ===================== PART 6 =====================
_h1('七、Part 6: Bellman 方程（330-403 行）')
_b('Part 6 计算两种值函数：V^π（随机策略下的值）和 V*（最优值函数）。')

_h2('330-336 行：解包')
_code('def part6_bellman(mdp):')
_code('    n_states = mdp["n_states"]')
_code('    ...')
_b('从 mdp 字典解包出 n_states、n_actions、gamma、R、P。')

_h2('339 行：随机策略')
_code('policy = np.ones((n_states, n_actions)) / n_actions')
_b('16×4 均匀概率矩阵，每个动作 25%。作为 Baseline 对比最优策略。')
_n('输出: V(0) = -3.9284（随机策略频繁踩陷阱）→ V*(0) = 3.3419（最优策略绕开陷阱）')

_h2('342-361 行：策略评估（核心）')
_code('V = np.zeros(n_states)')
_code('theta = 1e-6')
_code('max_iter = 1000')
_code('for i in range(max_iter):')
_code('    delta = 0')
_code('    for s in range(n_states):')
_code('        v_old = V[s]')
_code('        v_new = 0')
_code('        for a in range(n_actions):')
_code('            p_a = policy[s, a]')
_code('            if p_a == 0: continue')
_code('            bellman_sum = R[s][a]')
_code('            for s_next, prob in P[s][a].items():')
_code('                bellman_sum += gamma * prob * V[s_next]')
_code('            v_new += p_a * bellman_sum')
_code('        V[s] = v_new')
_code('        delta = max(delta, abs(v_old - v_new))')
_code('    if delta < theta: break')

_b('实现 Bellman 期望方程：')
_b('V^π(s) = Σ_a π(a|s)·[R(s,a) + γ·Σ_{s\'} P(s\'|s,a)·V^π(s\')]')

_h3('逐层拆解')
_b('① for i (外层)：迭代直到 V 稳定（133 次）')
_b('② delta：记录本轮所有状态中 V 的最大变化量')
_b('③ for s (中层)：对 16 个状态逐一更新')
_b('④ for a (内层1)：对 4 个动作做概率加权平均')
_b('⑤ for s_next (内层2)：对每个动作的随机转移结果求和')
_b('⑥ V[s] = v_new：更新值函数')
_b('⑦ if delta < theta：收敛则退出')
_n('V(0) = -3.9284 —— 随机策略频繁踩陷阱，值为负。')

_h2('365-378 行：最优值函数 V*')
_code('V_opt = np.zeros(n_states)')
_code('for i in range(max_iter):')
_code('    for s in range(n_states):')
_code('        q_values = []')
_code('        for a in range(n_actions):')
_code('            q = R[s][a]')
_code('            for s_next, prob in P[s][a].items():')
_code('                q += gamma * prob * V_opt[s_next]')
_code('            q_values.append(q)')
_code('        V_opt[s] = max(q_values)')
_b('核心区别：用 max 代替加权平均——"最优"就是"选最好的动作"。')
_n('V*(0) = 3.3419 —— 远高于随机策略的 -3.93。')

_h2('380-383 行：打印')
_code("print(f'V(0)={V[0]:.4f}, V*(0)={V_opt[0]:.4f}')")

_h2('385-402 行：可视化')
_code('fig, axes = plt.subplots(1, 2, figsize=(10, 4))')
_code('im0 = axes[0].imshow(V.reshape(4,4), cmap="RdYlBu_r", vmin=-1, vmax=1)')
_code('axes[0].set_title("V^π(s) — 随机策略")')
_b('左图 V^π（随机策略，多为蓝色负值），右图 V*（最优策略，终点附近红色正值）。')
_code('for r in range(4):')
_code('    for c in range(4):')
_code('        axes[0].text(c, r, f"{V[r*4+c]:.2f}", ha="center", va="center", fontsize=8)')
_b('在每个格子叠加数值。固定 vmin/vmax 让两张图颜色可比。')

# ===================== PART 7 =====================
_h1('八、Part 7: 策略迭代（410-496 行）')
_b('策略迭代 = 策略评估 + 策略改进，交替进行直到策略收敛。')

_h2('410-416 行：函数定义')
_code('def part7_policy_iteration(mdp):')
_code('    """策略迭代: 策略评估 → 策略改进 → 直到收敛"""')
_b('接受 mdp 字典，返回 (policy, V)。')

_h2('418-419 行：初始化')
_code('policy = np.random.randint(0, n_actions, size=n_states)')
_code('V = np.zeros(n_states)')
_bl('policy：一维 [16]，每个元素 0~3（随机动作）','确定性策略，与 Part 6 的概率矩阵不同')
_bl('V：值函数全零')

_h2('421-434 行：策略评估（内嵌函数）')
_code('def policy_evaluation(policy, V, theta=1e-6):')
_code('    for _ in range(1000):')
_code('        for s in range(n_states):')
_code('            v_old = V[s]')
_code('            a = policy[s]     # 确定性策略，只有一个动作')
_code('            v_new = R[s][a]')
_code('            for s_next, prob in P[s][a].items():')
_code('                v_new += gamma * prob * V[s_next]')
_code('            V[s] = v_new')
_code('        if delta < theta: break')
_code('    return V')
_b('与 Part 6 策略评估的区别：确定性策略直接取 policy[s] 作为唯一动作，不需要对 4 个动作做加权平均，计算量少 3/4。')
_code('V^π(s) = R(s,π(s)) + γ·Σ P(s\'|s,π(s))·V^π(s\')')

_h2('436-449 行：策略改进')
_code('def policy_improvement(policy, V):')
_code('    policy_stable = True')
_code('    for s in range(n_states):')
_code('        old_action = policy[s]')
_code('        q_values = []')
_code('        for a in range(n_actions):')
_code('            q = R[s][a]')
_code('            for s_next, prob in P[s][a].items():')
_code('                q += gamma * prob * V[s_next]')
_code('            q_values.append(q)')
_code('        policy[s] = int(np.argmax(q_values))')
_code('        if old_action != policy[s]:')
_code('            policy_stable = False')
_code('    return policy, policy_stable')

_b('核心思想：基于当前 V 贪心地选择更好的动作。')
_code('π\'(s) = argmax_a [R(s,a) + γ·Σ P(s\'|s,a)·V(s\')]')
_b('如果所有状态下的动作不再改变，policy_stable=True，迭代可以停止。')

_h2('451-459 行：主循环')
_code('for iteration in range(50):')
_code('    V = policy_evaluation(policy, V)')
_code('    policy, stable = policy_improvement(policy, V)')
_code('    if stable: break')
_code('    else: print(...)')
_b('运行输出：')
_code('''第 1 轮: 策略改进中...  ← 随机策略→评估→改进（大变）
第 2 轮: 策略改进中...  ← 微调
第 3 轮: 策略改进中...  ← 个别格子调整
第 4 轮: 收敛 [OK]      ← 策略稳定''')
_bx('策略迭代收敛非常快（4 轮），因为"评估+改进"在策略空间直接跳跃到更优区域。')

_h2('461-474 行：打印最优策略')
_code('action_symbols = {0:"↑", 1:"↓", 2:"←", 3:"→"}')
_code('for r in range(SIZE):')
_code('    for c in range(SIZE):')
_code('        if s == GOAL: row += " G "')
_code('        elif s == TRAP: row += " X "')
_code('        else: row += f" {action_symbols[policy[s]]} "')
_b('最优策略网格：')
_code('''  ↑   ↑   ↑   ↑
  ↑   X   ↑   ↑     ← 避开陷阱(1,1)
  ↑   ↑   ↑   ↑
  ↑   ↑   →   G     ← (3,2)右转进终点''')
_b('解读：引导智能体从 (0,0) 向右下角移动，绕开陷阱。')

_h2('476-494 行：可视化')
_code('fig, ax = plt.subplots(figsize=(5,5))')
_code('ax.imshow(V.reshape(4,4), cmap="RdYlBu_r", vmin=-1, vmax=1)')
_code('ax.text(c, r, action_symbols[policy[s]], fontsize=16)')
_b('imshow 显示 V 的颜色背景 + 叠加策略箭头。保存到 results/。')

# ===================== PART 8 =====================
_h1('九、Part 8: 值迭代（502-558 行）')
_b('值迭代直接迭代 Bellman 最优方程，不显式维护策略。')

_h2('502-511 行：定义和初始化')
_code('def part8_value_iteration(mdp):')
_code('    V = np.zeros(n_states)')
_code('    theta = 1e-6')
_b('值迭代只需要 V 数组，不需要 policy。收敛后从 V 提取策略。')

_h2('513-529 行：核心迭代')
_code("print(f'\\n[Part 8] 值迭代')")
_code('for iteration in range(1000):')
_code('    delta = 0')
_code('    for s in range(n_states):')
_code('        v_old = V[s]')
_code('        q_max = -np.inf')
_code('        for a in range(n_actions):')
_code('            q = R[s][a]')
_code('            for s_next, prob in P[s][a].items():')
_code('                q += gamma * prob * V[s_next]')
_code('            if q > q_max: q_max = q')
_code('        V[s] = q_max')
_code('        delta = max(delta, abs(v_old - V[s]))')
_code('    if delta < theta: break')

_b('实现 Bellman 最优方程：V_{k+1}(s) = max_a [R(s,a) + γ·Σ P·V_k(s\')]')

_h3('与策略迭代的关键区别')
_bl('策略迭代：评估→改进交替','4 轮收敛，每轮成本高（完整策略评估）')
_bl('值迭代：直接更新 V','133 轮收敛，每轮成本低（一次 max 操作）')
_b('q_max = -np.inf：初始化为负无穷，确保第一个 Q 值必定替换。')
_n('133 次迭代后收敛。比策略迭代轮次多得多，但每轮计算量远小于一轮完整策略评估。')

_h2('531-539 行：提取策略')
_code('policy = np.zeros(n_states, dtype=int)')
_code('for s in range(n_states):')
_code('    q_values = []')
_code('    for a in range(n_actions):')
_code('        q = R[s][a] + γ·Σ P·V[s_next]')
_code('        q_values.append(q)')
_code('    policy[s] = int(np.argmax(q_values))')
_b('当 V 收敛到 V* 后，贪心提取最优策略。与 Part 7 策略改进步骤完全一样。')

_h2('541-554 行：打印策略')
_b('输出与 Part 7 完全一致，验证两种方法等价。')

_h2('Part 7 vs Part 8 对比')
_code('''特征        策略迭代           值迭代
─────────────────────────────────────
更新方式    评估+改进交替       直接更新 V
收敛轮数    4 轮               133 轮
每轮成本    高（完整策略评估）  低（一次 max）
策略表示    显式维护            从 V 后提取
适用场景    状态空间小          状态空间较大''')

# ===================== EXTRA =====================
_h1('十、Extra: 收敛可视化 + 验证（564-697 行）')

_h2('564-593 行：带追踪的值迭代')
_code('V_track = []')
_code('for iteration in range(100):')
_code('    ... # 同值迭代')
_code('    V_track.append(V.copy())  # 深拷贝！')
_b('V_track 记录每轮 V 的快照。V.copy() 必须深拷贝，否则 V_track 里全是最新值。')

_h2('596-618 行：可视化')
_code('axes[0].plot(V_track[:, 0], label="Start")    # 起点')
_code('axes[0].plot(V_track[:, GOAL_IDX], label="Goal")  # 终点')
_code('axes[0].plot(V_track[:, TRAP_IDX], label="Trap")  # 陷阱')
_b('三条曲线：起点从 0 爬到 3.34，终点跳到 1.0 稳定，陷阱降到 -1.0 稳定。')
_code('deltas = [max|V[i+1]-V[i]| for i in range(len(V_track)-1)]')
_code('axes[1].plot(deltas)')
_code('axes[1].set_yscale("log")')
_b('下图展示每轮最大 ΔV，对数坐标呈指数衰减（近似直线），证明值迭代指数收敛。')

_h2('625 行：入口')
_code("if __name__ == '__main__':")
_b('直接运行时执行以下代码，被 import 时跳过。')

_h2('682-692 行：最终验证')
_code("print(f'策略迭代 V(0) = {s0_pi:.4f}')")
_code("print(f'值迭代   V(0) = {s0_vi:.4f}')")
_code("print(f'ΔV = {abs(s0_pi - s0_vi):.6f}')")
_b('ΔV = 0.000006 < 1e-4 → 两种方法收敛到同一最优值函数 ✓')

# ===================== OPTIMIZER =====================
_h1('附录：优化器概念详解')
_b('训练神经网络 = 调整参数让 Loss 最小。优化器决定"往哪走、走多远"。')

_h2('SGD — 随机梯度下降')
_code('θ = θ - lr * ∇L(θ)')
_b('每次沿最陡反方向迈固定步长。问题：峡谷地形来回震荡、所有参数统一 lr。')
_b('类比"蒙眼下山"——用脚探哪边最陡，往反方向迈步。')

_h2('Momentum — 动量法')
_code('v = γ*v + lr*∇L')
_code('θ = θ - v')
_b('加惯性（γ=0.9）。之前往东，这次梯度偏北一点，主要还往东。')
_bl('平缓区域加速：','梯度持续指向同一方向时累积动量')
_bl('峡谷减少震荡：','窄方向正负梯度相互抵消')
_bl('能冲过小坑：','小坑梯度挡不住累积动量')
_b('类比"滚球下山"——越滚越快，能冲过小凹陷。')

_h2('RMSProp — 自适应学习率')
_code('E[g²] = β*E[g²] + (1-β)*g²')
_code('θ = θ - lr / sqrt(E[g²]+ε) * g')
_b('梯度大的参数：分母大 → 学习率自动变小（陡坡小步）')
_b('梯度小的参数：分母小 → 学习率自动变大（缓坡大步）')
_b('类比"登山杖"——每个方向独立调步长。')

_h2('Adam — 集大成者')
_code('m_t = β₁*m_{t-1} + (1-β₁)*g_t     # 动量')
_code('v_t = β₂*v_{t-1} + (1-β₂)*g_t²    # RMS')
_code('θ_{t+1} = θ_t - lr * m_t / (sqrt(v_t)+ε)')
_bl('β₁=0.9：','动量系数（Momentum 部分）')
_bl('β₂=0.999：','RMS 衰减率（RMSProp 部分）')
_bl('lr=0.001：','默认学习率，多数任务直接可用')
_bl('偏差校正：','前几步 m_t 和 v_t 偏小，自动修正')
_bx('Adam = Momentum(知道往哪走) + RMSProp(知道走多大步)。是目前最主流默认优化器。你代码中使用的正是 Adam ✓')

_h2('对比总表')
_code('''特性        SGD    Momentum  RMSProp   Adam
惯性        ✗       ✓(0.9)    ✗        ✓(0.9)
自适应lr    ✗       ✗        ✓(0.999)  ✓(0.999)
收敛速度    慢      中        中-快     快
超参数敏感  高      中        低        很低''')

_h2('实际选择建议')
_bl('新手/大多数任务：Adam','lr=0.001 直接可用')
_bl('计算机视觉：','SGD+Momentum，泛化有时更好')
_bl('NLP/Transformers：','AdamW（Adam+正确权重衰减）')
_bl('资源受限：','SGD（省一半显存）')

# ===================== SUMMARY =====================
_h1('附录：模型参数与核心原理速查')

_h2('模型参数统计')
_code('''MLP (Part 3):          193 参数 (fc1:160 + fc2:33) = 772 字节
PowerPredictor (Part 4): 4,929 参数 (fc1:704 + fc2:4160 + fc3:65) = ~19KB
GridWorld (Part 5-8):    无参数（规划方法，不是学习方法）''')

_h2('运行输出全记录')
_code("""Part 1: 广播 [[11,21,31],[12,22,32],[13,23,33]]
Part 2: dy/dx=[7,9]  autograd 验证 ✓
Part 3: MLP(4->32->1)  193 params
Part 4: MAE=1.888kW  RMSE=2.374kW
Part 5: GridWorld  16 状态 x 4 动作  γ=0.9
Part 6: V(0)=-3.9284  V*(0)=3.3419
Part 7: 策略迭代 4 轮收敛
Part 8: 值迭代 133 轮收敛
验证:  ΔV=0.000006  一致 ✓""")

_h2('核心原理一句话')
_bl('Tensor：','多维数组 + GPU 加速 + 自动求导')
_bl('Autograd：','计算图 + 链式法则 → 自动计算梯度')
_bl('nn.Module：','PyTorch 模型基类（参数管理/设备迁移/序列化）')
_bl('训练循环：','zero_grad → forward → loss → backward → step')
_bl('MDP：','⟨S,A,P,R,γ⟩ 强化学习问题框架')
_bl('Bellman 方程：','V=R+γ·V\' 递归定义长期价值')
_bl('策略迭代：','评估→改进→循环，收敛快（4轮）')
_bl('值迭代：','直接迭代 V→max Q，收敛慢但每轮便宜（133轮）')
_bl('两者关系：','收敛到同一最优解 ✓')

# ========== SAVE ==========
out_dir = 'F:/CLAUDE/research/ems-platform/docs'
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'Week9_完整逐行精讲_全篇合并_v2.docx')
doc.save(out_path)
print('OK saved -> ' + out_path)
