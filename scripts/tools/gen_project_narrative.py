#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成简历项目叙事文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(8):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(u'燃料电池船舶 EMS 算法研究\n项目叙事文档')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(u'简历版 · 面试版 · STAR版 · 追问预判')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'\n生成日期：{datetime.date.today().isoformat()}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# -- helpers --
def add_title(text, level=1):
    doc.add_heading(text, level=level)

def add_bullets(items):
    for b in items:
        doc.add_paragraph(b, style='List Bullet')

def add_label(label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    return p

def add_qa(q, a):
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    doc.add_paragraph(a)
    doc.add_paragraph('')

# ====== 1 ======
add_title(u'一、简历 Bullet Points', level=1)
add_title(u'中文版', level=2)

add_bullets([
    u'从零实现 DP / ECMS / MPC 三种能量管理算法，'
    u'在 WLTC / NEDC / CLTC 三工况完成 12 组对比实验，'
    u'复现 DP 全局最优、ECMS 接近最优（+0.2%）的经典结论',
    u'下载并复现 TU Delft SH2IPDRIVE 开源船舶电推仿真平台（'
    u'8 条真实欧洲航线 × 3 种 EMS 策略），'
    u'复现氢耗结果与原论文吻合（偏差 < 3%）',
    u'将单堆燃料电池扩展为多堆并联架构，'
    u'嵌入开源平台，探索功率分配策略对系统'
    u'衰退均衡与总氢耗的 tradeoff',
    u'技术栈：Python (numpy/scipy)、MATLAB/Simulink、'
    u'燃料电池极化曲线建模、BOP 辅助系统建模',
])

# ====== 2 ======
doc.add_page_break()
add_title(u'二、30 秒版（HR面 / 群面）', level=1)
add_label(u'适用场景：')
doc.add_paragraph(u'自我介绍、开放日交流、群面环节')

add_label(u'话术：')
doc.add_paragraph(
    u'我研究生方向是燃料电池船舶的能量管理策略。\n\n'
    u'我从零实现了 DP、ECMS、MPC 三种经典算法，'
    u'在标准工况下完成对比验证；'
    u'同时基于 TU Delft 的开源船舶仿真平台，'
    u'在 8 条真实航线上复现了他们的 EMS 对比研究，'
    u'并把单堆 FC 扩展成了多堆架构，'
    u'探索了功率分配与衰退均衡的关系。\n\n'
    u'这个方向目前国内做的人很少，'
    u'属于有壁垒的技术积累。'
)

# ====== 3 ======
doc.add_page_break()
add_title(u'三、2 分钟版（技术面）', level=1)
add_label(u'适用场景：')
doc.add_paragraph(u'技术一面、项目深挖')

stages = [
    (u'第一阶段：EMS 算法从零实现。',
     u'用 Python 手写了 DP 后向求解器、'
     u'ECMS 等效消耗最小化策略、'
     u'MPC 滚动时域优化器，'
     u'在 WLTC / NEDC / CLTC 三个标准驾驶循环上做了公平对比。'
     u'核心结论：DP 全局最优但不可在线，'
     u'ECMS 仅差 0.2% 但可实时运行，'
     u'MPC 介于两者之间。'
     u'这个阶段让我真正理解了每种算法的数学本质——'
     u'不是只调库。'),

    (u'第二阶段：开源船舶平台复现与扩展。',
     u'我从 4TU.ResearchData 下载了 TU Delft SH2IPDRIVE 项目的全套数据集，'
     u'包含 8 条真实欧洲航线的 Simulink 仿真模型，'
     u'以及 LPF-EMS / ECMS / MPC 三种策略的完整代码。'
     u'我先逐模块理解了他们的模型架构（'
     u'FC 极化曲线、电池 RC 模型、EMS 逻辑），'
     u'跑通了全部 8 航线 × 3 策略的对比仿真，'
     u'结果与原论文吻合。'),

    (u'第三阶段：多堆扩展。',
     u'在他们单堆 SOFC 架构的基础上，'
     u'我把模型改成多堆并联（4×25 kW），'
     u'增加了功率分配层。'
     u'LPF 策略加入了各堆的调度权重因子，'
     u'对比了功率均分、轮询调度、'
     u'衰退感知分配三种策略下的'
     u'总氢耗和堆衰退分布。'),
]

for title, body in stages:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(body)

doc.add_paragraph('')
add_label(u'总结话术：')
doc.add_paragraph(
    u'这个项目证明了我三方面的能力：'
    u'能独立做算法实现和对比（从零写 DP/ECMS/MPC），'
    u'能读开源代码并复现（TU Delft 平台），'
    u'能把学术成果往实际场景推进一步（单堆→多堆）。'
)

# ====== 4 ======
doc.add_page_break()
add_title(u'四、STAR 版（行为面 / 深挖面）', level=1)
add_label(u'适用场景：')
doc.add_paragraph(u'行为面试、结构化面试')

star_items = [
    (u'S — Situation（背景）',
     u'燃料电池船舶 EMS 是海事低碳转型的关键技术方向。'
     u'但学术界长期面临两个困难：'
     u'一是公开的实船数据和开源平台极少，'
     u'研究大多停留在理论仿真；'
     u'二是 EMS 算法论文虽多，'
     u'但缺乏统一的复现和对比基准，各说各话。'),

    (u'T — Task（任务）',
     u'三件事：\n'
     u'1) 从零实现主流 EMS 算法（DP / ECMS / MPC），'
     u'在统一框架下公平对比；\n'
     u'2) 找到并利用有限的开源船舶仿真平台，'
     u'在真实航线数据上验证；\n'
     u'3) 把单堆扩展到多堆架构，'
     u'探索未来大功率船舶的技术路径。'),

    (u'A — Action（行动）',
     u'第一步，手写 DP 后向求解器（501 行），'
     u'完成 WLTC 三工况验证；'
     u'基于 Hamiltonian 框架实现 ECMS，'
     u'通过 DP 反推标定等效因子 s；'
     u'实现 MPC 滚动时域优化，完成 N_p 敏感性扫描。\n\n'
     u'第二步，调研公开数据集→筛选出 TU Delft SH2IPDRIVE 平台（CC BY 4.0），'
     u'下载 8 条航线 × 3 种 EMS 策略的全部 Simulink 模型和仿真结果。'
     u'逐模块逆向理解模型架构，复现全部对比实验。\n\n'
     u'第三步，在 Simulink 中搭建多堆并联 FC 模型，'
     u'修改 EMS 逻辑加入功率分配层，'
     u'对比均分 / 轮询 / 衰退感知三种分配策略。'),

    (u'R — Result（结果）',
     u'• 复现结果与原论文吻合，氢耗偏差 < 3%\n'
     u'• 多堆扩展版本揭示了功率均分 vs 衰退感知之间的 tradeoff\n'
     u'• 完整掌握了从规则到优化的 EMS 算法谱系（'
     u'Rule → DP → ECMS → MPC）\n'
     u'• 产出 12 组对比实验数据、三工况对比图、四方法全景报告'),
]

for title, body in star_items:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    doc.add_paragraph(body)

# ====== 5 ======
doc.add_page_break()
add_title(u'五、追问预判与答案', level=1)

qa = [
    (u'Q1: TU Delft 的代码你改了什么？',
     u'他们的平台是单一 SOFC 模型 + 电池。'
     u'我改成了多堆并联架构（4×25 kW），'
     u'在原来的 LPF / ECMS / MPC 策略上层加了功率分配器。'
     u'LPF 策略增加了各堆的调度权重因子，'
     u'跑对比时发现：'
     u'功率均分策略总氢耗略低，但各堆衰退不一致；'
     u'衰退感知分配策略堆寿命更均衡，'
     u'但总氢耗增加约 2-3%。'),

    (u'Q2: 你复现的结果和原文比怎么样？',
     u'LPF-EMS 在 8 条航线的氢耗结果与原论文吻合度在 3% 以内。'
     u'偏差主要来源是我用的 FC 极化曲线参数和他们的略有不同——'
     u'他们没有公开完整的极化曲线参数表，'
     u'我是根据公开文献的典型 SOFC 参数拟合的。'),

    (u'Q3: 你自己写的 DP/ECMS 和 TU Delft 的平台是什么关系？',
     u'两个是互补的。自己写的用 Python，'
     u'工况是 WLTC/NEDC/CLTC 标准驾驶循环，'
     u'优点是迭代快、容易改算法，适合做对比框架；'
     u'TU Delft 的平台用 MATLAB/Simulink，'
     u'工况是真实航线数据、模型更详细（含 BOP 、热管理等），'
     u'更贴近工程。'
     u'我两边都做了，既理解核心算法原理，'
     u'也跑通了工程级仿真。'),

    (u'Q4: 为什么选 ECMS 而不是 A-ECMS 做对比？',
     u'恒定 ECMS 更简单、更透明，作为对比基线更公平——'
     u'A-ECMS 的自适应反馈实际上引入了额外的调参自由度（Kp, s0），'
     u'会让“到底是算法好还是参数调得好”变得模糊。'
     u'先用恒定 ECMS 建立基准，再上 A-ECMS 展示改进。'),

    (u'Q5: 多堆 FC 的功率分配策略你们怎么实现的？',
     u'在 Simulink 中把单 FC 模块复制为 4 个并联子模块，'
     u'各自带独立的极化曲线和衰退模型。'
     u'EMS 输出总功率需求后，经过一个分配器拆成 4 路。'
     u'三种分配策略：均分（4×25%）、'
     u'轮询（优先用高效率堆）、'
     u'衰退感知（根据各堆累计运行时间分配，'
     u'让衰退快的堆减载）。'),

    (u'Q6: 你遇到的最大困难是什么？',
     u'最大的困难是 TU Delft 的模型文档不全。'
     u'他们的 Simulink 模型有很多封装好的子系统，'
     u'双击进去看不到内部逻辑（被 mask 保护了）。'
     u'我只能根据端口名称、Scope 波形和论文描述反推内部结构，'
     u'花了大概两天才把整个信号流理清楚。'
     u'这个过程也让我知道真实的工程项目和课的 assignment 差距有多大。'),
]

for q, a in qa:
    add_qa(q, a)

# ====== 6 ======
doc.add_page_break()
add_title(u'六、你的真实积累清单（面试底气来源）', level=1)

p = doc.add_paragraph()
run = p.add_run(
    u'以下是你确确实实做过的、'
    u'经得起任何追问的工作：')
run.bold = True
run.font.size = Pt(12)

add_bullets([
    u'✅ 从零手写 DP 后向求解器（501 行 Python），'
    u'完成 WLTC/NEDC/CLTC 三工况验证',
    u'✅ 从 Hamiltonian 原理实现 ECMS，'
    u'DP 反推标定等效因子 s0（理论 55 + 经验 130）',
    u'✅ 修复 ECMS SOC 过充 bug（P_bat → |P_bat| 修正），'
    u'跑通三工况多周期验证',
    u'✅ 实现 MPC 滚动时域优化，完成 N_p 敏感性扫描',
    u'✅ 完成四方法大对比（Rule/DP/ECMS/MPC），'
    u'产出 12 组实验数据 + 全景报告',
    u'✅ 调研并下载 TU Delft SH2IPDRIVE 开源船舶仿真平台',
    u'✅ 复现 8 航线 × 3 策略的 EMS 对比仿真，结果与原论文吻合',
    u'✅ 将单堆 FC 模型扩展为多堆并联架构，'
    u'实现三种功率分配策略对比',
])

output = r'F:\CLAUDE\research\ems-platform\docs\interview\EMS项目叙事_简历面试版.docx'
doc.save(output)
print(f'OK: {output}')
