#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 dqn_explained.py 逐行分析 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def p(text, bold=False, sz=11, color=None, indent=0):
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(3)
    if indent: pa.paragraph_format.left_indent = Cm(indent)
    run = pa.add_run(text); run.font.name='微软雅黑'; run.font.size=Pt(sz); run.bold=bold
    if color: run.font.color.rgb = color

def b(text, sz=11, color=None): p(text, bold=True, sz=sz, color=color)
def bullet(text, lv=0):
    pa = doc.add_paragraph(text, style='List Bullet')
    pa.paragraph_format.left_indent = Cm(1.5 + lv*0.8)

def code(lines, label=None):
    if label:
        pa = doc.add_paragraph(); run=pa.add_run(label); run.bold=True; run.font.size=Pt(10)
    for line in lines.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before=Pt(0); pa.paragraph_format.space_after=Pt(1)
        pa.paragraph_format.left_indent=Cm(1)
        run=pa.add_run(line); run.font.name='Consolas'; run.font.size=Pt(9.5); run.font.color.rgb=RGBColor(0x33,0x33,0x33)

def tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i,hd in enumerate(headers):
        run=t.rows[0].cells[i].paragraphs[0].add_run(hd); run.bold=True; run.font.size=Pt(10); run.font.name='微软雅黑'
    for rd in rows:
        row=t.add_row()
        for c,txt in enumerate(rd):
            run=row.cells[c].paragraphs[0].add_run(txt); run.font.size=Pt(10); run.font.name='微软雅黑'

def brk(): doc.add_page_break()
def note(text):
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    run = pa.add_run(f'  {text}')
    run.font.name = '微软雅黑'; run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

def explain(line_num, code_line, explanation):
    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(4)
    pa.paragraph_format.space_after = Pt(1)
    run = pa.add_run(f'L{line_num}  ')
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x99,0x99,0x99)
    run = pa.add_run(code_line)
    run.font.name = 'Consolas'; run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = RGBColor(0x1F,0x3A,0x5F)
    p(explanation, indent=1, sz=10)

# ======================== 封面 ========================
for _ in range(4): doc.add_paragraph('')
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t.add_run('DQN 逐行拆解 — 逐行精讲\n'); run.font.size=Pt(24); run.bold=True; run.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
run=t.add_run('dqn_explained.py'); run.font.size=Pt(18); run.font.color.rgb=RGBColor(0x66,0x66,0x66)
t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t2.add_run('从 Q-learning 出发，改一行就是 DQN'); run.font.size=Pt(13); run.font.color.rgb=RGBColor(0x66,0x66,0x66)
t3=doc.add_paragraph(); t3.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t3.add_run(f'生成日期：{datetime.date.today().isoformat()}'); run.font.size=Pt(10); run.font.color.rgb=RGBColor(0x99,0x99,0x99)
brk()

# ================================================================
h('文件概览', 1)
p('文件名: dqn_explained.py')
p('功能: 从 Q-learning 出发，一步步改成 DQN，核心是 Part 6 的手动追踪一步更新')
p('')
tbl(
    ['Part', '内容', '行号', '建议'],
    [
        ['Part 1', 'Q-learning 复习', '63-79', '已会，扫一眼'],
        ['Part 2', 'Q 表 → 神经网络', '85-163', '核心改变，仔细看'],
        ['Part 3', '两个坑（回放+目标网络）', '169-234', '理解为什么需要'],
        ['Part 4', '完整 DQN 边跑边看', '241-405', '跑 200 局看过程'],
        ['Part 5', 'QL vs DQN 同台对比', '411-530', '看 Q 值差异'],
        ['Part 6', '一步更新完整追踪', '536-650', '最重要！'],
    ]
)

brk()

# ================================================================
h('导入与环境（1-57 行）', 1)

explain(1, "#!/usr/bin/env python3", 'Shebang 行，告诉系统用 Python 3 运行。')
explain(2, "# -*- coding: utf-8 -*-", '文件编码声明，支持中文。')
explain(4, '"""DQN 逐行拆解 ..."""', '文件说明。核心卖点：从 Q-learning 出发，改一行就是 DQN。')
explain(17, 'import numpy as np', 'NumPy。和 Q-learning 一样，用来处理数组。')
explain(18, 'import torch', 'PyTorch。这是 DQN 和 QL 的根本区别——QL 不需要深度学习框架。')
explain(19, "import torch.nn as nn", '神经网络模块。Linear（全连接层）、ReLU、MSELoss。')
explain(20, "import torch.optim as optim", '优化器模块。Adam 代替了 QL 的手动 lr 更新。')
explain(21, 'import random', '随机数，用来做 ε-贪心探索。')
explain(22, 'import os', '路径操作。')
explain(24, "RESULTS_DIR = r'...'", '结果保存目录。')
explain(29, 'SIZE = 4', '网格大小 4×4。')
explain(30, 'N_STATES = SIZE * SIZE', '状态数 = 16。')
explain(31, 'N_ACTIONS = 4', '动作数 = 4。')
explain(32, 'GOAL = 15', '终点编号。')
explain(33, 'TRAP = 5', '陷阱编号。')
explain(34, 'GAMMA = 0.9', '折扣因子。')
explain(37, 'ACTION_DELTA = [(-1,0), (1,0), (0,-1), (0,1)]', '四个方向的坐标变化：↑ ↓ ← →。')
explain(38, "ACTION_NAMES = ['↑', '↓', '←', '→']", '动作编号对应的箭头符号，打印策略时用。')
explain(40, 'def is_valid(r, c):', '检查坐标是否在网格内。')
explain(43, 'def step(s, a):', '环境交互函数。和 Q-learning 完全一样。80% 走对方向，20% 滑走。')

brk()

# ================================================================
h('Part 2：核心改变 — Q 表换成神经网络（85-163 行）', 1)

p('这是 DQN 和 Q-learning 的第一个核心区别。理解了这个，DQN 就懂了一半。', bold=True, sz=12, color=RGBColor(0x1F,0x3A,0x5F))

explain(85, 'class TinyDQN(nn.Module):', '定义神经网络类，继承 PyTorch 的 nn.Module。和 Q 表的作用完全一样：输入状态 s，输出 4 个 Q 值。')
explain(95, 'def __init__(self):', '构造函数。初始化网络层。')
explain(96, 'super().__init__()', '调用父类构造函数。PyTorch 要求必须调用。')
explain(98, 'self.fc1 = nn.Linear(16, 32)', '第一层全连接：16 维输入 → 32 维隐藏。权重矩阵 [32×16] = 512 个参数。')
explain(99, 'self.relu = nn.ReLU()', 'ReLU 激活函数。把负数变成 0，引入非线性。')
explain(100, 'self.fc2 = nn.Linear(32, 4)', '第二层全连接：32 维 → 4 维输出（4 个 Q 值）。权重矩阵 [4×32] + bias [4] = 132 个参数。')
p('总参数 = 16×32 + 32 + 32×4 + 4 = 644。对比 Q 表只有 16×4 = 64 个。', sz=10, indent=1, color=RGBColor(0x1F,0x3A,0x5F))

explain(102, 'def forward(self, x):', '前向传播：输入 x → fc1 → ReLU → fc2 → 输出 4 个 Q 值。')
explain(104, 'return self.fc2(self.relu(self.fc1(x)))', '一行代码完成前向传播。x 流过各层，最终输出 4 个 Q 值。')
p('这和 Q 表的对应关系：net(格子0) → [Q↑, Q↓, Q←, Q→] 和 Q[0] → [Q↑, Q↓, Q←, Q→] 输出格式完全一样。', sz=10, indent=1, color=RGBColor(0x1F,0x3A,0x5F))

explain(106, 'def state_to_onehot(s):', '把状态编号（0-15）转成 one-hot 向量。为什么不用数字 0-15 直接输入？因为格子 3 不是格子 2 的"一半"，one-hot 告诉网络每个状态是独立的类别。')
explain(114, 'x = torch.zeros(16)', '创建长度为 16 的全零张量。')
explain(115, 'x[s] = 1.0', '把第 s 个位置设为 1。比如 s=3 → [0,0,0,1,0,0,...]。')
explain(116, 'return x.unsqueeze(0)', '加 batch 维度 [16] → [1, 16]。PyTorch 要求 batch 维度。')

explain(119, 'def part2_qtable_to_network():', '打印 Q 表和网络输出，直观对比两种方式。核心结论：输出格式完全一样，都是 [Q↑, Q↓, Q←, Q→]，不同在于内部存储方式。')

brk()

# ================================================================
h('Part 3：两个大坑（169-234 行）', 1)

p('如果把 Q 表直接换成网络，按 Q-learning 的方式训练，会出两个问题：', bold=True, sz=12)

h('坑 1：数据相关性（181-203 行）', 2)
p('Q-learning 每次只改 Q 表的一个格子，顺序数据不受影响。但网络一次更新影响所有参数，如果数据是顺序相关的（相邻几步高度相关），网络会过拟合到最近几步。', indent=0)
p('解决：经验回放。把经验存到缓冲区，训练时随机采样，打乱顺序。', bold=True, sz=10, color=RGBColor(0x1F,0x3A,0x5F), indent=0)

h('坑 2：目标不稳定（207-228 行）', 2)
p('target = r + γ·max Q(sp)。在 Q 表里改一个格子对其他格子影响小。在网络里一次梯度下降可能改变所有状态的 Q 值输出，导致 target 跟着变——"追自己的尾巴"。', indent=0)
p('解决：目标网络。再复制一个网络专门算 target，它不频繁更新。', bold=True, sz=10, color=RGBColor(0x1F,0x3A,0x5F), indent=0)

explain(231, "print('结论：DQN = Q-learning + 神经网络 + 经验回放 + 目标网络')", '核心公式。后面两个都是因为"用了神经网络"才需要的新技术。')

brk()

# ================================================================
h('Part 4：完整 DQN 边跑边看（241-405 行）', 1)

explain(241, 'def dqn_verbose(episodes=200):', '主训练函数。和 Q-learning 的框架完全一样：ε-贪心 → 执行 → 更新。只跑了 200 局（不是 5000），因为加了大量打印，只是为了让你看过程。')

explain(249, 'q_net = TinyDQN()', '创建在线网络。负责预测 Q 值，每步更新。')
explain(250, 'target_net = TinyDQN()', '创建目标网络。负责算 target，不频繁更新。')
explain(251, 'target_net.load_state_dict(q_net.state_dict())', '初始时两个网络参数一样。')

explain(253, 'optimizer = optim.Adam(q_net.parameters(), lr=0.01)', 'Adam 优化器。负责根据梯度更新 q_net 的参数。QL 没有这个——QL 直接改 Q 表的一个格子。')
explain(254, 'loss_fn = nn.MSELoss()', '均方误差损失函数。MSE = (Q_current - target)²。')

explain(256, 'replay_buffer = []', '经验回放缓冲区。QL 没有这个。')
explain(257, 'BUFFER_SIZE = 1000', '缓冲区容量。')
explain(258, 'BATCH_SIZE = 8', '每次训练从 buffer 随机抽 8 条。凑一批训练，比单条训练稳定。')

explain(260, 'epsilon = 1.0', '探索率。和 QL 一样。')

# Training loop main section
explain(265, 'for ep in range(1, episodes + 1):', '主循环。从第 1 局到 200 局。')

explain(286, 'if random.random() < epsilon:', 'ε-贪心选动作。和 QL 完全一样。')
explain(287, 'a = random.randint(0, N_ACTIONS - 1)', '探索：随机选一个动作。')
explain(290, 'with torch.no_grad():', '利用：用在线网络算 Q 值。torch.no_grad() 告诉 PyTorch 不用算梯度，省内存。')
explain(291, 'q_values = q_net(state_to_onehot(s))', '前向传播算 Q 值。QL 这一步是 q_values = Q[s]。')
explain(292, 'a = int(torch.argmax(q_values).item())', 'argmax 选最大的 Q 值。和 QL 一样。')

explain(295, 'sp, reward, done = step(s, a)', '执行动作。和 QL 完全一样。')

explain(298, 'replay_buffer.append((s, a, reward, sp, done))', '存经验到缓冲区。QL 没有这一步。')
explain(299, 'if len(replay_buffer) > BUFFER_SIZE:', '缓冲区满了就删最旧的。')

explain(305, 'if len(replay_buffer) >= BATCH_SIZE:', '经验够 8 条才开始训练。')

explain(306, 'batch = random.sample(replay_buffer, BATCH_SIZE)', '从缓冲区随机抽 8 条——经验回放的核心！')

explain(323, 'with torch.no_grad():', '算 target 时不计算梯度。用目标网络。')
explain(324, 'next_q = target_net(next_states)', '目标网络前向传播，算出下个状态的 Q 值。')
explain(325, 'max_next_q = torch.max(next_q, dim=1).values', '取每个状态的 max Q。和 QL 的 np.max(Q[sp]) 一样。')
explain(326, 'td_targets = rewards + GAMMA * max_next_q * (1 - dones)', '算 TD target。公式和 QL 完全一样！r + γ·max Q(sp)。')

explain(329, 'current_q = q_net(states)', '用在线网络算当前状态的 Q 值。')
explain(330, "current_q_a = current_q.gather(1, actions.unsqueeze(1)).squeeze()", '从 4 个 Q 值中挑出实际执行的那个动作对应的 Q 值。')

explain(332, 'loss = loss_fn(current_q_a, td_targets)', '算 loss = MSE(Q(s,a), target) — 和 QL 的 td_error 本质一样。')
explain(334, 'optimizer.zero_grad()', '清空上轮的梯度。')
explain(335, 'loss.backward()', '反向传播！计算每个参数的梯度。—— 这是 DQN 的灵魂一行。')
explain(336, 'optimizer.step()', '用梯度更新所有 644 个参数。QL 只改一个格子。')

explain(366, 'epsilon = max(0.01, epsilon * 0.995)', 'ε 衰减。和 QL 一样。')

explain(369, 'if ep % 20 == 0:', '每 20 局更新一次目标网络。')
explain(370, "target_net.load_state_dict(q_net.state_dict())", '把在线网络的参数复制给目标网络。保证 target 相对稳定。')

brk()

# ================================================================
h('Part 6：一步更新完整追踪（536-650 行）', 1)

p('这是整个文件最重要的部分。手动追踪一次 loss.backward() 的完整流程。', bold=True, sz=12, color=RGBColor(0x1F,0x3A,0x5F))

explain(536, 'def part6_peek_inside():', '用极简的设定，手动追踪 DQN 的一步更新。这是理解 DQN 本质最快的方式。')

explain(543, 'q_net = TinyDQN()', '创建在线网络。')
explain(544, 'target_net = TinyDQN()', '创建目标网络。')
explain(545, 'target_net.load_state_dict(q_net.state_dict())', '两个网络初始参数相同。')

explain(552, 's, a, r, sp, done = 0, 3, 0.0, 1, False', '设定场景：状态 s=0，动作 a=3（→），走到 sp=1，奖励 r=0，没结束。')

# Step 1
explain(555, 'with torch.no_grad():', '查 Q 值，不计算梯度。和查 Q 表一样。')
explain(556, 'q_s = q_net(state_to_onehot(s)).numpy()[0]', '前向传播，得到格子 0 的 4 个 Q 值。QL 这一步是 Q[0]。')
explain(557, 'q_sp = q_net(state_to_onehot(sp)).numpy()[0]', '格子 1 的 4 个 Q 值。')
explain(558, 'q_target_sp = target_net(state_to_onehot(sp)).numpy()[0]', '用目标网络算格子 1 的 Q 值（算 target 时用）。')

# Step 2
explain(568, 'target = r + GAMMA * np.max(q_target_sp)', '算 target = r + γ·max Q_target(sp)。和 QL 的 target 公式一模一样！')
p('QL 这一步是 target = r + GAMMA * np.max(Q[sp])。唯一区别：QL 用 Q 表，DQN 用目标网络。', sz=10, indent=1, color=RGBColor(0x1F,0x3A,0x5F))

# Step 3
explain(579, 'loss = (current_q_val - target) ** 2', '算 MSE loss = (Q - target)²。QL 算的是 td_error = target - Q。')
p('QL 用 td_error 直接改格子：Q[s][a] += lr × td_error', sz=10, indent=1)
p('DQN 用 loss 做梯度下降：loss.backward() + optimizer.step()', sz=10, indent=1, color=RGBColor(0x1F,0x3A,0x5F))

# Step 4
explain(588, "print('🔑 Q-learning: Q[0][3] += lr × (target - Q[0][3]) = 改 1 个格子')", 'QL 只改一个格子。')
explain(589, "print('🔑 DQN: loss.backward() + optimizer.step() = 改全部 644 个参数')", 'DQN 改全部参数。这是最核心的区别。')

# Gradient info
explain(613, 'loss_val.backward()', '反向传播！计算每个参数的梯度。这就是 DQN 的灵魂——自动求导。')
explain(616, "print('【梯度信息：每个参数要改多少】')", '打印每个参数层的梯度范数。')
explain(623, "print(f'{name}: 梯度范数 = {g_norm:.6f}（{grad.numel()} 个参数）')", '显示每层参数的梯度大小。')
explain(624, "print(f'总梯度范数 = {total_grad:.6f}')", '所有参数梯度的总和。')
explain(625, "print('644 个参数同时往【让 Q(格子0,右箭头) 更接近 target】的方向调整')", '这就是 DQN 的核心直觉。')

explain(627, 'optimizer.step()', '执行梯度下降，更新参数。')

# After update
explain(635, "print('【第 5 步：更新后的 Q 值】')", '看更新前后的 Q 值变化。')
explain(636, "print(f'Q(格子0, →) 更新前: {current_q_val:.4f}')", '更新前 Q 值。')
explain(637, "print(f'Q(格子0, →) 更新后: {q_s_after[3]:.4f}')", '更新后 Q 值。')
explain(638, "print(f'变化: {q_s_after[3] - current_q_val:+.4f}')", '变化量。')

explain(640, "print(f'注意格子{sp}的 Q 值也变了')", '参数共享的效果！更新格子 0 的 Q(→) 时，格子 1 的 Q(→) 也跟着变了。')
explain(646, "print('【这就是 DQN 的核心直觉】')", '一次更新 = 644 个参数同时微调 → 泛化 + 更难训练。')

brk()

# ================================================================
h('总结：QL vs DQN 核心对照表', 1)
tbl(
    ['功能', 'Q-learning', 'DQN'],
    [
        ['存 Q 值', 'Q = np.zeros((16, 4))', 'q_net = TinyDQN()'],
        ['参数数', '64（每个格子独立）', '644（参数共享）'],
        ['状态输入', '直接用数字 s', 'one-hot 向量 [16 维]'],
        ['选动作（利用）', 'a = argmax Q[s]', 'a = argmax q_net(x)'],
        ['算 target', 'r + γ·max Q[sp]', 'r + γ·max Q_target(sp)'],
        ['更新', 'Q[s][a] += lr × td_error', 'loss.backward() + optimizer.step()'],
        ['改多少个格子', '1 个', '644 个（全部）'],
        ['经验回放', '不需要', '需要（缓冲区 + 随机采样）'],
        ['目标网络', '不需要', '需要（每 20 局同步一次）'],
    ]
)
p('')
b('核心结论：DQN = Q-learning + 神经网络 + 经验回放 + 目标网络', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
p('公式没变，只是 Q 表换成了网络。但换了网络之后，必须加经验回放（解决数据相关性）和目标网络（解决目标不稳定）才能正常工作。')

brk()

h('附录：Part 6 输出解读', 1)
p('跑 python scripts/dqn_explained.py，拉到最底下看 Part 6 的输出，你应该看到类似这样的内容：', sz=10)
code('''【第 1 步：网络前向传播】
  net(格子0) = [-0.094  0.093  0.073  0.069]
    Q(格子0, →) = 0.0692  ← 我们要更新这个值

【第 2 步：算 target】
  target = 0.0 + 0.9 × 0.1357 = 0.1221

【第 3 步：算 loss】
  loss = (0.0692 - 0.1221)² = 0.0028

【第 4 步：梯度下降】
  Q-learning: Q[0][3] += lr × (target - Q[0][3]) = 改 1 个格子
  DQN:        loss.backward() + optimizer.step() = 改全部 644 个参数

  fc1.weight: 梯度范数 = 0.1829（512 个参数）
  fc1.bias:   梯度范数 = 0.1829（32 个参数）
  fc2.weight: 梯度范数 = 0.4162（128 个参数）
  fc2.bias:   梯度范数 = 0.4094（4 个参数）
  总梯度范数 = 1.1914

【第 5 步：更新后的 Q 值】
  Q(格子0, →) 更新前: -0.0205 → 更新后: 0.0566  (+0.0771)
  Q(格子1, →) 更新前: 0.2047  → 更新后: 0.2375  (+0.0328)  ← 格子1也被影响了！''')

p('')
note(f'生成日期：{datetime.date.today().isoformat()}')

# 保存
os.makedirs(OUT_DIR, exist_ok=True)
path = os.path.join(OUT_DIR, 'dqn_explained_逐行精讲.docx')
doc.save(path)
print(f'OK: {path}')
