#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 week10_dqn_gridworld.py 逐行精讲 doc"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def h(text,level=1):
    hd=doc.add_heading(text,level=level)
    for r in hd.runs: r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
def p(text,bold=False,sz=11,color=None,indent=0):
    pa=doc.add_paragraph()
    if indent: pa.paragraph_format.left_indent=Cm(indent)
    run=pa.add_run(text); run.font.name='微软雅黑'; run.font.size=Pt(sz); run.bold=bold
    if color: run.font.color.rgb=color
def b(text,sz=11,color=None): p(text,bold=True,sz=sz,color=color)
def bullet(text,lv=0):
    pa=doc.add_paragraph(text,style='List Bullet')
    pa.paragraph_format.left_indent=Cm(1.5+lv*0.8)
def code(lines,label=None):
    if label:
        pa=doc.add_paragraph(); run=pa.add_run(label); run.bold=True; run.font.size=Pt(10)
    for line in lines.split('\n'):
        pa=doc.add_paragraph()
        pa.paragraph_format.space_before=Pt(0); pa.paragraph_format.space_after=Pt(0); pa.paragraph_format.left_indent=Cm(1)
        run=pa.add_run(line); run.font.name='Consolas'; run.font.size=Pt(9.5); run.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'
    for i,hd in enumerate(headers):
        run=t.rows[0].cells[i].paragraphs[0].add_run(hd); run.bold=True; run.font.size=Pt(10); run.font.name='微软雅黑'
    for rd in rows:
        row=t.add_row()
        for c,txt in enumerate(rd):
            run=row.cells[c].paragraphs[0].add_run(txt); run.font.size=Pt(10); run.font.name='微软雅黑'
def brk(): doc.add_page_break()

# ======================== 封面 ========================
for _ in range(5): doc.add_paragraph('')
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t.add_run('DQN GridWorld\n逐行精讲'); run.font.size=Pt(24); run.bold=True; run.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t2.add_run('scripts/week10_dqn_gridworld.py 逐行解释'); run.font.size=Pt(13); run.font.color.rgb=RGBColor(0x66,0x66,0x66)
t3=doc.add_paragraph(); t3.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=t3.add_run(f'生成日期：{datetime.date.today().isoformat()}'); run.font.size=Pt(10); run.font.color.rgb=RGBColor(0x99,0x99,0x99)
brk()

# ======================== 整体说明 ========================
h('整体说明', 1)
p('这个文件是 DQN 在 4×4 GridWorld 上的实现。和 Q-learning 对比：')
tbl(
    ['', 'Q-learning (week10_qlearning)', 'DQN (week10_dqn)'],
    [
        ['Q 值存储', 'Q 表 np.zeros((16,4))', '神经网络 DQN() 644 个参数'],
        ['选动作', 'argmax Q[s]（查表）', 'argmax Q_theta(s)（网络前向传播）'],
        ['更新方式', 'Q[s][a] += lr * (target - Q[s][a])', 'loss.backward() 梯度下降'],
        ['数据相关性', 'Q 表不受影响', '用 经验回放 解决'],
        ['目标不稳定', 'Q 表每次改一格，影响小', '用 目标网络 解决'],
    ]
)
p('')
b('DQN = Q-learning + 神经网络 + 经验回放 + 目标网络', sz=12, color=RGBColor(0x1F,0x3A,0x5F))
p('这篇文章逐行解释这四个部分在代码里是怎么实现的。')
brk()

# ======================== 第1部分：导入和环境 ========================
h('第一部分：导入和环境定义（第 1-36 行）', 1)

code('''1:  #!/usr/bin/env python3
2:  # -*- coding: utf-8 -*-''')
p('文件头。指定 Python3 解释器和 UTF-8 编码。')
brk()

code('''4:  \"\"\"
5:  Week 10 - DQN: 用神经网络代替 Q 表
6:  ==================================
7:  对照 Q-learning：
8:    Q-learning: Q[s][a] ← 一张 16×4 的表格
9:    DQN:        Q_theta(s, a) ← 一个神经网络
10:
11:  环境：同一个 4x4 GridWorld
12:  \"\"\"''')
p('模块文档字符串。用大白话说明这个文件在干什么。')
brk()

code('''14: import numpy as np
15: import random
16: import os
17: import torch
18: import torch.nn as nn
19: import torch.optim as optim''')
b('第 14-19 行：导入库')
bullet('numpy: 数据处理（和 Q-learning 一样）')
bullet('random: 随机数（ε-贪心、滑走方向）')
bullet('os: 文件路径操作')
bullet('torch: PyTorch，神经网络的基石')
bullet('torch.nn: 网络层（Linear, ReLU, MSELoss）')
bullet('torch.optim: 优化器（Adam），用来更新网络参数')
brk()

code('''21: RESULTS_DIR = r'F:\\CLAUDE\\research\\ems-platform\\results' ''')
p('结果保存路径。训练好的 Q 表、训练记录都会存到这里。')
brk()

code('''24: SIZE = 4
25: N_STATES = SIZE * SIZE
26: N_ACTIONS = 4
27: GOAL_IDX = 15
28: TRAP_IDX = 5
29: GAMMA = 0.9''')
b('第 24-29 行：环境常量（和 Q-learning 完全一样）')
p('这 6 个常量和你的 Week 9 Part 5 GridWorld 完全一致。')
bullet('SIZE=4: 4×4 网格')
bullet('N_STATES=16: 16 个格子')
bullet('N_ACTIONS=4: ↑↓←→ 四个动作')
bullet('GOAL_IDX=15: 终点在 (3,3)')
bullet('TRAP_IDX=5: 陷阱在 (1,1)')
bullet('GAMMA=0.9: 折扣因子')
brk()

code('''30: ACTION_DELTA = [(-1, 0), (1, 0), (0, -1), (0, 1)]
31: ACTION_SYMBOLS = {0: '\\u2191', 1: '\\u2193', 2: '\\u2190', 3: '\\u2192'}''')
b('第 30-31 行：动作映射')
p('ACTION_DELTA: ↑=(行-1), ↓=(行+1), ←=(列-1), →=(列+1)')
p('ACTION_SYMBOLS: 打印策略时的箭头符号。')
brk()

code('''33: def is_valid(r, c):
34:     return 0 <= r < SIZE and 0 <= c < SIZE''')
b('第 33-34 行：边界检查')
p('判断坐标是否在网格内。撞墙时智能体留在原地。')
brk()

code('''36: def step(s, a):
37:     \"\"\"执行动作，返回 (s_next, reward, done)\"\"\"
38:     r, c = divmod(s, SIZE)''')
b('第 36-38 行：step 函数')
p('step(s, a) 接收状态 s 和动作 a，返回 (s_next, reward, done)。')
p('divmod(s, SIZE) 把一维索引 s 转换成 (行, 列) 坐标。')
brk()

code('''39:     if random.random() < 0.8:
40:         dr, dc = ACTION_DELTA[a]
41:     else:
42:         other = [i for i in range(N_ACTIONS) if i != a]
43:         dr, dc = ACTION_DELTA[random.choice(other)]''')
b('第 39-43 行：80%/20% 随机转移')
p('模拟真实环境的"不完美控制"：')
bullet('80% 概率：走选定的方向（如选了"→"就真的往右）')
bullet('20% 概率：随机滑到其他 3 个方向之一')
p('这个随机性和你的 Week 9 Part 5 P 表定义完全一致。')
p('Q-learning 和 DQN 不知道这个规则——它们只能通过"试了之后发现走歪了"来学习。')
brk()

code('''44:     nr, nc = r + dr, c + dc
45:     if not is_valid(nr, nc):
46:         nr, nc = r, c
47:     s_next = nr * SIZE + nc
48:     reward = 1.0 if s_next == GOAL_IDX else (-1.0 if s_next == TRAP_IDX else 0.0)
49:     done = (s_next == GOAL_IDX or s_next == TRAP_IDX)
50:     return s_next, reward, done''')
b('第 44-50 行：执行动作')
bullet('第 44 行：新坐标 = 原坐标 + 方向偏移')
bullet('第 45-46 行：撞墙就留在原地')
bullet('第 47 行：新坐标转回一维索引')
bullet('第 48 行：奖励——终点 +1，陷阱 -1，其他 0')
bullet('第 49 行：done——到达终点或陷阱就结束')
bullet('第 50 行：返回结果')

brk()

# ======================== 第2部分：DQN 神经网络 ========================
h('第二部分：DQN 神经网络定义（第 53-73 行）', 1)
p('这是 DQN 和 Q-learning 的第一个核心区别——用神经网络代替 Q 表。')

code('''53: class DQN(nn.Module):
54:     \"\"\"用神经网络代替 Q 表
55:     输入：状态 s（one-hot 编码，16 维）
56:     输出：4 个动作的 Q 值
57:     \"\"\"
58:     def __init__(self, state_dim=16, hidden=32):
59:         super().__init__()
60:         self.net = nn.Sequential(
61:             nn.Linear(state_dim, hidden),
62:             nn.ReLU(),
63:             nn.Linear(hidden, N_ACTIONS)
64:         )
65:     def forward(self, x):
66:         return self.net(x)''')
b('第 53-66 行：DQN 类')
p('这是整个 DQN 的核心——一个简单的全连接神经网络。')
p('')
b('第 53 行：class DQN(nn.Module)')
p('继承 PyTorch 的 nn.Module。这是 PyTorch 所有网络的基类，提供参数管理、设备迁移等功能。')
p('')
b('第 58-64 行：构造函数 __init__')
p('定义网络的"骨架"——有哪些层：')
bullet('nn.Linear(16, 32): 全连接层，16 维输入（one-hot 状态）→ 32 维隐藏层')
bullet('nn.ReLU(): 激活函数，max(0, x)，引入非线性。没有 ReLU，两层线性网络等价于一层')
bullet('nn.Linear(32, 4): 全连接层，32 维 → 4 维输出（4 个动作的 Q 值）')

b('对比 Q 表：')
p('  Q 表: 存储 16×4 = 64 个数值')
p('  这个网络: W1(16×32=512) + b1(32) + W2(32×4=128) + b2(4) = 676 个参数')
p('  网络参数比 Q 表多 10 倍，但优势是：')
bullet('可以输入连续状态（Q 表只能处理离散格子索引）')
bullet('可以泛化——没见过的状态也能输出 Q 值')

p('')
b('第 65-66 行：forward')
p('定义前向传播——数据怎么流过网络。输入 x (16维向量) → fc1 → ReLU → fc2 → 4个Q值。')
p('在 PyTorch 中，调用 model(s) 会自动执行 forward。')
brk()

code('''68: def state_to_tensor(s):
69:     \"\"\"把状态 s（0-15）转成 one-hot 向量\"\"\"
70:     x = torch.zeros(16)
71:     x[s] = 1.0
72:     return x.unsqueeze(0)''')
b('第 68-72 行：状态编码')
p('为什么需要 one-hot？', bold=True)
p('状态 s 是一个数字（0-15）。如果直接把数字 7 输入网络，网络会以为"7 比 6 大"——')
p('但状态 6 和状态 7 之间没有大小关系，只是两个不同的格子。')
p('')
p('One-hot 编码解决了这个问题：把数字转成一个只有 1 位是 1 其他都是 0 的向量。')
p('')
code('''s=0: [1,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0]
s=5: [0,0,0,0,0,1,0,0,0,0,0,0,0,0,0,0]
s=15:[0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,1]''')
p('')
bullet('第 70 行：创建 16 维全零向量')
bullet('第 71 行：第 s 位设为 1')
bullet('第 72 行：unsqueeze(0) 增加 batch 维度 [1, 16]（PyTorch 网络要求输入有 batch 维）')

brk()

# ======================== 第3部分：DQN 训练 ========================
h('第三部分：DQN 训练函数（第 75-164 行）', 1)

code('''75: def dqn(episodes=5000, lr=0.01, epsilon_start=1.0,
76:         epsilon_end=0.01, epsilon_decay=0.998):''')
b('第 75-76 行：函数定义')
p('5 个参数：')
bullet('episodes=5000: 训练局数（和 Q-learning 一样）')
bullet('lr=0.01: 学习率（Q-learning 是 0.1，这里更小，因为网络对步长更敏感）')
bullet('epsilon_start=1.0: 初始探索率（全随机）')
bullet('epsilon_end=0.01: 最终探索率（1% 随机）')
bullet('epsilon_decay=0.998: 探索率衰减速度')
brk()

code('''77:     print(f'DQN \\u5f00\\u59cb\\u8bad\\u7ec3: {episodes} \\u5c40')
78:     print(f'  \\u795e\\u7ecf\\u7f51\\u7edc: 16\\u7ef4\\u8f93\\u5165 \\u2192 32\\u9690\\u85cf \\u2192 4\\u8f93\\u51fa')
79:     print(f'  Q \\u8868\\u53c2\\u6570: 16\\u00d74 = 64 \\u4e2a\\u683c\\u5b50')
80:     print(f'  \\u795e\\u7ecf\\u7f51\\u7edc\\u53c2\\u6570: 676 = {16*32+32+32*4+4} \\u4e2a\\u53c2\\u6570')
81:     print(f'  \\u5b66\\u4e60\\u7387 lr = {lr}')
82:     print()''')
b('第 77-82 行：打印信息')
p('打印 DQN 的配置信息，方便和 Q-learning 对比。')
brk()

code('''84:     q_network = DQN()
85:     target_network = DQN()
86:     target_network.load_state_dict(q_network.state_dict())
87:     optimizer = optim.Adam(q_network.parameters(), lr=lr)
88:     loss_fn = nn.MSELoss()''')
b('第 84-88 行：初始化网络和优化器')
p('')
b('第 84 行：q_network = DQN()')
p('在线网络（也叫主网络）。负责：')
bullet('选动作：argmax Q_online(s)')
bullet('预测 Q 值：用于算 loss')
bullet('每步都更新梯度')
p('')
b('第 85 行：target_network = DQN()')
p('目标网络。和在线网络结构一样，但参数更新慢。负责：')
bullet('计算 target = r + gamma * max Q_target(sp)')
bullet('参数冻结，不频繁更新')
p('')
b('第 86 行：target_network.load_state_dict(...)')
p('初始化时让两个网络参数一致。之后在线网络更新，目标网络隔一段时间才复制一次。')
p('')
b('第 87 行：optimizer = optim.Adam(...)')
p('Adam 优化器——Week 9 Part 4 用过。负责梯度下降更新网络参数。')
p('')
b('第 88 行：loss_fn = nn.MSELoss()')
p('均方误差损失: (y_pred - y_target)²。DQN 的损失函数就是让 Q_theta(s,a) 逼近 target。')
brk()

code('''90:     replay_buffer = []
91:     BUFFER_SIZE = 10000
92:     BATCH_SIZE = 32''')
b('第 90-92 行：经验回放缓冲区')
p('这是 DQN 的核心改进之一，Q-learning 不需要这个。')
bullet('replay_buffer: Python 列表，存经验 (s, a, r, sp, done)')
bullet('BUFFER_SIZE=10000: 最多存 10000 条经验，满了就删最旧的')
bullet('BATCH_SIZE=32: 每次训练从缓冲区随机抽 32 条')
p('')
b('类比：')
p('  Q-learning 是"当场学了就忘"——走一步更新一次 Q 表，然后这步的经验就不要了。')
p('  DQN 是"先记笔记，回头随机翻笔记学习"——经验存起来，以后还能反复学。')
brk()

code('''94:     epsilon = epsilon_start
95:     episode_rewards = []
96:     episode_steps = []''')
b('第 94-96 行：初始化训练变量')
bullet('epsilon: 探索率，初始 1.0，逐渐衰减到 0.01')
bullet('episode_rewards: 记录每局总奖励，画训练曲线用')
bullet('episode_steps: 记录每局步数')
brk()

code('''98:     for ep in range(1, episodes + 1):
99:         s = 0
100:         total_reward = 0
101:         steps = 0''')
b('第 98-101 行：训练主循环')
p('外层循环：每局从起点 (0,0) 出发，走到终点或陷阱为止。')
brk()

code('''103:         while True:
104:             if random.random() < epsilon:
105:                 a = random.randint(0, N_ACTIONS - 1)
106:             else:
107:                 with torch.no_grad():
108:                     q_values = q_network(state_to_tensor(s))
109:                     a = int(torch.argmax(q_values).item())''')
b('第 103-109 行：ε-贪心选动作（和 Q-learning 一样，但查的是网络不是表）')
p('')
b('Q-learning 版本：')
code('  a = argmax Q[s]           # 查 Q 表的一行')
p('')
b('DQN 版本：')
p('这一段的执行流程：')
bullet('第 104 行：随机数 < epsilon → 探索，否则利用')
bullet('第 105 行：探索时随机选一个动作')
bullet('第 107 行：with torch.no_grad() 不追踪梯度（推理模式，省内存）')
bullet('第 108 行：q_network(state_to_tensor(s)) 网络前向传播，得到 4 个 Q 值')
bullet('第 109 行：argmax 选 Q 值最大的动作')
p('')
b('核心区别：', color=RGBColor(0xC0,0x39,0x2B))
p('  Q-learning: Q[s] → 查表，O(1) 时间，不涉及计算')
p('  DQN: q_network(s) → 矩阵乘法 + ReLU + 矩阵乘法，需要计算')
p('  但 DQN 的优势是：输入 s 可以是任意值（连续状态），Q 表只能查预设好的离散索引')
brk()

code('''111:             s_next, reward, done = step(s, a)
112:             replay_buffer.append((s, a, reward, s_next, done))
113:             if len(replay_buffer) > BUFFER_SIZE:
114:                 replay_buffer.pop(0)''')
b('第 111-114 行：执行并存储经验')
bullet('第 111 行：step(s, a) 执行动作，返回下一步和奖励（和 Q-learning 一样）')
bullet('第 112 行：把经验 (s,a,r,sp,done) 存到缓冲区')
bullet('第 113-114 行：缓冲区满了（>10000），删最旧的经验')
p('Q-learning 没有这一步——它走完一步就立刻更新 Q 表，然后丢弃经验。')
brk()

code('''116:             if len(replay_buffer) >= BATCH_SIZE:
117:                 batch = random.sample(replay_buffer, BATCH_SIZE)''')
b('第 116-117 行：从缓冲区采样')
p('缓冲区至少有 32 条经验后，才开始训练。')
p('random.sample 随机抽 32 条——这就是经验回放的核心。')
p('')
b('为什么不能按顺序学？')
p('按顺序学：s0→s4→s5(Trap)→s0→s1... 相邻步高度相关')
p('随机采样：打破相关性，更像监督学习的独立同分布假设')
brk()

code('''119:                 states = torch.zeros(BATCH_SIZE, 16)
120:                 next_states = torch.zeros(BATCH_SIZE, 16)
121:                 actions = torch.zeros(BATCH_SIZE, dtype=torch.long)
122:                 rewards = torch.zeros(BATCH_SIZE)
123:                 dones = torch.zeros(BATCH_SIZE)''')
b('第 119-123 行：准备 batch 张量')
p('把 32 条经验拆成 5 个 PyTorch 张量（网络需要 tensor 格式）：')
bullet('states: [32, 16] — 32 条经验的起始状态')
bullet('next_states: [32, 16] — 32 条经验的下一个状态')
bullet('actions: [32] — 选的动作索引')
bullet('rewards: [32] — 获得的奖励')
bullet('dones: [32] — 是否结束')
brk()

code('''125:                 for i, (s, a, r, ns, d) in enumerate(batch):
126:                     states[i] = state_to_tensor(s)
127:                     next_states[i] = state_to_tensor(ns)
128:                     actions[i] = a
129:                     rewards[i] = r
130:                     dones[i] = 1.0 if d else 0.0''')
b('第 125-130 行：填充 batch')
p('把采样出的 32 条经验逐一转换成张量。')
p('dones 用 1.0/0.0 表示结束/未结束，后面算 target 时会用到。')
brk()

code('''132:                 with torch.no_grad():
133:                     next_q = target_network(next_states)
134:                     max_next_q = torch.max(next_q, dim=1).values
135:                     td_targets = rewards + GAMMA * max_next_q * (1 - dones)''')
b('第 132-135 行：用目标网络计算 target')
p('这是 DQN 的核心更新公式，也是和 Q-learning 的第二个关键区别。')
p('')
b('用目标网络，不用在线网络：')
p('第 132 行：torch.no_grad() — 目标网络不计算梯度，冻结')
p('第 133 行：target_network(next_states) — 目标网络算下一步的 Q 值')
p('第 134 行：torch.max(next_q, dim=1).values — 取每个状态的 max Q')
p('第 135 行：td_targets = r + gamma * max(Q(sp)) — Q-learning 的 target 公式')
p('  (1 - dones)：终止状态的 target 就是 r 本身，不加未来回报')
p('')
b('对比 Q-learning：', color=RGBColor(0x1F,0x3A,0x5F))
code('''Q-learning: target = r + gamma * max(Q[sp])     # 用 Q 表，查表
DQN:        target = r + gamma * max(Q_target(sp))  # 用目标网络，前向传播''')
p('公式一模一样，只是 Q 表换成了目标网络。')
brk()

code('''137:                 current_q = q_network(states)
138:                 current_q = current_q.gather(1, actions.unsqueeze(1)).squeeze()''')
b('第 137-138 行：获取当前 Q 值')
p('第 137 行：q_network(states) → 输出 [32, 4]，每个状态 4 个动作的 Q 值')
p('第 138 行：.gather(1, actions) → 从 4 个 Q 值中挑出实际选的那个动作的 Q 值')
p('')
p('举例：', bold=True)
code('''states[0] = 格子3 → q_network输出 [0.2, 0.5, 0.3, 0.8]  ← 4个动作的Q值
actions[0] = 3（→） → current_q[0] = 0.8  ← 只取"→"的Q值''')
brk()

code('''140:                 loss = loss_fn(current_q, td_targets)''')
b('第 140 行：计算损失')
p('loss = MSE(current_q, td_targets) — 让网络输出的 Q 值逼近 target')
p('')
b('对比 Q-learning：', color=RGBColor(0x1F,0x3A,0x5F))
code('''Q-learning:  Q[s][a] += lr * (td_target - Q[s][a])
            # 直接把 Q 表的一个格子往 target 方向拉一点

DQN:        loss = MSE(Q_theta(s,a), td_target)
            loss.backward()  # 算出所有参数的梯度
            optimizer.step() # 更新所有参数
            # 一次更新改变整个网络的输出''')
brk()

code('''142:                 optimizer.zero_grad()
143:                 loss.backward()
144:                 optimizer.step()''')
b('第 142-144 行：梯度下降更新网络')
p('这三行是 PyTorch 的标准训练三步曲：')
bullet('第 142 行：optimizer.zero_grad() — 清零梯度。PyTorch 默认累积梯度，不清零会导致上一次的梯度叠加到这一次。')
bullet('第 143 行：loss.backward() — 反向传播，自动计算所有参数对 loss 的梯度。这一行执行链式法则，从输出层到输入层逐层计算偏导。')
bullet('第 144 行：optimizer.step() — 用 Adam 优化器更新参数。θ = θ - lr · (梯度方向)')
p('')
b('和 Q-learning 的更新对比：', color=RGBColor(0x1F,0x3A,0x5F))
p('  Q-learning: 5 行代码，只改 1 个数值')
p('  DQN: 5 行代码，更新 676 个参数')
p('  这就是 DQN 更慢但更能泛化的原因——每次更新影响整个 Q 函数', bold=True)
brk()

code('''146:             total_reward += reward
147:             steps += 1
148:             s = s_next
149:             if done:
150:                 break''')
b('第 146-150 行：更新状态')
p('记录奖励和步数，移动到下一个状态。如果到达终点或陷阱，结束本局。')
brk()

code('''152:         episode_rewards.append(total_reward)
153:         episode_steps.append(steps)
154:         epsilon = max(epsilon_end, epsilon * epsilon_decay)''')
b('第 152-154 行：记录训练数据 + ε 衰减')
p('第 154 行：epsilon 每局乘以 0.998（衰减），但最低 0.01。')
bullet('前期（ε≈1.0）：几乎全随机探索，智能体到处乱走')
bullet('中期（ε≈0.3-0.1）：逐渐利用学到的知识，但还在探索')
bullet('后期（ε≈0.01）：几乎只利用，偶尔随机一下')
brk()

code('''156:         if ep % 200 == 0:
157:             target_network.load_state_dict(q_network.state_dict())''')
b('第 156-157 行：更新目标网络（硬更新）')
p('这是 DQN 的第二个核心改进——目标网络。')
p('')
bullet('每 200 局，把在线网络的参数复制给目标网络')
bullet('其余时间目标网络冻结不动')
p('')
b('为什么需要目标网络？', color=RGBColor(0xC0,0x39,0x2B))
p('如果没有目标网络，第 135 行会用在线网络算 target：')
code('''# 错误做法：用同一个网络算 target
td_targets = r + gamma * max(q_network(next_states))
#       ↑                                   ↑
#   更新 Q_theta          依赖 Q_theta 的值
#   target 和预测用的是同一个网络！靶子在动！''')
p('')
p('每 200 局复制一次，相当于把靶子固定下来，让在线网络去瞄准固定靶。', bold=True)
brk()

code('''159:             avg_reward = np.mean(episode_rewards[-100:])
160:             avg_steps = np.mean(episode_steps[-100:])
161:             print(f'  \\u7b2c {ep:4d}/{episodes} \\u5c40 | \\u03b5={epsilon:.3f} | '
162:                   f'\\u5e73\\u5747\\u5956\\u52b1={avg_reward:+.4f} | \\u5e73\\u5747\\u6b65\\u6570={avg_steps:.1f}')''')
b('第 159-162 行：打印训练进度')
p('每 200 局打印一次：')
bullet('avg_reward: 最近 100 局平均奖励。正数越大说明越少掉陷阱、越多到终点。')
bullet('avg_steps: 最近 100 局平均步数。越小说明路径越短。')
brk()

code('''164:     policy = np.zeros(N_STATES, dtype=int)
165:     with torch.no_grad():
166:         for s in range(N_STATES):
167:             q = q_network(state_to_tensor(s))
168:             policy[s] = int(torch.argmax(q).item())
169:     return policy, episode_rewards, episode_steps''')
b('第 164-169 行：提取策略并返回')
p('训练完成后，对每个状态 s 跑一次网络，取 Q 值最大的动作作为策略。')
p('')
b('和 Q-learning 提取策略对比：', color=RGBColor(0x1F,0x3A,0x5F))
code('''Q-learning: policy[s] = argmax Q[s]         # 查 Q 表
DQN:        policy[s] = argmax Q_network(s)  # 跑一次网络''')
p('结果一样，只是数据来源不同。')

brk()

# ======================== 第4部分：主程序 ========================
h('第四部分：主程序（第 171-227 行）', 1)

code('''171: def save_results(policy, rewards, steps):
172:     os.makedirs(RESULTS_DIR, exist_ok=True)
173:     path = os.path.join(RESULTS_DIR, 'week10_dqn_training.csv')
174:     with open(path, 'w', encoding='utf-8') as f:
175:         f.write('episode,reward,steps\\n')
176:         for i, (r, s) in enumerate(zip(rewards, steps)):
177:             f.write(f'{i+1},{r:.4f},{s}\\n')
178:     print(f'\\u8bad\\u7ec3\\u6570\\u636e\\u5df2\\u4fdd\\u5b58: {path}')''')
b('第 171-178 行：保存训练结果')
p('把每局的奖励和步数存成 CSV 文件，可以拖到 Excel 里画训练曲线。')

brk()

code('''180: if __name__ == \\'__main__\\':''')
b('第 180 行：入口')
p('当直接运行这个文件时执行以下代码；被 import 时不执行。')

brk()

code('''181:     print('=' * 65)
182:     print('  DQN - \\u7528\\u795e\\u7ecf\\u7f51\\u7edc\\u4ee3\\u66ff Q \\u8868')
183:     print('  \\u73af\\u5883: 4x4 GridWorld, \\u548c Q-learning \\u4e00\\u6a21\\u4e00\\u6837')
184:     print('  \\u533a\\u522b:')
185:     print('    Q-learning: Q[s][a] \\u2190 16\\u00d74 = 64 \\u53c2\\u6570\\u7684\\u8868\\u683c')
186:     print('    DQN:        Q_theta(s,a) \\u2190 676 \\u53c2\\u6570\\u7684\\u795e\\u7ecf\\u7f51\\u7edc')
187:     print('  \\u5176\\u4ed6\\uff08\\u03b5-\\u8d2a\\u5fc3\\u3001\\u516c\\u5f0f\\u3001\\u5956\\u52b1\\uff09\\u5b8c\\u5168\\u4e00\\u6837')
188:     print('=' * 65)''')
b('第 181-188 行：打印文件说明')
p('用最直接的方式告诉读者：这个文件在干什么，和 Q-learning 什么关系。')

brk()

code('''190:     policy, rewards, steps = dqn(
191:         episodes=5000,
192:         lr=0.01,
193:         epsilon_start=1.0,
194:         epsilon_end=0.01,
195:         epsilon_decay=0.998
196:     )''')
b('第 190-196 行：训练')
p('调用 dqn() 函数开始训练。5000 局，和 Q-learning 一样。')

brk()

code('''198:     print()
199:     print('=' * 50)
200:     print('DQN \\u8bad\\u7ec3\\u7ed3\\u679c')
201:     print('=' * 50)''')
b('第 198-201 行：打印结果标题')

brk()

code('''204:     for r in range(SIZE):
205:         row = \\'  |\\'
206:         for c in range(SIZE):
207:             s = r * SIZE + c
208:             if s == GOAL_IDX:
209:                 row += \\' G |\\'
210:             elif s == TRAP_IDX:
211:                 row += \\' X |\\'
212:             else:
213:                 row += f\\' {ACTION_SYMBOLS[policy[s]]} |\\'
214:         print(row)
215:         print(\\'  \\' + \\'-\\' * 19)''')
b('第 204-215 行：打印最优策略网格')
p('和 Q-learning 一样，用箭头画出最优策略。')

brk()

code('''217:     print(f\\'\\u6700\\u540e 100 \\u5c40\\u5e73\\u5747\\u5956\\u52b1: {np.mean(rewards[-100:]):+.4f}\\')
218:     print(f\\'\\u6700\\u540e 100 \\u5c40\\u5e73\\u5747\\u6b65\\u6570: {np.mean(steps[-100:]):.1f}\\')
219:     save_results(policy, rewards, steps)''')
b('第 217-219 行：打印训练统计 + 保存结果')

brk()

code('''221:     print()
222:     print('=' * 65)
223:     print('  \\u4e09\\u4e2a\\u65b9\\u6cd5\\u5bf9\\u6bd4:')
224:     print('    \\u503c\\u8fed\\u4ee3 (DP):    64 \\u4e2a Q \\u8868\\u683c\\u5b50,   133 \\u8f6e\\u76f4\\u63a5\\u7b97\\u5b8c')
225:     print('    Q-learning:     64 \\u4e2a Q \\u8868\\u683c\\u5b50,   5000 \\u5c40\\u8bd5\\u51fa\\u6765')
226:     print('    DQN:            644 \\u4e2a\\u7f51\\u7edc\\u53c2\\u6570,   5000 \\u5c40\\u8bd5\\u51fa\\u6765')
227:     print()
228:     print('  DQN \\u5728\\u8fd9\\u4e2a\\u5c0f\\u95ee\\u9898\\u4e0a\\u6ca1\\u6709\\u4f18\\u52bf\\uff08\\u751a\\u81f3\\u66f4\\u6162\\uff09')
229:     print('  \\u4f46\\u771f\\u5b9e\\u95ee\\u9898\\u72b6\\u6001\\u7a7a\\u95f4\\u5de8\\u5927\\uff08\\u56fe\\u50cf\\u3001\\u8fde\\u7eed\\u503c\\uff09')
230:     print('  Q \\u8868\\u5b58\\u4e0d\\u4e0b\\uff0c\\u53ea\\u80fd\\u7528\\u795e\\u7ecf\\u7f51\\u7edc')''')
b('第 221-230 行：三种方法对比总结')
p('用最简洁的方式总结三种方法的区别。')

# ======================== 附录：整体对比 ========================
brk()
h('附录：DQN vs Q-learning 代码逐行对照', 1)

p('这两段代码做的是同一件事——用 ε-贪心选动作，执行，更新：')
p('')
b('Q-learning（选动作 + 更新）：')
code('''# 选动作
a = argmax Q[s] if random() > epsilon else random_action()

# 执行
sp, r, done = step(s, a)

# 更新（直接改 Q 表）
Q[s][a] += lr * (r + gamma * max(Q[sp]) - Q[s][a])''')

b('DQN（选动作 + 更新）：')
code('''# 选动作（查网络，不是查表）
a = argmax q_network(s) if random() > epsilon else random_action()

# 执行
sp, r, done = step(s, a)
buffer.append((s,a,r,sp,done))  # 多了一步：存经验

# 更新（梯度下降，不是改表）
batch = sample(buffer, 32)       # 经验回放
target = r + gamma * max(target_network(sp))  # 目标网络
loss = MSE(q_network(s,a), target)
loss.backward()
optimizer.step()''')
p('')
b('框架完全一样，只是"查表"换成了"跑网络"，"直接改表"换成了"梯度下降"。', sz=12, color=RGBColor(0x1F,0x3A,0x5F))

# ======================== 保存 ========================
path = os.path.join(OUT_DIR, 'Week10_DQN_GridWorld_逐行精讲.docx')
doc.save(path)
print(f'OK: {path}')
