#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 Week9 Part4 MLP 功率预测 逐行解释文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Consolas'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 辅助函数
def add_code(text, indent=0):
    """添加代码行（灰色背景用缩进代替，实际上保持等宽即可）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # 设置等宽字体
    for run in p.runs:
        run.font.name = 'Consolas'
    run = p.add_run('    ' * indent + text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p

def add_comment(text):
    """注释说明行，绿色"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    return p

def add_section(title):
    doc.add_heading(title, level=2)

def add_text(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(11)
    return p

def add_line(lineno, code, explanation):
    """一行代码 + 逐行解释"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0)

    # 行号
    run = p.add_run(f'{lineno:4d}  ')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 代码
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    if code.strip().startswith('#'):
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    else:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # 解释
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run = p2.add_run(f'└─ {explanation}')
    run.font.name = '微软雅黑'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ══════════════════════════════════════════
# 封面
# ══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Week 9 · Part 4\nMLP 功率预测 — 逐行解释')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(f'scripts/week9_complete.py  Lines 166–269\n生成日期：{datetime.date.today().isoformat()}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ══════════════════════════════════════════
# 总体说明
# ══════════════════════════════════════════
add_section('本节在学什么')
add_text(
    'Part 4 的目标是用一个 MLP（多层感知机）做燃料电池功率预测。\n'
    '具体任务：给定过去 10 个时刻的 FC 功率值，预测下一个时刻的功率。\n\n'
    '这是一个经典的时序预测问题，用全连接网络（而非 RNN/LSTM）实现，'
    '目的是演示 PyTorch 的完整训练流程：数据准备 → 模型定义 → 训练 → 评估 → 可视化。'
)

add_section('代码整体结构')
add_text('整个 Part 4 包含 104 行代码，分为 5 个逻辑块：')
add_text('① 生成模拟功率数据 (L178-181)')
add_text('② 构建时序样本 (L184-201)')
add_text('③ 定义 MLP 模型 (L204-218)')
add_text('④ 训练循环 (L225-239)')
add_text('⑤ 评估 + 可视化 (L242-269)')

doc.add_page_break()

# ══════════════════════════════════════════
# 逐行解释
# ══════════════════════════════════════════
add_section('逐行代码解析')

# L166-175: 函数定义 + import
add_line(166, '# Part 4: MLP 功率预测（燃料电池功率预测）',
         '注释标记。本 Part 的内容是 MLP 功率预测。')
add_line(167, '# ══════════════════════════════════════════════════',
         '装饰分隔线，便于在脚本中快速定位。')
add_line(168, 'def part4_power_prediction():',
         '函数定义。整个 Part 4 封装在这个函数里。之所以不放在全局，是为了配合 --part 参数单独调用。')
add_line(169, '    """', '文档字符串开始。')
add_line(170, '    用 PyTorch MLP 做 FC 功率预测。', '一句话说明本函数的功能。')
add_line(171, '    任务：根据历史功率序列，预测下一时刻的燃料电池功率。', '具体任务描述：这是一个时序预测任务。')
add_line(172, '    """', '文档字符串结束。')
add_line(173, '    import torch', '导入 PyTorch 主库。虽然文件开头已经 import 过，但在函数内部再 import 是安全的，Python 会缓存。')
add_line(174, '    import torch.nn as nn', '导入神经网络模块，提供了 Linear、Module 等核心类。')
add_line(175, '    import torch.nn.functional as F', '导入函数式 API（relu 等激活函数），方便在 forward 中直接用 F.relu()。')

doc.add_paragraph('')
add_text('──────────────────── ① 生成模拟数据 ────────────────────')

add_line(178, '    np.random.seed(42)',
         '固定随机种子为 42，确保每次运行生成的模拟数据相同，结果可复现。')
add_line(179, '    t = np.linspace(0, 100, 1000)',
         '生成 1000 个时间点，均匀分布在 [0, 100] 区间。t 的值为：0, 0.1, 0.2, ..., 99.9。')
add_line(180, '    power = 30 + 15 * np.sin(0.1 * t) + 5 * np.sin(0.5 * t) + np.random.randn(1000) * 2',
         '生成模拟的燃料电池功率序列。公式拆解：'
         '30 是直流偏置（基础功率 30 kW）；'
         '15*sin(0.1*t) 是低频大振幅波动（模拟航行工况变化，周期约 63 步）；'
         '5*sin(0.5*t) 是高频小振幅波动（模拟海浪/负载扰动）；'
         'randn*2 是高斯噪声（模拟传感器噪声和随机扰动）。')
add_line(181, '    power = np.clip(power, 10, 80)',
         '将功率限制在 [10, 80] kW 范围内，模拟 FC 的实际工作区间——FC 不允许低于最低功率（否则效率低），也不能超过额定功率。')

doc.add_paragraph('')
add_text('──────────────────── ② 构建时序样本 ────────────────────')

add_line(184, '    def create_sequences(data, seq_len=10):',
         '定义辅助函数：将一维时间序列转换成 (输入序列, 目标值) 的样本对。'
         'seq_len=10 表示用过去 10 个时刻预测下一时刻。')
add_line(185, '        X, y = [], []',
         '初始化两个空列表：X 存输入序列，y 存要预测的目标值。')
add_line(186, '        for i in range(len(data) - seq_len):',
         '循环遍历数据，确保 i+seq_len 不越界。如果 data 有 N 个点，则生成 N-seq_len 个样本。')
add_line(187, '            X.append(data[i:i + seq_len])',
         '取第 i 到 i+seq_len-1 共 seq_len 个点作为输入特征。')
add_line(188, '            y.append(data[i + seq_len])',
         '取第 i+seq_len 这个点（紧接在输入序列后面的值）作为预测目标。')
add_line(189, '        return np.array(X), np.array(y)',
         '返回 numpy 数组。X 的形状为 (样本数, seq_len)，y 的形状为 (样本数,)。')

add_line(191, '    SEQ_LEN = 10',
         '定义超参数：用过去 10 步预测下一步。这个值可以调——更大会给模型更多历史信息但维度更高。')
add_line(192, '    X, y = create_sequences(power, SEQ_LEN)',
         '调用函数，将 1000 个点的功率序列变成 990 个样本（1000-10=990）。'
         'X 的形状：(990, 10)，y 的形状：(990,)。')
add_line(193, '    split = int(0.8 * len(X))',
         '计算训练集切分点：80% 的数据用于训练，20% 用于测试。990 × 0.8 = 792。')
add_line(194, '    X_train, X_test = X[:split], X[split:]',
         '按 8:2 切分输入数据：前 792 个样本训练，后 198 个样本测试。')
add_line(195, '    y_train, y_test = y[:split], y[split:]',
         '对应切分目标值。')

add_line(197, '    X_train_t = torch.tensor(X_train, dtype=torch.float32).unsqueeze(-1)',
         '将 numpy 数组转 PyTorch Tensor，类型 float32。'
         'unsqueeze(-1) 在最后加一维：从 (792, 10) 变成 (792, 10, 1)。'
         '为什么要加这一维？因为 nn.Linear 期望输入为 (batch, features)，'
         'MLP 会把最后一维作为特征维度处理，这里加维是为了保持灵活性。'
         '实际上对于全连接网络，后面在 forward 中会展开，所以 (792, 10) 就够了，'
         '但保持 3D 形状便于后续扩展到 CNN 等需要通道维度的模型。')
add_line(198, '    y_train_t = torch.tensor(y_train, dtype=torch.float32).unsqueeze(-1)',
         '训练目标值：(792,) → (792, 1)。加维是为了匹配模型输出的形状。')
add_line(199, '    X_test_t = torch.tensor(X_test, dtype=torch.float32).unsqueeze(-1)',
         '测试输入：(198, 10) → (198, 10, 1)。')
add_line(200, '    y_test_t = torch.tensor(y_test, dtype=torch.float32).unsqueeze(-1)',
         '测试目标值：(198,) → (198, 1)。')
add_line(201, '    print(f"Train: {X_train_t.shape}, Test: {X_test_t.shape}")',
         '打印形状确认：Train: torch.Size([792, 10, 1]), Test: torch.Size([198, 10, 1])。')

doc.add_paragraph('')
add_text('──────────────────── ③ 定义 MLP 模型 ────────────────────')

add_line(204, '    class PowerPredictor(nn.Module):',
         '定义模型类，继承 PyTorch 的 nn.Module。所有 PyTorch 模型都应该继承这个基类。')
add_line(205, '        def __init__(self, seq_len, hidden=64):',
         '构造函数。接受两个参数：seq_len（输入维度，即用多少步历史）和 hidden（隐藏层神经元数，默认 64）。')
add_line(206, '            super().__init__()',
         '调用父类 nn.Module 的构造函数，这是必须的——nn.Module 内部会注册参数、管理状态。')
add_line(207, '            self.fc1 = nn.Linear(seq_len, hidden)',
         '第一层全连接：输入 seq_len 维 → 输出 hidden 维（64 维）。'
         'nn.Linear 内部有可训练的权重矩阵 W (64×10) 和偏置 b (64)。')
add_line(208, '            self.fc2 = nn.Linear(hidden, hidden)',
         '第二层全连接：64 维 → 64 维。深层网络比单层能学到更复杂的映射关系。')
add_line(209, '            self.fc3 = nn.Linear(hidden, 1)',
         '输出层：64 维 → 1 维，输出预测的功率值（一个标量）。')
add_line(210, '            self.dropout = nn.Dropout(0.1)',
         'Dropout 层：训练时以 10% 的概率随机将神经元置零，防止过拟合。'
         '相当于每次训练都在训练一个不同结构的子网络（集成学习的思想）。')

add_line(212, '        def forward(self, x):',
         '前向传播函数。这是 nn.Module 的核心：定义数据从输入到输出的计算过程。')
add_line(213, '            x = x.view(x.size(0), -1)',
         '将输入展平。x 的形状从 (batch, seq_len, 1) 变成 (batch, seq_len)。'
         'view() 是 PyTorch 的张量重塑方法，-1 表示自动推断该维度大小。'
         '这一步很关键：因为 nn.Linear 需要 2D 输入 (batch, features)。')
add_line(214, '            x = F.relu(self.fc1(x))',
         '第一层 → ReLU 激活。self.fc1(x) 做线性变换，F.relu() 将负值截断为 0，'
         '引入非线性。如果没有激活函数，多层线性 = 一层线性，就没有"深度"的意义了。')
add_line(215, '            x = self.dropout(x)',
         '第一层输出后加 Dropout，随机丢弃 10% 的神经元（只在训练时生效）。')
add_line(216, '            x = F.relu(self.fc2(x))',
         '第二层 → ReLU 激活。再经过一次线性变换 + 非线性。')
add_line(217, '            x = self.fc3(x)',
         '输出层：直接输出预测值（没有激活函数，因为这是回归任务——输出可以是任意实数）。')
add_line(218, '            return x',
         '返回预测结果。形状为 (batch, 1)。')

add_line(220, '    model = PowerPredictor(SEQ_LEN, hidden=64)',
         '实例化模型。输入维度 = 10（SEQ_LEN），隐藏层 = 64 个神经元。'
         '此时模型参数是随机初始化的。')
add_line(221, '    optimizer = torch.optim.Adam(model.parameters(), lr=0.001)',
         '使用 Adam 优化器。lr=0.001 是学习率。'
         'Adam 是自适应矩估计优化器，结合了 Momentum 和 RMSProp 的优点，'
         '在大多数任务上表现稳定，是最常用的深度学习优化器。')
add_line(222, '    loss_fn = nn.MSELoss()',
         '损失函数：均方误差 MSE。计算预测值与真实值之间差的平方的平均值。'
         '选 MSE 因为这是回归任务的标配——平方项会放大大误差的惩罚。')

doc.add_paragraph('')
add_text('──────────────────── ④ 训练循环 ────────────────────')

add_line(225, '    n_epochs = 200',
         '训练 200 轮。每一轮意味着模型看到全部训练数据一次。')
add_line(226, '    train_losses = []',
         '列表，记录每轮训练的 loss，用于后续绘制训练曲线。')
add_line(227, '    for epoch in range(n_epochs):',
         '主训练循环。epoch 从 0 到 199，共 200 轮。')
add_line(228, '        model.train()',
         '将模型切换到训练模式。这会影响 Dropout 和 BatchNorm 的行为：'
         '训练模式下 Dropout 生效，评估模式下 Dropout 关闭。')
add_line(229, '        optimizer.zero_grad()',
         '清空梯度。PyTorch 的梯度是累加的（不像某些框架会自动清零），'
         '所以每轮训练前必须手动清零。忘了这个 bug 很常见。')
add_line(230, '        y_pred = model(X_train_t)',
         '前向传播：将训练数据送入模型，得到预测值。'
         '等价于调用 PowerPredictor.forward()，但 PyTorch 会附加一些 hook。')
add_line(231, '        loss = loss_fn(y_pred, y_train_t)',
         '计算损失：用 MSE 比较预测值和真实值，得到一个标量。')
add_line(232, '        loss.backward()',
         '反向传播：计算损失对每个参数的梯度。'
         '这是 PyTorch 的 Autograd 机制自动完成的——从 loss 开始反向遍历计算图，'
         '计算 d(loss)/d(每个参数)。')
add_line(233, '        optimizer.step()',
         '更新参数：优化器根据梯度和学习率，更新模型权重。'
         '等效于：param = param - lr * grad。')
add_line(234, '        train_losses.append(loss.item())',
         '记录本轮的 loss 值。loss.item() 将 PyTorch 标量张量转换为 Python float。')

add_line(235, '        if (epoch + 1) % 50 == 0:',
         '每 50 轮（50/100/150/200）打印一次训练和测试 loss。'
         '不加 if 的话每轮都打印太刷屏。')
add_line(236, '            model.eval()',
         '切换到评估模式：关闭 Dropout，确保每次前向结果稳定。')
add_line(237, '            with torch.no_grad():',
         '在评估时禁用梯度计算，节省内存和计算量。'
         '在 torch.no_grad() 上下文中的操作不会构建计算图。')
add_line(238, '                test_loss = loss_fn(model(X_test_t), y_test_t)',
         '用测试集计算 loss。注意这里没有调用 backward() 和 step()，'
         '只是评估模型在未见过的数据上的表现。')
add_line(239, '            print(f"Epoch {epoch+1:3d}: train_loss={loss.item():.6f}, test_loss={test_loss.item():.6f}")',
         '打印 epoch、训练 loss、测试 loss。格式化为 3 位整数和 6 位小数。')

doc.add_paragraph('')
add_text('──────────────────── ⑤ 评估 + 可视化 ────────────────────')

add_line(242, '    model.eval()',
         '训练完成后，切换为评估模式。后续的所有前向都不使用 Dropout。')
add_line(243, '    with torch.no_grad():',
         '禁用梯度计算，加快推理速度。这个上下文管理器内的所有操作都不会跟踪梯度。')
add_line(244, '        y_pred = model(X_test_t).numpy().flatten()',
         '预测：model(X_test_t) 输出形状 (198, 1)。.numpy() 转成 numpy 数组，'
         '.flatten() 展平成一维 (198,)，方便与 y_test 比较。')
add_line(245, '        y_true = y_test',
         '真实值（原始的 numpy 数组，没有转成 Tensor）。')
add_line(246, '    mae = np.mean(np.abs(y_pred - y_true))',
         '计算平均绝对误差 MAE。衡量预测偏差的平均大小，单位是 kW。'
         '比 MSE 更直观——"平均偏差 X kW"。')
add_line(247, '    rmse = np.sqrt(np.mean((y_pred - y_true) ** 2))',
         '计算均方根误差 RMSE。相比于 MAE，RMSE 对大误差更敏感。'
         '如果 RMSE >> MAE，说明存在少数大偏差点。')
add_line(248, '    print(f"\\nTest MAE: {mae:.3f} kW, RMSE: {rmse:.3f} kW")',
         '打印结果。')

add_line(251, '    fig, axes = plt.subplots(2, 1, figsize=(12, 8))',
         '创建 2×1 的子图网格，整图尺寸 12×8 英寸。'
         'axes[0] 画训练 loss，axes[1] 画预测值对比。')
add_line(252, '    axes[0].plot(train_losses)',
         '绘制训练 loss 曲线。x 轴是 epoch，y 轴是 MSE loss。'
         '正常训练时 loss 应该不断下降然后趋于平稳。')
add_line(253, '    axes[0].set_xlabel("Epoch")',
         'x 轴标签。')
add_line(254, '    axes[0].set_ylabel("MSE Loss")',
         'y 轴标签。')
add_line(255, '    axes[0].set_title("Training Loss")',
         '子图标题。')
add_line(256, '    axes[0].grid(True)',
         '显示网格线，让曲线更容易读。')
add_line(257, '    axes[1].plot(y_true[:200], label="True", alpha=0.7)',
         '绘制真实功率值（前 200 个测试点），蓝色实线，透明度 0.7。'
         '只画前 200 点是因为太多了看不清。')
add_line(258, '    axes[1].plot(y_pred[:200], label="Predicted", alpha=0.7)',
         '绘制预测功率值（前 200 个测试点），橙色实线，透明度 0.7。')
add_line(259, '    axes[1].set_xlabel("Time Step")',
         'x 轴标签：时间步。')
add_line(260, '    axes[1].set_ylabel("FC Power (kW)")',
         'y 轴标签：FC 功率，单位 kW。')
add_line(261, '    axes[1].set_title(f"FC Power Prediction (MAE={mae:.2f} kW, RMSE={rmse:.2f} kW)")',
         '子图标题，里面嵌入 MAE 和 RMSE 的值。')
add_line(262, '    axes[1].legend()',
         '显示图例（"True" 和 "Predicted"）。')
add_line(263, '    axes[1].grid(True)',
         '网格线。')

add_line(264, '    plt.tight_layout()',
         '自动调整子图间距，防止标签重叠。')
add_line(265, '    plt.savefig(os.path.join(RESULTS_DIR, "week9_complete_mlp_power_prediction.png"), dpi=150)',
         '保存图片到 results/ 目录，文件名 week9_complete_mlp_power_prediction.png，分辨率 150 DPI。')
add_line(266, '    plt.close()',
         '关闭图像，释放内存。在脚本中如果不关闭，循环生成多图会内存泄漏。')
add_line(267, '    print(f"Saved: {RESULTS_DIR}/week9_complete_mlp_power_prediction.png")',
         '打印保存路径。')

add_line(269, '    return model',
         '返回训练好的模型对象。目前虽然没用到返回值，但保留入口方便后续直接调用模型做推理。')

doc.add_page_break()

# ══════════════════════════════════════════
# 知识要点总结
# ══════════════════════════════════════════
add_section('核心知识点总结')

p = doc.add_paragraph()
run = p.add_run('1. 时序预测的滑动窗口法')
run.bold = True
run.font.name = '微软雅黑'
add_text(
    '将一维时间序列转换成 (seq_len 步历史 → 第 seq_len+1 步) 的监督学习样本。\n'
    '这是最简单的时序处理方法，缺点是固定窗口大小，无法捕捉超长依赖。'
)

p = doc.add_paragraph()
run = p.add_run('2. MLP 的结构设计')
run.bold = True
run.font.name = '微软雅黑'
add_text(
    '输入层(10) → 隐藏层(64) → Dropout → 隐藏层(64) → 输出层(1)\n'
    '三层全连接 + ReLU + Dropout，是分类/回归任务的标准结构。'
)

p = doc.add_paragraph()
run = p.add_run('3. 训练循环的标准模板')
run.bold = True
run.font.name = '微软雅黑'
add_text(
    'model.train() → zero_grad() → y_pred = model(x) → loss = loss_fn(y_pred, y_true) '
    '→ loss.backward() → optimizer.step()\n'
    '这是 PyTorch 训练的固定配方。理解这 6 步 = 理解 PyTorch 训练。'
)

p = doc.add_paragraph()
run = p.add_run('4. 评估模式的必要性')
run.bold = True
run.font.name = '微软雅黑'
add_text(
    'model.eval() + torch.no_grad() 缺一不可：\n'
    'eval() 关闭 Dropout → 保证输出确定\n'
    'no_grad() 不建计算图 → 节省内存 + 加速推理'
)

doc.add_paragraph('')
add_section('面试追问预备')
add_text(
    'Q: 为什么用过去 10 步而不是 5 步或 20 步？\n'
    '→ SEQ_LEN 是超参数，越大的窗口提供越多历史信息但维度越高容易过拟合。'
    '通常用自相关分析（ACF/PACF）来选择最优窗口大小。\n\n'
    'Q: 这个 MLP 在时序预测上的根本缺陷是什么？\n'
    '→ MLP 没有时序结构：它把 10 个历史点当作独立的 10 个特征，'
    '不"知道"它们之间有先后顺序。RNN/LSTM 或 Transformer 才能建模时序依赖。\n\n'
    'Q: 为什么用 sin 生成模拟数据？\n'
    '→ 模拟数据是为了快速验证模型代码的正确性。'
    '在实际项目中，用的必然是真实传感器数据（如之前下载的 Shifts 数据集或 TU Delft 功率剖面）。'
)

output = r'F:\CLAUDE\research\ems-platform\docs\notes\Week9_Part4_MLP_逐行解释.docx'
doc.save(output)
print(f'OK: {output}')
