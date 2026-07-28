#!/usr/bin/env python3
"""生成《EMS 仿真器 — 实施任务书》DOCX"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1a, 0x47, 0x8a)

def heading(s, level=1):
    return doc.add_heading(s, level=level)

def para(s, bold=False):
    p = doc.add_paragraph(s)
    if bold:
        for r in p.runs:
            r.bold = True
    return p

def bullet(s):
    doc.add_paragraph(s, style='List Bullet')

def code_block(s):
    p = doc.add_paragraph()
    r = p.add_run(s)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)

# ═══════════════ 封面 ═══════════════
doc.add_heading('EMS 仿真器\n实施任务书', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('把"一堆脚本"变成"一个能跑的系统"')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('编制日期：2026-07-27 | 预期耗时：4-5天 | 可复用代码：~2880 行').font.size = Pt(9)

doc.add_page_break()

# ═══════════════ 1. 目标 ═══════════════
heading('1. 目标', level=1)
para(
    '把现在的分散脚本（day8_dp_ems.py, day9_ecms_ems.py, mpc_ems_optimized.py, mpc_ems_ekf.py）'
    '重构为一个统一的 EMS 仿真器。\n\n'
    '现在：每个脚本自己读数据、自己跑仿真、自己出图——做一次四方法对比要手动跑4个脚本。\n'
    '之后：一个命令跑所有方法，结果自动归档，加新算法只需要写一个 engine。'
)

heading('1.1 一句话定义', level=2)
para(
    '"输入工况曲线 → 选择EMS算法 → 仿真器自动跑完全程 → 输出氢耗/SOC/效率对比报告"',
    bold=True
)

heading('1.2 成功后长什么样', level=2)
code_block(
    '# 单算法仿真\n'
    'python sim/run.py --method mpc --cycle wltc --ekf\n\n'
    '# 全部算法对比\n'
    'python sim/run.py --compare all\n\n'
    '# 自定义工况\n'
    'python sim/run.py --method ecms --cycle custom --load my_data.csv\n\n'
    '# 参数扫描\n'
    'python sim/run.py --method ecms --scan s --range 50 200 --step 10'
)

# ═══════════════ 2. 架构 ═══════════════
doc.add_page_break()
heading('2. 架构设计', level=1)

heading('2.1 目录结构', level=2)
code_block(
    'scripts/simulator/\n'
    '├── run.py                    # 入口：解析参数，调度仿真\n'
    '├── config/\n'
    '│   └── vehicle.yaml          # 整车参数（FC/电池/SOC限值等）\n'
    '├── core/\n'
    '│   ├── simulation.py         # 仿真循环（时间步进+日志+出图）\n'
    '│   ├── metrics.py            # 统一指标计算（氢耗/SOC/效率/RMSE）\n'
    '│   └── cycle_loader.py       # 工况加载（WLTC/NEDC/CLTC/TU Delft）\n'
    '├── engines/\n'
    '│   ├── __init__.py           # engine 注册表\n'
    '│   ├── base.py               # Engine 基类（统一接口）\n'
    '│   ├── dp.py                 # DP 求解器（封装现有代码）\n'
    '│   ├── ecms.py               # ECMS 求解器（封装现有代码）\n'
    '│   ├── mpc.py                # MPC 求解器（封装现有代码）\n'
    '│   ├── mpc_ekf.py            # MPC + EKF\n'
    '│   └── ppo.py                # PPO（第11周完成后加入）\n'
    '├── estimators/\n'
    '│   └── ekf.py                # EKF SOC 估计（从mpc_ems_ekf.py提取）\n'
    '└── results/                  # 运行结果自动生成'
)

heading('2.2 核心设计：Engine 接口', level=2)
para('每个算法都实现同一个接口，仿真循环不关心具体算法是什么：')
code_block(
    'class Engine:\n'
    '    """EMS 算法引擎基类"""\n'
    '    def __init__(self, config: dict):\n'
    '        """初始化：加载参数，分配资源"""\n'
    '        pass\n\n'
    '    def reset(self, soc_init: float):\n'
    '        """开始新仿真：重置状态"""\n'
    '        pass\n\n'
    '    def step(self, t: float, P_load: float, SoC: float) -> dict:\n'
    '        """每一步的控制决策：返回 P_fc"""\n'
    '        pass\n\n'
    '    def name(self) -> str:\n'
    '        """算法名称（用于报告/图例）"""\n'
    '        pass'
)

heading('2.3 核心设计：仿真循环', level=2)
para('simulation.py 做的事情：')
code_block(
    'def run_simulation(engine, cycle, config):\n'
    '    engine.reset(config["soc_init"])\n'
    '    for t, P_load in cycle:\n'
    '        action = engine.step(t, P_load, SoC)\n'
    '        SoC = update_soc(SoC, P_load, action["P_fc"])\n'
    '        log(t, P_load, action["P_fc"], SoC, ...)\n'
    '    metrics = compute_metrics(log)\n'
    '    plot_results(log, metrics)\n'
    '    return metrics'
)

# ═══════════════ 3. 实施步骤 ═══════════════
doc.add_page_break()
heading('3. 实施步骤', level=1)

heading('3.1 第1天：搭骨架（配置+循环+CLI）', level=2)
para('目标：跑通一个"假的"仿真循环，确认架构能跑。')
bullet('创建目录结构')
bullet('把硬编码的FC参数提取到 config/vehicle.yaml')
bullet('写 core/cycle_loader.py（加载 WLTC/NEDC/CLTC 三工况）')
bullet('写 core/metrics.py（氢耗计算 + SOC终值 + 等效氢耗 + FC平均效率）')
bullet('写 engines/base.py（Engine 基类）')
bullet('写 run.py（argparse 解析 --method --cycle --compare 等参数）')
bullet('验证：写一个 DummyEngine 打印日志，确认循环能跑到结束')

heading('3.2 第2-3天：封装现有算法', level=2)
para('目标：把现有代码装进 Engine 接口，跑通真实仿真。')

para('engines/dp.py', bold=True)
bullet('DP 是"批量"算法（不是逐时间步决策），所以比较特殊')
bullet('实现方式：step() 第一次调用时算完整个 DP，然后缓存结果')
bullet('复用的代码：day8_dp_ems.py 中的 solver 核心逻辑')

para('engines/ecms.py', bold=True)
bullet('ECMS 是逐时间步的，直接封装')
bullet('每步计算代价函数 → 网格搜索最优 P_fc')
bullet('复用代码：day9_ecms_ems.py 中的 ecms_sim() 核心循环')

para('engines/mpc.py', bold=True)
bullet('MPC 也是逐时间步的')
bullet('每步：构建预测序列 → 优化 → 取第一步')
bullet('复用代码：mpc_ems_optimized.py 中的 mpc_sim() 核心循环')

para('engines/mpc_ekf.py', bold=True)
bullet('在 MPC 基础上串联 EKF SOC 估计')
bullet('提取 EKF 逻辑到 estimators/ekf.py')
bullet('复用代码：mpc_ems_ekf.py')

heading('3.3 第4天：对比功能 + TU Delft 验证', level=2)
bullet('run.py 增加 --compare all 模式：串行跑所有 engine，汇总结果')
bullet('cycle_loader.py 增加 TU Delft 数据加载（从 Results_ch3.xlsx 提取）')
bullet('验证：一次命令跑完四方法对比，确认结果与之前手动跑的吻合')

heading('3.4 第5天：README + 简历包装', level=2)
bullet('写 scripts/simulator/README.md（架构说明+使用指南+结果展示）')
bullet('准备面试 Demo 命令（3-4条命令，现场可演示）')
bullet('更新简历项目描述')

# ═══════════════ 4. 复用率 ═══════════════
heading('4. 代码复用率估算', level=1)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['Engine', '现有行数', '可复用行数', '需新写行数']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)

data = [
    ['DP', '517', '~350', '~150'],
    ['ECMS', '628', '~400', '~150'],
    ['MPC', '718', '~450', '~200'],
    ['MPC+EKF', '1017', '~700', '~250'],
]
for i, row in enumerate(data):
    for j, v in enumerate(row):
        table.rows[i+1].cells[j].text = v
        for p in table.rows[i+1].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

para('')
para('总计：现有 ~2880 行，可直接复用 ~1900 行（66%），需新写 ~750 行。', bold=True)
para(
    '新写的部分主要是：Engine 封装层（150行）、仿真循环（200行）、'
    'config加载器（50行）、cycle_loader（100行）、metrics（100行）、CLI（150行）。'
)

# ═══════════════ 5. 完成后效果 ═══════════════
doc.add_page_break()
heading('5. 完成后能展示什么', level=1)

heading('5.1 终端演示（面试时现场跑）', level=2)
para('面试官问"你做过什么"时，你可以当场打开终端：')
code_block(
    '$ python sim/run.py --method mpc --cycle wltc --ekf\n'
    '>> Loading config: config/vehicle.yaml\n'
    '>> Loading cycle: WLTC (1800s)\n'
    '>> Running MPC+EKF ...\n'
    '>> Done. Results saved to results/20260727_mpc_ekf_wltc/\n'
    '>> H2: 0.2198 kg | SOC_end: 0.572 | SOC_RMSE: 0.0024\n\n'
    '$ python sim/run.py --compare all\n'
    '>> Running DP ... done.\n'
    '>> Running ECMS ... done.\n'
    '>> Running MPC ... done.\n'
    '>> Running MPC+EKF ... done.\n'
    '>> Comparison saved to results/compare_20260727/'
)

heading('5.2 简历上的描述', level=2)
para('简历上的写法（约3行）：', bold=True)
bullet('搭建了燃料电池混动系统 EMS 仿真器（Python），集成 DP/ECMS/MPC/PPO 四种能量管理算法，支持任意工况输入与统一指标对比')
bullet('在 WLTC/NEDC/CLTC 三工况及 TU Delft 实船数据集上完成验证，DP 相比规则控制降低氢耗 19.2%')
bullet('集成 EKF SOC 在线估计模块，SOC 跟踪 RMSE 0.0024（比开环提高 4.8 倍），可模拟传感器噪声和偏置场景')

para('')
para('和现在的区别：', bold=True)
bullet('现在：脚本散落，叙事弱，面试官看到的是"跑实验的记录"')
bullet('之后：统一入口 + 标准化接口 + 对比功能，面试官看到的是"他能搭系统"')

# ═══════════════ 6. 注意 ═══════════════
heading('6. 注意事项', level=1)
bullet('不要重写算法——所有算法逻辑直接从现有脚本复制，只封装接口')
bullet('DP 是"批量"算法（需全工况），其他是"逐时间步"算法 —— 仿真循环要同时支持两种模式')
bullet('config/vehicle.yaml 的参数值从现有脚本中提取，不要无中生有')
bullet('每次提交一个可运行的版本，不要写到一半跑不起来')
bullet('PPO engine 留空即可（等第11周 PPO 实现后再填）')

# ═══════════════ 7. 交付物清单 ═══════════════
heading('7. 完整交付物清单', level=1)
para('最后一个步骤完成后，你的仓库里应该多出这些文件：')
bullet('scripts/simulator/run.py')
bullet('scripts/simulator/config/vehicle.yaml')
bullet('scripts/simulator/core/__init__.py')
bullet('scripts/simulator/core/simulation.py')
bullet('scripts/simulator/core/metrics.py')
bullet('scripts/simulator/core/cycle_loader.py')
bullet('scripts/simulator/engines/__init__.py')
bullet('scripts/simulator/engines/base.py')
bullet('scripts/simulator/engines/dp.py')
bullet('scripts/simulator/engines/ecms.py')
bullet('scripts/simulator/engines/mpc.py')
bullet('scripts/simulator/engines/mpc_ekf.py')
bullet('scripts/simulator/engines/ppo.py（占位，空壳）')
bullet('scripts/simulator/estimators/__init__.py')
bullet('scripts/simulator/estimators/ekf.py')
bullet('scripts/simulator/README.md')

# ═══════════════ 8. 时间线 ═══════════════
heading('8. 时间线一览', level=1)

table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Light Grid Accent 1'
for i, h in enumerate(['天次', '任务', '产出']):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)

days = [
    ['第1天', '搭骨架：config + cycle_loader + CLI + 仿真循环 + Dummy验证', '能跑通一个假循环'],
    ['第2天', '封装 DP + ECMS', '两个真实算法能跑'],
    ['第3天', '封装 MPC + MPC+EKF', '四种算法都能跑'],
    ['第4天', '对比功能 + TU Delft 验证', '一次命令跑完全部对比'],
    ['第5天', 'README + 简历包装', '项目完整可展示'],
]
for i, row in enumerate(days):
    for j, v in enumerate(row):
        table2.rows[i+1].cells[j].text = v
        for p in table2.rows[i+1].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

para('')
para('如果只有晚上和周末做：第1-2天用一个周末，第3-5天用下一个周末，10天内完成。', bold=True)

# ═══════════════ 9. 面试话术 ═══════════════
heading('9. 面试话术速查', level=1)

heading('9.1 一句话版', level=2)
para(
    '"我搭了一个燃料电池混动系统的 EMS 仿真器，把 DP/ECMS/MPC 统一到同一个框架下，'
    '一条命令就能跑完对比、出报告。"'
)

heading('9.2 两分钟版', level=2)
para(
    '"我之前是分别在跑不同的算法脚本，后来发现每换一种工况就要改好几个地方，'
    '对比结果还要手动汇总，很麻烦。\n'
    '所以我用了一个周末重新搭了这个仿真器——给每个算法封装了统一的 Engine 接口，'
    '把整车参数抽到 YAML 配置文件里，写了一个仿真循环负责时间步进和数据记录。'
    '现在只要一条命令就能跑任意算法、任意工况，结果自动保存到带时间戳的文件夹。\n'
    '这个改造让我做对比实验的效率提升了很多，而且代码结构清晰后，加新算法（比如刚做的 PPO）'
    '只需要写一个 Engine 类注册进去就行，仿真循环完全不用改。"'
)

# ═══════════════ 10. 对比 ═══════════════
heading('10. 做之前 vs 做之后', level=1)
table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Light Grid Accent 1'
for i, h in enumerate(['维度', '现在', '之后']):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)

compare = [
    ['运行方式', 'python day8_dp_ems.py（脚本名无规律）', 'python sim/run.py --method dp（统一入口）'],
    ['加新算法', '新建一个 .py 从头写', '写一个 Engine 类，注册就行'],
    ['加新工况', '改代码里的路径', 'CSV 放 cycles/ 目录下，自动识别'],
    ['对比结果', '肉眼看多个终端输出', '自动汇总成对比表 + 对比图'],
]
for i, row in enumerate(compare):
    for j, v in enumerate(row):
        table3.rows[i+1].cells[j].text = v
        for p in table3.rows[i+1].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

# ── 结尾 ──
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('\n—— 任务书完 ——').bold = True
para('编制日期：2026-07-27')

out_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    'docs',
    'EMS仿真器_实施任务书.docx'
)
doc.save(out_path)
print('Done:', out_path)
