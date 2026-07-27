#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 DQN / SAC 概念精讲文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def para(text, bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    return p

def code_block(lines, label=None):
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
    for line in lines.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
    for row_data in rows:
        row = table.add_row()
        for c, txt in enumerate(row_data):
            cell = row.cells[c]
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
    return table

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph('')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('DQN / SAC 概念精讲')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t2.add_run('从值迭代到深度 Q 网络 · 从策略梯度到最大熵')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t3.add_run(f'\n生成日期：{datetime.date.today().isoformat()}')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

para('')
para('—— 接在 Week 9 之后的概念扩展，帮你理解 DRL 两大流派', size=10, color=RGBColor(0x99, 0x99, 0x99))

page_break()

# ═══════════════════════════════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════════════════════════════
heading('目录', 1)
para('第一部分：DQN — 深度 Q 网络（基于值的流派）', bold=True)
para('  1.1 从 Q-learning 到 DQN：为什么需要深度学习')
para('  1.2 DQN 的两大创新：经验回放 + 目标网络')
para('  1.3 DQN 算法流程')
para('  1.4 DQN 的局限')
para('')
para('第二部分：SAC — Soft Actor-Critic（基于策略的流派）', bold=True)
para('  2.1 从策略梯度到 Actor-Critic')
para('  2.2 最大熵强化学习')
para('  2.3 SAC 算法流程')
para('  2.4 SAC 为什么是当前主流')
para('')
para('第三部分：对比总表', bold=True)
para('第四部分：在 EMS 中怎么选', bold=True)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第一部分：DQN
# ═══════════════════════════════════════════════════════════════════
heading('第一部分：DQN — 深度 Q 网络', 1)
para('— Value-Based 流派的代表算法', color=RGBColor(0x66, 0x66, 0x66))

heading('1.1 从 Q-learning 到 DQN', 2)

para('回忆 Week 9 的值迭代 (Value Iteration)：', bold=True)
para('V*(s) = max_a [ R(s,a) + γ · Σ P(s\'|s,a) · V*(s\') ]')
para('')
para('Q-learning 把这个公式改成不需要知道 P 的形式：', bold=True)
para('Q*(s,a) = R(s,a) + γ · max_a\' Q*(s\', a\')')
para('')
para('关键区别：', bold=True)
bullet('值迭代需要知道 P（模型已知）→ 这是 Model-Based 方法')
bullet('Q-learning 不需要 P → 这是 Model-Free 方法')
bullet('Q-learning 用一个 Q 表存储每个 (s,a) 的值，直接在环境中采样更新')
para('')
para('但 Q 表有个致命问题：', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('GridWorld 只有 16 个状态 × 4 个动作 = 64 个格子。')
para('但真实问题的状态空间是巨大的：')
bullet('围棋：10¹⁷⁰ 个状态')
bullet('EMS：SOC(150) × 功率(60) = 9000 个状态 — Q 表勉强可行')
bullet('连续状态（如温度、电压）：Q 表根本没法存')

para('')
para('DQN 的核心思想：用神经网络近似 Q 函数', bold=True, size=12, color=RGBColor(0x1F, 0x3A, 0x5F))
para('')
para('Q(s,a) ≈ Q_θ(s,a)    ← θ 是神经网络参数')
para('')
para('用神经网络拟合 Q 函数，任意连续状态都能输入，不必离散化。')

page_break()

heading('1.2 DQN 的两大创新', 2)

heading('创新一：经验回放 (Experience Replay)', 3)
para('问题：强化学习的数据是时序相关的。智能体连续采样导致样本之间高度相关，神经网络会"忘记"之前学过的经验。')
para('')
para('解法：', bold=True)
bullet('智能体把每一步的 (s, a, r, s\') 存到一个回放缓冲区 (Replay Buffer)')
bullet('训练时从缓冲区随机采样一个小 batch（打乱时序相关性）')
bullet('类比：学习历史不能只看最近几天，要随机翻以前的记录看')
para('')
para('经验回放还有一个好处：数据复用。一份经验可以训练多次，提高样本效率。')

heading('创新二：目标网络 (Target Network)', 3)
para('问题：DQN 的更新目标是 r + γ · max Q_θ(s\', a\')，其中 Q_θ 既是"被更新者"又是"更新目标"。')
para('这就像一个人追着自己的影子跑——目标一直在变，训练不稳定。')
para('')
para('解法：', bold=True)
bullet('维护两个 Q 网络：主网络 Q_θ 和目标网络 Q_θ⁻')
bullet('θ⁻ 每隔 C 步从 θ 复制，其他时间冻结不变')
bullet('更新目标：r + γ · max Q_θ⁻(s\', a\')  ← 目标相对稳定')
para('')
para('类比：射箭时靶子不能动。目标网络就是那个固定靶，主网络是射出去的箭。')

page_break()

heading('1.3 DQN 算法流程', 2)
para('用伪代码一目了然：')
para('')

code_block('''初始化 Q 网络 Q_θ，目标网络 Q_θ⁻ ← Q_θ
初始化经验回放缓冲区 D

for episode in range(M):
    获取初始状态 s
    for step in range(T):
        # ε-贪心选动作
        if random() < ε:
            a ← 随机动作
        else:
            a ← argmax Q_θ(s, a)

        执行 a，得到 r, s'
        存储 (s, a, r, s') 到 D

        if D 中样本足够:
            从 D 随机采样一个 batch (s, a, r, s')
            计算目标 y = r + γ · max Q_θ⁻(s', a')
            损失 L = (y - Q_θ(s, a))²
            梯度下降更新 θ

        s ← s'

    每 C 步: θ⁻ ← θ''' )

para('')
para('其中 ε-贪心 (ε-greedy) 是探索策略：', bold=True)
bullet('ε 概率随机探索（发现新可能）')
bullet('1-ε 概率选当前最优（利用已学知识）')
bullet('训练过程中 ε 逐渐衰减（从 1.0 降到 0.01）')

page_break()

heading('1.4 DQN 的局限', 2)

para('DQN 在 Atari 游戏上取得了巨大成功（2013 年 DeepMind），但它有几个固有局限：')
para('')

add_table(
    ['局限', '原因', '后果'],
    [
        ['只能处理离散动作', 'argmax Q(s,a) 需要遍历所有动作', '无法直接用于连续控制（如发动机功率、扭矩）'],
        ['Q 值过估计', 'max 操作导致 Q 值系统性偏高', '策略次优，训练不稳定'],
        ['样本效率低', 'on-policy 性质（早期 DQN 是 off-policy，但需要大量采样）', '训练慢，需要上千万帧'],
        ['对奖励设计敏感', 'Q 值范围差异大，网络难收敛', '需要精细调参'],
    ]
)

para('')
para('这些局限推动了后续改进：Double DQN（解决过估计）、Dueling DQN（分离状态价值和动作优势）、PER（优先经验回放）等。')
para('但更重要的是，它们催生了另一条技术路线——Based Policy-Based 方法。')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第二部分：SAC
# ═══════════════════════════════════════════════════════════════════
heading('第二部分：SAC — Soft Actor-Critic', 1)
para('— Policy-Based 流派的集大成者', color=RGBColor(0x66, 0x66, 0x66))

heading('2.1 从策略梯度到 Actor-Critic', 2)

para('DQN 的问题是"先算 Q 值，再从 Q 值选动作"。如果动作空间是连续的（比如分配 0-30kW 功率），就没法遍历所有动作求 argmax。')
para('')
para('策略梯度 (Policy Gradient) 的思路完全不同：', bold=True)
para('不经过 Q 值，直接学一个策略网络 π_θ(a|s)，输入状态 s，输出动作 a（或动作分布）。')
para('')
para('策略梯度定理 (REINFORCE)：', bold=True)
para('∇J(θ) = E[ ∇log π_θ(a|s) · G ]')
para('')
para('其中 G 是从当前步到结束的累积奖励。直观理解：如果一个轨迹获得了高回报，就增大这条轨迹上所有动作的概率。')
para('')
para('但 REINFORCE 有个问题：G 的方差极大。一条轨迹好不代表每一步都好。')
para('')
para('Actor-Critic 架构的改进：', bold=True)
bullet('Actor = 策略网络 π_θ(s) → 输出动作（"演员"）')
bullet('Critic = 价值网络 V_φ(s) 或 Q_φ(s,a) → 评估动作好坏（"评委"）')
bullet('Actor 根据 Critic 的反馈更新，Critic 根据实际奖励更新')
para('')
para('类比：')
para('  REINFORCE = 演员自己看整部电影的结果来改演技（方差大）')
para('  Actor-Critic = 演员演完一幕，导演当场说"这幕演得好/不好"（方差小）')

page_break()

heading('2.2 最大熵强化学习', 2)

para('SAC (Soft Actor-Critic) 的核心创新：在标准 RL 目标上加了一个"熵"项。', bold=True)
para('')
para('标准 RL 目标：', bold=True)
para('  J = E[ Σ γ^t · r_t ]          ← 只最大化累积奖励')
para('')
para('SAC 的目标：', bold=True)
para('  J = E[ Σ γ^t · (r_t + α·H(π(·|s_t))) ]')
para('  其中 H(π(·|s_t)) = -Σ π(a|s) · log π(a|s)  ← 熵')
para('')
para('熵 (Entropy) 衡量不确定性。熵越大 → 动作分布越均匀 → 智能体越"随机"：')
bullet('低熵：每次都选同一个动作（确定性策略）')
bullet('高熵：动作选择多样化（随机策略）')
para('')
para('为什么要加熵项？三个好处：', bold=True)
para('')
para('① 鼓励探索：熵正则化让智能体不会过早陷入确定性策略，保持对环境的探索。')
para('')
para('② 鲁棒性：学过多种做法的策略，在环境变化时不容易崩溃。（类比：一个厨师会多种做法，遇到缺材料也能应变）')
para('')
para('③ 多模态：当多个动作同样好时，SAC 会均匀分配概率，而不是随机选一个。')

para('')
para('α（温度系数）控制探索程度：', bold=True)
bullet('α = 0：退化为标准 RL，不鼓励探索')
bullet('α 很大：智能体几乎随机行动')
bullet('SAC 会自动调节 α：目标熵 = -dim(A)（动作维数的负数）')

page_break()

heading('2.3 SAC 算法流程', 2)

para('SAC 使用 5 个网络（比 DQN 复杂得多）：', bold=True)
para('')

add_table(
    ['网络', '符号', '作用'],
    [
        ['Actor', 'π_θ', '输入状态 s，输出动作分布（均值 + 方差）'],
        ['Critic Q1', 'Q_φ1', '评估 (s,a) 的价值'],
        ['Critic Q2', 'Q_φ2', '评估 (s,a) 的价值（取最小值防过估计）'],
        ['目标 Q1', 'Q_φ1⁻', '稳定训练（同 DQN 的目标网络）'],
        ['目标 Q2', 'Q_φ2⁻', '稳定训练'],
    ]
)

para('')

code_block('''初始化策略 π_θ，两个 Q 网络 Q_φ1, Q_φ2，目标 Q⁻ ← Q
初始化经验回放缓冲区 D

for each step:
    a ~ π_θ(a|s)          # 从策略分布采样动作
    执行 a, 得到 r, s'
    存储 (s, a, r, s') 到 D

    从 D 采样 batch
    # 1. 更新 Q 网络（最小化 Bellman 残差）
    目标值 y = r + γ · ( min(Q_φ⁻(s', a'), Q_φ₂⁻(s', a')) - α·log π_θ(a'|s') )
    损失 L_Q = (Q_φ1(s,a) - y)² + (Q_φ2(s,a) - y)²
    梯度下降更新 φ1, φ2

    # 2. 更新策略（最大化 V + 熵）
    L_π = α·log π_θ(a|s) - min(Q_φ1(s,a), Q_φ2(s,a))
    梯度下降更新 θ

    # 3. 更新 α（自动调节）
    L_α = -α · (log π_θ(a|s) + H_target)
    梯度下降更新 α

    # 4. 软更新目标网络
    φ⁻ ← τ·φ + (1-τ)·φ⁻     # τ 很小（如 0.005）''' )

para('')
para('关键技巧：', bold=True)
bullet('双 Q 网络（Double Q）：取两个 Q 的最小值，避免 DQN 中的过估计问题')
bullet('重参数化技巧：a = f_θ(s, ε) 而不是直接从分布采样，使梯度可以回传')
bullet('软更新：目标网络每次只移动一小步（τ=0.005），比 DQN 的硬复制更稳定')

page_break()

heading('2.4 SAC 为什么是当前主流', 2)
para('SAC 是目前连续控制任务中最主流的 off-policy 算法，原因：')
para('')

add_table(
    ['优点', '说明'],
    [
        ['样本效率高', 'Off-policy + 经验回放，一份数据用多次'],
        ['训练稳定', '最大熵 + 双 Q + 软更新，三者共同抑制训练崩溃'],
        ['自动调探索', 'α 自动调节，不需要手动调 ε 衰减'],
        ['连续/离散都行', 'SAC 原生支持连续动作，离散化后也支持离散'],
        ['超参数鲁棒', '相比 PPO 对学习率、batch size 等更不敏感'],
    ]
)

para('')
para('SAC 的不足：', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
bullet('实现复杂度高（5 个网络，代码量是 DQN 的 2-3 倍）')
bullet('对 CPU 内存需求大（每个网络都有优化器状态）')
bullet('调试困难（不知道是 Q 没学好还是 π 没学好）')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第三部分：对比总表
# ═══════════════════════════════════════════════════════════════════
heading('第三部分：三者对比总表', 1)
para('你在 Week 9 学的值迭代 (VI)、以及 DQN、SAC，本质上都在求解同一个问题：', bold=True)
para('')
para('只是求解路径不同，最终都收敛到同一个最优策略 V* / Q*。')
para('')

add_table(
    ['维度', '值迭代 (Part 8)', 'DQN', 'SAC'],
    [
        ['流派', 'Dynamic Programming', 'Value-Based', 'Actor-Critic'],
        ['是否需要模型 P', '✅ 需要', '❌ 不需要', '❌ 不需要'],
        ['动作空间', '离散', '离散', '连续 / 离散'],
        ['核心公式', 'V = max Q', 'Q = r + γ·max Q⁻', 'π = argmax(Q + αH)'],
        ['函数近似', '无（查表）', '神经网络 Q_θ', '神经网络 π_θ + Q_φ'],
        ['探索方式', '无（已知 P）', 'ε-贪心', '策略本身的熵'],
        ['样本效率', '不需要数据（已知模型）', '低', '中-高'],
        ['实现难度', '⭐⭐（一周能写）', '⭐⭐⭐', '⭐⭐⭐⭐⭐'],
        ['适合场景', '模型已知、小状态', '离散控制、游戏', '连续控制、机器人'],
    ]
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第四部分：在 EMS 中怎么选
# ═══════════════════════════════════════════════════════════════════
heading('第四部分：在 EMS 中怎么选', 1)
para('回到你的实际场景——燃料电池 EMS 能量管理。这些算法怎么用？')
para('')

heading('4.1 为什么现有项目用 DP（不是 RL）', 2)
bullet('DP 有模型保证全局最优，可做基准')
bullet('DP 不需要训练，不需要调超参数')
bullet('DP 的劣势是需要已知工况，不能在线使用')
para('')

heading('4.2 如果未来要做 RL-EMS', 2)
para('根据你的 STATUS.md 计划，只做 PPO（不做 DQN/SAC 实现）。选 PPO 而不是 DQN/SAC 的理由：')
para('')

add_table(
    ['算法', '选或不选', '理由'],
    [
        ['PPO', '✅ 主选', 'On-policy + clipped objective = 训练稳定简单，连续动作原生支持，工业落地最多'],
        ['SAC', '❌ 暂不', '虽然样本效率更高，但实现和调参成本高，秋招时间线来不及深度调优'],
        ['DQN', '❌ 跳过', '只支持离散动作，EMS 功率分配是连续/高散化问题，argmax 效率低'],
    ]
)

para('')
para('面试时如果被问到"为什么选 PPO 不选 SAC/DQN"，可以这样回答：', bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))
para('')
para('"PPO 在训练稳定性和实现复杂度之间取得了最好的平衡。SAC 样本效率更高但需要精细调节 5 个网络，在秋招时间线下优先跑通 PPO 闭环。'
     'DQN 只能处理离散动作，而 EMS 中功率分配本质上是连续控制问题——离散化会损失精度，且动作空间随离散粒度指数增长。"',
     indent=1)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 附录：面试 Q&A
# ═══════════════════════════════════════════════════════════════════
heading('附录：面试可能追问的问题', 1)

para('Q1: DQN 为什么需要经验回放？', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('强化学习的数据是时序相关的，连续采样会导致样本之间高度相关，神经网络会产生灾难性遗忘（catastrophic forgetting）。经验回放通过随机采样打乱相关性，同时提高数据利用率。')

para('')
para('Q2: DQN 的目标网络为什么能稳定训练？', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('如果没有目标网络，每一步更新的目标 y = r + γ·max Q_θ(s\', a\') 中的 Q_θ 在不断变化。这相当于"追着移动的靶子射箭"，目标不稳定导致训练发散。目标网络冻结参数 C 步，让更新目标相对稳定，训练更平滑。')

para('')
para('Q3: SAC 的"最大熵"是什么意思？', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('熵衡量随机性。SAC 在奖励基础上加了一个熵奖励，鼓励策略保持随机。好处是探索更充分、策略更鲁棒。α 控制探索力度，SAC 能自动调节 α 到目标熵值。')

para('')
para('Q4: PP O 和 SAC 怎么选？', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('PPO 更简单稳定，是 on-policy 方法，对超参数不敏感，适合快速落地。SAC 样本效率更高，适合数据获取成本高的场景（如真实机器人），但实现和调参更复杂。工业界 PPO 使用更广泛。')

para('')
para('Q5: DRL 在 EMS 中落地的主要挑战？', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('三个挑战：① 安全约束 — RL 探索阶段可能产生不安全动作（如过放电池），需要安全层或约束优化；② 样本效率 — 真实系统中采集百万步数据不现实，需要数字孪生或 sim-to-real；③ 泛化性 — 训练时的工况分布与真实运行分布可能不同，策略可能失效。')

# ═══════════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════════
path = os.path.join(OUT_DIR, 'DQN_SAC_概念精讲.docx')
doc.save(path)
print(f'OK: {path}')
