#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成优化器（SGD/Momentum/RMSProp/Adam）逐行精讲文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def para(text, bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    return p

def code_block(lines, label=None):
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
    for line in lines.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
    for row_data in rows:
        row = table.add_row()
        for c, txt in enumerate(row_data):
            cell = row.cells[c]
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
    return table

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph('')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('优化器精讲\nSGD · Momentum · RMSProp · Adam')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t2.add_run('把深度学习优化器彻底讲清楚')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t3.add_run(f'\n生成日期：{datetime.date.today().isoformat()}')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

para('')
para('—— 接 Week 9 Part 4 优化器部分，逐行理解四个优化器从简单到复杂的进化链路', size=10, color=RGBColor(0x99, 0x99, 0x99))

page_break()

# ═══════════════════════════════════════════════════════════════════
# 0. 一个共同的问题
# ═══════════════════════════════════════════════════════════════════
heading('首先：所有优化器在解决同一个问题', 1)

para('训练神经网络 = 调整参数 θ 让损失函数 L(θ) 最小。', bold=True)
para('')
para('优化器的每一步都做同一件事：')
para('')
code_block('θ ← θ - lr · 方向(∇L)    # ∇L = 梯度（导数）')
para('')
para('不同优化器的区别只在"方向"怎么算：', bold=True)
bullet('SGD：方向 = 梯度本身')
bullet('Momentum：方向 = 梯度 + 历史动量的衰减')
bullet('RMSProp：方向 = 梯度 / 自适应步长（每个参数独立调步长）')
bullet('Adam：方向 = 上面两个加在一起')
para('')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. SGD
# ═══════════════════════════════════════════════════════════════════
heading('第一章：SGD — 随机梯度下降', 1)

heading('1.1 核心公式', 2)
code_block('θ = θ - lr · ∇L(θ)')
para('')
para('其中：', bold=True)
bullet('θ：模型参数（权重和偏置）')
bullet('lr：学习率（步长），如 0.01')
bullet('∇L(θ)：损失函数对 θ 的梯度（指向损失增加最快的方向）')
para('')

heading('1.2 直觉理解：蒙眼下山', 2)
para('想象你被蒙住眼睛站在山上，要走到山谷：', bold=True)
para('')
bullet('用脚探一下哪个方向最陡（算梯度）')
bullet('往反方向迈一小步（梯度下降）')
bullet('再探、再走、重复')
para('')
para('问题：', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
bullet('步长固定：陡坡时步子太大可能跳过最低点，缓坡时步子太小半天走不到')
bullet('所有参数同一学习率：不同参数应该有不同的步长')
bullet('峡谷地形震荡：窄方向梯度正负交替，来回震荡浪费步数')

heading('1.3 在 Week 9 Part 4 中的代码', 2)
para('虽然 Part 4 最终用了 Adam，但如果换 SGD 就是这一行：')
code_block('optimizer = torch.optim.SGD(model.parameters(), lr=0.001)')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. Momentum
# ═══════════════════════════════════════════════════════════════════
heading('第二章：Momentum — 动量法', 1)

heading('2.1 核心公式', 2)
code_block('v = γ·v + lr · ∇L       # γ 通常 = 0.9')
code_block('θ = θ - v')
para('')
para('相比 SGD 多了一项 v（速度/动量）：', bold=True)
bullet('γ·v：历史动量的衰减保留（惯性项）')
bullet('lr · ∇L：当前梯度（推动力）')
bullet('θ = θ - v：用动量 v 而不是直接用梯度更新')

heading('2.2 直觉理解：滚球下山', 2)
para('还是蒙眼下山，但这次脚下踩着一个滚球：', bold=True)
para('')
bullet('球有惯性——之前往东滚，这次梯度偏北一点，主要还往东')
bullet('平缓区域：梯度持续指向同一方向，动量累积→越滚越快')
bullet('峡谷震荡：窄方向梯度正负交替，加总后抵消→震荡减小')
bullet('小坑冲过去：小坑的梯度挡不住累积的动量')
para('')
para('比喻：滚球下山 vs 蒙眼下山', bold=True)
para('  SGD：每一步都是"站在原地探路→迈步"')
para('  Momentum：每一步是"带着之前的冲劲继续跑"')
para('')

heading('2.3 Momentum 解决了 SGD 的什么问题？', 2)

add_table(
    ['SGD 的问题', 'Momentum 的解法'],
    [
        ['峡谷震荡', '正负振荡在动量中相互抵消，运动方向更平滑'],
        ['平缓区域慢', '持续同向梯度累积动量，加速前进'],
        ['可能卡在小局部极小值', '累积动量可以冲过小坑'],
    ]
)

para('')
para('Momentum 的 PyTorch 代码：')
code_block('optimizer = torch.optim.SGD(model.parameters(), lr=0.001, momentum=0.9)')
para('注意：PyTorch 的 SGD 加上 momentum 参数就是 Momentum 优化器。')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. RMSProp
# ═══════════════════════════════════════════════════════════════════
heading('第三章：RMSProp — 自适应学习率', 1)

heading('3.1 核心公式', 2)
code_block('E[g²] = β·E[g²] + (1-β)·g²     # β 通常 = 0.999')
code_block('θ = θ - lr / √(E[g²] + ε) · g')
para('')
para('符号说明：', bold=True)
bullet('g = ∇L：当前梯度')
bullet('E[g²] = 梯度平方的指数移动平均（梯度的"历史大小"）')
bullet('β = 衰减率（0.999），控制记忆长度')
bullet('ε = 1e-8，防除零')
bullet('lr / √(E[g²] + ε)：每个参数独立的学习率')

heading('3.2 直觉理解：登山杖', 2)
para('想象你下山时每只脚都有一根登山杖，各自独立探路：', bold=True)
para('')
bullet('梯度大的方向（陡坡）：E[g²] 大 → 分母大 → 步长自动变小（陡坡小步走）')
bullet('梯度小的方向（缓坡）：E[g²] 小 → 分母小 → 步长自动变大（缓坡大步走）')
para('')

heading('3.3 具体例子看区别', 2)
para('假设参数 θ₁ 的梯度一直是 100，θ₂ 的梯度一直是 0.01：')
para('')

add_table(
    ['', 'SGD（同一步长）', 'RMSProp（自适应）'],
    [
        ['θ₁ 梯度 = 100', '步长 = lr × 100 → 剧烈震荡', 'E[g²] 大 → lr/10 → 步长适中'],
        ['θ₂ 梯度 = 0.01', '步长 = lr × 0.01 → 几乎不动', 'E[g²] 小 → lr/0.1 → 步长放大'],
    ]
)

para('')
para('RMSProp 让 θ₁ 走得稳、θ₂ 走得快——每个参数有自己的节奏。', bold=True)

heading('3.4 RMSProp 解决了 Momentum 没解决的问题', 2)
para('Momentum 解决了"方向"问题（往哪走），但没解决"步长"问题（走多大）。', bold=True)
para('')
bullet('有些参数需要大步（平坦区域）')
bullet('有些参数需要小步（陡峭区域）')
bullet('RMSProp 让每个参数独立调节自己的步长')

para('')
para('RMSProp 的 PyTorch 代码：')
code_block('optimizer = torch.optim.RMSprop(model.parameters(), lr=0.001, alpha=0.999)')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. Adam
# ═══════════════════════════════════════════════════════════════════
heading('第四章：Adam — 集大成者', 1)
para('— 当前最主流的默认优化器', color=RGBColor(0x66, 0x66, 0x66))

heading('4.1 Adam = Momentum + RMSProp', 2)
para('Adam 把前面两个优化器的优点合在一起：', bold=True)

add_table(
    ['部件', '来源', '作用'],
    [
        ['动量 m_t', 'Momentum', '知道往哪个方向走（惯性）'],
        ['自适应步长 v_t', 'RMSProp', '知道走多大步（每个参数独立）'],
        ['偏差校正', 'Adam 原创', '解决前几步的估计偏差'],
    ]
)

para('')

heading('4.2 完整公式（逐步拆解）', 2)

para('第一步：计算梯度', bold=True)
code_block('g_t = ∇L(θ_{t-1})   # 和前面一样')
para('')

para('第二步：动量（一阶矩估计）', bold=True)
code_block('m_t = β₁·m_{t-1} + (1-β₁)·g_t   # β₁ = 0.9')
para('这步和 Momentum 几乎一样：历史梯度的加权平均。')
para('')

para('第三步：自适应步长（二阶矩估计）', bold=True)
code_block('v_t = β₂·v_{t-1} + (1-β₂)·g_t²  # β₂ = 0.999')
para('这步和 RMSProp 一样：历史梯度平方的加权平均。')
para('')

para('第四步：偏差校正', bold=True)
code_block('m̂_t = m_t / (1 - β₁^t)   # t 是当前步数')
code_block('v̂_t = v_t / (1 - β₂^t)')
para('为什么需要偏差校正？', bold=True)
para('初始化时 m₀ = 0, v₀ = 0。前几步 m_t 和 v_t 严重偏向 0，除以 (1 - βᵗ) 把它们"拉回来"。')
para('第 1 步：β₁¹ = 0.9，1 - 0.9 = 0.1 → m̂ = m₁ / 0.1，放大 10 倍')
para('第 10 步：β₁¹⁰ ≈ 0.35，1 - 0.35 = 0.65 → m̂ = m₁₀ / 0.65，放大 ~1.5 倍')
para('第 100 步：β₁¹⁰⁰ ≈ 0.00003 → 校正基本消失（1 / 0.99997 ≈ 1）')
para('')

para('第五步：参数更新', bold=True)
code_block('θ_t = θ_{t-1} - lr · m̂_t / (√v̂_t + ε)')
para('')
para('最终形式 = SGD的骨架 + Momentum的方向 + RMSProp的步长 + 偏差校正', bold=True)

heading('4.3 直观理解：一个聪明的登山者', 2)
para('Adam 就像一个经验丰富的登山者：')
para('')
bullet('记得之前往哪个方向走（动量 → 保持方向）')
bullet('路陡就小步，路平就大步（自适应 → 调步长）')
bullet('刚开始走得谨慎，慢慢放开步子（偏差校正 → 前几步修正）')
para('')
para('他不一定每一步走得最快，但综合下来最早到山脚。')

heading('4.4 在 Week 9 Part 4 中的代码', 2)
para('这就是你在 week9_complete.py 中实际使用的：')
code_block('optimizer = torch.optim.Adam(model.parameters(), lr=0.001)')
para('')
para('默认参数 β₁=0.9, β₂=0.999, ε=1e-8——大多数任务直接可用，不需要调。')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. 对比总表
# ═══════════════════════════════════════════════════════════════════
heading('第五章：四优化器终极对比', 1)

add_table(
    ['特性', 'SGD', 'Momentum', 'RMSProp', 'Adam'],
    [
        ['核心思想', '梯度下降', '惯性加速', '自适应步长', '前两者结合'],
        ['公式', 'θ -= lr·g', 'v=γv+lr·g\nθ -= v', 'E=βE+(1-β)g²\nθ -= lr/√E·g', 'm=β₁m+(1-β₁)g\nv=β₂v+(1-β₂)g²\n偏差校正\nθ -= lr·m/√v'],
        ['记忆内容', '无', '一阶动量 v', '二阶动量 E[g²]', '一阶 + 二阶'],
        ['每个参数独立步长', '✗', '✗', '✓', '✓'],
        ['惯性/动量', '✗', '✓(γ=0.9)', '✗', '✓(β₁=0.9)'],
        ['收敛速度', '慢', '中', '中-快', '快'],
        ['超参数敏感度', '高（lr 最关键）', '中', '低', '很低（默认 lr=0.001 即可）'],
        ['内存开销', '1×（参数数）', '2×（参数+动量）', '2×（参数+二阶矩）', '3×（参数+动量+二阶矩）'],
        ['PyTorch 代码', 'SGD(lr)', 'SGD(lr,\nmomentum=0.9)', 'RMSprop(lr,\nalpha=0.999)', 'Adam(lr)'],
    ]
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. 进化脉络
# ═══════════════════════════════════════════════════════════════════
heading('第六章：四优化器的进化脉络', 1)
para('这张图帮你理解四个优化器之间的联系：')
para('')

code_block('''SGD                     # 蒙眼下山 — 最朴素
   │
   ├── +惯性 ──→ Momentum      # 滚球下山 — 解决震荡
   │
   └── +自适应 ─→ RMSProp      # 登山杖 — 解决步长
                    │
                    └── 两者合并 ──→ Adam  # 集大成者，默认选择''')
para('')
para('Adam 是这条进化链的终点吗？', bold=True)
para('不一定。近年来有一些新优化器：')
bullet('AdamW：修复 Adam 的权重衰减实现方式（PyTorch 推荐用 AdamW 替代 Adam）')
bullet('LAMB：针对大 batch size 的 Adam 变体')
bullet('LION：Google 2023 年提出的新优化器，内存更省')
para('')
para('但对你当前的学习阶段：Adam 已经够用且是默认选择。', bold=True)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. 面试 Q&A
# ═══════════════════════════════════════════════════════════════════
heading('第七章：面试常问问题', 1)

para('Q1: "Adam 和 SGD 有什么区别？什么时候用哪个？"', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('Adam 自动调学习率 + 动量加速，默认 lr=0.001 大多数任务直接可用。SGD 需要精细调学习率，对新手不友好。但 SGD+Momentum 在某些 CV 任务上泛化性更好。建议：新手直接用 Adam，资深玩家可以试 SGD+Momentum。')
para('')
para('Q2: "Adam 的三个超参数 β₁, β₂, ε 分别控制什么？"', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('β₁=0.9：动量衰减率。越大历史梯度贡献越大，更新越平滑。β₂=0.999：二阶矩衰减率。越大自适应步长变化越慢。ε=1e-8：防除零。这三个默认值在大多数情况下不需要改。')
para('')
para('Q3: "Momentum 的 γ 为什么通常设 0.9？"', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('γ=0.9 意味着最近 10 步的梯度贡献了约 65% 的动量（1-0.9¹⁰≈0.65）。如果 γ 太小（如 0.5），动量记忆太短，惯性效果不明显。如果 γ 太大（如 0.99），动量记忆太长，更新过于平滑可能反应迟钝。0.9 是经验最优值。')
para('')
para('Q4: "RMSProp 和 Adagrad 有什么区别？"', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('Adagrad 是 RMSProp 的前身。Adagrad 累积所有历史梯度的平方（不加权），导致步长持续衰减到零（学不动了）。RMSProp 用指数移动平均替代累积和，步长不会衰减到零，因此可以持续学习。')
para('')
para('Q5: "为什么你的 MLP 训练中用了 Adam 而不是 SGD？"', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para('（面试中结合项目的回答）"因为 Adam 是当前最主流的默认优化器，lr=0.001 直接可用，不需要花时间调参。我们的目标是验证 MLP 做功率预测的可行性，Adam 能快速收敛到满意结果。如果后期需要部署到嵌入式设备（内存受限），可能会换成 SGD 来节省一半的优化器状态内存。"')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 8. 在代码中观察区别
# ═══════════════════════════════════════════════════════════════════
heading('第八章：动手试：在代码中观察四个优化器的区别', 1)

para('如果你想亲眼看到四个优化器的行为差异，Week 9 的 MLP 代码稍微改一下就行：')
para('')

code_block('''# 把你的 Adam 换成下面四个之一，观察收敛速度
optimizer = torch.optim.SGD(model.parameters(), lr=0.001)
# optimizer = torch.optim.SGD(model.parameters(), lr=0.001, momentum=0.9)
# optimizer = torch.optim.RMSprop(model.parameters(), lr=0.001)
# optimizer = torch.optim.Adam(model.parameters(), lr=0.001)

for epoch in range(n_epochs):
    ... # 训练循环不变''')

para('')
para('你会发现：', bold=True)
bullet('SGD：收敛最慢，loss 下降曲线最曲折')
bullet('Momentum：比 SGD 快，震荡减少')
bullet('RMSProp：收敛更快，曲线平滑')
bullet('Adam：最快，最平滑，几乎不用调参')

# ═══════════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════════
path = os.path.join(OUT_DIR, '优化器精讲_SGD_Momentum_RMSProp_Adam.docx')
doc.save(path)
print(f'OK: {path}')
