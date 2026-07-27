#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Part 8 值迭代 大白话精讲文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

OUT_DIR = r'F:\CLAUDE\research\ems-platform\docs\notes'
CODE_FILE = r'F:\CLAUDE\research\ems-platform\scripts\week9_complete.py'

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def para(text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    return p

def code_block(lines, label=None):
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
    for line in lines.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def table_row(table, cells, bold=False):
    row = table.add_row()
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.name = '微软雅黑'
        run.font.size = Pt(10)
        run.bold = bold

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph('')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('Week 9 · Part 8\n值迭代（Value Iteration）')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t2.add_run('大白话逐行精讲 + 手算例子')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t3.add_run(f'\n生成日期：{datetime.date.today().isoformat()}')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第1章：值迭代到底是什么
# ═══════════════════════════════════════════════════════════════════
heading('第1章：值迭代到底是什么', 1)

heading('1.1 一个直觉故事', 2)
para('想象你是一个刚到一个陌生城市的人，想知道"从我住的酒店走到火车站最快怎么走"。')
para('')
para('策略迭代（Part 7）的做法：', bold=True)
para('"我先试探着往一个方向走，看看能不能到——哦能到，但绕路了。那我换个方向试试……"')
para('  每走一次就重新评估整个路线，然后调整方向。反复试几次就找到最优路线了。')
para('')
para('值迭代（Part 8）的做法：', bold=True)
para('"我不急着走。我先在脑子里把每个路口"值多少钱"算清楚——')
para('  火车站值 10 分，火车站旁边的路口值 9 分，再远一点的值 8 分……')
para('  等我所有路口的分数都算准了，我看哪个方向分数最高就往哪走。"')

heading('1.2 核心思想一句话', 2)
para('值迭代不问"我现在该往哪走"，它只问"每个格子值多少分"。')
para('分算准了，方向自然就知道了。', bold=True, size=12, color=RGBColor(0xC0, 0x39, 0x2B))

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第2章：手算演示（4个格子的极简世界）
# ═══════════════════════════════════════════════════════════════════
heading('第2章：手算演示', 1)
para('为了让你真正看懂，我们用一个 4 格子的世界来手算：')
para('')

# 画格子
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
cells_data = [
    ('[0] 起点', '[1]'),
    ('[2]', '[3] 目标 G\n+1 分'),
]
for r in range(2):
    for c in range(2):
        cell = table.cell(r, c)
        cell.text = cells_data[r][c]
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(12)

para('')
para('规则：', bold=True)
bullet('γ = 0.9（折扣因子，未来分数打 9 折）')
bullet('走到 G 得 +1 分，其他格子走一步得 0 分')
bullet('撞墙原地不动')
bullet('动作：↑ ↓ ← →（为简化，假设 100% 走对）')

heading('2.1 初始化', 2)
para('一开始所有格子都不知道自己值多少分：')
code_block('V = [0, 0, 0, 0]   ← 格子 0,1,2,3 的价值都是 0')

heading('2.2 第 1 轮迭代', 2)
para('对每个格子，问：如果我能选最优动作，我能得多少分？')
para('')
para('格子 3（目标 G）：', bold=True)
code_block('上、下、左、右四个方向的 Q 值都是 0（因为周围没有奖励）')
para('但等等——格子 3 本身就是目标，走到 G 有 +1 的即时奖励！', bold=True)
para('格子 3 的奖励函数 R[3][*]，停在 G 会得到 +1 分。')
para('所以 V[3] = max(0, 0, 0, 0, +1) = 1')
para('')
para('格子 0（起点）：', bold=True)
code_block('往右 → Q = R[0][右] + 0.9 × V[1] = 0 + 0.9 × 0 = 0')
para('所有方向都是 0，所以 V[0] = max(0,0,0,0) = 0')
para('')
para('第 1 轮结果：', bold=True)
code_block('V = [0, 0, 0, 1]')
para('只有目标 G 知道了自己是 1 分。')

heading('2.3 第 2 轮迭代', 2)
para('格子 1：', bold=True)
code_block('往下 → Q = 0 + 0.9 × V[3] = 0 + 0.9 × 1 = 0.9')
para('所以 V[1] = 0.9')
para('')
para('格子 2：', bold=True)
code_block('往右 → Q = 0 + 0.9 × V[3] = 0 + 0.9 × 1 = 0.9')
para('所以 V[2] = 0.9')
para('')
para('格子 0：', bold=True)
code_block('往右 → Q = 0 + 0.9 × V[1] = 0 + 0.9 × 0.9 = 0.81')
para('往下 → Q = 0 + 0.9 × V[2] = 0 + 0.9 × 0.9 = 0.81')
para('所以 V[0] = max(0, 0.81, 0.81, 0) = 0.81')
para('')
para('第 2 轮结果：', bold=True)
code_block('V = [0.81, 0.9, 0.9, 1]')

heading('2.4 第 3 轮', 2)
code_block('V[1] = 0.9 × V[3]  = 0.9 × 1  = 0.9    ← 不变')
code_block('V[2] = 0.9 × V[3]  = 0.9 × 1  = 0.9    ← 不变')
code_block('V[0] = 0.9 × V[1]  = 0.9 × 0.9 = 0.81   ← 不变')
para('')
para('V 不再变化 → 收敛！✅', bold=True, size=12, color=RGBColor(0x27, 0xAE, 0x60))

heading('2.5 从 V 提取策略', 2)
para('现在 V 算准了，回头看每个格子哪个动作最好：')
para('')
para('格子 0：', bold=True)
code_block('Q(右) = 0 + 0.9 × V[1] = 0.9 × 0.9 = 0.81')
code_block('Q(下) = 0 + 0.9 × V[2] = 0.9 × 0.9 = 0.81')
code_block('Q(左) = 0 + 0.9 × V[0] = 0')
code_block('Q(上) = 0 + 0.9 × V[0] = 0')
para('→ 格子 0 的最优动作是 右 或 下（这里取第一个，右）')
para('')
para('格子 1：', bold=True)
code_block('Q(下) = 0 + 0.9 × V[3] = 0.9    → 最大')
para('→ 格子 1 的最优动作是 下')
para('')
para('格子 2：', bold=True)
code_block('Q(右) = 0 + 0.9 × V[3] = 0.9    → 最大')
para('→ 格子 2 的最优动作是 右')
para('')
para('最终策略：', bold=True)
code_block('格子 0 → 右     格子 1 → 下     格子 2 → 右     格子 3 → G')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第3章：看实际代码（GridWorld 4×4）
# ═══════════════════════════════════════════════════════════════════
heading('第3章：实际代码逐行看', 1)
para('回到 week9_complete.py 的真实代码。')

heading('3.1 主循环', 2)
para('第 530-545 行：')
code_block('''for iteration in range(1000):
    delta = 0
    for s in range(n_states):      # 遍历 16 个格子
        v_old = V[s]               # 记下旧的分数
        q_max = -np.inf            # 准备找最好的 Q 值
        for a in range(n_actions): # 试上下左右
            q = R[s][a]            # 基础分
            for s_next, prob in P[s][a].items():
                q += gamma * prob * V[s_next]  # 加上"未来分"
            if q > q_max:
                q_max = q          # 只保留最大的！
        V[s] = q_max               # V(s) = max Q(s,a)
        delta = max(delta, abs(v_old - V[s]))
    if delta < theta:              # 所有格子分数都不变了
        break                      # 收敛！''')

heading('3.2 关键区别：max vs 平均', 2)
para('对比 Part 6（策略评估）和 Part 8（值迭代）：')
para('')

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['', 'Part 6 策略评估', 'Part 8 值迭代']):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10)

rows_data = [
    ['公式', 'V(s) = Σ π(a|s) · Q(s,a)', 'V(s) = max_a Q(s,a)'],
    ['含义', '所有动作的加权平均', '只取最好的动作'],
    ['像什么', '"大家投票，取平均意见"', '"让最懂的人做决定"'],
]
for data in rows_data:
    row = table.add_row()
    for c, txt in enumerate(data):
        cell = row.cells[c]
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'

heading('3.3 收敛后提取策略', 2)
para('第 547-555 行：')
code_block('''# V 已经收敛了，现在一次性提取策略
policy = np.zeros(n_states, dtype=int)
for s in range(n_states):
    q_values = []
    for a in range(n_actions):      # 对每个动作算 Q 值
        q = R[s][a]
        for s_next, prob in P[s][a].items():
            q += gamma * prob * V[s_next]
        q_values.append(q)
    policy[s] = int(np.argmax(q_values))  # 选最好的动作''')
para('')
para('整个迭代过程中没有 policy 变量。收敛后只花 16 步（每个格子算一次）就得到最优策略。', bold=True)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第4章：结果解读
# ═══════════════════════════════════════════════════════════════════
heading('第4章：结果长什么样', 1)

heading('4.1 策略图', 2)
para('运行后输出类似：')
code_block('''  ↑  ↑  →  G
  ↑  →  →  ↓
  ↑  ↑  ↓  ↓
  ↑  →  →  ↓''')
para('每个箭头表示该格子"往哪走最好"。所有格子都指向通往 G 的路径。')

heading('4.2 收敛曲线', 2)
para('Extra 部分画了两张图：')
para('')
para('图 1：V(s) 随迭代次数的变化', bold=True)
bullet('起点 V(0) 从 0 开始，逐渐上升到 ~0.52')
bullet('目标 V(15) 瞬间跳到 1，然后不变')
bullet('陷阱 V(5) 变成负数（-1 左右）')
para('')
para('图 2：每轮最大变化量 ΔV', bold=True)
bullet('y 轴是对数尺度')
bullet('ΔV 指数衰减——画出来是一条直线')
bullet('这是因为 Bellman 算子是压缩映射，每次压缩 γ=0.9 倍')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第5章：Part 7 vs Part 8
# ═══════════════════════════════════════════════════════════════════
heading('第5章：Part 7 vs Part 8 对比', 1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

data = [
    ['对比维度', '策略迭代 (Part 7)', '值迭代 (Part 8)'],
    ['核心公式', '策略评估: V = E_π[...]\n策略改进: π = argmax Q', 'V = max_a Q(s,a)'],
    ['显式策略', '每一轮都维护一个策略', '迭代过程中没有策略'],
    ['收敛轮数', '少（本例 4 轮）', '多（本例 ~133 轮）'],
    ['每轮代价', '大（内循环要算到 V 收敛）', '小（只更新一次 V）'],
    ['提取策略', '改进步骤中自动得到', '收敛后反推'],
    ['保证', '单调改进，一定收敛到最优', '压缩映射，一定收敛到 V*'],
    ['状态空间', '小到中等', '任意大小（更通用）'],
]
for r, row_data in enumerate(data):
    if r == 0:
        cells = table.rows[0].cells
    else:
        row = table.add_row()
        cells = row.cells
    for c, txt in enumerate(row_data):
        cell = cells[c]
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        if r == 0:
            run.bold = True

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第6章：常见疑问
# ═══════════════════════════════════════════════════════════════════
heading('第6章：常见疑问', 1)

para('Q: 值迭代每轮都只取 max，会不会错过次优解？', bold=True, size=11, color=RGBColor(0xC0, 0x39, 0x2B))
para('不会。Bellman 最优方程有数学保证：反复应用 V := max_a Q(s,a) 一定收敛到 V*，即真正的最优值函数。从 V* 提取的策略就是最优策略。')
para('')

para('Q: 为什么值迭代轮数这么多（133 轮）而策略迭代只要 4 轮？', bold=True, size=11, color=RGBColor(0xC0, 0x39, 0x2B))
para('策略迭代每轮把 V 算到完全收敛才改进，一步迈得大。值迭代每次只走一小步（一轮 Bellman 更新），所以需要更多步。')
para('')

para('Q: 那值迭代不是更差吗？', bold=True, size=11, color=RGBColor(0xC0, 0x39, 0x2B))
para('不一定。虽然轮数多，但每轮计算量小。对 GridWorld 这样 16 个状态的小问题，两者没区别。对大问题（几万、几百万状态），值迭代因为每轮简单所以更实用。')
para('')

para('Q: 这个跟 EMS 有什么关系？', bold=True, size=11, color=RGBColor(0xC0, 0x39, 0x2B))
para('EMS 中的 DP 求解器也是用 Bellman 最优方程：状态 = SOC × 功率请求，动作 = FC 功率分配，奖励 = -氢耗。只是 GridWorld 的 16 个状态变成了 ~9000 个（150 个 SOC × 60 个功率请求）。数学原理完全一样。')

page_break()

# ═══════════════════════════════════════════════════════════════════
# 第7章：代码与文档导航
# ═══════════════════════════════════════════════════════════════════
heading('第7章：如果你还是觉得模糊', 1)
para('把 Part 8 和 Part 6 对照着看：')
bullet('Part 6 是"给定一个策略，算它值多少分"——期望方程')
bullet('Part 8 是"不管策略，直接算每个格子最高能值多少分"——最优方程')
para('')
para('两者的代码几乎一模一样，区别只有一行：', bold=True)
code_block('Part 6:  V[s] = sum(p_a * q)     # 加权平均')
code_block('Part 8:  V[s] = max(q_values)    # 取最大值')
para('')
para('如果你想实际操作一下：')
para('')
para('运行命令：', bold=True)
code_block('cd F:\\CLAUDE\\research\\ems-platform')
code_block('python scripts/week9_complete.py --part 8', label='终端执行：')
para('')
para('或者用 Python：', bold=True)
code_block('from scripts.week9_complete import part5_mdp_gridworld, part8_value_iteration')
code_block('mdp = part5_mdp_gridworld()')
code_block('result = part8_value_iteration(mdp)')
code_block('print("最优策略:", result["policy"])')
code_block('print("V(0) =", result["V"][0])')

# ═══════════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════════
path = os.path.join(OUT_DIR, 'Week9_Part8_值迭代_大白话精讲.docx')
doc.save(path)
print(f'OK: {path}')
