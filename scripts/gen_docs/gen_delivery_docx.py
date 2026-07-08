#!/usr/bin/env python3
"""生成 EMS/BMS 投递清单 DOCX 文档"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# 封面
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EMS/BMS 算法方向 · 投递清单（南方十城70家）')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('整理日期：2026-06-18 | 211本+双非硕 | 控制算法方向')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 使用说明
doc.add_heading('使用说明', level=2)
for item in [
    '投递时间：2026年8月（秋招提前批）~ 10月',
    '投递节奏：先C档保底 → 再B档稳 → 最后A档冲刺',
    '投递渠道：优先官网/公众号，其次BOSS直聘/牛客网',
    '简历关键词：DP全局最优、ECMS实时优化、SOC过充BUGFIX、多工况验证、MPC、C++',
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# 表格辅助函数
def add_table(doc, data, col_widths=None):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    return table

def section_header(doc, title, color=(180, 0, 0)):
    doc.add_heading(title, level=2)

# A档
section_header(doc, 'A档 · 冲一冲（20家）')
doc.add_paragraph('项目写好了有机会，简历关可能卡但值得试试')

add_table(doc, [
    ['#', '公司', '城市', '岗位方向', '预计年薪', '投递入口'],
    ['1', '比亚迪 (BYD)', '深圳', 'BMS/电控算法', '20-30万', 'job.byd.com 校园招聘'],
    ['2', '汇川技术', '深圳', '电机控制/能源算法', '22-35万', 'inovance.zhiye.com/Campus'],
    ['3', '欣旺达动力', '深圳', 'BMS/域控', '18-28万', 'sunwoda.com → 加入我们'],
    ['4', '豪鹏科技 (001283)', '深圳', 'BMS软件/AI仿真', '22-37万', 'highpowertech.com'],
    ['5', '英飞源', '深圳', 'EMS/BMS/算法', '20-30万', 'inflytech.com → 加入我们'],
    ['6', 'EcoFlow正浩创新', '深圳', 'BMS软件', '22-35万', 'jobs.ecoflow.com'],
    ['7', '高特电子 (301669)', '杭州', '储能BMS算法', '20-30万', 'cngt.com → 招贤纳士'],
    ['8', '科工电子 (874963)', '杭州', '储能BMS/EMS', '18-28万', 'kegongdz.com → 加入我们'],
    ['9', '高泰昊能', '杭州', '动力BMS/EMS', '18-26万', 'qualtech.com.cn'],
    ['10', '华塑科技 (301157)', '杭州', '储能BMS', '16-24万', 'huasucn.com'],
    ['11', 'Ampace新能安', '厦门', 'BMS/EMS算法', '18-35万(18薪)', 'ampace.hotjob.cn'],
    ['12', '东风汽车集团', '武汉', 'BMS算法/能量管理', '20-35万', 'dfmc.com.cn → 人才招聘'],
    ['13', '楚能新能源', '武汉', 'BMS/EMS系统开发', '18-36万', 'cnenergy.com.cn'],
    ['14', '骆驼集团(武汉光谷)', '武汉', 'BMS工程师-软件', '20-28万', '一览电池英才网'],
    ['15', '经纬恒润', '西安/南京/武汉', 'BMS/VCU算法', '22-45万', 'zhaopin.hirain.com'],
    ['16', '三一集团', '长沙', 'BMS算法工程师', '30-60万', 'sany.com.cn → 加入三一'],
    ['17', '华自科技 (300490)', '长沙', '储能EMS/算法', '15-20万', 'cshnac.com'],
    ['18', '中车四方(科创二十所)', '北京', '能量管理算法', '20-35万', '高校人才网搜索'],
    ['19', '亿纬锂能', '惠州', 'BMS/控制算法', '18-28万', 'wecruit.hotjob.cn'],
    ['20', '阳光电源', '合肥', '储能EMS/控制', '20-35万', 'sungrowpower.zhiye.com'],
])
doc.add_paragraph()

# B档
section_header(doc, 'B档 · 稳一稳（30家）')
doc.add_paragraph('大概率能拿面试，重点投递对象')

b_data = [
    ['#', '公司', '城市', '岗位方向', '预计年薪', '投递入口'],
    ['21', '新能德(NVT)', '东莞', 'BMS算法', '16-24万', 'nvt.com.cn → 加入我们'],
    ['22', 'ATL(新能源科技)', '东莞', 'BMS算法', '18-26万', 'atlbattery.com'],
    ['23', '华宝新能', '深圳', 'BMS/EMS', '15-22万', 'hellobms.com'],
    ['24', '拓邦股份 (002139)', '深圳', '电池控制/BMS', '14-20万', 'topband.com.cn'],
    ['25', '德赛电池 (000049)', '惠州/深圳', 'BMS/EMS', '14-22万', 'desaybattery.com'],
    ['26', '华阳集团 (002906)', '惠州', 'BMS/控制算法', '14-20万', 'foryougroup.com'],
    ['27', '首航新能源', '深圳', '储能BMS/EMS', '13-18万', 'solarinverter.com.cn'],
    ['28', '古瑞瓦特', '深圳', '储能控制', '13-18万', 'growatt.com'],
    ['29', '广汽埃安', '广州', 'VCU/BMS控制', '18-25万', 'campus.gac-nio.com'],
    ['30', '小鹏汽车', '广州', '能量管理/控制', '20-30万', 'campus.xiaopeng.com'],
    ['31', '菲利斯太阳能', '广州', 'BMS软件', '14-20万', 'BOSS直聘搜索'],
    ['32', '正泰电源', '杭州/上海', '储能EMS', '14-20万', 'chintpowersource.com'],
    ['33', '禾迈股份 (688032)', '杭州', '储能控制', '14-20万', 'hoymiles.com'],
    ['34', '昱能科技 (688348)', '嘉兴', '储能EMS', '13-18万', 'ynf-tech.com'],
    ['35', '谱地新能源', '杭州', 'BMS/嵌入式', '12-16万', 'BOSS/智联搜索'],
    ['36', '索克曼能源', '厦门', '软件（UPS/BMS）', '10-12万', '厦门人才网'],
    ['37', '北辰星储能', '厦门', 'BMS/EMS技术', '8-11万', 'BOSS直聘搜索'],
    ['38', '武汉菱电电控 (688667)', '武汉', 'EMS/电控', '15-22万', 'lingdian.com.cn'],
    ['39', '华昱欣科技', '武汉/全国', '控制软件/嵌入式', '15-22万', '2026届校招'],
    ['40', '彼欧(Opmobility)', '武汉', 'BMS工程师', '25-40K×15薪', '猎聘搜索'],
    ['41', '中步擎天新能源', '武汉', '储能系统', '10-16万', 'BOSS直聘搜索'],
    ['42', '华思系统', '长沙', 'BMS/EMS嵌入式', '15-22万', '2026届校招'],
    ['43', '红太阳新能源', '长沙', '光储BMS/EMS', '18-24万', '北极星招聘'],
    ['44', '湖南奕航新能源', '长沙', 'BMS/EMS', '8-15万', '一览电池英才网'],
    ['45', '中创新航 (CALB)', '常州', 'BMS算法校招', '13-18万', 'calb.com.cn'],
    ['46', '派能科技', '上海', '储能EMS/BMS', '18-26万', 'pylontech.com.cn'],
    ['47', '固德威', '苏州', '储能EMS', '16-24万', 'goodwe.zhiye.com'],
    ['48', '国轩高科', '合肥/南京', 'BMS算法', '18-26万', 'gotion.com.cn'],
    ['49', '飞毛腿动力科技', '福州', 'BMS算法/SOX算法', '15-30K/月', '福州人才网/一览电池'],
    ['50', '东南汽车', '福州', '助理电控（BMS）', '6-11K/月', 'BOSS直聘搜索'],
]
add_table(doc, b_data)
doc.add_paragraph()

# C档
section_header(doc, 'C档 · 保底（20家）')
doc.add_paragraph('先拿offer稳住心态，再冲更好的')

add_table(doc, [
    ['#', '公司', '城市', '岗位方向', '预计年薪', '投递入口'],
    ['51', '和而泰 (002402)', '深圳', 'BMS/控制', '12-16万', 'szhittech.com'],
    ['52', '英威腾 (002334)', '深圳', '储能控制', '13-20万', 'invt.com.cn'],
    ['53', '科陆电子 (002121)', '深圳', '储能EMS', '13-18万', 'clou.com.cn'],
    ['54', '盛弘股份 (300693)', '深圳', '储能BMS', '14-20万', 'sinexcel.com'],
    ['55', '科士达 (002518)', '深圳', '储能EMS', '12-18万', 'kstar.com.cn'],
    ['56', '朗科智能 (300543)', '深圳', 'BMS/控制', '13-18万', 'longood.com'],
    ['57', '南科动力', '深圳', '氢燃料控制算法', '16-24万', 'BOSS直聘搜索'],
    ['58', '深圳南山热电', '深圳', '储能EMS', '14-20万', '北极星招聘'],
    ['59', '广州奥鹏能源', '广州', 'BMS软件', '12-16万', 'BOSS直聘搜索'],
    ['60', '智光电气 (002169)', '广州', '储能EMS', '12-18万', 'gzzg.com.cn'],
    ['61', '杉杉股份 (600884)', '宁波/上海', 'BMS/能源管理', '14-22万', 'shanshantech.com'],
    ['62', '亿晶光电 (600537)', '常州', '储能系统', '12-16万', 'egingpv.com'],
    ['63', '武汉珠和辰仁', '武汉', '储能系统', '10-20万', 'BOSS直聘搜索'],
    ['64', '中步擎天新能源', '武汉', '储能系统', '8-16万', 'BOSS直聘搜索'],
    ['65', '智狐能源', '福州', 'EMS（研发方向）', '20K+/月', '北极星招聘'],
    ['66', '索克曼能源', '厦门', '软件工程师', '8-12K/月', '厦门人才网'],
    ['67', '北辰星储能', '厦门', '技术研发', '7-9K/月', 'BOSS直聘搜索'],
    ['68', '拓普菲斯新能源', '厦门', '储能系统（BMS）', '12-18K/月', 'BOSS直聘搜索'],
    ['69', '湖南奕航新能源', '长沙', 'BMS/EMS', '8-15万', '一览电池英才网'],
    ['70', '谱地新能源', '杭州', 'BMS/嵌入式', '12-16万', 'BOSS/智联搜索'],
])

doc.add_paragraph()

# 城市分布
doc.add_heading('城市分布一览', level=2)
add_table(doc, [
    ['城市', '公司数', '定位'],
    ['深圳/东莞/惠州', '21家', '主战场（岗位最多）'],
    ['杭州', '8家', '第二战场（BMS产业密集）'],
    ['武汉', '8家', '第三选择（东风+楚能）'],
    ['长沙', '5家', '三一集团薪资高'],
    ['广州/佛山', '5家', ''],
    ['苏州/常州/宁波', '5家', ''],
    ['厦门', '4家', 'Ampace一家顶几家'],
    ['合肥/上海（近）', '4家', ''],
    ['福州', '3家', '岗位偏少'],
])
doc.add_paragraph()

# 时间表
doc.add_heading('投递时间表', level=2)
timeline = [
    ('2026年7月（现在~）', [
        '写好简历，按岗位方向准备2-3个版本',
        '整理好项目代码（GitHub repo README写清楚）',
        '关注各公司提前批通知',
    ]),
    ('2026年8月（秋招启动）', [
        'C档20家先投 → 拿1-2个保底offer',
        '同时投B档部分公司',
    ]),
    ('2026年9月（秋招高峰）', [
        'B档全部投完，争取更多面试',
        '开始投A档（简历已有保底，心态稳）',
        '面完后复盘，持续迭代面试话术',
    ]),
    ('2026年10月（秋招中后期）', [
        'A档剩余公司',
        '收到满意offer后停止海投',
    ]),
]
for period, items in timeline:
    p = doc.add_paragraph()
    run = p.add_run(period)
    run.bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# 各城市核心目标
doc.add_heading('各城市核心目标总结', level=2)
add_table(doc, [
    ['城市', '最值得冲的公司', '策略'],
    ['深圳', '比亚迪 / 汇川 / 豪鹏', '主战场，大量投'],
    ['杭州', '高特电子（刚上市）', 'BMS产业密集，深耕'],
    ['厦门', 'Ampace新能安（18薪）', '精投这一家'],
    ['武汉', '东风汽车 / 楚能新能源', '国企+民企都有'],
    ['长沙', '三一集团（30-60万）', '三一薪资有惊喜'],
    ['福州', '飞毛腿动力', '岗位少，顺带投'],
    ['广州', '小鹏 / 广汽埃安', '不如深圳密集'],
])
doc.add_paragraph()

# 招聘平台
doc.add_heading('常用招聘平台', level=2)
add_table(doc, [
    ['平台', '网址', '用途'],
    ['公司官网校招', '见各公司链接', '最权威，优先使用'],
    ['BOSS直聘', 'zhipin.com', '回复快'],
    ['牛客网', 'nowcoder.com', '面经+笔试刷题'],
    ['猎聘', 'liepin.com', '中高端岗位'],
    ['北极星招聘', 'hr.bjx.com.cn', '新能源垂直招聘'],
    ['一览电池英才网', 'jdjob88.com', '电池行业垂直'],
    ['应届生求职网', 'yingjiesheng.com', '校招信息汇总'],
    ['智联招聘', 'zhaopin.com', '综合平台'],
    ['厦门人才网', 'xmrc.com.cn', '厦门本地'],
])

# 保存
output_path = os.path.join('F:\\CLAUDE\\research\\ems-platform\\docs', 'EMS_BMS_投递清单70家_南方十城.docx')
doc.save(output_path)
print(f'[OK] 已生成: {output_path}')
