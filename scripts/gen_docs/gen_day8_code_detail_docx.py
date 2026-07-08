# -*- coding: utf-8 -*-
"""生成 Day8 DP 代码详细解读报告 (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')

def add_code_block(doc, code_text):
    """添加代码块样式"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)

def add_note(doc, text):
    """添加注释/提示框"""
    p = doc.add_paragraph()
    run = p.add_run('💡 ')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_heading_numbered(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def main():
    doc = Document()

    # ── 样式设置 ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    # ── 标题 ──
    title = doc.add_heading('day8_dp_ems.py 代码详细解读', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('日期：2026-06-11\n').font.size = Pt(10)
    meta.add_run('总行数：501 行\n').font.size = Pt(10)
    meta.add_run('核心功能：后向 DP + 前向 Rollout + 规则控制器对比\n').font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 一、整体架构
    # ══════════════════════════════════════════
    add_heading_numbered(doc, '整体架构', level=1)

    doc.add_paragraph('day8_dp_ems.py 的主流程如下：')

    arch = doc.add_paragraph()
    arch.paragraph_format.left_indent = Inches(0.5)
    run = arch.add_run(
        '加载工况(vehicle_power)  →  后向DP(backward_dp)\n'
        '                         →  前向Rollout(forward_rollout)\n'
        '                         →  规则对比(run_rule_controller)\n'
        '                         →  指标打印 + 对比图'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph('依赖库：')
    add_code_block(doc, 'import numpy as np      # 矩阵运算、插值\nimport pandas as pd     # CSV 读写\nimport matplotlib        # 画图')

    # ══════════════════════════════════════════
    # 二、参数区
    # ══════════════════════════════════════════
    add_heading_numbered(doc, '参数区（L29-58）', level=1)
    doc.add_paragraph('先看一遍，不用深究。这些是车辆物理参数，用来算"车子以某个速度跑需要多少功率"。')

    add_code_block(doc,
        'MASS = 1500          # 车重 kg\n'
        'G = 9.81             # 重力加速度\n'
        'F_R = 0.015          # 滚动阻力系数\n'
        'RHO = 1.225          # 空气密度\n'
        'CD = 0.32            # 风阻系数\n'
        'AREA = 2.2           # 迎风面积 m²\n'
        'ETA_DRIVE = 0.90     # 传动效率'
    )

    # ── 参数表 ──
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    params = [
        ('SOC 网格数 N_SOC', '150'),
        ('FC 功率网格数 N_PFC', '60'),
        ('SOC 参考值 SOC_REF', '0.6'),
        ('SOC 维持惩罚 α', '100'),
        ('终端 SOC 惩罚 β', '10000'),
        ('SOC 范围', '0.2 ~ 0.9'),
        ('FC 功率范围', '0 ~ 30 kW'),
    ]
    for i, (k, v) in enumerate(params):
        table.rows[i+1].cells[0].text = k
        table.rows[i+1].cells[1].text = v
    table.rows[0].cells[0].text = '参数名'
    table.rows[0].cells[1].text = '值'

    # ══════════════════════════════════════════
    # 三、模型函数
    # ══════════════════════════════════════════
    add_heading_numbered(doc, '模型函数', level=1)

    # 3.1 FC效率
    add_heading_numbered(doc, 'fc_efficiency() — FC 效率查表', level=2)
    doc.add_paragraph('燃料电池的效率不是恒定的。低功率时效率低，过高功率效率也会下降。15kW时效率最高（55%）。')
    doc.add_paragraph('所以能量管理的目标之一就是让FC尽可能在高效区间工作。')

    add_code_block(doc,
        'def fc_efficiency(P_fc):\n'
        '    """FC 效率曲线查表"""\n'
        '    return np.interp(P_fc, PFC_EFF_BP, ETA_FC)'
    )

    table = doc.add_table(rows=10, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'FC 功率 (kW)'
    table.rows[0].cells[1].text = '效率'
    for i, (p, e) in enumerate(zip(
        [0, 2, 5, 8, 10, 15, 20, 25, 30],
        [0, 0.28, 0.40, 0.48, 0.50, 0.55, 0.53, 0.48, 0.40]
    )):
        table.rows[i+1].cells[0].text = str(p)
        table.rows[i+1].cells[1].text = f'{e:.0%}'

    # 3.2 氢耗
    add_heading_numbered(doc, 'fc_hydrogen_flow() — FC 功率 → 氢耗', level=2)
    doc.add_paragraph('公式：氢耗(g/s) = P_fc(W) / (效率 × 氢气热值(J/kg)) × 1000')

    add_code_block(doc,
        'def fc_hydrogen_flow(P_fc):\n'
        '    is_scalar = np.isscalar(P_fc)     # 记下输入是标量还是数组\n'
        '    P_fc = np.atleast_1d(np.asarray(P_fc, dtype=float))\n'
        '    eta = fc_efficiency(P_fc)\n'
        '    with np.errstate(divide="ignore", invalid="ignore"):\n'
        '        mdot = P_fc * 1000 / (eta * LHV_H2) * 1000\n'
        '    mdot[~np.isfinite(mdot)] = 0      # nan/inf → 0\n'
        '    mdot[P_fc == 0] = 0               # 关机时氢耗=0\n'
        '    return float(mdot[0]) if is_scalar else mdot'
    )

    doc.add_paragraph('关键细节：')
    doc.add_paragraph('• P_fc * 1000 — kW 转 W', style='List Bullet')
    doc.add_paragraph('• LHV_H2 = 120e6 — 氢气低热值 (J/kg)', style='List Bullet')
    doc.add_paragraph('• * 1000 — kg 转 g', style='List Bullet')
    doc.add_paragraph('• ~np.isfinite 处理除以零产生的 nan/inf', style='List Bullet')
    doc.add_paragraph('• is_scalar 判断让函数既能传标量（逐点调用）又能传数组（向量化调用）', style='List Bullet')

    # 3.3 车辆动力学
    add_heading_numbered(doc, 'vehicle_power() — 车速 → 功率需求', level=2)
    doc.add_paragraph('输入：车速序列（km/h），输出：功率需求序列（kW）。一共三步：')

    add_heading_numbered(doc, '第一步：算加速度', level=3)
    add_code_block(doc,
        'v_ms = v_kmh / 3.6                    # km/h → m/s\n'
        'a = np.zeros_like(v_ms)               # 初始化加速度数组\n'
        'a[1:-1] = (v_ms[2:] - v_ms[:-2]) / (2 * dt)  # 中心差分\n'
        'a[0] = (v_ms[1] - v_ms[0]) / dt               # 第一个点\n'
        'a[-1] = (v_ms[-1] - v_ms[-2]) / dt            # 最后一个点\n'
        'a = np.clip(a, -3, 3)                 # 限制加速度范围'
    )
    add_note(doc, '中心差分原理：第i个点的加速度 ≈ (v[i+1] - v[i-1]) / (2×Δt)，没用到自己')

    add_heading_numbered(doc, '第二步：算轮边功率', level=3)
    doc.add_paragraph('车子要克服三个力才能动：')
    doc.add_paragraph('• 滚动阻力 F_rr = MASS × G × F_R（常数，与速度无关）', style='List Bullet')
    doc.add_paragraph('• 空气阻力 F_aero = 0.5 × RHO × CD × AREA × v²（与速度平方成正比）', style='List Bullet')
    doc.add_paragraph('• 加速惯性力 F_inertia = MASS × a（F = ma）', style='List Bullet')

    add_code_block(doc,
        'F_rr = MASS * G * F_R\n'
        'F_aero = 0.5 * RHO * CD * AREA * v_ms ** 2\n'
        'F_inertia = MASS * a\n'
        'P_wheel = (F_rr + F_aero + F_inertia) * v_ms  # P = F × v'
    )

    add_heading_numbered(doc, '第三步：加上传动效率', level=3)
    add_code_block(doc,
        'P_load = np.maximum(P_wheel / ETA_DRIVE / 1000, 0)  # /1000: W→kW\n'
        'P_load[v_kmh < 0.5] = 0                              # 静止时功率=0'
    )
    add_note(doc, 'np.maximum(..., 0) 确保功率不为负（制动能量回收在本模型中未考虑）')

    # 3.4 SOC转移
    add_heading_numbered(doc, 'state_transition() — 电池 SOC 如何变化', level=2)
    doc.add_paragraph('输入：当前SOC、FC输出功率、负载需求功率')
    doc.add_paragraph('输出：下一时刻的SOC')

    add_heading_numbered(doc, '核心逻辑', level=3)
    add_code_block(doc, 'P_bat = P_load_k - P_fc    # 正=放电，负=充电')

    doc.add_paragraph('然后根据简单电池电路模型（内阻模型）求解电流：')
    add_code_block(doc,
        'P_w = P_bat * 1000\n'
        'Delta = V_oc**2 - 4 * R_INT * P_w    # 一元二次方程判别式\n'
        'valid = Delta >= 0\n'
        'I = (V_oc - np.sqrt(Delta[valid])) / (2 * R_INT)'
    )

    add_heading_numbered(doc, '为什么有 Delta ？', level=3)
    doc.add_paragraph('从功率平衡公式推导：')
    doc.add_paragraph('P_bat = V_terminal × I = (V_oc - I×R_int) × I')
    doc.add_paragraph('→ R_int × I² - V_oc × I + P_bat = 0')
    doc.add_paragraph('→ I = (V_oc ± √(V_oc² - 4×R_int×P_bat)) / (2×R_int)')
    doc.add_paragraph('')
    doc.add_paragraph('Delta < 0 表示电池无法提供这么大功率（物理不可行），此时代码保留原SOC。')

    add_heading_numbered(doc, 'SOC更新（安时积分法）', level=3)
    add_code_block(doc, 'SOC_next = SOC_k - I / (Q_BAT * 3600) * dt')
    doc.add_paragraph('I > 0（放电）→ SOC 下降；I < 0（充电）→ SOC 上升。')

    # ══════════════════════════════════════════
    # 四、后向DP
    # ══════════════════════════════════════════
    add_heading_numbered(doc, '后向 DP — 算法核心', level=1)
    doc.add_paragraph('这是整个程序最核心的部分，面试也最喜欢问。')

    add_heading_numbered(doc, '两个关键表格', level=2)

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = '变量'
    table.rows[0].cells[1].text = '含义'
    table.rows[1].cells[0].text = 'J[k][i]'
    table.rows[1].cells[1].text = '在时刻k、SOC为状态i时，从此刻到终点所需的最小累积氢耗'
    table.rows[2].cells[0].text = 'pi[k][i]'
    table.rows[2].cells[1].text = '在时刻k、SOC为状态i时，应输出的最优FC功率'

    add_heading_numbered(doc, '核心递推方程（面试白板就写这个）', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'J[k][i] = min_{p_fc} [ g(p_fc) + α×(SOC_next-SOC_ref)² + J[k+1][lookup(SOC_next)] ]'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('↓ 拆解：')
    doc.add_paragraph('• J[k][i] — 当前状态的最小代价', style='List Bullet')
    doc.add_paragraph('• g(p_fc) — 这一步的氢耗', style='List Bullet')
    doc.add_paragraph('• α×(SOC_next-SOC_ref)² — SOC偏离参考值的即时惩罚', style='List Bullet')
    doc.add_paragraph('• J[k+1][lookup(SOC_next)] — 查表得到"下一个状态到终点的最优代价"', style='List Bullet')

    add_heading_numbered(doc, '三层循环逻辑', level=2)
    doc.add_paragraph('虽然代码里向量化了，但理解时按三层 for 循环想：')
    add_code_block(doc,
        '从最后一个时刻往前推：\n'
        '  for 每个时刻 k (从 N-1 到 0):\n'
        '    for 每个 SOC 状态 i:\n'
        '      for 每个 FC 功率 (0~30kW, 60个网格):\n'
        '        算: 控制后的 SOC_next\n'
        '        算: 这一步的氢耗 g\n'
        '        查: SOC_next 对应的未来最优代价 J_next\n'
        '        算: 总代价 = g + SOC惩罚 + J_next\n'
        '      选: 总代价最小的 FC 功率 → 存入 pi[k][i]'
    )

    add_heading_numbered(doc, '向量化优化', level=2)
    doc.add_paragraph('实际代码用向量化代替了最内层循环，60个FC功率一次算完：')
    add_code_block(doc,
        '# 一次算所有 FC 功率的 SOC_next\n'
        'SOC_next_all = state_transition(soc, PFC_GRID, P_load_k, DT)\n'
        '\n'
        '# 筛选可行的控制\n'
        'feasible = (SOC_next_all >= SOC_MIN) & (SOC_next_all <= SOC_MAX)\n'
        '\n'
        '# 未来代价插值 + SOC惩罚\n'
        'J_future = np.interp(SOC_next_all[feasible], SOC_GRID, J_next_k)\n'
        'J_future += ALPHA * (SOC_next_all[feasible] - SOC_REF) ** 2\n'
        '\n'
        '# 总代价 = 氢耗 + 未来代价，不可行的设为 inf\n'
        'total = np.full(N_PFC, np.inf)\n'
        'total[feasible] = g[feasible] + J_future\n'
        '\n'
        '# 选最小\n'
        'pi[k, i] = PFC_GRID[np.argmin(total)]'
    )

    # ══════════════════════════════════════════
    # 五、前向Rollout
    # ══════════════════════════════════════════
    add_heading_numbered(doc, '前向 Rollout', level=1)
    doc.add_paragraph('和DP后向不同，这是从前向后跑仿真：')
    add_code_block(doc,
        'SOC[0] = SOC_0\n'
        'for k in range(N):\n'
        '    # 查策略表：当前 SOC 对应的最优 FC 功率\n'
        '    pfc = np.interp(SOC[k], SOC_GRID, pi[k, :])\n'
        '\n'
        '    # 更新 SOC\n'
        '    SOC[k+1] = state_transition(SOC[k], pfc, P_load[k], DT)\n'
        '\n'
        '    # 累计氢耗\n'
        '    M_H2[k] = fc_hydrogen_flow(pfc) * DT'
    )
    doc.add_paragraph('整个过程：给定初始SOC → 查表得FC功率 → 算电池功率 → 更新SOC → 到下一点……')

    # ══════════════════════════════════════════
    # 六、规则控制器
    # ══════════════════════════════════════════
    add_heading_numbered(doc, '规则控制器', level=1)
    doc.add_paragraph('作为对比基线，逻辑是工程师经验规则：')
    add_code_block(doc,
        '如果 负载很小(<1kW):\n'
        '    如果 SOC还没满 → FC最小功率\n'
        '    如果 SOC满了 → 关闭FC\n'
        '否则如果 SOC过低(<0.4):  FC多发电，给电池充电\n'
        '否则如果 SOC过高(>0.8):  尽量用电池，少用FC\n'
        '否则（正常范围）:          FC跟随负载功率'
    )
    doc.add_paragraph('这部分不需要逐行看，理解逻辑就行。')

    # ══════════════════════════════════════════
    # 七、主程序
    # ══════════════════════════════════════════
    add_heading_numbered(doc, '主程序流程', level=1)
    add_code_block(doc,
        'def main():\n'
        '    1. 加载工况 → vehicle_power()   # 车速→功率\n'
        '    2. 后向DP  → backward_dp()      # 核心计算\n'
        '    3. 前向仿真 → forward_rollout() # 用策略表跑\n'
        '    4. 规则对比 → run_rule_controller()\n'
        '    5. 打印指标 → print_metrics()\n'
        '    6. 保存CSV + 画图'
    )

    # ══════════════════════════════════════════
    # 八、看懂≠会写
    # ══════════════════════════════════════════
    add_heading_numbered(doc, '看懂 ≠ 会写', level=1)

    doc.add_paragraph('读完上面这些，你应该能回答三个问题：')
    doc.add_paragraph('1. backward_dp() 的三层循环分别遍历什么？', style='List Number')
    doc.add_paragraph('2. state_transition() 里的判别式 Delta 在算什么？', style='List Number')
    doc.add_paragraph('3. 为什么 DP 比规则控制器好？本质原因是什么？', style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('建议做这个练习：')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph('关掉 day8_dp_ems.py，打开一个空白文件，从零开始写：')
    doc.add_paragraph('1. state_transition() 函数', style='List Number')
    doc.add_paragraph('2. backward_dp() 的核心循环体（不考虑向量化，用三层 for 写也行）', style='List Number')
    doc.add_paragraph()
    doc.add_paragraph('写不出来就回头看，然后关掉再写。写出来为止。')

    # ── 保存 ──
    os.makedirs(DOCS_DIR, exist_ok=True)
    out_path = os.path.join(DOCS_DIR, 'Day8_dp_ems代码详细解读.docx')
    doc.save(out_path)
    print(f'[OK] 已保存: {out_path}')

if __name__ == '__main__':
    main()
