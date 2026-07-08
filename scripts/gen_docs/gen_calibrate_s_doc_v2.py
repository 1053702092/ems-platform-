#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 calibrate_s_from_dp 逐行分析 DOCX（中文文件名）"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, 'docs')
os.makedirs(DOCS, exist_ok=True)

doc = Document()

# 样式
sty = doc.styles['Normal']
sty.font.name = 'Calibri'
sty.font.size = Pt(11)
sty.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
sty.paragraph_format.line_spacing = 1.35
sty.paragraph_format.space_after = Pt(4)
for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
cs = doc.styles.add_style('CB', WD_STYLE_TYPE.PARAGRAPH)
cs.font.name = 'Consolas'; cs.font.size = Pt(9)
cs.paragraph_format.space_before = Pt(2)
cs.paragraph_format.space_after = Pt(2)
cs.paragraph_format.left_indent = Cm(0.5)

def code(t): doc.add_paragraph(t, style='CB')
def bold(p, t): r = p.add_run(t); r.bold = True; return r
def sep():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('─' * 50); r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA); r.font.size = Pt(8)

# 封面
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('calibrate_s_from_dp.py'); r.font.size = Pt(28); r.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DP 反推标定 ECMS 等效因子 — 逐行原理解析')
r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('从 Pontryagin 最小原理到等效因子的完整推导\n')
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
doc.add_page_break()

# 目录
doc.add_heading('目录', level=1)
for t in ['概述与顶层直觉',
           '核心数学原理：PMP ↔ ECMS 的桥梁',
           'λ 到 s 的换算推导',
           '函数级逐行分析',
           '输出图表含义',
           '工程意义与应用',
           '附录']:
    doc.add_paragraph(t, style='List Number')
doc.add_page_break()

# 1
doc.add_heading('1. 概述与顶层直觉', level=1)
doc.add_heading('1.1 要解决什么问题', level=2)
doc.add_paragraph('在燃料电池汽车能量管理策略（EMS）中，有两种典型方法：')
tbl = doc.add_table(rows=3, cols=3); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = '方法'
tbl.rows[0].cells[1].text = '特点'
tbl.rows[0].cells[2].text = '是否需要未来信息'
tbl.rows[1].cells[0].text = 'DP（动态规划）'
tbl.rows[1].cells[1].text = '全局最优，但必须知道完整工况后才能计算（非因果）'
tbl.rows[1].cells[2].text = '✅ 是'
tbl.rows[2].cells[0].text = 'ECMS（等效消耗最小化）'
tbl.rows[2].cells[1].text = '实时决策，只靠当前状态和参数 s 决定功率分配（因果）'
tbl.rows[2].cells[2].text = '❌ 否'
doc.add_paragraph('ECMS 的核心调优参数是等效因子 s，它决定了“1 kWh 电池电能值多少克氢”。s 设大了→ ECMS 倾向于少用电池（太保守），s 设小了→ ECMS 过度消耗电池（无法维持 SOC）。找到合适的 s 是 ECMS 工程应用的关键。')
doc.add_heading('1.2 核心洞察', level=2)
doc.add_paragraph('DP 虽然不能实时运行，但它隐含地给出了一个关键信息——电池电能的“边际价值”：“当前 SOC 再高 1 单位，全程总氢耗还能省多少？”这个边际价值在最优控制理论中叫做 costate（协态变量）λ。而 λ 与 ECMS 的等效因子 s 之间有一个确定的数学关系，因此我们可以：')
p = doc.add_paragraph(); bold(p, '跑一次 DP → 提取 λ → 换算成 s₀ → 直接用 ECMS')
doc.add_page_break()

# 2
doc.add_heading('2. 核心数学原理：PMP ↔ ECMS 的桥梁', level=1)
doc.add_heading('2.1 Pontryagin 最小原理（PMP）的 Hamiltonian', level=2)
doc.add_paragraph('PMP 是解决最优控制问题的核心工具。对于 EMS 问题，Hamiltonian 定义为：')
code('H = m_dot_H2  +  λ · SOC_dot')
p = doc.add_paragraph(); p.add_run('其中：').bold = True
doc.add_paragraph('m_dot_H2 — 燃料电池的瞬时氢耗 [g/s]', style='List Bullet')
doc.add_paragraph('λ — costate（协态变量），表示“当前 SOC 增加 1 单位，未来最优总代价能减少多少”', style='List Bullet')
doc.add_paragraph('SOC_dot — 电池 SOC 的变化率 [1/s]', style='List Bullet')
doc.add_paragraph('PMP 说：最优控制（功率分配）在每个时刻使 H 取最小值。')
doc.add_heading('2.2 ECMS 的瞬时优化目标', level=2)
code('H_eq = m_dot_H2  +  s · |P_bat| / 3600')
doc.add_paragraph('ECMS 用一个常数等效因子 s，把电池功率 P_bat 折算成“虚拟氢耗”，然后每个时刻都最小化实际氢耗 + 虚拟氢耗之和。')
doc.add_heading('2.3 两边的桥梁', level=2)
doc.add_paragraph('PMP 和 ECMS 本质上在做同一件事：把“用电池”的远期代价折算到当前决策中。区别在于 PMP 用 λ（时变、隐式、最优），ECMS 用 s（常数、显式、需调参）。关键洞察：把 PMP 的 λ·SOC_dot 项展开，就能发现它与 s·|P_bat|/3600 项在数学形式上是对偶的。')
doc.add_page_break()

# 3
doc.add_heading('3. λ 到 s 的换算推导', level=1)
doc.add_heading('3.1 SOC 变化率的展开', level=2)
doc.add_paragraph('电池的 SOC 动态由电流决定：')
code('SOC_dot  =  -I / (Q_bat · 3600)')
doc.add_paragraph('其中 Q_bat = 50 Ah 是电池容量，3600 是 Ah → As 的换算。')
doc.add_paragraph('电池功率与电流的关系：')
code('P_bat = V_oc · I / 1000        [kW]')
doc.add_paragraph('代入 SOC_dot：')
code('SOC_dot  =  -P_bat · 1000 / (V_oc · Q_bat · 3600)')
doc.add_heading('3.2 代入 PMP 的 Hamiltonian', level=2)
code('λ · SOC_dot  =  λ · [ -P_bat · 1000 / (V_oc · Q_bat · 3600) ]\n              =  [ -λ · 1000 / (V_oc · Q_bat) ]  ·  P_bat / 3600')
doc.add_heading('3.3 与 ECMS 对比', level=2)
doc.add_paragraph('ECMS 的等效氢耗项是 s · |P_bat| / 3600，忽略绝对值（方向由符号决定），对比可得：')
code('s  ≈  -λ · 1000 / (V_oc · Q_bat)')
doc.add_paragraph('这就是代码第 80 行的公式来源。')
doc.add_heading('3.4 数值换算', level=2)
doc.add_paragraph('代码中使用近似值：')
doc.add_paragraph('V_oc_approx = 352 V（SOC ≈ 0.6 时的开路电压）', style='List Bullet')
doc.add_paragraph('Q_bat = 50 Ah', style='List Bullet')
doc.add_paragraph('则换算关系为：', style='List Bullet')
code('s = -λ · 1000 / (352 · 50) = -λ / 17.6')
doc.add_paragraph('注意负号：当 λ 为负值时（这是典型情况——电池电量越高未来越省），s 为正，物理意义清晰：用电池要消耗正值的“虚拟氢耗”。')
doc.add_page_break()

# 4
doc.add_heading('4. 函数级逐行分析', level=1)

# 4.1
doc.add_heading('4.1 compute_costate — 沿最优轨迹计算 costate λ', level=2)
code('''def compute_costate(J, SOC_GRID, opt_soc_traj):
    N = opt_soc_traj.shape[0] - 1
    dSOC = SOC_GRID[1] - SOC_GRID[0]
    lambdas = np.zeros(N)
    for k in range(N):
        i = np.argmin(np.abs(SOC_GRID - opt_soc_traj[k]))
        i_l = max(0, i - 1)
        i_r = min(N_SOC - 1, i + 1)
        if i_l == i_r:
            lambdas[k] = 0
        else:
            lambdas[k] = (J[k, i_r] - J[k, i_l]) / (SOC_GRID[i_r] - SOC_GRID[i_l])
    return lambdas''')
p = doc.add_paragraph(); bold(p, '原理：')
doc.add_paragraph('J[k][i] 是 DP 的代价矩阵——“从时刻 k、SOC 网格点 i 出发，到终点最少要花多少总代价”。沿最优 SOC 轨迹对 J 取 SOC 方向的偏导数，就得到 costate λ。')
code('λ_k  ≈  (J[k, i+1] - J[k, i-1]) / (SOC_grid[i+1] - SOC_grid[i-1])')
p = doc.add_paragraph(); bold(p, '为什么用中心差分？')
doc.add_paragraph('中心差分的截断误差是 O(h²)，而前向/后向差分只有 O(h)。SOC 网格间距约 0.0047（(0.9-0.2)/150），中心差分能显著提高精度。边界处（i=0 或 i=N_SOC-1）退化为单侧差分。')

# 4.2
doc.add_heading('4.2 costate_to_s — λ → 等效因子 s 换算', level=2)
code('''def costate_to_s(lambda_k):
    V_oc_approx = 352.0
    s_k = -lambda_k * 1000 / (V_oc_approx * Q_BAT)
    return s_k''')
doc.add_paragraph('使用常数 V_oc 近似是简化处理。实际开路电压随 SOC 变化（代码中 OCV_LU 数组从 320V → 380V）。但系数 1000/(352×50)=1/17.6 作为整体比例因子，影响的是 s 的绝对值而非相对变化趋势，用常数近似足够了。')

# 4.3
doc.add_heading('4.3 backward_dp — 后向动态规划', level=2)
code('''def backward_dp(P_load, SOC_0=0.6):
    J = zeros((N+1, N_SOC))     # 代价矩阵
    pi = zeros((N, N_SOC))      # 策略矩阵
    J[N, :] = BETA * (SOC_GRID - SOC_0)**2  # 终端惩罚
    for k in range(N-1, -1, -1):
        for i in range(N_SOC):
            SOC_next_all = state_transition(soc, PFC_GRID, P_load_k, DT)
            # feasible = SOC_next 在 [SOC_MIN, SOC_MAX] 内
            total = g + J_future + α·(SOC_next - SOC_ref)²
            J[k,i] = min(total)
            pi[k,i] = argmin(total)
    return J, pi''')
p = doc.add_paragraph(); bold(p, '核心 DP 递推方程：')
code('''J[k][i] = min_{P_fc} [
    H2_flow(P_fc) · DT               # 当前氢耗
    + J[k+1][interp(SOC_next)]          # 未来最优代价
    + α · (SOC_next - SOC_ref)²            # SOC 维持惩罚
]''')
p = doc.add_paragraph(); bold(p, '关键设计要点：')
doc.add_paragraph('向量化内层循环：对每个 SOC 状态 i，一次计算 60 种 P_fc 候选的 SOC_next', style='List Bullet')
doc.add_paragraph('SOC 惩罚 α·(SOC_next - SOC_ref)² 放在“控制结果”上，而不是对当前 SOC 状态惩罚', style='List Bullet')
doc.add_paragraph('终端惩罚 BETA=10000强制终点 SOC 回到 0.6', style='List Bullet')
doc.add_paragraph('J[k][i] = inf 表示该状态不可达，DP 会自动避开', style='List Bullet')

# 4.4
doc.add_heading('4.4 forward_rollout — 前向 Rollout', level=2)
code('''def forward_rollout(P_load, pi, SOC_0=0.6):
    SOC[0] = SOC_0
    for k in range(N):
        P_fc[k] = interp(SOC[k], SOC_GRID, pi[k, :])
        SOC[k+1] = state_transition(SOC[k], P_fc[k], P_load[k])
        M_H2[k] = fc_hydrogen_flow(P_fc[k]) * DT
    return {SOC, P_fc, P_bat, m_H2_cumul}''')
doc.add_paragraph('前向 Rollout 就是把策略表 π 用在实际仿真中：从真实初始 SOC₀ 出发，每步查表输出 P_fc，状态更新后继续，直到仿真结束。由于 SOC 是连续变量而策略表是离散网格点，用 numpy.interp 做线性插值——比最近邻法更平滑。')

# 4.5
doc.add_heading('4.5 main — 主流程整合', level=2)
doc.add_paragraph('主流程分四个步骤：')
p = doc.add_paragraph(); bold(p, 'Step 1 — 后向 DP')
doc.add_paragraph('调用 backward_dp(P_load) 计算代价矩阵 J 和策略矩阵 π。J 的形状是 (N+1, N_SOC)，记录了从“每个时刻 × 每个 SOC 状态”出发到终点的最优累计代价。')
p = doc.add_paragraph(); bold(p, 'Step 2 — 前向 Rollout')
doc.add_paragraph('调用 forward_rollout(P_load, π) 得到最优 SOC 轨迹 opt_soc，记录 DP 在每个时刻的最优 SOC。')
p = doc.add_paragraph(); bold(p, 'Step 3 — 计算 Costate λ')
doc.add_paragraph('调用 compute_costate(J, SOC_GRID, opt_soc) 沿最优轨迹用有限差分算出每个时刻的 λ。λ > 0表示“当前 SOC 越高，未来代价越大”（电池应该放电），λ < 0表示“当前 SOC 越高，未来代价越小”（电池应该充电或省着用）。')
p = doc.add_paragraph(); bold(p, 'Step 4 — 换算为等效因子 s')
doc.add_paragraph('调用 costate_to_s(λ) 得到时变的等效因子 s(t) 信号。取中位数作为推荐的恒定 s₀（中位数对极端值鲁棒）。')

doc.add_page_break()

# 5
doc.add_heading('5. 输出图表含义', level=1)
tbl2 = doc.add_table(rows=5, cols=3); tbl2.style = 'Light Grid Accent 1'
tbl2.rows[0].cells[0].text = '子图'; tbl2.rows[0].cells[1].text = '内容'; tbl2.rows[0].cells[2].text = '分析要点'
data = [
    ('1', 'SOC 轨迹', 'DP 最优 SOC 是否维持在 [0.3, 0.8] 区间？终点是否回到 SOC_ref=0.6？'),
    ('2', 'Costate λ', 'λ 随时间波动——加速段 λ 负值大（电池值钱），刹车段 λ 接近 0。分布范围说明电池的边际价值变化幅度。'),
    ('3', '等效因子 s', 's(t) 在 50-250 之间波动。橙色虚线是中位数 s₀。绿色阴影带是 [25%, 75%] 分位数区间。'),
    ('4', 's 直方图', '分布形状决定 s₀ 的置信度——单峰集中分布说明 DP 反推的 s 值一致性好；多峰或宽分布说明工况复杂、单一固定 s 不够用。'),
]
for i, (c1, c2, c3) in enumerate(data):
    tbl2.rows[i+1].cells[0].text = c1
    tbl2.rows[i+1].cells[1].text = c2
    tbl2.rows[i+1].cells[2].text = c3

doc.add_heading('5.1 统计输出指标解读', level=2)
doc.add_paragraph('Costate λ 均值/标准差/范围 — 衡量电池边际价值的典型量级和波动性', style='List Bullet')
doc.add_paragraph('等效因子 s 均值/中位数/标准差/范围 — 推荐 s₀ = median(s_valid)，因为当 λ 接近 0 时 s ≈ 0 是极端值，中位数对此类异常值更鲁棒', style='List Bullet')
doc.add_paragraph('DP 总氢耗 — 基准线，后续 ECMS 的氢耗应该与之对比来评估策略优劣', style='List Bullet')
doc.add_paragraph('推荐 s₀ vs 经验校准 s₀（abs 公式）= 130 g/kWh — 两者接近验证了方法的有效性', style='List Bullet')

doc.add_page_break()

# 6
doc.add_heading('6. 工程意义与应用', level=1)
doc.add_heading('6.1 为什么这很重要', level=2)
doc.add_paragraph('传统上 ECMS 的等效因子 s 通过试错法（参数扫描）标定——试 10 个不同的 s 值、跑 10 次仿真、看哪个氢耗最低且 SOC 维持得最好。这工作量大且依赖经验。DP 反推法提供了理论依据：不需要反复试跑 ECMS，一次 DP 就能告诉你 s 应该设成多少。')
doc.add_heading('6.2 时变特性 → 自适应 ECMS（A-ECMS）', level=2)
doc.add_paragraph('从输出图可以看到，最优等效因子 s 是随时间变化的——这意味着固定一个 s₀ 不可能在所有工况下都最优。这正是自适应 ECMS（A-ECMS）的理论基础：')
code('s(k) = s₀ · (1 + Kp · (SOC_ref - SOC(k)))')
doc.add_paragraph('当 SOC 偏低时 s 增大（更珍惜电池、倾向充电），当 SOC 偏高时 s 减小（更倾向放电）。DP 反推的 s 分布宽度（IQR）可以作为 A-ECMS 的 Kp 增益设计参考——分布越宽，说明 s 需要调整的幅度越大，Kp 可以设得更大。')
doc.add_heading('6.3 工程使用建议', level=2)
doc.add_paragraph('对每个目标工况（WLTC/NEDC/CLTC）分别跑一次 DP 反推，获得对应的 s₀', style='List Bullet')
doc.add_paragraph('用 s₀ 作为 ECMS 的初始值，在实车或 HiL 上做 ±20% 的微调扫描', style='List Bullet')
doc.add_paragraph('s 的分布宽度越大的工况，越值得使用 A-ECMS 而非固定 s', style='List Bullet')
doc.add_paragraph('对比 DP 氢耗与 ECMS 氢耗的差距，可以评估实时策略的“最优性差距”', style='List Bullet')

doc.add_page_break()

# 7
doc.add_heading('7. 附录：完整代码', level=1)
with open(os.path.join(ROOT, 'scripts', 'calibrate_s_from_dp.py'), 'r', encoding='utf-8') as f:
    for line in f.read().split('\n'):
        code(line)

# 保存
output = os.path.join(DOCS, 'calibrate_s_from_dp_analysis.docx')
doc.save(output)
print(f'[OK] {output}')
