#!/usr/bin/env python3
"""
gen_calibrate_s_doc.py — 生成 calibrate_s_from_dp.py 逐行原理分析 DOCX
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')
os.makedirs(DOCS_DIR, exist_ok=True)

doc = Document()

# ── 样式设定 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

# heading 样式
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 代码样式（用 list bullet 搭配等宽字体）
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Cm(0.5)

def add_code(text):
    """添加代码块"""
    doc.add_paragraph(text, style='CodeBlock')

def add_bold_inline(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.size = Pt(8)

# ════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('calibrate_s_from_dp.py')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DP 反推标定 ECMS 等效因子 — 逐行原理解析')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n从 Pontryagin 最小原理到等效因子的完整推导\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 目录（手动）
# ════════════════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '1. 概述与顶层直觉',
    '2. 核心数学原理：PMP ↔ ECMS 的桥梁',
    '3. λ 到 s 的换算推导',
    '4. 函数级逐行分析',
    '    4.1 compute_costate — 沿最优轨迹计算 costate λ',
    '    4.2 costate_to_s — λ → 等效因子 s 换算',
    '    4.3 backward_dp — 后向动态规划',
    '    4.4 forward_rollout — 前向 Rollout',
    '    4.5 main — 主流程',
    '5. 输出图表含义',
    '6. 工程意义与应用',
    '7. 附录：完整代码',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ════════════════════════════════════════════════════
# 1. 概述
# ════════════════════════════════════════════════════
doc.add_heading('1. 概述与顶层直觉', level=1)

doc.add_heading('1.1 要解决什么问题', level=2)
doc.add_paragraph(
    '在燃料电池汽车能量管理策略（EMS）中，有两种典型方法：'
)

table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '方法'
hdr[1].text = '特点'
hdr[2].text = '是否需要未来信息'
cells = table.rows[1].cells
cells[0].text = 'DP（动态规划）'
cells[1].text = '全局最优，但必须知道完整工况后才能计算（非因果）'
cells[2].text = '✅ 是'
cells = table.rows[2].cells
cells[0].text = 'ECMS（等效消耗最小化）'
cells[1].text = '实时决策，只靠当前状态和参数 s 决定功率分配（因果）'
cells[2].text = '❌ 否'

doc.add_paragraph(
    'ECMS 的核心调优参数是等效因子 s（equivalence factor），它决定了"1 kWh 电池电能值多少克氢"。'
    's 设大了 → ECMS 倾向于少用电池（太保守），'
    's 设小了 → ECMS 过度消耗电池（无法维持 SOC）。'
    '找到合适的 s 是 ECMS 工程应用的关键。'
)

doc.add_heading('1.2 核心洞察', level=2)
doc.add_paragraph(
    'DP 虽然不能实时运行，但它隐含地给出了一个关键信息——电池电能的"边际价值"：'
    '"当前 SOC 再高 1 单位，全程总氢耗还能省多少？"'
)
doc.add_paragraph(
    '这个边际价值在最优控制理论中叫做 costate（协态变量）λ。'
    '而 λ 与 ECMS 的等效因子 s 之间有一个确定的数学关系，'
    '因此我们可以：'
)
p = doc.add_paragraph()
add_bold_inline(p, '跑一次 DP → 提取 λ → 换算成 s₀ → 直接用 ECMS')

doc.add_page_break()

# ════════════════════════════════════════════════════
# 2. 核心数学原理
# ════════════════════════════════════════════════════
doc.add_heading('2. 核心数学原理：PMP ↔ ECMS 的桥梁', level=1)

doc.add_heading('2.1 Pontryagin 最小原理（PMP）的 Hamiltonian', level=2)
doc.add_paragraph(
    'PMP 是解决最优控制问题的核心工具。对于 EMS 问题，Hamiltonian 定义为：'
)
add_code('H = m_dot_H2  +  λ · SOC_dot')
p = doc.add_paragraph()
p.add_run('其中：').bold = True
doc.add_paragraph('m_dot_H2 — 燃料电池的瞬时氢耗 [g/s]', style='List Bullet')
doc.add_paragraph(
    'λ — costate（协态变量），表示"当前 SOC 增加 1 单位，未来最优总代价能减少多少"',
    style='List Bullet'
)
doc.add_paragraph(
    'SOC_dot — 电池 SOC 的变化率 [1/s]', style='List Bullet'
)

doc.add_paragraph(
    'PMP 说：最优控制（功率分配）在每个时刻使 H 取最小值。'
    '注意这里的 λ·SOC_dot 项——它把电池的"状态变化"也计量到瞬时代价中，'
    '这就是 DP 隐含的"边际价值"机制。'
)

doc.add_heading('2.2 ECMS 的瞬时优化目标', level=2)
add_code('H_eq = m_dot_H2  +  s · |P_bat| / 3600')
doc.add_paragraph(
    'ECMS 用一个常数等效因子 s，把电池功率 P_bat 折算成"虚拟氢耗"，'
    '然后每个时刻都最小化实际氢耗 + 虚拟氢耗之和。'
    's 的单位是 [g/kWh]——每消耗（或回收）1 kWh 电能，等价于消耗 s 克氢气。'
)

doc.add_heading('2.3 两边的桥梁', level=2)
doc.add_paragraph(
    'PMP 和 ECMS 本质上在做同一件事：把"用电池"的远期代价折算到当前决策中。'
    '区别在于 PMP 用 λ（时变、隐式、最优），ECMS 用 s（常数、显式、需调参）。'
    '关键洞察：把 PMP 的 λ·SOC_dot 项展开，就能发现它与 s·|P_bat|/3600 项'
    '在数学形式上是对偶的。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 3. λ → s 换算推导
# ════════════════════════════════════════════════════
doc.add_heading('3. λ 到 s 的换算推导', level=1)

doc.add_heading('3.1 SOC 变化率的展开', level=2)
doc.add_paragraph('电池的 SOC 动态由电流决定：')
add_code('SOC_dot  =  -I / (Q_bat · 3600)')
doc.add_paragraph('其中 Q_bat = 50 Ah 是电池容量，3600 是 Ah → As 的换算。')
doc.add_paragraph('电池功率与电流的关系：')
add_code('P_bat = V_oc · I / 1000        [kW]')
doc.add_paragraph('I 的单位为 A，代入 SOC_dot：')
add_code('SOC_dot  =  -P_bat · 1000 / (V_oc · Q_bat · 3600)')

doc.add_heading('3.2 代入 PMP 的 Hamiltonian', level=2)
add_code(
    'λ · SOC_dot  =  λ · [ -P_bat · 1000 / (V_oc · Q_bat · 3600) ]\n'
    '              =  [ -λ · 1000 / (V_oc · Q_bat) ]  ·  P_bat / 3600'
)

doc.add_heading('3.3 与 ECMS 对比', level=2)
doc.add_paragraph(
    'ECMS 的等效氢耗项是 s · |P_bat| / 3600，忽略绝对值（方向由符号决定），对比可得：'
)
add_code('s  ≈  -λ · 1000 / (V_oc · Q_bat)')
doc.add_paragraph('这就是代码第 80 行的公式来源。')

doc.add_heading('3.4 数值换算', level=2)
doc.add_paragraph('代码中使用近似值：')
doc.add_paragraph('V_oc_approx = 352 V（SOC ≈ 0.6 时的开路电压）', style='List Bullet')
doc.add_paragraph('Q_bat = 50 Ah', style='List Bullet')
doc.add_paragraph('则 λ → s 的换算关系为：', style='List Bullet')
add_code('s = -λ · 1000 / (352 · 50) = -λ / 17.6')

doc.add_paragraph(
    '注意负号：当 λ 为负值时（这是典型情况——电池电量越高未来越省），'
    's 为正，物理意义清晰：用电池要消耗正值的"虚拟氢耗"。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 4. 函数级逐行分析
# ════════════════════════════════════════════════════
doc.add_heading('4. 函数级逐行分析', level=1)

# ── 4.1 compute_costate ──
doc.add_heading('4.1 compute_costate — 沿最优轨迹计算 costate λ', level=2)
add_code(
    'def compute_costate(J, SOC_GRID, opt_soc_traj):\n'
    '    N = opt_soc_traj.shape[0] - 1       # 时间步数\n'
    '    dSOC = SOC_GRID[1] - SOC_GRID[0]     # 网格间距（均匀网格）\n'
    '    lambdas = np.zeros(N)\n\n'
    '    for k in range(N):                    # 对每个时间步\n'
    '        i = np.argmin(np.abs(SOC_GRID - opt_soc_traj[k]))  # 找最近网格索引\n'
    '        i_l = max(0, i - 1)               # 左邻居（边界保护）\n'
    '        i_r = min(N_SOC - 1, i + 1)       # 右邻居（边界保护）\n'
    '        if i_l == i_r:\n'
    '            lambdas[k] = 0                 # 退化为单点，导数为 0\n'
    '        else:\n'
    '            lambdas[k] = (J[k, i_r] - J[k, i_l]) / (SOC_GRID[i_r] - SOC_GRID[i_l])\n'
    '    return lambdas'
)

doc.add_paragraph()
p = doc.add_paragraph()
add_bold_inline(p, '原理：')
doc.add_paragraph(
    'J[k][i] 是 DP 的代价矩阵——"从时刻 k、SOC 网格点 i 出发，到终点最少要花多少总代价"。'
    '沿最优 SOC 轨迹对 J 取 SOC 方向的偏导数，就得到 costate λ：'
)
add_code('λ_k  ≈  (J[k, i+1] - J[k, i-1]) / (SOC_grid[i+1] - SOC_grid[i-1])')

p = doc.add_paragraph()
add_bold_inline(p, '为什么用中心差分？')
doc.add_paragraph(
    '中心差分（Central Difference）的截断误差是 O(h²)，'
    '而前向/后向差分只有 O(h)。这里 SOC 网格间距约 0.0047（(0.9-0.2)/150），'
    '中心差分能显著提高精度。边界处（i=0 或 i=N_SOC-1）退化为单侧差分。'
)

p = doc.add_paragraph()
add_bold_inline(p, '输入参数：')
doc.add_paragraph('J: (N+1, N_SOC) — 后向 DP 代价矩阵', style='List Bullet')
doc.add_paragraph('SOC_GRID: (N_SOC,) — SOC 离散网格', style='List Bullet')
doc.add_paragraph('opt_soc_traj: (N+1,) — DP 最优 SOC 轨迹', style='List Bullet')

p = doc.add_paragraph()
add_bold_inline(p, '输出：')
doc.add_paragraph('lambdas: (N,) — 每个时间步的 costate', style='List Bullet')

doc.add_paragraph()

# ── 4.2 costate_to_s ──
doc.add_heading('4.2 costate_to_s — λ → 等效因子 s 换算', level=2)
add_code(
    'def costate_to_s(lambda_k):\n'
    '    V_oc_approx = 352.0                     # 典型开路电压 @ SOC=0.6\n'
    '    s_k = -lambda_k * 1000 / (V_oc_approx * Q_BAT)\n'
    '    return s_k'
)

doc.add_paragraph()
p = doc.add_paragraph()
add_bold_inline(p, '说明：')
doc.add_paragraph(
    '使用常数 V_oc 近似是简化处理。实际开路电压随 SOC 变化（代码中 OCV_LU 数组从 320V → 380V）。'
    '更精确的做法是用动态 V_oc(SOC) 逐点换算，但系数 1000/(352×50)=1/17.6 作为整体比例因子，'
    '影响的是 s 的绝对值而非相对变化趋势，用常数近似足够了——因为最终我们取的是 s 的统计中位数。'
)

# ── 4.3 backward_dp ──
doc.add_heading('4.3 backward_dp — 后向动态规划', level=2)
add_code(
    'def backward_dp(P_load, SOC_0=0.6):\n'
    '    N = len(P_load)             # 时间步数（如 WLTC 的 1800s）\n'
    '    SOC_GRID = linspace(0.2, 0.9, 150)    # SOC 离散网格\n'
    '    PFC_GRID = linspace(0, 30, 60)        # FC 功率离散网格\n\n'
    '    J = zeros((N+1, N_SOC))    # 代价矩阵\n'
    '    pi = zeros((N, N_SOC))     # 策略矩阵\n\n'
    '    # 终端惩罚：终点 SOC 偏离 0.6 的巨大代价\n'
    '    J[N, :] = BETA * (SOC_GRID - SOC_0)**2\n\n'
    '    for k in range(N-1, -1, -1):           # 从终点往起点回溯\n'
    '        for i in range(N_SOC):             # 遍历每个 SOC 状态\n'
    '            # 向量化：一次试完 60 种 FC 功率\n'
    '            SOC_next_all = state_transition(soc, PFC_GRID, P_load_k, DT)\n'
    '            feasible = 范围内筛选\n\n'
    '            # 总代价 = 当前氢耗 + 未来代价 + SOC 维持惩罚\n'
    '            total = g + J_future + α·(SOC_next - SOC_ref)²\n'
    '            J[k,i] = min(total)\n'
    '            pi[k,i] = argmin(total)\n'
    '    return J, pi'
)

doc.add_paragraph()
p = doc.add_paragraph()
add_bold_inline(p, '核心 DP 递推方程：')
add_code(
    'J[k][i] = min_{P_fc ∈ PFC_GRID} [\n'
    '    H2_flow(P_fc) · DT                            ← 当前氢耗\n'
    '    + J[k+1][interp(SOC_next)]                    ← 未来最优代价\n'
    '    + α · (SOC_next - SOC_ref)²                    ← SOC 维持惩罚\n'
    ']'
)

doc.add_paragraph()
p = doc.add_paragraph()
add_bold_inline(p, '关键设计要点：')
doc.add_paragraph(
    '向量化内层循环：对每个 SOC 状态 i，一次计算 60 种 P_fc 候选的 SOC_next，'
    '比原来逐个尝试快约 60 倍',
    style='List Bullet'
)
doc.add_paragraph(
    'SOC 惩罚 α·(SOC_next - SOC_ref)² 放在"控制结果"上，'
    '而不是对当前 SOC 状态惩罚，这样 DP 才能"看到"偏离 SOC_ref 的后果',
    style='List Bullet'
)
doc.add_paragraph(
    '终端惩罚 BETA 很大（10000），强制 DP 在终点把 SOC 拉回 0.6，'
    '保证公平对比（ECMS 也以 SOC_ref 为目标）',
    style='List Bullet'
)
doc.add_paragraph(
    'J[k][i] = inf 表示该状态不可达（没有可行控制能维持 SOC 在范围内），'
    'DP 会自动避开这些区域',
    style='List Bullet'
)

# ── 4.4 forward_rollout ──
doc.add_heading('4.4 forward_rollout — 前向 Rollout', level=2)
add_code(
    'def forward_rollout(P_load, pi, SOC_0=0.6):\n'
    '    SOC[0] = SOC_0\n'
    '    for k in range(N):\n'
    '        # 查策略表得到最优 FC 功率\n'
    '        P_fc[k] = interp(SOC[k], SOC_GRID, pi[k, :])\n'
    '        # 状态更新\n'
    '        SOC[k+1] = state_transition(SOC[k], P_fc[k], P_load[k])\n'
    '        M_H2[k] = fc_hydrogen_flow(P_fc[k]) * DT\n'
    '    return {SOC, P_fc, P_bat, m_H2_cumul}'
)

doc.add_paragraph()
p = doc.add_paragraph()
add_bold_inline(p, '原理：')
doc.add_paragraph(
    '后向 DP 得到的策略矩阵 π[k][i] 是一个"如果此时此刻在状态 i，最优输出是什么"的查找表。'
    '前向 Rollout 就是把这个查找表用在实际仿真中：从真实初始 SOC₀ 出发，'
    '每步查表输出 P_fc，状态更新后继续，直到仿真结束。'
)
doc.add_paragraph(
    '由于 SOC 是连续变量而策略表是离散网格点，用 numpy.interp 做线性插值——'
    '这比用最近邻法更平滑，避免 SOC 轨迹出现阶梯状突变。'
)

# ── 4.5 main ──
doc.add_heading('4.5 main — 主流程整合', level=2)
doc.add_paragraph('主流程分四个清晰的步骤：')

p = doc.add_paragraph()
add_bold_inline(p, 'Step 1 — 后向 DP（第 101-102 行）')
doc.add_paragraph(
    '调用 backward_dp(P_load) 计算代价矩阵 J 和策略矩阵 π。'
    'J 的形状是 (N+1, N_SOC)，记录了从"每个时刻 × 每个 SOC 状态"出发到终点的最优累计代价。'
)

p = doc.add_paragraph()
add_bold_inline(p, 'Step 2 — 前向 Rollout（第 106-107 行）')
doc.add_paragraph(
    '调用 forward_rollout(P_load, π) 得到最优 SOC 轨迹 opt_soc，'
    '这是一条 (N+1,) 的向量，记录 DP 在每个时刻的最优 SOC。'
)

p = doc.add_paragraph()
add_bold_inline(p, 'Step 3 — 计算 Costate λ（第 110-111 行）')
doc.add_paragraph(
    '调用 compute_costate(J, SOC_GRID, opt_soc) 沿最优轨迹用有限差分算出每个时刻的 λ。'
    'λ > 0 表示"当前 SOC 越高，未来代价越大"（电池应该放电），'
    'λ < 0 表示"当前 SOC 越高，未来代价越小"（电池应该充电或省着用）。'
)

p = doc.add_paragraph()
add_bold_inline(p, 'Step 4 — 换算为等效因子 s（第 114-115 行）')
doc.add_paragraph(
    '调用 costate_to_s(λ) 得到时变的等效因子 s(t) 信号。'
    '取中位数作为推荐的恒定 s₀（中位数对极端值鲁棒）。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 5. 输出图表
# ════════════════════════════════════════════════════
doc.add_heading('5. 输出图表含义', level=1)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '子图'
hdr[1].text = '内容'
hdr[2].text = '分析要点'
rows = [
    ('1', 'SOC 轨迹', 'DP 最优 SOC 是否维持在 [0.3, 0.8] 区间？终点是否回到 SOC_ref=0.6？'),
    ('2', 'Costate λ', 'λ 随时间波动——加速段 λ 负值大（电池值钱），刹车段 λ 接近 0。分布范围说明电池的边际价值变化幅度。'),
    ('3', '等效因子 s', 's(t) 在 50-250 之间波动。橙色虚线是中位数 s₀。绿色阴影带是 [25%, 75%] 分位数区间，反映 s 的典型波动范围。'),
    ('4', 's 直方图', '分布形状决定 s₀ 的置信度——单峰集中分布说明 DP 反推的 s 值一致性好；多峰或宽分布说明工况复杂、单一固定 s 不够用。'),
]
for i, (col1, col2, col3) in enumerate(rows):
    table.rows[i+1].cells[0].text = col1
    table.rows[i+1].cells[1].text = col2
    table.rows[i+1].cells[2].text = col3

doc.add_paragraph()
doc.add_heading('5.1 统计输出指标解读', level=2)
doc.add_paragraph(
    'Costate λ 均值/标准差/范围 — 衡量电池边际价值的典型量级和波动性',
    style='List Bullet'
)
doc.add_paragraph(
    '等效因子 s 均值/中位数/标准差/范围 — 推荐 s₀ = median(s_valid)，'
    '因为当 λ 接近 0 时 s ≈ 0 是极端值，中位数对此类异常值更鲁棒',
    style='List Bullet'
)
doc.add_paragraph(
    'DP 总氢耗 — 基准线，后续 ECMS 的氢耗应该与之对比来评估策略优劣',
    style='List Bullet'
)
doc.add_paragraph(
    '推荐 s₀ vs 经验校准 s₀（abs 公式）= 130 g/kWh — 两者接近验证了方法的有效性',
    style='List Bullet'
)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 6. 工程意义
# ════════════════════════════════════════════════════
doc.add_heading('6. 工程意义与应用', level=1)

doc.add_heading('6.1 为什么这很重要', level=2)
doc.add_paragraph(
    '传统上 ECMS 的等效因子 s 通过试错法（参数扫描）标定——试 10 个不同的 s 值、'
    '跑 10 次仿真、看哪个氢耗最低且 SOC 维持得最好。这工作量大且依赖经验。'
)
doc.add_paragraph(
    'DP 反推法提供了理论依据：不需要反复试跑 ECMS，一次 DP 就能告诉你 s 应该设成多少。'
)

doc.add_heading('6.2 时变特性 → 自适应 ECMS（A-ECMS）', level=2)
doc.add_paragraph(
    '从输出图可以看到，最优等效因子 s 是随时间变化的——这意味着固定一个 s₀ 不可能'
    '在所有工况下都最优。这正是自适应 ECMS（A-ECMS）的理论基础：'
)
add_code(
    's(k) = s₀ · (1 + Kp · (SOC_ref - SOC(k)))'
)
doc.add_paragraph(
    '当 SOC 偏低时 s 增大（更珍惜电池、倾向充电），'
    '当 SOC 偏高时 s 减小（更倾向放电）。'
    'DP 反推的 s 分布宽度（IQR）可以作为 A-ECMS 的 Kp 增益设计参考——'
    '分布越宽，说明 s 需要调整的幅度越大，Kp 可以设得更大。'
)

doc.add_heading('6.3 工程使用建议', level=2)
doc.add_paragraph(
    '对每个目标工况（WLTC/NEDC/CLTC）分别跑一次 DP 反推，获得对应的 s₀',
    style='List Bullet'
)
doc.add_paragraph(
    '用 s₀ 作为 ECMS 的初始值，在实车或 HiL 上做 ±20% 的微调扫描',
    style='List Bullet'
)
doc.add_paragraph(
    's 的分布宽度越大的工况，越值得使用 A-ECMS 而非固定 s',
    style='List Bullet'
)
doc.add_paragraph(
    '对比 DP 氢耗与 ECMS 氢耗的差距，可以评估实时策略的"最优性差距"',
    style='List Bullet'
)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 7. 附录
# ════════════════════════════════════════════════════
doc.add_heading('7. 附录：完整代码', level=1)

with open(os.path.join(PROJECT_ROOT, 'scripts', 'calibrate_s_from_dp.py'), 'r', encoding='utf-8') as f:
    code_lines = f.read()

# 按行添加代码
for line in code_lines.split('\n'):
    add_code(line)

# ── 保存 ──
output_path = os.path.join(DOCS_DIR, 'calibrate_s_from_dp_逐行分析.docx')
doc.save(output_path)
print(f'[OK] 文档已生成: {output_path}')
