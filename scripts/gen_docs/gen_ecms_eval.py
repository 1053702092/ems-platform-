# -*- coding: utf-8 -*-
"""生成 ECMS Week 5 结果评判/评估报告 .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import pandas as pd
import numpy as np

DOCS_DIR = r'F:\CLAUDE\research\ems-platform\docs'
RESULTS_DIR = r'F:\CLAUDE\research\ems-platform\results'

# ── helpers ──
def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def p(doc, text, bold=False, size=10.5, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def tbl(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table

def verdict(doc, text, level='pass'):
    """ verdict badge """
    colors = {
        'pass': RGBColor(0, 128, 0),
        'warn': RGBColor(200, 120, 0),
        'fail': RGBColor(200, 0, 0),
    }
    labels = {'pass': '✅ 通过', 'warn': '⚠️ 需改进', 'fail': '❌ 未达标'}
    para = doc.add_paragraph()
    run = para.add_run(f'{labels.get(level, "")} {text}')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para


def main():
    os.makedirs(DOCS_DIR, exist_ok=True)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ================================================================
    # 0. Load data
    # ================================================================
    df_scan = pd.read_csv(os.path.join(RESULTS_DIR, 'Day9_ecms_scan_wltc.csv'))
    df_ecms = pd.read_csv(os.path.join(RESULTS_DIR, 'Day9_ecms_wltc.csv'))
    df_dp = pd.read_csv(os.path.join(RESULTS_DIR, 'dp_ems_wltc.csv'))
    df_rule = pd.read_csv(os.path.join(RESULTS_DIR, 'Day7_ems_sim_wltc.csv'))

    # 计算关键指标
    rule_h2 = 0.2831  # 已知值
    dp_h2 = df_dp['m_H2_cumul_kg'].iloc[-1]
    ecms_h2 = df_ecms['m_H2_cumul_kg'].iloc[-1]
    rule_soc_end = df_rule['SOC'].iloc[-1]
    dp_soc_end = df_dp['SOC'].iloc[-1]
    ecms_soc_end = df_ecms['SOC'].iloc[-1]

    # 扫描趋势数据
    s_min_row = df_scan.loc[df_scan['H2_kg'].idxmin()]
    s_max_row = df_scan.loc[df_scan['H2_kg'].idxmax()]
    s_mid_row = df_scan[df_scan['s_factor'] == 180].iloc[0]

    # ================================================================
    # TITLE
    # ================================================================
    title = doc.add_heading('ECMS Week 5 结果评估报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('评估日期：2026-06-12 | 评估对象：ECMS 等效消耗最小化策略\n'
                       '对比基准：DP 全局最优 | 规则控制器 Baseline | 工况：WLTC (1800s)')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()

    # ================================================================
    # 1. 总体结论
    # ================================================================
    h(doc, '1. 总体结论', level=1)

    p(doc, '本周完成 ECMS 理论学习 + Python 实现 + 参数扫描 + 三方法对比。', bold=True)

    # 六维度评分卡
    tbl(doc,
        ['评估维度', '评分', '权重', '说明'],
        [
            ['算法正确性',    '⭐⭐⭐☆  (3/5)', '20%', '框架正确，但 SOC 惩罚系数不足导致过充'],
            ['与 DP 接近度',  '⭐☆☆☆☆  (1/5)', '20%', '氢耗差距大（0.5796 vs 0.2287），Week 6 用 A-ECMS 改善'],
            ['SOC 维持能力',  '⭐☆☆☆☆  (1/5)', '15%', 'SOC_end=0.890，严重过充（目标 0.6）'],
            ['代码质量',      '⭐⭐⭐⭐☆  (4/5)', '20%', '结构清晰，复用 day8 组件，注释充分'],
            ['可视化',        '⭐⭐⭐⭐⭐  (5/5)', '15%', '五合一对比图 + 扫描趋势图，信息丰富'],
            ['文档完整性',    '⭐⭐⭐⭐☆  (4/5)', '10%', '原理文档详细，八股文 6 题，缺少 Week 5 小报告'],
        ])
    p(doc, f'  综合评分：约 2.7 / 5.0  — 代码和可视化优秀，核心算法调优留 Week 6')

    # ================================================================
    # 2. 三方法对比结果
    # ================================================================
    h(doc, '2. 三方法对比结果', level=1)
    p(doc, 'WLTC 工况下 Rule / DP / ECMS 的核心指标对比：')

    tbl(doc,
        ['指标', '规则控制器', 'DP（全局最优）', 'ECMS（s=120）', 'ECMS vs DP', '是否达标'],
        [
            ['总氢耗 (kg)',
             f'{rule_h2:.4f}',
             f'{dp_h2:.4f}',
             f'{ecms_h2:.4f}',
             f'+{(ecms_h2/dp_h2-1)*100:.1f}%',
             '❌ 未达标（差2.5倍）'],
            ['SOC 初→终',
             f'0.60→{rule_soc_end:.3f}',
             f'0.60→{dp_soc_end:.3f}',
             f'0.60→{ecms_soc_end:.3f}',
             f'偏差 {(ecms_soc_end-0.6):+.3f}',
             '❌ 严重过充'],
            ['FC 平均效率',
             '40.4%',
             f'{df_dp["P_fc_kW"].apply(lambda x: np.interp(x, [0,2,5,8,10,15,20,25,30],[0,0.28,0.40,0.48,0.50,0.55,0.53,0.48,0.40]) if x>=0 else 0).mean():.1%}',
             f'{df_ecms["P_fc_kW"].apply(lambda x: np.interp(x, [0,2,5,8,10,15,20,25,30],[0,0.28,0.40,0.48,0.50,0.55,0.53,0.48,0.40]) if x>=0 else 0).mean():.1%}',
             '—',
             '⚠️ 偏低'],
            ['FC >50% 占比',
             '20.8%',
             '40.5%',
             '10.3%',
             '-30.2 pp',
             '❌ 远低于 DP'],
            ['SOC 维持',
             '启发式',
             'α+β惩罚',
             '恒定s（无维持）',
             '—',
             '❌ 无SOC维持机制'],
        ])

    p(doc,
        '核心发现：ECMS 在恒定 s=120 时，ECMS 倾向于让电池放电替代 FC 发电（s 小→电池"便宜"），'
        '导致 SOC 一路冲到 0.89。虽然氢耗降低了（相比更大 s 值），但牺牲了 SOC 平衡，'
        '在实际应用中是不可接受的。'
        '这就是为什么标准 ECMS 必须配合自适应机制（A-ECMS）使用。')

    # ================================================================
    # 3. 参数扫描趋势分析
    # ================================================================
    h(doc, '3. 等效因子 s 扫描趋势分析', level=1)
    p(doc, 's 从 120 到 250，步长 5，共 27 组仿真。关键趋势：')

    h(doc, '3.1 氢耗趋势', level=2)
    p(doc,
        f'• s=120 时氢耗最低：{s_min_row["H2_kg"]:.4f} kg（但 SOC_end={s_min_row["SOC_end"]:.3f}，过充）\n'
        f'• s=250 时氢耗最高：{s_max_row["H2_kg"]:.4f} kg（SOC_end={s_max_row["SOC_end"]:.3f}，仍过充）\n'
        f'• 趋势：氢耗随 s 单调递增，符合理论预期（s 越大→电池越贵→多用 FC→氢耗越高）\n'
        f'• 但在整个扫描范围内 SOC_end 始终 ≈ 0.89，说明恒定 s 无法将 SOC 压回 0.6 附近')

    h(doc, '3.2 FC 效率趋势', level=2)
    p(doc,
        f'• s=120 时 FC 平均效率最高：{s_min_row["FC_eff_mean"]:.1%}（FC 少工作，只在高效区）\n'
        f'• s=250 时 FC 平均效率最低：{s_max_row["FC_eff_mean"]:.1%}（FC 被迫多做功）\n'
        f'• 趋势：FC 效率随 s 递减，与理论一致\n'
        f'• FC>50% 占比从 10.3%（s=120）微升至 15.6%（s≥180），变化不大')

    h(doc, '3.3 SOC 终值趋势', level=2)
    p(doc,
        '• 在整个 s ∈ [120, 250] 范围内，SOC_end 始终 ≈ 0.890（不变！）\n'
        '• 这是因为：SOC 惩罚项（500×(SOC_next-0.6)²）在当前电池参数下不足以对抗等效因子的驱动力\n'
        '• 根本原因：电池容量大（50Ah）、SOC 变化慢，单步优化看不到长期 SOC 后果\n'
        '• 这也说明：恒定 s 不是好的 SOC 维持策略，必须有自适应反馈')

    h(doc, '3.4 分段结论', level=2)
    tbl(doc,
        ['s 范围', '氢耗', 'FC 效率', 'SOC_end', '评价'],
        [
            ['120~140', '最低 (0.58~0.61 kg)', '最高 (39~41%)', '0.89 过充', '氢耗最优但 SOC 失控'],
            ['145~165', '中等 (0.62~0.65 kg)', '中等 (37~39%)', '0.89 过充', '折中但无 SOC 维持'],
            ['170~250', '较高 (0.66~0.67 kg)', '较低 (35~36%)', '0.89 过充', 'FC 效率低，无优势'],
        ])

    verdict(doc, 's 扫描趋势整体符合理论预期（氢耗↑s↑，FC效率↓s↑），但 SOC 维持全面失败。', 'warn')

    # ================================================================
    # 4. 与理论预期的对比
    # ================================================================
    h(doc, '4. 与理论预期的对比', level=1)
    p(doc, '根据 Week 5 学习文档中的预期，检验实际结果：')

    tbl(doc,
        ['预期指标', '理论预期', '实际结果', '差距', '是否达标'],
        [
            ['ECMS 氢耗 vs DP', 'DP 的 102~107%', f'{ecms_h2/dp_h2:.1%}', f'+{(ecms_h2/dp_h2-1)*100:.0f}%', '❌ 远高于预期'],
            ['ECMS 氢耗 vs Rule', '比 Rule 低 15~20%', f'{(1-ecms_h2/rule_h2)*100:.1f}%', '负增长', '❌ 反而更差'],
            ['SOC_end', '接近 0.6 (±0.05)', f'{ecms_soc_end:.3f}', f'+{ecms_soc_end-0.6:.3f}', '❌ 严重偏离'],
            ['FC 平均效率', '43~46%', f'{df_ecms["P_fc_kW"].apply(lambda x: np.interp(x, [0,2,5,8,10,15,20,25,30],[0,0.28,0.40,0.48,0.50,0.55,0.53,0.48,0.40]) if x>=0 else 0).mean():.1%}', '-8~10pp', '❌ 偏低'],
            ['FC>50% 占比', '35~42%', '10.3%', '-25~32pp', '❌ 远低于预期'],
            ['计算速度', '<1s', '<0.5s', '达标', '✅ 满足实时性'],
        ])

    verdict(doc, '除计算速度外，所有性能指标均未达预期。根本原因：恒定 s 无法维持 SOC → ECMS 持续过度放电。', 'fail')

    # ================================================================
    # 5. 根因分析
    # ================================================================
    h(doc, '5. 根因分析', level=1)

    h(doc, '5.1 直接原因：SOC 惩罚系数不足', level=2)
    p(doc,
        '当前 SOC 惩罚：500 × (SOC_next - 0.6)²\n'
        '• 当 SOC_next = 0.65 时，惩罚 = 500 × 0.0025 = 1.25 g/s\n'
        '• 但等效因子项 s×P_bat/3600 = 120×10/3600 ≈ 0.33 g/s（10kW 放电时）\n'
        '• 惩罚项仅当 SOC 偏离 0.6 超过约 0.02 时才超过等效因子项\n'
        '• 在 1Hz 采样下，SOC 变化极慢（每步约 0.00016），惩罚积累需要很长时间才能生效\n\n'
        '结论：SOC 惩罚的时空尺度与等效因子项不匹配——前者是慢变量（SOC），后者是快变量（功率）。')

    h(doc, '5.2 根本原因：恒定 s 没有反馈机制', level=2)
    p(doc,
        '标准 ECMS 的设计初衷是"已知工况下的基准测试"，恒定 s 配合 DP 离线标定。\n'
        '在我们的场景中：\n'
        '• 没有用 DP 反推 s*(t)（Week 5 跳过了这一步）\n'
        '• 没有 SOC 反馈自适应（这是 Week 6 的任务）\n'
        '• 惩罚项只是"权宜之计"，无法替代自适应的长期 SOC 维持能力\n\n'
        '这也解释了为什么文献中标准 ECMS 几乎总是配合 A-ECMS 使用——'
        '恒定 s 在实际驾驶中无法保证 SOC 平衡。')

    h(doc, '5.3 次要原因：电池模型参数', level=2)
    p(doc,
        '• SOC 变化幅度极小（每步约 0.00016，1800 步总共约 0.29）\n'
        '• 这意味着即使不给惩罚，SOC 也不会快速恶化\n'
        '• 但这也意味着惩罚项更难"感知"到 SOC 偏移\n'
        '• 如果用更小的电池（如 20Ah），SOC 变化会更显著，惩罚更容易生效\n'
        '→ 但在实际项目中 50Ah 是合理的，不应改参数来适配策略')

    # ================================================================
    # 6. 改进方案（Week 6 路线图）
    # ================================================================
    h(doc, '6. Week 6 改进方案', level=1)
    p(doc, '两个方向并行，优先解决 SOC 维持问题：')

    h(doc, '方案 A：A-ECMS（自适应 ECMS）— 推荐 ⭐', level=2)
    p(doc,
        '核心改动：s(k) = s₀ × [1 + Kp × (SOC_ref - SOC(k))]\n'
        '• SOC < 0.6 时 → s 增大 → 电池放电"变贵" → ECMS 多用 FC → SOC 回升\n'
        '• SOC > 0.6 时 → s 减小 → 电池放电"变便宜" → ECMS 多用电池 → SOC 回落\n'
        '• Kp 是唯一调参（建议范围 20~60）\n\n'
        '预期效果：\n'
        '  - SOC_end 从 0.89 → 0.55~0.63（接近 0.6）\n'
        '  - 氢耗略有增加（因 SOC 维持牺牲少量经济性），但仍在合理范围\n'
        '  - FC>50% 占比改善')

    h(doc, '方案 B：DP 离线标定 s₀', level=2)
    p(doc,
        '用现有 DP 结果反推最优 s*(t) 序列：\n'
        '  s*(t) = -H_fc(P_fc*) / P_bat*\n'
        '• 取 s*(t) 的均值作为 A-ECMS 的 s₀\n'
        '• 这样 s₀ 本身就是"最优策略下的等效因子"，比理论公式更准\n'
        '• 配合方案 A 的自适应，效果更好')

    h(doc, '方案 C：增大 SOC 惩罚系数（不推荐）', level=2)
    p(doc,
        '直接把 SOC 惩罚从 500 加到 5000 或更高。\n'
        '问题：\n'
        '• 惩罚太大 → FC 功率跳变剧烈（不连续）\n'
        '• 惩罚太小 → 仍然压不住 SOC\n'
        '• 这是"手动调参"思路，不如自适应机制优雅\n'
        '→ 可以作为 Week 6 调参过程中的参考，但不作为最终方案')

    # ================================================================
    # 7. 改进优先级
    # ================================================================
    h(doc, '7. 改进优先级排序', level=1)
    tbl(doc,
        ['优先级', '改进项', '预期效果', '难度', '预计时间'],
        [
            ['P0 ⭐', '实现 A-ECMS（Kp 调 SOC）', 'SOC_end→0.6附近，氢耗<0.65kg', '低', '1天'],
            ['P1', 'DP 反推标定 s₀', 's₀ 更准确，减少 Kp 调参范围', '低', '0.5天'],
            ['P2', '优化 SOC 惩罚辅助项', '减少 FC 功率跳变', '中', '0.5天'],
            ['P3', 'NEDC/CLTC 工况验证', '验证泛化性', '低', '0.5天'],
            ['P4', 'C++ 基础练习', 'Week 5 并行任务', '中', '1天'],
        ])

    # ================================================================
    # 8. 数据附录
    # ================================================================
    h(doc, '8. 数据附录', level=1)

    h(doc, '8.1 扫描结果明细（s vs 各指标）', level=2)
    tbl(doc,
        ['s(g/kWh)', 'H2(kg)', 'SOC_end', 'FC_eff', 'FC>50%'],
        [[f'{r["s_factor"]:.0f}', f'{r["H2_kg"]:.4f}', f'{r["SOC_end"]:.3f}',
          f'{r["FC_eff_mean"]:.1%}', f'{r["FC_eff_gt50"]:.1%}']
         for _, r in df_scan.iterrows()])

    h(doc, '8.2 ECMS 最优解（按不同标准）', level=2)
    best_h2 = df_scan.loc[df_scan['H2_kg'].idxmin()]
    best_soc = df_scan.loc[(df_scan['SOC_end'] - 0.6).abs().idxmin()]
    best_eff = df_scan.loc[df_scan['FC_eff_mean'].idxmax()]
    tbl(doc,
        ['优化目标', '最优 s', 'H2(kg)', 'SOC_end', 'FC_eff'],
        [
            ['最小氢耗',     f'{best_h2["s_factor"]:.0f}', f'{best_h2["H2_kg"]:.4f}', f'{best_h2["SOC_end"]:.3f}', f'{best_h2["FC_eff_mean"]:.1%}'],
            ['SOC最接近0.6', f'{best_soc["s_factor"]:.0f}', f'{best_soc["H2_kg"]:.4f}', f'{best_soc["SOC_end"]:.3f}', f'{best_soc["FC_eff_mean"]:.1%}'],
            ['FC效率最高',   f'{best_eff["s_factor"]:.0f}', f'{best_eff["H2_kg"]:.4f}', f'{best_eff["SOC_end"]:.3f}', f'{best_eff["FC_eff_mean"]:.1%}'],
        ])
    p(doc,
        '三个目标的最优 s 完全不同（最小氢耗→120，SOC最准→250，FC最高→120），'
        '说明恒定 s 无法同时满足多个目标——这正是自适应 ECMS 存在的理由。')

    h(doc, '8.3 输出文件清单', level=2)
    tbl(doc,
        ['文件', '类型', '说明'],
        [
            ['scripts/day9_ecms_ems.py', '代码', 'ECMS 完整实现（标准+A-ECMS）'],
            ['results/Day9_ecms_wltc.csv', '数据', 'ECMS 最优 s=120 的仿真结果'],
            ['results/Day9_ecms_scan_wltc.csv', '数据', 's=120~250 全扫描结果'],
            ['results/Day9_ecms_scan_wltc.png', '图表', 's 扫描趋势图（3面板）'],
            ['results/Day9_ECMS_compare_wltc.png', '图表', 'Rule vs DP vs ECMS 五合一对比'],
            ['docs/ECMS_原理与实现_Week5学习文档.docx', '文档', 'ECMS 理论 + 推导 + 代码框架'],
            ['docs/ECMS_Week5_结果评估报告.docx', '文档', '本文件'],
        ])

    # ================================================================
    # Save
    # ================================================================
    out_path = os.path.join(DOCS_DIR, 'ECMS_Week5_结果评估报告.docx')
    doc.save(out_path)
    print(f'[OK] 已生成: {out_path}')
    print(f'     大小: {os.path.getsize(out_path) / 1024:.0f} KB')


if __name__ == '__main__':
    main()
