# -*- coding: utf-8 -*-
"""
生成 MPC-EMS 两个文件的逐行分析 .docx
  - scripts/mpc_ems.py （基础版）
  - scripts/mpc_ems_optimized.py （优化版）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, sys
import numpy as np

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')
SCRIPTS_DIR = os.path.join(PROJECT_ROOT, 'scripts')

# ── helpers ──
def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def p(doc, text, bold=False, size=10.5, color=None, indent=False):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Cm(0.75)
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def code_block(doc, text, size=8.5):
    """添加代码块样式的段落"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(50, 50, 50)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def note(doc, text, label="📌 注意"):
    """添加提示框"""
    para = doc.add_paragraph()
    run = para.add_run(f'{label}：{text}')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0, 70, 130)
    run.font.italic = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def key_point(doc, text):
    """关键要点"""
    para = doc.add_paragraph()
    run = para.add_run(f'✦ {text}')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(180, 60, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def tbl(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table

def tbl2(doc, headers, rows):
    """左对齐表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table


# ================================================================
# 正文构建
# ================================================================
os.makedirs(DOCS_DIR, exist_ok=True)

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── 封面 ──
p(doc, '', size=14)
h(doc, 'MPC-EMS 模型预测控制能量管理', level=0)

p(doc, '逐行代码原理分析文档', bold=True, size=14, color=RGBColor(0, 70, 130))

p(doc, '', size=8)
p(doc, '文件清单：', bold=True, size=11)
p(doc, '  • scripts/mpc_ems.py           — MPC 基础版（网格搜索 + receding horizon）')
p(doc, '  • scripts/mpc_ems_optimized.py  — MPC 优化版（SOC 软约束 + 终点惩罚 + FC 功率变化惩罚）')
p(doc, '', size=4)
p(doc, '生成日期：2026-07-07', size=9, color=RGBColor(120, 120, 120))
p(doc, '研究项目：燃料电池 EMS 能量管理策略对比', size=9, color=RGBColor(120, 120, 120))

doc.add_page_break()

# ========================================================================
# 第一部分：概览
# ========================================================================
h(doc, '第一部分：MPC-EMS 整体架构概览', level=1)

h(doc, '1.1 什么是 MPC（模型预测控制）？', level=2)
p(doc, '模型预测控制（MPC, Model Predictive Control）是一种基于模型的最优控制方法。其核心思想是：')
p(doc, '  ① 在每一时刻 k，利用系统模型预测未来 N_p 步的系统行为', indent=True)
p(doc, '  ② 在当前时刻求解一个有限时域的开环最优控制问题', indent=True)
p(doc, '  ③ 执行最优控制序列中的第一步', indent=True)
p(doc, '  ④ 在下一时刻 k+1，用最新测量值更新状态，重复①②③', indent=True)
p(doc, '  ⑤ 这种"边走边看边规划"的方式称为 Receding Horizon（滚动时域控制）', indent=True)

h(doc, '1.2 本工程中 MPC 的特色', level=2)
p(doc, '本项目中的 MPC 应用于燃料电池混合动力系统（FC + Battery），有以下特点：')
p(doc, '  • 使用网格搜索（Grid Search）替代传统的二次规划求解——简单直观，适合教学', indent=True)
p(doc, '  • 利用已知工况信息（WLTC/NEDC/CLTC）作为未来功率预测的"完美预测"', indent=True)
p(doc, '  • 将 DP 全局最优解作为基准，对比 MPC 的次优性', indent=True)
p(doc, '  • 在优化版中引入了 SOC 软约束、终点欠差惩罚、FC 功率变化惩罚等工程修正', indent=True)

h(doc, '1.3 MPC 与 DP / ECMS / Rule 的定位对比', level=2)

tbl2(doc, ['方法', '原理', '最优性', '在线实时性', '依赖'],
    [
        ['Rule\n规则控制器', '基于 SOC 阈值的 if-else 逻辑', '非最优', '✅ 强', '无'],
        ['DP\n动态规划', '全局后向迭代（Bellman 方程）', '✅ 全局最优', '❌ 弱（需离线算）', '已知完整工况'],
        ['ECMS\n等效氢耗最小化', '瞬时等效最小化（s × P_bat 折算）', '近似最优', '✅ 强', '等效因子 s'],
        ['MPC\n模型预测控制', '滚动时域网格搜索', '介于 DP 和 ECMS', '中（取决于 N_p）', '未来 N_p 步预测'],
    ]
)

p(doc, '', size=4)
key_point(doc, 'MPC 是连接离线最优（DP）与在线实时（ECMS/规则）之间的桥梁：它用有限时域预测逼近全局最优。')

doc.add_page_break()

# ========================================================================
# 第二部分：mpc_ems.py 逐行分析
# ========================================================================
h(doc, '第二部分：mpc_ems.py — 逐行分析', level=1)
p(doc, '文件路径：scripts/mpc_ems.py  |  总行数：571 行', bold=True)

# ── 第1-15行：文件头 ──
h(doc, '2.1 文件头与文档字符串（第 1-15 行）', level=2)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['1', '# -*- coding: utf-8 -*-', 'Python 编码声明，确保源文件中可以包含中文字符。'],
        ['2-14', '""" """ 文档字符串',
         'MPC 模块的功能说明、用法示例和依赖信息。\n'
         '用法行（7-10）展示了四种调用方式：\n'
         '  (a) 无参数 → 默认 WLTC + N_p=50\n'
         '  (b) --cycle nedc → 切换工况\n'
         '  (c) --np 30 → 修改预测时域\n'
         '  (d) --compare → 四方法对比模式'],
        ['15', '# 无内容', '空行，分隔文档和代码。'],
    ]
)

note(doc, 'docstring 中的用法说明本身就是对 MPC 超参数（--cycle、--np）的文档化，是良好的 Python 工程实践。')

# ── 第16-35行：导入 ──
h(doc, '2.2 模块导入与路径配置（第 16-35 行）', level=2)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['16-21', 'import os, sys, argparse\nimport numpy as np\nimport pandas as pd\nimport matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt',
         '标准科学计算库。\nmatplotlib.use("Agg") 非交互后端，用于服务器/批处理环境，不显示图形窗口。'],
        ['23-26', 'PROJECT_ROOT = ...\nRESULTS_DIR = ...\nSCRIPTS_DIR = ...\nsys.path.insert(0, SCRIPTS_DIR)',
         '工程路径配置：\n  • PROJECT_ROOT = 工程根目录（通过 __file__ 向上两级）\n  • 将 SCRIPTS_DIR 加入 sys.path，使后续 import 可以找到 day8_dp_ems 等模块。'],
        ['29-36', 'from day8_dp_ems import (\n    fc_hydrogen_flow, ..., load_drive_cycle, ...,\n    SOC_MIN, SOC_MAX, ..., OCV_LU, Q_BAT, R_INT,\n)',
         '从 day8_dp_ems 复用核心组件：\n  • fc_hydrogen_flow — 燃料电池氢耗映射函数\n  • fc_efficiency — FC 效率计算\n  • vehicle_power — 从速度计算功率需求\n  • state_transition — 系统状态转移方程\n  • load_drive_cycle — 加载工况数据\n  • SOC_MIN/SOC_MAX/PFC_MIN/PFC_MAX — 系统边界约束\n  • DT — 采样时间（1s）\n  • LHV_H2 — 氢气低热值\n  • SOC_BP/OCV_LU/Q_BAT/R_INT — 电池模型参数'],
        ['38-39', '# 从 day9_ecms_ems 复用常量\n# 注意：...N_PFC=60',
         '注释说明 N_PFC 同时用于 DP 网格密度和 MPC/ECMS 控制搜索网格。'],
    ]
)

note(doc, 'mpc_ems.py 大量复用 day8_dp_ems 的模型函数，体现了代码复用原则——MPC、DP、ECMS 共享同一套车辆模型和电池参数。', label='🔗 架构说明')

# ── 第41-54行：MPC 参数 ──
h(doc, '2.3 MPC 参数定义（第 41-54 行）', level=2)

tbl2(doc, ['行号', '代码/参数', '含义与设计说明'],
    [
        ['44', 'N_P_DEFAULT = 50', '默认预测时域 N_p=50 步（即 50 秒）。这是 MPC 最重要的超参数：\n'
         '  • N_p 越大，控制器看到的未来越多，决策越接近全局最优\n'
         '  • 但计算量线性增长（每步 need 仿真 horizon 步）\n'
         '  • N_p=50 是 WLTC 全程 ≈1800s 的 ~1/36，属于中等时域'],
        ['45', 'S_MPC = 130.0', '等效因子（Equivalent Factor），单位 [g/kWh]。\n'
         '功能：将电池功率消耗折算为"未来的氢耗"，加入代价函数。\n'
         '物理意义：如果 MPC 只用氢耗做代价，它会优先用电池放电（白捡的能量），'
         '导致 SOC 快速下降。加入 S_MPC × |P_bat| 项后，\n'
         '  "用电" 和 "烧氢" 在代价上等价，实现 charge-sustaining。'],
        ['46', 'W_SOC = 500.0', 'SOC 维持惩罚权重。当 SOC 偏离参考值 > 0.05 时，\n'
         '在代价函数中加入 W_SOC × (dev)² 项。\n'
         '  • 权重 500 相对于氢耗（≈0.01-0.05 g/s·步）较大\n'
         '  • 确保 SOC 不会长期偏离目标'],
        ['47', 'BETA_TERM = 1000.0', '终端 SOC 惩罚系数 β_term。\n'
         '仅在仿真后 30% 阶段生效，对终端 SOC 偏差施加额外惩罚。\n'
         '设计动机：防止 MPC 在接近工况结束时选择"透支电池"策略。'],
        ['48', 'PFC_GRID = np.linspace(PFC_MIN, PFC_MAX, N_PFC)',
         '控制量搜索网格：在 [PFC_MIN, PFC_MAX] 上均匀取 N_PFC=60 个离散值。\n'
         'MPC 枚举这 60 个候选值，选择使代价最小的 P_fc。\n'
         '网格越密，最优解越精确，但计算量线性增长。'],
        ['49', 'SOC_REF = 0.6', '目标 SOC 值（60%）。\n燃料电池混动系统通常维持 SOC 在 50-70%，\n'
         '60% 是在充放电能力与寿命之间的平衡点。'],
    ]
)

key_point(doc, 'MPC 的核心调优参数是 N_p（预测时域）和 S_MPC（等效因子），前者控制"看得多远"，后者控制"用电多贵"。')

# ── 第51-54行：氢耗预计算 ──
h(doc, '2.4 氢耗预计算（第 51-54 行）', level=2)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['54', 'H2_GRID = fc_hydrogen_flow(PFC_GRID)',
         '预计算所有网格点对应的氢耗率 [g/s]。\n'
         '加速原理：在主仿真循环中，每个候选 P_fc 的氢耗需要被多次引用，\n'
         '预计算避免了重复调用 fc_hydrogen_flow（内部包含分段插值/查表，\n'
         '计算开销相对较大）。\n'
         'H2_GRID shape = (N_PFC,) = (60,)，每格对应一个 g/s 值。'],
    ]
)

# ── 第56-77行：单步仿真 ──
h(doc, '2.5 单步状态转移函数：mpc_step_soc（第 56-77 行）', level=2)

p(doc, '功能：根据当前 SOC、燃料电池功率、负载功率，计算下一步 SOC。', bold=True)

tbl2(doc, ['行号', '代码', '原理分析'],
    [
        ['60-63', 'def mpc_step_soc(soc_k, p_fc, p_load_k, dt=DT):\n    """..."""',
         '函数签名，与 day8_dp_ems.state_transition 相同，但返回 float 而非 ndarray。\n'
         '返回 float 的理由：在 N_p 步滚动仿真中，每次只有 1 个候选控制量的 1 步预测，\n'
         '逐元素运算比 array 运算更轻量（避免 numpy 广播的开销）。'],
        ['65', 'p_bat = p_load_k - p_fc', '电池功率 = 负载功率 - 燃料电池功率。\n'
         '  • p_bat > 0：电池放电（提供不足的功率）\n'
         '  • p_bat < 0：电池充电（FC 功率过剩）\n'
         '  • p_bat ≈ 0：功率平衡'],
        ['66', 'v_oc = np.interp(soc_k, SOC_BP, OCV_LU)',
         '开路电压查找：利用 SOC-OCV 曲线（SOC_BP 为断点，OCV_LU 为对应电压值）。\n'
         'np.interp 是线性插值，SOC_BP 和 OCV_LU 来自 day8 的电池参数辨识。'],
        ['67', 'p_w = p_bat * 1000.0', '功率从 [kW] 转换为 [W]，因为电池内阻 R_INT 的单位是 Ω，\n'
         '计算电流需要以 W 为单位。1 kW = 1000 W。'],
        ['69', 'delta = v_oc ** 2 - 4 * R_INT * p_w',
         '判别式 Δ，来自电池等效电路模型的一元二次方程求解：\n'
         '  P_bat = V_oc × I - I² × R_int\n'
         '  → R_int × I² - V_oc × I + P_bat = 0\n'
         '  → I = (V_oc ± √(V_oc² - 4×R_int×P_bat)) / (2×R_int)\n'
         '判别式 Δ 必须 ≥ 0，否则说明负载需求超过电池物理极限。'],
        ['70-72', 'if delta < 0:\n    return np.clip(soc_k, SOC_MIN, SOC_MAX)',
         '物理不可行时，保持当前 SOC 并 clip 在边界内。\n'
         '这是基础版的简化策略——相当于"忽略不可行候选"，\n'
         '不做额外处理。优化版（mpc_ems_optimized.py）改为返回 None，将筛选交给上层。'],
        ['74-77', 'i = (v_oc - np.sqrt(delta)) / (2 * R_INT)\n'
         'i = np.clip(i, -300, 300)\n'
         'soc_next = soc_k - i / (Q_BAT * 3600) * dt\n'
         'return np.clip(soc_next, SOC_MIN, SOC_MAX)',
         '电流计算：取 Δ≥0 时的小根（物理上电池放电电流取正）。\n'
         '  • 电流限制 |i| ≤ 300A（保护电池）\n'
         '  • SOC 更新：ΔSOC = -I / (Q_bat × 3600) × Δt\n'
         '    - Q_bat 单位 Ah，×3600 转换为 As\n'
         '    - 除以 3600 将 Ah 转为 As（库仑计数）\n'
         '  • 最终 clip 到 [SOC_MIN, SOC_MAX] 保证数值安全'],
    ]
)

note(doc, '第 70 行的"保持 SOC"策略有一定风险：如果多个候选都不可行，控制器可能"卡"在不可行状态。优化版通过返回 None 来区别处理。')

# ── 第79-194行：MPC 仿真 ──
h(doc, '2.6 MPC 主仿真函数：mpc_sim（第 79-194 行）', level=2)
p(doc, '这是 MPC 的核心——网格搜索 + Receding Horizon 的完整实现。', bold=True)

h(doc, '2.6.1 函数签名与文档（第 83-104 行）', level=3)

p(doc, '函数签名：')
code_block(doc, 'def mpc_sim(P_load, SOC_0=0.6, N_p=N_P_DEFAULT, w_soc=W_SOC,\n            beta_term=BETA_TERM, soc_ref=SOC_REF):')

tbl2(doc, ['参数', '类型', '默认值', '说明'],
    [
        ['P_load', 'array (N,)', '-', '未来整个工况的功率需求 [kW]（MPC 内部只用 k 到 k+N_p 段）'],
        ['SOC_0', 'float', '0.6', '初始 SOC'],
        ['N_p', 'int', '50', '预测时域步数（最重要的超参数）'],
        ['w_soc', 'float', '500.0', 'SOC 偏离惩罚权重'],
        ['beta_term', 'float', '1000.0', '终端 SOC 惩罚系数'],
        ['soc_ref', 'float', '0.6', 'SOC 参考值'],
    ]
)

p(doc, '返回字典包含：time, SOC, SOC_end, P_fc_kW, P_bat_kW, m_H2_g, m_H2_cumul_kg, fc_efficiency。')

h(doc, '2.6.2 初始化（第 106-119 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['106-112', 'N = len(P_load)\nSOC = np.zeros(N+1)\nP_fc = np.zeros(N)\nP_bat = np.zeros(N)\nm_H2 = np.zeros(N)\nSOC_pred_history = []',
         'N: 总仿真步数。SOC 多分配 1 个元素（SOC[N] 为终点 SOC）。\n'
         'SOC_pred_history 记录每步的预测轨迹，用于事后分析（当前未在 main 中使用）。'],
        ['113', 'SOC[0] = SOC_0', '设定初始 SOC。'],
        ['116', 'penalty_start = int(N * 0.7)', '终端惩罚生效起点：后 30% 阶段。\n'
         '设计动机：仿真前期让 MPC 自由探索，最后 30% 强制引导 SOC 回到参考值附近。\n'
         '对 WLTC（N≈1800）：penalty_start ≈ 1260，即最后 540 秒开始强约束。'],
        ['118-119', 'print(f\'[MPC] N_p={N_p}...\')', '控制台输出，便于运行时监控。'],
    ]
)

h(doc, '2.6.3 主循环：滚动时域（第 121-176 行）', level=3)

p(doc, '这是 MPC 的核心循环，对每个时刻 k 执行以下操作：', bold=True)

p(doc, '', size=2)

p(doc, '步骤① — 取未来 N_p 步的功率预测（第 124-126 行）：', bold=True)
tbl2(doc, ['行号', '代码', '说明'],
    [
        ['125', 'horizon = min(N_p, N - k)', '在工况末尾时 horizon < N_p，避免越界。'],
        ['126', 'p_load_pred = P_load[k : k + horizon]',
         '取未来功率段。\n★ 这是一个"已知工况"（已知 future disturbance）的假设——\n'
         '在实际应用中，这里应接入预测模块的预测值，而非真实值。'],
    ]
)

note(doc, '在真实车辆上，未来的功率需求来自导航/GPS/交通信息预测。这里用真实值替代，属于"上帝视角"的离线仿真。')

p(doc, '步骤② — 枚举所有候选控制并仿真 N_p 步（第 128-167 行）：', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['129-130', 'J_best = np.inf\nbest_j = 0', '初始化最优代价和对应的网格索引。'],
        ['132', 'for j in range(N_PFC):', '遍历 60 个候选 P_fc 值。'],
        ['133-134', 'p_fc_cand = PFC_GRID[j]\nh2_cand = H2_GRID[j]',
         '取出候选控制量及其预计算的氢耗率。'],
        ['137-138', 'soc_pred = soc_k\nJ_total = 0.0',
         '初始化预测 SOC（从当前 SOC 开始前向仿真）和累计代价。'],
        ['140', 'for i in range(horizon):', '对 horizon 内的每一步进行仿真。'],
        ['141-142', 'p_load_i = p_load_pred[i]\np_bat_i = p_load_i - p_fc_cand',
         '当前步的负载功率和电池功率。\n注意：这里假设在 horizon 内 P_fc 恒定不变（每个候选值固定运行 horizon 步）。\n'
         '这是一种简化，实际 MPC 可以允许控制量在 horizon 内变化（多变量优化），\n'
         '但计算量会指数级增长。'],
        ['145', 'J_total += h2_cand * DT',
         '累加氢耗代价：h2_cand [g/s] × DT [s] = 本步耗氢量 [g]。'],
        ['148-151', '# ★ 关键修正：等效能量平衡惩罚\n...\nJ_total += S_MPC * abs(p_bat_i) / 3600.0 * DT',
         '这是 MPC 实现中最关键的修正。\n'
         '原理：如果不加这一项，MPC 会优先用电池供电（因为电池不产生氢耗），\n'
         '导致 SOC 快速下降。加入 S_MPC×|P_bat|/3600 后：\n'
         '  • |P_bat|/3600 将 kW 转换为 kWh/s\n'
         '  • S_MPC [g/kWh] 将电功率折算为等效氢耗 [g/s]\n'
         '  • 物理含义：用电 = 以后要用氢补回来，等价于当前就计代价\n'
         '  • 这本质上与 ECMS 的等效氢耗最小化思想一致！'],
        ['154', 'soc_pred = mpc_step_soc(soc_pred, p_fc_cand, p_load_i)',
         '执行一步状态转移，更新预测 SOC。'],
        ['157-159', 'soc_dev = soc_pred - soc_ref\nif abs(soc_dev) > 0.05:\n    J_total += W_SOC * soc_dev ** 2 * DT',
         'SOC 维持惩罚：\n  • 当 SOC 偏离参考值超过 5% 阈值时触发\n  • 二次型惩罚 (dev²) 使偏差越大边际成本越高\n  • DT 乘子使惩罚与步长成比例\n'
         '死区 0.05 的作用：避免控制器为极小偏差反复切换控制策略。'],
        ['162-163', 'if i == horizon - 1 and k >= penalty_start:\n    J_total += beta_term * (soc_pred - soc_ref) ** 2',
         '终端惩罚：仅在 horizon 最后一步且仿真进入后 30% 阶段时生效。\n'
         '  • beta_term=1000 大于 w_soc=500，表明终点 SOC 约束更严格\n'
         '  • 二次型确保偏差剧烈受罚'],
        ['165-167', 'if J_total < J_best:\n    J_best = J_total\n    best_j = j',
         '更新最优候选：保留使累计代价最小的 P_fc 索引。'],
    ]
)

p(doc, '步骤③ — 执行最优控制的第一步（第 169-176 行）：', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['170-173', 'P_fc[k] = PFC_GRID[best_j]\nP_bat[k] = P_load[k] - P_fc[k]\nm_H2[k] = fc_hydrogen_flow(P_fc[k]) * DT\nSOC[k+1] = mpc_step_soc(soc_k, P_fc[k], P_load[k])',
         '执行最优控制的第一步（Receding Horizon 的核心）：\n'
         '  • 记录当前时刻的 FC 功率、电池功率、氢耗\n'
         '  • 更新 SOC 到下一步\n'
         '  • 下一时刻 k+1，将从新 SOC 重新开始整个流程'],
        ['175-176', 'if k % 300 == 0:\n    print(...)',
         '每 300 步输出一次进度，便于监控长时间仿真（WLTC 约 1800 步，输出 6 次）。'],
    ]
)

h(doc, '2.6.4 后处理与返回（第 178-194 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['179', 'print(f\'[MPC] 完成，SOC_end = {SOC[-1]:.3f}\')',
         '输出终点 SOC，是评估 charge-sustaining 能力的关键指标。\n理想情况：SOC_end ≈ SOC_0 = 0.6。'],
        ['182-183', 'p_fc_arr = P_fc\neff_arr = fc_efficiency(p_fc_arr)',
         '批量计算 FC 效率。fc_efficiency 内部包含 P_fc → 效率的查表/插值映射。'],
        ['185-194', 'return {...}',
         '返回字典包含全部仿真结果，供后续分析、绘图和保存。'],
    ]
)

doc.add_page_break()

# ── 第197-230行：N_p 敏感性分析 ──
h(doc, '2.7 N_p 敏感性分析：mpc_n_p_scan（第 197-230 行）', level=2)

p(doc, '功能：扫描不同预测时域 N_p 对氢耗和 SOC_end 的影响。', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['200-208', 'def mpc_n_p_scan(P_load, N_p_values=None, SOC_0=0.6):\n    """..."""',
         '默认扫描 N_p = [10, 20, 30, 50, 80, 120, 200]。\n'
         '涵盖从极短（10s）到极长（200s）的时域，用于观察：\n'
         '  • N_p 增大是否持续改善氢耗（收敛趋势）\n'
         '  • 是否存在"最佳时域"（N_p 过大可能引入噪声或终点效应）'],
        ['218-222', 'for n_p in N_p_values:\n    if n_p > len(P_load): continue\n    res = mpc_sim(P_load, SOC_0=SOC_0, N_p=n_p)',
         '对每个 N_p 运行完整 MPC 仿真，记录氢耗和终点 SOC。'],
        ['223-228', 'results.append({...})', '收集结果到列表，最终转换为 DataFrame。'],
        ['230', 'return pd.DataFrame(results)', '返回 DataFrame 便于后续绘图和 csv 保存。'],
    ]
)

key_point(doc, 'N_p 敏感性分析揭示了 MPC 的核心权衡：更长的预测时域 = 更好的性能 + 更多的计算量。实际应用中需找到"够好"而非"最好"的 N_p。')

# ── 第232-328行：四方法对比可视化 ──
h(doc, '2.8 四方法对比图：plot_four_way（第 232-328 行）', level=2)

p(doc, '功能：绘制 Rule / DP / ECMS / MPC 四种方法的五合一对比图。', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['236-238', 'def plot_four_way(t, v, P_load, rule, dp, ecms, mpc_result, cycle_name="wltc"):',
         '输入：时间 t、速度 v、负载功率 P_load、四种方法的仿真结果字典。'],
        ['244', 'fig, axes = plt.subplots(5, 1, figsize=(16, 14), sharex=True)',
         '5×1 子图布局：\n  (1) 速度 + SOC 对比\n  (2) 功率分配对比\n'
         '  (3) SOC 对比\n  (4) 累计氢耗对比\n  (5) FC 效率直方图'],
        ['246-247', 'colors / linestyles',
         '四种方法的区分配色：橘/绿/蓝/红，线型分别用 -- / - / -. / :，确保灰度打印也能区分。'],
        ['249-271', '子图 (1) 速度 + SOC',
         '双 y 轴：左轴速度（蓝色透明），右轴 SOC（四种颜色四条线）。\n'
         '从速度- SOC 的对照可以看出急加速时 SOC 的下降。'],
        ['274-286', '子图 (2) 功率分配',
         '阴影填充负载功率，实线为 DP 和 MPC 的 FC 功率。\n'
         '绿色/橙色填充区分别表示电池放电/充电，直观展示功率分担关系。'],
        ['289-298', '子图 (3) SOC 对比',
         '四条 SOC 轨迹直接对比。灰色虚线为 SOC_ref=0.6 参考线。\n'
         '可以直观看出哪种方法 SOC 维持最好、哪种过度波动。'],
        ['301-309', '子图 (4) 累计氢耗对比',
         '四条氢耗曲线的斜率反映瞬时效率，终值反映总氢耗。\n'
         '图例中标注最终数值，便于直接量化对比。'],
        ['312-322', '子图 (5) FC 效率直方图',
         '将 FC 效率分布画成直方图，标注均值。\n'
         '帮助理解不同控制策略下 FC 工作点的分布差异。\n'
         '高效区间（如 40-55%）占比越高越好。'],
        ['324-328', 'plt.tight_layout()\nplt.savefig(...)\nplt.close()',
         '保存高分辨率 PNG 到 RESULTS_DIR，关闭图形释放内存。'],
    ]
)

# ── 第331-377行：指标打印 ──
h(doc, '2.9 指标打印：print_four_way_metrics（第 331-377 行）', level=2)

p(doc, '功能：格式化打印四种方法的对比指标表格。', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['341-354', '提取各方法的氢耗、SOC_end、FC 效率',
         '使用 get() 方法保证兼容性——dp/ecms 可能没有 fc_efficiency 字段。'],
        ['356-365', 'rows = [\n    (\'总氢耗 (kg)\', ...),\n    (\'SOC 初值→终值\', ...),\n    ...\n]',
         '定义对比指标行：总氢耗、SOC 变化、FC 平均效率、FC 最大功率、总能量需求。'],
        ['367-369', 'for row in rows:\n    print(...)',
         '按列对齐格式输出，形成整齐的表格。'],
        ['373-376', '相对 DP 的氢耗差距',
         '计算 (±% )：\n  MPC 相对 DP 的氢耗增量百分比\n  Rule 的差距最大（+30-60%）\n  MPC 的差距居中（+5-20%）\n  这是核心对比结论。'],
    ]
)

# ── 第379-412行：N_p 敏感性图 ──
h(doc, '2.10 N_p 敏感性绘图：plot_np_sensitivity（第 379-412 行）', level=2)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['383-384', 'def plot_np_sensitivity(np_df, dp_H2, cycle_name="wltc"):',
         '绘制两行子图：上为氢耗 vs N_p，下为 SOC_end vs N_p。'],
        ['391-396', 'ax1.plot(np_df[\'N_p\'], np_df[\'H2_kg\'], \'ro-\', ...)\nax1.axhline(y=dp_H2, ...)',
         '红色虚线连接 N_p-H2 点，绿色横线标注 DP 最优值。\n'
         '理想特征：N_p 增大 → H2 下降，趋近 DP 水平。'],
        ['400-406', 'ax2.plot(np_df[\'N_p\'], np_df[\'SOC_end\'], \'bo-\', ...)\nax2.axhline(y=SOC_REF, ...)',
         '蓝色线显示 SOC_end 随 N_p 的变化。\n'
         '理想情况：SOC_end 随 N_p 增大趋近 0.6。'],
    ]
)

# ── 第414-570行：主程序 ──
h(doc, '2.11 主程序 main（第 414-570 行）', level=2)

h(doc, '2.11.1 参数解析（第 418-425 行）', level=3)
tbl2(doc, ['行号', '代码', '说明'],
    [
        ['419-424', 'parser = argparse.ArgumentParser(...)\nparser.add_argument(\'--cycle\', ...)\nparser.add_argument(\'--np\', ...)\nparser.add_argument(\'--scan\', ...)\nparser.add_argument(\'--compare\', ...)\nparser.add_argument(\'--plot-only\', ...)',
         'argparse 命令行参数解析：\n  --cycle: 工况选择（wltc/nedc/cltc）\n  --np: 预测时域\n  --scan: 是否跑 N_p 扫描\n  --compare: 四方法对比\n  --plot-only: 仅查看已有结果（跳过仿真）'],
    ]
)

h(doc, '2.11.2 流程控制（第 427-565 行）', level=3)

p(doc, '主程序的执行流程分为 8 个步骤：')
p(doc, '  步骤 1（第 436-439 行）：加载工况数据并计算功率需求')
p(doc, '  步骤 2（第 442-443 行）：运行规则控制器作为 baseline')
p(doc, '  步骤 3（第 446-451 行）：运行 DP 作为全局最优基准')
p(doc, '  步骤 4（第 454-461 行）：可选运行 ECMS 作对比')
p(doc, '  步骤 5（第 464-465 行）：运行 MPC 主仿真')
p(doc, '  步骤 6（第 468-542 行）：打印指标并绘图')
p(doc, '  步骤 7（第 545-556 行）：保存 MPC 结果到 CSV')
p(doc, '  步骤 8（第 559-563 行）：可选执行 N_p 敏感性扫描')

key_point(doc, '主程序在 --compare 模式下运行四种方法并对比，这是整个项目"EMS 方法全景对比"的核心产出。')

notebook_pages = """
主要逻辑：

第 436-439 行：加载工况
  使用 day8_dp_ems.load_drive_cycle 读取 WLTC/NEDC/CLTC 的速度-时间序列，
  再通过 vehicle_power 计算每个时刻的功率需求 P_load。

第 442-443 行：规则控制器
  run_rule_controller 是基于 SOC 阈值的简单 if-else 策略。
  虽然是"最低级"的控制，但提供了一个重要的 reference baseline。

第 447-451 行：DP 全局最优
  调用 backward_dp 和 forward_rollout：
  ┌─────────────┐     ┌─────────────┐
  │ backward_dp │  →  │forward_ro. │
  │ (后向迭代)  │     │ (前向推出)  │
  └─────────────┘     └─────────────┘
  DP 给出全局最优控制策略，是所有其它方法的性能上限。

第 454-461 行：ECMS
  仅在 --compare 模式下运行。day9_ecms_ems.ecms_sim 使用瞬时等效因子方法。

第 464-465 行：MPC
  调用 mpc_sim，传入 P_load、初始 SOC=0.6、用户指定的 N_p。

第 469-542 行：输出对比
  --compare 模式 → 4 方法指标 + 5 合一图
  非 compare 模式 → 3 方法指标（Rule/DP/MPC）+ 4 合一图

第 545-556 行：保存结果
  将 MPC 仿真结果（时间、速度、功率、SOC、氢耗）写入 CSV。
  路径格式：results/mpc_ems_{cycle}_np{n_p}.csv

第 559-563 行：N_p 扫描
  如果指定 --scan，对 [10,20,30,50,80,120,200] 逐一运行 MPC 仿真。
"""

doc.add_page_break()

# ========================================================================
# 第三部分：mpc_ems_optimized.py 逐行分析
# ========================================================================
h(doc, '第三部分：mpc_ems_optimized.py — 逐行分析', level=1)
p(doc, '文件路径：scripts/mpc_ems_optimized.py  |  总行数：719 行', bold=True)
p(doc, '在 mpc_ems.py 基础上引入的三大改进：', bold=True)
p(doc, '  1️⃣ SOC 软约束与死区（deadband）机制', indent=True)
p(doc, '  2️⃣ 终点 SOC 欠差惩罚（route-end penalty）', indent=True)
p(doc, '  3️⃣ FC 功率变化惩罚（slew rate penalty）——抑制功率跳变、保护燃料电池寿命', indent=True)

h(doc, '3.1 与基础版的共同部分', level=2)
p(doc, '第 1-39 行（模块导入、路径配置、复用 day8 核心组件）与 mpc_ems.py 完全相同，此处不重复分析。')
p(doc, '关键不同从参数定义（第 41-56 行）开始。', bold=True)

h(doc, '3.2 优化版 MPC 参数（第 41-55 行）', level=2)

tbl2(doc, ['行号', '参数', '基础版值', '优化版值', '变化分析'],
    [
        ['44', 'N_P_DEFAULT', '50', '50', '不变'],
        ['45', 'S_MPC', '130', '130', '不变（等效因子已经较为理想）'],
        ['46', 'W_SOC', '500', '1200', '⬆ 从 500→1200，SOC 惩罚大幅加强，'
         '因为优化版增加了 deadband 机制（0.015 死区），\n'
         '在死区内无惩罚，但超出后以更高权重补偿'],
        ['47', 'BETA_TERM', '1000', '5000', '⬆ 从 1000→5000，终端 SOC 约束更强，'
         '确保终点 SOC 更接近参考值'],
        ['48', 'SOC_DEADBAND', '-', '0.015', '🆕 SOC 死区：在 SOC_ref ± 0.015 范围内不做惩罚，\n'
         '避免控制器为极小偏差频繁切换（减少 FC 启停抖动）'],
        ['49', 'SOC_SOFT_MIN', '-', '0.57', '🆕 SOC 软下限：SOC 低于 0.57 时启动重罚。\n'
         '作用：防止 MPC"透支"电池来换取更低的原始氢耗'],
        ['50', 'W_SOC_LOW', '-', '20000', '🆕 低 SOC 软约束惩罚权重（20000，非常大）。\n'
         '一旦 SOC 低于 0.57，惩罚项 W_SOC_LOW × (0.57 - SOC)² 将快速变大，\n'
         '迫使控制器提高 P_fc 为电池充电。'],
        ['51', 'SOC_FINAL_TOL', '-', '0.01', '🆕 终点 SOC 允许偏差容限。\n'
         '要求 SOC_end ≥ SOC_REF - 0.01 = 0.59'],
        ['52', 'W_FINAL_SOC', '-', '80000', '🆕 终点 SOC 不足惩罚权重（80000）。\n'
         '一旦 SOC_end 低于 0.59，惩罚极大，几乎强制满足终点约束。'],
        ['53', 'W_PFC_SLEW', '-', '0.001', '🆕 FC 功率变化率惩罚权重。\n'
         '每次候选 P_fc 与上一步 P_fc 的差值平方计入代价。\n'
         '0.001 看似很小，但乘以 (ΔP_fc)² 后可在大幅跳变时产生显著惩罚。'],
    ]
)

note(doc, '优化版引入了 6 个新参数，参数总量比基础版翻了一倍。这体现了"从教学演示到工程优化"的演进——更多约束 = 更好性能 + 更多调参工作量。')

h(doc, '3.3 单步仿真改进（第 63-85 行）', level=2)
p(doc, '函数 mpc_step_soc 与基础版核心相同，但有一处关键改进：', bold=True)

tbl2(doc, ['行号', '基础版', '优化版', '改进说明'],
    [
        ['70-72 vs 77-78', 'delta < 0 时 clip 保持 SOC',
         'delta < 0 或 SOC 越界时 return None',
         '基础版用 clip"掩盖"了物理不可行，可能导致选用实际上不可行的候选。\n'
         '优化版返回 None，调用方将该候选的代价设为 ∞，\n'
         '彻底排除不可行解。这保证了控制决策的安全性。'],
        ['83-84', '无', 'if not np.isfinite(soc_next) or soc_next < SOC_MIN or soc_next > SOC_MAX:\n    return None',
         '新增数值稳定性检查：如果 soc_next 出现 NaN/Inf 或越界，\n'
         '也返回 None。这是防御性编程的体现。'],
    ]
)

h(doc, '3.4 SOC 等效氢耗修正：soc_equivalent_h2（第 87-98 行）', level=2)
p(doc, '这是优化版新增的辅助函数，用于报告层面的公平比较。', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['88-91', 'def soc_equivalent_h2(raw_h2_kg, soc_end, soc_ref=SOC_REF, s_factor=S_MPC):\n    """将终端 SOC 偏差折算成等效氢耗"""',
         '设计动机：两个控制器 A（SOC_end=0.60, H2=0.5kg）和 B（SOC_end=0.50, H2=0.45kg），\n'
         '表面 B 更省氢，但 B 消耗了更多电池储能。\n'
         'soc_equivalent_h2 将"少充的电/多放的电池"折算为等效氢耗，\n'
         '使 comparison 公平。'],
        ['96', 'delta_soc = soc_ref - soc_end',
         'SOC 偏差：SOC 偏低 → delta_soc > 0 → 需要"补"等效氢耗。'],
        ['97', 'e_bat_kwh = Q_BAT * np.mean(OCV_LU) * delta_soc / 1000.0',
         '将 SOC 偏差转换为电池能量偏差 [kWh]：\n'
         '  Q_BAT [Ah] × OCV_mean [V] × ΔSOC / 1000 [Wh→kWh]'],
        ['98', 'return raw_h2_kg + s_factor * e_bat_kwh / 1000.0',
         '等效氢耗 = 原始氢耗 + 电池能量偏离的经济代价。\n'
         's_factor [g/kWh] / 1000 [g→kg] × e_bat_kwh [kWh] = 等效氢耗增量 [kg]'],
    ]
)

key_point(doc, 'SOC 等效氢耗修正让"公平比较"成为可能——对 charge-sustaining 和 charge-depleting 策略在同一尺度上对比。')

h(doc, '3.5 SOC 跟踪惩罚：soc_tracking_penalty（第 100-128 行）', level=2)
p(doc, '优化版将 SOC 相关的所有惩罚项封装为一个独立函数，这是重要的架构改进。', bold=True)

tbl2(doc, ['行号', '代码/惩罚项', '说明'],
    [
        ['101-105', 'def soc_tracking_penalty(soc, is_terminal, is_route_end,\n    w_soc=W_SOC, beta_term=BETA_TERM, ...):',
         '函数签名，包含所有 SOC 惩罚参数，共 9 个参数。\n'
         '两个标志位 is_terminal / is_route_end 区分不同阶段的约束。'],
        ['114', 'abs_dev = abs(soc - soc_ref)', '计算绝对偏差。'],
        ['115', 'excess = max(abs_dev - soc_deadband, 0.0)',
         '死区处理：只有当偏差 > 0.015 时才产生惩罚。\n'
         'excess 是"超出死区的偏差量"。'],
        ['116', 'penalty = w_soc * excess ** 2 * DT',
         '常规 SOC 跟踪惩罚：\n  w_soc=1200（基础版为 500，增强 2.4 倍）\n  二次型确保偏离越大惩罚越重'],
        ['118-119', 'low_gap = max(soc_soft_min - soc, 0.0)\npenalty += w_soc_low * low_gap ** 2 * DT',
         'SOC 软下限惩罚：\n  soc_soft_min=0.57\n  w_soc_low=20000（远大于 w_soc=1200）\n'
         '效果：当 SOC 低于 0.57 时，惩罚瞬间变重，强制 FC 充电'],
        ['121-122', 'if is_terminal:\n    penalty += beta_term * excess ** 2',
         '滚动窗口终端惩罚（同基础版，但系数从 1000 提升到 5000）。'],
        ['124-126', 'if is_route_end:\n    final_shortfall = max((soc_ref - soc_final_tol) - soc, 0.0)\n    penalty += w_final_soc * final_shortfall ** 2',
         '真实工况终点惩罚（优化版新增）：\n'
         '  soc_final_tol=0.01，要求 SOC_end ≥ 0.59\n'
         '  w_final_soc=80000，几乎强制满足该约束\n'
         '与 is_terminal 的区别：is_terminal 是滚动窗口的末端，\n'
         '  is_route_end 是整个仿真工况的真正终点。'],
    ]
)

note(doc, 'soc_tracking_penalty 函数的设计体现了现代 MPC 中"软约束（soft constraint）"的思想——不硬性禁止 SOC 过低，而是通过极高的惩罚系数来"劝退"。')

h(doc, '3.6 MPC 主仿真改进：mpc_sim（第 130-291 行）', level=2)
p(doc, '优化版的 mpc_sim 在架构上同基础版一致（网格搜索 + 滚动时域），但有三处关键改进。')

h(doc, '3.6.1 改进点 1：FC 功率变化惩罚（第 194 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['185', 'p_fc_prev = P_fc[k-1] if k > 0 else np.clip(P_load[k], PFC_MIN, PFC_MAX)',
         '记录上一步最优 FC 功率。对于 k=0，用当前负载功率作为初始值（默认 FC 跟随负载）。'],
        ['194', 'J_total += w_pfc_slew * (p_fc_cand - p_fc_prev) ** 2',
         'FC 功率变化惩罚：\n'
         '  w_pfc_slew=0.001 看似很小，但 (ΔP_fc)² 项在跳变大时显著。\n'
         '  假设 P_fc 从 10kW 跳到 50kW：Δ=40，平方=1600，代价=1.6\n'
         '  而此时氢耗约 0.03g/s × 1s = 0.03\n'
         '  1.6 相对于 0.03 是非常大的代价！\n'
         '物理意义：抑制 FC 功率的剧烈波动，保护燃料电池膜电极寿命。'],
    ]
)

h(doc, '3.6.2 改进点 2：不可行候选处理（第 210-213 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['210-213', 'soc_pred_next = mpc_step_soc(soc_pred, p_fc_cand, p_load_i)\n'
         'if soc_pred_next is None:\n    J_total = np.inf\n    break',
         '当候选控制量在某步导致物理不可行（电池功率超过极限、SOC 越界），\n'
         '直接将该候选的代价置为 ∞ 并跳出仿真循环。\n'
         '与基础版的对比：基础版用 clip 掩盖不可行 → 可能导致选中"表面上可行"但实际不安全的控制。\n'
         '优化版直接剔除 → 更安全、更诚实。'],
    ]
)

h(doc, '3.6.3 改进点 3：SOC 惩罚分离为函数调用（第 217-231 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['217-218', 'is_terminal = i == horizon - 1\nis_route_end = k + i + 1 >= N',
         '两个标志位的计算：\n  is_terminal：当前仿真步是否为滚动窗口的最后一步\n  is_route_end：当前仿真步是否到达工况的真正终点'],
        ['219-231', 'J_total += soc_tracking_penalty(\n    soc_pred,\n    is_terminal=is_terminal,\n    is_route_end=is_route_end,\n    ...\n)',
         '调用独立的 SOC 惩罚函数（第 3.5 节详述）。\n'
         '相比基础版第 155-163 行的内联代码，这种分离：\n'
         '  (a) 提高了代码可读性\n'
         '  (b) 便于独立测试 SOC 惩罚逻辑\n'
         '  (c) 参数集中管理，便于调参'],
    ]
)

h(doc, '3.6.4 改进点 4：后备策略——无可行候选时的容错（第 238-247 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['238-247', 'if best_j is None:\n'
         '    one_step_feasible = []\n'
         '    for j, p_fc_cand in enumerate(PFC_GRID):\n'
         '        soc_next = mpc_step_soc(soc_k, p_fc_cand, P_load[k])\n'
         '        if soc_next is not None:\n'
         '            one_step_feasible.append((abs(soc_next - soc_ref), j))\n'
         '    if one_step_feasible:\n'
         '        best_j = min(one_step_feasible)[1]\n'
         '    else:\n'
         '        best_j = int(np.argmin(np.abs(PFC_GRID - np.clip(P_load[k], PFC_MIN, PFC_MAX))))',
         '容错后备策略（fallback）：\n'
         '  场景：如果所有 60 个候选在 N_p 步仿真中都因物理不可行被剔除\n'
         '    （best_j 保持 None），说明当前系统状态（SOC_k、P_load_k）非常不利。\n'
         '  此时改用单步策略：\n'
         '    1. 对每个候选 P_fc，只仿真 1 步\n'
         '    2. 选使 |SOC_next - SOC_ref| 最小的候（平衡 SOC 优先）\n'
         '    3. 如果单步也全部不可行，选最接近 P_load 的 P_fc\n'
         '  这种分级容错确保控制器在极端情况下也不会崩溃。'],
    ]
)

key_point(doc, '后备容错（fallback）是工程级 MPC 的重要特征——仿真环境下不太会出现，但实际系统中电池电压异常、传感器噪声等都可能导致正常求解失败。')

h(doc, '3.6.5 改进点 5：SOC 更新安全处理（第 252-253 行）', level=3)

tbl2(doc, ['行号', '基础版', '优化版'],
    [
        ['252-253', 'SOC[k+1] = mpc_step_soc(...)\n    直接使用返回值',
         'SOC[k+1] = soc_k if soc_next is None else soc_next\n    （保护性赋值）'],
    ]
)

p(doc, '如果 mpc_step_soc 返回 None（物理不可行），基础版会直接将 None 赋值给 SOC，\n'
       '导致后续步骤的数值错误。优化版在 SOC 不可用时保持当前值，增加了一层防御。')

h(doc, '3.6.6 等效氢耗记录（第 258-262 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['259-260', 'raw_h2_kg = np.cumsum(m_H2)[-1] / 1000\nh2_eq_kg = soc_equivalent_h2(..., SOC[-1], ...)',
         '记录原始氢耗和 SOC 修正等效氢耗两个指标，使不同方法的比较更加公平。'],
        ['261-262', 'print(f\'...H2_raw={raw_h2_kg:.4f} kg, SOC_end={SOC[-1]:.3f}, H2_eq={h2_eq_kg:.4f} kg\')',
         '输出行同时显示两个氢耗指标。'],
    ]
)

h(doc, '3.6.7 返回结构扩展（第 268-291 行）', level=3)

tbl2(doc, ['行号', '字段', '说明'],
    [
        ['276-277', 'H2_raw_kg\nH2_eq_kg', '新增的原始氢耗和等效氢耗，方便上层调用直接获取这两个关键指标。'],
        ['279-291', 'config: { N_p, s_factor, w_soc, beta_term, soc_deadband, soc_soft_min, w_soc_low, soc_final_tol, w_final_soc, w_pfc_slew }',
         '将 MPC 的全部配置参数嵌入返回结果中。\n'
         '设计动机：便于事后追溯——看到 CSV 结果时能直接知道使用了哪些参数，\n'
         '而不用回头找脚本中的参数定义。这是可复现性（reproducibility）的良好工程实践。'],
    ]
)

# ── 第294-330行：N_p 扫描改进 ──
h(doc, '3.7 N_p 扫描改进：mpc_n_p_scan（第 294-330 行）', level=2)

tbl2(doc, ['行号', '基础版', '优化版', '改进'],
    [
        ['297', 'def mpc_n_p_scan(P_load, N_p_values=None, SOC_0=0.6):',
         'def mpc_n_p_scan(P_load, N_p_values=None, SOC_0=0.6, **mpc_kwargs):',
         '优化版加 **mpc_kwargs，传入所有 MPC 超参数（s_factor, w_soc 等），\n'
         '使 N_p 扫描可以用与主仿真完全相同的超参数集。'],
        ['319', 'res = mpc_sim(...) （无额外参数）',
         'res = mpc_sim(..., **mpc_kwargs)',
         '传递超参数到 mpc_sim，保证扫描结果的一致性。'],
        ['322-324', '仅记录 H2_kg 和 SOC_end',
         '额外记录 H2_eq_kg',
         'N_p 扫描结果中也包含等效氢耗列，便于综合评估。'],
    ]
)

# ── 第430-487行：四方法对比绘图 ──
h(doc, '3.8 四方法对比图改进（第 336-428 行）', level=2)
p(doc, 'plot_four_way 基本与基础版相同，仅文件命名有变化：', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['425', 'f\'FourWay_compare_optimized_{cycle_name}.png\'',
         '优化版的结果文件添加 _optimized 后缀，与基础版结果区分。'],
    ]
)

# ── 第434-487行：指标打印改进 ──
h(doc, '3.9 指标打印改进：print_four_way_metrics（第 434-487 行）', level=2)

tbl2(doc, ['行号', '代码/字段', '说明'],
    [
        ['449', 'mpc_SOC_end = mpc_result.get(\'SOC_end\', mpc_result[\'SOC\'][-1])',
         '兼容性处理：优化版返回的 dict 有 SOC_end 字段，基础版可能没有。'
         '使用 get() 保证函数可以复用两种版本的返回结果。'],
        ['450-453', 'rule_H2_eq = soc_equivalent_h2(rule_H2, rule_SOC_end)\n...\nmpc_H2_eq = ...',
         '新增 SOC 修正等效氢耗的计算。\n'
         '因为 Rule 和 DP 没有返回 H2_eq_kg，需要在此处计算。\n'
         'MPC 优化版虽然已经算了 H2_eq_kg，但此处也可以从原始指标重算，\n'
         '保证计算一致。'],
        ['464-465', '(\'SOC修正氢耗 (kg)\', ...)',
         '新的指标行：SOC 修正氢耗——公平比较的氢耗指标。\n'
         '在优化版报告中这是必输指标。'],
        ['483-486', 'print(\'  相对 DP 的 SOC 修正氢耗差距:\')\nprint(f\'    MPC:   {mpc_H2_eq-dp_H2_eq)/dp_H2_eq*100:+.1f}%\')',
         '增加了 SOC 修正氢耗的差距分析。\n'
         '相比原始氢耗差距，修正氢耗差距更能反映算法本身的优劣\n'
         '（排除了 SOC 透支的影响）。'],
    ]
)

# ── 第493-524行 ──
h(doc, '3.10 N_p 敏感性绘图改进（第 493-524 行）', level=2)

tbl2(doc, ['行号', '代码/改进', '说明'],
    [
        ['501-503', 'ax1.plot(np_df[\'N_p\'], np_df[\'H2_kg\'], \'ro-\', label=\'MPC raw\')\nif \'H2_eq_kg\' in np_df.columns:\n    ax1.plot(np_df[\'N_p\'], np_df[\'H2_eq_kg\'], \'mo--\', label=\'MPC SOC-corrected\')',
         '新增等效氢耗曲线（紫色虚线），与原始氢耗曲线（红色实线）对比。\n'
         '两者之间的差距反映了 SOC 透支/节省的程度。'],
        ['521', 'f\'MPC_np_sensitivity_optimized_{cycle_name}.png\'',
         '文件名添加 _optimized 后缀。'],
    ]
)

# ── 第527-718行：主程序改进 ──
h(doc, '3.11 主程序改进：main（第 527-718 行）', level=2)

h(doc, '3.11.1 参数扩展（第 531-545 行）', level=3)
p(doc, '相比基础版的 5 个参数，优化版主程序扩展到了 12 个：', bold=True)

tbl2(doc, ['参数', '默认值', '说明'],
    [
        ['--cycle', 'wltc', '工况选择（同基础版）'],
        ['--np', '50', '预测时域（同基础版）'],
        ['--s-factor', '130.0', '🆕 等效因子命令行可调'],
        ['--w-soc', '1200', '🆕 SOC 惩罚权重可调'],
        ['--beta-term', '5000', '🆕 终端惩罚系数可调'],
        ['--soc-soft-min', '0.57', '🆕 SOC 软下限可调'],
        ['--w-soc-low', '20000', '🆕 低 SOC 惩罚可调'],
        ['--soc-final-tol', '0.01', '🆕 终点容差可调'],
        ['--w-final-soc', '80000', '🆕 终点惩罚可调'],
        ['--w-pfc-slew', '0.001', '🆕 FC 变化惩罚可调'],
        ['--scan', 'False', 'N_p 扫描（同基础版）'],
        ['--compare', 'False', '四方法对比（同基础版）'],
    ]
)

note(doc, '12 个参数使优化版可以灵活调参而无需修改代码，支持批量的超参数扫描实验。', label='🔧 工程化改进')

h(doc, '3.11.2 参数传递机制（第 549-558 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['549-558', 'mpc_kwargs = {\n    \'s_factor\': args.s_factor,\n    \'w_soc\': args.w_soc,\n    ...\n}',
         '将所有 MPC 超参数打包为字典，通过 **mpc_kwargs 传递给 mpc_sim。\n'
         '设计优势：\n'
         '  • 新增参数只需在 argparse 和 mpc_kwargs 中各加一行\n'
         '  • mpc_sim 无需修改函数签名即可接收新参数（靠 **kwargs，虽然实际 mpc_sim 有显式参数）\n'
         '  • 便于后续自动化——可以直接从 JSON/YAML 配置文件读参数字典'],
    ]
)

h(doc, '3.11.3 结果保存扩展（第 678-704 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['678-691', 'df_mpc = pd.DataFrame({\n    ..., \'H2_eq_kg\': mpc_result[\'H2_eq_kg\'],\n})\n...\ndf_mpc.to_csv(...)',
         'CSV 保存增加 H2_eq_kg 列，方便后续分析。'],
        ['693-704', 'summary = {\n    \'cycle\': cycle,\n    \'N_p\': n_p,\n    \'H2_raw_kg\': ...,\n    \'SOC_end\': ...,\n    \'SOC_delta_ref_minus_end\': SOC_REF - mpc_result[\'SOC_end\'],\n    \'H2_eq_kg\': ...,\n    **mpc_result[\'config\'],\n}\nsummary_path = f\'mpc_ems_optimized_{cycle}_np{n_p}_summary.csv\'',
         '🆕 独立的汇总文件（_summary.csv）：\n'
         '  • 包含所有关键指标：原始氢耗、终点 SOC、SOC 偏差、等效氢耗\n'
         '  • 通过 **mpc_result[\'config\'] 自动带上所有超参数\n'
         '  • 每行一个实验结果，方便多组实验横向对比\n'
         '  • 这对批处理/超参数扫描实验非常有用'],
    ]
)

key_point(doc, '优化版主程序新增的 _summary.csv 文件对超参数搜索实验至关重要——它让"不同参数下的结果对比"不再需要手动拼接数据。')

doc.add_page_break()

# ========================================================================
# 第四部分：基础版 vs 优化版 全面对比
# ========================================================================
h(doc, '第四部分：基础版 vs 优化版 — 全面对比', level=1)

h(doc, '4.1 改进总览', level=2)

tbl(doc, ['维度', '基础版 (mpc_ems.py)', '优化版 (mpc_ems_optimized.py)', '改进说明'],
    [
        ['代码行数', '571 行', '719 行', '+148 行（+26%）'],
        ['参数数量', '5 个', '12 个', '+7 个新参数'],
        ['SOC 惩罚', 'w_soc=500, 无死区', 'w_soc=1200, deadband=0.015', '更强、更平滑'],
        ['终端惩罚', 'beta=1000, 后30%生效', 'beta=5000, 可调', '约束更强、更灵活'],
        ['SOC 软下限', '无', '0.57 + w=20000', '防 SOC 过度衰减'],
        ['终点欠差罚', '无', 'tol=0.01 + w=80000', '强制 ends near SOC_ref'],
        ['FC 功率跳变罚', '无', 'w=0.001 + 二次型', '保护 FC 寿命'],
        ['不可行处理', 'clip 保持 SOC', 'return None + 后备', '更安全、更诚实'],
        ['等效氢耗指标', '无', 'soc_equivalent_h2()', '公平比较基准'],
        ['配置追踪', '无', '返回 config dict', '可复现性提升'],
        ['汇总文件', '仅 CSV', 'CSV + _summary.csv', '实验管理友好'],
    ]
)

h(doc, '4.2 性能预期差异', level=2)

p(doc, '优化版在以下方面预期优于基础版：', bold=True)
p(doc, '  ✅ SOC 维持更精准（终点 SOC 更接近 0.6）')
p(doc, '  ✅ FC 功率波动更小（变化率惩罚抑制快速切换）')
p(doc, '  ✅ 极端情况下更有韧性（后备策略 + 不可行剔除）')
p(doc, '  ✅ 结果可复现性更高（配置参数嵌入输出）')
p(doc, '  ⚠️ 但原始氢耗可能略微升高（因为 SOC 软下限阻止了"透支电池"策略）')

p(doc, '这正是"原始氢耗 vs 等效氢耗"对比的价值——优化版可能原始氢耗高一点，')
p(doc, '但等效氢耗更低，因为终点的 SOC_equity 被纳入了评估。', bold=True)

h(doc, '4.3 使用场景建议', level=2)

tbl2(doc, ['场景', '推荐版本', '理由'],
    [
        ['MPC 原理教学/入门', '基础版 (mpc_ems.py)',
         '代码更简洁，核心 MPC 逻辑不被大量工程细节掩盖'],
        ['EMS 性能优化/实际应用', '优化版 (mpc_ems_optimized.py)',
         'SOC 维持更好、FC 运行更平稳、容错更强'],
        ['超大规模参数扫描实验', '优化版',
         '支持 CLI 参数调优、自动记录配置、_summary.csv 便于批量对比'],
        ['阅读/理解 MPC 核心思想', '基础版',
         '网格搜索 + receding horizon 的架构一目了然'],
    ]
)

doc.add_page_break()

# ========================================================================
# 第五部分：核心算法图解
# ========================================================================
h(doc, '第五部分：核心算法原理解析', level=1)

h(doc, '5.1 Receding Horizon（滚动时域）流程', level=2)

code_block(doc, '''┌─────────────────────────────────────────────────────────────┐
│                   Receding Horizon MPC                         │
├─────────────────────────────────────────────────────────────┤
│  时刻 k:                                                      │
│  ┌─ Future Prediction (N_p steps) ─────────────────────────┐ │
│  │  P_load_pred = [P_k, P_{k+1}, ..., P_{k+N_p-1}]         │ │
│  │                                                          │ │
│  │  Grid Search:                                            │ │
│  │  for each p_fc ∈ PFC_GRID:                              │ │
│  │    ├─ Simulate horizon steps                             │ │
│  │    ├─ Compute cost = Σ H₂ + Σ s*|P_bat| + SOC_penalty   │ │
│  │    └─ Save candidate cost                                │ │
│  │  ──────────────────────────────────────────────────────── │ │
│  │  Select p_fc* = argmin(cost)                             │ │
│  └──────────────────────────────────────────────────────────┘ │
│  │                                                            │
│  │  Apply p_fc*(k) to system (第1步)                         │
│  │  SOC_{k+1} = f(SOC_k, p_fc*, P_load_k)                   │
│  │                                                            │
│  └───────────→  k = k+1, repeat ───────────────────────────┘''')

h(doc, '5.2 代价函数构成示意', level=2)

p(doc, '基础版代价函数：')
code_block(doc, '''J_k = Σ_{i=0}^{horizon-1} [
     H₂(p_fc)*Δt                  ← 氢耗
     + S_MPC * |P_bat_i|/3600*Δt  ← 等效电池能量
     + W_SOC * dev_i² * Δt        ← SOC 维持惩罚 (|dev|>0.05时)
] + β_term * dev_N²                ← 终端惩罚 (后30%阶段)''')

p(doc, '优化版代价函数：')
code_block(doc, '''J_k = Σ_{i=0}^{horizon-1} [
     H₂(p_fc)*Δt                  ← 氢耗
     + S_MPC * |P_bat_i|/3600*Δt  ← 等效电池能量
     + soc_tracking_penalty(       ← SOC 综合惩罚
         dev, low_gap, is_terminal, is_route_end, ...)
] + w_pfc_slew * (p_fc - p_fc_prev)²  ← FC 功率变化惩罚''')

h(doc, '5.3 优化版 SOC 惩罚函数结构', level=2)

code_block(doc, '''soc_tracking_penalty(soc, is_terminal, is_route_end):
┌─ penalty = 0
├─ excess = max(|soc - soc_ref| - deadband(0.015), 0)
├─ penalty += W_SOC(1200) * excess² * Δt    ← 常规SOC跟踪
├─ low_gap = max(soft_min(0.57) - soc, 0)
├─ penalty += W_SOC_LOW(20000) * low_gap² * Δt  ← SOC软下限
├─ if is_terminal:
│    penalty += BETA_TERM(5000) * excess²    ← 滚动终端惩罚
├─ if is_route_end:
│    penalty += W_FINAL_SOC(80000) * shortfall²  ← 终点欠差罚
└─ return penalty''')

h(doc, '5.4 等效氢耗修正原理', level=2)

code_block(doc, '''公平比较公式：
  H₂_eq = H₂_raw + s_factor * ΔE_bat / 1000

其中：
  ΔE_bat [kWh] = Q_bat[Ah] * OCV_mean[V] * (SOC_ref - SOC_end) / 1000
  s_factor [g/kWh] = 等效因子（S_MPC = 130 g/kWh）

物理含义：
  SOC_end < SOC_ref → ΔE_bat > 0 → 多用了电池能量 → H₂_eq > H₂_raw
  SOC_end > SOC_ref → ΔE_bat < 0 → 节省了电池能量 → H₂_eq < H₂_raw''')

doc.add_page_break()

# ========================================================================
# 第六部分：参考文献与扩展阅读
# ========================================================================
h(doc, '第六部分：参考文献与扩展阅读', level=1)

h(doc, '6.1 MPC 理论基础', level=2)
p(doc, '  [1] Camacho, E. F., & Bordons, C. (2013). Model Predictive Control (2nd ed.). Springer.')
p(doc, '  [2] Mayne, D. Q., Rawlings, J. B., Rao, C. V., & Scokaert, P. O. M. (2000). '
       'Constrained model predictive control: Stability and optimality. Automatica, 36(6), 789-814.')
p(doc, '  [3] Rawlings, J. B., Mayne, D. Q., & Diehl, M. (2017). Model Predictive Control: '
       'Theory, Computation, and Design (2nd ed.). Nob Hill Publishing.')

h(doc, '6.2 MPC 在 EMS 中的应用', level=2)
p(doc, '  [4] Bordons, C., del Real, A. J., & Arce, A. (2011). Model Predictive Control of '
       'Powersplit Hybrid Electric Vehicles. In: Proceedings of the 18th IFAC World Congress.')
p(doc, '  [5] Hemi, H., Ghouili, J., & Cheriti, A. (2014). A real time fuzzy logic power '
       'management strategy for a fuel cell vehicle. Energy Conversion and Management, 80, 163-173.')
p(doc, '  [6] Feroldi, D., & Carignano, M. (2016). Sizing for fuel cell hybrid vehicles: '
       'A comparison study. International Journal of Hydrogen Energy, 41(32), 14247-14258.')

h(doc, '6.3 本工程配套文档', level=2)
p(doc, '  • docs/MPC_原理与实现_第7周学习笔记.docx — MPC 理论知识')
p(doc, '  • docs/MPC_第7周学习报告.docx — 仿真实验报告')
p(doc, '  • docs/MPC_精细化原理解析.docx — 更详细的 MPC 原理解析')

p(doc, '', size=6)
p(doc, '—— 文档完 ——', bold=True, size=11, color=RGBColor(100, 100, 100))
p(doc, '生成于 2026-07-07 | 基于 mpc_ems.py (v1) 和 mpc_ems_optimized.py (v2)', size=9, color=RGBColor(140, 140, 140))

# ── 保存 ──
output_path = os.path.join(DOCS_DIR, 'MPC_EMS_逐行代码原理分析.docx')
doc.save(output_path)
print(f'[OK] 文档已保存: {output_path}')
