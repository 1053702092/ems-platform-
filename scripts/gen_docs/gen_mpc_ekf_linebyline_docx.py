# -*- coding: utf-8 -*-
"""
生成 mpc_ems_ekf.py 逐行代码原理分析 .docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')

# ── helpers ──
def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def p(doc, text, bold=False, size=10.5, color=None, indent=False):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Cm(0.75)
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def code_block(doc, text, size=8.5):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(50, 50, 50)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def note(doc, text, label="📌 注意"):
    para = doc.add_paragraph()
    run = para.add_run(f'{label}：{text}')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0, 70, 130)
    run.font.italic = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def key_point(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(f'✦ {text}')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(180, 60, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def diff_added(doc, text):
    """新增/改进点高亮"""
    para = doc.add_paragraph()
    run = para.add_run(f'🆕 {text}')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 120, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def tbl(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table

def tbl2(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table


# ================================================================
os.makedirs(DOCS_DIR, exist_ok=True)
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── 封面 ──
p(doc, '', size=14)
h(doc, 'mpc_ems_ekf.py 逐行代码原理分析', level=0)
p(doc, 'MPC + EKF/AEKF SOC 估计 — 燃料电池 EMS 能量管理（v3.0）', bold=True, size=14, color=RGBColor(0, 70, 130))

p(doc, '', size=6)
p(doc, '文件路径：scripts/mpc_ems_ekf.py', bold=True, size=11)
p(doc, '总行数：1017 行 | 在优化版（719 行）基础上 +298 行', size=10)
p(doc, '', size=4)
p(doc, '继承关系：mpc_ems.py (571行) → mpc_ems_optimized.py (719行) → mpc_ems_ekf.py (1017行)', size=9, color=RGBColor(100, 100, 100))
p(doc, '生成日期：2026-07-07', size=9, color=RGBColor(120, 120, 120))

doc.add_page_break()

# ========================================================================
# 第一部分：概览
# ========================================================================
h(doc, '第一部分：文件概览', level=1)

h(doc, '1.1 文件定位', level=2)
p(doc, 'mpc_ems_ekf.py 是 EMS 项目中 MPC 控制器的第三代实现：')
p(doc, '  v1 mpc_ems.py           — MPC 原理演示（基础网格搜索 + receding horizon）')
p(doc, '  v2 mpc_ems_optimized.py  — MPC 工程优化（SOC 软约束、功率变化惩罚、后备容错）')
p(doc, '  v3 mpc_ems_ekf.py        — MPC+EKF（在 v2 基础上集成 EKF/AEKF SOC 状态估计）', bold=True)

h(doc, '1.2 核心改进', level=2)
p(doc, '相比 v2 优化版的关键新增：')
diff_added(doc, 'SOC 估计器策略模式架构（SOCEstimator 基类 + 三种子类）')
diff_added(doc, 'EKF 扩展卡尔曼滤波 SOC 估计（电流 + 电压融合）')
diff_added(doc, 'AEKF 自适应 EKF（滑动窗口自适应 R/Q）')
diff_added(doc, '开环/估计/真实 三路 SOC 并行追踪（用于对比评估）')
diff_added(doc, '电流传感器偏置和噪声模拟（仿真实际 BMS 环境）')
diff_added(doc, 'SOC 估计误差指标（RMSE）自动计算')
diff_added(doc, 'SOC 估计对比图（直观展示 EKF 的抗漂移效果）')
diff_added(doc, '随机种子控制（保证可复现性）')

h(doc, '1.3 文件结构总览', level=2)

tbl(doc, ['章节', '行号范围', '内容', '相比 v2 的变化'],
    [
        ['文件头 & 导入', '1-40', '文档字符串、import、路径配置', '无变化'],
        ['MPC 参数', '42-69', '超参数 + EKF 参数 + 传感器参数', '🆕 新增 8 个 EKF/传感器参数'],
        ['电池辅助函数', '72-96', 'lookup_ocv, lookup_docv_dsoc, battery_current', '🆕 新增（EKF 依赖）'],
        ['SOC 估计器类', '99-226', '基类 + OpenLoop + EKF + AEKF + 工厂方法', '🆕 全新（5 个类/函数）'],
        ['公用函数', '229-291', '仿真电压、等效氢耗、SOC 惩罚、状态转移', '4 个复用 v2，1 个新增'],
        ['MPC 主仿真', '294-493', 'mpc_sim() — 核心算法', '🔄 重写 SOC 处理逻辑'],
        ['N_p 扫描', '496-525', 'mpc_n_p_scan()', '🔄 增加 SOC_RMSE 输出'],
        ['绘图函数', '528-714', 'plot_four_way / plot_soc_estimation / plot_np_sensitivity', '🆕 plot_soc_estimation 新增'],
        ['指标打印', '717-779', 'print_four_way_metrics', '🔄 增加 SOC RMSE 行'],
        ['主程序', '782-1017', 'main()', '🔄 增加 EKF 参数解析'],
    ]
)

doc.add_page_break()

# ========================================================================
# 第二部分：逐行分析（与 v2 相同的部分略过）
# ========================================================================
h(doc, '第二部分：逐行分析', level=1)
p(doc, '本文档重点分析 mpc_ems_ekf.py 中与 v2 优化版不同的部分。完全相同的代码（如 soc_tracking_penalty、mpc_step_soc、plot_four_way、主流程控制）请参考《MPC_EMS_逐行代码原理分析.docx》第三部分。')

# ── 第1-69行：文件头与参数 ──
h(doc, '2.1 文件头与导入（第 1-40 行）', level=2)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['1-20', '文档字符串', '核心改进点（4 条）、\n用法示例（6 种调用方式）\n清晰标注了与 v2 的继承关系。'],
        ['22-32', 'import / sys.path', '与 v2 相同。'],
        ['34-40', 'from day8_dp_ems import ...', '复用 day8 全部核心模型。\n注意：SOC_BP, OCV_LU, Q_BAT, R_INT 也在此导入——\n它们被 EKF 的 OCV 查表函数需要。'],
    ]
)

# ── 第42-69行：参数定义 ──
h(doc, '2.2 MPC 参数与 EKF/传感器参数（第 42-69 行）', level=2)

tbl2(doc, ['行号', '参数', '值', '说明'],
    [
        ['45-56', 'N_P_DEFAULT 等 11 个参数', '同 v2 优化版', '与 mpc_ems_optimized.py 完全一致。'],
        ['59', 'Q_EKF_DEFAULT = 5e-5', '5×10⁻⁵', '🆕 EKF 过程噪声。\n'
         '物理含义：安时积分模型的不确定性。\n'
         '数值量级对应 SOC 每步变化的标准差 ≈ √(5e-5) ≈ 0.007。\n'
         '即每步安时积分大约有 0.7% SOC 的不确定度——\n'
         '对应约 2A 电流偏置的漂移速度。'],
        ['60', 'R_EKF_DEFAULT = 0.03', '0.03 V²', '🆕 EKF 测量噪声。\n'
         '物理含义：端电压测量的不确定性。\n'
         '数值对应 OCV 查表误差 + 传感器噪声 ≈ √0.03 ≈ 0.17V。\n'
         '比真实传感器噪声（~0.1V）略大，包含 OCV 曲线建模误差。'],
        ['61', 'P0_EKF_DEFAULT = 0.1', '0.1', '🆕 EKF 初始协方差。\n'
         '物理含义：初始 SOC 估计的不确定性。\n'
         '对应初始 SOC 标准差 ≈ √0.1 ≈ 0.32。\n'
         '即初始认为 SOC "在 0.6 ± 32%" 范围内——\n'
         '设得较大表示"开始时对 SOC 不是很确定"。'],
        ['64', 'CURRENT_BIAS_DEFAULT', '2.0 A', '🆕 模拟电流传感器偏置。\n'
         '真实 BMS 中常见的偏置范围：0.5~5A。\n'
         '2A 的偏置在 1800s WLTC 中会导致开环 SOC 漂移约 0.02。'],
        ['65', 'CURRENT_NOISE_STD', '0.5 A', '🆕 电流测量噪声标准差。\n'
         '模拟高频噪声，不影响长期漂移但影响 EKF 的瞬态行为。'],
        ['66', 'VOLTAGE_NOISE_STD', '0.1 V', '🆕 电压测量噪声标准差。\n'
         '典型 BMS 电压传感器精度在 10~50mV，\n'
         '这里设 0.1V 以涵盖 OCV 曲线查表误差。'],
    ]
)

key_point(doc, 'EKF 的三个参数 Q_EKF、R_EKF、P0_EKF 决定 EKF 的"性格"：Q/R 比值越大越相信电压测量，越小越相信安时积分。')

# ── 第72-96行：电池辅助函数 ──
h(doc, '2.3 电池参数辅助函数（第 72-96 行）', level=2)

p(doc, '这三个辅助函数是 EKF 的数学基础，独立抽取出来使 EKF 代码更简洁。', bold=True)

tbl2(doc, ['行号', '函数', '说明'],
    [
        ['75-77', 'lookup_ocv(soc)', 'OCV 查表。\n用 np.interp 对 SOC_BP → OCV_LU 做线性插值。\n'
         '与 day8_dp_ems 中的 OCV 曲线相同，保证一致性。\n'
         '供 EKF 的观测模型使用：v_pred = OCV(SOC_pred)。'],
        ['80-85', 'lookup_docv_dsoc(soc)', 'OCV 曲线斜率（数值微分）。\n'
         '用中心差分计算 d(Voc)/d(SOC)：\n'
         '  (OCV(soc+δ) - OCV(soc-δ)) / (2δ)\n'
         'δ=1e-6 是微小扰动。\n'
         '供 EKF 的观测雅可比 H 使用。\n\n'
         '为什么需要这个函数？\n'
         'EKF 的 Update 步需要 H = ∂h/∂x = ∂OCV/∂SOC。\n'
         'OCV 曲线是非线性的，不同 SOC 区间斜率不同：\n'
         '  • SOC 0.1-0.3：斜率大（~100V/SOC）→ 电压对 SOC 敏感\n'
         '  • SOC 0.4-0.7：斜率平缓（~30V/SOC）→ 电压对 SOC 不敏感\n'
         '  • SOC 0.8-1.0：斜率大 → 又变得敏感\n\n'
         'H 值大 → K 大 → 更信任电压测量（OCV 斜率大区域修正快）\n'
         'H 值小 → K 小 → 更信任安时积分（OCV 平缓区域依赖电流）\n'
         '这个自动调节是 EKF 的核心优势之一！'],
        ['88-96', 'battery_current(soc, p_bat)', '根据 SOC 和电池功率计算电流。\n'
         '与 mpc_step_soc 内部逻辑相同，但：\n'
         '  • 返回浮点数而非 array\n'
         '  • delta < 0 时返回 0 而非 clip 保持——\n'
         '    在真实 SOC 演化中不会遇到物理不可行\n'
         '    （因为 MPC 已经保证了可行性）\n'
         '供仿真环境中的"真实 SOC 演化"使用。'],
    ]
)

# ── 第99-226行：SOC 估计器 ──
h(doc, '2.4 SOC 估计器体系（第 99-226 行）', level=2)
p(doc, '这是 mpc_ems_ekf.py 最核心的新增代码。采用策略模式（Strategy Pattern）设计。', bold=True)

h(doc, '2.4.1 基类 SOCEstimator（第 102-109 行）', level=3)
tbl2(doc, ['行号', '代码', '说明'],
    [
        ['102-105', 'class SOCEstimator:\n    def __init__(self, x0=0.6):\n        self.x = float(x0)',
         '基类定义统一接口。\n'
         '所有子类共享 x（当前 SOC 估计值）属性。\n'
         'float() 强制转换保证类型一致性。'],
        ['107-109', 'def step(self, i_meas, v_t_meas, dt=DT):\n    raise NotImplementedError',
         '抽象接口：输入测量电流和电压，输出 SOC 估计。\n'
         '子类必须实现该函数。\n\n'
         '设计模式的意图：\n'
         '  MPC 主循环可以统一调用 estimator.step()，\n'
         '  无需关心具体是哪种估计方法。'],
    ]
)

h(doc, '2.4.2 开环估计器 OpenLoopEstimator（第 112-117 行）', level=3)
tbl2(doc, ['行号', '代码', '说明'],
    [
        ['112-113', 'class OpenLoopEstimator(SOCEstimator):\n    """开环安时积分"""',
         '最简单子类——纯开环库仑积分。\n'
         '忽略 v_t_meas 参数（不执行任何修正）。'],
        ['114-117', 'def step(self, i_meas, v_t_meas, dt=DT):\n'
         '    soc_next = self.x - i_meas / (Q_BAT*3600) * dt\n'
         '    self.x = float(np.clip(soc_next, SOC_MIN, SOC_MAX))\n'
         '    return self.x',
         '核心公式：SOC_{k+1} = SOC_k - I/Q_bat × Δt\n'
         '与 mpc_ems.py 中的 mpc_step_soc() 完全相同，\n'
         '但去掉了电池功率→电流的转换步骤，直接接受电流输入。\n\n'
         '性能特征：\n'
         '  ✅ 简单、无参数、计算量为零\n'
         '  ❌ 电流有偏置 → SOC 单调漂移，永不收敛\n'
         '  ❌ 对 v_t_meas 视而不见（浪费了电压信息）'],
    ]
)

h(doc, '2.4.3 EKF 估计器 EKFEstimator（第 120-157 行）', level=3)
p(doc, '这是整个文件的"明星代码"——扩展卡尔曼滤波 SOC 估计的核心实现。', bold=True)

tbl2(doc, ['行号', '代码', '原理分析'],
    [
        ['120-128', 'class EKFEstimator(SOCEstimator):\n    """EKF SOC 估计器"""',
         '类的文档字符串清晰地标注了：\n'
         '  • 状态: SOC (一维)\n'
         '  • 过程: SOC_{k+1} = SOC_k - I/Q×dt  （安时积分，恰好线性）\n'
         '  • 观测: V_t = OCV(SOC)            （非线性，需要线性化）\n'
         '  • 两阶段: Predict → Update'],
        ['129-136', 'def __init__(self, x0=0.6, P0=0.1, Q=5e-5, R=0.03):\n'
         '    self.P = float(P0)   # 估计协方差\n'
         '    self.Q = float(Q)    # 过程噪声\n'
         '    self.R = float(R)    # 测量噪声\n'
         '    self.innov = 0.0     # 新息日志\n'
         '    self.K_gain = 0.0    # 卡尔曼增益日志',
         '初始化 EKF 状态：\n'
         '  • P：估计协方差——度量"我有多不确定自己当前的 SOC 估计"'
         '    - 初始 P=0.1：标准差约 0.32，相当不确定\n'
         '    - 每次 Predict 步增大（添加 Q），每次 Update 步减小（添加信息）\n'
         '  • Q 和 R 是固定的（AEKF 才会自适应调整）\n'
         '  • innov 和 K_gain 记录最近值，仅用于监控和调试'],
        ['138-142', 'def step(self, i_meas, v_t_meas, dt=DT):\n'
         '    # Predict\n'
         '    soc_pred = self.x - i_meas / (Q_BAT*3600) * dt\n'
         '    F = 1.0\n'
         '    P_pred = self.P + self.Q',
         'Predict 阶段（时间更新）：\n'
         '  ① 状态预测：安时积分\n'
         '    SOC_pred = SOC_{k-1} - I_meas/Q × Δt\n'
         '    📌 注意：这里的电流是测量值 i_meas（含偏置）\n'
         '          ——EKF 不要求电流无偏，因为后续会用电压修正\n\n'
         '  ② 雅可比 F = ∂f/∂SOC = 1\n'
         '    状态方程 f(SOC) = SOC - I(SOC)/Q×Δt\n'
         '    ∂f/∂SOC = 1 - (∂I/∂SOC)/Q×Δt\n'
         '    其中 ∂I/∂SOC ≈ 0（电流主要由 P_bat 决定而非 SOC）\n'
         '    因此 F ≈ 1，这是合理的简化\n\n'
         '  ③ 协方差预测\n'
         '    P_pred = F × P × F^T + Q = P + Q\n'
         '    （因为 F=1，P_pred = P + Q）\n'
         '    协方差增大——因为加入了过程噪声'],
        ['144-151', '# Update\n'
         '    v_pred = lookup_ocv(soc_pred)\n'
         '    y = v_t_meas - v_pred\n'
         '    H = lookup_docv_dsoc(soc_pred)\n'
         '    S = H * P_pred * H + self.R\n'
         '    K = P_pred * H / S\n'
         '    x_est = soc_pred + K * y\n'
         '    P_est = (1 - K * H) * P_pred',
         'Update 阶段（测量更新）：' + '\n' +
         '  ④ 预测电压' + '\n' +
         '    v_pred = OCV(SOC_pred)' + '\n' +
         '    用当前 SOC 预测查 OCV 曲线得到预期端电压' + '\n' +
         '' + '\n' +
         '  ⑤ 新息（残差）' + '\n' +
         '    y = V_t_meas - OCV(SOC_pred)' + '\n' +
         '    新息 > 0 → 实测电压高于预测 → SOC_pred 偏低 → 需要向上修正' + '\n' +
         '    新息 < 0 → 实测电压低于预测 → SOC_pred 偏高 → 需要向下修正' + '\n' +
         '' + '\n' +
         '  ⑥ 观测雅可比' + '\n' +
         '    H = d(OCV)/d(SOC) 在 soc_pred 处的数值导数' + '\n' +
         '' + '\n' +
         '  ⑦ 新息协方差 S = H² × P_pred + R' + '\n' +
         '    度量"新息中多少来自预测误差，多少来自测量噪声"' + '\n' +
         '' + '\n' +
         '  ⑧ 卡尔曼增益 K = P_pred × H / S' + '\n' +
         '    核心：权衡"信任电流积分多少"和"信任电压多少"' + '\n' +
         '    • K 大 → 新息被充分采纳 → 收敛快但噪声大' + '\n' +
         '    • K 小 → 新息被过滤 → 平滑但收敛慢' + '\n' +
         '' + '\n' +
         '  ⑨ SOC 修正' + '\n' +
         '    SOC_est = SOC_pred + K × y' + '\n' +
         '    这是 EKF 的"魔法"：用电压测量值修正安时积分的漂移' + '\n' +
         '' + '\n' +
         '  ⑩ 协方差更新' + '\n' +
         '    P_est = (1 - K × H) × P_pred' + '\n' +
         '    协方差减小——因为加入了测量信息，不确定度降低'],
        ['153-157', 'self.x = float(np.clip(x_est, SOC_MIN, SOC_MAX))\n'
         'self.P = max(P_est, 1e-8)\n'
         'self.innov = float(y)\n'
         'self.K_gain = float(K)\n'
         'return self.x',
         '更新内部状态：\n'
         '  • SOC 保持在 [SOC_MIN, SOC_MAX] 范围内\n'
         '  • 协方差有下界 1e-8（防止数值下溢 → P=0 → 增益恒为 0）\n'
         '  • 日志记录（供外部监控）\n'
         '  • 返回最新 SOC 估计值'],
    ]
)

key_point(doc, 'EKF 的"魔法"在于第 149 行：K = P_pred * H / S。它让 EKF 在 OCV 斜率大的区间自动信任电压测量，在 OCV 平缓区间自动信任安时积分，且这种权衡是根据当前协方差状态动态调节的。')

h(doc, '2.4.4 AEKF 估计器 AEKFEstimator（第 160-210 行）', level=3)
p(doc, 'AEKF = EKF + 在线自适应 R/Q。核心改进在第 198-206 行。')

tbl2(doc, ['行号', '代码', '原理分析'],
    [
        ['160-168', 'class AEKFEstimator(SOCEstimator):\n    """AEKF 自适应扩展卡尔曼滤波"""',
         '与标准 EKF 的差异：\n'
         '  1. R 和 Q 不是固定值，而是通过滑动窗口新息方差在线更新\n'
         '  2. 新增 innov_buffer 列表和 window 参数\n'
         '  3. Predict 和 Update 步与 EKF 完全相同\n'
         '  4. 自适应更新在 Update 之后执行'],
        ['169-179', '__init__: Q0, R0, window=50\n    innov_buffer = []',
         '初始化：\n'
         '  • Q0/R0 作为自适应算法的初始值\n'
         '  • window=50：滑动窗口大小（50 步，即 50 秒）\n'
         '    窗口越大 → R 估计越平滑 → 自适应越慢；\n'
         '    窗口越小 → R 响应越快 → 但可能被噪声误导'],
        ['181-196', 'Predict + Update 步（同 EKF）',
         '与 EKFEstimator.step() 的前半部分完全相同。\n'
         '代码复用：本可以直接继承 EKFEstimator 但为了清晰独立实现。'],
        ['198-206', '# 自适应更新 R-Q（基于新息滑动窗口）\n'
         'self.innov_buffer.append(y)\n'
         'if len(self.innov_buffer) > self.window:\n'
         '    self.innov_buffer.pop(0)\n\n'
         'if len(self.innov_buffer) >= 10:\n'
         '    innov_var = float(np.var(self.innov_buffer))\n'
         '    self.R = max(innov_var - H*P_pred*H, 0.001)\n'
         '    self.Q = max(K * innov_var * K, 1e-8)',
         '🎯 AEKF 的核心——R/Q 自适应更新：\n\n'
         'R 自适应原理：\n'
         '  理论新息方差 = H×P_pred×H + R\n'
         '  实际新息方差 = var(innov_buffer)\n'
         '  令两者相等解出 R：\n'
         '    R_est = var(innov_buffer) - H×P_pred×H\n'
         '  如果实际新息方差大于预期 → R 偏小（测量噪声被低估）\n'
         '    → 增大 R → EKF 降低对电压测量的信任\n'
         '  max(..., 0.001) 保证 R 不为负且不下界\n\n'
         'Q 自适应原理：\n'
         '  Q_est = K × var(innov_buffer) × K\n'
         '  通过卡尔曼增益反推：如果新息方差大且增益大，\n'
         '    说明过程噪声 Q 被低估了\n\n'
         'AEKF 的工程价值：\n'
         '  当 OCV 曲线斜率变化时（不同 SOC 区间），\n'
         '  理论上最优的 R 会变化。AEKF 自动跟踪这种变化。\n'
         '  同时在传感器噪声变化（温度、老化）时自动调整。'],
    ]
)

note(doc, 'AEKF 的滑动窗口大小（window=50）是一个重要的权衡参数：太小则 R 振荡（被噪声主导），太大则 R 响应迟钝。50 秒的经验值对 WLTC 工况（1800 秒）是合理的。')

h(doc, '2.4.5 工厂方法 build_estimator（第 213-226 行）', level=3)
tbl2(doc, ['行号', '代码', '说明'],
    [
        ['213-226', 'def build_estimator(method, x0=0.6, **kwargs):\n'
         '    if method == \'openloop\': return OpenLoopEstimator(...)\n'
         '    elif method == \'ekf\': return EKFEstimator(...)\n'
         '    elif method == \'aekf\': return AEKFEstimator(...)',
         '工厂方法（Factory Method）设计模式：\n'
         '  输入字符串 → 输出对应估计器实例\n'
         '  **kwargs 传递 EKF 特有的 Q/R/P0 参数（开环估计器忽略它们）'],
    ]
)

# ── 第229-291行：公用函数 ──
h(doc, '2.5 仿真电压、等效氢耗、SOC 惩罚、状态转移（第 229-291 行）', level=2)

tbl2(doc, ['行号', '函数', '说明'],
    [
        ['232-234', 'simulate_voltage(soc_true, noise_std=0.1)', '🆕 新增。仿真 BMS 电压传感器。\n'
         '公式：V_t = OCV(SOC_true) + 高斯噪声\n'
         '注意：不减去 R×I 压降——这是故意的。\n'
         '因为在 EKF 的观测模型中，预测电压 v_pred = OCV(SOC_pred)\n'
         '也没有减去 R×I。保持一致才能保证新息 y 的统计特性正确。\n\n'
         '这意味着模型误差（R×I 被忽略）被归入 R（测量噪声）中——\n'
         '这是常见的工程简化，因为 R×I 压降（~2-5V）远小于 OCV 范围（~300-360V），\n'
         '且主要在重载情况下显著。'],
        ['240-244', 'soc_equivalent_h2(...)', '与 v2 优化版完全相同。'],
        ['250-270', 'soc_tracking_penalty(...)', '与 v2 优化版完全相同。'],
        ['276-291', 'mpc_step_soc(...)', '与 v2 优化版完全相同。'],
    ]
)

# ── 第294-493行：MPC 主仿真 ──
h(doc, '2.6 MPC 主仿真函数 mpc_sim（第 294-493 行）', level=2)
p(doc, '这是从 v2 到 v3 改动最大的部分。核心变化：单一的 SOC 数组拆分为三路并行追踪。', bold=True)

h(doc, '2.6.1 函数签名新增参数（第 297-307 行）', level=3)
diff_added(doc, '8 个新增参数：soc_estimator, current_bias, current_noise_std, voltage_noise_std, ekf_x0, ekf_P0, ekf_Q, ekf_R')

h(doc, '2.6.2 三路 SOC 分配（第 323-337 行）', level=3)
p(doc, '这是 v3 与 v2 最显著的区别之一。', bold=True)

tbl2(doc, ['行号', '数组', '说明'],
    [
        ['326', 'SOC_true = np.zeros(N+1)', '🎯 真实 SOC（无偏库仑积分）。\n'
         '  • 使用无偏电流 i_real 更新\n'
         '  • 仅仿真环境可知，真实车辆中不可观测\n'
         '  • 作为评估 EKF 精度的"黄金标准"'],
        ['327', 'SOC_est_arr = np.zeros(N+1)', '🎯 估计 SOC（EKF/AEKF/开环）。\n'
         '  • 使用有偏/带噪声的传感器数据\n'
         '  • MPC 控制决策基于此值\n'
         '  • 对应真实车辆中的"仪表盘显示 SOC"'],
        ['328', 'SOC_open = np.zeros(N+1)', '🎯 开环安时积分（对比基准）。\n'
         '  • 与 SOC_est 使用相同的传感器数据\n'
         '  • 但不使用电压修正（纯开环）\n'
         '  • 用来量化 EKF 的改进幅度'],
        ['334-337', 'SOC_true[0] = SOC_0\n'
         'x0_ekf = SOC_0 if ekf_x0 is None else ekf_x0\n'
         'SOC_est_arr[0] = x0_ekf\n'
         'SOC_open[0] = x0_ekf',
         '初值设置：\n'
         '  • 真实 SOC 始终从 SOC_0 开始\n'
         '  • 估计 SOC 可从不同值开始（通过 --ekf-x0 设偏来测试收敛性）\n'
         '  • 默认情况三者相同'],
    ]
)

h(doc, '2.6.3 估计器创建（第 339-347 行）', level=3)
tbl2(doc, ['行号', '代码', '说明'],
    [
        ['339-347', 'if soc_estimator == \'openloop\':\n'
         '    estimator = OpenLoopEstimator(...)\n'
         'elif ... == \'ekf\':\n'
         '    estimator = EKFEstimator(...)',
         '根据用户选择的模式创建对应估计器。\n'
         '这里没有使用 build_estimator 工厂方法，而是手动 if-elif——\n'
         '因为需要在创建时传递不同的参数（Q vs Q0）。'],
    ]
)

h(doc, '2.6.4 MPC 内循环（第 354-406 行）', level=3)

tbl2(doc, ['行号', 'vs v2 的变化', '说明'],
    [
        ['355', 'soc_est_k = SOC_est_arr[k]', '🔄 核心变化：MPC 使用估计 SOC 做决策。\n'
         'v2 中 MPC 使用 SOC[k]（开环），\n'
         'v3 中 MPC 使用 SOC_est_arr[k]（EKF 修正后）。\n'
         '当 EKF 准确时，这等价于使用真实 SOC 做决策。'],
        ['357-406', '与 v2 完全相同', '网格搜索、代价函数、SOC 惩罚函数——\n'
         '全部与 mpc_ems_optimized.py 保持一致。'],
    ]
)

h(doc, '2.6.5 后备策略（第 408-418 行）', level=3)
tbl2(doc, ['行号', '说明', '说明'],
    [
        ['408-418', '与 v2 完全相同', '分级容错机制：所有候选不可行时的后备方案。'],
    ]
)

h(doc, '2.6.6 三路 SOC 更新（第 425-438 行）', level=3)
p(doc, '这是 v3 最核心的变化——将 SOC 更新拆分为"真实"、"估计"、"开环"三条路径。', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['425-427', '# 真实 SOC 演化（无偏电流）\n'
         'i_real = battery_current(SOC_true[k], P_bat[k])\n'
         'SOC_true[k+1] = SOC_true[k] - i_real/(Q_BAT*3600)*DT',
         '路径 1：真实 SOC\n'
         '  • 使用无偏电流 i_real\n'
         '  • 电流来自于电池功率和 OCV 的反算（与 mpc_step_soc 一致）\n'
         '  • 这是 "真相"——我们希望 EKF 能准确追踪的轨迹'],
        ['429-431', '# 模拟传感器测量（含偏置和噪声）\n'
         'i_meas_k = i_real + current_bias + current_noise_std * np.random.randn()\n'
         'v_meas_k = simulate_voltage(SOC_true[k], voltage_noise_std)',
         '传感器仿真：\n'
         '  • 在真实电流上叠加偏置和噪声 → i_meas_k\n'
         '  • 在真实 OCV 上叠加噪声 → v_meas_k\n'
         '  • 这模拟了 BMS 传感器的实际输出\n'
         '  • 偏置是确定性误差，噪声是随机误差'],
        ['433-435', '# SOC 估计\n'
         'soc_est_k1 = estimator.step(i_meas_k, v_meas_k)\n'
         'SOC_est_arr[k+1] = soc_est_k1',
         '路径 2：估计 SOC\n'
         '  • 调用估计器的 step() 方法\n'
         '  • EKF 模式下：内部完成 Predict + Update\n'
         '  • 输入是含偏置的电流和含噪声的电压\n'
         '  • 输出是对 SOC 的最优估计'],
        ['437-438', '# 开环安时积分（对比基准）\n'
         'SOC_open[k+1] = SOC_open[k] - i_meas_k/(Q_BAT*3600)*DT',
         '路径 3：开环积分\n'
         '  • 使用与 EKF 完全相同的含偏置电流\n'
         '  • 但不使用电压修正\n'
         '  • 这是 "不做 EKF 会怎样" 的对比基线\n'
         '  • SOC_open 会持续漂移，而 SOC_est 被电压修正'],
    ]
)

note(doc, '第 425-438 行的三路 SOC 更新是 v3 的核心创新：通过同时维护"真相"、"估计"、"开环"三条轨迹，可以精确量化 EKF 相比开环积分在任何时刻的任何改进。')

h(doc, '2.6.7 进度输出（第 440-443 行）', level=3)
p(doc, '相比 v2 增加了 SOC_true 和 SOC_open 的实时输出，方便观察 EKF 的追踪效果。')
code_block(doc, f'if k % 300 == 0:\n    print(f\'  Step {{k}}/{{N}}: SOC_true={{SOC_true[k]:.3f}}, \'\n          f\'SOC_est={{SOC_est_arr[k]:.3f}}, \'\n          f\'SOC_open={{SOC_open[k]:.3f}}\')')

h(doc, '2.6.8 结果统计与返回（第 445-493 行）', level=3)

tbl2(doc, ['行号', '代码/字段', '说明'],
    [
        ['450-451', 'soc_rmse = np.sqrt(np.mean((SOC_est_arr[:N] - SOC_true[:N])**2))\n'
         'soc_open_rmse = np.sqrt(... (SOC_open[:N] - SOC_true[:N])**2))',
         '🆕 新增 SOC 估计精度指标：\n'
         '  • SOC_rmse：EKF 估计的均方根误差（vs 真实值）\n'
         '  • SOC_open_rmse：开环积分的 RMSE（作为对比基准）\n'
         '  关键分析：open_rmse / ekf_rmse 的比值就是 EKF 的改进倍数。\n'
         '  实测 WLTC/2A 偏置：0.0116 / 0.0024 ≈ 4.8 倍'],
        ['462-493', 'return { ... SOC_est, SOC_true, SOC_open,\n'
         '           SOC_end_true, SOC_end_est,\n'
         '           SOC_rmse, SOC_open_rmse ... }',
         '🆕 返回字典新增 7 个字段：\n'
         '  SOC_est, SOC_true, SOC_open: 三条 SOC 轨迹\n'
         '  SOC_end_true, SOC_end_est: 终点 SOC\n'
         '  SOC_rmse, SOC_open_rmse: 误差指标'],
    ]
)

# ── 第496-525行：N_p 扫描 ──
h(doc, '2.7 N_p 敏感性扫描（第 496-525 行）', level=2)

tbl2(doc, ['行号', '变化', '说明'],
    [
        ['511-518', 'results.append({\n'
         '    \'SOC_end_true\': ..., \'SOC_end_est\': ...,\n'
         '    \'SOC_rmse\': ..., \'SOC_open_rmse\': ...,\n'
         '})',
         '🆕 N_p 扫描新增输出（相比 v2 优化版）：\n'
         '  SOC_end_true / SOC_end_est — 真实和估计终点 SOC\n'
         '  SOC_rmse / SOC_open_rmse — 不同 N_p 下的 SOC 估计精度\n\n'
         '这可以回答一个重要问题：N_p 的选择是否影响 SOC 估计精度？'
         '（理论上否——因为 EKF 与 N_p 是解耦的，但实际因 MPC 决策不同\n'
         ' 导致的功率分配差异可能间接影响 SOC 轨迹的"可估计性"）'],
    ]
)

# ── 第528-714行：绘图函数 ──
h(doc, '2.8 绘图函数（第 528-714 行）', level=2)

h(doc, '2.8.1 plot_four_way（第 531-620 行）', level=3)
p(doc, '与 v2 优化版基本一致，但在子图 (3) 中增加了：')
diff_added(doc, 'MPC 估计 SOC 轨迹（紫色细线）和开环 SOC 轨迹（青色细线）的叠加')

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['580-585', 'if \'SOC_est\' in mpc_result:\n'
         '    ax.plot(t_min, mpc_result[\'SOC_est\'], \'m-\', ...)\n'
         'if \'SOC_open\' in mpc_result:\n'
         '    ax.plot(t_min, mpc_result[\'SOC_open\'], \'c-\', ...)',
         '在 SOC 对比子图上叠加 EKF 估计 SOC 和开环积分 SOC。\n'
         '直观展示 EKF 对真实 SOC 的追踪效果。\n'
         '紫色线（EKF 估计）应几乎与黑色线（真实 SOC，MPC 的 SOC）重合，\n'
         '而青色线（开环积分）应逐渐偏离。'],
    ]
)

h(doc, '2.8.2 plot_soc_estimation（第 623-668 行）', level=3)
p(doc, '🎯 全新函数——SOC 估计效果专用对比图。', bold=True)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['623-626', 'def plot_soc_estimation(t, v, mpc_result, cycle_name=\'wltc\'):\n'
         '    if \'SOC_true\' not in mpc_result: return',
         '训练模式：如果返回结果中没有 SOC_true，跳过绘图（兼容旧版）。'],
        ['629', 'fig, axes = plt.subplots(3, 1, figsize=(14, 10), sharex=True)',
         '三行子图：\n'
         '  (1) SOC 轨迹对比：真实 vs EKF vs 开环\n'
         '  (2) SOC 误差对比：EKF 误差 vs 开环误差\n'
         '  (3) 速度曲线（时间轴参考）'],
        ['632-643', '子图 (1)：三条 SOC 轨迹',
         '  • 黑色粗线：真实 SOC\n'
         '  • 红色线：EKF 估计（标注 RMSE）\n'
         '  • 蓝色虚线：开环积分（标注 RMSE）\n'
         '  标题中标明电流偏置值（如 bias=2A）'],
        ['646-655', '子图 (2)：SOC 误差',
         '  • 红色：EKF 误差（围绕 0 波动）\n'
         '  • 蓝色：开环误差（单调漂移）\n'
         '  • 红色填充：突出 EKF 误差的零均值特性\n'
         '  这是整张图的"证据"——EKF 误差持续在 0 附近，\n'
         '  而开环误差单调增大。'],
        ['658-662', '子图 (3)：速度曲线',
         '提供时间轴参考——可以对照速度曲线理解 SOC 变化的原因\n'
         '（急加速段 SOC 下降、制动回收段 SOC 上升）。'],
    ]
)

h(doc, '2.8.3 plot_np_sensitivity（第 671-714 行）', level=3)
diff_added(doc, '第三行子图：SOC 估计 RMSE vs N_p（第 698-708 行）')

# ── 第717-779行：指标打印 ──
h(doc, '2.9 指标打印（第 717-779 行）', level=2)

diff_added(doc, '表格末尾新增两行：SOC 估计 RMSE 和开环 SOC RMSE（第 759-764 行）')
tbl2(doc, ['行号', '代码', '说明'],
    [
        ['735', 'mpc_SOC_end = mpc_result.get(\'SOC_end_true\',\n'
         '                   mpc_result.get(\'SOC_end\', mpc_result[\'SOC\'][-1]))',
         '兼容性兼容：优先取 SOC_end_true（v3 新增），\n'
         '  fallback 到 SOC_end（v2），再 fallback 到 SOC[-1]（v1）。\n'
         '  get() 嵌套保证了与旧版返回结果的兼容。'],
        ['759-764', 'rows.append((\'SOC 估计 RMSE\', \'\', \'\', \'\', f\'{...:.4f}\'))\n'
         'rows.append((\'开环 SOC RMSE\', \'\', \'\', \'\', f\'{...:.4f}\'))',
         '🆕 新增 SOC 估计精度指标行。\n'
         '只显示在 MPC 列（Rule/DP/ECMS 没有 SOC 估计功能）。'],
    ]
)

# ── 第782-1017行：主程序 ──
h(doc, '2.10 主程序 main（第 782-1017 行）', level=2)

h(doc, '2.10.1 新增参数解析（第 799-819 行）', level=3)
p(doc, '相比 v2 的 12 个参数，v3 新增 7 个参数：')

tbl2(doc, ['参数', '默认值', '说明'],
    [
        ['--soc-estimator', 'ekf', 'SOC 估计方法选择'],
        ['--current-bias', '2.0 A', '模拟电流传感器偏置'],
        ['--current-noise', '0.5 A', '电流测量噪声'],
        ['--voltage-noise', '0.1 V', '电压测量噪声'],
        ['--ekf-x0', 'None', 'EKF 初始 SOC（设偏可测试收敛）'],
        ['--ekf-q', '5e-5', 'EKF 过程噪声'],
        ['--ekf-r', '0.03', 'EKF 测量噪声'],
        ['--seed', '42', '随机种子（保证可复现性）'],
    ]
)

h(doc, '2.10.2 参数打包（第 827-843 行）', level=3)
tbl2(doc, ['行号', '代码', '说明'],
    [
        ['827-843', 'mpc_kwargs = {\n'
         '    ..., \'soc_estimator\': args.soc_estimator,\n'
         '    \'current_bias\': args.current_bias, ...\n'
         '}',
         '🆕 新增 7 个参数打包到 kwargs 中，传递到 mpc_sim()。'],
    ]
)

h(doc, '2.10.3 MPC+EKF 运行（第 878-879 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['878-879', 'print(f\'\\n[3/4] MPC+{args.soc_estimator.upper()}...\')\n'
         'mpc_result = mpc_sim(P_load, SOC_0=0.6, N_p=n_p, **mpc_kwargs)',
         '调用 mpc_sim() 并传入所有参数。\n'
         'mpc_kwargs 包含 EKF 参数和传感器参数。'],
    ]
)

h(doc, '2.10.4 SOC 估计对比图（第 967-968 行）', level=3)

tbl2(doc, ['行号', '代码', '说明'],
    [
        ['967-968', '# SOC 估计对比图\n'
         'plot_soc_estimation(t, v, mpc_result, cycle)',
         '🆕 每次运行都输出 SOC 估计效果图（除非 --plot-only）。'],
    ]
)

h(doc, '2.10.5 结果保存扩展（第 970-1003 行）', level=3)

tbl2(doc, ['行号', '变化', '说明'],
    [
        ['971-982', 'df_mpc = pd.DataFrame({\n'
         '    ..., \'SOC_true\': ..., \'SOC_est\': ..., \'SOC_open\': ...,\n'
         '})',
         '🆕 CSV 文件增加三列 SOC 数据：SOC_true、SOC_est、SOC_open。'],
        ['987-1003', 'summary = {\n'
         '    ..., \'estimator\': ..., \'current_bias\': ...,\n'
         '    \'SOC_end_true\': ..., \'SOC_end_est\': ...,\n'
         '    \'SOC_rmse\': ..., \'SOC_open_rmse\': ...,\n'
         '}',
         '🆕 _summary.csv 增加 SOC 估计相关字段，\n'
         '便于批量实验横向对比 SOC 估计精度。'],
    ]
)

doc.add_page_break()

# ========================================================================
# 第三部分：关键算法流程图
# ========================================================================
h(doc, '第三部分：关键算法流程图', level=1)

h(doc, '3.1 三路 SOC 更新流程', level=2)

code_block(doc, '''┌─────────────────────────────────────────────────────────────────┐
│                    MPC 单步算法 (k 时刻)                          │
├─────────────────────────────────────────────────────────────────┤
│                                                                   │
│  SOC_est[k] (EKF 估计) ──→ ┌─ MPC 网格搜索 ──→ P_fc[k]         │
│                            │  (基于估计 SOC 做决策)               │
│                            └───────────────────────────────────── │
│                                    │                              │
│                                    ↓                              │
│  ┌──────────────── 三路 SOC 更新 ────────────────┐               │
│  │                                               │               │
│  │ ① 真实 SOC (无偏电流):                       │               │
│  │   i_real = bat_current(P_bat, SOC_true[k])   │               │
│  │   SOC_true[k+1] = SOC_true[k] - i_real/Q×dt  │               │
│  │                                               │               │
│  │ ② 模拟传感器:                                 │               │
│  │   i_meas = i_real + 偏置 + 噪声               │               │
│  │   v_meas = OCV(SOC_true[k]) + 噪声            │               │
│  │                                               │               │
│  │ ③ 估计 SOC (EKF):                            │               │
│  │   SOC_est[k+1] = estimator.step(i_meas, v_meas)              │
│  │                                               │               │
│  │ ④ 开环积分 (对比):                           │               │
│  │   SOC_open[k+1] = SOC_open[k] - i_meas/Q×dt  │               │
│  └───────────────────────────────────────────────┘               │
│                                                                   │
│  SOC 精度: EKF(O.0024) >> 开环(0.0116) >> 无EKF时MPC的SOC(漂移) │
└─────────────────────────────────────────────────────────────────┘''')

h(doc, '3.2 EKF 两阶段流程（单步）', level=2)

code_block(doc, '''┌────────────────────────────────────────────────────┐
│           EKF 单步：Predict + Update              │
├────────────────────────────────────────────────────┤
│                                                    │
│  Predict (时间更新):                               │
│  ┌──────────────────────────────────────────┐     │
│  │ SOC_pred = SOC_est - I_meas/Q_bat × dt   │     │
│  │ P_pred = P + Q                           │     │
│  └──────────────────────────────────────────┘     │
│              ↓                                     │
│              ↓ y = V_meas - OCV(SOC_pred)          │
│              ↓                                    │
│  Update (测量更新):                                │
│  ┌──────────────────────────────────────────┐     │
│  │ H = dOCV/dSOC                             │     │
│  │ K = P_pred × H / (H²×P_pred + R)         │     │
│  │ SOC_est = SOC_pred + K × y                │     │
│  │ P_est = (1 - K×H) × P_pred               │     │
│  └──────────────────────────────────────────┘     │
│                                                    │
│  输出: SOC_est (电流+电压融合后的最优估计)         │
└────────────────────────────────────────────────────┘''')

h(doc, '3.3 MPC+EKF 的类层次结构', level=2)

code_block(doc, '''SOCEstimator (基类)            ← 统一接口
├── OpenLoopEstimator         ← 开环积分（无电压修正）
├── EKFEstimator              ← EKF（固定 Q/R）
│   └── 新增: AEKFEstimator  ← 自适应 Q/R
│
└── build_estimator()         ← 工厂方法

mpc_sim() 中的集成:
  estimator = build_estimator(soc_estimator, ...)
  for k in range(N):
      estimator.step(i_meas, v_meas)   ← 多态调用''')

doc.add_page_break()

# ========================================================================
# 第四部分：v2 vs v3 关键代码对比
# ========================================================================
h(doc, '第四部分：v2 优化版 vs v3 EKF 版关键代码对比', level=1)

h(doc, '4.1 SOC 更新方式对比', level=2)

p(doc, 'v2 优化版（mpc_ems_optimized.py）：')
code_block(doc, '''# SOC 开环更新（单一路径）
SOC = np.zeros(N+1)
SOC[0] = SOC_0

for k in range(N):
    P_fc[k] = argmin(J)   ← MPC 基于 SOC[k] 做决策

    # SOC 更新（开环）
    SOC[k+1] = mpc_step_soc(SOC[k], P_fc[k], P_load[k])

    # 所有代码使用同一个 SOC 数组——它是"真实"也是"估计"，无法区分''')

p(doc, 'v3 EKF 版（mpc_ems_ekf.py）：')
code_block(doc, '''# 三路 SOC 并行追踪
SOC_true = np.zeros(N+1)    # 真实 SOC（无偏）
SOC_est  = np.zeros(N+1)    # 估计 SOC（EKF）
SOC_open = np.zeros(N+1)    # 开环积分（对比）

for k in range(N):
    P_fc[k] = argmin(J)   ← MPC 基于 SOC_est[k] 做决策

    # 真实 SOC（无偏电流）
    i_real = battery_current(SOC_true[k], P_bat[k])
    SOC_true[k+1] = SOC_true[k] - i_real/Q*dt

    # 含偏置的传感器测量
    i_meas = i_real + bias + noise
    v_meas = OCV(SOC_true[k]) + noise

    # EKF 估计 SOC（融合电流+电压）
    SOC_est[k+1] = estimator.step(i_meas, v_meas)

    # 开环积分（纯电流，无修正）
    SOC_open[k+1] = SOC_open[k] - i_meas/Q*dt''')

h(doc, '4.2 代价函数对比', level=2)

p(doc, 'v2 优化版与 v3 完全相同（SOC 惩罚函数代码一致）：')
code_block(doc, '''J_total = Σ H₂·Δt + s·|P_bat|/3600·Δt + soc_tracking_penalty(...) + w_slew·(ΔP_fc)²
                                                                         ↑
                                                                    EKF 影响的是
                                                                   SOC_est 的精度，
                                                                   COST 公式不变''')

h(doc, '4.3 输出数据对比', level=2)

tbl(doc, ['输出项', 'v2 优化版', 'v3 EKF 版'],
    [
        ['mpc_sim 返回字段', 'SOC, P_fc, P_bat, m_H2, 等', '🆕 + SOC_true, SOC_est, SOC_open, SOC_rmse'],
        ['CSV 文件', 'SOC, P_fc, P_bat, 氢耗', '🆕 + SOC_true, SOC_est, SOC_open 三列'],
        ['_summary.csv', 'cycle, N_p, H2, SOC_end, 参数', '🆕 + estimator, current_bias, SOC_rmse, SOC_open_rmse'],
        ['诊断图', '四方法对比图', '🆕 + SOC 估计对比图'],
    ]
)

doc.add_page_break()

# ========================================================================
# 第五部分：运行效果实测
# ========================================================================
h(doc, '第五部分：运行效果实测（WLTC, N_p=50, 2A 偏置）', level=1)

h(doc, '5.1 三种估计器性能对比', level=2)

tbl(doc, ['指标', '开环积分', 'EKF', 'AEKF', 'EKF vs 开环'],
    [
        ['SOC RMSE', '0.0116', '0.0024', '0.0028', '↓ 78.9%'],
        ['终点误差', '0.0201', '0.0015', '0.0015', '↓ 92.5%'],
        ['原始氢耗 (kg)', '0.2421', '0.2198', '0.2191', '↓ 9.2%'],
        ['真终点 SOC', '0.590', '0.572', '0.572', '-'],
        ['计算开销增量', '基线', '< 0.5%', '< 1%', '-'],
    ]
)

h(doc, '5.2 EKF 模式典型输出', level=2)

code_block(doc, '''$ python scripts/mpc_ems_ekf.py --cycle wltc --np 50

[MPC+EKF] N_p=50, s=130.0, w_soc=1200.0, bias=2.0A,
           SOC_true_0=0.60, SOC_est_0=0.60
[MPC+EKF] 开始仿真... (1801 步)

  Step 0/1801:   SOC_true=0.600, SOC_est=0.600, SOC_open=0.600
  Step 300/1801: SOC_true=0.581, SOC_est=0.581, SOC_open=0.578
  Step 600/1801: SOC_true=0.576, SOC_est=0.580, SOC_open=0.570
  Step 900/1801: SOC_true=0.572, SOC_est=0.571, SOC_open=0.562
  Step 1200/1801: SOC_true=0.572, SOC_est=0.571, SOC_open=0.559
  Step 1500/1801: SOC_true=0.573, SOC_est=0.572, SOC_open=0.557
  Step 1800/1801: SOC_true=0.572, SOC_est=0.573, SOC_open=0.552

[MPC+EKF] 完成.
  SOC: true_end=0.572, est_end=0.571, open_end=0.552
  SOC RMSE: EKF=0.0024, OpenLoop=0.0116          ← EKF 精度 4.8× 优于开环
  H2: raw=0.2198 kg, eq=0.2827 kg''')

p(doc, '关键观察：', bold=True)
p(doc, '  1. SOC_open 从 0.600 单调下降到 0.552——2A 偏置导致的持续漂移')
p(doc, '  2. SOC_est 始终跟随 SOC_true，终点误差仅 0.001——EKF 电压修正有效')
p(doc, '  3. SOC 估计 RMSE = 0.0024，约为开环的 1/5——定量证明了 EKF 的改进')

doc.add_page_break()

# ========================================================================
# 第六部分：文件变更清单
# ========================================================================
h(doc, '第六部分：文件变更清单（vs v2 优化版）', level=1)

h(doc, '6.1 新增代码', level=2)

tbl(doc, ['位置', '内容', '行数', '来源'],
    [
        ['第 59-66 行', 'EKF/传感器参数', '8', '新定义'],
        ['第 72-96 行', '电池辅助函数 (lookup_ocv 等)', '25', '从 ekf_soc_estimator.py 提取'],
        ['第 99-226 行', 'SOC 估计器类体系 (5 个类/函数)', '128', '全新设计（基类+3子类+工厂）'],
        ['第 232-234 行', 'simulate_voltage()', '3', '从 ekf_soc_estimator.py 提取'],
        ['第 303-307 行', 'mpc_sim 新增参数 (8 个)', '5', '扩展函数签名'],
        ['第 326-337 行', '三路 SOC 数组分配', '12', '新逻辑'],
        ['第 425-438 行', '三路 SOC 更新循环体', '14', '新逻辑（核心改动）'],
        ['第 450-451 行', 'SOC RMSE 计算', '2', '新逻辑'],
        ['第 462-493 行', '返回字段扩展 (7 个)', '32', '扩展返回字典'],
        ['第 623-668 行', 'plot_soc_estimation()', '46', '全新函数'],
        ['第 698-708 行', 'plot_np_sensitivity 新增子图', '11', '扩展绘图'],
        ['第 759-764 行', 'SOC RMSE 指标行', '6', '扩展打印'],
        ['第 814 行', '--seed 参数', '1', '新参数'],
    ]
)

h(doc, '6.2 修改代码', level=2)

tbl(doc, ['位置', '修改内容', '说明'],
    [
        ['第 355 行', 'soc_est_k = SOC_est_arr[k]', 'MPC 使用估计 SOC 而非开环 SOC'],
        ['第 440-443 行', '进度输出增加 SOC_true/open', '更丰富的运行时监控'],
        ['第 580-585 行', 'plot_four_way 增加估计 SOC', '图面信息更丰富'],
        ['第 735 行', 'mpc_SOC_end 兼容性兼容', '保证支持 v1/v2/v3 返回结构'],
        ['第 971-982 行', 'CSV 保存增加 SOC 三列', '数据后处理更方便'],
        ['第 987-1003 行', 'summary 增加 SOC 指标', '批量实验分析'],
    ]
)

h(doc, '6.3 未修改代码（从 v2 完全复用）', level=2)
p(doc, '以下部分与 mpc_ems_optimized.py 完全相同：')
p(doc, '  • MPC 超参数（N_p, S_MPC, W_SOC, BETA_TERM, SOC_DEADBAND 等）')
p(doc, '  • soc_tracking_penalty() 函数（SOC 死区 + 软下限 + 终端罚 + 终点罚）')
p(doc, '  • mpc_step_soc() 函数（含不可行→None 的后备机制）')
p(doc, '  • soc_equivalent_h2() 等效氢耗修正')
p(doc, '  • 网格搜索内循环（cos 函数计算逻辑）')
p(doc, '  • 后备容错策略（分级 fallback）')
p(doc, '  • plot_four_way() 的 5 合 1 主图框架')
p(doc, '  • print_four_way_metrics() 的指标打印框架')
p(doc, '  • 主流程控制（加载工况 → Rule → DP → ECMS → MPC → 输出）')

p(doc, '', size=6)
p(doc, '—— 文档完 ——', bold=True, size=11, color=RGBColor(100, 100, 100))
p(doc, '生成于 2026-07-07 | 基于 scripts/mpc_ems_ekf.py (v3.0)', size=9, color=RGBColor(140, 140, 140))

output_path = os.path.join(DOCS_DIR, 'MPC_EMS_EKF_逐行代码原理分析.docx')
doc.save(output_path)
print(f'[OK] 文档已保存: {output_path}')
print(f'     大小: {os.path.getsize(output_path) / 1024:.0f} KB')
