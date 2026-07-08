# -*- coding: utf-8 -*-
"""生成 MPC 算法原理解析 DOCX 文档（面向小白的通俗版）"""

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUTPUT_DIR = os.path.join(PROJECT_ROOT, 'docs')

def set_cell_shading(cell, color):
    """设置表格单元格底色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)

def add_code_block(doc, code_text):
    """添加代码块（灰色底 + 等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 添加底色
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear',
    })
    shading.append(shd)
    return p

def build_doc():
    doc = Document()

    # ── 全局样式 ──
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing = 1.35

    # ── 标题 ──
    title = doc.add_heading('MPC（模型预测控制）算法原理解析', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—— 基于 EMS 项目的逐行代码解读 · 通俗版')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── 前言 ──
    doc.add_heading('前言：为什么 MPC 难懂？', level=1)
    doc.add_paragraph(
        'MPC（Model Predictive Control，模型预测控制）难懂，是因为它同时包含了三个概念：'
        '"模型" + "预测" + "控制"，而且这三个东西是揉在一起同时发生的。'
        '很多教程一上来就讲数学公式，反而把直观逻辑淹没了。'
    )
    doc.add_paragraph(
        '这份文档的目标是：用你项目里的真实代码，把 MPC 的每一块掰开揉碎讲清楚。'
        '读完你应该能回答三个问题：'
    )
    for q in ['① MPC 到底在算什么？', '② 为什么它叫"滚动时域"？', '③ 它比 Rule/DP/ECMS 好在哪？']:
        doc.add_paragraph(q, style='List Bullet')

    # ══════════════════════════════════════════════
    # 第一章
    # ══════════════════════════════════════════════
    doc.add_heading('一、MPC 一句话概括', level=1)
    doc.add_paragraph(
        'MPC = 每走一步，先往前看 N 步，把所有可能的选择都试一遍，选最好的那一步走，'
        '然后下一步重新来。'
    )

    doc.add_heading('一个类比：下棋', level=2)
    table = doc.add_table(rows=5, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.columns[0].cells:
        set_cell_shading(cell, 'E8F0FE')
    headers = [('MPC 概念', '下棋类比')]
    data = [
        ('当前时刻 k', '当前棋盘局面'),
        ('预测模型（Model）', '你记住的棋谱/套路'),
        ('预测时域 N_p', '你想往后想几步（3步？10步？）'),
        ('代价函数 J', '这一步走完，局势是好是坏'),
        ('滚动时域', '走一步，重新想后面的'),
    ]
    for i, (a, b) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = a
        row.cells[1].text = b
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 第二章
    # ══════════════════════════════════════════════
    doc.add_heading('二、MPC 三步走（代码逐行解读）', level=1)

    # ── 第1步 ──
    doc.add_heading('第 1 步：预测——往前看 N_p 步', level=2)
    doc.add_paragraph(
        'MPC 的第一步是"知道未来要发生什么"。'
        '在你的代码里，负载功率 P_load 是预先知道的工况数据（WLTC/NEDC/CLTC 驾驶循环），'
        '所以 MPC 可以"预知"未来的负载需求。'
    )
    doc.add_paragraph('代码对应：')
    add_code_block(doc,
        'horizon = min(N_p, N - k)                          # 决定看多远\n'
        'p_load_pred = P_load[k: k + horizon]                # 取出未来 horizon 步的负载'
    )
    doc.add_paragraph(
        '举例：如果 N_p = 50，当前在 k=100，那么 MPC 会取出 '
        'P_load[100:150] 作为"已知的未来之路"。'
    )

    # ── 第2步 ──
    doc.add_heading('第 2 步：优化——试所有候选，选最好的', level=2)
    doc.add_paragraph(
        '这是 MPC 的核心。代码中通过「网格搜索」来枚举所有可能的燃料电池输出功率，'
        '对每一个候选值，模拟它未来 horizon 步的 SOC 轨迹，计算总代价，取最小的那个。'
    )

    doc.add_heading('2.1 枚举候选功率', level=3)
    add_code_block(doc,
        'for j in range(N_PFC):                              # 遍历所有候选功率\n'
        '    p_fc_cand = PFC_GRID[j]                          # 例如 [0, 5, 10, ..., 40] kW\n'
        '    h2_cand = H2_GRID[j]                             # 对应的氢耗'
    )
    doc.add_paragraph(
        '注意：这里 N_PFC 是离散化的功率网格点数（比如 50 个点），PFC_GRID[j] 是第 j 个候选功率值。'
        '你的代码把 0~40kW 分成了等间距的离散点，然后逐个尝试。'
    )

    doc.add_heading('2.2 模拟未来轨迹', level=3)
    add_code_block(doc,
        'soc_pred = soc_est_k                                 # 从当前 SOC 开始\n'
        'J_total = 0.0\n'
        'for i in range(horizon):                             # 往前推 horizon 步\n'
        '    p_bat_i = p_load_i - p_fc_cand                   # 电池功率 = 负载 - FC 功率\n'
        '    \n'
        '    # 代价 = 氢耗 + 等效电池能量 + SOC 偏离惩罚\n'
        '    J_total += h2_cand * DT                           # 氢耗代价\n'
        '    J_total += s_factor * abs(p_bat_i) / 3600.0 * DT # 电池"虚拟"代价\n'
        '    J_total += soc_tracking_penalty(soc_pred, ...)    # SOC 偏离惩罚\n'
        '    \n'
        '    soc_pred_next = mpc_step_soc(soc_pred, p_fc_cand, p_load_i)  # SOC 前推一步\n'
        '    if soc_pred_next is None:                        # 物理不可行？抛弃\n'
        '        feasible = False\n'
        '        break\n'
        '    soc_pred = soc_pred_next'
    )

    doc.add_heading('2.3 代价函数的组成（为什么这么算？）', level=3)
    doc.add_paragraph('代价函数 J 由三部分组成，对应三个控制目标：')

    table2 = doc.add_table(rows=4, cols=3, style='Table Grid')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table2.columns[0].cells:
        set_cell_shading(cell, 'E8F0FE')

    t2_data = [
        ('代价项', '代码', '物理含义'),
        ('氢耗 J₁', 'h2_cand * DT', '直接消耗的氢气质量——省油'),
        ('电池等效 J₂', 's_factor * |P_bat| / 3600 * DT', '电池充放电的"虚拟氢耗"——\n避免过度用电（ECMS 的核心思想）'),
        ('SOC 惩罚 J₃', 'soc_tracking_penalty(...)', '让 SOC 维持在 0.6 附近——\n保电池健康 + 终端不亏电'),
    ]
    for i, (a, b, c) in enumerate(t2_data):
        for j, val in enumerate([a, b, c]):
            table2.rows[i].cells[j].text = val
            for p in table2.rows[i].cells[j].paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph(
        '三个目标互相牵制：想省氢（J₁）就得用电池，但用电池会偏离 SOC（J₃），'
        '而且电池的电也不是白给的（J₂）。MPC 就是在这三个目标之间找最优平衡点。'
    )

    doc.add_heading('2.4 选最优', level=3)
    add_code_block(doc,
        'if feasible and J_total < J_best:\n'
        '    J_best = J_total\n'
        '    best_j = j                                       # 记录最优候选的索引'
    )
    doc.add_paragraph(
        '在所有候选功率中，选总代价最小的那个。如果所有候选都不可行（物理限制导致 SOC 越界），'
        '代码还有一个「后备策略」——退化为只看一步。'
    )

    # ── 第3步 ──
    doc.add_heading('第 3 步：执行——只走第一步，然后滚动', level=2)
    add_code_block(doc,
        '# ── 执行最优控制 ──\n'
        'P_fc[k] = PFC_GRID[best_j]                           # 只执行第 k 步的功率\n'
        'P_bat[k] = P_load[k] - P_fc[k]                       # 电池功率 = 余缺\n'
        'm_H2[k] = fc_hydrogen_flow(P_fc[k]) * DT            # 记录氢耗'
    )
    doc.add_paragraph(
        '注意：这里的 PFC_GRID[best_j] 是在第 2 步中选出的最优功率，'
        '但 MPC 只执行当前这一步（k 时刻），不会把整个 horizon 的方案都执行完。'
        '这是 MPC 和 DP（动态规划）最本质的区别。'
    )
    doc.add_paragraph(
        '为什么只走一步？因为真实世界有噪声和模型误差，你算好的未来轨迹会慢慢偏离实际。'
        '每步重新算，才能不断修正。这就是「滚动时域」（Receding Horizon）的含义。'
    )

    # ══════════════════════════════════════════════
    # 第三章
    # ══════════════════════════════════════════════
    doc.add_heading('三、一个完整的 MPC 步进示例', level=1)
    doc.add_paragraph('假设：N_p = 3（为了简化），SOC 当前 = 0.60，未来 3 步负载 = [30, 25, 35] kW')

    table3 = doc.add_table(rows=7, cols=6, style='Table Grid')
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for cell in table3.columns[0].cells:
        set_cell_shading(cell, 'E8F0FE')

    t3_data = [
        ('候选 FC 功率\n(kW)', '氢耗\n(kg)', '第1步后 SOC\n(负载30kW)', '第2步后 SOC\n(负载25kW)', '第3步后 SOC\n(负载35kW)', '总代价 J'),
        ('5', '0.001', '0.58', '0.57', '0.54', '0.042'),
        ('10', '0.002', '0.59', '0.59', '0.56', '0.035'),
        ('15', '0.003', '0.60', '0.61', '0.59', '0.028'),
        ('20', '0.004', '0.61', '0.62', '0.61', '0.033'),
        ('25', '0.005', '0.62', '0.64', '0.63', '0.045'),
        ('30', '0.006', '0.63', '0.65', '0.65', '0.058'),
    ]
    # 高亮最优行
    for i, row_data in enumerate(t3_data):
        for j, val in enumerate(row_data):
            table3.rows[i].cells[j].text = val
            for p in table3.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
        if i == 4:  # 第5行（索引4）= 15kW = 最优
            for cell in table3.rows[i].cells:
                set_cell_shading(cell, 'D4EDDA')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('→ 最优选择：P_fc = 15 kW（绿色高亮行），因为它总代价最小')
    run.bold = True
    run.font.color.rgb = RGBColor(0x15, 0x80, 0x3A)

    doc.add_paragraph(
        '但 MPC 只在当前步执行 P_fc = 15 kW。下一步 k+1，SOC 可能实际变成了 0.598（因为有噪声），'
        'MPC 重新做一遍上述计算。这就是滚动。'
    )

    # ══════════════════════════════════════════════
    # 第四章
    # ══════════════════════════════════════════════
    doc.add_heading('四、MPC vs 其他方法', level=1)

    table4 = doc.add_table(rows=5, cols=4, style='Table Grid')
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table4.columns[0].cells:
        set_cell_shading(cell, 'E8F0FE')

    t4_data = [
        ('方法', '核心思想', '需要预知全程？', '适合场景'),
        ('Rule 规则', '写死 if-else，"SOC 低就充电"', '不需要', '简单、可靠、但次优'),
        ('DP 动态规划', '从终点倒着算，全局最优', '需要！', '离线的基准对比',
        ),
        ('ECMS 等效消耗', '每步算氢电价，选最经济的', '不需要', '近视但快，在线可用'),
        ('MPC 模型预测', '每步往前看 N 步，滚动优化', '只需要未来 N 步', '在线控制 + 规划平衡'),
    ]
    for i, row_data in enumerate(t4_data):
        for j, val in enumerate(row_data):
            table4.rows[i].cells[j].text = val
            for p in table4.rows[i].cells[j].paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
        if i == 4:
            for cell in table4.rows[i].cells:
                set_cell_shading(cell, 'FFF3CD')

    doc.add_paragraph('')
    doc.add_heading('关键区别：MPC 是 DP 和 ECMS 的折中', level=2)
    for text in [
        'DP：知道全程（600 步），算全局最优——但现实中你不可能预知 600 秒后的负载',
        'ECMS：只知道当前 1 步，反应快但「近视」，遇到长上坡来不及准备',
        'MPC：知道未来 N_p 步（如 50 步），在局部视野下做最优决策——兼顾了前瞻性和实时性',
    ]:
        doc.add_paragraph(text, style='List Bullet')

    # ══════════════════════════════════════════════
    # 第五章
    # ══════════════════════════════════════════════
    doc.add_heading('五、MPC 的重要参数', level=1)

    table5 = doc.add_table(rows=6, cols=3, style='Table Grid')
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table5.columns[0].cells:
        set_cell_shading(cell, 'E8F0FE')

    t5_data = [
        ('参数', '含义', '调大 → 效果'),
        ('N_p', '预测时域长度', '看得更远 → 全局性更好，但计算量更大\n你的代码扫了 [10,20,30,50,80,120,200]'),
        ('s_factor\n(等效因子)', '电池电量的「虚拟氢价」', 's↑ → 更不愿用电池 → SOC 维持好但氢耗高'),
        ('W_soc\n(SOC 惩罚权重)', 'SOC 偏离 0.6 的重视程度', 'W↑ → SOC 更"僵硬"守在 0.6\nW↓ → SOC 浮动范围大'),
        ('PFC_GRID\n(候选点数)', '离散化的精细程度', '点越多 → 精度高但计算慢'),
    ]
    for i, (a, b, c) in enumerate(t5_data):
        for j, val in enumerate([a, b, c]):
            table5.rows[i].cells[j].text = val
            for p in table5.rows[i].cells[j].paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph('')

    # ══════════════════════════════════════════════
    # 第六章
    # ══════════════════════════════════════════════
    doc.add_heading('六、你代码里的 MPC + EKF 组合', level=1)
    doc.add_paragraph(
        '你的项目在 MPC 的基础上还加了 EKF（扩展卡尔曼滤波），这是两个独立模块的串联：'
    )

    doc.add_heading('MPC 和 EKF 各管各的事', level=2)
    table6 = doc.add_table(rows=3, cols=2, style='Table Grid')
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    t6_data = [
        ('MPC 控制模块', 'EKF 估计模块'),
        ('输入：当前 SOC 估计值 + 未来 N_p 步负载\n输出：P_fc（燃料电池功率决策）', '输入：电流 + 电压传感器测量\n输出：修正后的 SOC 估计值'),
        ('回答："现在该发多少电？"', '回答："电池还剩多少电？"'),
    ]
    for i, (a, b) in enumerate(t6_data):
        for j, val in enumerate([a, b]):
            table6.rows[i].cells[j].text = val
            for p in table6.rows[i].cells[j].paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
        if i == 0:
            for cell in table6.rows[i].cells:
                set_cell_shading(cell, 'E8F0FE')

    doc.add_paragraph('')
    doc.add_paragraph('整体循环流程：')
    add_code_block(doc,
        'for k in range(N):  # 每个时间步\n'
        '    # ① EKF 修正 SOC（用电流+电压融合）\n'
        '    soc_est_k = SOC_est_arr[k]\n'
        '    \n'
        '    # ② MPC 决策（基于修正后的 SOC）\n'
        '    ... 网格搜索、模拟、选最优 ...\n'
        '    P_fc[k] = PFC_GRID[best_j]\n'
        '    \n'
        '    # ③ 真实 SOC 演化\n'
        '    SOC_true[k+1] = SOC_true[k] - i_real / Q * dt\n'
        '    \n'
        '    # ④ 传感器测量（含偏置和噪声）\n'
        '    i_meas_k = i_real + bias + noise\n'
        '    v_meas_k = simulate_voltage(...)\n'
        '    \n'
        '    # ⑤ EKF 更新 SOC 估计\n'
        '    SOC_est_arr[k+1] = estimator.step(i_meas_k, v_meas_k)'
    )

    doc.add_paragraph(
        '为什么需要 EKF？因为安时积分（OpenLoop）在有电流偏置时 SOC 会越漂越远，'
        '而 EKF 用电压测量做"锚定"，定期修正漂移。你用 --current-bias 参数模拟传感器故障，'
        '就是测试 EKF 的抗漂移能力。'
    )

    # ══════════════════════════════════════════════
    # 第七章
    # ══════════════════════════════════════════════
    doc.add_heading('七、常见问题 FAQ', level=1)

    faqs = [
        ('MPC 和 DP 什么关系？',
         'DP 从终点倒推，需要知道全程负载，算全局最优解，是「上帝视角」。'
         'MPC 从当前往前推 N 步，只需要知道未来 N 步的负载，是「有限视野」。'
         'DP 不能在线使用，MPC 可以。你的代码把 DP 结果当「理论下限」做对比。'),
        ('为什么 N_p 不是越大越好？',
         'N_p 大 → 看得远 → 但计算量正比于 N_p × N_PFC（候选点数）。'
         '而且预测得越远，模型误差越大，远了也不准。'
         '你的 N_p 扫描就是找这个平衡点。'),
        ('s_factor 是什么？',
         '它是 ECMS 中的"等效因子"——把电池的 1kWh 电能折算成多少克氢气。'
         's 越大，MPC 越倾向于用燃料电池直接供电（而不是先用电池再用 FC 充电），'
         '因为电池电量的"虚拟价格"变贵了。'),
        ('MPC 一定是全局最优吗？',
         '不是。MPC 是「有限时域下的局部最优」。因为只看 N_p 步，'
         '可能错过需要更远视野才能看到的更优策略。但在工程实践中，'
         'N_p 足够大（如 50~100 步）时，MPC 的效果非常接近全局最优。'),
        ('代码里为什么叫"网格搜索"？',
         '因为 PFC_GRID 是一个离散的功率网格，代码枚举了网格上的每一个点。'
         '更高效的 MPC 会用优化求解器（如 CVXOPT、qpOASES），但网格搜索直观、可控、'
         '适合教学和原型验证。'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f'Q: {q}')
        run.bold = True
        run.font.color.rgb = RGBColor(0x19, 0x60, 0xD2)
        doc.add_paragraph(f'A: {a}')

    # ══════════════════════════════════════════════
    # 附录
    # ══════════════════════════════════════════════
    doc.add_heading('附录：关键代码位置速查', level=1)

    table_a = doc.add_table(rows=8, cols=3, style='Table Grid')
    table_a.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table_a.columns[0].cells:
        set_cell_shading(cell, 'E8F0FE')

    a_data = [
        ('功能', '文件', '行号'),
        ('MPC 主仿真循环', 'scripts/mpc_ems_ekf.py', '297–493'),
        ('网格搜索 + 代价计算', 'scripts/mpc_ems_ekf.py', '361–406'),
        ('后备策略（不可行时）', 'scripts/mpc_ems_ekf.py', '408–418'),
        ('SOC 一步转移', 'scripts/mpc_ems_ekf.py', '276–291'),
        ('SOC 跟踪惩罚函数', 'scripts/mpc_ems_ekf.py', '250–271'),
        ('EKF SOC 估计器', 'scripts/mpc_ems_ekf.py', '120–157'),
        ('AEKF 自适应估计器', 'scripts/mpc_ems_ekf.py', '160–210'),
    ]
    for i, (a, b, c) in enumerate(a_data):
        for j, val in enumerate([a, b, c]):
            table_a.rows[i].cells[j].text = val
            for p in table_a.rows[i].cells[j].paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph('')
    doc.add_paragraph('— END —')
    doc.add_paragraph(
        '说明：本文档为 EMS 能源管理研究项目的一部分。'
        '代码来自 scripts/mpc_ems_ekf.py。'
    )

    # ── 保存 ──
    out_path = os.path.join(OUTPUT_DIR, 'MPC_算法原理解析_通俗版.docx')
    doc.save(out_path)
    print(f'[OK] 文档已生成: {out_path}')
    return out_path

if __name__ == '__main__':
    build_doc()
