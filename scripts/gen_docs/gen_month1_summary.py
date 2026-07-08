# -*- coding: utf-8 -*-
"""生成第1个月（Month 1）总结报告 .docx"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

PROJECT_ROOT = r'F:\CLAUDE\research\ems-platform'
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')
RESULTS_DIR = os.path.join(PROJECT_ROOT, 'results')

def set_cell_border(cell, **kwargs):
    """Helper to set cell borders (simplified)"""
    pass  # Skip detailed border styling for simplicity

def add_heading_zh(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def add_paragraph_zh(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p

def add_table_zh(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table

def main():
    os.makedirs(DOCS_DIR, exist_ok=True)
    doc = Document()

    # ── Global style ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ═══════════════════════════════════════════
    # Cover / Title
    # ═══════════════════════════════════════════
    title = doc.add_heading('EMS 平台项目 — 第1个月总结报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('完成日期：2026-06-11 | 作者：SE\n项目周期：6个月（2026-04 ~ 2026-10）')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacer

    # ═══════════════════════════════════════════
    # 1. 项目概述
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '1. 项目概述', level=1)
    add_paragraph_zh(doc,
        '本项目围绕燃料电池（FC）混合动力汽车的能量管理系统（EMS）展开，'
        '目标是在6个月内系统学习传统控制策略与强化学习算法，'
        '并完成一个完整的EMS算法对比项目，为秋招储能/车企算法岗位做准备。'
    )
    add_paragraph_zh(doc,
        '定位：A（EMS/BMS算法）为主，C（AI+新能源）为辅。\n'
        '就业面：储能/车企/电池/Tier1（最广）+ AI能源交叉（增长型）。'
    )

    # ═══════════════════════════════════════════
    # 2. 第1个月目标与完成情况
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '2. 第1个月目标与完成情况', level=1)
    add_paragraph_zh(doc,
        '第1个月主题：工程底座搭建 + DP（动态规划）入门与深度分析。\n'
        '计划完成4周内容，实际全部达成。'
    )

    add_heading_zh(doc, '2.1 每周任务回顾', level=2)

    add_heading_zh(doc, 'Week 1-2: 工程底座搭建', level=3)
    add_paragraph_zh(doc,
        '• 环境配置：Python 3.13 + numpy/pandas/matplotlib + MATLAB R2024b/Simulink\n'
        '• 车辆动力学模型实现（三力模型：滚动阻力 + 空气阻力 + 惯性力）\n'
        '• 电池模型：OCV-R_int 等效电路模型 + SOC 状态转移\n'
        '• FC 效率曲线：查表 + 氢耗计算（LHV_H2 = 120 MJ/kg）\n'
        '• 工况数据：WLTC / NEDC / CLTC 下载与预处理\n'
        '• 产出：scripts/download_drive_cycles.py + 工况CSV文件'
    )

    add_heading_zh(doc, 'Week 3: 规则控制器 + DP原理学习', level=3)
    add_paragraph_zh(doc,
        '• 规则控制器实现：基于SOC区间的启发式策略\n'
        '• DP理论学习：Bellman最优性原理、后向递归、前向Rollout\n'
        '• DP核心算法实现：backward_dp() + forward_rollout()\n'
        '• DP vs 规则控制器对比：WLTC工况氢耗降低19.2%\n'
        '• 产出：scripts/day8_dp_ems.py (501行) + 对比可视化'
    )

    add_heading_zh(doc, 'Week 4: DP深度分析', level=3)
    add_paragraph_zh(doc,
        '• 参数敏感性分析：α（SOC维持惩罚）、β（终端惩罚）、网格密度\n'
        '• 三工况补充：CLTC数据下载 + WLTC/NEDC/CLTC 三工况对比\n'
        '• 结果量化：α/β/网格密度对氢耗和SOC终值的影响\n'
        '• 练习版脚本：exercise_day8_dp_ems.py（5个填空题，自测用）\n'
        '• 产出：敏感性分析报告 + 全景对比图 + 原始数据CSV'
    )

    # ═══════════════════════════════════════════
    # 3. 核心成果数据
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '3. 核心成果数据', level=1)

    add_heading_zh(doc, '3.1 DP vs 规则控制器 — WLTC工况', level=2)
    add_paragraph_zh(doc, 'WLTC（ Worldwide Harmonized Light Vehicle Test Cycle，1800秒）：')

    add_table_zh(doc,
        ['指标', '规则控制器', 'DP', '改善率'],
        [
            ['总氢耗 (kg)', '0.2831', '0.2287', '↓19.2% ✅'],
            ['SOC 初值→终值', '0.60 → 0.614', '0.60 → 0.574', '偏差4.3%'],
            ['FC 平均效率', '37.8%', '45.7%', '+7.9 pp'],
            ['FC 高效(>50%)占比', '20.8%', '40.5%', '翻倍 (+19.7 pp)'],
            ['FC 最大功率 (kW)', '25.0', '25.0', '—'],
            ['总能量需求 (kWh)', '4.01', '4.01', '—'],
        ]
    )

    add_heading_zh(doc, '3.2 三工况DP基准对比', level=2)
    add_table_zh(doc,
        ['工况', '时长', '能量需求', 'DP氢耗', '改善率(vs规则)', 'SOC终值'],
        [
            ['WLTC', '1800s', '4.01 kWh', '0.2287 kg', '19.2%', '0.574'],
            ['NEDC', '1181s', '1.66 kWh', '0.0990 kg', '31.4%', '0.574'],
            ['CLTC', '1800s', '2.11 kWh', '0.1448 kg', '30.2%', '0.575'],
        ]
    )
    add_paragraph_zh(doc,
        '注：NEDC改善率最高（31.4%），因其低速段多，规则控制器效率低，DP优化空间大。'
    )

    add_heading_zh(doc, '3.3 敏感性分析结论', level=2)
    add_table_zh(doc,
        ['参数', '测试范围', '结论'],
        [
            ['α (SOC维持惩罚)', '10, 50, 100, 200, 500', '10~500范围内结果稳定，默认100合理'],
            ['β (终端惩罚)', '1000, 5000, 10000, 50000, 100000', '全范围对氢耗影响<1%，默认10000合理'],
            ['网格密度', '50×20, 100×40, 150×60, 200×80', '密度增加结果收敛，150×60精度/速度最佳'],
        ]
    )

    # ═══════════════════════════════════════════
    # 4. 技术栈与知识点
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '4. 技术栈与知识点掌握', level=1)

    add_heading_zh(doc, '4.1 Python工程', level=2)
    add_table_zh(doc,
        ['技能', '掌握程度', '应用场景'],
        [
            ['NumPy向量化计算', '⭐⭐⭐⭐⭐', 'DP内层循环向量化、状态转移批量计算'],
            ['Pandas数据处理', '⭐⭐⭐⭐', '工况CSV读写、结果分析'],
            ['Matplotlib可视化', '⭐⭐⭐⭐', '五合一对比图、敏感性分析图'],
            ['argparse CLI', '⭐⭐⭐⭐', 'scripts/dp_ems.py --cycle nedc'],
            ['项目结构组织', '⭐⭐⭐⭐', 'scripts/ docs/ results/ 三层分离'],
        ]
    )

    add_heading_zh(doc, '4.2 控制理论', level=2)
    add_table_zh(doc,
        ['概念', '理解深度', '应用'],
        [
            ['Bellman最优性原理', '⭐⭐⭐⭐⭐', 'DP后向递推基础'],
            ['动态规划（DP）', '⭐⭐⭐⭐⭐', '全局最优求解，后向+前向'],
            ['等效消耗最小化（ECMS）', '⭐⭐⭐ 学习中', '第2个月主题'],
            ['模型预测控制（MPC）', '⭐⭐⭐ 学习中', '第2个月主题'],
            ['规则控制器设计', '⭐⭐⭐⭐', 'SOC区间分段启发式策略'],
        ]
    )

    add_heading_zh(doc, '4.3 车辆系统建模', level=2)
    add_paragraph_zh(doc,
        '• 车辆纵向动力学：三力模型（F_rr + F_aero + F_inertia），中心差分求加速度\n'
        '• 电池等效电路模型：OCV查表 + R_int内阻 + SOC状态转移\n'
        '• FC效率曲线：查表插值 + 氢耗计算（P_fc → mdot_H2）\n'
        '• 功率分配逻辑：P_bat = P_load - P_fc，SOC约束闭环'
    )

    # ═══════════════════════════════════════════
    # 5. 代码资产清单
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '5. 代码资产清单', level=1)

    add_table_zh(doc,
        ['文件', '行数', '功能', '状态'],
        [
            ['scripts/day8_dp_ems.py', '~512', 'DP完整实现（后向+前向+规则对比+绘图）', '✅ 完成'],
            ['scripts/exercise_day8_dp_ems.py', '~510', 'DP练习版（5个填空题）', '✅ 完成'],
            ['scripts/gen_sensitivity_report.py', '~100', '敏感性分析报告生成', '✅ 完成'],
            ['scripts/download_drive_cycles.py', '~60', 'WLTC/NEDC/CLTC数据下载', '✅ 完成'],
            ['scripts/analyze_results.py', '~88', '结果分析脚本（终端打印）', '✅ 完成'],
            ['results/dp_ems_wltc.csv', '—', 'WLTC DP结果数据', '✅ 完成'],
            ['results/dp_ems_nedc.csv', '—', 'NEDC DP结果数据', '✅ 完成'],
            ['results/DP_vs_Rule_wltc.png', '—', 'WLTC五合一对比图', '✅ 完成'],
            ['results/DP_vs_Rule_nedc.png', '—', 'NEDC五合一对比图', '✅ 完成'],
            ['results/DP_sensitivity_analysis.png', '—', '敏感性分析全景图', '✅ 完成'],
            ['docs/Day8_DP成果评估报告.docx', '—', 'DP成果评估文档', '✅ 完成'],
            ['docs/Day8_DP参数敏感性分析报告.docx', '—', '敏感性分析报告', '✅ 完成'],
            ['docs/Day8_dp_ems代码说明.docx', '—', 'DP代码逐段解释', '✅ 完成'],
            ['docs/Day8_后向DP核心逻辑通俗解释.docx', '—', 'DP原理通俗版', '✅ 完成'],
        ]
    )

    # ═══════════════════════════════════════════
    # 6. 面试表达准备（全4周八股文）
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '6. 面试表达准备（全4周）', level=1)

    add_heading_zh(doc, '6.1 DP项目三段式叙事', level=2)
    add_paragraph_zh(doc, '【30秒版 — HR面/群面】')
    add_paragraph_zh(doc,
        '「我用动态规划为燃料电池混动系统优化能量管理，'
        '在WLTC工况下实现19.2%的氢耗降低，SOC维持稳定。」'
    )
    add_paragraph_zh(doc, '【2分钟版 — 技术面】')
    add_paragraph_zh(doc,
        '「我基于Bellman最优性原理设计了一个后向DP求解器，'
        'SOC网格150×60的状态空间，通过后向递归计算最优代价函数、'
        '前向Rollout得到最优控制序列。与规则控制器对比，'
        'WLTC工况氢耗从0.283kg降到0.229kg，降低19.2%。'
        '同时FC高效区间(>50%)占比从20.8%提升到40.5%，翻了一倍。'
        'SOC终值0.574，偏差仅4.3%。用的是Python numpy手写，'
        '不依赖第三方优化库，代码501行。」'
    )
    add_paragraph_zh(doc, '【5分钟版 — 深挖】')
    add_paragraph_zh(doc,
        '• 为什么要用DP？——全局最优 vs 局部最优\n'
        '• DP的局限？——需要已知工况、计算量大、无法在线\n'
        '• SOC惩罚系数怎么调的？——α=100的调参过程\n'
        '• 如果工况变了会怎样？——NEDC也做了，对比结果说明泛化性'
    )

    add_heading_zh(doc, '6.2 八股文精选（12题，覆盖4周全部知识点）', level=2)
    add_paragraph_zh(doc, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    add_paragraph_zh(doc, '【Week 1-2 工程底座篇】（4题）', bold=True)
    add_paragraph_zh(doc, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    add_paragraph_zh(doc,
        'Q1: 燃料电池混动车辆的功率分配逻辑是什么？\n'
        'A: 总功率需求 P_load 由燃料电池 P_fc 和电池 P_bat 分担：'
        'P_bat = P_load - P_fc。目标是使 P_fc 工作在高效区间，'
        '同时维持 SOC 在合理范围（如 0.2~0.9）。电池起"削峰填谷"作用。\n\n'
        'Q2: 为什么用三力模型而不是更复杂的车辆动力学模型？\n'
        'A: 三力模型（滚动阻力 + 空气阻力 + 惯性力）已经能准确估计 '
        'WLTC/NEDC 等标准工况的功率需求。更复杂的模型（如轮胎模型、'
        '传动系统模型）对能量管理优化影响不大，反而增加计算负担。'
        '在 1Hz 采样下，中心差分求加速度足够精确。\n\n'
        'Q3: 电池的 OCV-R_int 等效电路模型原理？\n'
        'A: 电池开路电压 OCV 通过 SOC 查表得到（OCV_LU 数组）。'
        '实际端电压 V = OCV - I×R_int。由功率 P_bat = V×I 联立求解电流 I：'
        'R_int×I² - OCV×I + P_bat = 0。取物理可行的根（Δ=OCV²-4×R_int×P_bat ≥ 0）。'
        '然后 SOC_{k+1} = SOC_k - I/(Q_bat×3600)×dt。\n\n'
        'Q4: WLTC / NEDC / CLTC 三种工况有什么区别？\n'
        'A: WLTC（1800s，全球统一轻型车测试循环，涵盖高速段）；'
        'NEDC（1181s，旧欧洲标准，以低速为主）；CLTC（1800s，中国工况，'
        '加减速频繁）。三种工况覆盖了不同的驾驶风格，用于验证EMS策略的泛化能力。'
    )

    add_paragraph_zh(doc, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    add_paragraph_zh(doc, '【Week 3 DP入门篇】（4题）', bold=True)
    add_paragraph_zh(doc, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    add_paragraph_zh(doc,
        'Q5: 动态规划（DP）求解EMS问题的基本思路？\n'
        'A: 将连续时间和状态离散化，构造状态空间（SOC网格）和控制空间（FC功率网格）。'
        '用Bellman最优性原理从终点反向递推代价函数 J[k][i]，'
        '得到每个时刻每个SOC状态的最优FC功率策略 π[k][i]。'
        '然后前向Rollout仿真得到实际轨迹。时间复杂度 O(N×S×A)，'
        '其中 N=1800（时刻数），S=150（SOC网格），A=60（FC功率网格）。\n\n'
        'Q6: 后向DP和前向DP有什么区别？为什么用后向？\n'
        'A: 后向DP从终点往回算，利用Bellman原理，每个状态只存最优后续决策，'
        '天然适合求全局最优。前向DP需要枚举所有路径，复杂度指数增长。'
        '前向Rollout只是查表仿真，不是求解过程。\n\n'
        'Q7: DP中的SOC惩罚系数α和终端惩罚β分别起什么作用？\n'
        'A: α 惩罚SOC偏离参考值（如0.6）的即时代价，α越大SOC越稳定；'
        'β 惩罚终点SOC与参考值的偏差，β越大终点SOC越接近目标值。'
        '两者共同确保SOC在整个循环中维持在合理区间。\n\n'
        'Q8: 为什么DP结果优于规则控制器？\n'
        'A: 规则控制器是基于启发式的局部策略（如SOC低时FC多发），'
        '无法预见未来的功率需求。DP通过全局优化，可以在合适的时候让电池放电、'
        'FC工作在高效区，从而在保证SOC约束的前提下最小化总氢耗。'
        '实测WLTC改善19.2%，FC高效区间占比从20.8%提升到40.5%。'
    )

    add_paragraph_zh(doc, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    add_paragraph_zh(doc, '【Week 4 DP深度篇】（4题）', bold=True)
    add_paragraph_zh(doc, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    add_paragraph_zh(doc,
        'Q9: DP参数敏感性分析的结果说明了什么？\n'
        'A: α在10~500范围内、β在1000~100000范围内，DP结果（氢耗、SOC终值）'
        '变化都<1%，说明DP对这两个惩罚系数不敏感，默认值鲁棒。'
        '网格密度从50×20到200×80结果收敛，150×60是精度和速度的最佳平衡点。\n\n'
        'Q10: 三种工况下DP改善率不同（WLTC 19.2% < CLTC 30.2% < NEDC 31.4%），为什么？\n'
        'A: NEDC改善率最高因为其以低速段为主，规则控制器在低功率区间效率很低'
        '（FC频繁启停、工作在低效区），DP的优化空间更大。'
        'WLTC包含高速段，规则控制器在高功率下也能保持一定效率，DP优势相对较小。\n\n'
        'Q11: DP的局限性和在实际EMS中的应用限制？\n'
        'A: ①需要已知完整工况（offline方法），无法应对实时变化的驾驶条件；'
        '②计算量大（O(N×S×A)），复杂系统实时性差；'
        '③状态空间离散化引入精度损失；④多目标优化（氢耗+FC寿命+舒适性）难以直接扩展。'
        '因此DP主要用于benchmark对比，实际在线EMS用ECMS/MPC/RL。\n\n'
        'Q12: 你的DP实现有什么特色？\n'
        'A: ①纯Python+numpy手写，不依赖cvxpy/casadi等优化库；'
        '②内层循环向量化（一次计算所有PFC_GRID的SOC_next），兼顾可读性和速度；'
        '③包含规则控制器作为baseline；④支持WLTC/NEDC/CLTC三工况一键切换；'
        '⑤参数化设计，方便调参实验。'
    )

    # ═══════════════════════════════════════════
    # 7. 经验总结与反思
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '7. 经验总结与反思', level=1)

    add_heading_zh(doc, '7.1 做得好的地方', level=2)
    add_paragraph_zh(doc,
        '✓ 完整实现了DP前后向算法，不依赖第三方优化库，代码自包含\n'
        '✓ 三工况（WLTC/NEDC/CLTC）全覆盖，验证了泛化性\n'
        '✓ 参数敏感性分析严谨（α/β/网格密度），结论有说服力\n'
        '✓ 练习版脚本（exercise版）设计巧妙，5个填空自测\n'
        '✓ 代码注释详细（501行含中文注释），适合面试讲解\n'
        '✓ 面试叙事提前准备（30秒/2分钟/5分钟三版本）'
    )

    add_heading_zh(doc, '7.2 需要改进的地方', level=2)
    add_paragraph_zh(doc,
        '△ SOC终端0.574与参考值0.6有微小偏差 → 后续可调β或增加终端约束\n'
        '△ 部分区间FC功率0 → 频繁启停对FC寿命不利（需加启停惩罚）\n'
        '△ 目前无再生制动 → 下阶段可加入制动能量回收模型\n'
        '△ 电池模型较简化（R_int等效电路）→ 后续可考虑更精确的Thevenin模型'
    )

    add_heading_zh(doc, '7.3 下阶段重点', level=2)
    add_paragraph_zh(doc,
        '→ 第2个月：ECMS原理与实现 + C++基础\n'
        '→ 重点掌握等效因子自适应调优\n'
        '→ 继续强化Python→C++的过渡能力\n'
        '→ 秋招投递策略：8-10月边学边投，不赶进度'
    )

    # ═══════════════════════════════════════════
    # 8. 里程碑时间线
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '8. 里程碑时间线', level=1)

    add_table_zh(doc,
        ['日期', '里程碑', '状态'],
        [
            ['2026-04', '项目启动，环境搭建', '✅'],
            ['2026-04 中旬', 'Week 1-2: 工程底座完成', '✅'],
            ['2026-04 底', 'Week 3: DP实现 + 规则对比', '✅'],
            ['2026-05 底', 'Week 4: 敏感性分析 + 三工况', '✅'],
            ['2026-05-29', 'Month 1总结报告生成', '✅'],
            ['2026-06 中旬', 'Week 5: ECMS开始', '⏳ 即将开始'],
            ['2026-07 底', 'Week 8: 传统EMS四方法对比', '📌 可投递'],
            ['2026-08 底', 'Week 12: RL-EMS完成', '📌 简历强度够'],
            ['2026-09-10', '秋招投递期', '📌 目标30-50家'],
        ]
    )

    # ═══════════════════════════════════════════
    # 9. 附录：关键文件索引
    # ═══════════════════════════════════════════
    add_heading_zh(doc, '9. 附录：关键文件索引', level=1)
    add_paragraph_zh(doc,
        '代码：scripts/day8_dp_ems.py | scripts/exercise_day8_dp_ems.py\n'
        '数据：results/dp_ems_wltc.csv | results/dp_ems_nedc.csv\n'
        '图表：results/DP_vs_Rule_wltc.png | results/DP_sensitivity_analysis.png\n'
        '文档：docs/Day8_DP成果评估报告.docx | docs/Day8_DP参数敏感性分析报告.docx\n'
        '模板：ems-plan/status.md（进度跟踪）'
    )

    # ── Save ──
    out_path = os.path.join(DOCS_DIR, 'Month1_总结报告.docx')
    doc.save(out_path)
    print(f'[OK] 已生成: {out_path}')

if __name__ == '__main__':
    main()
