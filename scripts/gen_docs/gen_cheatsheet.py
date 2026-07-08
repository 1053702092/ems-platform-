# -*- coding: utf-8 -*-
"""生成 C++ 速查表 + Python 重点补漏 .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DOCS_DIR = r'F:\CLAUDE\research\ems-platform\docs'

def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def p(doc, text, bold=False, size=10, mono=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if mono:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    else:
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return para

def tbl(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return table

def code_block(doc, lines):
    """Add code block with Consolas font"""
    for line in lines:
        p(doc, line, mono=True)


def main():
    os.makedirs(DOCS_DIR, exist_ok=True)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ================================================================
    # 封面
    # ================================================================
    title = doc.add_heading('C++ + Python 速查手册', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('EMS 算法岗专用 | 只看面试/工程需要的\n'
                       'C++ 部分：看懂代码 + LeetCode Easy | Python 部分：补 NumPy 基础')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # ================================================================
    # PART 1: C++ 速查
    # ================================================================
    h(doc, 'Part 1: C++ 速查手册', level=1)
    p(doc, '目标：看懂 EMS/C++ 代码，能写 LeetCode Easy 题。\n'
           '学习顺序：从左到右，按需查阅。')

    # 1.1 基础语法对照
    h(doc, '1.1 Python → C++ 语法对照', level=2)
    p(doc, '最实用的对照表——把你熟悉的 Python 翻译成 C++：')
    tbl(doc,
        ['功能', 'Python', 'C++'],
        [
            ['注释', '# 这是注释', '// 这是注释\n/* 多行注释 */'],
            ['变量声明', 'x = 5', 'int x = 5;'],
            ['浮点数', 'y = 3.14', 'double y = 3.14;'],
            ['布尔值', 'flag = True', 'bool flag = true;'],
            ['字符串', 'name = "hello"', 'string name = "hello";'],
            ['打印', 'print(x)', 'cout << x << endl;'],
            ['条件判断', 'if x > 0:\n    pass', 'if (x > 0) {\n    // ...\n}'],
            ['循环 for', 'for i in range(10):', 'for (int i = 0; i < 10; i++)'],
            ['循环 while', 'while x > 0:', 'while (x > 0) { ... }'],
            ['函数定义', 'def f(x):\n    return x*2', 'double f(double x) {\n    return x * 2;\n}'],
            ['列表', 'lst = [1,2,3]', 'vector<int> lst = {1,2,3};'],
            ['字典', 'd = {"a":1}', 'map<string,int> d;\nd["a"] = 1;'],
            ['列表追加', 'lst.append(4)', 'lst.push_back(4);'],
            ['列表长度', 'len(lst)', 'lst.size()'],
            ['导入模块', 'import numpy', '#include <vector>\nusing namespace std;'],
        ])

    # 1.2 常用类型
    h(doc, '1.2 常用数据类型', level=2)
    tbl(doc,
        ['类型', '说明', '示例', '对应 Python'],
        [
            ['int', '整数', 'int x = 42;', 'x = 42'],
            ['double', '双精度浮点', 'double d = 3.14;', 'd = 3.14'],
            ['float', '单精度浮点', 'float f = 1.0f;', '（少见）'],
            ['bool', '布尔', 'bool ok = true;', 'ok = True'],
            ['char', '单个字符', 'char c = \'A\';', 'c = "A"[0]'],
            ['string', '字符串', 'string s = "abc";', 's = "abc"'],
            ['vector<T>', '动态数组（最常用！）', 'vector<double> v;', 'list / np.array'],
            ['map<K,V>', '字典/哈希表', 'map<string,int> m;', 'dict'],
            ['set<T>', '去重集合', 'set<int> s;', 'set()'],
            ['pair<K,V>', '二元组', 'pair<int,double> p;', 'tuple'],
        ])

    # 1.3 vector 常用操作（面试高频）
    h(doc, '1.3 vector 常用操作（面试必考）', level=2)
    code_block(doc, [
        'vector<double> v;              // 创建空数组',
        'v.push_back(1.0);              // 尾部追加',
        'v.pop_back();                  // 尾部删除',
        'v.size();                      // 元素个数',
        'v.empty();                     // 是否为空',
        'v[0]; v.at(0);                 // 访问（at 会边界检查）',
        'v.front(); v.back();           // 首/尾元素',
        'v.clear();                     // 清空',
        '',
        '// 初始化',
        'vector<int> a(10);             // 10个0',
        'vector<int> b(10, 5);          // 10个5',
        'vector<int> c = {1,2,3,4,5};   // 列表初始化',
        '',
        '// 遍历',
        'for (int i = 0; i < v.size(); i++) { cout << v[i]; }',
        'for (auto x : v) { cout << x; }  // 范围for（C++11）',
        '',
        '// 排序',
        'sort(v.begin(), v.end());      // 升序',
        'sort(v.begin(), v.end(), greater<int>());  // 降序',
        '',
        '// 查找',
        'auto it = find(v.begin(), v.end(), 42);   // 返回迭代器',
        'if (it != v.end()) { /* 找到了 */ }',
    ])

    # 1.4 map 常用操作
    h(doc, '1.4 map 常用操作', level=2)
    code_block(doc, [
        'map<string, double> m;         // 创建字典',
        'm["key"] = 3.14;               // 插入/修改',
        'm.count("key");                // 是否存在（0 或 1）',
        'm.find("key");                 // 返回迭代器',
        'm.erase("key");                // 删除',
        'm.size();                      // 元素个数',
        '',
        '// 遍历',
        'for (auto& kv : m) {',
        '    cout << kv.first << " = " << kv.second << endl;',
        '}',
        '',
        '// 判断 key 是否存在（推荐）',
        'if (m.count("key")) { /* exists */ }',
    ])

    # 1.5 函数
    h(doc, '1.5 函数定义', level=2)
    code_block(doc, [
        '// 基本形式',
        '返回类型 函数名(参数列表) {',
        '    // 函数体',
        '    return 值;',
        '}',
        '',
        '// 示例：计算 FC 氢耗',
        'double fcHydrogenFlow(double P_fc) {',
        '    // P_fc [kW], eta_fc [效率], LHV [J/kg]',
        '    double eta_fc = 0.50;',
        '    double LHV_H2 = 120e6;',
        '    if (P_fc < 0.01) return 0.0;',
        '    double mdot = P_fc * 1000 / (eta_fc * LHV_H2) * 1000;  // g/s',
        '    return mdot;',
        '}',
        '',
        '// 传引用（可以修改传入的变量）',
        'void updateSOC(double& soc, double delta) {',
        '    soc += delta;  // 直接修改调用者的变量',
        '}',
        '',
        '// 传指针（另一种修改方式）',
        'void updateSOC_ptr(double* soc, double delta) {',
        '    *soc += delta;',
        '}',
    ])

    # 1.6 类基础
    h(doc, '1.6 类（Class）基础', level=2)
    code_block(doc, [
        'class Battery {',
        'private:',
        '    double soc;           // 私有成员',
        '    double capacity;      // 容量 Ah',
        '',
        'public:',
        '    // 构造函数',
        '    Battery(double init_soc, double cap)',
        '        : soc(init_soc), capacity(cap) {}',
        '',
        '    // 成员函数',
        '    void setSOC(double s) { soc = s; }',
        '    double getSOC() const { return soc; }',
        '',
        '    // 状态转移',
        '    void step(double P_bat_kW, double dt) {',
        '        // ... SOC 更新逻辑',
        '    }',
        '};',
        '',
        '// 使用',
        'Battery bat(0.6, 50.0);     // SOC=0.6, 50Ah',
        'bat.step(10.0, 1.0);        // 10kW 放电 1 秒',
        'cout << bat.getSOC();       // 读取 SOC',
    ])

    # 1.7 LeetCode Easy 模板
    h(doc, '1.7 LeetCode C++ 模板（Easy 题通用）', level=2)
    code_block(doc, [
        '#include <bits/stdc++.h>     // 万能头文件（竞赛/刷题用）',
        'using namespace std;',
        '',
        '// 示例：Two Sum（两数之和）',
        'vector<int> twoSum(vector<int>& nums, int target) {',
        '    unordered_map<int, int> mp;  // 值→索引',
        '    for (int i = 0; i < nums.size(); i++) {',
        '        int need = target - nums[i];',
        '        if (mp.count(need)) return {mp[need], i};',
        '        mp[nums[i]] = i;',
        '    }',
        '    return {};',
        '}',
        '',
        '// 示例：Reverse String',
        'void reverseString(vector<char>& s) {',
        '    int i = 0, j = s.size() - 1;',
        '    while (i < j) {',
        '        swap(s[i++], s[j--]);',
        '    }',
        '}',
    ])

    # 1.8 常见坑
    h(doc, '1.8 C++ 常见坑（面试/写代码时注意）', level=2)
    tbl(doc,
        ['坑', '说明', '正确写法'],
        [
            ['数组越界', 'v[10] 不检查边界，越界会 crash', '用 v.at(10) 或先判断 size'],
            ['未初始化', 'int x; 值是随机的', 'int x = 0;'],
            ['传值 vs 传引用', 'void f(vector<int> v) 是拷贝！', 'void f(const vector<int>& v)'],
            ['指针空悬', 'delete 后还使用指针', 'delete 后置 nullptr'],
            ['size_t vs int', 'v.size() 返回 unsigned', 'for (int i = 0; i < (int)v.size(); i++)'],
            ['浮点比较', '0.1 + 0.2 != 0.3', '用 abs(a-b) < 1e-6'],
            ['string vs char*', 'C 风格字符串容易溢出', '优先用 std::string'],
        ])

    # ================================================================
    # PART 2: Python 重点补漏
    # ================================================================
    doc.add_page_break()
    h(doc, 'Part 2: Python 重点补漏', level=1)
    p(doc, '针对 EMS 项目代码中用到的高频 Python 特性，逐个击破。')

    # 2.1 NumPy 核心
    h(doc, '2.1 NumPy 核心（项目最常用）', level=2)
    p(doc, '你的代码里 80% 的 Python 就是 NumPy，把这几个吃透：')
    tbl(doc,
        ['功能', 'Python 写法', '说明'],
        [
            ['创建数组', 'np.array([1,2,3])', '从 list 创建 ndarray'],
            ['等差数列', 'np.linspace(0, 30, 60)', '60个点，0到30等分'],
            ['全零/全一', 'np.zeros(10) / np.ones(10)', '初始化用'],
            ['数组运算', 'a + b, a * 2, np.sqrt(a)', '元素级运算，不需要循环'],
            ['索引', 'a[0], a[1:5], a[a>0]', '支持切片和布尔索引'],
            ['求和/均值', 'np.sum(a), np.mean(a)', '聚合函数'],
            ['插值', 'np.interp(x, xp, fp)', '查表插值（SOC→OCV 用这个）'],
            ['裁剪', 'np.clip(a, 0, 30)', '限制范围'],
            ['条件替换', 'np.where(cond, a, b)', '条件选择'],
            ['布尔掩码', 'mask = a > 5; a[mask]', '筛选元素'],
            ['向量化', 'a * 1000 / (b + 1e-9)', '整数组运算，不用 for 循环'],
        ])
    p(doc, '重点练习：把下面 C 风格的 for 循环改成 NumPy 向量化：')
    code_block(doc, [
        'import numpy as np',
        '',
        '# ❌ 慢：Python for 循环',
        'a = [1, 2, 3, 4, 5]',
        'b = []',
        'for x in a:',
        '    b.append(x * 2 + 1)',
        '',
        '# ✅ 快：NumPy 向量化',
        'a = np.array([1, 2, 3, 4, 5])',
        'b = a * 2 + 1',
        '',
        '# 复杂条件',
        'result = np.where(a > 3, a * 2, a + 1)',
        '# 等价于: a>3 的元素乘2，否则加1',
    ])

    # 2.2 Pandas 基础
    h(doc, '2.2 Pandas 基础（数据处理）', level=2)
    code_block(doc, [
        'import pandas as pd',
        '',
        '# 读取 CSV',
        'df = pd.read_csv("data.csv")',
        '',
        '# 取列',
        'times = df["time"].values       # 转 numpy 数组',
        'speeds = df["speed_kmh"].values',
        '',
        '# 写 CSV',
        'df.to_csv("output.csv", index=False)',
        '',
        '# 基本统计',
        'df.describe()                    # 均值/最大/最小',
        'df["col"].mean()                 # 列均值',
        'df["col"].max()                  # 最大值',
    ])

    # 2.3 Python 函数
    h(doc, '2.3 Python 函数进阶', level=2)
    code_block(doc, [
        '# 基本函数',
        'def fc_hydrogen_flow(P_fc):',
        '    """FC 功率→氢耗 [g/s]"""',
        '    eta = np.interp(P_fc, PFC_EFF_BP, ETA_FC)',
        '    with np.errstate(divide="ignore", invalid="ignore"):',
        '        mdot = P_fc * 1000 / (eta * LHV_H2) * 1000',
        '    mdot[~np.isfinite(mdot)] = 0',
        '    mdot[P_fc == 0] = 0',
        '    return mdot',
        '',
        '# 默认参数',
        'def state_transition(SOC_k, P_fc, P_load_k, dt=1.0):',
        '    """dt 有默认值 1.0"""',
        '    ...',
        '',
        '# 返回多个值（Python 特色）',
        'def get_stats(data):',
        '    return data.mean(), data.max(), data.min()',
        '',
        'avg, mx, mn = get_stats(arr)    # 直接解包',
    ])

    # 2.4 Python 常见坑
    h(doc, '2.4 Python 常见坑', level=2)
    tbl(doc,
        ['坑', '说明', '正确做法'],
        [
            ['可变默认参数', 'def f(x=[]): x.append(1) — 列表默认值是共享的！', '用 None，内部再初始化'],
            ['整数除法', '3/2=1.5 ✅, 3//2=1 ❌', '需要浮点时用 /，不要用 //'],
            ['缩进', 'Python 用缩进表示代码块', '统一用 4 空格，不要 tab'],
            ['变量作用域', 'for 循环里的变量外面也能访问', '注意变量名不要冲突'],
            ['import *', 'from x import * 污染命名空间', 'import numpy as np'],
            ['深拷贝 vs 浅拷贝', 'b = a 只是引用，不是拷贝', 'import copy; b = copy.deepcopy(a)'],
            ['字符串 format', 'f"{x:.2f}" 比 "%s" % x 更现代', '优先用 f-string'],
        ])

    # 2.5 调试技巧
    h(doc, '2.5 Python 调试技巧（项目实战用）', level=2)
    code_block(doc, [
        '# 1. 打印调试',
        'print(f"[DEBUG] SOC={soc:.3f}, P_fc={pfc:.1f}, P_load={pload:.1f}")',
        '',
        '# 2. assert 断言（快速检查）',
        'assert SOC_MIN <= soc_next <= SOC_MAX, f"SOC out of range: {soc_next}"',
        '',
        '# 3. 查看变量类型',
        'print(type(x))',
        'print(x.shape)    # numpy 数组的形状',
        '',
        '# 4. 条件断点',
        'if k % 300 == 0:',
        '    print(f"k={k}/{N}, progress={pct:.0f}%")',
        '',
        '# 5. 计时',
        'import time',
        't0 = time.time()',
        '# ... 你的代码',
        'print(f"耗时: {time.time()-t0:.2f}s")',
    ])

    # 2.6 你的项目中的 Python 模式
    h(doc, '2.6 你的项目中常见的 Python 代码模式', level=2)
    p(doc, '把这几段读熟，你的代码就能完全看懂：')
    code_block(doc, [
        '# 模式1：向量化状态转移（day8_dp_ems.py 第 119 行）',
        'def state_transition(SOC_k, P_fc, P_load_k, dt=1.0):',
        '    is_scalar = np.isscalar(P_fc)',
        '    P_fc = np.atleast_1d(np.asarray(P_fc, dtype=float))',
        '    P_bat = P_load_k - P_fc',
        '    V_oc = np.interp(SOC_k, SOC_BP, OCV_LU)',
        '    SOC_next = np.full_like(P_fc, SOC_k)',
        '    mask = np.abs(P_bat) >= 0.01',
        '    if mask.any():',
        '        P_w = P_bat[mask] * 1000',
        '        Delta = V_oc**2 - 4 * R_INT * P_w',
        '        valid = Delta >= 0',
        '        idx = np.where(mask)[0][valid]',
        '        if len(idx) > 0:',
        '            I = (V_oc - np.sqrt(Delta[valid])) / (2 * R_INT)',
        '            I = np.clip(I, -300, 300)',
        '            SOC_next[idx] = SOC_k - I / (Q_BAT * 3600) * dt',
        '    return float(SOC_next[0]) if is_scalar else SOC_next',
        '',
        '# 模式2：查表插值',
        'eta = np.interp(P_fc, PFC_EFF_BP, ETA_FC)',
        '# PFC_EFF_BP = [0, 2, 5, 8, ...] 断点',
        '# ETA_FC      = [0, 0.28, 0.40, 0.48, ...] 对应值',
        '# np.interp 自动线性插值',
        '',
        '# 模式3：布尔掩码筛选',
        'feasible = (SOC_next >= SOC_MIN) & (SOC_next <= SOC_MAX)',
        'total[feasible] = g[feasible] + J_future',
        '',
        '# 模式4：argmin 找最优',
        'best_idx = np.argmin(total)',
        'J[k, i] = total[best_idx]',
        'pi[k, i] = PFC_GRID[best_idx]',
    ])

    # 2.7 Python 学习路径
    h(doc, '2.7 你的 Python 补漏学习路径', level=1)
    tbl(doc,
        ['优先级', '内容', '时间', '练习方式'],
        [
            ['P0 ⭐', 'NumPy 数组创建 + 向量化运算', '1天', '把 for 循环改成向量化'],
            ['P0 ⭐', 'np.interp / np.clip / np.where', '0.5天', '重读 day8_dp_ems.py'],
            ['P1', '函数定义 + 默认参数 + 返回值', '0.5天', '写3个简单函数'],
            ['P1', 'Pandas DataFrame 读写', '0.5天', '处理 CSV 数据'],
            ['P2', '类基础（class / __init__）', '1天', '写一个简单的 Battery 类'],
            ['P2', 'argparse CLI 用法', '0.5天', '看 day8_dp_ems.py 的 main()'],
            ['P3', '调试技巧（print / assert）', '0.5天', '在代码里加调试输出'],
        ])

    # ================================================================
    # Save
    # ================================================================
    out_path = os.path.join(DOCS_DIR, 'Cpp_Python_速查手册.docx')
    doc.save(out_path)
    print(f'[OK] 已生成: {out_path}')
    print(f'     大小: {os.path.getsize(out_path) / 1024:.0f} KB')

if __name__ == '__main__':
    main()
