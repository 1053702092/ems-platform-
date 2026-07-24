#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate docx: 深海多堆燃料电池系统工程因素全览 - 包含完整开题总结表"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Style ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── Helpers ──
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    return table

def bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ══════════════════════════════════════════════════
#   COVER
# ══════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(60)
r = t.add_run('深海多堆燃料电池系统\n工程因素全览')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('多目标优化框架下的工程因素映射 · 2026-07').font.size = Pt(12)
doc.add_page_break()

# ══════════════════════════════════════════════════
#  CHAPTER 1: 开题总结表 (核心)
# ══════════════════════════════════════════════════
doc.add_heading('一、四个优化目标到工程因素的全映射', level=1)
doc.add_paragraph(
    '本研究的四个优化目标（等效氢耗、功率波动频率、电堆启停次数、四堆效率偏差）'
    '分别映射到 7 个顶层工程维度。下表是开题总结对照表：'
)

add_table(doc,
    ['优化目标', '工程维度', '直接工程指标', '物理单位', '因素类型', '参考权重\n(深海)'],
    [
        ['1. 等效氢耗最小', '经济性',  '续航里程',               'km',           '目标', '★★★★☆'],
        ['',                '经济性',  '燃料成本',               '¥/任务',       '目标', ''],
        ['',                '储运',    '储氢/储氧系统体积',       'm³',       '约束', ''],
        ['',                '热管理',  '散热系统功耗占比',        '% of BOP',     '约束', ''],
        ['',                '安全性',  '储氢量→泄漏风险池',       '风险分',       '约束', ''],
        ['2. 功率波动最小', '耐久性',  '电堆机械/热寿命',         'h',            '目标', '★★★★★'],
        ['',                '耐久性',  'BOP 附件寿命',            'h',            '目标', ''],
        ['',                '声学',    '声辐射/振动特征',         'dB re 1µPa','目标(军用)', ''],
        ['',                '电气',    'DC/DC 可靠性 + EMC',      'MTBF/dBµV', '约束', ''],
        ['',                '动态',    '负载跟随响应时间',        's',            '约束', ''],
        ['3. 启停次数最少', '耐久性',  '碳腐蚀寿命(最大杀手)',    '次→h',       '目标', '★★★★★'],
        ['',                '安全性',  '任务可靠性(启动失败率)',   '[0,1]',        '约束', ''],
        ['',                '能耗',    '冷启动+待机功耗',         '% SOC/次',     '约束', ''],
        ['',                '消耗品',  '吹扫 N₂ 消耗量',     'L/次',         '约束', ''],
        ['4. 效率偏差最小', '维护性',  '堆间均衡老化→维护周期', 'h/更换',       '目标', '★★★★☆'],
        ['',                '安全性',  '故障预警灵敏度',           'h 提前量',     '目标', ''],
        ['',                '物流',    '备件库存策略复杂度',       '种类/年',      '约束', ''],
        ['',                '可扩展',  '能否推广到 6/8 堆',        '是/否',        '设计验证', ''],
    ],
    col_widths=[2.5, 1.5, 3.5, 2, 2, 1.8]
)

doc.add_paragraph()
doc.add_heading('二、各目标详解', level=1)

doc.add_heading('2.1 等效氢耗最小 → 续航 + 燃料成本 + 储氢体积', level=2)
doc.add_paragraph(
    '等效氢耗将电池净电耗折算为等效氢气消耗，统一度量后最小化。'
    '这是最直观的经济性指标——同样储氢量下跑得更远。'
    '在深海场景中，补氢几乎不可能，续航即任务半径。'
    '然而，单纯追求低氢耗往往需要频繁调整功率适配负载，与目标 2 和 3 冲突。'
    '此外，氢耗低 → 储氢需求小 → 储罐体积/重量减小 → 更多舱体空间留给有效载荷。'
)

doc.add_heading('2.2 功率波动频率最小 → 寿命 + 声学 + DC/DC 可靠性', level=2)
doc.add_paragraph('燃料电池“怕变不怕稳”。频繁的功率波动带来连锁损伤链：')
bullet(doc, '机械应力：MEA 反复膨胀收缩 → 膜电极机械疲劳')
bullet(doc, '气体饥饿：电流突变时供气滞后 → 反极腐蚀催化剂')
bullet(doc, '热循环：功率变化 → 温度梯度冲击 → 密封失效加速')
bullet(doc, 'DC/DC 变换器 IGBT 开关损耗随波动频次上升 → 键合线疲劳脱落')
bullet(doc, '声学：功率波动 = 电磁噪声 + 泵阀机械噪声频谱变化 (水下隐身穿透)')
doc.add_paragraph(
    '深海场景中热管理惯性更大，功率波动的危害被放大。'
    '典型约束：功率变化率 ≤ 3–5 kW/s。'
    '对军用潜航器而言，声学特征可能是比寿命更优先的指标。'
)

doc.add_heading('2.3 电堆启停次数最少 → 寿命 + 可靠性 + 辅助能耗', level=2)
doc.add_paragraph('启停是电堆衰减的最大来源 (50–60%)：')
bullet(doc, '启动阶段：低电压 + 高电位 → 碳载体腐蚀 (Carbon Corrosion)，Pt 催化剂脱落')
bullet(doc, '停机阶段：阳极残留 H₂ + O₂ → 氢-空界面 → 阳极高电位腐蚀')
bullet(doc, '数据：一次启停 ≈ 连续运行数小时的衰减量')
bullet(doc, '深海冷启动耗时 15–30 分钟，消耗 2–5% SOC，且存在启动失败 → 任务失败风险')
bullet(doc, '工程妥协：深海潜航器通常不关机，全程低功率 idle 保温')

doc.add_heading('2.4 四堆效率偏差最小 → 均衡老化 + 故障预警 + 维护经济性', level=2)
doc.add_paragraph('多堆系统特有的目标，核心是短板效应：')
bullet(doc, '最差的一个堆决定整个系统的输出能力')
bullet(doc, '均衡老化 → 各堆同时到寿、同时更换 → 维护成本大幅降低')
bullet(doc, '偏差扩大是密封渗漏/催化剂中毒/冷却不均的先兆信号')
bullet(doc, '良好偏差控制意味着控制策略可推广到 6/8/12 堆架构')

doc.add_page_break()

# ══════════════════════════════════════════════════
#  CHAPTER 3: 深海环境特殊性
# ══════════════════════════════════════════════════
doc.add_heading('三、深海环境的特殊性——附加工程考量', level=1)
doc.add_paragraph('地面车载的框架只是起点。深海环境引入了车载不存在的硬约束。')

doc.add_heading('3.1 高压环境 (2000m+ → 20MPa+)', level=2)
add_table(doc,
    ['工程问题', '物理机制', '后果', '应对措施'],
    [
        ['气体压缩性非线性', '高压下 H₂/O₂ 密度剧变', '流量控制精度下降', '温度/压力补偿模型'],
        ['MEA 压差失衡', '外静水压 vs 内气体压差超标', '膜破裂或密封失效', '压力补偿系统(≤1s)'],
        ['高压气液分离', '高压下气泡更小更难分离', '水淹 → 反极', '高压气液分离器'],
        ['密封蠕变加速', '持续高压 O-ring 蠕变', '海水渗入 → 短路', 'FFKM / 双道密封'],
    ],
    col_widths=[3.5, 4, 3.5, 4]
)

doc.add_heading('3.2 氧气供应策略', level=2)
doc.add_paragraph('深海没有空气 (21% O₂)，氧气来源分两路线：')
add_table(doc,
    ['维度', '方案 A: 携带纯氧', '方案 B: 海水溶解氧提取'],
    [
        ['实现', '高压氧罐 (液态/气态)', '膜接触器萃取海水溶解 O₂'],
        ['能量密度', '高 (体积占比 15–20%)', '低 (萃取能耗高)'],
        ['续航', '储氧量封顶', '理论上无限'],
        ['安全', '纯氧泄漏 → 助燃风险', '常压氧 → 安全'],
        ['成熟度', '成熟 (已工程化)', '前沿 (实验室)'],
        ['适用', '浅海<500m / 短期', '深海长航时 AUV (远期)'],
    ],
    col_widths=[2.5, 5.5, 6]
)

doc.add_heading('3.3 低温海水热管理 + 压力瞬变', level=2)
doc.add_paragraph('深海海水 2–4°C vs 电堆 60–80°C，大温差是双刃剑：')
bullet(doc, '散热效率极高 → 散热器体积可缩小', '✅ 优势：')
bullet(doc, '冷启动困难 → 需电加热或 H₂ 催化燃烧预热 (15–30 min)', '❌ 陷阱：')
bullet(doc, '局部过冷 → 冷凝水堆积 → 流道水淹', '❌ 陷阱：')
doc.add_paragraph(
    '浮潜过程压力瞬变：上浮/下潜时环境压力快速变化，'
    '>2 m/s 潜航速度产生 >20 kPa/s 压力变化率 → '
    '压力补偿器响应必须 ≤ 1s，否则 MEA 压差超标 → 膜破裂。'
)

doc.add_page_break()

# ══════════════════════════════════════════════════
#  CHAPTER 4: 约束清单
# ══════════════════════════════════════════════════
doc.add_heading('四、系统集成工程约束清单', level=1)
doc.add_paragraph('以下约束在优化模型中应作为硬边界（不参与 Pareto 优化）：')

add_table(doc,
    ['分类', '约束项', '典型值', '超出后果'],
    [
        ['性能', '电堆输出功率范围',   '20%–100% 额定', '<20% 水淹；>100% 热失控'],
        ['性能', '功率变化率',         '≤ 3–5 kW/s', '供气滞后 → 反极'],
        ['性能', '负载跟随时间',       '≤ 2 s',   '母线电压崩溃'],
        ['性能', '过载能力',           '120% (30s)',  '紧急避障不可用'],
        ['热管理', '电堆温度范围',     '60–80°C',  '低温低效；高温膜脱水'],
        ['热管理', '堆内温度均匀性',   '≤ 5°C',  '热点 → 非均匀老化'],
        ['热管理', '冷启动能耗',       '≤ 5% SOC', '影响续航余量'],
        ['安全', '阳极-阴极压差',      '≤ 50 kPa', 'MEA 机械破裂'],
        ['安全', '单电池最低电压',     '≥ 0.55 V', '反极 → 不可逆损伤'],
        ['安全', '舱内 H₂ 浓度',  '≤ 0.4% LFL','爆炸风险 (深海无逃生通道)'],
        ['安全', '绝缘电阻',           '≥ 1 MΩ', '海水渗入短路'],
        ['物理', '储氢量上限',         '任务前确定',    '不可逾越'],
        ['物理', '系统总重/体积',      '舱体容量决定',  '装不进去'],
        ['电气', '母线电压波动',       '≤ ±5%', '推进电机控制品质'],
        ['电气', 'EMC 电磁兼容',       'MIL-STD-461', '干扰声纳/通信'],
    ],
    col_widths=[1.8, 3.5, 3.5, 5.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
#  CHAPTER 5: 场景权重对比
# ══════════════════════════════════════════════════
doc.add_heading('五、优化权重的场景迁移', level=1)
doc.add_paragraph('地面车载 vs 深海潜航器的优先级显著不同。核心判断：深海 = 安全 > 寿命 > 效率。')

add_table(doc,
    ['优化目标', '地面车辆', '深海潜航器', '权重漂移原因'],
    [
        ['等效氢耗最小',   '★★★★★', '★★★★☆', '补氢不可能，但安全约束更硬'],
        ['功率波动最小',   '★★★☆☆', '★★★★★', '热管理惯性大 + 声学隐身要求'],
        ['启停最少',       '★★★★☆', '★★★★★', '冷启动能耗高 + 启动失败不可返回'],
        ['效率偏差最小',   '★★☆☆☆', '★★★★☆', '偏差是静默失效先行指标'],
    ],
    col_widths=[2.5, 2.5, 2.5, 6]
)

doc.add_paragraph()
doc.add_heading('5.1 额外工程维度的加权', level=2)
add_table(doc,
    ['工程维度', '地面车辆', '深海潜航器', '在优化框架中的位置'],
    [
        ['安全性 (泄漏/绝缘/压差)',   '★★★☆☆', '★★★★★', '硬约束，不参与 Pareto'],
        ['声学特征/振动噪声',         '☆☆☆☆☆', '★★★★★ (军用)', '军用: 第6目标; 民用: 可忽略'],
        ['堆内温度均匀性',            '★★★☆☆', '★★★★☆', '建议第 5 优化目标'],
        ['负载跟随响应速度',          '★★★★☆', '★★★★☆', '硬约束'],
        ['系统复杂度',                '★★★☆☆', '★★★☆☆', '设计约束'],
        ['可维护性 (模块化更换)',     '★★☆☆☆', '★★★★☆', '全生命周期成本影响'],
    ],
    col_widths=[4, 2.5, 2.5, 5.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
#  CHAPTER 6: 推荐优化框架
# ══════════════════════════════════════════════════
doc.add_heading('六、推荐的完整优化框架', level=1)
doc.add_paragraph('综合以上分析，建议将框架扩展如下：')

doc.add_heading('6.1 优化目标 (Pareto 集)', level=2)
add_table(doc,
    ['#', '目标', '类型', '隶属工程维度', '适用场景'],
    [
        ['1', '等效氢耗最小',              '必选', '经济性',           '所有场景'],
        ['2', '功率波动频率最小',          '必选', '耐久性+声学',     '所有场景'],
        ['3', '电堆启停次数最少',          '必选', '耐久性+安全性',   '所有场景'],
        ['4', '四堆效率偏差最小',          '必选', '维护性+预警',     '所有场景'],
        ['5', '堆内温度均匀性最大化',      '建议', '热管理',          '所有场景 (深海尤甚)'],
        ['6', '声学特征优化 (声压谱)',     '可选', '声学隐身',        '军用/敏感任务'],
    ],
    col_widths=[1, 3.5, 1.5, 2.5, 3.5]
)

doc.add_heading('6.2 硬约束 (不参与优化)', level=2)
add_table(doc,
    ['类别', '约束项', '数量'],
    [
        ['安全性', 'H₂ 浓度 / 绝缘 / 单池最低电压 / 压差', '4'],
        ['动态性', '负载跟随时间 / 过载能力 / 母线波动', '3'],
        ['热边界', '温度窗口 / 冷启动时间 / 温度均匀性上限', '3'],
        ['物理',  '储氢量 / 总重量 / 总体积', '3'],
        ['电气',  'EMC 限值 / DC/DC 效率下限', '2'],
    ],
    col_widths=[2, 10, 1.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
#  APPENDIX
# ══════════════════════════════════════════════════
doc.add_heading('附录：工程因素 → 优化数学模型映射', level=1)
doc.add_paragraph('将工程因素转化为数学优化模型的建议映射关系：')

add_table(doc,
    ['工程因素', '优化变量类型', '数学模型表示', '备注'],
    [
        ['续航里程',              '目标函数', 'J₁ = ∫ Ḥ₂_eq dt', '等效氢耗累计'],
        ['电堆寿命 (机械+碳腐蚀)',  '目标函数', 'J₂ = 1/(αN_start + β∫|dP/dt|dt)', '启停+波动联合建模'],
        ['系统可靠性',             '目标函数', 'J₃ = σ_eff⁻¹ · (1 − P_fail)', '偏差→预警'],
        ['堆内温度均匀性',         '目标(新增)', 'J₄ = max(T_ij) − min(T_ij)', '第 5 目标'],
        ['声学特征 (军用水下)',    '目标(军用)', 'J₅ = ∫ S(f) · W(f) df', '声压谱密×听觉加权'],
        ['功率范围',               '约束', 'P_min ≤ P ≤ P_max', '硬件'],
        ['功率变化率',             '约束', '|dP/dt| ≤ ΔP_max', '硬件'],
        ['负载跟随时间',           '约束', 't_response ≤ 2 s', '动态硬约束'],
        ['电堆温度窗口',           '约束', 'T_min ≤ T ≤ T_max', '硬件'],
        ['阳极-阴极压差',           '约束', '|P_an − P_ca| ≤ 50 kPa', '安全红线'],
        ['单电池最低电压',          '约束', 'V_cell ≥ 0.55 V', '安全红线'],
        ['舱内 H₂ 浓度',       '约束', 'C_H2 ≤ 0.4% LFL', '安全红线'],
        ['绝缘电阻',               '约束', 'R_ins ≥ 1 MΩ', '安全红线'],
        ['储氢量上限',             '约束', '∫ Ḥ_total dt ≤ V_H₂', '物理上限'],
        ['系统总重/体积',           '约束', 'W ≤ W_max, V ≤ V_max', '舱体限制'],
        ['过载能力',               '约束', 'P_max ≥ 120% · P_rated (30s)', '紧急避障'],
        ['母线电压波动',            '约束', 'ΔV_bus ≤ ±5%', '电机控制品质'],
        ['冷启动能耗',             '约束', 'E_start ≤ 5% SOC', '电池容量限制'],
    ],
    col_widths=[3, 2.5, 5.5, 3]
)

doc.add_paragraph()
doc.add_paragraph('— 文档结束 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──
out = r'F:\CLAUDE\research\ems-platform\docs\深海多堆FC系统工程因素全览.docx'
doc.save(out)
print(f'OK generated: {out}')
print(f'File size: {os.path.getsize(out)/1024:.1f} KB')
