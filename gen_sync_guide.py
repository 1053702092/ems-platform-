# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        r = p.add_run(t); r.font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

doc = Document()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('EMS-PLAN 双设备同步指南'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('跨设备切换 Claude Code 工作流程'); r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
doc.add_paragraph()

# Section 1
h1(doc, '一、设备切换清单')
h2(doc, '设备A 关机前（下班前）')
tx(doc, '顺序执行以下步骤：')
bl(doc, '打开 Git Bash，进入项目目录', bp='Step 1: ')
tx(doc, '  cd F:/CLAUDE/research/ems-platform')
bl(doc, '更新 STATUS.md 中的进度（修改勾选框）', bp='Step 2: ')
bl(doc, '提交代码变更', bp='Step 3: ')
tx(doc, '  git add -A && git commit -m "update: 进度说明"')
bl(doc, '同步 Claude 记忆到仓库', bp='Step 4: ')
tx(doc, '  python sync_memory.py --save')
bl(doc, '推送到 GitHub', bp='Step 5: ')
tx(doc, '  git push')
bl(doc, '关闭 Claude Code', bp='Step 6: ')
doc.add_paragraph()

h2(doc, '设备B 开机后（开始工作）')
tx(doc, '顺序执行以下步骤：')
bl(doc, '打开 Git Bash，进入项目目录', bp='Step 1: ')
tx(doc, '  cd <项目目录>')
bl(doc, '拉取最新代码', bp='Step 2: ')
tx(doc, '  git pull')
bl(doc, '恢复 Claude 记忆', bp='Step 3: ')
tx(doc, '  python sync_memory.py --load')
bl(doc, '打开 Claude Code，开始工作', bp='Step 4: ')
tx(doc, '  code .')
bl(doc, '跟我说：「继续 EMS-PLAN，看 STATUS.md」', bp='Step 5: ')
doc.add_paragraph()

# Section 2
h1(doc, '二、Git 同步说明')
t = doc.add_table(rows=6, cols=3)
t.style = 'Light Grid Accent 1'
data = [
    ('内容', '是否提交 Git', '原因'),
    ('代码 (.py/.m)', '✅ 提交', '核心工作产物'),
    ('文档 (.md/.docx)', '✅ 提交', '进度和设计记录'),
    ('Simulink 模型 (.slx)', '❌ 不提交', '二进制文件太大，手动拷贝'),
    ('Claude 记忆', '✅ 通过 sync_memory.py', '保存到 .claude-memory/ 目录'),
    ('实验数据 (.mat/.csv)', '❌ 不提交', '可重新生成'),
]
for i, (a, b, c) in enumerate(data):
    t.rows[i].cells[0].text = a; t.rows[i].cells[1].text = b; t.rows[i].cells[2].text = c
    for cell in t.rows[i].cells:
        for p2 in cell.paragraphs:
            for r2 in p2.runs: r2.font.size = Pt(9)
    if i == 0:
        for cell in t.rows[i].cells:
            shading(cell, '2F5496')
            for p2 in cell.paragraphs:
                for r2 in p2.runs: r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r2.bold = True
doc.add_paragraph()

tx(doc, '注意：.slx 文件已在 .gitignore 中排除。如果你把模型拷贝到新电脑，需要手动从原电脑拷贝 Energy.slx 到同样的目录位置。')
doc.add_paragraph()

# Section 3
h1(doc, '三、Claude 记忆同步说明')
tx(doc, 'Claude Code 的记忆文件位于系统目录：')
mem_path = 'C:\\Users\\10537\\.claude\\projects\\F--CLAUDE-research\\memory\\'
p = doc.add_paragraph()
r = p.add_run('  ' + mem_path); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
tx(doc, '这些文件记录了你的偏好、项目上下文和技术决策。跨设备同步方式：')
bl(doc, 'sync_memory.py --save：将系统记忆复制到仓库 .claude-memory/ 目录中', bp='保存：')
bl(doc, 'sync_memory.py --load：将仓库 .claude-memory/ 中的记忆恢复到系统目录', bp='加载：')
bl(doc, '记忆文件通过 git push/pull 在两台设备间同步', bp='同步：')
doc.add_paragraph()

# Section 5
h1(doc, '四、STATUS.md 进度跟踪')
tx(doc, 'STATUS.md 是跨设备同步的核心——它记录了当前学到哪、下一步做什么。每次切换设备前更新它。')
bl(doc, '用复选框标记已完成的任务 (- [x] 已完成)')
bl(doc, '记录当前学到哪个阶段')
bl(doc, '列出待办事项和遇到的问题')
doc.add_paragraph()

# Section 6
h1(doc, '五、环境要求')
t = doc.add_table(rows=5, cols=2)
t.style = 'Light Grid Accent 1'
env_data = [
    ('工具', '说明'),
    ('MATLAB R2024b + Simulink', '运行 Energy.slx 模型必需'),
    ('Python 3.12+', '运行策略代码和脚本'),
    ('numpy/pandas/matplotlib', 'pip install 安装'),
    ('VS Code', '代码编辑器'),
]
for i, (a, b) in enumerate(env_data):
    t.rows[i].cells[0].text = a; t.rows[i].cells[1].text = b
    for cell in t.rows[i].cells:
        for p2 in cell.paragraphs:
            for r2 in p2.runs: r2.font.size = Pt(9)
    if i == 0:
        for cell in t.rows[i].cells:
            shading(cell, '2F5496')
            for p2 in cell.paragraphs:
                for r2 in p2.runs: r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r2.bold = True

doc.add_paragraph()
r = doc.add_paragraph().add_run('=' * 50 + '\nEMS-PLAN 双设备同步指南\n生成日期：2026-06-02\n' + '=' * 50)
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save('F:/CLAUDE/research/ems-platform/docs/双设备同步指南.docx')
print('OK')
