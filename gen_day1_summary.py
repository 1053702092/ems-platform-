# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        r = p.add_run(t); r.font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def tbl(doc, hd, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(hd))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

doc = Document()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Day1 学习总结'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('第1周第1天 - pandas数据处理 + matplotlib画图'); r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
doc.add_paragraph()

h1(doc, '一、numpy 数值计算')
tx(doc, 'numpy是Python科学计算的基础库，提供数组操作和数学函数。')

h2(doc, 'np.linspace')
tx(doc, '在指定区间内均匀生成等间隔的数。')
bl(doc, 'np.linspace(0, pi, 200) - 在0到pi之间取200个均匀分布的数')
bl(doc, '常用场景：生成连续的时间轴、x轴坐标、平滑曲线')

h2(doc, 'np.arange')
tx(doc, '按步长生成等差数列（类似Python的range()）。')
bl(doc, 'np.arange(0, 1800, 1) - [0,1,2,...,1799]，步长为1')

h2(doc, 'np.sin 配合 np.linspace')
tx(doc, '用sin函数生成平滑的加速/减速曲线。')
bl(doc, 'np.sin(np.linspace(0, pi, N)) * A - 先从0升到A再降到0')
bl(doc, 'np.sin(np.linspace(0, 2*pi, N)) * A + B - 在B的上下波动')

h2(doc, 'np.clip')
tx(doc, '把数组的值限制在指定范围内。')
bl(doc, 'np.clip(v, 0, None) - v中所有负数变成0，正数不变')

h2(doc, 'np.random.seed 和 np.random.normal')
tx(doc, '两个完全不同的概念：')
bl(doc, 'np.random.seed(42)：固定随机种子，让随机结果可复现', bp='seed：')
bl(doc, 'np.random.normal(0, 0.5, n)：生成n个服从正态分布的随机数', bp='normal：')
bl(doc, 'seed控制normal的输出，但seed不影响normal的功能')

h2(doc, '数组切片')
tx(doc, 'Python列表/数组的区间访问语法。')
bl(doc, 'v[0:200] = ... - 给数组第0到第199个元素赋值')
bl(doc, '左闭右开：包含起始索引，不包含结束索引')

doc.add_paragraph()

h1(doc, '二、pandas 数据处理')
tx(doc, 'pandas是Python最主流的数据处理库，核心数据结构是DataFrame（类似Excel表格）。')

h2(doc, '读取CSV')
bl(doc, "pd.read_csv('文件名.csv') - 读取CSV文件为DataFrame")

h2(doc, '数据预览')
bl(doc, 'df.head() - 查看前5行（快速预览数据）')
bl(doc, 'df.describe() - 生成统计摘要（行数/均值/标准差/最大最小/四分位数）')
bl(doc, 'list(df.columns) - 查看所有列名')
bl(doc, 'len(df) - 查看行数')

h2(doc, '数据筛选')
bl(doc, "df[df['speed'] > 50] - 筛选车速>50的所有行")
bl(doc, "df['power_demand'].mean() - 计算功率列的平均值")
bl(doc, "df['power_demand'].max() - 计算功率列的最大值")

h2(doc, '创建和保存DataFrame')
bl(doc, "pd.DataFrame({'列名1': 数据1, '列名2': 数据2}) - 用字典创建表格")
bl(doc, "df.to_csv('文件名.csv', index=False) - 保存CSV")

doc.add_paragraph()

h1(doc, '三、matplotlib 画图')

h2(doc, '基础折线图')
bl(doc, "plt.plot(x, y, 'b-', linewidth=1) - 画蓝色实线折线图")
bl(doc, "alpha=0.7 - 设置线条透明度")

h2(doc, '常用设置')
bl(doc, 'plt.figure(figsize=(12, 6)) - 设置画布尺寸')
bl(doc, 'plt.xlabel() / plt.ylabel() - 设置轴标签')
bl(doc, 'plt.title() - 设置图标题')
bl(doc, 'plt.grid(True, alpha=0.3) - 显示网格线')

h2(doc, '双子图 (subplots)')
bl(doc, 'fig, axes = plt.subplots(2, 1, figsize=(12, 6)) - 创建2行1列子图')
bl(doc, 'axes[0]操作第1个子图，axes[1]操作第2个子图')
bl(doc, 'plt.tight_layout() - 自动调整子图间距')

h2(doc, '双y轴图 (twinx)')
bl(doc, 'ax1 = plt.subplot() - 左y轴（蓝色）')
bl(doc, 'ax2 = ax1.twinx() - 共享x轴的右y轴（红色）')
bl(doc, "ax1.tick_params(axis='y', labelcolor='b') - y轴刻度变蓝")

h2(doc, '保存图片')
bl(doc, "plt.savefig('文件名.png', dpi=150) - 保存PNG图片")

doc.add_paragraph()

h1(doc, '四、Python基础语法')

h2(doc, 'f-string格式化字符串')
bl(doc, "f'文本{变量}文本' - 花括号里嵌入变量或表达式")
bl(doc, "print(f'数据: {len(df)}行') - 输出：数据: 1800行")
bl(doc, '相比+拼接更直观简洁')

h2(doc, 'range / np.arange / np.linspace 区别')
tbl(doc, ['函数', '特点', '示例'],
[['range()', 'Python内置，整数，左闭右开', 'range(0,10,2) -> [0,2,4,6,8]'],
['np.arange()', 'numpy，支持小数，左闭右开', 'np.arange(0,1,0.2) -> [0,0.2,0.4,0.6,0.8]'],
['np.linspace()', 'numpy，指定个数，包含两端', 'np.linspace(0,1,5) -> [0,0.25,0.5,0.75,1]']])

doc.add_paragraph()
r = doc.add_paragraph().add_run('='*50+'\n'+'Day1 学习总结\n'+'生成日期：2026-06-03\n'+'='*50)
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

fname = 'F:/CLAUDE/research/ems-platform/docs/Day1学习总结.docx'
doc.save(fname)
print('OK:', fname)
