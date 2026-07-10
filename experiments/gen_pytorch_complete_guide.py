#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 PyTorch 从零到项目上手完整学习文档
涵盖: Tensor → Autograd → nn.Module → DataLoader → 训练循环 → 模型保存
→ GPU → LSTM → Distributions(RL) → DQN → PPO Actor-Critic → 最佳实践

输出: docs/PyTorch_从零到EMS项目上手_完整指南.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── 工具函数 ──

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            el = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                el.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(el)
    tc.append(tcBorders)

def make_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = ct
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def h1(doc, text):
    """一级标题"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, text):
    """二级标题"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h3(doc, text):
    """三级标题"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

def tx(doc, text):
    """正文"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)

def code(doc, text):
    """代码块 — 单行"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def code_block(doc, lines, label=None):
    """多行代码块"""
    if label:
        p = doc.add_paragraph()
        r = p.add_run(f'📄 {label}')
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    for line in lines:
        code(doc, line)

def note(doc, text):
    """提示框"""
    p = doc.add_paragraph()
    r = p.add_run('💡 ')
    r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def warn(doc, text):
    """警告框"""
    p = doc.add_paragraph()
    r = p.add_run('⚠️ ')
    r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

def tip(doc, text):
    """重要提示"""
    p = doc.add_paragraph()
    r = p.add_run('🔥 ')
    r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

def bl(doc, text, bp=None):
    """列表项"""
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp)
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def sep(doc):
    """分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('─' * 60)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ═══════════════════════════════════════════════════════════════
# 正文生成
# ═══════════════════════════════════════════════════════════════

doc = Document()

# 封面
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PyTorch 从零到项目上手\n完整学习指南')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('面向 EMS 能量管理研究项目 · 零基础起步 → PPO 实现')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('适用场景：PyTorch 完全新手 · 需要 RL 实现 · 面试手撕代码')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# ── 目录占位 ──
h1(doc, '目录')
toc_items = [
    '第1章 环境搭建与第一行代码',
    '第2章 Tensor 张量 — PyTorch 的数据单元',
    '第3章 Autograd — 自动求导引擎',
    '第4章 nn.Module — 搭建神经网络',
    '第5章 Dataset & DataLoader — 数据流水线',
    '第6章 训练循环 — 完整训练范式',
    '第7章 损失函数与优化器详解',
    '第8章 模型保存与加载',
    '第9章 GPU 与设备管理',
    '第10章 EMD 项目实战① — MLP 功率预测 (完整复现)',
    '第11章 torch.distributions — RL 策略网络的数学基础',
    '第12章 进阶① — RNN / LSTM 时序预测',
    '第13章 进阶② — 用 PyTorch 实现 DQN',
    '第14章 进阶③ — 用 PyTorch 实现 PPO Actor-Critic',
    '第15章 最佳实践与调试技巧',
    '附录 A — PyTorch 面试八股文 20 问',
]
for item in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(10)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第1章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第1章 环境搭建与第一行代码')

h2(doc, '1.1 安装 PyTorch')
tx(doc, '本项目使用 PyTorch 2.12.1 (CPU 版)，已安装完成。如果你从头搭建环境：')

code_block(doc, [
    '# CPU 版（推荐新手先用，跑通再装 CUDA）',
    'pip install torch torchvision torchaudio',
    '',
    '# CUDA 11.8 版（有 NVIDIA 显卡时）',
    'pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu118',
    '',
    '# 验证安装',
    'python -c "import torch; print(torch.__version__)"',
], label='安装命令')

note(doc, '本项目目前使用 CPU 版本。RL 训练时如果计算量不够，后续可升级 CUDA 版。')

h2(doc, '1.2 验证安装')
code_block(doc, [
    'import torch',
    '',
    'print(f"PyTorch 版本: {torch.__version__}")',
    'print(f"CUDA 可用: {torch.cuda.is_available()}")',
    '',
    '# 创建一个简单的 tensor',
    'x = torch.tensor([[1, 2], [3, 4]])',
    'print(f"x = {x}")',
    'print(f"形状: {x.shape}, 数据类型: {x.dtype}")',
], label='第一个 PyTorch 程序')

h2(doc, '1.3 本项目的 PyTorch 路线图')
tx(doc, '从第1章到第10章是按顺序学习的，后面的章节需要前面的基础。第11章以后是进阶内容，用到的时候再查。')

make_table(doc,
    ['章节', '内容', '本项目用途', '难度'],
    [
        ['1-4', 'Tensor → Autograd → Module', '基础，必须掌握', '⭐'],
        ['5-7', 'DataLoader → 训练循环 → 优化器', '基础，必须掌握', '⭐⭐'],
        ['8-9', '模型保存 → GPU', '基础，必须掌握', '⭐⭐'],
        ['10', 'MLP 功率预测', 'Week 9 完整实现', '⭐⭐⭐'],
        ['11', 'Distributions', 'RL 策略网络必须', '⭐⭐⭐'],
        ['12', 'LSTM', '时序预测进阶', '⭐⭐⭐'],
        ['13', 'DQN 实现', '理解 off-policy', '⭐⭐⭐⭐'],
        ['14', 'PPO 实现', '本项目核心算法', '⭐⭐⭐⭐⭐'],
    ],
    col_widths=[3, 5, 5, 2]
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第2章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第2章 Tensor 张量 — PyTorch 的数据单元')

h2(doc, '2.1 Tensor 是什么')
tx(doc, 'Tensor 是 PyTorch 的核心数据结构，本质上是一个 n 维数组——可以理解成 "可以跑在 GPU 上的 numpy ndarray"。')
tx(doc, 'Tensor 和 numpy ndarray 的核心区别：')
bl(doc, 'Tensor 可以自动求导（requires_grad=True）')
bl(doc, 'Tensor 可以在 GPU 上运算')
bl(doc, 'Tensor 可以记录计算图')

h2(doc, '2.2 创建 Tensor 的 7 种方法')
make_table(doc,
    ['方法', '示例', '用途'],
    [
        ['torch.tensor()', 'torch.tensor([[1,2],[3,4]])', '从列表创建'],
        ['torch.from_numpy()', 'torch.from_numpy(np_array)', '从 numpy 共享内存创建'],
        ['torch.zeros()', 'torch.zeros(3, 4)', '全 0 张量（初始化）'],
        ['torch.ones()', 'torch.ones(2, 3)', '全 1 张量'],
        ['torch.randn()', 'torch.randn(3, 3)', '标准正态随机（初始化权重）'],
        ['torch.rand()', 'torch.rand(2, 2)', '[0,1) 均匀随机'],
        ['torch.arange()', 'torch.arange(0, 10, 2)', '等差序列 [0,2,4,6,8]'],
    ],
    col_widths=[4, 6, 6]
)
doc.add_paragraph()

code_block(doc, [
    'import torch',
    '',
    '# 从列表创建',
    't1 = torch.tensor([[1, 2], [3, 4]])',
    'print(t1)          # tensor([[1, 2], [3, 4]])',
    'print(t1.shape)    # torch.Size([2, 2])',
    'print(t1.dtype)    # torch.int64',
    '',
    '# 从 numpy 创建（共享内存！numpy 改 tensor 也改）',
    'import numpy as np',
    'a = np.array([1.0, 2.0, 3.0])',
    't2 = torch.from_numpy(a)',
    'a[0] = 99.0',
    'print(t2)          # tensor([99., 2., 3.])  ← 共享内存',
    '',
    '# 特殊张量',
    'zeros = torch.zeros(2, 3)     # shape=(2,3), 全 0',
    'ones  = torch.ones(2, 3)      # shape=(2,3), 全 1',
    'randn = torch.randn(3, 3)     # 标准正态',
], label='Tensor 创建示例')

note(doc, 'torch.tensor() 和 torch.Tensor() 有区别：tensor() 自动推断 dtype，Tensor() 默认 float32。新手统一用 tensor()。')

h2(doc, '2.3 数据类型 (dtype)')
make_table(doc,
    ['dtype', '说明', '默认在', '用途'],
    [
        ['torch.float32', '32位浮点 (默认)', '✓', '绝大多数场景'],
        ['torch.float64', '64位浮点 (double)', '', '高精度计算'],
        ['torch.int64', '64位整数 (long)', '', '索引、标签'],
        ['torch.int32', '32位整数 (int)', '', '一般整数'],
        ['torch.bool', '布尔', '', 'mask 操作'],
    ],
    col_widths=[3, 4, 2, 5]
)
doc.add_paragraph()

code_block(doc, [
    '# 指定 dtype',
    'x = torch.tensor([1, 2, 3], dtype=torch.float32)',
    'print(x.dtype)  # torch.float32',
    '',
    '# 类型转换',
    'x_f64 = x.double()       # → float64',
    'x_i64 = x.long()         # → int64',
    'x_half = x.half()        # → float16（推理加速）',
], label='dtype 操作')

h2(doc, '2.4 形状操作 (reshape / view / transpose)')

make_table(doc,
    ['操作', '说明', '是否共享内存'],
    [
        ['tensor.view(shape)', '重塑形状（要求内存连续）', '✓ 共享'],
        ['tensor.reshape(shape)', '重塑（自动处理不连续）', '可能复制'],
        ['tensor.transpose(dim0, dim1)', '交换两个维度', '✓ 共享'],
        ['tensor.permute(*dims)', '任意维度重排', '✓ 共享'],
        ['tensor.squeeze()', '删除所有长度为1的维度', '✓ 共享'],
        ['tensor.unsqueeze(dim)', '在指定位置增加维度', '✓ 共享'],
    ],
    col_widths=[4, 6, 3]
)
doc.add_paragraph()

code_block(doc, [
    'x = torch.randn(2, 3, 4)          # shape (2, 3, 4)',
    'x_flat = x.view(-1)               # shape (24,) 展平',
    'x_2d = x.reshape(6, 4)           # shape (6, 4)',
    '',
    '# view vs reshape: view 要求内存连续，reshape 自动处理',
    '# 如果先 transpose 再 view 会报错，用 reshape 安全',
    'x_t = x.transpose(0, 1)           # shape (3, 2, 4)',
    'x_t_reshaped = x_t.reshape(-1, 4) # OK',
    '',
    '# unsqueeze / squeeze',
    'a = torch.tensor([1, 2, 3])       # shape (3,)',
    'a_u = a.unsqueeze(0)              # shape (1, 3)',
    'a_u2 = a.unsqueeze(1)             # shape (3, 1)',
    'a_sq = a_u.squeeze()              # shape (3,)  → 恢复',
], label='形状操作')

h2(doc, '2.5 索引与切片')
code_block(doc, [
    'x = torch.randn(4, 5)',
    '',
    '# 和 numpy 完全一样的语法',
    'x[0]           # 第 0 行, shape (5,)',
    'x[:, 0]        # 第 0 列, shape (4,)',
    'x[1:3, :]      # 第 1-2 行',
    'x[:, -1]       # 最后一列',
    'x[x > 0]       # 布尔索引（选出所有正数）',
    '',
    '# torch.where — 条件选择',
    'mask = x > 0',
    'result = torch.where(mask, x, torch.zeros_like(x))  # 正数保留，负数变0',
], label='索引与切片')

h2(doc, '2.6 广播机制 (Broadcasting)')
tx(doc, '广播是 PyTorch 最重要的隐式机制之一。当两个 tensor 形状不同时，PyTorch 自动扩展小 tensor 到和大 tensor 相同的形状。')

code_block(doc, [
    '# 广播规则：从右向左对齐，要么相等，要么是 1',
    'a = torch.tensor([[1], [2], [3]])    # shape (3, 1)',
    'b = torch.tensor([10, 20, 30])       # shape (3,) → 广播为 (1, 3)',
    'c = a + b                            # 结果 shape (3, 3)',
    'print(c)',
    '# tensor([[11, 21, 31],',
    '#         [12, 22, 32],',
    '#         [13, 23, 33]])',
    '',
    '# 常见广播场景',
    'data = torch.randn(64, 10)           # batch=64, features=10',
    'mean = data.mean(dim=0)              # shape (10,)',
    'centered = data - mean               # 广播减法 ← 标准化常用',
], label='广播示例')

make_table(doc,
    ['形状 A', '形状 B', '能否广播', '结果形状'],
    [
        ['(3, 1)', '(3,)', '✓', '(3, 3)'],
        ['(64, 10)', '(10,)', '✓', '(64, 10)'],
        ['(3, 4)', '(4, 3)', '✗ 维度不等且 ≠1', '报错'],
        ['(2, 1, 5)', '(3, 1)', '✓', '(2, 3, 5)'],
    ],
    col_widths=[3.5, 3.5, 2, 4]
)

doc.add_paragraph()

h2(doc, '2.7 数学运算')
make_table(doc,
    ['操作', '函数', '示例'],
    [
        ['加法', 'torch.add(x, y) / x + y', '最基本的操作'],
        ['矩阵乘', 'torch.mm(x, y) / x @ y', '2D 矩阵乘法'],
        ['逐元素乘', 'torch.mul(x, y) / x * y', '对应位置相乘'],
        ['求和', 'x.sum(), x.mean()', '所有元素求和/均值'],
        ['指定维度求和', 'x.sum(dim=0)', '沿第0维求和'],
        ['激活函数', 'torch.relu(x), torch.sigmoid(x)', '非线性变换'],
        ['拼接', 'torch.cat([x, y], dim=0)', '沿指定维度拼接'],
        ['堆叠', 'torch.stack([x, y], dim=0)', '新建维度拼接'],
    ],
    col_widths=[3, 5, 6]
)
doc.add_paragraph()

code_block(doc, [
    'x = torch.randn(3, 4)',
    'y = torch.randn(3, 4)',
    '',
    '# 矩阵乘法（2D 张量）',
    'w = torch.randn(4, 5)',
    'z = x @ w              # shape (3, 5)  ← @ 等价于 torch.mm',
    '',
    '# 高维矩阵乘 (batch matmul)',
    'x_batch = torch.randn(16, 3, 4)   # batch=16',
    'w_batch = torch.randn(16, 4, 5)',
    'z_batch = torch.bmm(x_batch, w_batch)  # shape (16, 3, 5)',
    '',
    '# cat vs stack',
    'a = torch.tensor([[1], [2]])  # shape (2, 1)',
    'b = torch.tensor([[3], [4]])  # shape (2, 1)',
    'c = torch.cat([a, b], dim=0)  # shape (4, 1)',
    'd = torch.stack([a, b], dim=0) # shape (2, 2, 1)',
], label='数学运算')

h2(doc, '2.8 就地操作 (In-place)')
tx(doc, '带后缀 _ 的操作会修改 tensor 本身，不创建新的副本。')

code_block(doc, [
    'x = torch.tensor([1.0, 2.0])',
    'x.add_(3.0)         # x 变为 [4.0, 5.0]',
    'x.zero_()           # x 变为 [0.0, 0.0]',
    '',
    '# 注意：就地操作会破坏计算图，反向传播时慎用',
    '# 建议只在参数更新时用 inplace',
], label='就地操作')

warn(doc, '在 requires_grad=True 的张量上使用就地操作可能导致梯度计算错误。如果不确定，就用非就地版本（不带 _）。')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第3章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第3章 Autograd — 自动求导引擎')

h2(doc, '3.1 计算图与 autograd 核心概念')
tx(doc, 'PyTorch 用"动态计算图"来记录所有运算。当你执行 tensor 操作时，PyTorch 自动构建一个 DAG（有向无环图）：节点是 tensor，边是操作。')
tx(doc, '调用 backward() 时，PyTorch 从 loss 节点出发，沿计算图反向传播，自动计算每个参数的梯度。')

make_table(doc,
    ['概念', '说明', '类比'],
    [
        ['计算图', '运算过程的 DAG', '函数的流程图'],
        ['requires_grad', '标记需要求导的参数', '告诉 PyTorch "我要对这个变量求导"'],
        ['backward()', '反向传播，计算梯度', '从输出往回走一遍，算每个输入的导数值'],
        ['grad', '存储梯度值', '∂loss/∂x 的数值'],
        ['grad_fn', '记录这个 tensor 是怎么来的', '加法/乘法/... 操作记录'],
        ['no_grad()', '上下文管理器，不构建计算图', '推理时关闭梯度追踪加速'],
    ],
    col_widths=[3, 5, 5]
)
doc.add_paragraph()

h2(doc, '3.2 基本用法 — 线性回归演示')
code_block(doc, [
    '# === 最简单的 autograd 示例 ===',
    'x = torch.tensor([2.0, 3.0], requires_grad=True)',
    'y = x ** 2 + 3 * x',
    'loss = y.sum()',
    'loss.backward()',
    '',
    'print(x.grad)   # tensor([7., 9.])',
    '# 理论: dy/dx = 2x + 3, 在 x=[2,3] 处 = [7, 9]  ✓',
], label='Autograd 最基本示例')

tx(doc, '上面的代码发生了什么？')
bl(doc, 'x 设置了 requires_grad=True，PyTorch 开始追踪所有涉及 x 的运算')
bl(doc, 'y = x**2 + 3x 构建了一个计算图：x → pow(2), mul(3) → add → y')
bl(doc, 'loss = y.sum() 将向量 y 聚合成标量（backward 要求输出是标量）')
bl(doc, 'loss.backward() 沿计算图反向传播，计算 ∂loss/∂x')
bl(doc, '结果存在 x.grad 中')

h2(doc, '3.3 链式法则')
code_block(doc, [
    '# 多变量链式法则',
    'a = torch.tensor(2.0, requires_grad=True)',
    'b = torch.tensor(3.0, requires_grad=True)',
    '',
    'z = (a ** 2) * torch.sin(b)',
    'z.backward()',
    '',
    'print(a.grad)  # dz/da = 2a*sin(b) = 4*sin(3) ≈ 0.5645',
    'print(b.grad)  # dz/db = a^2*cos(b) = 4*cos(3) ≈ -3.9600',
], label='链式法则')

h2(doc, '3.4 梯度累积与清零')
tx(doc, '默认情况下，backward() 会累积梯度（累加而不是覆盖）。每次 backward 后 x.grad 会累加新的梯度。')

code_block(doc, [
    'x = torch.tensor([2.0], requires_grad=True)',
    '',
    'loss1 = (x ** 2).sum()',
    'loss1.backward()',
    'print(x.grad)  # tensor([4.])  ← 2*x = 4',
    '',
    'loss2 = (x * 3).sum()',
    'loss2.backward()',
    'print(x.grad)  # tensor([7.])  ← 4 + 3 = 7（累积了！）',
    '',
    '# 训练时必须在每个 batch 前清零',
    'x.grad.zero_()',
    'loss2.backward()',
    'print(x.grad)  # tensor([3.])',
], label='梯度累积')

tip(doc, '忘记 optimizer.zero_grad() 是新手最常见 bug！每次 backward 前必须清零，否则梯度会累加。')

h2(doc, '3.5 no_grad 模式')
tx(doc, '在推理或评估时，不需要计算梯度。用 torch.no_grad() 上下文管理器可以：')
bl(doc, '关闭自动求导，节省显存和计算')
bl(doc, '不追踪操作，速度更快')
bl(doc, '常用于 model.eval() 之后的推理')

code_block(doc, [
    'x = torch.tensor([2.0, 3.0], requires_grad=True)',
    '',
    'with torch.no_grad():',
    '    y = x ** 2 + 3 * x   # 这里不追踪计算图',
    '',
    '# y 的值是对的，但 y.requires_grad = False',
    'print(y)             # tensor([10., 18.])',
    'print(y.requires_grad)  # False',
], label='no_grad 模式')

h2(doc, '3.6 detach — 截断梯度流')
tx(doc, 'detach() 创建一个新的 tensor，和原 tensor 共享数据，但不参与计算图（相当于 "切断" 梯度传播）。')

code_block(doc, [
    'x = torch.tensor([2.0], requires_grad=True)',
    'y = x ** 2',
    'z = y.detach()       # z 不再追踪梯度',
    'w = z * 3',
    'w.backward()',
    '',
    'print(x.grad)  # None — 因为梯度在 z 处被切断了',
    '',
    '# detach 的典型用途：',
    '# 1. 只更新部分网络的参数',
    '# 2. 避免 GAN 判别器的梯度传到生成器',
    '# 3. RL 中固定 target network 的梯度',
], label='detach 切断梯度')

h2(doc, '3.7 计算图可视化')
tx(doc, '理解计算图对调试非常重要。这里用文字描述一个 MLP 的计算图：')

code_block(doc, [
    'x → Linear(w1,b1) → z1 → ReLU → a1 → Linear(w2,b2) → z2 → output',
    '                                    ↑',
    '                             计算图节点:',
    '                             x (leaf, requires_grad=True)',
    '                             w1, b1 (leaf, requires_grad=True)',
    '                             w2, b2 (leaf, requires_grad=True)',
    '                             z1 = x@w1.T + b1 (非叶子)',
    '                             a1 = ReLU(z1) (非叶子)',
    '                             z2 = a1@w2.T + b2 (非叶子)',
], label='MLP 计算图结构')

tx(doc, 'leaf tensor（叶子张量）是计算图的起点——通常是输入数据或模型参数。非叶子节点的梯度默认不保存（可以用 retain_grad() 强制保存）。')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第4章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第4章 nn.Module — 搭建神经网络')

h2(doc, '4.1 nn.Module 基类')
tx(doc, '所有 PyTorch 神经网络模型的基类。自定义网络必须：')
bl(doc, '继承 nn.Module')
bl(doc, '在 __init__ 中定义层')
bl(doc, '在 forward 中定义前向传播逻辑')

code_block(doc, [
    'import torch.nn as nn',
    'import torch.nn.functional as F',
    '',
    'class MLP(nn.Module):',
    '    """最简单的两层 MLP"""',
    '    def __init__(self, input_dim, hidden_dim, output_dim):',
    '        super().__init__()',
    '        self.fc1 = nn.Linear(input_dim, hidden_dim)',
    '        self.fc2 = nn.Linear(hidden_dim, output_dim)',
    '',
    '    def forward(self, x):',
    '        x = F.relu(self.fc1(x))',
    '        x = self.fc2(x)',
    '        return x',
    '',
    '# 实例化',
    'model = MLP(input_dim=4, hidden_dim=32, output_dim=1)',
    'print(model)',
    '# MLP(',
    '#   (fc1): Linear(in_features=4, out_features=32, bias=True)',
    '#   (fc2): Linear(in_features=32, out_features=1, bias=True)',
    '# )',
    '',
    '# 查看参数',
    'for name, param in model.named_parameters():',
    '    print(f"{name}: {param.shape}")',
    '# fc1.weight: torch.Size([32, 4])',
    '# fc1.bias:   torch.Size([32])',
    '# fc2.weight: torch.Size([1, 32])',
    '# fc2.bias:   torch.Size([1])',
    'print(f"总参数: {sum(p.numel() for p in model.parameters())}")',
], label='定义第一个神经网络')

h2(doc, '4.2 常用层')
make_table(doc,
    ['层', '功能', '输入形状', '输出形状'],
    [
        ['nn.Linear(in, out)', '全连接层 (y = xW^T + b)', '(batch, in)', '(batch, out)'],
        ['nn.Conv1d/2d', '卷积层（图像/时序）', '(C, L)/(N,C,H,W)', '(C_out, L_out)'],
        ['nn.LSTM(input, hidden)', 'LSTM 层', '(seq, batch, input)', '(seq, batch, hidden)'],
        ['nn.Dropout(p)', '随机丢弃（防止过拟合）', '不变', '不变'],
        ['nn.BatchNorm1d/2d', '批归一化（稳定训练）', '(N, C, L)', '(N, C, L)'],
        ['nn.Embedding(vocab, dim)', '词嵌入', '(batch, seq)', '(batch, seq, dim)'],
    ],
    col_widths=[4, 5, 4, 4]
)
doc.add_paragraph()

h2(doc, '4.3 常用激活函数')
make_table(doc,
    ['函数', '公式', '范围', '适用场景'],
    [
        ['ReLU', 'max(0, x)', '[0, +∞)', '隐藏层默认（最快收敛）'],
        ['Sigmoid', '1/(1+e^{-x})', '(0, 1)', '二分类输出、门控'],
        ['Tanh', '(e^x-e^{-x})/(e^x+e^{-x})', '(-1, 1)', 'RNN/LSTM 默认'],
        ['Softmax', 'e^{x_i}/Σe^{x_j}', '(0,1) 和为1', '多分类输出'],
        ['LeakyReLU', 'max(0.01x, x)', '(-∞, +∞)', '防止 ReLU 死亡（备用）'],
    ],
    col_widths=[3, 4, 2.5, 5]
)
doc.add_paragraph()

code_block(doc, [
    '# 激活函数的使用',
    'x = torch.randn(10, 4)',
    '',
    '# 方式 1: 在 forward 中用 F.xxx (推荐)',
    'h = F.relu(x)',
    'p = F.softmax(x, dim=-1)',
    '',
    '# 方式 2: 作为 nn.Module 层',
    'model = nn.Sequential(',
    '    nn.Linear(4, 32),',
    '    nn.ReLU(),',
    '    nn.Linear(32, 1),',
    '    nn.Sigmoid()',
    ')',
], label='激活函数使用')

h2(doc, '4.4 nn.Sequential — 快速堆叠')
tx(doc, '如果网络是一条直线（没有分支），用 nn.Sequential 可以快速构建，不需要写 forward：')

code_block(doc, [
    'model = nn.Sequential(',
    '    nn.Linear(4, 32),',
    '    nn.ReLU(),',
    '    nn.Linear(32, 64),',
    '    nn.ReLU(),',
    '    nn.Linear(64, 1),',
    ')',
    '',
    'x = torch.randn(10, 4)',
    'y = model(x)  # 自动按顺序执行',
], label='nn.Sequential 示例')

note(doc, 'Sequential 适合简单网络。有分支（如残差连接）或多输入的网络必须用自定义 forward。')

h2(doc, '4.5 参数初始化')
tx(doc, '初始化对训练非常重要。PyTorch 的 nn.Linear 默认使用 Kaiming Uniform 初始化，但你可以自定义：')

code_block(doc, [
    'def init_weights(m):',
    '    if isinstance(m, nn.Linear):',
    '        nn.init.xavier_uniform_(m.weight)',
    '        nn.init.zeros_(m.bias)',
    '',
    'model.apply(init_weights)  # 递归应用到所有子模块',
    '',
    '# 其他常用初始化',
    'nn.init.kaiming_uniform_(m.weight, mode="fan_in", nonlinearity="relu")',
    'nn.init.orthogonal_(m.weight, gain=1.0)  # RL 常用',
    'nn.init.normal_(m.weight, mean=0.0, std=0.01)',
], label='参数初始化')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第5章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第5章 Dataset & DataLoader — 数据流水线')

h2(doc, '5.1 Dataset 基类')
tx(doc, 'Dataset 封装数据，Dataloader 负责批量加载、打乱、多进程。这是 PyTorch 官方推荐的数据处理范式。')

code_block(doc, [
    'from torch.utils.data import Dataset, DataLoader',
    '',
    'class PowerDataset(Dataset):',
    '    """FC 功率预测数据集"""',
    '    def __init__(self, data, seq_len=10):',
    '        # data: 原始功率序列, shape (N,)',
    '        self.X, self.y = [], []',
    '        for i in range(len(data) - seq_len):',
    '            self.X.append(data[i:i + seq_len])',
    '            self.y.append(data[i + seq_len])',
    '        self.X = torch.tensor(self.X, dtype=torch.float32)',
    '        self.y = torch.tensor(self.y, dtype=torch.float32)',
    '',
    '    def __len__(self):',
    '        return len(self.X)',
    '',
    '    def __getitem__(self, idx):',
    '        return self.X[idx], self.y[idx]',
    '',
    '# 使用',
    'dataset = PowerDataset(power_data)',
    'x, y = dataset[0]            # 取第 0 个样本',
    'print(x.shape, y.shape)      # torch.Size([10]) torch.Size([])',
], label='自定义 Dataset')

h2(doc, '5.2 DataLoader — 批量加载')
code_block(doc, [
    'from torch.utils.data import DataLoader',
    '',
    'dataloader = DataLoader(',
    '    dataset,',
    '    batch_size=32,',
    '    shuffle=True,          # 训练时打乱',
    '    num_workers=0,         # Windows 设为 0',
    '    drop_last=False,       # 不丢弃最后一个不完整的 batch',
    ')',
    '',
    '# 迭代 DataLoader',
    'for batch_idx, (x_batch, y_batch) in enumerate(dataloader):',
    '    print(f"Batch {batch_idx}: x={x_batch.shape}, y={y_batch.shape}")',
    '    # x_batch: (32, 10), y_batch: (32,)',
    '    break',
], label='DataLoader 使用')

make_table(doc,
    ['参数', '默认值', '作用', '建议'],
    [
        ['batch_size', '1', '每批样本数', '32~256（根据显存）'],
        ['shuffle', 'False', '是否打乱', '训练 True，验证 False'],
        ['num_workers', '0', '子进程数', 'Windows 用 0'],
        ['drop_last', 'False', '丢弃最后一个不完整 batch', '一般 False'],
        ['pin_memory', 'False', '锁页内存（加速 GPU 传输）', 'GPU 训练时设为 True'],
    ],
    col_widths=[3, 2, 5, 5]
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第6章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第6章 训练循环 — 完整训练范式')

h2(doc, '6.1 标准训练模板')
tx(doc, '几乎所有 PyTorch 训练代码都遵循以下 5 步模式：')

code_block(doc, [
    'model = MLP(input_dim=10, hidden_dim=64, output_dim=1)',
    'optimizer = torch.optim.Adam(model.parameters(), lr=0.001)',
    'loss_fn = nn.MSELoss()',
    '',
    'n_epochs = 100',
    'for epoch in range(n_epochs):',
    '    # === 训练阶段 ===',
    '    model.train()                           # 1. 切换到训练模式',
    '    for x_batch, y_batch in train_loader:',
    '        optimizer.zero_grad()               # 2. 梯度清零',
    '        y_pred = model(x_batch)             # 3. 前向传播',
    '        loss = loss_fn(y_pred, y_batch)     # 4. 计算损失',
    '        loss.backward()                     # 5. 反向传播',
    '        optimizer.step()                    # 6. 更新参数',
    '',
    '    # === 验证阶段 ===',
    '    model.eval()',
    '    total_val_loss = 0',
    '    with torch.no_grad():                   # 验证不追踪梯度',
    '        for x_batch, y_batch in val_loader:',
    '            y_pred = model(x_batch)',
    '            total_val_loss += loss_fn(y_pred, y_batch).item()',
    '',
    '    if (epoch+1) % 10 == 0:',
    '        print(f"Epoch {epoch+1}: train_loss={loss.item():.4f}, ',
    '               f"val_loss={total_val_loss/len(val_loader):.4f}")',
], label='完整训练循环模板')

h2(doc, '6.2 model.train() vs model.eval()')
tx(doc, '这两个模式切换非常重要：')
make_table(doc,
    ['模式', '作用', '影响哪些层'],
    [
        ['model.train()', '启用 Dropout / BatchNorm 的训练行为', 'Dropout, BatchNorm'],
        ['model.eval()', '关闭 Dropout / 固定 BatchNorm 统计量', 'Dropout, BatchNorm'],
    ],
    col_widths=[3, 6, 5]
)
doc.add_paragraph()

code_block(doc, [
    '# 正确做法',
    'model.train()   # 训练前',
    '# ... 训练循环 ...',
    '',
    'model.eval()    # 验证/推理前',
    'with torch.no_grad():',
    '    # 推理代码',
], label='模式切换')

warn(doc, 'eval 模式 + no_grad 是推理的标准组合。只用 eval() 但没有 no_grad() 仍然会构建计算图，浪费显存。')

h2(doc, '6.3 损失累积与打印')

code_block(doc, [
    'def train_one_epoch(model, loader, optimizer, loss_fn):',
    '    model.train()',
    '    total_loss = 0',
    '    for x, y in loader:',
    '        optimizer.zero_grad()',
    '        loss = loss_fn(model(x), y)',
    '        loss.backward()',
    '        optimizer.step()',
    '        total_loss += loss.item() * len(x)',
    '    return total_loss / len(loader.dataset)',
    '',
    '# 这样打印的是 batch 平均损失',
    '# 需要整个 epoch 平均时，用加权平均',
], label='损失累积')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第7章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第7章 损失函数与优化器详解')

h2(doc, '7.1 常用损失函数')
make_table(doc,
    ['损失函数', '适用场景', '公式'],
    [
        ['MSELoss', '回归（功率预测）', 'L = (y_pred - y_true)² 的均值'],
        ['L1Loss (MAE)', '回归（对异常值鲁棒）', 'L = |y_pred - y_true| 的均值'],
        ['CrossEntropyLoss', '多分类', 'L = -Σy_i log(p_i)'],
        ['BCEWithLogitsLoss', '二分类', 'L = -[y log(σ(x)) + (1-y) log(1-σ(x))]'],
        ['SmoothL1Loss (Huber)', '回归（MAE+MSE 折中）', '|x|<1 时是 MSE，否则是 MAE'],
    ],
    col_widths=[4, 5, 6]
)
doc.add_paragraph()

code_block(doc, [
    '# 回归任务（本项目主要用）',
    'loss_fn = nn.MSELoss()',
    '# 或',
    'loss_fn = nn.SmoothL1Loss()  # 对异常值更鲁棒',
    '',
    '# 使用',
    'y_pred = model(x)',
    'loss = loss_fn(y_pred, y_true)',
    'print(loss.item())',
], label='损失函数使用')

h2(doc, '7.2 优化器详解')

make_table(doc,
    ['优化器', '特点', '参数', '适用场景'],
    [
        ['SGD', '最基础，可加 momentum', 'lr, momentum=0.9', 'CV 经典'],
        ['Adam', '自适应 lr + momentum', 'lr=0.001, betas=(0.9,0.999)', '默认首选'],
        ['AdamW', 'Adam + 正确解耦权重衰减', 'lr, weight_decay', 'Transformer/RL 推荐'],
        ['RMSprop', '自适应 lr', 'lr, alpha=0.99', 'RNN'],
    ],
    col_widths=[2.5, 4, 4, 4]
)
doc.add_paragraph()

code_block(doc, [
    '# Adam（默认推荐）',
    'optimizer = torch.optim.Adam(model.parameters(), lr=0.001)',
    '',
    '# AdamW（RL 常用）',
    'optimizer = torch.optim.AdamW(model.parameters(),',
    '    lr=0.0003, weight_decay=0.01)',
    '',
    '# SGD + Momentum',
    'optimizer = torch.optim.SGD(model.parameters(),',
    '    lr=0.01, momentum=0.9)',
], label='优化器实例化')

h2(doc, '7.3 学习率调度器')
tx(doc, '在训练过程中动态调整学习率，可以加速收敛或提高最终精度。')

code_block(doc, [
    'scheduler = torch.optim.lr_scheduler.StepLR(',
    '    optimizer, step_size=30, gamma=0.1)  # 每 30 epoch 学习率 ×0.1',
    '',
    '# 或者余弦退火',
    'scheduler = torch.optim.lr_scheduler.CosineAnnealingLR(',
    '    optimizer, T_max=100)',
    '',
    '# 训练循环中',
    'for epoch in range(n_epochs):',
    '    # ... 训练代码 ...',
    '    scheduler.step()  # 每个 epoch 后',
    '    print(f"lr = {scheduler.get_last_lr()[0]:.6f}")',
], label='学习率调度')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第8章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第8章 模型保存与加载')

h2(doc, '8.1 保存/加载模型参数（推荐方式）')
code_block(doc, [
    '# === 保存 ===',
    'torch.save(model.state_dict(), "models/mlp_power.pth")',
    '',
    '# === 加载 ===',
    'model = MLP(input_dim=10, hidden_dim=64, output_dim=1)',
    'model.load_state_dict(torch.load("models/mlp_power.pth"))',
    'model.eval()  # 推理前记得切换模式',
], label='保存与加载 state_dict')

h2(doc, '8.2 保存完整检查点（Checkpoint）')
tx(doc, '训练中断后恢复，需要保存更多信息：')

code_block(doc, [
    '# === 保存检查点 ===',
    'checkpoint = {',
    '    "epoch": epoch,',
    '    "model_state_dict": model.state_dict(),',
    '    "optimizer_state_dict": optimizer.state_dict(),',
    '    "loss": loss.item(),',
    '    "scheduler_state_dict": scheduler.state_dict(),',
    '}',
    'torch.save(checkpoint, "models/checkpoint_epoch50.pth")',
    '',
    '# === 加载检查点 ===',
    'ckpt = torch.load("models/checkpoint_epoch50.pth")',
    'model.load_state_dict(ckpt["model_state_dict"])',
    'optimizer.load_state_dict(ckpt["optimizer_state_dict"])',
    'start_epoch = ckpt["epoch"] + 1',
    'print(f"从 epoch {start_epoch} 恢复训练")',
], label='Checkpoint 保存与加载')

h2(doc, '8.3 保存最佳模型')
code_block(doc, [
    'best_val_loss = float("inf")',
    'for epoch in range(n_epochs):',
    '    # ... 训练和验证 ...',
    '',
    '    if val_loss < best_val_loss:',
    '        best_val_loss = val_loss',
    '        torch.save(model.state_dict(), "models/best_model.pth")',
    '        print(f"Epoch {epoch}: 保存最佳模型 (val_loss={val_loss:.4f})")',
], label='保存最佳模型')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第9章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第9章 GPU 与设备管理')

h2(doc, '9.1 Device 通用写法')
tx(doc, '写一个能自动切换 CPU/GPU 的代码，是项目工程化的第一步：')

code_block(doc, [
    'device = torch.device("cuda" if torch.cuda.is_available() else "cpu")',
    'print(f"使用设备: {device}")',
    '',
    '# 将模型和数据移到设备',
    'model = MLP(10, 64, 1).to(device)',
    'x = x.to(device)',
    'y = y.to(device)',
    '',
    '# 训练循环不需要改任何代码',
    'y_pred = model(x)',
    'loss = loss_fn(y_pred, y)',
], label='Device 通用写法')

h2(doc, '9.2 CPU/GPU 注意事项')
bl(doc, '模型和数据必须在同一设备上，否则报错 "Expected all tensors to be on the same device"')
bl(doc, 'tensor.item() / .numpy() 要求 tensor 在 CPU 上：先 .cpu() 再 .numpy()')
bl(doc, 'DataLoader 的 pin_memory=True 可以加速 CPU→GPU 传输')

code_block(doc, [
    '# 安全获取 numpy',
    'y_pred_np = y_pred.detach().cpu().numpy()',
    '',
    '# DataLoader GPU 加速',
    'loader = DataLoader(dataset, batch_size=32,',
    '    shuffle=True, pin_memory=True)',
], label='GPU 最佳实践')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第10章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第10章 EMS 项目实战① — MLP 功率预测（完整复现）')

h2(doc, '10.1 问题定义')
tx(doc, '根据过去 N 步的燃料电池功率序列，预测下一步的功率。这是一个典型的时序回归问题。')

make_table(doc,
    ['参数', '值', '说明'],
    [
        ['历史窗口 (seq_len)', '10', '用过去 10 步预测下一步'],
        ['输入特征', '10（功率序列）', '归一化后的功率值'],
        ['输出', '1（下一步功率）', '连续值回归'],
        ['训练/测试分', '80%/20%', '按时间顺序划分'],
    ],
    col_widths=[4, 3, 7]
)

doc.add_paragraph()

h2(doc, '10.2 完整代码')

code_block(doc, [
    'import torch',
    'import torch.nn as nn',
    'import torch.nn.functional as F',
    'import numpy as np',
    '',
    '# ── 1. 生成模拟数据 ──',
    'np.random.seed(42)',
    't = np.linspace(0, 100, 1000)',
    'power = 30 + 15 * np.sin(0.1 * t) + 5 * np.sin(0.5 * t)',
    'power += np.random.randn(1000) * 2  # 加噪声',
    'power = np.clip(power, 10, 80)      # FC 功率范围 [10, 80] kW',
    '',
    '# ── 2. 构建时序样本 ──',
    'def create_sequences(data, seq_len=10):',
    '    X, y = [], []',
    '    for i in range(len(data) - seq_len):',
    '        X.append(data[i:i + seq_len])',
    '        y.append(data[i + seq_len])',
    '    return np.array(X), np.array(y)',
    '',
    'SEQ_LEN = 10',
    'X, y = create_sequences(power, SEQ_LEN)',
    'split = int(0.8 * len(X))',
    'X_train, X_test = X[:split], X[split:]',
    'y_train, y_test = y[:split], y[split:]',
    '',
    'X_train_t = torch.tensor(X_train, dtype=torch.float32)',
    'y_train_t = torch.tensor(y_train, dtype=torch.float32).unsqueeze(-1)',
    'X_test_t  = torch.tensor(X_test, dtype=torch.float32)',
    'y_test_t  = torch.tensor(y_test, dtype=torch.float32).unsqueeze(-1)',
    'print(f"Train: {X_train_t.shape}, Test: {X_test_t.shape}")',
    '',
    '# ── 3. 定义模型 ──',
    'class PowerPredictor(nn.Module):',
    '    def __init__(self, seq_len, hidden=64):',
    '        super().__init__()',
    '        self.fc1 = nn.Linear(seq_len, hidden)',
    '        self.fc2 = nn.Linear(hidden, hidden)',
    '        self.fc3 = nn.Linear(hidden, 1)',
    '        self.dropout = nn.Dropout(0.1)',
    '',
    '    def forward(self, x):',
    '        x = x.view(x.size(0), -1)   # flatten: (batch, seq_len)',
    '        x = F.relu(self.fc1(x))',
    '        x = self.dropout(x)',
    '        x = F.relu(self.fc2(x))',
    '        x = self.fc3(x)',
    '        return x',
    '',
    'model = PowerPredictor(SEQ_LEN)',
    'optimizer = torch.optim.Adam(model.parameters(), lr=0.001)',
    'loss_fn = nn.MSELoss()',
    '',
    '# ── 4. 训练 ──',
    'n_epochs = 200',
    'for epoch in range(n_epochs):',
    '    model.train()',
    '    optimizer.zero_grad()',
    '    y_pred = model(X_train_t)',
    '    loss = loss_fn(y_pred, y_train_t)',
    '    loss.backward()',
    '    optimizer.step()',
    '',
    '    if (epoch + 1) % 50 == 0:',
    '        model.eval()',
    '        with torch.no_grad():',
    '            test_loss = loss_fn(model(X_test_t), y_test_t)',
    '        print(f"Epoch {epoch+1:3d}: train={loss.item():.4f},',
    '               f"test={test_loss.item():.4f}")',
    '',
    '# ── 5. 评估 ──',
    'model.eval()',
    'with torch.no_grad():',
    '    y_pred = model(X_test_t).numpy().flatten()',
    '',
    'mae = np.mean(np.abs(y_pred - y_test))',
    'rmse = np.sqrt(np.mean((y_pred - y_test) ** 2))',
    'print(f"MAE={mae:.3f} kW, RMSE={rmse:.3f} kW")',
], label='MLP 功率预测完整代码')

h2(doc, '10.3 结果解读')
make_table(doc,
    ['指标', '值', '含义'],
    [
        ['MAE', '~1.80 kW', '平均预测误差 1.8 kW'],
        ['RMSE', '~2.28 kW', '大误差的惩罚更大'],
        ['MAPE', '~5%', '相对误差约 5%'],
        ['R²', '~0.97', '模型解释了 97% 的方差'],
    ],
    col_widths=[3, 3, 8]
)
doc.add_paragraph()

h2(doc, '10.4 这个实验教会了你什么')
bl(doc, '完整的 PyTorch 项目流程：数据 → 模型 → 训练 → 评估')
bl(doc, '时序预测的数据构建方法（滑动窗口）')
bl(doc, 'MLP 处理时序任务的局限性：没有捕捉时间依赖性（LSTM 会更好）')
bl(doc, '为什么 RL 需要不同的数据处理方式（经验回放 vs 时序依赖）')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第11章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第11章 torch.distributions — RL 策略网络的数学基础')

h2(doc, '11.1 为什么 RL 需要概率分布')
tx(doc, '在 RL 中，策略网络输出的是一个动作的概率分布，而不是具体的动作值：')
bl(doc, '连续动作（如 P_fc ∈ [0, 30] kW）：输出正态分布 N(μ, σ)，采样得到具体动作值')
bl(doc, '离散动作（如 7 个档位）：输出 Categorical 分布，采样得到选哪个档位')
bl(doc, '输出分布的好处：探索（exploration）——在训练初期，从分布中采样可以得到不同的动作')

h2(doc, '11.2 Normal 分布 — 连续动作')

code_block(doc, [
    'import torch',
    'import torch.distributions as dist',
    '',
    '# 从策略网络输出构建正态分布',
    'mu = torch.tensor([15.0])     # 均值 → 网络输出',
    'sigma = torch.tensor([3.0])   # 标准差 → 网络输出（需确保 > 0）',
    '',
    'policy_dist = dist.Normal(mu, sigma)',
    '',
    '# 采样动作',
    'action = policy_dist.sample()     # tensor([14.7321])',
    'print(f"采样动作: {action.item():.2f} kW")',
    '',
    '# 计算该动作的对数概率（训练 Actor 时用）',
    'log_prob = policy_dist.log_prob(action)',
    'print(f"对数概率: {log_prob.item():.4f}")',
    '',
    '# 计算熵（衡量分布的不确定性）',
    'entropy = policy_dist.entropy()',
    'print(f"熵: {entropy.item():.4f}")',
], label='Normal 分布基本操作')

h2(doc, '11.3 安全处理标准差')

code_block(doc, [
    '# 标准做法：网络输出 log_std，然后 exp 得到 std',
    '# 保证 std 始终 > 0',
    '',
    'class PolicyNetwork(nn.Module):',
    '    def __init__(self, state_dim, action_dim, hidden=64):',
    '        super().__init__()',
    '        self.fc1 = nn.Linear(state_dim, hidden)',
    '        self.fc_mu = nn.Linear(hidden, action_dim)',
    '        self.fc_log_std = nn.Linear(hidden, action_dim)',
    '',
    '    def forward(self, x):',
    '        x = F.relu(self.fc1(x))',
    '        mu = self.fc_mu(x)',
    '        log_std = self.fc_log_std(x)',
    '        log_std = torch.clamp(log_std, -20, 2)  # 限制范围',
    '        std = log_std.exp()',
    '        return dist.Normal(mu, std)',
    '',
    '# 使用',
    'policy = PolicyNetwork(state_dim=3, action_dim=1)',
    's = torch.randn(1, 3)  # [SOC, P_load, SOC_error]',
    'action_dist = policy(s)',
    'a = action_dist.sample()',
    'print(f"动作: {a.item():.2f}")',
], label='策略网络 — 输出分布')

h2(doc, '11.4 Categorical 分布 — 离散动作')

code_block(doc, [
    '# 离散动作（DQN 用）',
    'logits = torch.tensor([1.0, 2.0, 0.5, 3.0])  # 原始分数（logits）',
    'cat_dist = dist.Categorical(logits=logits)',
    '',
    '# 采样动作索引',
    'action = cat_dist.sample()     # tensor(3) — 选 logits 最大的索引',
    'print(f"选中的动作: {action.item()}")',
    '',
    '# 每个动作的概率',
    'probs = F.softmax(logits, dim=-1)',
    'print(f"动作概率: {probs}")',
    '',
    '# 对数概率',
    'log_prob = cat_dist.log_prob(action)',
    'print(f"log_prob: {log_prob.item():.4f}")',
], label='Categorical 分布')

h2(doc, '11.5 Squash 动作到有效范围')

code_block(doc, [
    '# RL 中常用 tanh 将动作限制到 [-1, 1]',
    '# 然后缩放实际范围',
    '',
    '# 从 Normal 采样',
    'raw_action = policy_dist.sample()  # 无界',
    '',
    '# Squash 到 [-1, 1]',
    'squashed_action = torch.tanh(raw_action)',
    '',
    '# 缩放到实际范围 [0, 30] kW',
    'PFC_MIN, PFC_MAX = 0, 30',
    'actual_action = PFC_MIN + (squashed_action + 1) / 2 * (PFC_MAX - PFC_MIN)',
    '',
    '# 注意：squash 后要对数概率要修正（SAC 需要）',
    '# log π(a|s) = log π(u|s) - Σ log(1 - tanh(u)² + 1e-6)',
    'log_prob_corrected = log_prob - torch.log(1 - squashed_action**2 + 1e-6).sum()',
], label='动作范围处理')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第12章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第12章 进阶① — RNN / LSTM 时序预测')

h2(doc, '12.1 为什么 LSTM 比 MLP 更适合时序')
tx(doc, 'MLP 预测功率时，把 10 步历史当作 "10 个独立特征"——它不知道这 10 个数在时间上是连续的。LSTM 天然处理序列，通过 hidden state 传递时序信息。')

make_table(doc,
    ['对比维度', 'MLP', 'LSTM'],
    [
        ['输入处理', '展平为向量，丢失时序结构', '保持序列维度 (seq, batch, feature)'],
        ['时间依赖性', '不建模', '通过 hidden state 显式建模'],
        ['长程依赖', '窗口外信息完全丢失', '理论上可以捕捉'],
        ['参数量', '少', '多（需要更多数据）'],
        ['训练难度', '简单', '较难（梯度消失/爆炸）'],
    ],
    col_widths=[3, 5, 6]
)
doc.add_paragraph()

h2(doc, '12.2 LSTM 核心概念')

code_block(doc, [
    'lstm = nn.LSTM(input_size=1,   # 每步特征数（功率：1维）',
    '               hidden_size=32, # 隐藏层维度',
    '               num_layers=1,   # LSTM 层数',
    '               batch_first=True)  # 输入形状: (batch, seq, feature)',
    '',
    '# 输入: (batch=32, seq_len=10, input_size=1)',
    'x = torch.randn(32, 10, 1)',
    '',
    '# 前向传播',
    'output, (h_n, c_n) = lstm(x)',
    '',
    'print(f"output: {output.shape}")    # (32, 10, 32) — 每步输出',
    'print(f"h_n:    {h_n.shape}")       # (1, 32, 32) — 最后一步的 hidden state',
    'print(f"c_n:    {c_n.shape}")       # (1, 32, 32) — 最后一步的 cell state',
    '',
    '# 取最后一步的输出做预测',
    'last_output = output[:, -1, :]      # (32, 32)',
    'predictor = nn.Linear(32, 1)',
    'y_pred = predictor(last_output)     # (32, 1)',
], label='LSTM 基本使用')

h2(doc, '12.3 LSTM 功率预测模型')

code_block(doc, [
    'class LSTMPowerPredictor(nn.Module):',
    '    def __init__(self, input_size=1, hidden_size=32,',
    '                 num_layers=1, output_size=1):',
    '        super().__init__()',
    '        self.lstm = nn.LSTM(input_size, hidden_size,',
    '                            num_layers, batch_first=True)',
    '        self.fc = nn.Linear(hidden_size, output_size)',
    '',
    '    def forward(self, x):',
    '        # x: (batch, seq_len, input_size)',
    '        output, _ = self.lstm(x)',
    '        # 取最后一步的 hidden state',
    '        last_hidden = output[:, -1, :]  # (batch, hidden_size)',
    '        y_pred = self.fc(last_hidden)   # (batch, 1)',
    '        return y_pred',
    '',
    '# 使用（注意输入需要 unsqueeze 第三维）',
    'X_train_3d = X_train_t.unsqueeze(-1)  # (N, 10) → (N, 10, 1)',
    'model = LSTMPowerPredictor()',
    'y_pred = model(X_train_3d)',
    'print(f"输入: {X_train_3d.shape} → 输出: {y_pred.shape}")',
], label='LSTM 功率预测模型')

note(doc, 'LSTM 通常比 MLP 预测更准（MAE 可能从 1.8 降到 1.2），但训练更慢。在 EMS 项目中，先跑通 MLP，再用 LSTM 优化。')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第13章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第13章 进阶② — 用 PyTorch 实现 DQN')

h2(doc, '13.1 DQN 架构')

code_block(doc, [
    'class DQN(nn.Module):',
    '    """DQN 网络 — 输入状态，输出每个离散动作的 Q 值"""',
    '    def __init__(self, state_dim, action_dim, hidden=64):',
    '        super().__init__()',
    '        self.fc1 = nn.Linear(state_dim, hidden)',
    '        self.fc2 = nn.Linear(hidden, hidden)',
    '        self.fc3 = nn.Linear(hidden, action_dim)',
    '',
    '    def forward(self, x):',
    '        x = F.relu(self.fc1(x))',
    '        x = F.relu(self.fc2(x))',
    '        return self.fc3(x)  # 输出每个动作的 Q 值',
    '',
    '# 实例化',
    'state_dim = 3   # [SOC, P_load, SOC_error]',
    'action_dim = 7  # 离散化: 0,5,10,15,20,25,30 kW',
    'online_net = DQN(state_dim, action_dim)',
    'target_net = DQN(state_dim, action_dim)',
    'target_net.load_state_dict(online_net.state_dict())',
    'target_net.eval()  # 目标网络不训练',
], label='DQN 网络定义')

h2(doc, '13.2 经验回放 (Experience Replay)')

code_block(doc, [
    'from collections import deque',
    'import random',
    '',
    'class ReplayBuffer:',
    '    """经验回放缓冲区"""',
    '    def __init__(self, capacity=100000):',
    '        self.buffer = deque(maxlen=capacity)',
    '',
    '    def push(self, state, action, reward, next_state, done):',
    '        # 存一条经验 (s, a, r, s\', done)',
    '        self.buffer.append((state, action, reward,',
    '                            next_state, done))',
    '',
    '    def sample(self, batch_size):',
    '        batch = random.sample(self.buffer, batch_size)',
    '        # 转置 batch 并堆叠成 tensor',
    '        states, actions, rewards, next_states, dones =',
    '            zip(*batch)',
    '        return (torch.tensor(np.array(states), dtype=torch.float32),',
    '                torch.tensor(actions, dtype=torch.long).unsqueeze(1),',
    '                torch.tensor(rewards, dtype=torch.float32).unsqueeze(1),',
    '                torch.tensor(np.array(next_states), dtype=torch.float32),',
    '                torch.tensor(dones, dtype=torch.float32).unsqueeze(1))',
    '',
    '    def __len__(self):',
    '        return len(self.buffer)',
    '',
    '# 使用',
    'buffer = ReplayBuffer(capacity=100000)',
    '# buffer.push(s, a, r, s_next, done)',
    '# batch = buffer.sample(batch_size=64)',
], label='经验回放实现')

h2(doc, '13.3 DQN 训练关键：选择动作 + 计算 loss')

code_block(doc, [
    '# ── ε-greedy 动作选择 ──',
    'def select_action(dqn, state, epsilon):',
    '    """ε-greedy 策略：以 ε 概率随机探索"""',
    '    if random.random() < epsilon:',
    '        return random.randrange(action_dim)  # 随机',
    '    else:',
    '        with torch.no_grad():',
    '            q_values = dqn(state)',
    '            return q_values.argmax().item()  # 贪心',
    '',
    '# ── DQN loss 计算 ──',
    'def compute_dqn_loss(online_net, target_net, batch, gamma=0.99):',
    '    states, actions, rewards, next_states, dones = batch',
    '',
    '    # Q(s, a) — 在线网络预测',
    '    q_values = online_net(states).gather(1, actions)',
    '',
    '    # target = r + γ * max_a\' Q_target(s\', a\')',
    '    with torch.no_grad():',
    '        next_q = target_net(next_states).max(1, keepdim=True)[0]',
    '        target = rewards + gamma * next_q * (1 - dones)',
    '',
    '    loss = F.mse_loss(q_values, target)',
    '    return loss',
], label='DQN 训练核心函数')

h2(doc, '13.4 DQN 完整训练循环')

code_block(doc, [
    '# 超参数',
    'BATCH_SIZE = 64',
    'GAMMA = 0.99',
    'EPSILON_START = 1.0',
    'EPSILON_END = 0.01',
    'EPSILON_DECAY = 0.995',
    'TARGET_UPDATE = 100  # 目标网络硬更新间隔',
    '',
    'online_net = DQN(state_dim, action_dim).to(device)',
    'target_net = DQN(state_dim, action_dim).to(device)',
    'target_net.load_state_dict(online_net.state_dict())',
    'optimizer = torch.optim.Adam(online_net.parameters(), lr=0.0001)',
    'buffer = ReplayBuffer(capacity=100000)',
    '',
    'epsilon = EPSILON_START',
    'for episode in range(M):',
    '    state = env.reset()',
    '    total_reward = 0',
    '    for t in range(T):',
    '        # 选择动作',
    '        action = select_action(online_net, state, epsilon)',
    '        next_state, reward, done = env.step(action)',
    '        buffer.push(state, action, reward, next_state, done)',
    '',
    '        # 训练',
    '        if len(buffer) > BATCH_SIZE:',
    '            batch = buffer.sample(BATCH_SIZE)',
    '            loss = compute_dqn_loss(online_net, target_net,',
    '                                     batch, GAMMA)',
    '            optimizer.zero_grad()',
    '            loss.backward()',
    '            optimizer.step()',
    '',
    '        state = next_state',
    '        total_reward += reward',
    '        if done: break',
    '',
    '    # 更新目标网络',
    '    if episode % TARGET_UPDATE == 0:',
    '        target_net.load_state_dict(online_net.state_dict())',
    '',
    '    # 衰减 epsilon',
    '    epsilon = max(EPSILON_END, epsilon * EPSILON_DECAY)',
], label='DQN 完整训练循环')

tip(doc, '这是面试高频题！能手写 DQN 训练循环（5 个核心组件：网络×2、buffer、epsilon、loss）是 RL 岗位的 baseline 能力。')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第14章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第14章 进阶③ — 用 PyTorch 实现 PPO Actor-Critic')

h2(doc, '14.1 PPO 的核心思想')
tx(doc, 'PPO (Proximal Policy Optimization) 是目前最主流的 RL 算法之一。它的核心是：在每次更新时，不要让新策略偏离旧策略太远。')
tx(doc, '实现方式：用 clipped surrogate objective 限制策略更新幅度。')

make_table(doc,
    ['组件', '功能', '输入', '输出'],
    [
        ['Actor (策略网络)', '输出动作分布', '状态 s', '动作的分布参数 (μ, σ)'],
        ['Critic (价值网络)', '评估状态的价值', '状态 s', '标量值 V(s)'],
        ['GAE (广义优势估计)', '计算优势 A(s,a)', '奖励序列 + V(s)', '优势估计'],
        ['Clip 机制', '限制策略更新幅度', '新旧策略概率比', '裁剪后的 loss'],
    ],
    col_widths=[4, 4, 4, 4]
)
doc.add_paragraph()

h2(doc, '14.2 共享参数 Actor-Critic')

code_block(doc, [
    'class ActorCritic(nn.Module):',
    '    """共享隐藏层的 Actor-Critic 网络（参数更少）"""',
    '    def __init__(self, state_dim, action_dim, hidden=64):',
    '        super().__init__()',
    '        # 共享层',
    '        self.fc1 = nn.Linear(state_dim, hidden)',
    '        self.fc2 = nn.Linear(hidden, hidden)',
    '',
    '        # Actor 头：输出动作分布',
    '        self.fc_mu = nn.Linear(hidden, action_dim)',
    '        self.fc_log_std = nn.Linear(hidden, action_dim)',
    '',
    '        # Critic 头：输出状态价值',
    '        self.fc_value = nn.Linear(hidden, 1)',
    '',
    '    def forward(self, x):',
    '        """前向传播，返回动作分布和状态价值"""',
    '        x = F.relu(self.fc1(x))',
    '        x = F.relu(self.fc2(x))',
    '',
    '        # Actor',
    '        mu = self.fc_mu(x)',
    '        log_std = torch.clamp(self.fc_log_std(x), -20, 2)',
    '        std = log_std.exp()',
    '        dist = dist.Normal(mu, std)',
    '',
    '        # Critic',
    '        value = self.fc_value(x)',
    '',
    '        return dist, value',
    '',
    'model = ActorCritic(state_dim=3, action_dim=1)',
    's = torch.randn(1, 3)',
    'action_dist, value = model(s)',
    'print(f"动作分布: μ={action_dist.mean.item():.2f},',
    '       f"σ={action_dist.stddev.item():.2f}")',
    'print(f"状态价值: {value.item():.3f}")',
], label='Actor-Critic 网络')

h2(doc, '14.3 PPO-Clip Loss')

code_block(doc, [
    'def compute_ppo_loss(model, old_model, batch, clip_eps=0.2):',
    '    """PPO-Clip loss"""',
    '    states, actions, old_log_probs, advantages, returns = batch',
    '',
    '    # 新策略的 log_prob 和 value',
    '    dist, values = model(states)',
    '    log_probs = dist.log_prob(actions).sum(dim=-1)',
    '    entropy = dist.entropy().sum(dim=-1).mean()',
    '',
    '    # 概率比 r(θ) = π_θ / π_θ_old',
    '    ratios = (log_probs - old_log_probs).exp()',
    '',
    '    # Clipped surrogate objective',
    '    surr1 = ratios * advantages',
    '    surr2 = torch.clamp(ratios, 1 - clip_eps, 1 + clip_eps) * advantages',
    '    actor_loss = -torch.min(surr1, surr2).mean()',
    '',
    '    # Critic loss（MSE）',
    '    critic_loss = F.mse_loss(values.squeeze(), returns)',
    '',
    '    # 总 loss',
    '    loss = actor_loss + 0.5 * critic_loss - 0.01 * entropy',
    '',
    '    return loss, {"actor": actor_loss.item(),',
    '                   "critic": critic_loss.item(),',
    '                   "entropy": entropy.item()}',
], label='PPO-Clip Loss 函数')

h2(doc, '14.4 GAE (广义优势估计)')

code_block(doc, [
    'def compute_gae(rewards, values, dones, gamma=0.99, lam=0.95):',
    '    """',
    '    计算 GAE 优势函数和 returns',
    '    GAE 是 TD-error 的指数加权平均',
    '    lam=0 → 1步TD, lam=1 → MC',
    '    """',
    '    advantages = []',
    '    gae = 0',
    '    for t in reversed(range(len(rewards))):',
    '        if t == len(rewards) - 1:',
    '            next_value = 0 if dones[t] else values[t+1]',
    '        else:',
    '            next_value = values[t+1]',
    '',
    '        delta = rewards[t] + gamma * next_value * (1-dones[t]) - values[t]',
    '        gae = delta + gamma * lam * (1-dones[t]) * gae',
    '        advantages.insert(0, gae)',
    '',
    '    advantages = torch.tensor(np.array(advantages), dtype=torch.float32)',
    '    returns = advantages + torch.tensor(values, dtype=torch.float32)',
    '',
    '    # 标准化优势（稳定训练，可选）',
    '    advantages = (advantages - advantages.mean()) / (advantages.std() + 1e-8)',
    '',
    '    return advantages, returns',
], label='GAE 计算')

h2(doc, '14.5 PPO 训练主循环')

code_block(doc, [
    '# 超参数',
    'LR = 0.0003',
    'GAMMA = 0.99',
    'GAE_LAM = 0.95',
    'CLIP_EPS = 0.2',
    'N_STEPS = 2048     # 每次收集多少步经验',
    'N_EPOCHS = 10      # 每个 batch 训练几轮',
    'BATCH_SIZE = 64',
    '',
    'model = ActorCritic(state_dim=3, action_dim=1)',
    'optimizer = torch.optim.Adam(model.parameters(), lr=LR)',
    '',
    'for iteration in range(N_ITERATIONS):',
    '    # ── 收集经验 ──',
    '    states, actions, rewards, dones, values, log_probs = [],',
    '        [], [], [], [], []',
    '    state = env.reset()',
    '    for _ in range(N_STEPS):',
    '        s_t = torch.tensor(state, dtype=torch.float32)',
    '        with torch.no_grad():',
    '            dist, value = model(s_t)',
    '            action = dist.sample()',
    '            log_prob = dist.log_prob(action).sum()',
    '',
    '        next_state, reward, done = env.step(action.numpy())',
    '        states.append(state)',
    '        actions.append(action)',
    '        rewards.append(reward)',
    '        dones.append(done)',
    '        values.append(value.item())',
    '        log_probs.append(log_prob.item())',
    '        state = next_state',
    '        if done: state = env.reset()',
    '',
    '    # ── 计算 GAE ──',
    '    advantages, returns = compute_gae(rewards, values, dones,',
    '                                       GAMMA, GAE_LAM)',
    '',
    '    # ── PPO 更新 ──',
    '    dataset = list(zip(states, actions, log_probs,',
    '                        advantages, returns))',
    '    for _ in range(N_EPOCHS):',
    '        random.shuffle(dataset)',
    '        for i in range(0, len(dataset), BATCH_SIZE):',
    '            batch = dataset[i:i+BATCH_SIZE]',
    '            # 转成 tensor 并移到 device',
    '            # ...',
    '            loss, info = compute_ppo_loss(model, batch)',
    '            optimizer.zero_grad()',
    '            loss.backward()',
    '            torch.nn.utils.clip_grad_norm_(model.parameters(), 0.5)',
    '            optimizer.step()',
], label='PPO 训练主循环')

tip(doc, 'PPO 是 EMS 项目的核心算法。上面的代码框架就是你第 11-12 周要跑通的内容。建议先理解整体流程，再逐行实现。')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第15章
# ═══════════════════════════════════════════════════════════════
h1(doc, '第15章 最佳实践与调试技巧')

h2(doc, '15.1 设置随机种子')
code_block(doc, [
    'def set_seed(seed=42):',
    '    """固定所有随机种子，保证可复现"""',
    '    torch.manual_seed(seed)',
    '    torch.cuda.manual_seed_all(seed)',
    '    np.random.seed(seed)',
    '    random.seed(seed)',
    '    torch.backends.cudnn.deterministic = True',
    '    torch.backends.cudnn.benchmark = False',
    '',
    'set_seed(42)',
], label='设置随机种子')

h2(doc, '15.2 常见错误与调试')
make_table(doc,
    ['错误信息', '原因', '解决办法'],
    [
        ['Expected all tensors on same device', '模型和输入在不同设备', '统一 .to(device)'],
        ['Trying to backward graph second time', '忘清零梯度', 'optimizer.zero_grad()'],
        ['CUDA out of memory', '显存不足', '减小 batch_size'],
        ['size mismatch', '输入输出形状不匹配', '检查网络定义和输入形状'],
        ['NaN in loss', '梯度爆炸/学习率太大', '减小 lr / 梯度裁剪'],
        ['RuntimeError: element 0 of tensors does not require grad', '参数没设梯度', '检查 requires_grad'],
    ],
    col_widths=[5, 4, 5]
)
doc.add_paragraph()

h2(doc, '15.3 梯度裁剪')
code_block(doc, [
    '# 防止梯度爆炸（RL 训练特别重要）',
    'torch.nn.utils.clip_grad_norm_(',
    '    model.parameters(), max_norm=1.0)',
    '# 在 loss.backward() 之后、optimizer.step() 之前调用',
], label='梯度裁剪')

h2(doc, '15.4 监控训练')
bl(doc, '每 epoch 打印 train_loss / val_loss')
bl(doc, '验证集上计算 MAE / RMSE / R²')
bl(doc, '保存最佳模型（val_loss 最低时）')
bl(doc, '画 loss 曲线 — 判断过拟合、欠拟合')
bl(doc, '用 TensorBoard 或 wandb 做高级监控')

code_block(doc, [
    'from torch.utils.tensorboard import SummaryWriter',
    '',
    'writer = SummaryWriter("runs/exp_001")',
    'for epoch in range(n_epochs):',
    '    # ... 训练 ...',
    '    writer.add_scalar("Loss/train", loss.item(), epoch)',
    '    writer.add_scalar("Loss/val", val_loss, epoch)',
    '    writer.add_histogram("weights/fc1", ',
    '        model.fc1.weight, epoch)',
    'writer.close()',
], label='TensorBoard 监控')

h2(doc, '15.5 PyTorch 训练代码模板 (速查)')

code_block(doc, [
    '# === 完整模板 ===',
    'device = torch.device("cuda" if torch.cuda.is_available() else "cpu")',
    '',
    'model = MyModel().to(device)',
    'optimizer = torch.optim.Adam(model.parameters(), lr=0.001)',
    'loss_fn = nn.MSELoss()',
    'scheduler = torch.optim.lr_scheduler.StepLR(optimizer, 30, 0.1)',
    '',
    'for epoch in range(n_epochs):',
    '    # 训练',
    '    model.train()',
    '    for x, y in train_loader:',
    '        x, y = x.to(device), y.to(device)',
    '        optimizer.zero_grad()',
    '        loss = loss_fn(model(x), y)',
    '        loss.backward()',
    '        torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)',
    '        optimizer.step()',
    '',
    '    # 验证',
    '    model.eval()',
    '    val_loss = 0',
    '    with torch.no_grad():',
    '        for x, y in val_loader:',
    '            x, y = x.to(device), y.to(device)',
    '            val_loss += loss_fn(model(x), y).item()',
    '',
    '    scheduler.step()',
    '    print(f"Epoch {epoch}: {loss.item():.4f} / {val_loss:.4f}")',
], label='训练代码完整模板')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 附录A
# ═══════════════════════════════════════════════════════════════
h1(doc, '附录 A — PyTorch 面试八股文 20 问')

qa_pairs = [
    ("Q1: Tensor 和 numpy 有什么区别？",
     "Tensor 支持 GPU 运算、自动求导（requires_grad）、计算图追踪。numpy 只有 CPU 数值计算。两者共享内存：from_numpy() 转换后修改一个会影响另一个。"),

    ("Q2: backward() 的输入为什么必须是标量？",
     "因为梯度是标量对张量的导数。如果输出是向量，需要传一个同形状的 gradient 参数，表示各分量的权重。通常 loss 是标量，所以不用传参数。"),

    ("Q3: optimizer.zero_grad() 作用是什么？忘记调会怎样？",
     "清零所有参数的梯度。PyTorch 默认累积梯度，不清零会导致每个 batch 的梯度全部累加，参数更新方向错误。"),

    ("Q4: model.train() 和 model.eval() 的区别？",
     "train() 启用 Dropout 和 BatchNorm 的训练行为；eval() 关闭 Dropout，BatchNorm 使用固定统计量。推理时必须 eval() + no_grad()。"),

    ("Q5: 什么是计算图？怎么查看？",
     "计算图是 PyTorch 记录运算过程的 DAG。节点是 tensor，边是操作。反向传播时沿图从 loss 走到每个参数。可以用 grad_fn 属性查看。"),

    ("Q6: no_grad() 和 detach() 的区别？",
     "no_grad() 是上下文管理器，范围内所有操作不追踪梯度。detach() 是对单个 tensor 截断梯度流，返回一个共享数据但不参与计算图的新 tensor。"),

    ("Q7: 什么是梯度爆炸/消失？怎么解决？",
     "梯度爆炸：loss→NaN，原因：学习率太大/网络太深。解决：梯度裁剪、减小 lr、BatchNorm。消失：梯度→0，原因：sigmoid 饱和/层数太深。解决：ReLU、残差连接。"),

    ("Q8: nn.Sequential 和自定义 forward 的适用场景？",
     "Sequential：线性堆叠的网络（无分支）。自定义 forward：有残差连接、多输入、多输出、条件分支的网络。"),

    ("Q9: 什么是过拟合？PyTorch 中怎么防止？",
     "模型在训练集上表现好但在测试集上差。防止：Dropout、权重衰减（weight_decay）、早停（Early Stopping）、数据增强、减少模型复杂度。"),

    ("Q10: 模型参数初始化有哪些方法？分别适用场景？",
     "Xavier：tanh/sigmoid 激活。Kaiming (He)：ReLU 激活（PyTorch Linear 默认）。Orthogonal：RL 策略网络常用。常数：bias 通常初始化为 0。"),

    ("Q11: DataLoader 的 num_workers 在 Windows 上为什么设为 0？",
     "Windows 的多进程数据加载（num_workers>0）可能与 PyTorch 的 multiprocessing 有兼容性问题，导致死锁或报错。Linux/macOS 可以安全使用 num_workers>0 加速。"),

    ("Q12: 什么是学习率调度器？常用的有哪些？",
     "在训练过程中动态调整学习率。常用：StepLR（每 N 轮 ×γ）、CosineAnnealingLR（余弦退火）、ReduceLROnPlateau（验证 loss 不降时减 lr）。"),

    ("Q13: 什么是梯度累积？什么场景需要？",
     "多个 batch 的梯度累加后再更新参数。用于显存不够但想用大 batch_size 的场景：每个小 batch 调 backward()，但不 step()，累积 N 次后再 step()。"),

    ("Q14: 怎么实现 Distributed Data Parallel (DDP)？",
     "多 GPU 训练标准方法：用 torch.distributed.launch 启动，每个进程一个 GPU。DistributedSampler 自动分数据。DDP 自动同步梯度。本项目目前不需要。"),

    ("Q15: PyTorch 和 TensorFlow 的核心设计区别？",
     "PyTorch：动态图（define-by-run），调试方便，Pythonic，学术圈主流。TF：静态图（define-then-run，2.x 也有 eager），部署生态好，工业界多。"),

    ("Q16: 为什么 RL 的 learning rate 通常比监督学习小？",
     "RL 的梯度本身噪声大（采样 + 自举估计），大 lr 容易发散。监督学习的梯度来自真实标签，噪声小。RL 典型 lr=3e-4，监督学习 1e-3。"),

    ("Q17: 什么是 replay buffer 的优先级采样（PER）？",
     "不是均匀采样，而是给 TD-error 大的经验更高的采样权重。优点是学习效率更高。缺点是实现复杂、引入 bias（需要重要性采样校正）。"),

    ("Q18: target network 在 DQN 中的作用？软更新和硬更新的区别？",
     "作用：固定 target Q 值，使训练目标稳定。硬更新：每 C 步将在线网络参数完全复制给目标网络。软更新（Polyak）：每一步 θ_target ← τ·θ_online + (1-τ)·θ_target，τ=0.005。"),

    ("Q19: 什么是 PyTorch 的 autograd 中的 hook？",
     "register_hook 可以在反向传播时拦截梯度。用于：梯度裁剪、打印梯度值、修改梯度（如 GAN 的梯度反转）。调试时可以用 hook 查看中间层的梯度。"),

    ("Q20: torch.jit.script 和 torch.onnx.export 的用途？",
     "模型部署相关：torch.jit.script 将 PyTorch 模型序列化为 TorchScript（可脱离 Python 运行）。torch.onnx.export 导出 ONNX 格式（跨框架/硬件部署）。本项目到第 5 个月才需要。"),
]

for q, a in qa_pairs:
    h2(doc, q)
    tx(doc, a)
    doc.add_paragraph()

# ── 结尾 ──
sep(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PyTorch 从零到项目上手 · 完整指南')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('生成日期：2026-07-10 · EMS 研究项目 · 秋招面试准备')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# — 保存 —
OUT = 'F:/CLAUDE/research/ems-platform/docs/PyTorch_从零到EMS项目上手_完整指南.docx'
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f'[OK] 文档已保存: {OUT}')
print(f'     文件大小: {os.path.getsize(OUT) / 1024:.1f} KB')
