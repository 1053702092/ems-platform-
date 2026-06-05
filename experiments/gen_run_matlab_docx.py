# -*- coding: utf-8 -*-
"""生成 run_ems_matlab.m 解释文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(t).font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

doc = Document()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('run_ems_matlab.m 逐段解释'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Day7 - MATLAB仿真运行脚本 | 文件: env/simulink_models/Use-Model/run_ems_matlab.m'); r.font.size = Pt(11)
doc.add_paragraph()

# ===== 作用 =====
h1(doc, '文件作用')
tx(doc, '在 MATLAB 中运行 EMS_hybrid_v1 仿真，从 simOut 读取 To Workspace 数据，保存 CSV 并输出统计结果。')
tx(doc, '直接在命令窗口执行 run_ems_matlab 即可一键仿真+出结果。')
doc.add_paragraph()

# ===== 执行流程 =====
h1(doc, '执行流程')
tbl(doc, ['步骤', '行号', '功能'],
[['1. 准备', '6-9', '设置路径，cd 到 Use-Model/ 目录'],
 ['2. 加载模型+数据', '12-15', 'load_system + csvread 读 WLTC'],
 ['3. 运行仿真', '18-19', 'sim() 跑 1800s'],
 ['4. 读取结果', '22-28', '用 get_sim_var 从 simOut 读变量'],
 ['5. 保存 CSV', '31-36', 'writetable 到 results/'],
 ['6. 输出统计', '39-44', 'trapz 积分算能量']])
doc.add_paragraph()

# ===== 关键代码逐段 =====
h1(doc, '关键代码逐段详解')

h2(doc, '1. cd + load_system')
cd(doc, "load_system('EMS_hybrid_v1');")
cd(doc, "data = csvread(fullfile(RESULTS_DIR, 'wltc_cycle.csv'), 1, 0);")
cd(doc, "assignin('base', 'sim_wltc', [data(:,1), data(:,2)]);")
tx(doc, '先加载模型，再读 WLTC 工况 CSV（跳过表头），拼成 [t, v] 矩阵存入 base workspace。')
tx(doc, 'From Workspace 块在仿真时从这里读取 sim_wltc 变量。')
doc.add_paragraph()

h2(doc, '2. 运行仿真')
cd(doc, "sim('EMS_hybrid_v1', 'StopTime', '1800');")
tx(doc, '启动 Simulink 仿真，运行 1800s（WLTC 全长）。仿真过程中 To Workspace 块自动把数据保存到 simOut 对象。')
tx(doc, '用无输出的 sim() 调用（不赋值），数据会保存在 base workspace 或 simOut 中。')
doc.add_paragraph()

h2(doc, '3. 读取结果（核心难点）')
cd(doc, 'P_load = get_sim_var(simOut, "sim_P_load");')
cd(doc, 'P_fc   = get_sim_var(simOut, "sim_P_fc");')
cd(doc, 'SOC    = get_sim_var(simOut, "sim_SOC");')
cd(doc, 'V_bat  = get_sim_var(simOut, "sim_V_bat");')
cd(doc, 'V_fc   = get_sim_var(simOut, "sim_V_fc");')
cd(doc, 'I_bat  = get_sim_var(simOut, "sim_I_bat");')
cd(doc, 'status = get_sim_var(simOut, "sim_status");')

tx(doc, '从 simOut 读取 7 个 To Workspace 变量。get_sim_var 是文件末尾的辅助函数，自动处理三种数据格式：')
tbl(doc, ['simOut 存储格式', '内部结构', '如何取值'],
[['Array', '普通 double 数组', '直接返回'],
 ['Timeseries', '带时间戳的复杂对象', '.Data 属性'],
 ['Structure', '旧版结构体', '.signals.values']])
tx(doc, '最后用 double(v(:)) 确保输出是双精度列向量，方便后面表格处理和积分。')
doc.add_paragraph()

h2(doc, '4. 保存 CSV')
cd(doc, 'T = table(t_vec, P_load, P_fc, SOC, V_bat, ...);')
cd(doc, 'writetable(T, csv_path);')
tx(doc, '用 MATLAB 的 table + writetable 直接写出 CSV。第一列是仿真时间（用 tout 或 linspace 生成）。')
doc.add_paragraph()

h2(doc, '5. 积分统计')
cd(doc, "fprintf('  总能量: %.2f kWh\\n', trapz(T.time, T.P_load_kW) / 3600);")
tx(doc, 'trapz(t, y) 用实际时间向量做数值积分，÷3600 把 W·s 转成 kWh。')
tx(doc, '注意：trapz 必须传入时间向量！如果漏掉 t，trapz(y) 会假设步长为 1s，对于 0.1s 步长的数据结果会差 10 倍。')
doc.add_paragraph()

# ===== get_sim_var 详解 =====
h1(doc, '辅助函数 get_sim_var 详解')
cd(doc, 'function v = get_sim_var(simOut, name)')
cd(doc, '    val = simOut.get(name);')
cd(doc, '    if isa(val, "timeseries")')
cd(doc, '        v = val.Data;')
cd(doc, '    elseif isstruct(val) && isfield(val, "signals")')
cd(doc, '        v = val.signals.values;')
cd(doc, '    else')
cd(doc, '        v = val;')
cd(doc, '    end')
cd(doc, '    v = double(v(:));')
cd(doc, 'end')

tx(doc, '这个函数解决一个核心问题：To Workspace 块在不同设置下存的数据格式不一样。')
tbl(doc, ['if 条件', '匹配的格式', 'MATLAB 版本'],
[['isa(val, "timeseries")', 'Timeseries 格式', 'R2016a+ 默认'],
 ['isstruct(val) 且有 signals', 'Structure 格式', '老版本'],
 ['else', 'Array 格式', '脚本中手动设置']])
tx(doc, 'double(v(:)) 的两步操作：')
bl(doc, 'v(:) 把行向量 [1x18001] 或矩阵转为列向量 [18001x1]', bp='压成一列：')
bl(doc, 'double() 确保是双精度，后续积分不会报类型错误', bp='转双精度：')
doc.add_paragraph()

# ===== 常见问题 =====
h1(doc, '常见问题')

h2(doc, 'Q1: 数据全是 0')
tx(doc, '原因：simOut 里没有对应的变量。检查 To Workspace 块的 VariableName 是否拼写正确，或连线是否断开。')

h2(doc, 'Q2: trapz 结果比预期大 10 倍')
tx(doc, '原因：没传时间向量 t。trapz(y) 假设步长=1，但 simOut 的数据步长是 0.1s。修复：用 trapz(t, y)。')

h2(doc, 'Q3: 变量名不存在报错')
tx(doc, '原因：To Workspace 块的 SaveFormat 不是 Array，或者 simOut 中变量名匹配失败。'
         '修复：在 build_ems_model 中设置 SaveFormat=Array，或在 get_sim_var 中增加新的格式判断。')

doc.add_paragraph()
doc.add_paragraph()
r = doc.add_paragraph().add_run('='*50+'\nrun_ems_matlab.m 逐段解释\n文件: env/simulink_models/Use-Model/run_ems_matlab.m\n生成日期：2026-06-05\n'+'='*50)
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

fname = 'F:/CLAUDE/research/ems-platform/docs/Day7_run_EMS_matlab_explain.docx'
doc.save(fname)
print('OK:', fname)
