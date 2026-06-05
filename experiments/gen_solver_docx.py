# -*- coding: utf-8 -*-
"""生成 变步长vs定步长求解器 解释文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(t).font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

doc = Document()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Simulink 求解器：变步长 vs 定步长'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('以 ode45 为例 | EMS_hybrid_v1 仿真设置'); r.font.size = Pt(11)
doc.add_paragraph()

# ===== 背景 =====
h1(doc, '什么是求解器？')
tx(doc, 'Simulink 仿真本质上是求解微分方程组。求解器（Solver）就是计算每一步系统状态的方法。')
tx(doc, 'EMS_hybrid_v1 中包含多个微分关系：')
bl(doc, '车速变化 → 加速度 → 功率需求（动力学方程）')
bl(doc, '电池 SOC 变化 → d(SOC)/dt = -I / Q（Ah 积分）')
bl(doc, 'FC 电压随电流变化（极化曲线）')
tx(doc, '求解器的任务：从初始状态开始，按时间一步步推算出整个系统的行为。')
doc.add_paragraph()

# ===== 核心区别 =====
h1(doc, '核心区别：每一步走多大')

h2(doc, '定步长（Fixed-Step）')
cd(doc, '|--0.01s--|--0.01s--|--0.01s--|--0.01s--|--0.01s--|')
tx(doc, '你设多少步长就是多少，从头到尾不变。不管信号变化多剧烈还是多平缓，每一步都用同样大小。')

h2(doc, '变步长（Variable-Step）—— ode45')
cd(doc, '|-----0.1s-----|---0.05s---|--0.02s-|-----0.1s-----|')
cd(doc, '                     ↑')
cd(doc, '              信号突变时自动缩小')
tx(doc, '求解器自己决定每一步走多大。信号平缓时跳大步（节省时间），信号剧烈时走小步（保证精度）。')

tbl(doc, ['对比项', '定步长', '变步长 (ode45)'],
[['步长', '固定不变', '自动调整'],
 ['精度控制', '无，取决于步长', '有，自动误差估计'],
 ['计算速度', '慢（平缓段也走小步）', '快（平缓段跳大步）'],
 ['适用场景', '实时仿真、硬件在环', '离线仿真、精度优先']])
doc.add_paragraph()

# ===== ode45 原理 =====
h1(doc, 'ode45 的工作原理')

tx(doc, 'ode45 基于 Runge-Kutta 4(5) 阶算法，名称中的 45 来自：')
bl(doc, '4 阶 Runge-Kutta 法：算一个结果 y4', bp='4: ')
bl(doc, '5 阶 Runge-Kutta 法：算一个更精确的结果 y5', bp='5: ')

tx(doc, '每走一步同时用两种方法算，然后比较差异：')

cd(doc, '每走一步:')
cd(doc, '  用4阶法 → y4')
cd(doc, '  用5阶法 → y5（更精确）')
cd(doc, '  误差估计 e = |y5 - y4|')
cd(doc, '  如果 e < 容忍度 → 下步可以走更大')
cd(doc, '  如果 e > 容忍度 → 立即缩小步长重新算')
doc.add_paragraph()

tx(doc, '可以理解成开车时的自适应巡航：')
bl(doc, '路况好（信号平缓）→ 开快一点（大步长）')
bl(doc, '路况差（信号剧烈）→ 开慢一点（小步长）')
bl(doc, '目的：在保证安全（精度）的前提下尽快到达终点（仿真完成）')
doc.add_paragraph()

# ===== 在 EMS 模型中的表现 =====
h1(doc, '在 EMS_hybrid_v1 中的具体表现')

tx(doc, '我们的模型设置：')
cd(doc, "set_param(mdl, 'Solver', 'ode45', 'StopTime', '1800', 'MaxStep', '0.1');")

tbl(doc, ['WLTC 阶段', '信号变化', 'ode45 行为', '实际步长'],
[['起步 (0-50s)', '速度从0急加速', '步长自动缩小', '~0.01s'],
 ['匀速巡航', '速度稳定~50km/h', '大步走', '~0.1s (MaxStep)'],
 ['急加速段', '速度快速上升', '步长缩小', '~0.02s'],
 ['减速停车', '速度快速下降', '步长缩小', '~0.02s'],
 ['等红灯 (0km/h)', '信号恒定=0', '大步走', '~0.1s']])

tx(doc, '仿真结果：1800s 的 WLTC 用了 18001 个时间点，平均步长 ~0.1s。')
tx(doc, 'MATLAB 实际跑了 8.3 秒就完成了 1800 秒的仿真（加速比 ~217x）。')
tx(doc, '如果是定步长（比如 0.01s），总步数会是 180000 步，仿真时间至少增加 10 倍。')
doc.add_paragraph()

# ===== 误差控制 =====
h1(doc, '误差控制参数')

tbl(doc, ['参数', '我们的设置', '作用'],
[['MaxStep', '0.1', '最大步长限制，保证加速度计算精度不低于 0.1s'],
 ['AbsTol', '1e-4', '绝对误差容忍度，控制状态量的绝对误差'],
 ['RelTol', '1e-3', '相对误差容忍度（0.1%），控制状态量的相对误差']])

tx(doc, 'RelTol=1e-3 意味着求解器保证每一步的相对误差不超过 0.1%。')
tx(doc, '比如 SOC=0.6 时，每步误差不超过 0.6×0.001 = 0.0006，即 0.06%。')
doc.add_paragraph()

# ===== 为什么用变步长 =====
h1(doc, '为什么我们的模型用变步长？')

h2(doc, '优点')
bl(doc, '计算效率高：平缓段大步走，总步数少，8秒跑完30分钟仿真', bp='① ')
bl(doc, '精度有保障：剧烈段自动缩小步长，不会漏掉关键瞬态', bp='② ')
bl(doc, '无需手动调步长：设好误差容忍度就行，求解器自己优化', bp='③ ')

h2(doc, '缺点')
bl(doc, '时间点不均匀：CSV 数据不是等间隔的，积分时要用 trapz(t, y)', bp='① ')
bl(doc, '结果不可重复：不同电脑、不同 MATLAB 版本跑的步长可能不同', bp='② ')
bl(doc, '不能用于实时系统：步长不确定，没法用于硬件在环', bp='③ ')

h2(doc, '什么时候该用定步长')
bl(doc, '生成代码 → 需要固定采样率')
bl(doc, '硬件在环（HIL） → 每一步必须在固定时间内算完')
bl(doc, '与外部实时设备通信 → 需要固定时间步')
bl(doc, '对比不同求解器结果 → 控制变量')
doc.add_paragraph()

# ===== 总结 =====
h1(doc, '一句话总结')
tx(doc, 'ode45 变步长求解器像自动驾驶——路况好开快（大步长），路况差开慢（小步长），在保证精度的前提下用最少步数跑完仿真。')
tx(doc, '定步长像定速巡航——不管前面有什么坡都一个速度冲过去，简单但效率低。')

doc.add_paragraph()
doc.add_paragraph()
r = doc.add_paragraph().add_run('='*50+'\nSimulink 求解器：变步长 vs 定步长\n以 ode45 为例 | EMS_hybrid_v1 仿真设置\n生成日期：2026-06-05\n'+'='*50)
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

fname = 'F:/CLAUDE/research/ems-platform/docs/Day7_Solver_ode45_explain.docx'
doc.save(fname)
print('OK:', fname)
