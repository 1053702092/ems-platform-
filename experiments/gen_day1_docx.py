# -*- coding: utf-8 -*-
"""生成 Day1 代码解释文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(t).font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

doc = Document()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Day1 代码逐段解释'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('第1周第1天 - Python数据处理入门 | 文件: experiments/day1_pandas_intro.py'); r.font.size = Pt(11)
doc.add_paragraph()

# ===== 整体结构 =====
h1(doc, '整体结构')
tx(doc, '脚本分为三部分：')
bl(doc, '生成模拟WLTC工况车速和功率数据，保存CSV')
bl(doc, 'pandas读取CSV，做统计和筛选')
bl(doc, 'matplotlib画双子图 + 双y轴图')
doc.add_paragraph()

# ===== 导入库 =====
h1(doc, '第一段：导入库')
cd(doc, 'import pandas as pd')
cd(doc, 'import matplotlib.pyplot as plt')
cd(doc, 'import numpy as np')
bl(doc, 'pandas：数据处理（读CSV、统计、筛选）')
bl(doc, 'matplotlib.pyplot：画图')
bl(doc, 'numpy：数值计算（数组、数学函数、随机数）')
doc.add_paragraph()

# ===== 生成工况数据 =====
h1(doc, '第二段：生成模拟工况数据')

h2(doc, '固定随机种子')
cd(doc, "np.random.seed(42)")
tx(doc, '固定随机种子，每次运行结果一致，确保可复现。')

h2(doc, '生成时间轴')
cd(doc, "t = np.arange(0, 1800, 1)  # 1800秒")
tx(doc, 'arange(起始, 结束, 步长) = [0,1,2,...,1799]，30分钟。')

h2(doc, '模拟车速曲线')
cd(doc, "v = np.zeros_like(t)")
cd(doc, "# 低速段: sin(0->pi)*15 = 0->15->0 km/h")
cd(doc, "v[0:200] = np.sin(np.linspace(0, np.pi, 200)) * 15")
cd(doc, "v[200:400] = 15 + np.sin(np.linspace(0, np.pi*2, 200)) * 5")
cd(doc, "v[400:600] = np.linspace(15, 0, 200)")
cd(doc, "# 中速段: sin(0->pi)*25+25 = 25->50->25 km/h")
cd(doc, "v[600:800] = np.sin(np.linspace(0, np.pi, 200)) * 25 + 25")
cd(doc, "v[800:1000] = 50 + np.sin(np.linspace(0, np.pi, 200)) * 5")
cd(doc, "v[1000:1200] = np.linspace(50, 0, 200)")
cd(doc, "# 高速段: 35->70->35 km/h")
cd(doc, "v[1200:1400] = np.sin(np.linspace(0, np.pi, 200)) * 35 + 35")
cd(doc, "v[1400:1600] = 70 + np.sin(np.linspace(0, np.pi, 200)) * 10")
cd(doc, "v[1600:1800] = np.linspace(70, 0, 200)")
cd(doc, "v = np.clip(v, 0, None)")

tbl(doc, ['段', '时间', '车速范围', '实现方式'],
[['低速', '0-600s', '0-15-0 km/h', 'sin(0->pi)*15 先升后降'],
['中速', '600-1200s', '25-50-0 km/h', 'sin(0->pi)*25+25 波动上升'],
['高速', '1200-1800s', '35-70-0 km/h', 'sin(0->pi)*35+35 高速波动']])

bl(doc, 'linspace(0, pi, 200)：在0~pi之间均匀取200个数', bp='linspace：')
bl(doc, 'v[0:200] = ...：数组切片，给第0~199个元素赋值', bp='切片：')
bl(doc, 'clip(v, 0, None)：把所有负数变成0', bp='clip：')
doc.add_paragraph()

h2(doc, '功率需求计算')
cd(doc, "P_demand = 0.005*v**2 + 0.1*v + 2 + np.random.normal(0,0.5,len(v))")
tbl(doc, ['项', '值', '物理意义'],
[['空气阻力', '0.005*v^2', '车速平方增长'],
['滚动阻力', '0.1*v', '线性增长'],
['基础负载', '2 kW', '空调/灯光等固定用电'],
['噪声', '+-0.5kW', '模拟驾驶波动']])
doc.add_paragraph()

h2(doc, '保存CSV')
cd(doc, "df_cycle = pd.DataFrame({'time':t, 'speed':v, 'power_demand':P_demand})")
cd(doc, "df_cycle.to_csv('results/wltc_sample.csv', index=False)")
tx(doc, 'pd.DataFrame(dict)用字典创建表格；to_csv保存CSV，index=False不写行号。')
doc.add_paragraph()

# ===== pandas操作 =====
h1(doc, '第三段：pandas基础操作')
cd(doc, "df = pd.read_csv('results/wltc_sample.csv')")
cd(doc, "df.head()       # 前5行")
cd(doc, "df.describe()   # 统计摘要")
cd(doc, "df[df['speed']>50]   # 筛选车速>50的行")
cd(doc, "df['power_demand'].mean()  # 平均功率")
tbl(doc, ['函数', '功能'],
[['read_csv()', '读取CSV为DataFrame表格'],
['head()', '显示前5行预览数据'],
['describe()', '统计摘要(均值/标准差/四分位数)'],
['df[条件]', '按条件筛选行'],
['mean()/max()', '计算列的平均值/最大值']])
doc.add_paragraph()

# ===== 画图 =====
h1(doc, '第四段：matplotlib画图')

h2(doc, '图1：双子图')
cd(doc, "fig, axes = plt.subplots(2, 1, figsize=(12, 6))")
cd(doc, "axes[0].plot(df['time'], df['speed'], 'b-')  # 上图车速")
cd(doc, "axes[1].plot(df['time'], df['power_demand'], 'r-')  # 下图功率")
cd(doc, "plt.tight_layout(); plt.savefig('results/wltc_sample_plot.png', dpi=150)")
tbl(doc, ['代码', '说明'],
[["subplots(2,1)", "创建2行1列子图"],
["plot(x,y,'b-')", "画折线图,b=蓝色,-=实线"],
["tight_layout()", "自动调整子图间距"],
["savefig('...png')", "保存为PNG图片"]])
doc.add_paragraph()

h2(doc, '图2：双y轴图')
cd(doc, "ax1 = plt.subplot()    # 左y轴: 车速(蓝)")
cd(doc, "ax2 = ax1.twinx()     # 右y轴: 功率(红)")
bl(doc, 'ax1：左y轴画车速（蓝色）', bp='twinx()：')
bl(doc, 'ax2：右y轴画功率（红色），共享x轴')
tx(doc, '结果图保存到 results/wltc_dual_axis.png')

doc.add_paragraph()
r = doc.add_paragraph().add_run('='*50+'\nDay1 代码逐段解释\n生成日期：2026-06-03\n'+'='*50)
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

fname = 'F:/CLAUDE/research/ems-platform/docs/Day1_code_explain.docx'
doc.save(fname)
print('OK:', fname)
