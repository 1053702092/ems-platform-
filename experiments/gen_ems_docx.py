# -*- coding: utf-8 -*-
"""生成 EMS Controller 模型解释文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(t).font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

doc = Document()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('EMS 规则控制器逐段解释'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Day7 - EMS顶层模型搭建 | 文件: env/simulink_models/ems_controller_fcn.m'); r.font.size = Pt(11)
doc.add_paragraph()

# ===== 整体结构 =====
h1(doc, '整体结构')
tx(doc, 'ems_controller_fcn.m 是一个基于规则的 EMS（能量管理）控制器，决定燃料电池和电池之间的功率分配。')
tx(doc, '函数签名：')
cd(doc, 'function [P_fc_ref, P_bat_ref, status] = ems_controller_fcn(P_load, SOC)')
tbl(doc, ['参数', '含义', '单位'],
[['P_load（输入）', '功率需求（正=驱动）', 'kW'],
 ['SOC（输入）', '电池荷电状态', '0-1'],
 ['P_fc_ref（输出）', '燃料电池功率参考值', 'kW'],
 ['P_bat_ref（输出）', '电池功率参考值（正=放电，负=充电）', 'kW'],
 ['status（输出）', '当前工作模式（1-4）', '编号']])
doc.add_paragraph()

# ===== 控制目标 =====
h1(doc, '控制目标')
tx(doc, 'EMS 控制器需要满足以下约束：')
bl(doc, '功率平衡：P_load = P_fc + P_bat（任何时候）', bp='① ')
bl(doc, 'SOC 保护：SOC 不能低于 30% 或高于 90%', bp='② ')
bl(doc, 'FC 效率：燃料电池在 3-25kW 范围内工作（低于 3kW 效率差）', bp='③ ')
bl(doc, 'FC 动态保护：避免燃料电池频繁启停和功率剧烈波动', bp='④ ')
doc.add_paragraph()

# ===== 参数 =====
h1(doc, '第一段：参数定义')
cd(doc, 'P_fc_min = 3;      % 最低功率 [kW]')
cd(doc, 'P_fc_max = 25;     % 额定最大功率 [kW]')
cd(doc, 'P_fc_max_peak = 30;% 峰值功率 [kW]')
cd(doc, 'SOC_min = 0.30;    % 最低SOC')
cd(doc, 'SOC_low = 0.40;    % 偏低SOC（进入充电维持）')
cd(doc, 'SOC_high = 0.80;   % 偏高SOC（允许深度放电）')
cd(doc, 'SOC_max = 0.90;    % 最高SOC')
tbl(doc, ['参数', '值', '含义'],
[['P_fc_min', '3 kW', 'FC 最低功率。低于此效率很低，尽量不让 FC 低于这个值运行'],
 ['P_fc_max', '25 kW', 'FC 额定最大功率。持续运行的功率上限'],
 ['P_fc_max_peak', '30 kW', 'FC 峰值功率。短时（~10s）可达到，I-V 曲线最高点'],
 ['SOC_min', '0.30', '绝对最低 SOC。低于此有损坏电池风险'],
 ['SOC_low', '0.40', 'SOC 偏低阈值。低于此进入充电维持模式'],
 ['SOC_high', '0.80', 'SOC 偏高阈值。高于此尽量用电池'],
 ['SOC_max', '0.90', 'SOC 最高。高于此停止充电，保护电池']])
doc.add_paragraph()

# ===== 四种模式 =====
h1(doc, '四种工作模式')
tx(doc, '控制器根据 P_load 和 SOC 分为四种工作模式。')

tbl(doc, ['模式', 'status', '颜色', '触发条件'],
[['FC-only', '1', '红', 'SOC 正常 + P_load 在 FC 高效区'],
 ['Hybrid', '2', '绿', 'P_load > P_fc_max 或 SOC 偏高'],
 ['Charging', '3', '橙', 'P_load < P_fc_min 或 SOC 偏低'],
 ['Idle', '4', '灰', '停车（P_load < 1kW）且 SOC 已满']])
doc.add_paragraph()

# ===== 情形0 =====
h1(doc, '情形0：停车/制动（P_load < 1kW）')
cd(doc, 'if P_load < 1.0')
cd(doc, '    if SOC < SOC_max   % 没充满')
cd(doc, '        P_fc_ref = P_fc_min;')
cd(doc, '        P_bat_ref = P_load - P_fc_min;')
cd(doc, '        status = 3;     % 充电模式')
cd(doc, '    else                % 已充满')
cd(doc, '        P_fc_ref = 0;')
cd(doc, '        P_bat_ref = 0;')
cd(doc, '        status = 4;     % 怠速停机')
cd(doc, '    end')
cd(doc, '    return;')
tx(doc, '车速很低或停车时，功率需求接近零。此时：')
bl(doc, '如果电池没充满（SOC < 90%）：FC 以最低功率 3kW 运行，多出来的电能给电池充电。相当于停车怠速充电。')
bl(doc, '如果电池已充满（SOC ≥ 90%）：FC 完全关闭（0kW），电池也不出力。')
doc.add_paragraph()

# ===== 情形1 =====
h1(doc, '情形1：SOC 过低——充电维持模式')
cd(doc, 'if SOC < SOC_low')
cd(doc, '    P_fc_ref = min(P_load + (SOC_target(SOC) * 5), P_fc_max);')
cd(doc, '    P_fc_ref = max(P_fc_ref, P_fc_min);')
cd(doc, '    P_bat_ref = P_load - P_fc_ref;')
cd(doc, '    status = 3;')
tx(doc, 'SOC 低于 40% 时，强制进入充电维持模式：')
bl(doc, 'FC 出力 = P_load + 充电额外功率')
bl(doc, 'SOC_target(SOC) 函数根据 SOC 深度决定充电强度：SOC 越低，额外充电功率越大')
bl(doc, '充电额外功率 = SOC_target × 5kW（最高加到 5kW）')

tbl(doc, ['SOC', 'SOC_target 系数', '额外充电功率'],
[['< 20%', '1.0', '+5kW'],
 ['20-30%', '0.7', '+3.5kW'],
 ['30-40%', '0.4', '+2kW'],
 ['> 40%', '0.2', '+1kW（过渡平滑）']])
tx(doc, '最终 FC 出力还要用 max/min 限制在 [P_fc_min, P_fc_max] 范围内。')
doc.add_paragraph()

# ===== 情形2 =====
h1(doc, '情形2：SOC 过高——电池优先模式')
cd(doc, 'if SOC > SOC_high')
cd(doc, '    P_fc_ref = max(P_load - 10, P_fc_min);')
cd(doc, '    P_fc_ref = min(P_fc_ref, P_fc_max);')
cd(doc, '    P_bat_ref = P_load - P_fc_ref;')
cd(doc, '    status = 2;')
tx(doc, 'SOC 高于 80% 时，尽量消耗电池电量：')
bl(doc, '从负载功率中减去 10kW（这 10kW 由电池提供），剩下的由 FC 提供')
bl(doc, 'FC 最低仍保持 3kW（不能完全关断，防止 FC 频繁启停）')
bl(doc, '这样能有效降低 SOC，为后续回收制动能量腾出空间')
doc.add_paragraph()

# ===== 情形3 =====
h1(doc, '情形3：正常 SOC——跟随模式')
tx(doc, 'SOC 在 [40%, 80%] 范围内，控制器按功率需求分情况处理：')
doc.add_paragraph()

h2(doc, '子情形3a：低负载（P_load ≤ P_fc_min）')
cd(doc, 'if P_load <= P_fc_min')
cd(doc, '    P_fc_ref = P_fc_min;')
cd(doc, '    P_bat_ref = P_load - P_fc_min;')
cd(doc, '    status = 3;')
tx(doc, '负载需求低于 FC 最低高效功率（3kW）。')
tx(doc, '让 FC 以最低 3kW 运行（保持效率），多余的电给电池充电。')
tx(doc, '例如 P_load=1kW → P_fc=3kW, P_bat=-2kW（充电）')
doc.add_paragraph()

h2(doc, '子情形3b：中等负载（P_fc_min < P_load ≤ P_fc_max）')
cd(doc, 'elseif P_load <= P_fc_max')
cd(doc, '    P_fc_ref = P_load;')
cd(doc, '    P_bat_ref = 0;')
cd(doc, '    status = 1;')
tx(doc, '负载在 FC 高效区（3-25kW），FC 完全跟随负载。')
tx(doc, '电池不参与（P_bat_ref=0）。这是最理想的工作点——FC 效率高，电池零损耗。')
doc.add_paragraph()

h2(doc, '子情形3c：高负载（P_load > P_fc_max）')
cd(doc, 'else')
cd(doc, '    P_fc_ref = P_fc_max;')
cd(doc, '    P_bat_ref = P_load - P_fc_max;')
cd(doc, '    status = 2;')
tx(doc, '负载超过 FC 额定上限（25kW），FC 满负荷运行。')
tx(doc, '不足部分由电池放电补充。')
tx(doc, '例如 P_load=40kW → P_fc=25kW, P_bat=15kW（放电）')
doc.add_paragraph()

# ===== 辅助函数 =====
h1(doc, '辅助函数：SOC_target')
cd(doc, 'function factor = SOC_target(SOC)')
cd(doc, "    if SOC < 0.2;   factor = 1.0;")
cd(doc, "    elseif SOC < 0.3; factor = 0.7;")
cd(doc, "    elseif SOC < 0.4; factor = 0.4;")
cd(doc, "    else;            factor = 0.2;")
cd(doc, '    end')
tx(doc, '这个函数决定充电维持模式下的充电强度。SOC 越低，充电系数越大，FC 多出的功率越多。')
doc.add_paragraph()

# ===== 决策流程图 =====
h1(doc, '决策流程图')
cd(doc, '   输入: P_load, SOC')
cd(doc, '      │')
cd(doc, '   ┌──┴──┐')
cd(doc, '   │停车?│ ← P_load < 1kW')
cd(doc, '   └──┬──┘')
cd(doc, '   YES│                       NO│')
cd(doc, '   ┌──┴──┐                  ┌──┴──┐')
cd(doc, '   │SOC  │                  │SOC  │ ← SOC < SOC_low?')
cd(doc, '   │<max?│                  └──┬──┘')
cd(doc, '   └──┬──┘                  YES│              NO│')
cd(doc, '  YES│    NO│               ┌──┴──┐        ┌──┴──┐')
cd(doc, ' 充电  Idle              充电维持        │SOC  │ ← SOC > SOC_high?')
cd(doc, '  mode  mode                              └──┬──┘')
cd(doc, '                                          YES│        NO│')
cd(doc, '                                       ┌──┴──┐   ┌──┴──┐')
cd(doc, '                                     电池优先   │P_load│')
cd(doc, '                                                       └──┬──┘')
cd(doc, '                                          ┌────────┼────────┐')
cd(doc, '                                          │小      │中      │大')
cd(doc, '                                      充电 mode  FC跟随  混合 mode')
doc.add_paragraph()

doc.add_paragraph()
r = doc.add_paragraph().add_run('='*50+'\nEMS 规则控制器——逐段解释\n文件: env/simulink_models/ems_controller_fcn.m\n生成日期：2026-06-05\n'+'='*50)
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

fname = 'F:/CLAUDE/research/ems-platform/docs/Day7_EMS_controller_explain.docx'
doc.save(fname)
print('OK:', fname)
