# -*- coding: utf-8 -*-
"""生成 DP 动态规划原理文档 — 用于 EMS 能量管理策略"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h3(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(t).font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

def nl(doc):
    doc.add_paragraph()

# ====================================================================
# 创建文档
# ====================================================================
doc = Document()

# ---------- 封面 ----------
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('动态规划（DP）原理'); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('在燃料电池混合动力 EMS 中的应用'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

nl(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('EMS-PLAN · 第3周 · 原理篇'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('生成日期：2026-06-06'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
doc.add_page_break()

# ====================================================================
# 1. 核心思想
# ====================================================================
h1(doc, '一、核心思想')
tx(doc, '动态规划（Dynamic Programming, DP）是一种求解多阶段决策过程最优化的数学方法。在 EMS 中，DP 将能量管理问题建模为：')
tx(doc, '"给定一个已知的行驶工况（如 WLTC），找出一条最优的 SOC 轨迹和功率分配方案，使全程的氢燃料消耗最小。"')
nl(doc)

tx(doc, '与规则控制器（基于 if-else 启发式规则）不同，DP 通过后向递推，考虑到了整个工况的全局信息，因此能给出理论上的最优解。')
nl(doc)

h2(doc, '1.1 DP 的直觉理解')
tx(doc, '想象你开车走一条已知的路线（工况已知），途中有上坡、下坡、拥堵各种路况。你要决定什么时候用发动机（燃料电池）、')
tx(doc, '什么时候用电池，才能让全程最省油。规则控制器像一个"经验司机"——上坡就加大油门，下坡就收油。')
tx(doc, '而 DP 像一个"看过导航路线的上帝视角"——它知道前面 30 公里哪里有长上坡、哪里有下坡，')
tx(doc, '提前规划好电池电量的使用策略（比如：上坡前把电池充满，下坡时多回收能量），从而实现全局最优。')

# ====================================================================
# 2. 问题建模
# ====================================================================
h1(doc, '二、问题建模')

tx(doc, '一个标准的离散时间最优控制问题包含以下要素：')
nl(doc)

tbl(doc, ['要素', '符号', '在 EMS 中的定义', '说明'],
[['时间步', 'k = 0,1,...,N-1', '工况的每个采样点', 'WLTC 共 1800 步（1s 分辨率）'],
 ['状态变量', 'x_k', 'SOC（电池荷电状态）', '连续变量，离散化为网格点'],
 ['控制变量', 'u_k', 'P_fc（燃料电池功率）', '决策变量，离散化为网格点'],
 ['外部扰动', 'w_k', 'P_load（功率需求）', '由工况决定，全程已知'],
 ['状态转移', 'f(x,u,w)', 'SOC_{k+1} = SOC_k - (V_oc·I_bat)/(Q_bat·3600)·Δt', '电池模型'],
 ['单步代价', 'g(x,u)', 'ṁ_H₂(P_fc)·Δt + α·(SOC - SOC_ref)²', '氢耗 + SOC 维持惩罚'],
 ['总代价', 'J', 'Σ g_k(x_k, u_k)', '全程累计代价，DP 最小化此值']])
nl(doc)

h2(doc, '2.1 状态变量：SOC')
tx(doc, 'SOC 是系统的核心状态变量，决定了电池的可用能量。DP 将 SOC 离散化为 100~200 个网格点：')
tbl(doc, ['参数', '值', '说明'],
[['SOC 范围', '[0.2, 0.9]', '低于 0.2 或高于 0.9 会损害电池寿命'],
 ['离散化网格数 N_s', '100~200', '网格越密，精度越高，计算量越大'],
 ['网格步长', '~0.005', '约 0.5% 的 SOC 分辨率']])
nl(doc)

h2(doc, '2.2 控制变量：P_fc')
tx(doc, 'P_fc 是每个时间步的决策变量，决定了 FC 和电池之间的功率分配：')
tbl(doc, ['参数', '值', '说明'],
[['P_fc 范围', '[0, 30] kW', '0-关闭，30-峰值功率'],
 ['离散化网格数 N_u', '30~60', '网格越密，寻优越精细'],
 ['网格步长', '0.5~1 kW', '0.5kW 步长已足够精确']])
nl(doc)

h2(doc, '2.3 外部扰动：工况')
tx(doc, 'P_load[k] 由行驶工况通过车辆动力学模型计算得到，在整个 DP 过程中是已知的（确定性 DP）。')
tx(doc, '这使得 DP 能够"事前"看到全局信息——这也是 DP 与实时控制器的根本区别。')
nl(doc)

h2(doc, '2.4 状态转移方程（电池模型）')
tx(doc, '给定当前状态 SOC_k、控制 P_fc_k 和扰动 P_load[k]，下一步 SOC_{k+1} 由电池模型确定：')
nl(doc)
tx(doc, '① 功率平衡：      P_bat_k = P_load[k] - P_fc_k')
tx(doc, '② 电池电流：      I_bat = (V_oc - sqrt(V_oc² - 4·R_int·P_bat)) / (2·R_int)')
tx(doc, '③ SOC 更新：      SOC_{k+1} = SOC_k - I_bat / (Q_bat·3600) · Δt')
nl(doc)

# ====================================================================
# 3. Bellman 最优性原理
# ====================================================================
h1(doc, '三、Bellman 最优性原理')
tx(doc, 'Bellman 最优性原理是 DP 的核心理论基础：')
nl(doc)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.right_indent = Inches(0.5)
r = p.add_run('"一个最优策略具有这样的性质：无论初始状态和初始决策如何，')
r.font.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p.add_run('\n')
r = p.add_run('对于由初始决策所确定的后续状态，后续的决策也必须构成一个最优策略。"')
r.font.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
nl(doc)

tx(doc, '用数学公式表达为 Bellman 最优性方程：')
nl(doc)
tx(doc, '  J_k(x_k) = min [ g_k(x_k, u_k) + J_{k+1}(x_{k+1}) ]')
tx(doc, '              u_k')
nl(doc)
tx(doc, '其中：')
bl(doc, 'J_k(x_k)：从状态 x_k 到终点的最小累计代价（"价值函数"）', bp='• ')
bl(doc, 'g_k(x_k, u_k)：当前步的即时代价（氢耗 + SOC 惩罚）', bp='• ')
bl(doc, 'J_{k+1}(x_{k+1})：查表得到的未来最优代价（已在前一步算好）', bp='• ')
bl(doc, 'min 的操作是枚举所有可行的控制 u_k，取其中总代价最小的', bp='• ')
nl(doc)

# ====================================================================
# 4. 算法流程
# ====================================================================
h1(doc, '四、算法流程')
tx(doc, '完整的 DP 算法分为两个阶段：后向 DP（Backward Induction）和前向 Rollout（Forward Simulation）。')

nl(doc)
h2(doc, '4.1 阶段一：后向 DP（离线计算）')
tx(doc, '从最后一个时间步 N-1 开始，倒序计算到 k=0。这是 DP 最耗时的部分。')
nl(doc)

cd(doc, '算法 1：后向 DP')
cd(doc, '─' * 40)
cd(doc, '边界条件：J_N(x_N) = 0 或终端SOC惩罚')
cd(doc, '')
cd(doc, 'for k = N-1 downto 0:')
cd(doc, '    for each SOC_k in 网格:')
cd(doc, '        for each P_fc in 网格:')
cd(doc, '            P_bat = P_load[k] - P_fc')
cd(doc, '            SOC_{k+1} = battery_model(SOC_k, P_bat)')
cd(doc, '            if SOC_{k+1} 超出边界:')
cd(doc, '                跳过（不可行控制）')
cd(doc, '            cost = H2_consumption(P_fc) * dt')
cd(doc, '                 + alpha * (SOC_k - SOC_ref)^2')
cd(doc, '            total = cost + J_{k+1}(SOC_{k+1})')
cd(doc, '            if total < min_cost:')
cd(doc, '                min_cost = total')
cd(doc, '                best_u = P_fc')
cd(doc, '        J_k(SOC_k) = min_cost')
cd(doc, '        π_k(SOC_k) = best_u')
cd(doc, '    end')
cd(doc, 'end')

nl(doc)
tx(doc, '时间复杂度：O(N × N_s × N_u)')
tx(doc, '以 WLTC 为例：N=1800，N_s=150，N_u=60 → 约 1620 万次枚举。Python 实现约需数秒。')
nl(doc)

h2(doc, '4.2 阶段二：前向 Rollout（在线仿真）')
tx(doc, '用后向 DP 得到的策略表 π_k(SOC_k) 从初始状态开始正向仿真：')
nl(doc)

cd(doc, '算法 2：前向 Rollout')
cd(doc, '─' * 40)
cd(doc, 'SOC_0 = 0.6   # 初始SOC')
cd(doc, '')
cd(doc, 'for k = 0 to N-1:')
cd(doc, '    P_fc_k = π_k(SOC_k)          # 查最优策略表')
cd(doc, '    P_bat_k = P_load[k] - P_fc_k')
cd(doc, '    SOC_{k+1} = battery_model(...)')
cd(doc, '    m_H2_k = H2_consumption(P_fc_k) * dt')
cd(doc, '    记录: SOC_k, P_fc_k, P_bat_k, m_H2_k')
cd(doc, 'end')

nl(doc)
tx(doc, '前向 rollout 只是查表，不需要枚举寻优，因此非常快（毫秒级）。')

# ====================================================================
# 5. 氢耗模型
# ====================================================================
h1(doc, '五、氢耗模型（代价函数核心）')
tx(doc, 'FC 的效率不是常数——不同功率点效率不同，DP 的价值就在于自动找到让 FC 工作在高效区的策略。')
nl(doc)

tx(doc, '氢气流速（代价）计算公式：')
nl(doc)
tx(doc, '  ṁ_H₂(P_fc) = P_fc / (η_fc(P_fc) × LHV_H₂)')
nl(doc)
tx(doc, '其中：')
bl(doc, 'LHV_H₂ = 120 MJ/kg — 氢的低热值', bp='• ')
bl(doc, 'η_fc(P_fc) — FC 效率曲线（上凸函数，中等负载效率最高）', bp='• ')

nl(doc)
tx(doc, '典型 FC 效率曲线：')
tbl(doc, ['P_fc (kW)', '效率 η_fc', '氢耗 (g/s)', '说明'],
[['0', '0%', '0.000', '关闭（无氢耗）'],
 ['3', '35%', '0.071', '最低功率点（效率低端）'],
 ['8', '48%', '0.139', '中等负载（效率较高）'],
 ['15', '55%', '0.227', '高效区（最优工作点）'],
 ['20', '53%', '0.315', '高效区边缘'],
 ['25', '48%', '0.434', '额定功率上限'],
 ['30', '40%', '0.625', '峰值功率（效率低）']])
nl(doc)

tx(doc, '注意：FC 在 15kW 附近效率最高（约 55%），过低或过高都会降低效率。DP 会自动')
tx(doc, '倾向于让 FC 工作在高效区，多余功率给电池充电，不足时由电池补充。这正是')
tx(doc, 'DP 比规则控制器节能的根本原因。')

# ====================================================================
# 6. SOC 维持惩罚
# ====================================================================
h1(doc, '六、SOC 维持惩罚')
tx(doc, '单纯最小化氢耗会导致 SOC 持续下降（用电池的电更"便宜"），因此需要加入 SOC 维持惩罚项：')
nl(doc)

tx(doc, '  g_k(x_k, u_k) = ṁ_H₂(P_fc_k)·Δt + α·(SOC_k - SOC_ref)²')
nl(doc)

tx(doc, '惩罚项的作用：')
bl(doc, 'α（惩罚系数）：调节 SOC 维持的强度。α 越大，SOC 越倾向于维持在 SOC_ref 附近', bp='① ')
bl(doc, 'SOC_ref（参考 SOC）：通常设为 0.6（中间值），允许一定程度波动', bp='② ')
bl(doc, '终端惩罚：最后一步加一个大的终端惩罚，强制 SOC_N 接近 SOC_0（电荷维持）', bp='③ ')
nl(doc)

tx(doc, '惩罚系数的选择需要权衡：')
tbl(doc, ['α 的取值', 'SOC 维持效果', '氢耗经济性'],
[['α 过小', 'SOC 漂移大，可能耗尽电池', '氢耗最低（过度依赖电池）'],
 ['α 适中', 'SOC 小幅波动 ±0.05', '氢耗接近最优'],
 ['α 过大', 'SOC 近乎恒定', '氢耗偏高（灵活性差）']])

# ====================================================================
# 7. 数值实现要点
# ====================================================================
h1(doc, '七、数值实现要点')

h2(doc, '7.1 插值处理')
tx(doc, '后向 DP 中，J_{k+1}(SOC_{k+1}) 查表时，SOC_{k+1} 不一定正好落在网格点上，')
tx(doc, '需要进行线性插值。同样，前向 rollout 中 π_k(SOC_k) 也需要插值。')
tx(doc, '使用 numpy.interp 即可高效实现。')
nl(doc)

h2(doc, '7.2 不可行控制')
tx(doc, '某些 (SOC_k, P_fc) 组合会导致 SOC_{k+1} 超出 [0.2, 0.9] 边界（违反电池保护约束）。')
tx(doc, '这些控制必须被标记为不可行，在 DP 枚举中跳过。')
nl(doc)

h2(doc, '7.3 SOC 终值约束（电荷维持）')
tx(doc, '实际应用中，我们要求 SOC 终值 ≈ SOC 初值（否则不公平——消耗电池的电不算"真省"）。')
tx(doc, '可以通过两种方式实现：')
bl(doc, '硬约束：强制 SOC_N ∈ [SOC_0-0.01, SOC_0+0.01]', bp='方法A ')
bl(doc, '软约束：J_N(SOC_N) = β·(SOC_N - SOC_0)² 作为边界条件', bp='方法B ')
nl(doc)
tx(doc, '本书采用方法B（软约束），因为硬约束可能引起数值问题（无可行解）。')

# ====================================================================
# 8. 规则 vs DP
# ====================================================================
h1(doc, '八、规则控制器 vs DP 对比')

tbl(doc, ['维度', '规则控制器', 'DP 最优控制器'],
[['原理', 'if-else 启发式规则', 'Bellman 方程全局寻优'],
 ['前瞻性', '无（只看当前步）', '完整工况已知'],
 ['最优性', '不一定最优', '全局最优（离散精度内）'],
 ['计算量', '实时（微秒级）', '离线（数秒），在线查表快'],
 ['实时性', '✅ 可实时运行', '❌ 需完整工况，不能在线'],
 ['鲁棒性', '对工况变化不敏感', '对工况变化需重算'],
 ['用途', '实际部署', '作为 benchmark 评价其他策略'],
 ['复杂度', '简单，几十行代码', '中等，需网格+插值处理']])
nl(doc)

tx(doc, 'DP 的核心价值：它给出了理论的"最优解天花板"。无论规则控制器、ECMS、MPC 还是')
tx(doc, '强化学习，它们的节能空间有多大，都以 DP 结果为基准来衡量。')
tx(doc, '如果一个策略比 DP 还"省"，那通常意味着它牺牲了 SOC 维持（在消耗电池的储备能量）。')

# ====================================================================
# 9. 预期结果
# ====================================================================
h1(doc, '九、预期结果')
tx(doc, '在 WLTC 工况上，DP 相对于规则控制器的预期改善：')
nl(doc)

tbl(doc, ['指标', '规则控制器（参考值）', 'DP（预期）', '改善幅度'],
[['总氢耗', '~0.35 kg', '~0.30 kg', '↓ 10-15%'],
 ['SOC 终值 vs 初值', '0.60 → 0.61', '0.60 → 0.60', '更精确维持'],
 ['FC 平均效率', '~45%', '~50%', '↑ 5 个百分点'],
 ['FC 功率波动', '较频繁', '更平滑', '有利 FC 寿命']])
nl(doc)

tx(doc, '注：以上为经验估计。实际结果受 FC 效率曲线参数、电池参数和惩罚系数 α 影响。')
nl(doc)

# ====================================================================
# 页脚
# ====================================================================
nl(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('─' * 50); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('动态规划（DP）原理 — EMS-PLAN 第3周\n'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
r = p.add_run('生成日期：2026-06-06 | F:/CLAUDE/research/ems-platform\n'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
r = p.add_run('参考文献：Bellman R. Dynamic Programming (1957), \nGuzzella L. Vehicle Propulsion Systems (2013)'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

# ====================================================================
# 保存
# ====================================================================
fname = 'F:/CLAUDE/research/ems-platform/docs/DP_dynamic_programming_principle.docx'
doc.save(fname)
print('OK:', fname)
