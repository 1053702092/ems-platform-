# -*- coding: utf-8 -*-
"""生成 build_ems_model.m 解释文档（更新版）"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(t).font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

doc = Document()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('build_ems_model.m 逐段解释（更新版）'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Day7 - EMS顶层模型搭建 | 文件: env/simulink_models/Use-Model/build_ems_model.m'); r.font.size = Pt(11)
doc.add_paragraph()

# ===== 作用 =====
h1(doc, '文件作用')
tx(doc, '用 MATLAB Simulink 程序化 API 自动搭建 EMS_hybrid_v1 模型。运行脚本直接生成完整的 Simulink 模型，无需手动拖拽。')
doc.add_paragraph()

# ===== 用法 =====
h1(doc, '三种调用方式')
tbl(doc, ['命令', '作用'],
[["build_ems_model", "创建新模型（已存在则直接打开）"],
 ["build_ems_model('open')", "只打开已存在的模型，不重建"],
 ["build_ems_model('rebuild')", "删除旧的，重新搭建"]])
doc.add_paragraph()

# ===== 模型架构 =====
h1(doc, '模型整体架构')
tx(doc, 'Simulink 模型从信号流角度分为 5 个主要部分：')
doc.add_paragraph()

tbl(doc, ['编号', '模块', '类型', '功能'],
[['①', 'WLTC Data', 'From Workspace', '加载 WLTC 工况 [t, v] 矩阵，输出速度列'],
 ['②', 'Vehicle Power', 'MATLAB Function', '车速+加速度 → 功率需求 P_load'],
 ['③', 'EMS Controller', 'MATLAB Function', '规则基能量管理，输出 P_fc_ref 和 P_bat_ref'],
 ['④', 'FC System', 'Subsystem', '燃料电池 I-V 查表 + DC/DC 效率'],
 ['⑤', 'Battery', 'MATLAB Function', 'R-int 电池模型，SOC/V_bat/I_bat']])
doc.add_paragraph()

tx(doc, '信号流：')
cd(doc, 'WLTC [t,v] → Demux → 速度 v')
cd(doc, '    ├→ Vehicle Power/1 (速度)')
cd(doc, '    └→ Derivative → Vehicle Power/2 (加速度)')
cd(doc, 'Vehicle Power → P_load → EMS Controller/1')
cd(doc, 'Constant 0.6 → SOC_init → EMS Controller/2')
cd(doc, 'EMS Controller → [P_fc_ref → FC System,  P_bat_ref → Battery]')
cd(doc, 'SOC 反馈: Battery/1 → Memory → Battery/2 (下一时刻 SOC_init)')
cd(doc, '各模块输出 → To Workspace (日志)')
doc.add_paragraph()

# ===== 开发者踩坑记录 =====
h1(doc, '开发者踩坑记录（重要）')
tx(doc, '以下是搭建过程中遇到的所有 API 兼容性问题和修复方案。')

h2(doc, '坑1：SubSystem 默认自带 In1 / Out1')
tx(doc, 'add_block 创建 SubSystem 时，Simulink 自动添加了默认的 In1 和 Out1 端口。如果不删除，后面加的 Inport/Outport 端口号会偏移一位，导致接线全错。')
cd(doc, '% 创建后立即删除默认端口')
cd(doc, 'sub = add_block("simulink/Ports & Subsystems/Subsystem", path);')
cd(doc, 'delete_block([sub "/In1"]);')
cd(doc, 'delete_block([sub "/Out1"]);')
tx(doc, '症状：FC 功率读到 388V（实际是电压值），SOC 不变化。')

h2(doc, '坑2：Lookup Table API 版本不兼容')
tx(doc, '不同 MATLAB 版本的 1-D Lookup Table 参数名不同。R2024b 中 Breakpoints1 不可用。')
tx(doc, '解决方案：用 MATLAB Function 块 + interp1 代替，与 Simulink 版本无关。')

h2(doc, '坑3：From Workspace 不需要 Demux')
tx(doc, 'From Workspace 块使用 [t, u] 格式，第一列是时间（内部插值用），输出只有第二列的数据。所以从 From Workspace 出来的信号已经是单列速度，不需要 Demux 拆分。')
tx(doc, '症状：Demux 报错"端口宽度超出有效维度"。')

h2(doc, '坑4：set_param(block, Script, code) 不稳定')
tx(doc, 'MATLAB Function 块的 Script 属性在某些版本中不能直接用 set_param 设置。')
tx(doc, '解决方案：三级降级策略。')

tbl(doc, ['优先级', '方法', '适用版本'],
[['1(优先)', 'Simulink.MATLABFunctionBlock.setMATLABFunctionCode', 'R2021a+'],
 ['2(回退)', 'set_param(block, "Script", code)', 'R2021a+ 部分版本'],
 ['3(保底)', 'sfroot().find(...) → chart.Script = code', '所有版本']])

h2(doc, '坑5：ModelWorkspace.save API 变更')
tx(doc, 'R2024b 中 ModelWorkspace.save 和 saveToSource 都报错。')
tx(doc, '解决方案：只保存到 base workspace（assignin），不保存到模型工作区。需要持久化的数据用单独 CSV 文件。')

h2(doc, '坑6：SOC 必须反馈不能是常量')
tx(doc, 'Battery 的 SOC_init 输入如果用 Constant（常量 0.6），SOC 永远不会更新——每个时间步都从 0.6 开始算。必须把 Battery 输出的 SOC 通过 Memory 块反馈到输入。')
cd(doc, 'Battery/1 (SOC_out) → Memory → Battery/2 (SOC_init)')
doc.add_paragraph()

# ===== 各模块详解 =====
h1(doc, '主函数各段详解')

h2(doc, '1. switch 结构 (action 参数)')
cd(doc, 'function build_ems_model(action)')
cd(doc, 'if nargin < 1, action = "build"; end')
tx(doc, '三种 action：open（只打开）、rebuild（删除重建）、build（默认，存在就跳过）。注意前两种用 return 提前退出，build 完成后也继续执行后面的新建代码。')

h2(doc, '2. 路径计算')
cd(doc, 'SCRIPT_DIR = fileparts(mfilename("fullpath"));')
cd(doc, 'MODEL_DIR = fileparts(SCRIPT_DIR);')
cd(doc, 'PROJECT_ROOT = fileparts(fileparts(MODEL_DIR));')
tx(doc, 'mfilename("fullpath") 返回当前 .m 文件的完整路径，fileparts 逐级上跳。3 层从 Use-Model/ 回到项目根目录。')

h2(doc, '3. 求解器设置')
cd(doc, "set_param(mdl, 'Solver', 'ode45', 'StopTime', '1800', 'MaxStep', '0.1');")
tx(doc, 'ode45：变步长龙格-库塔求解器。StopTime=1800s（WLTC 全长）。MaxStep=0.1s 保证加速度计算精度。')

h2(doc, '4. WLTC 数据源 (From Workspace)')
tx(doc, '加载 WLTC [t, v] 矩阵，VariableName=sim_wltc。数据在仿真前由 assign_wltc_data 写入 base workspace。')

h2(doc, '5. 三个 MATLAB Function 块')
tx(doc, 'Vehicle Power、EMS Controller、Battery 都是 MATLAB Function 块。创建后用 set_matlab_func_code 注入 .m 文件代码。端口由函数签名自动推导。')

h2(doc, '6. FC System 子系统')
tx(doc, '内部结构：')
bl(doc, 'P_fc_ref (输入) → Divide (I = P/V) → Saturation [0,300]A')
bl(doc, '→ I_to_V (MATLAB Function, interp1 查 I-V 表) → V_stack')
bl(doc, '→ Product (V × I) → Gain (×0.95 DC/DC 效率) → P_fc_actual (输出)')
bl(doc, '→ Unit Delay → V_feedback (打破代数环, IC=350V)')
bl(doc, '→ V_fc (电压输出)')

h2(doc, '7. SOC 反馈（关键修正）')
tx(doc, 'Battery 输出的 SOC 经过 Memory 块延迟一步后反馈回 Battery 的 SOC_init 输入。Memory 的 InitialCondition=0.6 让 SOC 从 60% 开始。没有这个反馈，SOC 永远停在初始值。')

h2(doc, '8. 数据记录（7 个 To Workspace）')
tbl(doc, ['变量名', '记录内容', '来源'],
[['sim_V_fc', 'FC 端电压', 'FC System/1'],
 ['sim_P_fc', 'FC 实际功率', 'FC System/2'],
 ['sim_SOC', '电池 SOC', 'Battery/1'],
 ['sim_V_bat', '电池端电压', 'Battery/2'],
 ['sim_I_bat', '电池电流', 'Battery/3'],
 ['sim_P_load', '负载功率需求', 'Vehicle Power'],
 ['sim_status', 'EMS 工作模式', 'EMS Controller/3']])
tx(doc, '全部排到模型右侧（x=960），纵向间距 45px，避免重叠。')

h2(doc, '9. 自动整理布局')
cd(doc, 'Simulink.BlockDiagram.arrangeSystem(mdl);')
cd(doc, 'lines = find_system(mdl, "FindAll", "on", "Type", "line");')
cd(doc, 'for i = 1:length(lines)')
cd(doc, '    Simulink.BlockDiagram.routeLine(lines(i));')
cd(doc, 'end')
tx(doc, 'build 结束后自动排模块位置、整理信号线走线。相当于在 Simulink 界面右键→Arrange System。')

doc.add_paragraph()

# ===== 使用流程 =====
h1(doc, '完整使用流程')
tx(doc, '1. 在 MATLAB 命令窗口执行：')
cd(doc, '>> cd("F:/CLAUDE/research/ems-platform/env/simulink_models/Use-Model")')
cd(doc, '>> build_ems_model')
tx(doc, '2. 运行仿真：')
cd(doc, '>> sim("EMS_hybrid_v1")')
tx(doc, '3. 查看结果：')
cd(doc, '>> plot(sim_P_load); hold on; plot(sim_P_fc); legend("Load","FC")')
doc.add_paragraph()

doc.add_paragraph()
r = doc.add_paragraph().add_run('='*50+'\nbuild_ems_model.m 逐段解释（更新版）\n文件: env/simulink_models/Use-Model/build_ems_model.m\n生成日期：2026-06-05\n'+'='*50)
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

fname = 'F:/CLAUDE/research/ems-platform/docs/Day7_build_EMS_model_explain_v2.docx'
doc.save(fname)
print('OK:', fname)
