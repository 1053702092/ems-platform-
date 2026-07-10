#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Week9 合并文档 (PyTorch 篇 + RL 篇)
替代原有的 5 个碎片化 docx 文件
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = 'F:/CLAUDE/research/ems-platform/docs'

# ── 工具函数 ──

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def make_table(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h3(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, text, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        r = p.add_run(text); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

def code_block(doc, lines, label=None):
    if label:
        p = doc.add_paragraph(); r = p.add_run(label); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    for line in lines:
        cd(doc, line)

def note(doc, t):
    p = doc.add_paragraph(); r = p.add_run('* '); r.font.size = Pt(10); r.italic = True
    r = p.add_run(t); r.font.size = Pt(10); r.italic = True; r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('─' * 50); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ═══════════════════════════════════════════════════════════════
# 文档 1: Week9_PyTorch入门_完整笔记.docx
# ═══════════════════════════════════════════════════════════════
def build_pytorch_doc():
    doc = Document()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Week 9 — PyTorch 入门\n完整学习笔记'); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('文件: scripts/week9_complete.py (Part 1-4)'); r.font.size = Pt(10)
    doc.add_paragraph()

    # ── Part 1 ──
    h1(doc, 'Part 1: Tensor 基础')
    tx(doc, 'Tensor = 可跑在 GPU 上的 numpy ndarray，是 PyTorch 的核心数据结构。')
    h2(doc, '创建 Tensor 的 7 种方法')
    make_table(doc, ['方法', '示例', '用途'], [
        ['torch.tensor()', 'tensor([[1,2],[3,4]])', '从列表创建'],
        ['torch.from_numpy()', 'from_numpy(np_array)', '从 numpy 共享内存创建'],
        ['torch.zeros()', 'zeros(3, 4)', '全 0 张量'],
        ['torch.ones()', 'ones(2, 3)', '全 1 张量'],
        ['torch.randn()', 'randn(3, 3)', '标准正态随机'],
        ['torch.rand()', 'rand(2, 2)', '[0,1) 均匀随机'],
        ['torch.arange()', 'arange(0, 10, 2)', '等差序列'],
    ])
    doc.add_paragraph()

    h2(doc, '形状操作')
    make_table(doc, ['操作', '说明', '是否共享内存'], [
        ['view(shape)', '重塑（要求内存连续）', '✓ 共享'],
        ['reshape(shape)', '重塑（自动处理不连续）', '可能复制'],
        ['transpose(d0, d1)', '交换两个维度', '✓ 共享'],
        ['squeeze()', '删除长度为1的维度', '✓ 共享'],
        ['unsqueeze(dim)', '增加维度', '✓ 共享'],
    ])
    doc.add_paragraph()

    h2(doc, '广播机制 (Broadcasting)')
    tx(doc, '从右向左对齐，要么维度相等、要么维度为1，否则报错。')
    tx(doc, '例: (3,1) + (3,) → (3,3) — 每个位置自动扩展。')

    h2(doc, '数据类型')
    code_block(doc, [
        'x = torch.tensor([1,2,3], dtype=torch.float32)  # 默认 float32',
        'x_f64 = x.double()     # → float64',
        'x_i64 = x.long()       # → int64',
        'x_half = x.half()      # → float16（推理加速）',
    ])
    doc.add_paragraph()

    # ── Part 2 ──
    h1(doc, 'Part 2: Autograd 自动求导')
    tx(doc, 'PyTorch 用动态计算图记录所有运算。backward() 沿图反向传播，自动计算梯度。')

    h2(doc, '核心概念')
    make_table(doc, ['概念', '说明'], [
        ['requires_grad=True', '标记需要求导的参数'],
        ['backward()', '反向传播，从 loss 计算每个参数的梯度'],
        ['grad', '存储梯度值 ∂loss/∂x'],
        ['no_grad()', '推理时关闭梯度追踪（省显存）'],
        ['detach()', '切断单个 tensor 的梯度传播'],
    ])
    doc.add_paragraph()

    h2(doc, '关键要点')
    bl(doc, 'backward 要求输出是标量，否则需传 gradient 参数')
    bl(doc, '梯度默认累积（累加），backward 前必须 zero_grad()')
    bl(doc, '就地操作（带 _ 后缀）会破坏计算图，慎用')
    doc.add_paragraph()

    # ── Part 3 ──
    h1(doc, 'Part 3: nn.Module + MLP')
    tx(doc, '所有 PyTorch 模型继承 nn.Module，需实现 __init__ 和 forward。')
    code_block(doc, [
        'class MLP(nn.Module):',
        '    def __init__(self, input_dim, hidden_dim, output_dim):',
        '        super().__init__()',
        '        self.fc1 = nn.Linear(input_dim, hidden_dim)',
        '        self.fc2 = nn.Linear(hidden_dim, output_dim)',
        '    def forward(self, x):',
        '        x = F.relu(self.fc1(x))',
        '        return self.fc2(x)',
    ], label='网络定义模板')
    doc.add_paragraph()

    make_table(doc, ['层', '功能'], [
        ['nn.Linear(in, out)', '全连接层 y = xW^T + b'],
        ['nn.Dropout(p)', '随机丢弃（防过拟合）'],
        ['nn.Sequential', '快速堆叠线性网络'],
        ['nn.LSTM(i, h)', 'LSTM 层'],
    ])
    doc.add_paragraph()

    # ── Part 4 ──
    h1(doc, 'Part 4: MLP 功率预测（FC 功率预测）')
    tx(doc, '任务：根据过去 10 步的 FC 功率序列，预测下一步功率值。')
    make_table(doc, ['参数', '值'], [
        ['窗口 (seq_len)', '10 步'],
        ['模型', '3 层 MLP (10→64→64→1) + Dropout(0.1)'],
        ['优化器', 'Adam(lr=0.001)'],
        ['损失', 'MSELoss'],
        ['训练', '200 epoch'],
    ])
    doc.add_paragraph()
    bl(doc, '结果: MAE≈1.85 kW, RMSE≈2.27 kW')
    bl(doc, '输出: results/week9_complete_mlp_power_prediction.png')
    note(doc, 'MLP 不建模时间依赖性，LSTM 效果更好 — 见第12章。')
    doc.add_paragraph()

    # ── 结尾 ──
    sep(doc)
    r = doc.add_paragraph().add_run('Week 9 PyTorch 入门 · 完整笔记\n2026-07-10 · EMS 研究项目')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

    path = os.path.join(OUT_DIR, 'Week9_PyTorch入门_完整笔记.docx')
    doc.save(path)
    print(f'[OK] {path}')


# ═══════════════════════════════════════════════════════════════
# 文档 2: Week9_RL基础_完整笔记.docx
# ═══════════════════════════════════════════════════════════════
def build_rl_doc():
    doc = Document()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Week 9 — 强化学习基础\n完整学习笔记'); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('文件: scripts/week9_complete.py (Part 5-8)'); r.font.size = Pt(10)
    doc.add_paragraph()

    # ── Part 5 ──
    h1(doc, 'Part 5: MDP 五元组 — GridWorld')
    tx(doc, '马尔可夫决策过程 (MDP) 是强化学习的数学框架，由五元组定义:')

    make_table(doc, ['元素', '符号', '说明', 'GridWorld 示例'], [
        ['状态', 'S', '智能体的所有可能状态', '4×4=16 个格子'],
        ['动作', 'A', '智能体可执行的动作', '↑↓←→ 4 种'],
        ['转移概率', 'P(s\'|s,a)', '执行动作后到下一状态的概率', '80% 目标方向, 20% 随机'],
        ['奖励', 'R(s,a)', '执行动作的立即回报', '终点+1, 陷阱-1, 其他0'],
        ['折扣因子', 'γ', '未来奖励的折扣', 'γ=0.9'],
    ])
    doc.add_paragraph()
    code_block(doc, [
        '起点 (0,0) → 目标是走到终点 (3,3) +1',
        '陷阱在 (1,1), 踩到得 -1 并停在原地',
        '动作有随机性: 80% 去目标方向, 20% 随机去其他方向',
    ])
    doc.add_paragraph()

    # ── Part 6 ──
    h1(doc, 'Part 6: Bellman 方程')
    h2(doc, 'Bellman 期望方程 (策略评估)')
    tx(doc, 'V^π(s) = Σ_a π(a|s)[R(s,a) + γ Σ_{s\'} P(s\'|s,a) V^π(s\')]')
    tx(doc, '—— 在固定策略 π 下，每个状态的价值 = 当前奖励 + 折扣后的未来价值。')

    h2(doc, 'Bellman 最优方程')
    tx(doc, 'V*(s) = max_a [R(s,a) + γ Σ_{s\'} P(s\'|s,a) V*(s\')]')
    tx(doc, '—— 每个状态下选最好的动作，递归定义最优值函数。')

    make_table(doc, ['结果', '随机策略 V(0)', '最优 V*(0)'], [
        ['GridWorld', '-3.93', '3.34'],
    ])
    doc.add_paragraph()

    h2(doc, '关键理解')
    bl(doc, 'V(s): 状态价值 = 在这个状态能获得的总回报的期望')
    bl(doc, 'Q(s,a): 动作价值 = 在这个状态执行这个动作能获得的总回报')
    bl(doc, 'V(s) = max_a Q(s,a) (最优时)')
    doc.add_paragraph()

    # ── Part 7 ──
    h1(doc, 'Part 7: 策略迭代')
    tx(doc, '交替执行：策略评估 → 策略改进 → 直到策略不再变化。')
    code_block(doc, [
        '初始化随机策略 π',
        '循环:',
        '  策略评估: 解 Bellman 期望方程 V^π (迭代法)',
        '  策略改进: 对每个状态 s, π(s)=argmax_a Q(s,a)',
        '  如果 π 不再变化 → 收敛到最优策略',
    ], label='策略迭代流程')
    doc.add_paragraph()

    make_table(doc, ['轮次', '动作'], [
        ['第1轮', '策略改进中... (随机→初步最优)'],
        ['第4轮', '收敛 [OK]'],
    ])
    tx(doc, '结果: 起点策略 ↑ (向上走), 最优策略避开了陷阱 (1,1)。')
    doc.add_paragraph()

    # ── Part 8 ──
    h1(doc, 'Part 8: 值迭代')
    tx(doc, '直接迭代 Bellman 最优方程，不维护显式策略，收敛后从 V* 提取策略。')
    code_block(doc, [
        '初始化 V(s)=0',
        '循环:',
        '  V(s) = max_a [R(s,a) + γ Σ P(s\'|s,a) V(s\')]',
        '直到 V 的变化 < θ',
        '提取策略: π(s)=argmax_a Q(s,a)',
    ], label='值迭代流程')
    doc.add_paragraph()

    # ── 策略迭代 vs 值迭代 ──
    h1(doc, '附录: 策略迭代 vs 值迭代')
    make_table(doc, ['对比维度', '策略迭代', '值迭代'], [
        ['核心思路', '评估→改进交替', '直接优化 V*'],
        ['收敛速度', '轮数少（但每轮计算重）', '轮数多（但每轮计算轻）'],
        ['GridWorld', '4 轮收敛', '133 轮收敛'],
        ['适用场景', '状态空间小', '状态空间大'],
        ['最终结果', '一致（V=3.3419）', '一致（V=3.3419）'],
    ])
    doc.add_paragraph()
    tx(doc, '验证: 两种方法收敛到同一最优值函数 V(0)=3.3419, ΔV=0.000006 ✅')
    doc.add_paragraph()

    # ── 面试准备 ──
    h1(doc, '面试 Q&A 快速参考')
    h3(doc, 'Q: "MDP 五元组是什么？"')
    tx(doc, '答: <S, A, P, R, γ> — 状态空间、动作空间、转移概率、奖励函数、折扣因子。')

    h3(doc, 'Q: "策略迭代和值迭代的区别？"')
    tx(doc, '答: 策略迭代是先评估当前策略再改进，收敛轮数少但每轮计算重。值迭代直接迭代最优值函数，轮数多但每轮计算轻。两者最终收敛到同一最优值。')

    h3(doc, 'Q: "Bellman 最优性原理是什么？"')
    tx(doc, '答: 一个最优策略的子策略也是最优的。DP 和 RL 都是基于这个原理递推求解。')

    doc.add_paragraph()
    sep(doc)
    r = doc.add_paragraph().add_run('Week 9 RL 基础 · 完整笔记\n2026-07-10 · EMS 研究项目')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

    path = os.path.join(OUT_DIR, 'Week9_RL基础_完整笔记.docx')
    doc.save(path)
    print(f'[OK] {path}')


if __name__ == '__main__':
    os.makedirs(OUT_DIR, exist_ok=True)
    build_pytorch_doc()
    build_rl_doc()
    print('\n完成！2 份合并文档已生成。')
