# -*- coding: utf-8 -*-
"""生成 DP 原理学习笔记 — 整合网络资源精华版本"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h3(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(t).font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def nl(doc):
    doc.add_paragraph()

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

# ====================================================================
doc = Document()

# ---------- 封面 ----------
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n')
r = p.add_run('动态规划（DP）入门笔记'); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('从零理解燃料电池混合动力 EMS 中的 DP'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
nl(doc); nl(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('基于网络精选资源整理'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('阿里云开发者社区 · Kaputt Engineers Blog · GitHub HEV_EMS_DP'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('生成日期：2026-06-06'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
doc.add_page_break()

# ====================================================================
# 第〇章：太长不看版
# ====================================================================
h1(doc, '第〇章 · 一句话说清 DP')
tx(doc, 'DP 就干一件事：')
nl(doc)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
r = p.add_run('「已知整个工况的速度曲线，找出一条最省油的功率分配方案。」')
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
nl(doc)

tx(doc, '怎么做呢？两步走：')
nl(doc)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run('Step 1 — 倒着算（后向 DP）：'); r.bold = True; r.font.size = Pt(11)
p.add_run('从终点往起点，算出每个状态下"后面怎么走最省油"的对照表。')
r.font.size = Pt(10)
nl(doc)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run('Step 2 — 正着跑（前向 Rollout）：'); r.bold = True; r.font.size = Pt(11)
p.add_run('从起点开始，每步查那张对照表，按最优方案开车。')
r.font.size = Pt(10)
nl(doc)

tx(doc, '结果：DP 给出的答案是"理论最优"（在离散精度内），其他策略（规则/ECMS/RL）都以它为 benchmark 来衡量差距。')
nl(doc)

# ====================================================================
# 第一章：为什么需要 DP？
# ====================================================================
h1(doc, '第一章 · 为什么需要 DP？')

h2(doc, '1.1 EMS 到底在优化什么？')
tx(doc, '能量管理策略（EMS）要解决一个核心矛盾：')
nl(doc)
tx(doc, '你有两个能源：燃料电池（FC）和电池。')
tx(doc, 'FC 效率随功率变化，15kW 时效率最高（~55%），低功率和高功率都费油。')
tx(doc, '电池像"水库"，可以存电/放电，但充放都有损耗。')
nl(doc)

tx(doc, 'EMS 的任务：在每个时刻，决定 P_fc 和 P_bat 各出多少，使得：')
bl(doc, '全程氢耗最小化')
bl(doc, 'SOC 最终不低于起始值（不能"作弊"消耗电池）')
bl(doc, '各部件不超限（FC 不超 30kW，SOC 在 [0.2, 0.9] 内）')
nl(doc)

h2(doc, '1.2 规则控制器的问题')
tx(doc, '你现在的规则控制器：')
bl(doc, 'SOC 低 → 多开 FC 充电')
bl(doc, 'SOC 高 → 多用电池')
bl(doc, '功率大 → FC 满负荷，电池补充')
nl(doc)
tx(doc, '问题：规则控制器"只看当前"，不知道前面是上坡还是下坡。')
tx(doc, '比如你马上要爬一个大长坡，规则控制器不会提前把电池充满。')
tx(doc, '爬坡时 FC 满负荷还功率不够，电池被迫深度放电，整体效率反而低。')
nl(doc)

h2(doc, '1.3 DP 的"上帝视角"')
tx(doc, 'DP 和规则控制器的根本区别：')
nl(doc)

tbl(doc, ['', '规则控制器', 'DP'],
[['信息来源', '只看当前时刻', '整个工况已知'],
 ['思考方式', 'if-else 即时反应', '全局优化，事前规划'],
 ['最优性', '不一定', '全局最优（离散精度内）'],
 ['实时性', '✅ 可在线运行', '❌ 需离线计算']])
nl(doc)

tx(doc, '打个比方：')
bl(doc, '规则控制器 = 一个没看过导航的司机，凭经验开车', bp='🔵 ')
bl(doc, 'DP = 一个看过全程路况的司机，提前知道哪里该加速、哪里该滑行', bp='🟢 ')
nl(doc)

# ====================================================================
# 第二章：DP 的直觉理解
# ====================================================================
h1(doc, '第二章 · DP 的直觉理解——收费公路类比')
tx(doc, '数学公式容易吓到人，先讲个故事理解 Bellman 最优性原理。')
nl(doc)

h2(doc, '2.1 故事：最短路径问题')
tx(doc, '想象你从北京开车到上海，沿途有多个收费站。每个收费站之间油耗不同。')
tx(doc, '你不知道哪条路最省油，但有一个"神通广大的导航员"——它就是 DP。')
nl(doc)

tx(doc, '导航员的做法是：')
bl(doc, '先把地图倒过来看，从上海往北京想', bp='① ')
bl(doc, '从上海出发（终点），到上海的成本 = 0', bp='② ')
bl(doc, '找所有能"一步到上海"的收费站，记录"从这站到上海的最小成本"', bp='③ ')
bl(doc, '再找能"两步到上海"的站，用刚才算好的结果', bp='④ ')
bl(doc, '以此类推，直到回到北京', bp='⑤ ')
nl(doc)

tx(doc, '关键洞察：当导航员算"从济南到上海的最省油路线"时，')
tx(doc, '他不需要关心"你从北京到济南是怎么开的"——过去的已经过去了。')
tx(doc, '他只需要知道"你现在在济南，后面怎么走最省油"。')
nl(doc)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
r = p.add_run('这就是 Bellman 最优性原理的核心：')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
r = p.add_run('"从当前位置到终点的最优路径，与你怎么到当前位置无关。"'); r.font.italic = True; r.font.size = Pt(10)
nl(doc)

h2(doc, '2.2 把故事翻译回 EMS')
nl(doc)

tbl(doc, ['故事中的概念', 'EMS 中的对应'],
[['北京到上海的路程', 'WLTC 工况（1800s）'],
 ['收费站位置', '每个时间步 k'],
 ['车的油量', 'SOC（电池电量）'],
 ['两个收费站之间的油耗', '单步代价 g_k (氢耗 + SOC惩罚)'],
 ['导航员记录的"从这站到上海的最小成本"', '价值函数 J_k(SOC_k)'],
 ['导航员推荐的"下一步怎么开"', '最优策略 π_k(SOC_k)']])

# ====================================================================
# 第三章：DP 的数学框架
# ====================================================================
h1(doc, '第三章 · 数学框架——其实就一个方程')
tx(doc, '整个 DP 就一个核心方程（Bellman 方程），懂了它 DP 就懂了：')
nl(doc)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.left_indent = Inches(0.5); p.paragraph_format.right_indent = Inches(0.5)
r = p.add_run('Jₖ(xₖ) = min [ gₖ(xₖ, uₖ) + Jₖ₊₁(xₖ₊₁) ]')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
nl(doc)

tx(doc, '翻译成人话：')
nl(doc)

tbl(doc, ['符号', '含义', '在 EMS 中'],
[['k', '当前时间步', 'WLTC 的第几秒'],
 ['xₖ', '当前状态', '当前 SOC'],
 ['uₖ', '当前决策', 'P_fc（FC 功率）'],
 ['gₖ(xₖ, uₖ)', '当前的代价', '这一步的氢耗 + SOC 偏离惩罚'],
 ['xₖ₊₁', '下一步的状态', '电池模型算出的新 SOC'],
 ['Jₖ₊₁(xₖ₊₁)', '未来的最优代价', '"查表——从新 SOC 到终点的最优成本"'],
 ['min', '挑最小的', '枚举所有可能的 P_fc，选总成本最低的']])
nl(doc)

tx(doc, '所以这个方程的物理含义是：')
nl(doc)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
r = p.add_run('"你现在做的决策，既要考虑这一步的消耗，也要考虑对未来的影响。"')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
nl(doc)

# ====================================================================
# 第四章：完整算法流程
# ====================================================================
h1(doc, '第四章 · 完整算法流程（结合代码理解）')

h2(doc, '4.1 算法总览')
tx(doc, 'DP = 后向递推 + 前向仿真，缺一不可：')
nl(doc)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('后向递推（离线，算得慢） → 最优策略表 → 前向仿真（在线，查表快）')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
nl(doc)

h2(doc, '4.2 后向 DP（Backward Induction）')
tx(doc, '从终点倒着算，填一张"价值表"和一张"策略表"。')
nl(doc)

cd(doc, '算法：后向 DP')
cd(doc, '═' * 45)
cd(doc, '输入：P_load[0:N]（工况功率需求）')
cd(doc, '      SOC_grid[0:N_s]（SOC 离散网格）')
cd(doc, '      P_fc_grid[0:N_u]（控制离散网格）')
cd(doc, '参数：电池模型、FC 效率曲线、惩罚系数')
cd(doc, '')
cd(doc, '# 初始化终点价值函数')
cd(doc, 'J[N][:] = 0')
cd(doc, '')
cd(doc, '# 倒着算（从最后一个时间步到第一个）')
cd(doc, 'for k = N-1 downto 0:')
cd(doc, '   for each SOC in SOC_grid:')
cd(doc, '       min_cost = inf')
cd(doc, '       for each P_fc in P_fc_grid:')
cd(doc, '           # ① 功率平衡')
cd(doc, '           P_bat = P_load[k] - P_fc')
cd(doc, '')
cd(doc, '           # ② 状态转移 → 预测下一步 SOC')
cd(doc, '           SOC_next = battery_model(SOC, P_bat)')
cd(doc, '')
cd(doc, '           # ③ 检查约束')
cd(doc, '           if SOC_next < 0.2 or SOC_next > 0.9:')
cd(doc, '               continue  # 不可行，跳过')
cd(doc, '')
cd(doc, '           # ④ 计算单步代价')
cd(doc, '           g = H2_consumption(P_fc) * dt')
cd(doc, '             + alpha * (SOC - SOC_ref)^2')
cd(doc, '')
cd(doc, '           # ⑤ 查未来代价（线性插值）')
cd(doc, '           J_next = interp1(SOC_grid, J[k+1][:], SOC_next)')
cd(doc, '')
cd(doc, '           # ⑥ 总代价 = 当前 + 未来')
cd(doc, '           total_cost = g + J_next')
cd(doc, '')
cd(doc, '           # ⑦ 取最小')
cd(doc, '           if total_cost < min_cost:')
cd(doc, '               min_cost = total_cost')
cd(doc, '               best_u = P_fc')
cd(doc, '')
cd(doc, '       J[k][SOC] = min_cost      # 价值表')
cd(doc, '       π[k][SOC] = best_u        # 策略表')
cd(doc, '')
cd(doc, '输出：J 表（价值函数），π 表（最优策略）')
nl(doc)

tx(doc, '关键点理解：')
bl(doc, 'J[k+1] 是"已经算好的"——因为 k 是从大到小遍历的', bp='① ')
bl(doc, 'interp1 做线性插值：SOC_next 不一定正好在网格点上', bp='② ')
bl(doc, 'SOC_next 超边界 = 这个 P_fc 决策不可行，直接跳过', bp='③ ')
bl(doc, 'H2_consumption(P_fc) 由 FC 效率曲线决定（见附录）', bp='④ ')
nl(doc)

h2(doc, '4.3 前向 Rollout（Forward Simulation）')
tx(doc, '后向 DP 算完后，用 π 表从起点正着跑一遍：')
nl(doc)

cd(doc, '算法：前向 Rollout')
cd(doc, '═' * 45)
cd(doc, '输入：π 表（最优策略），P_load[0:N]')
cd(doc, '      SOC_0 = 0.6（初始 SOC）')
cd(doc, '')
cd(doc, 'for k = 0 to N-1:')
cd(doc, '    # ① 查π表得到最优 P_fc')
cd(doc, '    P_fc[k] = interp1(SOC_grid, π[k][:], SOC[k])')
cd(doc, '')
cd(doc, '    # ② 功率平衡')
cd(doc, '    P_bat[k] = P_load[k] - P_fc[k]')
cd(doc, '')
cd(doc, '    # ③ 电池模型 → 更新 SOC')
cd(doc, '    SOC[k+1] = battery_model(SOC[k], P_bat[k])')
cd(doc, '')
cd(doc, '    # ④ 计算氢耗')
cd(doc, '    m_H2[k] = H2_consumption(P_fc[k]) * dt')
cd(doc, '')
cd(doc, '输出：SOC 轨迹、P_fc 序列、P_bat 序列、累计氢耗')
nl(doc)

tx(doc, '前向 Rollout 只需要查表，不需要再枚举寻优，所以跑得极快（毫秒级）。')
nl(doc)

h2(doc, '4.4 用阿里云文章的代码理解')
tx(doc, '阿里云的文章提供了一段核心 MATLAB 代码，我们逐段拆解：')
nl(doc)

h3(doc, '① 参数初始化')
cd(doc, "Ts = 1;                  % 时间步长 (s)")
cd(doc, "SoC_min = 0.2;           % 最小SOC")
cd(doc, "SoC_max = 0.8;           % 最大SOC")
cd(doc, "P_engine_max = 50;       % 发动机最大功率 (kW)")
cd(doc, "P_motor_max = 20;        % 电机最大功率 (kW)")
cd(doc, "battery_capacity = 5;    % 电池容量 (kWh)")
nl(doc)

h3(doc, '② SOC 网格离散化')
cd(doc, "SoC_grid = linspace(SoC_min, SoC_max, 101);  % SOC网格")
cd(doc, "n_SoC = length(SoC_grid);                   % SOC状态数")
tx(doc, '把 SOC 从 0.2 到 0.8 均匀分成 101 个点。网格越密，精度越高，计算越慢。')
nl(doc)

h3(doc, '③ 核心：逆向递归（最关键的 20 行）')
cd(doc, "for t = time_steps:-1:1       % ★ 从终点往起点倒着走")
cd(doc, "    for i = 1:n_SoC           % 遍历每个SOC状态")
cd(doc, "        current_SoC = SoC_grid(i);")
cd(doc, "        min_cost = inf;")
cd(doc, "        best_action = [];")
cd(doc, "        ")
cd(doc, "        for P_engine = 0:P_engine_max           ★ 枚举所有控制")
cd(doc, "            for P_motor = -P_motor_max:P_motor_max")
cd(doc, "                if (P_engine + P_motor >= 0) && ...")
cd(doc, "                   (next_SoC >= SoC_min) && ...")
cd(doc, "                   (next_SoC <= SoC_max)          ★ 检查约束")
cd(doc, "                    next_SoC = current_SoC + P_motor*Ts/battery_capacity;")
cd(doc, "                    next_i = interp1(..., next_SoC, 'nearest');  ★ 插值")
cd(doc, "                    fuel_cons = P_engine * Ts / 3600;  ★ 油耗")
cd(doc, "                    total_cost = fuel_cons + V(t+1, next_i);  ★ 总代价")
cd(doc, "                    ")
cd(doc, "                    if total_cost < min_cost")
cd(doc, "                        min_cost = total_cost;")
cd(doc, "                        best_action = [P_engine, P_motor];")
cd(doc, "                    end")
cd(doc, "                end")
cd(doc, "            end")
cd(doc, "        end")
cd(doc, "        V(t, i) = min_cost;           % 价值表")
cd(doc, "        policy{t, i} = best_action;    % 策略表")
cd(doc, "    end")
cd(doc, "end")
nl(doc)

h3(doc, '④ 正向仿真（用策略表开车）')
cd(doc, "for t = 1:time_steps")
cd(doc, "    [P_engine, P_motor] = deal(policy{t, initial_i});  % 查表")
cd(doc, "    SoC_history(t+1) = SoC_history(t) + P_motor*Ts/battery_capacity;")
cd(doc, "end")
nl(doc)

# ====================================================================
# 第五章：Kaputt Blog 的精华
# ====================================================================
h1(doc, '第五章 · Kaputt Engineers Blog 精华')

tx(doc, '来源：https://kaputtengineers.wixsite.com/home/post/dynamic-programming')
tx(doc, '配套代码：https://github.com/yeoleparesh/HEV_EMS_DP (96 stars)')
nl(doc)

tx(doc, '这篇博客从代码出发，用 P2 并联混动作为例子，特点是：')
bl(doc, '不讲复杂的数学公式，直接上代码')
bl(doc, '每一步都有配图：功率需求图、SOC 曲线图')
bl(doc, '分析了不同初始 SOC（0.3/0.5/0.7）对结果的影响')
nl(doc)

h2(doc, '5.1 与阿里云文章的区别')
nl(doc)

tbl(doc, ['', '阿里云文章', 'Kaputt Blog'],
[['车型', '串联 HEV', 'P2 并联 HEV'],
 ['控制变量', 'P_engine + P_motor', 'P_batt（电池功率）'],
 ['状态变量', 'SOC', 'SOC'],
 ['代码平台', 'MATLAB（片段）', 'MATLAB（完整可运行）'],
 ['配图', '无', '有（SOC 曲线、功率分配图）'],
 ['配套源码', '无', 'GitHub 仓库']])
nl(doc)

h2(doc, '5.2 该博客的核心代码逻辑')
tx(doc, '它的核心思想和我们的一致，但写法上更紧凑：')
nl(doc)

cd(doc, "% Step 4: 逆向递归寻找最优 P_batt")
cd(doc, "V = zeros(ns, N);            % 价值函数表")
cd(doc, "V(:, N) = 0;                 % 边界条件")
cd(doc, "for i = N-1:-1:1             % ★ 倒着走")
cd(doc, "    for j = 1:ns             % ★ 遍历每个 SOC")
cd(doc, "        P_batt_grid = linspace(lb, ub, 250);  % 电池功率网格")
cd(doc, "        P_eng = P_dem(i) - P_batt_grid;       % 发动机功率")
cd(doc, "        c2g = (ts * fl_wt_en * P_eng) ./ eng_eff(P_eng);  % 代价")
cd(doc, "        SOC_next = SOC_grid(j) - (ts .* P_batt_grid ./ (Q_batt*U_oc));")
cd(doc, "        V_nxt = interp1(SOC_grid, V(:, i+1), SOC_next);  % 插值查表")
cd(doc, "        [V(j,i), k] = min(c2g + V_nxt);")
cd(doc, "        u_opt(j,i) = P_batt_grid(k);  % 最优策略")
cd(doc, "    end")
cd(doc, "end")
nl(doc)

tx(doc, '和阿里云文章的原理完全一样，只是：')
bl(doc, '控制变量选择的是 P_batt（电池功率）而不是 P_eng', bp='• ')
bl(doc, '它用了向量化计算（一次算一整个网格的 SOC_next），比逐点循环更高效', bp='• ')
bl(doc, '随文章有配图，能看到不同初始 SOC 下的最优策略差异', bp='• ')
nl(doc)

# ====================================================================
# 第六章：核心概念的进一步解释
# ====================================================================
h1(doc, '第六章 · 几个容易卡住的概念')

h2(doc, '6.1 为什么叫"动态规划"？')
tx(doc, '"规划" = 在约束下找最优解')
tx(doc, '"动态" = 决策是一步步做出的（随时间变化）')
tx(doc, '合起来："随时间推进、分步骤做出最优决策的方法。"')
nl(doc)

h2(doc, '6.2 为什么倒着算？')
tx(doc, '这是初学者最困惑的地方。')
nl(doc)
tx(doc, '正着想：你在 k=0 时选择 P_fc[0]，但这个选择会影响 SOC[1]，进而影响 SOC[2]...')
tx(doc, '你没法在 k=0 时就准确知道"选择 5kW 还是 10kW 最终更省油"——因为你不知道后面会发生什么。')
nl(doc)
tx(doc, '倒着想：在 k=N-1（最后一步）时，选择很简单，因为后面没有别的事了。')
tx(doc, '在 k=N-2 时，你只需要考虑"这一步 + 最后一步"的总代价——最后一步的最优选择已经算好了。')
tx(doc, '一步步往前推，每一步的"未来代价"都是已知的。')
nl(doc)
tx(doc, '所以倒着算 = "把未知的未来变成已知的未来"，让你每一步都能做出真正的全局最优选择。')
nl(doc)

h2(doc, '6.3 什么是 J_k(x_k) 价值函数？')
tx(doc, 'J_k(x_k) = "从第 k 秒、SOC 为 x_k 的状态出发，到工况结束，最小可能的总氢耗。"')
nl(doc)
tx(doc, '它是一个二维表格：')
cd(doc, '          SOC=0.20  SOC=0.21  SOC=0.22  ...  SOC=0.90')
cd(doc, 'k=1799     0.000    0.000    0.000         0.000   ← 终点代价=0')
cd(doc, 'k=1798    0.012    0.011    0.010         0.015   ← 最后一步')
cd(doc, 'k=1797    0.025    0.023    0.021         0.032')
cd(doc, '...')
cd(doc, 'k=0       2.345    2.312    2.289         2.567   ← 起点')
nl(doc)
tx(doc, '表格中的每个数字代表："如果我在这个时刻处于这个 SOC，后面总共还要烧多少氢。"')
tx(doc, 'DP 的目标就是让 k=0, SOC=0.6 时的 J 值最小。')
nl(doc)

h2(doc, '6.4 什么是 π_k(x_k) 策略表？')
tx(doc, 'π_k(x_k) = "在第 k 秒、SOC 为 x_k 时，燃料电池应该输出多少功率。"')
nl(doc)
tx(doc, '这也是一个二维表格，存的是"最优的 P_fc 值"：')
cd(doc, '          SOC=0.20  SOC=0.21  ...  SOC=0.60  ...  SOC=0.90')
cd(doc, 'k=0       15.2 kW  14.8 kW       8.5 kW        3.0 kW')
cd(doc, 'k=1       16.1 kW  15.5 kW       9.2 kW        3.0 kW')
cd(doc, '...')
cd(doc, 'k=1799     3.0 kW   3.0 kW       3.0 kW        0.0 kW')
nl(doc)
tx(doc, '前向 Rollout 就是一行行查这张表，得到每个时刻的 P_fc 值。')
nl(doc)

h2(doc, '6.5 为什么需要 SOC 维持惩罚？')
tx(doc, '如果只最小化氢耗，DP 会发现"用电池最省氢"（因为电池不"烧"氢），')
tx(doc, '于是它会尽量用电池的电，SOC 一路下降到 0.2 甚至更低。')
tx(doc, '这不公平——你在消耗电池的储备能量，"作弊"降低氢耗。')
nl(doc)
tx(doc, '解决方法：在代价函数中加一项 (SOC - SOC_ref)² × α。')
bl(doc, 'SOC 偏离 0.6 越远，惩罚越大 → DP 会尽量维持 SOC 在 0.6 附近', bp='• ')
bl(doc, 'α 的大小决定"维持力度"：α 大 → SOC 几乎恒定；α 小 → SOC 波动更大、氢耗更低', bp='• ')
bl(doc, '一般 α 选到让 SOC 终值 ≈ SOC 初值（电荷维持）', bp='• ')

# ====================================================================
# 第七章：把 DP 联系回我们的项目
# ====================================================================
h1(doc, '第七章 · 联系我们的项目')

h2(doc, '7.1 我们已经有的东西')
bl(doc, 'WLTC 工况数据 ✅ → 可以算出 P_load[k] 作为 DP 的输入')
bl(doc, '规则控制器代码 ✅ → 作为对比基准')
bl(doc, '电池模型 ✅ → 用于 SOC 状态转移')
bl(doc, 'Python 环境 ✅ → numpy/scipy 做插值')
nl(doc)

h2(doc, '7.2 还需要的')
bl(doc, 'FC 效率曲线 η_fc(P_fc) → 用于计算氢耗代价')
bl(doc, 'SOC 网格和 P_fc 网格离散化参数')
bl(doc, 'DP 核心算法（后向 + 前向）')
nl(doc)

h2(doc, '7.3 对比方案')
tx(doc, '跑完之后生成对比图，你会看到：')
nl(doc)

cd(doc, '            规则控制器                 DP')
cd(doc, '  SOC:  ┌────────────────┐   ┌────────────────┐')
cd(doc, '        │  波动较大       │   │  平滑优化      │')
cd(doc, '        └────────────────┘   └────────────────┘')
cd(doc, '  氢耗:     0.35 kg              0.30 kg')
cd(doc, '  FC效率:    ~45%                 ~50%')
nl(doc)

tx(doc, 'DP 的改善预期：氢耗降低 10-15%，FC 平均效率提升约 5 个百分点。')
nl(doc)

# ====================================================================
# 第八章：推荐学习资源
# ====================================================================
h1(doc, '第八章 · 推荐学习资源（按顺序看）')

nl(doc)

tbl(doc, ['序号', '资源', '语言', '类型', '链接'],
[['1', '阿里云开发者社区', '中文', '教程+代码', 'https://developer.aliyun.com/article/1728349'],
 ['2', 'Kaputt Engineers Blog', '英文', '教程+代码+图', 'https://kaputtengineers.wixsite.com/home/post/dynamic-programming'],
 ['3', 'GitHub: HEV_EMS_DP', '英文', '完整MATLAB代码', 'https://github.com/yeoleparesh/HEV_EMS_DP'],
 ['4', 'YouTube: EMS分类讲解', '英文', '视频~30min', 'https://www.youtube.com/watch?v=Blo9vyV_QDE'],
 ['5', '自动化学报综述', '中文', 'PDF综述论文', 'https://www.aas.net.cn/fileZDHXB/journal/article/zdhxb/2016/3/PDF/zdhxb-42-3-321.pdf']])
nl(doc)

tx(doc, '建议阅读顺序：')
bl(doc, '第一步：看这篇文档的"收费公路类比"和"算法流程"（已在本文件中）', bp='📖 ')
bl(doc, '第二步：打开阿里云文章，对着它的代码看一遍', bp='💻 ')
bl(doc, '第三步：下载 GitHub 仓库，在自己电脑上跑一遍', bp='🏃 ')
bl(doc, '第四步：回来看我们项目的 scripts/dp_ems.py（后续代码实现）', bp='🔧 ')

# ====================================================================
# 附录
# ====================================================================
h1(doc, '附录：FC 效率曲线参考')
tx(doc, 'DP 中氢耗计算的依据是 FC 效率曲线。以下是典型 PEMFC 效率数据：')
nl(doc)

tbl(doc, ['P_fc (kW)', '效率 η_fc (%)', '氢耗 (g/s)', '说明'],
[['0', '0', '0.000', '关机（无氢耗）'],
 ['2', '28', '0.060', '低效区'],
 ['5', '40', '0.104', '提升中'],
 ['10', '50', '0.167', '接近高效区'],
 ['15', '55', '0.227', '★ 最佳效率点'],
 ['20', '53', '0.315', '高效区边缘'],
 ['25', '48', '0.434', '额定功率'],
 ['30', '40', '0.625', '峰值功率']])
nl(doc)

tx(doc, '氢耗公式：ṁ_H₂ = P_fc / (η_fc × LHV_H₂)，其中 LHV_H₂ = 120 MJ/kg')
nl(doc)

# ====================================================================
# 参考文献
# ====================================================================
h1(doc, '参考文献')

nl(doc)
h2(doc, '网络资源')
nl(doc)

tbl(doc, ['编号', '标题', '类型', '链接'],
[['[1]', '基于动态规划算法的混合动力汽车能量管理建模与计算（阿里云开发者社区）', '中文教程+代码', 'https://developer.aliyun.com/article/1728349'],
 ['[2]', 'Dynamic Programming for HEV Energy Management（Kaputt Engineers Blog）', '英文教程+代码+图', 'https://kaputtengineers.wixsite.com/home/post/dynamic-programming'],
 ['[3]', 'GitHub: yeoleparesh/HEV_EMS_DP — DP for HEV Energy Management', 'MATLAB 完整实现', 'https://github.com/yeoleparesh/HEV_EMS_DP'],
 ['[4]', 'L 22 Energy Management Strategy, Classification（YouTube）', '英文视频讲解', 'https://www.youtube.com/watch?v=Blo9vyV_QDE'],
 ['[5]', '混合动力电动汽车能量管理策略研究综述（自动化学报）', '中文综述 PDF', 'https://www.aas.net.cn/fileZDHXB/journal/article/zdhxb/2016/3/PDF/zdhxb-42-3-321.pdf'],
 ['[6]', 'Optimization of EMS for FCHEV Based on DP（MDPI Energies）', '英文论文(Open Access)', 'https://www.mdpi.com/1996-1073/15/12/4325']])
nl(doc)

h2(doc, '学术文献')
nl(doc)

tbl(doc, ['编号', '引用', '说明'],
[['[7]', 'Bellman, R. Dynamic Programming. Princeton University Press, 1957.', 'DP 经典著作，Bellman 最优性原理原始出处'],
 ['[8]', 'Guzzella, L. & Sciarretta, A. Vehicle Propulsion Systems. Springer, 2013.', '车辆动力学与能量管理经典教材']])
nl(doc)

h2(doc, '本项目文件')
nl(doc)

tbl(doc, ['文件', '说明'],
[['STATUS.md', 'EMS-PLAN 进度跟踪'],
 ['experiments/run_ems_simulation.py', 'EMS 仿真启动器（Python + MATLAB）'],
 ['env/simulink_models/Use-Model/ems_controller_fcn.m', '规则基 EMS 控制器'],
 ['env/simulink_models/Use-Model/battery_simple_fcn.m', '简化 R-int 电池模型'],
 ['results/wltc_cycle.csv', 'WLTC 工况数据']])
nl(doc)

# ====================================================================
# 页脚
# ====================================================================
nl(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('═' * 50); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DP 入门笔记 — EMS-PLAN 第3周\n'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
r = p.add_run('生成日期：2026-06-06 | 项目：F:/CLAUDE/research/ems-platform\n'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
r = p.add_run('参考文献 [1]~[8] 详见上方列表'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

# ====================================================================
fname = 'F:/CLAUDE/research/ems-platform/docs/DP入门笔记_网络资源精华版.docx'
doc.save(fname)
print('OK:', fname)
