# -*- coding: utf-8 -*-
"""生成 Battery 模型解释文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, t, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(t).font.size = Pt(10)
    else:
        r = p.add_run(t); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

doc = Document()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Battery 简化模型逐段解释'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Day7 - EMS顶层模型搭建 | 文件: env/simulink_models/battery_simple_fcn.m'); r.font.size = Pt(11)
doc.add_paragraph()

# ===== 整体结构 =====
h1(doc, '整体结构')
tx(doc, 'battery_simple_fcn.m 是一个简化锂离子电池模型（R-int 模型），用于 EMS 仿真中计算电池的电压、电流和 SOC 变化。')
tx(doc, '函数签名：')
cd(doc, 'function [SOC, V_bat, I_bat] = battery_simple_fcn(P_bat_kW, SOC_init, dt)')
tbl(doc, ['参数', '含义', '单位'],
[['P_bat_kW', '电池功率（正=放电，负=充电）', 'kW'],
 ['SOC_init', '当前荷电状态', '0-1'],
 ['dt', '时间步长', 's'],
 ['SOC（输出）', '更新后的荷电状态', '0-1'],
 ['V_bat（输出）', '端电压', 'V'],
 ['I_bat（输出）', '电流（正=放电，负=充电）', 'A']])
doc.add_paragraph()

# ===== 电池参数 =====
h1(doc, '第一段：电池参数')
cd(doc, 'Q_bat = 50;         % 容量 [Ah]')
cd(doc, 'V_nom = 350;        % 额定电压 [V]')
cd(doc, 'R_int = 0.05;       % 内阻 [Ohm]')
tx(doc, '参考 50Ah 锂离子电池包（约 96 节 3.7V 电芯串联 ≈ 355V）。')
tbl(doc, ['参数', '值', '说明'],
[['Q_bat', '50 Ah', '电池容量，决定 SOC 变化速度'],
 ['V_nom', '350 V', '额定电压，用于 OCV 曲线范围参考'],
 ['R_int', '0.05 Ohm', '内阻，决定充放电电压降']])
doc.add_paragraph()

# ===== OCV-SOC =====
h1(doc, '第二段：SOC → 开路电压（OCV-SOC 查表）')
tx(doc, 'OCV（开路电压）随 SOC 变化的曲线是电池的核心特性。这里用查表法近似：')
cd(doc, 'SOC_breakpoints = [0, 0.1, 0.2, 0.3, 0.5, 0.7, 0.8, 0.9, 1.0];')
cd(doc, 'V_ocv_lookup    = [320, 330, 338, 345, 352, 358, 362, 368, 380];')
cd(doc, 'V_oc = interp1(SOC_breakpoints, V_ocv_lookup, SOC_init, linear, extrap);')

h2(doc, 'OCV-SOC 曲线特征')
tx(doc, '锂离子电池的 OCV-SOC 曲线呈反 S 形：中间平缓、两端陡峭。')
tx(doc, '插值点密度反映这个特征——SOC 20%-80% 取点稀疏（区域平缓），两端取点较密（变化陡）。')

tbl(doc, ['SOC', 'OCV (V)', '说明'],
[['0% (空电)', '320', '接近放空，电压最低'],
 ['10%', '330', '电压开始急剧下降'],
 ['20%', '338', '过渡区'],
 ['30%', '345', '进入平台区'],
 ['50%', '352', '中间点，接近额定'],
 ['70%', '358', '平台区'],
 ['80%', '362', '开始上升'],
 ['90%', '368', '快速上升'],
 ['100% (满电)', '380', '充满']])

tx(doc, '当前 SOC 介于两个断点之间时，interp1 用线性插值计算对应的 OCV。')
tx(doc, 'SOC 超出 0-1 范围时，extrap 选项外推而非报错。')

doc.add_paragraph()

# ===== 电流计算 =====
h1(doc, '第三段：电流计算（核心）')
tx(doc, 'R-int 模型用二次方程求解电流。')

h2(doc, '模型方程')
tx(doc, '已知：端电压 V_t = V_oc - I × R_int，功率 P = V_t × I，代入得：')
cd(doc, 'P = (V_oc - I·R) · I  =  V_oc·I - I²·R')
tx(doc, '整理为一元二次方程的标准形式：')
cd(doc, 'I²R - V_oc·I + P = 0')
tx(doc, '其中 a=R, b=-V_oc, c=P。')

h2(doc, '判别式')
cd(doc, 'Δ = b² - 4ac = V_oc² - 4·R·P')
tbl(doc, ['判别式情况', '含义', '处理'],
[['Δ > 0', '功率在电池能力范围内', '两个实数解，取物理正确的那个'],
 ['Δ = 0', '刚好在最大功率点', '一个实数解，但数值不稳定'],
 ['Δ < 0', '功率超出电池物理极限', '限幅到最大功率的 99%']])

doc.add_paragraph()

h2(doc, '公式一：功率过小直接归零')
cd(doc, 'if abs(P_bat_kW) < 0.01')
cd(doc, '    I_bat = 0;  % 接近0')
tx(doc, '功率绝对值 < 0.01kW（10W），认为功率太小，直接设电流为 0，避免后续方程计算的小数误差。')

doc.add_paragraph()

h2(doc, '公式二：功率限幅（防止 Δ < 0）')
cd(doc, 'if Delta < 0')
cd(doc, '    P_w = V_oc^2 / (4 * R_int);')
tx(doc, '电池理论最大功率 P_max = V_oc² / 4R（负载电阻等于内阻时的最大功率传输）。')
tx(doc, '当需求功率超过 P_max 时，强制限幅到 P_max 的 99%，保留 1% 安全余量。')
tbl(doc, ['情形', '限幅公式'],
[['放电 (P_bat_kW > 0)', 'P_w = min(P_w, 0.99·P_max)'],
 ['充电 (P_bat_kW < 0)', 'P_w = -min(|P_w|, 0.99·P_max)']])

doc.add_paragraph()

h2(doc, '公式三：求根公式解电流')
cd(doc, 'I_bat = (V_oc - sqrt(Delta)) / (2 * R_int);')
tx(doc, '二次方程求根公式 I = (V_oc ± √Δ) / 2R。')
tx(doc, '为什么取减号而不是加号？')

tbl(doc, ['情形', 'V_oc 与 √Δ', 'V_oc - √Δ', '物理含义'],
[['放电 (P>0)', '√Δ < V_oc', '> 0, I>0', '正电流放电 ✅'],
 ['充电 (P<0)', '√Δ > V_oc', '< 0, I<0', '负电流充电 ✅']])

tx(doc, '如果取加号 (V_oc + √Δ) / 2R，无论充放电电流都是正的——物理上不可能。')

doc.add_paragraph()

h2(doc, '公式四：限流保护')
cd(doc, 'if P_bat_kW > 0')
cd(doc, '    I_bat = min(I_bat,  300);   % 放电: 电流不超过 +300A')
cd(doc, 'else')
cd(doc, '    I_bat = max(I_bat, -300);   % 充电: 电流不低于 -300A')

tbl(doc, ['情形', '限流规则', '例子'],
[['放电 (I>0)', 'I 不超过 +300A', 'I=350A → 限到 300A'],
 ['充电 (I<0)', 'I 不低于 -300A', 'I=-350A → 限到 -300A']])

doc.add_paragraph()

# ===== 端电压 =====
h1(doc, '第四段：端电压计算')
cd(doc, 'V_bat = V_oc - I_bat * R_int;')
tx(doc, '欧姆定律，非常简单。内阻 R_int 越大、电流 I 越大，端电压偏离 OCV 越多。')
tx(doc, '放电时 I>0，V_bat < V_oc（电压跌落）；充电时 I<0，V_bat > V_oc（电压抬升）。')
doc.add_paragraph()

# ===== SOC 更新 =====
h1(doc, '第五段：SOC 更新（Ah 积分法）')
cd(doc, 'SOC_change = -I_bat / (Q_bat * 3600) * dt;')
cd(doc, 'SOC = SOC_init + SOC_change;')
tbl(doc, ['变量', '含义', '单位'],
[['I_bat', '电流', 'A'],
 ['Q_bat * 3600', '容量换算为库仑', 'As (安秒)'],
 ['dt', '时间步长', 's']])

tx(doc, '放电时 I>0 → SOC_change < 0 → SOC 下降；充电时相反。')
tx(doc, '/ 3600 是因为 Ah 要换算成秒（1Ah = 3600As）。')

cd(doc, 'SOC = min(max(SOC, 0.05), 0.95);')
tx(doc, '最后用 min+max 钳位到 [0.05, 0.95]，防止 SOC 越界导致后续计算异常。')

doc.add_paragraph()

# ===== 完整流程图 =====
h1(doc, '算法流程')
tx(doc, '输入 P_bat, SOC_init, dt →')
tx(doc, '  |')
tx(doc, '  ├→ 查 OCV-SOC 表 → V_oc')
tx(doc, '  ├→ 解 I²R - V_oc·I + P = 0 → I_bat')
tx(doc, '  │    ├ Δ < 0 → 限功率 → 重新算')
tx(doc, '  │    └ Δ ≥ 0 → 求根公式')
tx(doc, '  ├→ 限流 ±300A')
tx(doc, '  ├→ V_bat = V_oc - I·R')
tx(doc, '  └→ SOC_new = SOC_init - I·dt / (Q·3600)')
tx(doc, '  |')
tx(doc, '输出 SOC, V_bat, I_bat')
doc.add_paragraph()

doc.add_paragraph()
r = doc.add_paragraph().add_run('='*50+'\nBattery 简化模型——R-int 模型逐段解释\n文件: env/simulink_models/battery_simple_fcn.m\n生成日期：2026-06-05\n'+'='*50)
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

fname = 'F:/CLAUDE/research/ems-platform/docs/Day7_Battery_model_explain.docx'
doc.save(fname)
print('OK:', fname)
