#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 week9_complete.py 逐行代码分析文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT = 'F:/CLAUDE/research/ems-platform/docs/Week9_完整代码逐行分析.docx'

def shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color})
    tc.append(el)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading(c, '2F5496')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ct
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def h3(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

def tx(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10)

def bl(doc, text, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        r = p.add_run(text); r.font.size = Pt(10)

def cd(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

def note(doc, t):
    p = doc.add_paragraph(); r = p.add_run('* '); r.font.size = Pt(10); r.italic = True
    r = p.add_run(t); r.font.size = Pt(10); r.italic = True; r.font.color.rgb = RGBColor(0x55,0x55,0x55)


doc = Document()

# ── 封面 ──
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Week 9 完整代码逐行分析'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('文件: scripts/week9_complete.py (698 行, 8 Part + Extra)'); r.font.size = Pt(11)
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# 整体结构
# ═══════════════════════════════════════════════════════════════
h1(doc, '整体结构')
tx(doc, 'week9_complete.py 将 PyTorch 入门和 RL 基础合并为一个文件，共 8 个 Part + Extra + 验证。')

tbl(doc, ['Part', '行号', '内容', '类型'],
[['Part 1', '40-77', 'Tensor 基础', 'PyTorch'],
 ['Part 2', '83-112', 'Autograd 自动求导', 'PyTorch'],
 ['Part 3', '118-146', 'nn.Module + MLP', 'PyTorch'],
 ['Part 4', '152-253', 'MLP 功率预测', 'PyTorch 实战'],
 ['Part 5', '259-324', 'MDP 五元组 — GridWorld', 'RL 基础'],
 ['Part 6', '330-404', 'Bellman 方程', 'RL 基础'],
 ['Part 7', '410-496', '策略迭代', 'RL 基础'],
 ['Part 8', '502-558', '值迭代', 'RL 基础'],
 ['Extra', '564-619', '收敛过程可视化', 'RL 图'],
 ['验证', '680-692', '策略迭代 vs 值迭代 一致性检查', '验证']])
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# 引言部分 (L1-35)
# ═══════════════════════════════════════════════════════════════
h1(doc, '引言与导入 (第 1-35 行)')

h2(doc, '第 1-20 行：文件头')
cd(doc, "#!/usr/bin/env python3")
cd(doc, "# -*- coding: utf-8 -*-")
cd(doc, '"""')
cd(doc, 'Week 9 — PyTorch + RL 基础 完整通关')
cd(doc, '====================================')
cd(doc, 'PyTorch 篇 (Part 1-4)')
cd(doc, '  Part 1: Tensor 基础')
cd(doc, '  Part 2: Autograd 自动求导')
cd(doc, '  Part 3: nn.Module + MLP')
cd(doc, '  Part 4: MLP 功率预测（FC 功率预测）')
cd(doc, '')
cd(doc, 'RL 基础篇 (Part 5-8)')
cd(doc, '  Part 5: MDP 五元组 — GridWorld')
cd(doc, '  Part 6: Bellman 方程')
cd(doc, '  Part 7: 策略迭代')
cd(doc, '  Part 8: 值迭代')
cd(doc, '')
cd(doc, '前置: pip install torch, numpy, matplotlib')
cd(doc, '输出: results/week9_complete_*.png')
cd(doc, '"""')
bl(doc, '#!/usr/bin/env python3：Unix/Linux 系统下直接运行的 shebang，Windows 忽略')
bl(doc, '# -*- coding: utf-8 -*-：声明文件编码为 UTF-8，支持中文')
bl(doc, '三引号字符串：模块文档字符串 (docstring)，描述文件用途、内容、前置条件和输出')
doc.add_paragraph()

h2(doc, '第 22-26 行：导入库')
cd(doc, "import os, sys, itertools")
cd(doc, "import numpy as np")
cd(doc, "import matplotlib")
cd(doc, "matplotlib.use('Agg')")
cd(doc, "import matplotlib.pyplot as plt")
bl(doc, 'os, sys：系统路径操作', bp='os/sys：')
bl(doc, 'itertools：生成网格坐标 (笛卡尔积) — Part 5 MDP 用', bp='itertools：')
bl(doc, 'numpy as np：数值计算库，RL 部分大量使用', bp='numpy：')
bl(doc, "matplotlib.use('Agg')：设置 Agg 后端（不弹出窗口，只保存图片），服务器/脚本环境必须加")
bl(doc, 'pyplot as plt：画图接口')
doc.add_paragraph()

h2(doc, '第 28-29 行：结果目录')
cd(doc, "RESULTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'results')")
cd(doc, 'os.makedirs(RESULTS_DIR, exist_ok=True)')
bl(doc, "os.path.dirname(__file__)：当前文件所在目录 (scripts/)")
bl(doc, "再上一层：项目根目录 (ems-platform/)")
bl(doc, "拼接 'results' → results/ 目录")
bl(doc, 'exist_ok=True：目录已存在时不报错')
doc.add_paragraph()

h2(doc, '第 31-34 行：中文字体设置')
cd(doc, "# 中文字体（Windows）")
cd(doc, "plt.rcParams['font.family'] = 'sans-serif'")
cd(doc, "plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'SimSun']")
cd(doc, "plt.rcParams['axes.unicode_minus'] = False")
bl(doc, '设置 matplotlib 使用微软雅黑/黑体/宋体显示中文')
bl(doc, "unicode_minus=False：解决负号显示为方块的问题")
note(doc, '这是在 Windows 上的配置，Linux/macOS 需要换成对应中文字体。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Part 1
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Part 1: Tensor 基础 (第 37-77 行)')
tx(doc, '目标：掌握 PyTorch Tensor 的创建、类型、形状操作、广播机制。')

h2(doc, '第 40-41 行：函数定义')
cd(doc, "def part1_tensor_basics():")
cd(doc, '    """PyTorch Tensor 的创建与基本操作."""')
bl(doc, '无参数，无返回值（返回 device 给调用者，但实际仅用于打印）')
doc.add_paragraph()

h2(doc, '第 42 行：导入 PyTorch')
cd(doc, "    import torch")
tx(doc, '在每个函数内部导入 torch（而不是在文件顶部），好处是：首次调用时才会实际加载 PyTorch 库。如果只改 RL 部分的代码，不需要等 PyTorch 加载。但也说明这并不是一个性能敏感的模块加载策略。')
doc.add_paragraph()

h2(doc, '第 44-46 行：从列表创建 Tensor')
cd(doc, "    t1 = torch.tensor([[1, 2], [3, 4]])")
cd(doc, '    print(f"t1:\\n{t1}, dtype={t1.dtype}, shape={t1.shape}")')
tx(doc, 'torch.tensor() 从 Python 列表创建张量，自动推断 dtype=int64。t1.shape 返回 torch.Size([2, 2])。')
note(doc, 'torch.tensor() 和 torch.Tensor() 有区别：tensor() 自动推断 dtype，Tensor() 默认 float32。')
doc.add_paragraph()

h2(doc, '第 48-51 行：从 numpy 创建（共享内存）')
cd(doc, "    a = np.array([1.0, 2.0, 3.0])")
cd(doc, "    t2 = torch.from_numpy(a)")
cd(doc, "    print(f't2: {t2}')")
tx(doc, 'torch.from_numpy() 和 numpy 数组共享内存——修改一个会影响另一个。这节省了复制开销，但需要注意数据安全性。')
doc.add_paragraph()

h2(doc, '第 53-57 行：特殊张量')
cd(doc, "    zeros = torch.zeros(2, 3)")
cd(doc, "    ones = torch.ones(2, 3)")
cd(doc, "    rand = torch.randn(3, 3)  # 标准正态")
cd(doc, "    print(f'zeros: {zeros.shape}, ones: {ones.shape}, randn: {rand.shape}')")
tbl(doc, ['函数', '值', '用途'],
[['zeros(2,3)', '全 0', '初始化、填充'],
 ['ones(2,3)', '全 1', '初始化'],
 ['randn(3,3)', '标准正态随机', '权重初始化']])
doc.add_paragraph()

h2(doc, '第 59-64 行：设备管理')
cd(doc, "    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')")
cd(doc, "    print(f'Device: {device}')")
cd(doc, "    t_gpu = torch.tensor([1, 2, 3], device=device)")
tx(doc, '这段代码实现了"自动切换 CPU/GPU"——有 CUDA 显卡就用 GPU，否则用 CPU。这是 PyTorch 写可移植代码的标准写法。')
note(doc, '本项目是 CPU 版本，但代码已做好兼容，以后升级 CUDA 不需要改代码。')
doc.add_paragraph()

h2(doc, '第 66-70 行：索引与形状操作')
cd(doc, "    t = torch.randn(4, 5)")
cd(doc, "    print(f't[0]: {t[0].shape}, t[:, :3]: {t[:, :3].shape}')")
cd(doc, "    print(f't.view(-1): {t.view(-1).shape}')       # 展平")
cd(doc, "    print(f't.reshape(2, 10): {t.reshape(2, 10).shape}')  # 重塑")
bl(doc, 't[0]：取第 0 行，shape 从 (4,5) → (5,)', bp='索引：')
bl(doc, 't[:, :3]：取所有行 + 前 3 列 → (4,3)')
bl(doc, 'view(-1)：-1 表示自动计算，展平成 1D → (20,)', bp='view vs reshape：')
bl(doc, 'reshape(2,10)：重排为 2 行 10 列 → (2,10)')
note(doc, 'view 要求内存连续，reshape 更安全（自动处理不连续情况）。')
doc.add_paragraph()

h2(doc, '第 72-75 行：广播机制')
cd(doc, "    a = torch.tensor([[1], [2], [3]])  # (3,1)")
cd(doc, "    b = torch.tensor([10, 20, 30])     # (3,)")
cd(doc, "    print(f'broadcast add:\\n{a + b}')")
tx(doc, '广播规则：从右向左对齐维度，要么相等要么为 1。')
tx(doc, 'a 是 (3,1)，b 是 (3,)。PyTorch 自动把 a 横向复制为 (3,3)，把 b 纵向复制为 (3,3)，再相加。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Part 2
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Part 2: Autograd 自动求导 (第 80-112 行)')
tx(doc, '目标：理解 PyTorch 的自动求导引擎——计算图、backward()、梯度累积、no_grad。')

h2(doc, '第 87-96 行：基本 backward 演示')
cd(doc, "    x = torch.tensor([2.0, 3.0], requires_grad=True)")
cd(doc, "    y = x ** 2 + 3 * x")
cd(doc, "    loss = y.sum()")
cd(doc, "    loss.backward()")
cd(doc, "    print(f'gradient dy/dx: {x.grad}')")
bl(doc, 'requires_grad=True：告诉 PyTorch "我要对这个变量求导"，开始追踪所有涉及 x 的运算', bp='核心概念：')
bl(doc, '计算图构建：x → pow(2) 和 mul(3) → add → y → sum → loss')
bl(doc, "loss.backward()：从 loss 出发反向传播，自动计算 ∂loss/∂x")
bl(doc, '理论值：y = x² + 3x, dy/dx = 2x + 3, 在 x=[2,3] 处 = [7,9] → 验证通过 ✓')
doc.add_paragraph()

h2(doc, '第 98-99 行：梯度清零')
cd(doc, "    x.grad.zero_()")
tx(doc, 'PyTorch 默认累积梯度（累加），不清零会导致多个 batch 的梯度混合。每次 backward 前必须 zero_grad()。')
tip_text = doc.add_paragraph()
r = tip_text.add_run('🔥 ')
r.font.size = Pt(10)
r = tip_text.add_run('忘记 optimizer.zero_grad() 是新手最常见 bug！')
r.font.size = Pt(10); r.bold = True; r.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
doc.add_paragraph()

h2(doc, '第 101-107 行：链式法则')
cd(doc, "    a = torch.tensor(2.0, requires_grad=True)")
cd(doc, "    b = torch.tensor(3.0, requires_grad=True)")
cd(doc, "    z = (a ** 2) * torch.sin(b)")
cd(doc, "    z.backward()")
cd(doc, "    print(f'dz/da = {a.grad:.4f}')  # 2*a*sin(b) = 4*0.1411")
cd(doc, "    print(f'dz/db = {b.grad:.4f}')  # a^2*cos(b) = 4*(-0.99)")
tx(doc, '多变量的链式法则验证。z = a²·sin(b)，∂z/∂a = 2a·sin(b)，∂z/∂b = a²·cos(b)。手动计算验证结果。')
doc.add_paragraph()

h2(doc, '第 109-112 行：no_grad 模式')
cd(doc, "    with torch.no_grad():")
cd(doc, "        y_eval = x ** 2 + 3 * x")
cd(doc, "        print(f'no_grad: {y_eval}')")
tx(doc, 'torch.no_grad() 上下文管理器关闭梯度追踪。推理/验证时：省显存、加速、不构建计算图。必须搭配 model.eval() 使用。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Part 3
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Part 3: nn.Module + MLP (第 115-146 行)')
tx(doc, '目标：用 PyTorch 的 nn.Module 构建一个多层感知机。')

h2(doc, '第 118-134 行：定义 MLP 类')
cd(doc, "    class MLP(nn.Module):")
cd(doc, "        def __init__(self, input_dim, hidden_dim, output_dim):")
cd(doc, "            super().__init__()")
cd(doc, "            self.fc1 = nn.Linear(input_dim, hidden_dim)")
cd(doc, "            self.fc2 = nn.Linear(hidden_dim, output_dim)")
cd(doc, "")
cd(doc, "        def forward(self, x):")
cd(doc, "            x = F.relu(self.fc1(x))")
cd(doc, "            x = self.fc2(x)")
cd(doc, "            return x")

tbl(doc, ['代码', '说明'],
[['class MLP(nn.Module)', '继承 nn.Module — 所有网络模型都必须继承这个基类'],
 ["super().__init__()", '调用父类构造函数，注册所有子模块'],
 ["self.fc1 = nn.Linear(4, 32)", '全连接层: y = xW^T + b，输入4维→输出32维'],
 ["F.relu(self.fc1(x))", 'ReLU 激活函数: max(0, x)，给网络引入非线性'],
 ['forward()', '前向传播函数——叫 model(x) 时自动调用']])
doc.add_paragraph()

h2(doc, '第 136-139 行：实例化')
cd(doc, "    model = MLP(input_dim=4, hidden_dim=32, output_dim=1)")
cd(doc, "    print(f'Model:\\n{model}')")
cd(doc, "    print(f'Parameters: {sum(p.numel() for p in model.parameters())}')")
tx(doc, '创建输入4维、隐藏层32维、输出1维的 MLP。总参数量 = 4×32 + 32(bias) + 32×1 + 1(bias) = 128+32+32+1 = 193。')
doc.add_paragraph()

h2(doc, '第 141-144 行：前向传播')
cd(doc, "    x = torch.randn(10, 4)  # batch=10, features=4")
cd(doc, "    y = model(x)")
cd(doc, "    print(f'Input: {x.shape}, Output: {y.shape}')")
tx(doc, '输入 10 个样本（batch=10），每个 4 个特征。输出 10 个值，形状 (10,1)。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Part 4
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Part 4: MLP 功率预测 (第 149-253 行)')
tx(doc, '目标：完整的 PyTorch 实战——用 MLP 根据历史功率预测下一步燃料电池功率。这是后续 RL-EMS 项目的数据预处理基础。')

h2(doc, '第 161-165 行：生成模拟数据')
cd(doc, "    np.random.seed(42)")
cd(doc, "    t = np.linspace(0, 100, 1000)")
cd(doc, "    power = 30 + 15 * np.sin(0.1 * t) + 5 * np.sin(0.5 * t) + np.random.randn(1000) * 2")
cd(doc, "    power = np.clip(power, 10, 80)")
tx(doc, '用正弦波叠加模拟燃料电池功率信号：均值30kW、振幅15kW和5kW两种波动、加高斯噪声(σ=2kW)、限幅在[10,80]kW（实际 FC 功率范围）。')
doc.add_paragraph()

h2(doc, '第 167-184 行：构建时序样本')
cd(doc, "    def create_sequences(data, seq_len=10):")
cd(doc, "        X, y = [], []")
cd(doc, "        for i in range(len(data) - seq_len):")
cd(doc, "            X.append(data[i:i + seq_len])")
cd(doc, "            y.append(data[i + seq_len])")
cd(doc, "        return np.array(X), np.array(y)")
tx(doc, '滑动窗口法构建训练样本：用过去 10 步 (t-9,...,t) → 预测第 t+1 步。总共 1000-10 = 990 个样本。')
tx(doc, '80% 训练 (792个), 20% 测试 (198个)。然后转为 float32 tensor，用 unsqueeze(-1) 加最后一维供 LSTM 使用。')
doc.add_paragraph()

h2(doc, '第 187-202 行：定义 PowerPredictor 模型')
cd(doc, "    class PowerPredictor(nn.Module):")
cd(doc, "        def __init__(self, seq_len, hidden=64):")
cd(doc, "            self.fc1 = nn.Linear(seq_len, hidden)     # 10→64")
cd(doc, "            self.fc2 = nn.Linear(hidden, hidden)      # 64→64")
cd(doc, "            self.fc3 = nn.Linear(hidden, 1)           # 64→1")
cd(doc, "            self.dropout = nn.Dropout(0.1)            # 10%随机丢弃")
cd(doc, "")
cd(doc, "        def forward(self, x):")
cd(doc, "            x = x.view(x.size(0), -1)   # (batch,10,1)→(batch,10)")
cd(doc, "            x = F.relu(self.fc1(x))")
cd(doc, "            x = self.dropout(x)          # 防过拟合")
cd(doc, "            x = F.relu(self.fc2(x))")
cd(doc, "            x = self.fc3(x)              # 线性输出，回归任务")
cd(doc, "            return x")
tx(doc, '3 层 MLP (10→64→64→1) + Dropout。view(x.size(0), -1) 把 (batch, 10, 1) 展平成 (batch, 10) 供 Linear 层使用。')
note(doc, 'Dropout 只在训练时生效，model.eval() 时自动关闭。')
doc.add_paragraph()

h2(doc, '第 204-206 行：优化器与损失函数')
cd(doc, "    optimizer = torch.optim.Adam(model.parameters(), lr=0.001)")
cd(doc, "    loss_fn = nn.MSELoss()")
tbl(doc, ['组件', '选型', '原因'],
[['优化器', 'Adam(lr=0.001)', '自适应学习率，默认首选，不需要调 momentum'],
 ['损失函数', 'MSELoss', '回归任务，预测值和真实值的平方误差']])
doc.add_paragraph()

h2(doc, '第 209-223 行：训练循环')
cd(doc, "    for epoch in range(200):")
cd(doc, "        model.train()                  # 切换到训练模式")
cd(doc, "        optimizer.zero_grad()          # 梯度清零")
cd(doc, "        y_pred = model(X_train_t)      # 前向传播")
cd(doc, "        loss = loss_fn(y_pred, y_train_t)  # 算损失")
cd(doc, "        loss.backward()                # 反向传播")
cd(doc, "        optimizer.step()               # 更新参数")
tx(doc, '标准的 6 步训练循环：train → zero_grad → forward → loss → backward → step。每 50 epoch 打印一次训练/测试损失。')
note(doc, 'model.eval() + no_grad() 是验证时的标准组合，缺一不可。')
doc.add_paragraph()

h2(doc, '第 225-232 行：评估')
cd(doc, "    y_pred = model(X_test_t).numpy().flatten()")
cd(doc, "    mae = np.mean(np.abs(y_pred - y_test))")
cd(doc, "    rmse = np.sqrt(np.mean((y_pred - y_test) ** 2))")
tx(doc, 'MAE ≈ 1.8 kW：平均偏差不到 2kW。RMSE ≈ 2.3 kW：大误差惩罚更大。在功率范围 [10,80]kW 中，相对误差约 5%。')
doc.add_paragraph()

h2(doc, '第 234-251 行：可视化')
tx(doc, '画两个子图：上图为训练损失曲线（判断是否收敛），下图为预测值 vs 真实值对比（前 200 步）。保存到 results/ 目录。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Part 5
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Part 5: MDP 五元组 — GridWorld (第 256-324 行)')
tx(doc, '目标：用 4×4 网格世界理解马尔可夫决策过程的五个核心元素。')

h2(doc, '第 259-280 行：MDP 参数')
cd(doc, "    SIZE = 4")
cd(doc, "    n_states = SIZE * SIZE              # 16 个状态")
cd(doc, "    actions = ['↑', '↓', '←', '→']       # 4 个动作")
cd(doc, "    gamma = 0.9                         # 折扣因子")
cd(doc, "    GOAL = (3, 3); TRAP = (1, 1)        # 终点(+1) 和 陷阱(-1)")
tbl(doc, ['MDP 元素', 'GridWorld 对应', '值'],
[['S (状态)', '4×4 网格的每个格子', '16 个'],
 ['A (动作)', '上下左右移动', '4 个'],
 ["P (转移概率)", '80%目标方向+20%随机', '随机性'],
 ['R (奖励)', '终点+1, 陷阱-1, 其他0', '稀疏奖励'],
 ['γ (折扣)', '未来奖励的折扣', '0.9']])
doc.add_paragraph()

h2(doc, '第 282-286 行：辅助函数')
cd(doc, "    def pos_to_idx(r, c): return r * SIZE + c    # (行,列)→0~15的索引")
cd(doc, "    def is_valid(r, c): return 0 <= r < SIZE and 0 <= c < SIZE")
doc.add_paragraph()

h2(doc, '第 288-312 行：构建 R(s,a) 和 P(s\'|s,a)')
tx(doc, '用嵌套字典表示奖励和转移概率。核心循环遍历 4×4 网格的每个状态，对每个状态遍历 4 个动作：')
bl(doc, '终点和陷阱：吸收态，停在原地，奖励 +1/-1')
bl(doc, '普通格子：执行动作后有 80% 概率去目标方向，20% 概率均分到其他 3 个方向')
bl(doc, '如果撞墙（移出网格），停在原地')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Part 6
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Part 6: Bellman 方程 (第 327-404 行)')
tx(doc, '目标：用迭代法求解 Bellman 期望方程和最优方程，理解 V(s) 的物理意义。')

h2(doc, '第 338-361 行：策略评估 (Bellman 期望方程)')
cd(doc, "    policy = np.ones((n_states, n_actions)) / n_actions  # 均匀随机策略")
cd(doc, "    V = np.zeros(n_states)                               # 初始化 V=0")
cd(doc, "    for i in range(max_iter):                            # 迭代到收敛")
cd(doc, "        for s in range(n_states):                        # 对每个状态")
cd(doc, "            v_new = Σ_a π(a|s)[R(s,a) + γ Σ P(s'|s,a) V(s')]")
cd(doc, "            V[s] = v_new")
tx(doc, '每个状态的 V(s) = 所有动作的加权平均价值。迭代 133 轮收敛。')
tx(doc, '结果：起点 V(0) ≈ -3.93（随机策略下大概率掉陷阱）。')
doc.add_paragraph()

h2(doc, '第 363-378 行：最优值函数 V* (Bellman 最优方程)')
cd(doc, "    V_opt[s] = max_a [R(s,a) + γ Σ P(s'|s,a) V_opt(s')]")
tx(doc, '和策略评估的唯一区别：用 max 替代加权平均。每个状态选最好的动作。')
tx(doc, '结果：起点 V*(0) ≈ 3.34（最优策略下能避开陷阱走到终点）。')
doc.add_paragraph()

h2(doc, '第 385-402 行：可视化')
tx(doc, '画两个并排的 heatmap：左边是随机策略 V^π（值偏低，陷阱附近负值），右边是最优 V*（值普遍更高，终点附近最高）。数字标在每个格子中。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Part 7
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Part 7: 策略迭代 (第 407-496 行)')
tx(doc, '目标：实现策略迭代——交替执行策略评估和策略改进，直到策略收敛。')

h2(doc, '第 421-434 行：策略评估')
cd(doc, "    def policy_evaluation(policy, V):")
cd(doc, "        for s: V[s] = R[s][a] + γ Σ P(s'|s,a) V[s']")
tx(doc, '和 Part 6 一样解 Bellman 期望方程，区别是这里的策略是确定的（每个状态只有一个动作），不需要对所有动作加权平均。')
doc.add_paragraph()

h2(doc, '第 436-449 行：策略改进')
cd(doc, "    def policy_improvement(policy, V):")
cd(doc, "        for s: policy[s] = argmax_a [R + γ Σ P V]")
cd(doc, "        if 任何状态动作变了: policy_stable = False")
tx(doc, '贪心提升——对每个状态，计算所有 4 个动作的 Q 值，选最大的那个。如果没有任何状态的动作改变，说明策略已收敛。')
doc.add_paragraph()

h2(doc, '第 451-459 行：主循环')
cd(doc, "    for iteration in range(50):")
cd(doc, "        V = policy_evaluation(policy, V)")
cd(doc, "        policy, stable = policy_improvement(policy, V)")
cd(doc, "        if stable: break")
tx(doc, 'GridWorld 仅需 4 轮收敛到最优策略。每轮先完全评估当前策略，再改进。')
doc.add_paragraph()

h2(doc, '第 461-496 行：展示与可视化')
tx(doc, '打印最优策略图标（↑↓←→），画出最优策略图并用颜色表示各状态价值。终点标 G，陷阱标 X。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Part 8
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Part 8: 值迭代 (第 500-558 行)')
tx(doc, '目标：直接迭代 Bellman 最优方程，收敛后提取策略。')

h2(doc, '第 514-529 行：值迭代主循环')
cd(doc, "    for iteration in range(1000):")
cd(doc, "        V[s] = max_a [R[s][a] + γ Σ P(s'|s,a) V[s']]")
tx(doc, '和策略迭代的区别：不维护显式策略，直接迭代 V*。每轮对所有状态计算 max_a Q(s,a)。')
tx(doc, 'GridWorld 需 133 轮收敛（比策略迭代轮数多，但每轮计算更轻）。')
doc.add_paragraph()

h2(doc, '第 531-557 行：提取策略')
cd(doc, "    for s: policy[s] = argmax_a [R + γ Σ P V]")
tx(doc, 'V* 收敛后，从最优值函数中提取策略——和策略改进完全一样的 argmax 操作。最终策略应与策略迭代结果一致。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# Extra + 验证
# ═══════════════════════════════════════════════════════════════
h1(doc, 'Extra: 收敛过程可视化 (第 562-619 行)')
tx(doc, '记录值迭代每轮所有状态的 V(s)，画两条曲线：')
bl(doc, '上图：起点/终点/陷阱的 V(s) 随迭代次数的变化曲线')
bl(doc, '下图：对数坐标下的最大变化量 ΔV，观察收敛速度')
tx(doc, '可以看到起点 V(0) 从 0 逐渐上升到 ~3.34，陷阱 V(TRAP) 从 0 下降到 -1。')
doc.add_paragraph()

h1(doc, '验证: 策略迭代 vs 值迭代 一致性 (第 680-692 行)')
cd(doc, "    s0_pi = 3.3419  # 策略迭代 V(0)")
cd(doc, "    s0_vi = 3.3419  # 值迭代 V(0)")
cd(doc, "    ΔV = 0.000006   # 差值 < 1e-4")
tx(doc, '两种方法收敛到同一最优值函数，验证了实现正确性。这是面试常考点：策略迭代和值迭代最终结果相同，只是收敛路径不同。')
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# 主程序
# ═══════════════════════════════════════════════════════════════
h1(doc, '主程序入口 (第 625-697 行)')
cd(doc, "if __name__ == '__main__':")
tx(doc, '按顺序调用 Part 1→8，然后运行 Extra 收敛图，最后验证一致性。')
bl(doc, "if __name__ == '__main__'：Python 标准写法——直接运行此文件时执行，被 import 时不执行")
bl(doc, 'Part 5 的 mdp 字典对象被多个函数共享（Part 6-8 和 Extra 都用同一个 MDP）')
bl(doc, '验证逻辑：如果 ΔV < 1e-4 打印 [OK]，否则提示检查')
doc.add_paragraph()


# ── 结尾 ──
sep_p = doc.add_paragraph()
r = sep_p.add_run('─' * 50); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
r = doc.add_paragraph().add_run('Week 9 完整代码逐行分析 · 文件: scripts/week9_complete.py (698 行)\n生成日期：2026-07-10 · EMS 研究项目')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f'[OK] {OUT}')
